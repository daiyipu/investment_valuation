#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
定增项目企业基本面分析报告 - 主程序
从Tushare获取数据生成企业基本面分析报告

使用方法：
    python main.py --stock 002276.SZ --name 万马高分子
    python main.py --stock 603296.SH --name 华勤技术 --output 自定义报告.docx
"""

import os
import sys
import yaml
import argparse
import pandas as pd
from datetime import datetime
from typing import Dict, Any, Optional

# 添加当前目录到路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import tushare as ts
from utils.tushare_client import TushareClient
from utils.financial_scoring import FinancialScoringEngine
from utils.industry_analyzer import IndustryAnalyzer


class DingZengRiskReportGenerator:
    """定增项目企业基本面分析报告生成器"""

    def __init__(self, stock_code: str, stock_name: str = None, config_path: str = None):
        """初始化报告生成器

        Args:
            stock_code: 股票代码，如 002276.SZ
            stock_name: 企业名称（可选，不传则从Tushare自动获取）
            config_path: 配置文件路径
        """
        if config_path is None:
            config_path = os.path.join(os.path.dirname(__file__), 'config.yaml')

        with open(config_path, 'r', encoding='utf-8') as f:
            self.config = yaml.safe_load(f)

        # 获取Tushare Token (优先使用环境变量)
        tushare_token = self.config['tushare'].get('token', '') or os.environ.get('TUSHARE_TOKEN', '')
        if not tushare_token:
            raise ValueError("请设置TUSHARE_TOKEN环境变量或在config.yaml中配置token")

        # 初始化Tushare客户端
        self.ts_client = TushareClient(tushare_token)
        self.pro = self.ts_client.get_pro()

        # 用传入参数覆盖config中的项目信息
        self.stock_code = stock_code
        self.stock_name = stock_name
        self.project_info = {
            'stock_code': stock_code,
            'name': stock_name or '',
            'report_date': datetime.now().strftime('%Y-%m-%d'),
            'version': 'v1.0',
        }

        # 将项目信息写入config，以便各章节通过config.get('project')访问
        self.config['project'] = self.project_info

        # 初始化评分引擎（使用config中的评分权重配置）
        self.scoring_engine = FinancialScoringEngine(self.config['scoring'])

        # 行业分析器延迟初始化（需要先获取公司信息来确定行业）
        self.industry_analyzer = None

    def _init_industry_analyzer(self) -> None:
        """根据申万三级行业分类查找同行并初始化行业分析器"""
        try:
            company_info = self.pro.stock_basic(
                ts_code=self.stock_code,
                fields='ts_code,name,industry,market,list_date'
            )
            industry_name = ''
            if company_info is not None and len(company_info) > 0:
                industry_name = company_info.iloc[0].get('industry', '')

            peer_codes = [self.stock_code]
            sw_industry_name = ''

            # 优先使用申万三级行业分类查找同行
            peer_codes, sw_industry_name = self._find_sw_l3_peers(industry_name)

            if len(peer_codes) <= 1 and industry_name:
                # SW三级未找到，回退到一级行业分类
                try:
                    peers = self.pro.stock_basic(
                        industry=industry_name,
                        fields='ts_code,name,industry',
                        list_status='L'
                    )
                    if peers is not None and len(peers) > 0:
                        other_peers = peers[peers['ts_code'] != self.stock_code]['ts_code'].tolist()
                        peer_codes.extend(other_peers[:10])
                except Exception as e:
                    print(f"获取一级行业公司列表时出错: {e}")

            industry_config = {
                'industry': {
                    'name': sw_industry_name or industry_name,
                    'codes': peer_codes,
                }
            }
            self.industry_analyzer = IndustryAnalyzer(self.pro, industry_config)
            print(f"行业分析器已初始化: {sw_industry_name or industry_name}，同行公司 {len(peer_codes)} 家")

        except Exception as e:
            print(f"初始化行业分析器时出错: {e}，使用config中的默认配置")
            self.industry_analyzer = IndustryAnalyzer(self.pro, self.config['industry'])

    def _find_sw_l3_peers(self, industry_name: str) -> tuple:
        """通过申万三级行业分类查找同行公司（使用index_member_all一次查询）

        Returns:
            (peer_codes, sw_l3_name)
        """
        peer_codes = [self.stock_code]
        sw_l3_name = ''
        self._sw_industry_info = {}

        try:
            # 一次查询获取所有SW2021行业成分（is_new='Y'仅当前成分）
            all_members = self.pro.index_member_all(is_new='Y')
            if all_members is None or all_members.empty:
                return peer_codes, ''

            # 找到目标公司所属的SW三级行业
            target = all_members[all_members['ts_code'] == self.stock_code]
            if target.empty:
                return peer_codes, ''

            row = target.iloc[0]
            l3_code = row['l3_code']
            sw_l3_name = row['l3_name']

            # 保存完整行业信息供后续使用
            self._sw_industry_info = {
                'l1_name': row['l1_name'],
                'l2_name': row['l2_name'],
                'l3_name': sw_l3_name,
                'l3_code': l3_code,
            }

            # 获取同三级行业的所有公司
            peers = all_members[all_members['l3_code'] == l3_code]
            peer_codes = peers['ts_code'].tolist()

            print(f"申万三级行业匹配: {row['l1_name']} > {row['l2_name']} > {sw_l3_name} ({l3_code})，同行 {len(peer_codes)} 家")

        except Exception as e:
            print(f"申万三级行业分类查找出错: {e}")

        return peer_codes, sw_l3_name

    def _fetch_research_reports(self) -> dict:
        """从东方财富获取行业研报数据"""
        try:
            fetcher_config = self.config.get('report_fetcher', {})
            if not fetcher_config:
                print("  未配置report_fetcher，跳过研报获取")
                return {}

            from utils.report_fetcher import ReportFetcher
            fetcher = ReportFetcher(fetcher_config)

            # 优先使用index_member_all查到的SW行业名
            sw_info = getattr(self, '_sw_industry_info', {})
            sw_l1_name = sw_info.get('l1_name', '')
            sw_l2_name = sw_info.get('l2_name', '')
            sw_l3_name = sw_info.get('l3_name', '')

            if not sw_l1_name:
                try:
                    company_info = self.pro.stock_basic(
                        ts_code=self.stock_code, fields='ts_code,industry')
                    if company_info is not None and len(company_info) > 0:
                        sw_l1_name = company_info.iloc[0].get('industry', '')
                except Exception:
                    pass

            # 用L1行业名搜行业研报，L2/L3名作为补充关键词
            search_keywords = [kw for kw in [sw_l1_name, sw_l2_name, sw_l3_name] if kw]
            industry_name = sw_l1_name

            result = fetcher.fetch_reports_for_chapter(
                stock_code=self.stock_code,
                industry_name=industry_name,
                sw_l1_name=', '.join(search_keywords),
            )

            n_reports = len(result.get('reports', []))
            report_text_len = len(result.get('report_text', ''))
            n_errors = len(result.get('download_errors', []))
            print(f"  获取研报 {n_reports} 篇，LLM参考文本 {report_text_len} 字"
                  + (f"，正文提取失败 {n_errors} 篇" if n_errors else ""))
            return result

        except Exception as e:
            print(f"  获取行业研报数据时出错: {e}")
            return {}

    def load_all_data(self) -> Dict[str, Any]:
        """加载所有需要的数据

        Returns:
            包含所有分析数据的字典
        """
        print("=" * 60)
        print("开始加载数据...")
        print("=" * 60)

        data = {}

        # 1. 获取公司基本信息（同时获取公司名称）
        print("\n[1/15] 获取公司基本信息...")
        data['company_info'] = self._get_company_info()

        # 用Tushare返回的真实名称更新project_info
        if not self.stock_name and data['company_info'].get('name'):
            self.stock_name = data['company_info']['name']
            self.project_info['name'] = self.stock_name

        # 2. 初始化行业分析器（依赖公司行业信息）
        print("[2/15] 初始化行业分析器...")
        self._init_industry_analyzer()

        # 2.1 保存SW行业层级信息供各章节使用
        data['sw_industry'] = getattr(self, '_sw_industry_info', {})

        # 2.5 获取行业研报数据（东方财富）
        print("[2.5/15] 获取行业研报数据...")
        data['research_reports'] = self._fetch_research_reports()

        # 3. 获取财务指标数据
        print("[3/15] 获取财务指标数据...")
        data['financial_indicators'] = self._get_financial_indicators()

        # 4. 获取财务报表数据
        print("[4/15] 获取财务报表数据...")
        data['financial_statements'] = self._get_financial_statements()

        # 5. 获取行业对比数据
        print("[5/15] 获取行业对比数据...")
        data['industry_data'] = self._get_industry_comparison()

        # 5.5 获取同行财务报表数据（用于第五章占比分析和行业百分位）
        print("[5.5/15] 获取同行财务报表数据...")
        data['peer_statements'] = self._get_peer_financial_statements()

        # 将行业数据和财务报表传递给评分引擎
        self.scoring_engine.set_industry_data(data['industry_data'])
        self.scoring_engine.set_financial_statements(data['financial_statements'])

        # 6. 计算财务评分
        print("[6/15] 计算财务评分...")
        data['scoring_result'] = self._calculate_financial_score(data['financial_indicators'])

        # 7. 获取估值数据
        print("[7/15] 获取估值数据...")
        data['valuation_data'] = self._get_valuation_data()

        # 8. 获取主营业务和历史估值
        print("[8/15] 获取主营业务和历史估值...")
        data['main_business'] = self._get_main_business()
        data['historical_valuation'] = self._get_historical_valuation()

        # 9. 获取公司详细信息（stock_company）
        print("[9/15] 获取公司详细信息...")
        data['stock_company'] = self._get_stock_company()

        # 10. 获取股东数据
        print("[10/15] 获取股东数据...")
        data['top10_holders'] = self._get_top10_holders()
        data['top10_floatholders'] = self._get_top10_floatholders()
        data['stk_holdernumber'] = self._get_stk_holdernumber()

        # 11. 获取管理层信息
        print("[11/15] 获取管理层信息...")
        data['stk_managers'] = self._get_stk_managers()
        data['stk_rewards'] = self._get_stk_rewards()

        # 12. 获取区域业务数据
        print("[12/15] 获取区域业务数据...")
        data['region_business'] = self._get_region_business()

        # 13. 获取分红数据
        print("[13/15] 获取分红数据...")
        data['dividend'] = self._get_dividend()

        # 14. 获取限售股解禁数据
        print("[14/15] 获取限售股解禁数据...")
        data['share_float'] = self._get_share_float()

        # 15. 获取大股东交易数据
        print("[15/15] 获取大股东交易数据...")
        data['stk_holdertrade'] = self._get_stk_holdertrade()

        print("\n数据加载完成!")
        return data

    def _get_company_info(self) -> Dict[str, Any]:
        """获取公司基本信息"""
        company_info = self.pro.stock_basic(
            ts_code=self.stock_code,
            fields='ts_code,name,industry,market,list_date'
        )

        if company_info is None or len(company_info) == 0:
            print(f"警告: 无法获取股票 {self.stock_code} 的基本信息")
            return {}

        info = company_info.iloc[0].to_dict()

        try:
            name_en = self.pro.namechange(ts_code=self.stock_code, fields='name_en,short_name')
            if name_en is not None and len(name_en) > 0:
                info['name_en'] = name_en.iloc[0].get('name_en', '')
                info['short_name'] = name_en.iloc[0].get('short_name', '')
        except Exception as e:
            print(f"获取公司英文名称时出错: {e}")

        return info

    def _get_stock_company(self) -> Dict[str, Any]:
        """获取公司详细信息（stock_company接口）"""
        try:
            df = self.pro.stock_company(ts_code=self.stock_code)
            if df is not None and len(df) > 0:
                return df.iloc[0].to_dict()
        except Exception as e:
            print(f"获取公司详细信息时出错: {e}")
        return {}

    def _get_top10_holders(self) -> pd.DataFrame:
        """获取前十大股东数据"""
        try:
            df = self.pro.top10_holders(ts_code=self.stock_code)
            if df is not None and len(df) > 0:
                return df
        except Exception as e:
            print(f"获取前十大股东数据时出错: {e}")
        return pd.DataFrame()

    def _get_top10_floatholders(self) -> pd.DataFrame:
        """获取前十大流通股东数据"""
        try:
            df = self.pro.top10_floatholders(ts_code=self.stock_code)
            if df is not None and len(df) > 0:
                return df
        except Exception as e:
            print(f"获取前十大流通股东数据时出错: {e}")
        return pd.DataFrame()

    def _get_stk_holdernumber(self) -> pd.DataFrame:
        """获取股东户数变化数据"""
        try:
            df = self.pro.stk_holdernumber(ts_code=self.stock_code)
            if df is not None and len(df) > 0:
                return df
        except Exception as e:
            print(f"获取股东户数数据时出错: {e}")
        return pd.DataFrame()

    def _get_stk_managers(self) -> pd.DataFrame:
        """获取管理层信息"""
        try:
            df = self.pro.stk_managers(ts_code=self.stock_code)
            if df is not None and len(df) > 0:
                return df
        except Exception as e:
            print(f"获取管理层信息时出错: {e}")
        return pd.DataFrame()

    def _get_stk_rewards(self) -> pd.DataFrame:
        """获取管理层薪酬数据"""
        try:
            df = self.pro.stk_rewards(ts_code=self.stock_code)
            if df is not None and len(df) > 0:
                return df
        except Exception as e:
            print(f"获取管理层薪酬数据时出错: {e}")
        return pd.DataFrame()

    def _get_region_business(self) -> pd.DataFrame:
        """获取主营业务区域分布数据"""
        try:
            df = self.pro.fina_mainbz(
                ts_code=self.stock_code,
                type='R',
                start_date=f"{int(datetime.now().strftime('%Y')) - 2}0101"
            )
            if df is not None and len(df) > 0:
                return df
        except Exception as e:
            print(f"获取区域业务数据时出错: {e}")
        return pd.DataFrame()

    def _get_dividend(self) -> pd.DataFrame:
        """获取分红数据"""
        try:
            df = self.pro.dividend(ts_code=self.stock_code)
            if df is not None and len(df) > 0:
                return df
        except Exception as e:
            print(f"获取分红数据时出错: {e}")
        return pd.DataFrame()

    def _get_share_float(self) -> pd.DataFrame:
        """获取限售股解禁数据"""
        try:
            df = self.pro.share_float(ts_code=self.stock_code)
            if df is not None and len(df) > 0:
                return df
        except Exception as e:
            print(f"获取限售股解禁数据时出错: {e}")
        return pd.DataFrame()

    def _get_stk_holdertrade(self) -> pd.DataFrame:
        """获取大股东交易数据"""
        try:
            df = self.pro.stk_holdertrade(ts_code=self.stock_code)
            if df is not None and len(df) > 0:
                return df
        except Exception as e:
            print(f"获取大股东交易数据时出错: {e}")
        return pd.DataFrame()

    def _get_financial_indicators(self) -> pd.DataFrame:
        """获取财务指标数据"""
        end_date = datetime.now().strftime('%Y%m%d')
        start_year = int(datetime.now().strftime('%Y')) - 5
        start_date = f"{start_year}0101"

        df = self.pro.fina_indicator(
            ts_code=self.stock_code,
            start_date=start_date,
            end_date=end_date
        )

        if df is not None and len(df) > 0:
            df = df[df['end_date'].str.contains('1231', na=False)]

        return df if df is not None else pd.DataFrame()

    def _get_financial_statements(self) -> Dict[str, pd.DataFrame]:
        """获取财务报表数据"""
        end_date = datetime.now().strftime('%Y%m%d')
        start_year = int(datetime.now().strftime('%Y')) - 5
        start_date = f"{start_year}0101"

        statements = {}

        statements['balance_sheet'] = self.pro.balancesheet(
            ts_code=self.stock_code,
            start_date=start_date,
            end_date=end_date
        )

        statements['income_statement'] = self.pro.income(
            ts_code=self.stock_code,
            start_date=start_date,
            end_date=end_date
        )

        statements['cashflow'] = self.pro.cashflow(
            ts_code=self.stock_code,
            start_date=start_date,
            end_date=end_date
        )

        return statements

    def _get_peer_financial_statements(self) -> Dict[str, pd.DataFrame]:
        """获取同行公司的三大财务报表原始数据（用于第五章占比分析的行业百分位计算）"""
        peer_statements = {
            'peer_balance_sheet': pd.DataFrame(),
            'peer_income_statement': pd.DataFrame(),
            'peer_cashflow': pd.DataFrame(),
        }

        if self.industry_analyzer is None:
            return peer_statements

        industry_codes = self.industry_analyzer.industry_codes
        peer_codes = [c for c in industry_codes if c != self.stock_code]
        if not peer_codes:
            return peer_statements

        end_date = datetime.now().strftime('%Y%m%d')
        start_year = int(datetime.now().strftime('%Y')) - 5
        start_date = f"{start_year}0101"

        import time

        for api_name, result_key in [
            ('balancesheet', 'peer_balance_sheet'),
            ('income', 'peer_income_statement'),
            ('cashflow', 'peer_cashflow'),
        ]:
            df_list = []
            api_func = getattr(self.pro, api_name)
            fields_str = 'ts_code,end_date,' + self._get_statement_fields(api_name)

            for code in peer_codes:
                try:
                    df = api_func(ts_code=code, start_date=start_date, end_date=end_date, fields=fields_str)
                    if df is not None and not df.empty:
                        # 只保留年报
                        df = df[df['end_date'].astype(str).str.contains('1231', na=False)]
                        if not df.empty:
                            df_list.append(df)
                except Exception:
                    pass
                time.sleep(0.3)

            if df_list:
                cleaned = [d.dropna(axis=1, how='all') for d in df_list if not d.empty]
                if cleaned:
                    peer_statements[result_key] = pd.concat(cleaned, ignore_index=True)

        print(f"  同行报表数据获取完成: BS {len(peer_statements['peer_balance_sheet'])}条, "
              f"IS {len(peer_statements['peer_income_statement'])}条, "
              f"CF {len(peer_statements['peer_cashflow'])}条")
        return peer_statements

    @staticmethod
    def _get_statement_fields(api_name: str) -> str:
        """获取报表API需要拉取的字段列表（Tushare实际字段名）"""
        fields_map = {
            'balancesheet': 'money_cap,notes_receiv,accounts_receiv,prepayment,oth_receiv,inventories,oth_cur_assets,'
                            'total_cur_assets,lt_eqt_invest,invest_real_estate,fix_assets,cip,intan_assets,goodwill,'
                            'lt_amor_exp,defer_tax_assets,oth_nca,total_nca,total_assets,'
                            'st_borr,notes_payable,acct_payable,adv_receipts,payroll_payable,taxes_payable,'
                            'int_payable,oth_payable,non_cur_liab_due_1y,oth_cur_liab,total_cur_liab,'
                            'lt_borr,bond_payable,lease_liab,total_ncl,total_liab,'
                            'cap_rese,surplus_rese,undistr_porfit,total_hldr_eqy_exc_min_int,minority_int',
            'income': 'total_revenue,total_cogs,revenue,biz_tax_surchg,sell_exp,admin_exp,fin_exp,'
                      'assets_impair_loss,rd_exp,operate_profit,total_profit,income_tax,n_income,'
                      'n_income_attr_p,basic_eps',
            'cashflow': 'c_fr_sale_sg,recp_tax_rends,n_cashflow_act,'
                        'c_pay_acq_const_fiolta,n_cashflow_inv_act,'
                        'c_recp_borrow,c_prepay_amt_borr,c_pay_dist_dpcp_int_exp,n_cash_flows_fnc_act,'
                        'c_cash_equ_end_period',
        }
        return fields_map.get(api_name, '')

    def _get_industry_comparison(self) -> pd.DataFrame:
        """获取行业对比数据"""
        if self.industry_analyzer is None:
            return pd.DataFrame()

        industry_codes = self.industry_analyzer.industry_codes
        end_date = datetime.now().strftime('%Y%m%d')
        start_year = int(datetime.now().strftime('%Y')) - 5
        start_date = f"{start_year}0101"

        df_list = []
        for code in industry_codes:
            try:
                df = self.pro.fina_indicator(
                    ts_code=code,
                    start_date=start_date,
                    end_date=end_date
                )
                if df is not None and len(df) > 0:
                    df = df[df['end_date'].str.contains('1231', na=False)]
                    if not df.empty:
                        df_list.append(df)
            except Exception as e:
                print(f"获取 {code} 财务指标时出错: {e}")

        if df_list:
            cleaned = [d.dropna(axis=1, how='all') for d in df_list if not d.empty]
            if cleaned:
                return pd.concat(cleaned, ignore_index=True)
        return pd.DataFrame()

    def _calculate_financial_score(self, financial_data: pd.DataFrame = None) -> Dict[str, Any]:
        """计算财务评分"""
        if financial_data is None:
            financial_data = self._get_financial_indicators()

        if financial_data.empty:
            return {
                'total_score': 0,
                'grade': '暂无数据',
                'dimensions': {}
            }

        return self.scoring_engine.calculate_score(financial_data)

    def _get_valuation_data(self) -> Dict[str, Any]:
        """获取估值数据"""
        data = {}

        end_date = datetime.now().strftime('%Y%m%d')
        start_date = (datetime.now() - pd.Timedelta(days=365)).strftime('%Y%m%d')

        df_daily = self.pro.daily(ts_code=self.stock_code, start_date=start_date, end_date=end_date)

        if df_daily is not None and len(df_daily) > 0:
            data['current_price'] = df_daily.iloc[-1]['close']

        try:
            df_val = self.pro.daily_basic(ts_code=self.stock_code, start_date=start_date, end_date=end_date)
            if df_val is not None and len(df_val) > 0:
                latest = df_val.iloc[-1]
                data['pe_ratio'] = latest.get('pe_ttm', latest.get('pe', None))
                data['pb_ratio'] = latest.get('pb', None)
                data['market_cap'] = latest.get('total_mv', None)
                data['float_market_cap'] = latest.get('float_mv', None)
                data['turnover_rate'] = latest.get('turnover_rate', None)
        except Exception as e:
            print(f"获取估值数据时出错: {e}")

        return data

    def _get_main_business(self) -> pd.DataFrame:
        """获取主营业务构成数据"""
        try:
            df_product = self.pro.fina_mainbz(
                ts_code=self.stock_code,
                type='P',
                start_date=f"{int(datetime.now().strftime('%Y')) - 2}0101"
            )
            if df_product is not None and len(df_product) > 0:
                return df_product
        except Exception as e:
            print(f"获取主营业务数据时出错: {e}")

        return pd.DataFrame()

    def _get_historical_valuation(self) -> pd.DataFrame:
        """获取历史PE/PB数据（近3年）"""
        try:
            end_date = datetime.now().strftime('%Y%m%d')
            start_date = (datetime.now() - pd.Timedelta(days=365 * 3)).strftime('%Y%m%d')

            df = self.pro.daily_basic(
                ts_code=self.stock_code,
                start_date=start_date,
                end_date=end_date,
                fields='ts_code,trade_date,pe,pe_ttm,pb,ps,total_mv,float_mv,turnover_rate'
            )
            return df if df is not None else pd.DataFrame()
        except Exception as e:
            print(f"获取历史估值数据时出错: {e}")

        return pd.DataFrame()

    def generate_report(self, output_path: str = None) -> str:
        """生成完整的风险评估报告

        Args:
            output_path: 输出文件路径

        Returns:
            生成的报告文件路径
        """
        print("\n" + "=" * 60)
        print("开始生成报告...")
        print("=" * 60)

        # 加载所有数据
        all_data = self.load_all_data()

        # 生成各章节内容
        chapters = []

        print("\n生成第一章：项目概述...")
        from chapters.chapter01_overview import Chapter01Overview
        ch01 = Chapter01Overview(self.config, all_data)
        chapters.append(ch01.generate())

        print("生成第二章：公司基本情况...")
        from chapters.chapter02_company import Chapter02Company
        ch02 = Chapter02Company(self.config, all_data)
        chapters.append(ch02.generate())

        print("生成第三章：公司业务模式...")
        from chapters.chapter03_business import Chapter03Business
        ch03 = Chapter03Business(self.config, all_data)
        chapters.append(ch03.generate())

        print("生成第四章：行业研究...")
        from chapters.chapter04_industry import Chapter04Industry
        ch04 = Chapter04Industry(self.config, all_data)
        chapters.append(ch04.generate())

        print("生成第五章：公司财务情况...")
        from chapters.chapter05_financial import Chapter05Financial
        ch05 = Chapter05Financial(self.config, all_data, self.scoring_engine)
        chapters.append(ch05.generate())

        print("生成第六章：公司估值...")
        from chapters.chapter06_valuation import Chapter06Valuation
        ch06 = Chapter06Valuation(self.config, all_data)
        chapters.append(ch06.generate())

        print("生成第七章：风险概述...")
        from chapters.chapter07_risk_summary import Chapter07RiskSummary
        ch07 = Chapter07RiskSummary(self.config, all_data)
        chapters.append(ch07.generate())

        # 生成Word文档
        print("\n生成Word文档...")
        company_name = self.stock_name or all_data.get('company_info', {}).get('name', '未知')
        if output_path is None:
            output_dir = os.path.join(os.path.dirname(__file__), self.config['report']['output_dir'])
            os.makedirs(output_dir, exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = os.path.join(output_dir, f"{company_name}定增项目企业基本面分析报告_{timestamp}.docx")

        self._generate_word_document(chapters, output_path, company_name)

        print(f"\n报告生成完成: {output_path}")
        return output_path

    def _generate_word_document(self, chapters: list, output_path: str, company_name: str = '') -> None:
        """生成Word文档"""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 设置文档标题
        title_text = f"{company_name}定增项目企业基本面分析报告"
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加报告信息
        info_para = doc.add_paragraph()
        info_para.add_run(f"报告日期: {self.project_info['report_date']}").bold = True
        info_para.add_run(f"\n股票代码: {self.stock_code}")
        info_para.add_run(f"\n版本: {self.project_info['version']}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # 添加各章节内容
        for chapter in chapters:
            for element in chapter:
                if element['type'] == 'heading':
                    doc.add_heading(element['content'], level=element.get('level', 1))
                elif element['type'] == 'paragraph':
                    doc.add_paragraph(element['content'])
                elif element['type'] == 'table':
                    self._add_table_to_doc(doc, element['data'], element.get('headers', []))
                elif element['type'] == 'image':
                    if element.get('path'):
                        try:
                            doc.add_picture(element['path'], width=Inches(6))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            print(f"  插入图片失败: {e}")

        doc.save(output_path)

    def _add_table_to_doc(self, doc, data: list, headers: list) -> None:
        """添加表格到Word文档"""
        if not data:
            return

        table = doc.add_table(rows=1 + len(data), cols=len(headers) if headers else len(data[0]))
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        for i, header in enumerate(headers if headers else data[0].keys()):
            header_cells[i].text = str(header)

        num_cols = len(headers) if headers else len(data[0])
        for row_idx, row_data in enumerate(data):
            row_cells = table.rows[row_idx + 1].cells
            if isinstance(row_data, dict):
                for col_idx, value in enumerate(row_data.values()):
                    if col_idx >= num_cols:
                        break
                    row_cells[col_idx].text = str(value) if value is not None else ''
            else:
                for col_idx, value in enumerate(row_data):
                    if col_idx >= num_cols:
                        break
                    row_cells[col_idx].text = str(value) if value is not None else ''


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='定增项目企业基本面分析报告生成器',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python main.py --stock 002276.SZ --name 万马高分子
  python main.py --stock 603296.SH --name 华勤技术 --output 华勤技术报告.docx
  python main.py --stock 300735.SZ  # 企业名称自动从Tushare获取

股票代码格式:
  上交所: 600xxx.SH, 601xxx.SH, 603xxx.SH
  深交所: 000xxx.SZ, 002xxx.SZ, 300xxx.SZ
  北交所: 8xxxxx.BJ
        """
    )
    parser.add_argument('--stock', type=str, required=True,
                        help='股票代码（必填），如 002276.SZ')
    parser.add_argument('--name', type=str, default=None,
                        help='企业名称（可选，不填则自动从Tushare获取）')
    parser.add_argument('--output', type=str, default=None,
                        help='输出文件名（可选，默认自动生成）')

    args = parser.parse_args()

    print("=" * 60)
    print("定增项目企业基本面分析报告生成程序")
    print("=" * 60)
    print(f"股票代码: {args.stock}")
    print(f"企业名称: {args.name or '(自动获取)'}")
    print()

    # 创建报告生成器
    generator = DingZengRiskReportGenerator(
        stock_code=args.stock,
        stock_name=args.name
    )

    # 生成报告
    output_path = generator.generate_report(args.output)

    print(f"\n报告已生成: {output_path}")


if __name__ == "__main__":
    main()
