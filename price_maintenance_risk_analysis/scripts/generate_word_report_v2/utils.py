# -*- coding: utf-8 -*-
"""
通用工具函数模块

包含所有Word文档格式化和图表生成的工具函数
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os


# ==================== Word文档格式化函数 ====================

def setup_chinese_font(document):
    """设置中文字体 - V2版本"""
    # 正文：仿宋-GB2312，四号（14pt）
    document.styles['Normal'].font.name = '仿宋_GB2312'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    document.styles['Normal'].font.size = Pt(14)


def add_title(document, text, level=1):
    """添加标题 - V2版本"""
    if level == 0:  # 报告标题，使用段落而不是heading
        para = document.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = '方正公文小标宋_GBK'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
            run.font.size = Pt(22)  # 二号（22pt）
            run.font.bold = True
        return para
    else:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            # 根据级别设置字体和大小
            if level == 1:  # 一级标题
                run.font.name = '方正公文小标宋_GBK'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
                run.font.size = Pt(16)  # 三号（16pt）
            elif level == 2:  # 二级标题
                run.font.name = '方正公文小标宋_GBK'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
                run.font.size = Pt(15)  # 小三（15pt）
            else:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(14)
        return heading


def add_paragraph(document, text, bold=False, font_size=14):
    """添加段落 - V2版本（正文：仿宋-GB2312，四号）"""
    para = document.add_paragraph(text)
    for run in para.runs:
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(font_size)
        if bold:
            run.font.bold = True
    return para


def add_image(document, image_path, width=Inches(5)):
    """添加图片到文档"""
    if os.path.exists(image_path):
        document.add_picture(image_path, width=width)
        # 设置图片居中
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    else:
        print(f"⚠️ 图片不存在: {image_path}")
        return False


def add_table_data(document, headers, data, font_size=12):
    """添加表格 - V2版本（表格：宋体，默认小四号12pt）

    Args:
        document: Word文档对象
        headers: 表头列表
        data: 数据列表
        font_size: 字体大小（单位：磅），默认12pt（小四号），10.5pt为五号
    """
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 设置表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        para = header_cells[i].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.size = Pt(font_size)
        para.runs[0].font.name = '宋体'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            para = row_cells[i].paragraphs[0]
            para.runs[0].font.size = Pt(font_size)
            para.runs[0].font.name = '宋体'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if isinstance(cell_data, (int, float)):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def add_section_break(document):
    """添加分页符"""
    document.add_page_break()


# ==================== 图表生成函数 ====================
# 注意：以下函数需要从 generate_word_report_v2.py 中提取具体实现
# 这里只列出函数签名，实际实现需要从原文件复制

def generate_break_even_chart(issue_price, current_price, lockup_period, save_path):
    """生成盈亏平衡价格敏感性曲线"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


def generate_lockup_sensitivity_charts_split(issue_price, current_price, save_dir):
    """生成锁定期敏感性分析图表"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


def generate_time_window_analysis_chart(price_series, save_dir):
    """生成不同时间窗口的风险指标对比图表"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


def generate_tornado_chart_split(issue_price, current_price, lockup_period, volatility, drift, save_dir):
    """生成龙卷风图 - 参数敏感性排序（拆分版）"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


def generate_discount_scenario_charts_split(ma120, current_price, volatility, drift, lockup_period, save_dir):
    """生成发行价折扣情景图表（统一版）"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


def generate_sensitivity_3d_charts_split(issue_price, current_price, volatility, drift, lockup_period, save_dir):
    """生成3D曲面敏感性分析图表"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


def generate_monte_carlo_charts_split(sim_prices, issue_price, current_price, save_dir):
    """生成蒙特卡洛模拟图表（拆分版）"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


def generate_stress_test_charts_split(stress_test_results, save_dir):
    """生成压力测试图表"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


def generate_var_charts_split(var_results, save_dir):
    """生成VaR分析图表"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


# ==================== 其他辅助函数 ====================

def generate_relative_valuation_charts_split(current_metrics, industry_avg, peer_companies_data, save_dir):
    """生成相对估值分析图表（拆分版本）"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


def generate_dcf_charts_split(dcf_results, save_dir):
    """生成DCF估值图表"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


def generate_market_data_charts_split(market_data, save_dir):
    """生成市场数据图表"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass


def generate_industry_index_charts_split(industry_data, save_dir):
    """生成行业指数图表"""
    # TODO: 从 generate_word_report_v2.py 提取实现
    pass
