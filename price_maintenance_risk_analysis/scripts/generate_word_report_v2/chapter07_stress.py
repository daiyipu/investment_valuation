# -*- coding: utf-8 -*-
"""
第七章 - 压力测试

功能：生成压力测试章节内容，包括PE回归压力测试、经济面极端情况压力测试、多重敏感性指标极端情况压力测试
"""

import sys
import os
import numpy as np
import matplotlib.pyplot as plt
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 添加路径
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, PROJECT_DIR)

from utils.font_manager import get_font_prop

# 获取中文字体
font_prop = get_font_prop()


def add_title(document, text, level=1):
    """添加标题"""
    if level == 1:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = '方正公文小标宋_GBK'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
            run.font.size = Pt(16)
    elif level == 2:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = '方正公文小标宋_GBK'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
            run.font.size = Pt(15)
    else:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(14)
    return heading


def add_paragraph(document, text, bold=False, font_size=14):
    """添加段落"""
    para = document.add_paragraph(text)
    for run in para.runs:
        run.font.name = '仿宋-GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋-GB2312')
        run.font.size = Pt(font_size)
        if bold:
            run.bold = True
    return para


def add_image(document, image_path, width=Inches(5)):
    """添加图片到文档"""
    if os.path.exists(image_path):
        last_paragraph = document.add_paragraph()
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = last_paragraph.add_run()
        run.add_picture(image_path, width=width)


def add_table_data(document, headers, data, font_size=12):
    """添加表格"""
    table = document.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 设置表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(font_size)
                run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置数据行
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(font_size)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section_break(document):
    """添加分页符"""
    document.add_page_break()


def generate_stress_test_chart(scenarios, returns, save_path):
    """生成压力测试图表"""
    plt.figure(figsize=(10, 6))
    colors = ['#e74c3c' if r < 0 else '#2ecc71' for r in returns]
    bars = plt.barh(range(len(scenarios)), returns, color=colors, alpha=0.7)
    plt.xlabel('收益率 (%)', fontproperties=font_prop, fontsize=12)
    plt.ylabel('情景', fontproperties=font_prop, fontsize=12)
    plt.title('压力测试情景分析', fontproperties=font_prop, fontsize=14, fontweight='bold')
    plt.yticks(range(len(scenarios)), scenarios, fontproperties=font_prop)
    plt.xticks(fontproperties=font_prop)
    plt.axvline(x=0, color='black', linestyle='-', linewidth=1)
    plt.grid(True, alpha=0.3, axis='x')

    for bar, ret in zip(bars, returns):
        plt.text(bar.get_width() + (1 if ret > 0 else -1),
                bar.get_y() + bar.get_height()/2,
                f'{ret:.1f}%', ha='left' if ret > 0 else 'right',
                va='center', fontsize=10, fontproperties=font_prop)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_chapter(context):
    """
    生成第七章：压力测试

    参数:
        context: 包含所有必要数据的字典
            - document: Word文档对象
            - project_params: 项目参数
            - market_data: 市场数据
            - current_metrics_val: 当前指标（PE等）
            - industry_stats_val: 行业统计数据
            - peer_companies_val: 同行业公司数据
            - IMAGES_DIR: 图片保存目录
    """
    document = context['document']
    project_params = context['project_params']
    market_data = context['market_data']
    current_metrics_val = context['current_metrics_val']
    industry_stats_val = context['industry_stats_val']
    peer_companies_val = context['peer_companies_val']
    IMAGES_DIR = context['IMAGES_DIR']

    # ==================== 七、压力测试 ====================
    add_title(document, '七、压力测试', level=1)

    add_paragraph(document, '本章节从多个维度模拟极端市场情况下的定增项目表现，包括估值回归风险、经济面极端情况、以及多重风险因素叠加的最差情景。')
    add_paragraph(document, '通过全面压力测试，评估项目在极端环境下的风险承受能力。')

    # ==================== 7.1 PE回归压力测试 ====================
    add_title(document, '7.1 PE回归压力测试', level=2)

    add_paragraph(document, '本节分析PE估值回归到行业极端情况时的情景，评估估值回归风险。')
    add_paragraph(document, '通过模拟PE回归到行业Q1（25分位，即下四分位数），评估最坏情况下估值回调对投资收益的影响。')

    # 计算回归情景
    current_price_rel = project_params['current_price']
    pe_q1 = None
    return_pe_stress = 0

    if current_metrics_val['pe'] is None:
        add_paragraph(document, '当前PE数据缺失（可能为非交易日或数据未更新），无法进行PE回归压力测试。')
        add_paragraph(document, '将以PB倍数回归作为替代压力测试指标。')
        # PB回归替代
        if current_metrics_val['pb'] and current_metrics_val['pb'] > 0 and industry_stats_val['pb']['q1'] > 0:
            pb_q1 = industry_stats_val['pb']['q1']
            bps = current_price_rel / current_metrics_val['pb']
            target_price_q1 = bps * pb_q1
            return_q1 = (target_price_q1 - current_price_rel) / current_price_rel * 100
            scenario_headers_stress = ['情景', '当前PB', '回归后PB', '当前价格(元)', '目标价格(元)', '预期收益率(%)']
            scenario_data_stress = [
                ['当前估值', f"{current_metrics_val['pb']:.2f}", f"{current_metrics_val['pb']:.2f}",
                 f"{current_price_rel:.2f}", f"{current_price_rel:.2f}", "0.00"],
                [f'PB→行业Q1({pb_q1:.2f}倍)', f"{current_metrics_val['pb']:.2f}", f"{pb_q1:.2f}",
                 f"{current_price_rel:.2f}", f"{target_price_q1:.2f}", f"{return_q1:+.2f}"],
            ]
            add_table_data(document, scenario_headers_stress, scenario_data_stress)
            add_paragraph(document, f'• 如果PB回归到行业Q1({pb_q1:.2f}倍)，目标价格为{target_price_q1:.2f}元，预期收益{return_q1:+.2f}%')
            return_pe_stress = return_q1
        else:
            add_paragraph(document, 'PB数据同样不可用，跳过估值回归压力测试。')
    elif current_metrics_val['pe'] < 0:
        pe_val = current_metrics_val['pe']
        add_paragraph(document, f'当前PE为{pe_val:.2f}倍（负值），公司处于亏损状态，EPS为负。')
        add_paragraph(document, '负PE的PE回归分析在经济学意义上不适用（负EPS×行业正PE=负目标价）。')
        add_paragraph(document, '将以PB倍数回归作为替代压力测试指标。')
        if current_metrics_val['pb'] and current_metrics_val['pb'] > 0 and industry_stats_val['pb']['q1'] > 0:
            pb_q1 = industry_stats_val['pb']['q1']
            bps = current_price_rel / current_metrics_val['pb']
            target_price_q1 = bps * pb_q1
            return_q1 = (target_price_q1 - current_price_rel) / current_price_rel * 100
            scenario_headers_stress = ['情景', '当前PB', '回归后PB', '当前价格(元)', '目标价格(元)', '预期收益率(%)']
            scenario_data_stress = [
                ['当前估值', f"{current_metrics_val['pb']:.2f}", f"{current_metrics_val['pb']:.2f}",
                 f"{current_price_rel:.2f}", f"{current_price_rel:.2f}", "0.00"],
                [f'PB→行业Q1({pb_q1:.2f}倍)', f"{current_metrics_val['pb']:.2f}", f"{pb_q1:.2f}",
                 f"{current_price_rel:.2f}", f"{target_price_q1:.2f}", f"{return_q1:+.2f}"],
            ]
            add_table_data(document, scenario_headers_stress, scenario_data_stress)
            add_paragraph(document, f'• 如果PB回归到行业Q1({pb_q1:.2f}倍)，目标价格为{target_price_q1:.2f}元，预期收益{return_q1:+.2f}%')
            return_pe_stress = return_q1
        else:
            add_paragraph(document, 'PB数据同样不可用，跳过估值回归压力测试。')
    else:
        # PE = 股价 / EPS → EPS = 股价 / PE
        eps_rel = current_price_rel / current_metrics_val['pe']

        # 获取行业PE分位数数据
        pe_q1 = industry_stats_val['pe']['q1']  # 25分位（下四分位数）

        # 计算PE回归到行业Q1的极端情景
        target_price_q1 = eps_rel * pe_q1
        return_q1 = (target_price_q1 - current_price_rel) / current_price_rel * 100

        # 构建压力测试表格
        scenario_headers_stress = ['情景', '当前PE', '回归后PE', '当前价格(元)', '目标价格(元)', '预期收益率(%)']
        scenario_data_stress = [
            ['当前估值', f"{current_metrics_val['pe']:.2f}", f"{current_metrics_val['pe']:.2f}",
             f"{current_price_rel:.2f}", f"{current_price_rel:.2f}", "0.00"],
            [f'PE→行业Q1({pe_q1:.2f}倍)', f"{current_metrics_val['pe']:.2f}", f"{pe_q1:.2f}",
             f"{current_price_rel:.2f}", f"{target_price_q1:.2f}", f"{return_q1:+.2f}"],
        ]

        add_table_data(document, scenario_headers_stress, scenario_data_stress)

        add_paragraph(document, '')
        add_paragraph(document, 'PE回归压力测试分析：', bold=True)
        # 分析当前PE在行业中的位置
        current_pe = current_metrics_val['pe']
        pe_position = (peer_companies_val['pe'] < current_pe).sum() / len(peer_companies_val) * 100

        add_paragraph(document, f'• 当前PE({current_pe:.2f}倍)位于行业{pe_position:.1f}%分位')
        add_paragraph(document, f'• 行业Q1 PE({pe_q1:.2f}倍)为25分位数，代表行业较低估值水平')
        add_paragraph(document, f'• 如果PE回归到行业Q1，目标价格为{target_price_q1:.2f}元，预期收益{return_q1:+.2f}%')
        # 风险提示
        add_paragraph(document, '风险评估：', bold=True)
        if return_q1 > 0:
            add_paragraph(document, f' 即使在PE回归到行业Q1的极端情况下，预期收益仍为正({return_q1:+.2f}%)，估值安全边际较高')
        elif return_q1 > -10:
            add_paragraph(document, f' 在PE回归到行业Q1的极端情况下，预期收益为负({return_q1:+.2f}%)，存在一定估值回调风险，但风险可控')
        else:
            add_paragraph(document, f' 在PE回归到行业Q1的极端情况下，预期收益大幅为负({return_q1:+.2f}%)，估值回调风险较高，需谨慎投资')
        # 计算定增收益影响
        issue_price = project_params['issue_price']
        return_pe_stress = (target_price_q1 - issue_price) / issue_price * 100
        add_paragraph(document, '对定增收益的影响：')
        add_paragraph(document, f'• 发行价：{issue_price:.2f}元')
        add_paragraph(document, f'• PE回归Q1后的定增收益率：{return_pe_stress:+.2f}%')
        if return_pe_stress < 0:
            add_paragraph(document, f'•  估值回归将导致定增亏损{abs(return_pe_stress):.2f}%')
        else:
            add_paragraph(document, f'•  估值回归后定增仍能保持盈利')

    # ==================== 7.2 经济面极端情况压力测试 ====================
    add_paragraph(document, '')
    add_title(document, '7.2 经济面极端情况压力测试', level=2)

    add_paragraph(document, '本节模拟经济面极端情况（历史危机、黑天鹅事件）对定增项目的影响。')
    add_paragraph(document, '包括2008年金融危机、2020年疫情、行业政策收紧、个股重大利空等多种极端情景。')

    # 7.2.1 压力情景定义
    add_title(document, '7.2.1 压力情景定义', level=3)

    add_paragraph(document, '本节定义六种极端经济情景，基于历史事件和假设性风险构建。每种情景假设股价')
    add_paragraph(document, '在当前基础上下跌一定幅度，并伴随波动率飙升，用于评估定增项目在极端情况')
    add_paragraph(document, '下的最大损失和抗风险能力。')

    # 定义完整的压力测试情景（与notebook一致）
    stress_scenarios = {
        '市场危机_2008': {
            'description': '模拟2008年金融危机，股价下跌60%',
            'price_drop': 0.60,
            'volatility_spike': 2.0
        },
        '市场危机_2020': {
            'description': '模拟2020年疫情，股价下跌40%',
            'price_drop': 0.40,
            'volatility_spike': 1.5
        },
        '行业政策收紧': {
            'description': '行业监管政策收紧，股价下跌25%',
            'price_drop': 0.25,
            'volatility_spike': 1.2
        },
        '个股重大利空': {
            'description': '公司业绩暴雷，股价下跌35%',
            'price_drop': 0.35,
            'volatility_spike': 1.8
        },
        '流动性危机': {
            'description': '市场流动性枯竭，股价下跌20%并波动率飙升',
            'price_drop': 0.20,
            'volatility_spike': 2.5
        },
        '极端熊市': {
            'description': '极端熊市情景，股价下跌50%',
            'price_drop': 0.50,
            'volatility_spike': 1.3
        }
    }

    # 添加情景说明表格
    add_paragraph(document, '各极端情景的具体定义如下：')
    add_paragraph(document, '')

    scenario_headers = ['情景名称', '情景描述', '股价跌幅', '波动率倍数']
    scenario_table_data = [
        ['市场危机_2008', '2008年金融危机级别', '-60%', '2.0x'],
        ['市场危机_2020', '2020年疫情级别', '-40%', '1.5x'],
        ['极端熊市', '极端熊市情景', '-50%', '1.3x'],
        ['个股重大利空', '公司业绩暴雷', '-35%', '1.8x'],
        ['行业政策收紧', '行业监管收紧', '-25%', '1.2x'],
        ['流动性危机', '市场流动性枯竭', '-20%', '2.5x']
    ]
    add_table_data(document, scenario_headers, scenario_table_data)
    add_paragraph(document, '情景说明：')
    add_paragraph(document, '• 市场危机_2008：模拟2008年全球金融危机，A股市场暴跌约60%的极端情景')
    add_paragraph(document, '• 市场危机_2020：模拟2020年新冠疫情冲击，A股市场下跌约40%的情景')
    add_paragraph(document, '• 极端熊市：假设进入极端熊市周期，股价下跌50%，波动率上升30%')
    add_paragraph(document, '• 个股重大利空：公司突发重大利空（如财务造假、重大诉讼等），股价下跌35%')
    add_paragraph(document, '• 行业政策收紧：行业遭遇强监管政策（如教培、互联网金融等），股价下跌25%')
    add_paragraph(document, '• 流动性危机：市场流动性枯竭（如钱荒、信用违约等），股价下跌20%，波动率飙升150%')
    add_paragraph(document, '风险提示：')
    add_paragraph(document, '• 以上情景均为压力测试假设，不代表对未来市场的预测')
    add_paragraph(document, '• 实际极端情况可能比假设情景更严重或更轻微')
    add_paragraph(document, '• 投资决策应综合考虑多种情景下的风险敞口')

    # 定义完整的压力测试情景
    stress_scenarios = {
        '市场危机_2008': {
            'description': '模拟2008年金融危机，股价下跌60%',
            'price_drop': 0.60,
            'volatility_spike': 2.0
        },
        '市场危机_2020': {
            'description': '模拟2020年疫情，股价下跌40%',
            'price_drop': 0.40,
            'volatility_spike': 1.5
        },
        '行业政策收紧': {
            'description': '行业监管政策收紧，股价下跌25%',
            'price_drop': 0.25,
            'volatility_spike': 1.2
        },
        '个股重大利空': {
            'description': '公司业绩暴雷，股价下跌35%',
            'price_drop': 0.35,
            'volatility_spike': 1.8
        },
        '流动性危机': {
            'description': '市场流动性枯竭，股价下跌20%并波动率飙升',
            'price_drop': 0.20,
            'volatility_spike': 2.5
        },
        '极端熊市': {
            'description': '极端熊市情景，股价下跌50%',
            'price_drop': 0.50,
            'volatility_spike': 1.3
        }
    }

    # 计算各情景的结果
    scenario_names = []
    scenario_descriptions = []
    scenario_prices_list = []
    scenario_returns_list = []
    scenario_pnl_list = []

    current_price = market_data['current_price']
    issue_price = project_params['issue_price']
    issue_shares = project_params['issue_shares']

    for name, params in stress_scenarios.items():
        stressed_price = current_price * (1 - params['price_drop'])
        pnl_percent = ((stressed_price - issue_price) / issue_price) * 100
        pnl_amount = (stressed_price - issue_price) * issue_shares

        scenario_names.append(name)
        scenario_descriptions.append(params['description'])
        scenario_prices_list.append(stressed_price)
        scenario_returns_list.append(pnl_percent)
        scenario_pnl_list.append(pnl_amount)

    # 按照下跌幅度排序（从最严重到最轻）
    sorted_indices = sorted(range(len(scenario_names)), key=lambda i: scenario_prices_list[i])
    scenario_names = [scenario_names[i] for i in sorted_indices]
    scenario_descriptions = [scenario_descriptions[i] for i in sorted_indices]
    scenario_prices_list = [scenario_prices_list[i] for i in sorted_indices]
    scenario_returns_list = [scenario_returns_list[i] for i in sorted_indices]
    scenario_pnl_list = [scenario_pnl_list[i] for i in sorted_indices]

    # 保存图表
    stress_chart_path = os.path.join(IMAGES_DIR, '02_stress_test.png')
    generate_stress_test_chart(scenario_names, scenario_returns_list, stress_chart_path)

    # 二、压力测试结果
    add_title(document, '7.2.2 压力测试结果', level=3)

    # 生成详细的结果表格
    stress_table = []
    for i, name in enumerate(scenario_names):
        stress_table.append([
            name,
            scenario_descriptions[i],
            f'{scenario_prices_list[i]:.2f}',
            f'{scenario_returns_list[i]:+.2f}%',
            f'{scenario_pnl_list[i]/10000:+.2f}'
        ])

    add_table_data(document, ['情景', '描述', '受压后价格(元)', '收益率(%)', '盈亏(万元)'], stress_table)

    add_paragraph(document, '图表 7.1: 经济面极端情况压力测试')
    add_image(document, stress_chart_path)


    # ==================== 7.3 多重敏感性指标极端情况压力测试 ====================
    add_title(document, '7.3 多重敏感性指标极端情况压力测试', level=2)

    add_paragraph(document, '本节分析当多个敏感性指标同时发生极端情况时的最差情景。')
    add_paragraph(document, '通过组合最不利的参数（平价发行 + 高波动率 + 负向漂移率），评估项目的风险承受边界。')

    add_title(document, '7.3.1 极端情景组合定义', level=3)

    # 定义多重极端情景
    # 同时考虑：
    # 1. 平价发行（发行价等于MA20价格，无折价）
    # 2. 高波动率（当前波动率 × 1.5）
    # 3. 负向漂移率（当前漂移率 × 1.5，使情况更糟）

    current_vol_120d = market_data.get('volatility_120d', market_data.get('volatility', 0.30))
    current_drift_120d = market_data.get('annual_return_120d', market_data.get('drift', 0.08))
    current_price_multi = market_data['current_price']
    lockup_months = project_params['lockup_period']

    # 获取MA20价格，用于正确计算发行价
    ma20_price = market_data.get('ma_20', current_price_multi)
    premium_rate_normal = -0.10  # 正常情况：10%折价（九折）

    # 情景1：平价发行 + 高波动 + 负向漂移（三重打击）
    premium_rate_extreme = 0.0  # 0%溢价（平价发行，即发行价=MA20价格）
    vol_multiplier_extreme = 1.5  # 波动率放大1.5倍
    drift_extreme = current_drift_120d * 1.5  # 当前漂移率放大1.5倍（使情况更糟）

    # 正常情景：使用MA20的九折作为发行价
    issue_price_normal = ma20_price * (1 + premium_rate_normal)

    # 极端情景：平价发行（发行价等于MA20价格）
    issue_price_extreme = ma20_price * (1 + premium_rate_extreme)

    vol_extreme = current_vol_120d * vol_multiplier_extreme
    drift_rate_extreme = drift_extreme

    # 计算三重打击情景下的结果
    lockup_drift_extreme = drift_rate_extreme * (lockup_months / 12)
    lockup_vol_extreme = vol_extreme * np.sqrt(lockup_months / 12)

    # 使用蒙特卡洛模拟计算
    n_sim_extreme = 5000
    np.random.seed(42)
    returns_extreme = np.random.normal(lockup_drift_extreme, lockup_vol_extreme, n_sim_extreme)
    final_prices_extreme = current_price_multi * np.exp(returns_extreme)

    # 计算收益率
    returns_pnl_extreme = (final_prices_extreme - issue_price_extreme) / issue_price_extreme
    annualized_returns_extreme = (1 + returns_pnl_extreme) ** (12 / lockup_months) - 1

    # 统计结果（使用锁定期直接收益率，避免年化放大导致数值失真）
    mean_return_extreme = returns_pnl_extreme.mean() * 100
    median_return_extreme = np.median(returns_pnl_extreme) * 100
    profit_prob_extreme = (returns_pnl_extreme > 0).mean() * 100
    var_95_extreme = np.percentile(returns_pnl_extreme, 5) * 100
    var_99_extreme = np.percentile(returns_pnl_extreme, 1) * 100
    worst_loss_extreme = np.min(returns_pnl_extreme) * 100

    # 计算正常情景（使用正常发行价和正常参数）
    vol_normal = current_vol_120d
    drift_rate_normal = current_drift_120d

    lockup_drift_normal = drift_rate_normal * (lockup_months / 12)
    lockup_vol_normal = vol_normal * np.sqrt(lockup_months / 12)

    # 使用蒙特卡洛模拟计算正常情景
    np.random.seed(42)
    returns_normal = np.random.normal(lockup_drift_normal, lockup_vol_normal, n_sim_extreme)
    final_prices_normal = current_price_multi * np.exp(returns_normal)

    # 计算正常情景收益率
    returns_pnl_normal = (final_prices_normal - issue_price_normal) / issue_price_normal

    # 统计正常情景结果（使用锁定期直接收益率）
    mean_return_normal = returns_pnl_normal.mean() * 100
    median_return_normal = np.median(returns_pnl_normal) * 100
    profit_prob_normal = (returns_pnl_normal > 0).mean() * 100
    var_95_normal = np.percentile(returns_pnl_normal, 5) * 100
    var_99_normal = np.percentile(returns_pnl_normal, 1) * 100
    worst_loss_normal = np.min(returns_pnl_normal) * 100

    # 构建对比表格
    add_paragraph(document, '极端情景参数设置：')
    extreme_scenario_headers = ['参数', '正常值', '极端值', '变化幅度']
    extreme_scenario_data = [
        ['发行价', f'{issue_price_normal:.2f}元（MA20九折）', f'{issue_price_extreme:.2f}元（MA20平价）', '折价→平价'],
        ['波动率', f'{vol_normal*100:.2f}%', f'{vol_extreme*100:.2f}%', f'×{vol_multiplier_extreme:.1f}'],
        ['漂移率', f'{drift_rate_normal*100:+.2f}%', f'{drift_rate_extreme*100:+.2f}%', f'×{vol_multiplier_extreme:.1f}'],
        ['窗口期', '120日', '120日', '保持不变'],
        ['锁定期', f'{lockup_months}个月', f'{lockup_months}个月', '保持不变']
    ]
    add_table_data(document, extreme_scenario_headers, extreme_scenario_data)

    add_paragraph(document, '')
    add_title(document, '7.3.2 极端情景模拟结果', level=3)

    extreme_results_headers = ['指标', '正常情景', '极端情景（三重打击）', '差异']
    extreme_results_data = [
        ['预期收益', f'{mean_return_normal:+.2f}%', f'{mean_return_extreme:+.2f}%', f'{mean_return_extreme - mean_return_normal:+.2f}%'],
        ['中位数收益', f'{median_return_normal:+.2f}%', f'{median_return_extreme:+.2f}%', f'{median_return_extreme - median_return_normal:+.2f}%'],
        ['盈利概率', f'{profit_prob_normal:.1f}%', f'{profit_prob_extreme:.1f}%', f'{profit_prob_extreme - profit_prob_normal:+.1f}%'],
        ['95% VaR', f'{var_95_normal:+.2f}%', f'{var_95_extreme:+.2f}%', f'{var_95_extreme - var_95_normal:+.2f}%'],
        ['99% VaR', f'{var_99_normal:+.2f}%', f'{var_99_extreme:+.2f}%', f'{var_99_extreme - var_99_normal:+.2f}%'],
        ['最差情况', f'{worst_loss_normal:+.2f}%', f'{worst_loss_extreme:+.2f}%', f'{worst_loss_extreme - worst_loss_normal:+.2f}%']
    ]
    add_table_data(document, extreme_results_headers, extreme_results_data)

    add_paragraph(document, '')
    add_title(document, '7.3.3 极端情景分析', level=3)

    # 风险等级评估
    if profit_prob_extreme >= 40:
        extreme_risk_level = "中等风险 - 即使三重打击仍有40%以上盈利概率"
        risk_emoji_extreme = "🟡"
    elif profit_prob_extreme >= 20:
        extreme_risk_level = "高风险 - 盈利概率显著降低，但仍有20%以上机会"
        risk_emoji_extreme = "🟠"
    else:
        extreme_risk_level = "极高风险 - 三重打击下盈利概率低于20%，需极度谨慎"
        risk_emoji_extreme = ""

    add_paragraph(document, f'{risk_emoji_extreme} 极端情景评级: {extreme_risk_level}')
    add_paragraph(document, '')

    add_paragraph(document, '关键发现：')
    add_paragraph(document, f'• 在平价发行（发行价等于MA20价格{ma20_price:.2f}元）、高波动（×1.5）、负向漂移（{drift_extreme*100:+.2f}%）的三重打击下：')
    add_paragraph(document, f'  - 预期收益率为{mean_return_extreme:+.2f}%')
    add_paragraph(document, f'  - 盈利概率为{profit_prob_extreme:.1f}%')
    add_paragraph(document, f'  - 95% VaR为{var_95_extreme:+.2f}%，表示95%置信度下最大损失为{abs(var_95_extreme):.2f}%')
    add_paragraph(document, f'  - 最差情况可能损失{abs(worst_loss_extreme):.2f}%')
    
    # ==================== 7.4 压力测试综合结论 ====================
    add_paragraph(document, '')
    add_title(document, '7.4 压力测试综合结论', level=2)

    add_paragraph(document, '本节综合前面7.1、7.2、7.3的压力测试结果，总结定增项目在各类极端情况下的风险表现。')
    add_paragraph(document, '')

    add_title(document, '7.4.1 压力测试全景汇总', level=3)

    # 计算7.2经济面极端情况的统计数据
    # 最差情景是排序后的第一个（价格最低）
    worst_scenario_idx = 0  # 已按价格从低到高排序
    worst_scenario_name = scenario_names[worst_scenario_idx]
    worst_loss_percent = scenario_returns_list[worst_scenario_idx]
    worst_loss = scenario_pnl_list[worst_scenario_idx]

    # 计算盈利情景数量
    profit_scenarios = sum(1 for ret in scenario_returns_list if ret > 0)
    total_scenarios = len(scenario_names)

    # 创建综合汇总表
    summary_headers = ['压力测试类型', '测试情景', '最差结果', '风险评估', '风险等级']
    summary_data = [
        ['7.1 估值回归压力测试',
         f'PE回归至行业Q1({pe_q1:.2f}倍)' if pe_q1 else 'PB回归至行业Q1（PE不可用）',
         f'{return_pe_stress:+.2f}%',
         '估值回归风险' if return_pe_stress < 0 else '估值安全边际充足',
         '高风险' if return_pe_stress <= -10 else '中等风险' if return_pe_stress <= 0 else '低风险'],
        ['7.2 经济面极端情况', f'{worst_scenario_name}（股价下跌{int(stress_scenarios[worst_scenario_name]["price_drop"]*100)}%）',
         f'{worst_loss_percent:+.2f}%', f'最大亏损{abs(worst_loss):.2f}万元',
         '高风险' if profit_scenarios <= total_scenarios * 0.4 else '中等风险' if profit_scenarios <= total_scenarios * 0.7 else '低风险'],
        ['7.3 多重敏感性指标极端', '平价发行+高波动+负向漂移', f'{worst_loss_extreme:+.2f}%',
         f'盈利概率{profit_prob_extreme:.1f}%',
         '极高风险' if profit_prob_extreme < 20 else '高风险' if profit_prob_extreme < 40 else '中等风险' if profit_prob_extreme < 60 else '低风险']
    ]
    add_table_data(document, summary_headers, summary_data)

    add_paragraph(document, '')
    add_section_break(document)

    return context
