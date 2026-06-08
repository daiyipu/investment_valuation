# -*- coding: utf-8 -*-
"""
第二章：相对估值分析

本章节生成报告的相对估值分析部分，包括：
- 2.1 估值指标对比（PE、PB、PS）
- 2.1.1 同行公司名单
- 2.2 估值偏离度分析
- 2.3 PE历史分位数趋势分析
- 2.4 PE估值分析（绝对估值与修正估值）
"""

import os
import json
import time


import pandas as pd
import numpy as np
from datetime import datetime, timedelta
from docx.shared import Inches

# 导入工具函数
from module_utils import (
    add_title, add_paragraph, add_table_data, add_image, add_section_break,
    generate_relative_valuation_charts_split
)


def generate_chapter(context):
    """
    生成第二章：相对估值分析

    Args:
        context: 包含以下键的字典:
            - document: Word文档对象
            - project_params: 项目参数（包含stock_code等）
            - IMAGES_DIR: 图片目录

    Returns:
        更新后的context字典
    """
    # 从context中提取数据
    document = context['document']
    project_params = context['project_params']
    IMAGES_DIR = context['IMAGES_DIR']

    stock_code = project_params.get('stock_code', '')  # 从project_params获取（与V2一致）
    stock_name = context.get('stock_name', stock_code)
    _script_dir = os.path.dirname(os.path.abspath(__file__))
    DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(_script_dir)), 'data')
    cache_file = os.path.join(DATA_DIR, f"{stock_code.replace('.', '_')}_relative_valuation.json")

    # ==================== 二、相对估值分析 ====================
    add_title(document, '二、相对估值分析', level=1)

    add_paragraph(document, f'本章节通过相对估值法（参数法），将{stock_name}与行业内可比公司进行对比分析。')
    sw_l3_name = context.get('industry_data', {}).get('sw_l3_name', '行业')
    add_paragraph(document, f'选取申万三级分类"{sw_l3_name}"行业的同行公司，对比PE、PS、PB等估值倍数。')

    add_title(document, '2.1 估值指标对比', level=2)

    # 使用 Tushare 数据获取估值指标
    # 优先从DB读取缓存（仅当天创建的有效）
    _cache_loaded = False
    today_str = datetime.now().strftime('%Y%m%d')

    try:
        from utils.db_manager import ValuationDB
        db = ValuationDB()
        if db.is_cache_valid(stock_code):
            cached = db.load_relative_valuation(stock_code)
            if cached:
                current_metrics_val = cached.get('current_metrics', {})
                peer_data = cached.get('peer_companies', [])
                peer_companies_val = pd.DataFrame(peer_data) if peer_data else pd.DataFrame()
                sw_index_pe = cached.get('sw_index_pe')
                sw_index_pb = cached.get('sw_index_pb')
                sw_index_ps = cached.get('sw_index_ps')
                target_index_code = cached.get('target_index_code')
                target_industry_l3 = cached.get('target_industry_l3')
                trade_date = cached.get('trade_date')
                print(f"  使用DB缓存的相对估值数据（交易日期: {trade_date}）")
                _cache_loaded = True
    except Exception:
        pass

    # 回退到JSON文件缓存
    if not _cache_loaded and os.path.exists(cache_file):
        try:
            with open(cache_file, 'r', encoding='utf-8') as f:
                cached = json.load(f)
            if cached.get('cache_date') == today_str:
                current_metrics_val = cached['current_metrics']
                peer_companies_val = pd.DataFrame(cached['peer_companies'])
                sw_index_pe = cached.get('sw_index_pe')
                sw_index_pb = cached.get('sw_index_pb')
                sw_index_ps = cached.get('sw_index_ps')
                target_index_code = cached.get('target_index_code')
                target_industry_l3 = cached.get('target_industry_l3')
                trade_date = cached.get('trade_date')
                print(f"  使用文件缓存的相对估值数据（交易日期: {trade_date}）")
                _cache_loaded = True
        except (json.JSONDecodeError, KeyError):
            pass

    if not _cache_loaded:
        sw_index_pe = None
        sw_index_pb = None
        sw_index_ps = None
        target_index_code = None
        target_industry_l3 = None

        try:
            ts_token = os.environ.get('TUSHARE_TOKEN', '')

            if ts_token:
                import tushare as ts

                pro = ts.pro_api(ts_token)

                # 获取目标公司的估值数据（自动往前推1-2天直到找到交易日）
                trade_date = None
                df_target = None

                for days_back in range(1, 16):  # 尝试往前推1-15天（覆盖长假）
                    test_date = (datetime.now() - timedelta(days=days_back)).strftime('%Y%m%d')
                    try:
                        df_target = pro.daily_basic(
                            ts_code=stock_code,
                            trade_date=test_date,
                            fields='ts_code,trade_date,close,pe_ttm,pb,ps_ttm,total_mv'
                        )
                        if not df_target.empty:
                            trade_date = test_date
                            break
                    except:
                        continue

                if df_target is None or df_target.empty:
                    raise ValueError("未获取到目标公司数据（请检查网络或交易日历）")

                current_metrics_val = {
                    'pe': float(df_target.iloc[0]['pe_ttm']) if pd.notna(df_target.iloc[0]['pe_ttm']) else None,
                    'pb': float(df_target.iloc[0]['pb']) if pd.notna(df_target.iloc[0]['pb']) else None,
                    'ps': float(df_target.iloc[0]['ps_ttm']) if pd.notna(df_target.iloc[0]['ps_ttm']) else None
                }
                print(f" 获取相对估值数据成功，交易日期: {trade_date}")

                # 获取申万三级行业分类的同行公司（与 notebook 一致）
                df_industry = pro.index_member_all(ts_code=stock_code)
                if df_industry.empty:
                    raise ValueError("未获取到行业分类")

                # 过滤：只保留该股票的记录（防止API返回过多数据）
                df_industry = df_industry[df_industry['ts_code'] == stock_code]
                if df_industry.empty:
                    raise ValueError(f"未找到{stock_code}的行业分类记录")

                # 显示所有行业分类记录，方便调试
                df_industry = df_industry.sort_values('in_date', ascending=False)
                print(f" 获取到{len(df_industry)}条行业分类记录:")

                for idx, row in df_industry.head(5).iterrows():
                    print(f"   [{idx}] {row['in_date']}: 一级={row.get('index_name', 'N/A')}, L1={row.get('l1_name', 'N/A')}, L2={row.get('l2_name', 'N/A')}, L3={row.get('l3_name', 'N/A')}")
                    print(f"        L1代码={row.get('l1_code', 'N/A')}, L2代码={row.get('l2_code', 'N/A')}, L3代码={row.get('l3_code', 'N/A')}")

                latest_industry = df_industry.iloc[0]

                # 调试输出
                print(f"\n 使用最新记录:")
                print(f"   股票代码: {stock_code}")
                print(f"   申万一级: {latest_industry.get('index_name', 'N/A')}")
                print(f"   申万三级代码: {latest_industry['l3_code']}")
                print(f"   申万三级名称: {latest_industry['l3_name']}")

                target_index_code = latest_industry['l3_code']  # 申万三级行业指数代码
                target_industry_l3 = latest_industry['l3_name']  # 行业名称

                # 获取该三级行业的所有成分股
                print(f"\n 正在使用指数代码 {target_index_code} 查询成分股...")

                df_peers = pro.index_member_all(l3_code=target_index_code)
                print(f" 获取到 {len(df_peers)} 条成分股记录")

                df_peers = df_peers[df_peers['ts_code'] != stock_code]

                # 获取同行公司基本信息
                peer_codes = df_peers['ts_code'].unique().tolist()
                print(f" 过滤后剩余 {len(peer_codes)} 个同行公司")

                peer_basic = pro.stock_basic(ts_code=','.join(peer_codes[:30]),
                                           fields='ts_code,name,market')

                peer_stocks_all = pd.merge(df_peers, peer_basic, on='ts_code', how='left')
                peer_stocks_all = peer_stocks_all.drop_duplicates(subset=['ts_code'])

                # 限制数量并排序（扩充到30家）
                peer_stocks_all = peer_stocks_all.head(30)
                peer_names_dict = dict(zip(peer_stocks_all['ts_code'], peer_stocks_all['name_x']))

                # 获取同行公司的估值数据（一次全市场查询替代逐只循环）
                peer_data_list = []
                try:
                    df_all = pro.daily_basic(
                        trade_date=trade_date,
                        fields='ts_code,pe_ttm,pb,ps_ttm,total_mv'
                    )
                    if not df_all.empty:
                        peer_codes = peer_stocks_all['ts_code'].tolist()
                        df_peers = df_all[df_all['ts_code'].isin(peer_codes)].copy()
                        # 补充名称
                        df_peers['name'] = df_peers['ts_code'].map(peer_names_dict)
                        peer_data_list = [df_peers]
                        print(f"  一次批量获取 {len(df_peers)}/{len(peer_codes)} 家同行估值数据")
                except Exception as e:
                    # 降级：逐只查询
                    print(f"  批量查询失败({e})，降级逐只查询...")
                    for peer_code in peer_stocks_all['ts_code'].tolist():
                        peer_name = peer_names_dict.get(peer_code, peer_code)
                        try:
                            df_peer = pro.daily_basic(
                                ts_code=peer_code,
                                trade_date=trade_date,
                                fields='ts_code,pe_ttm,pb,ps_ttm,total_mv'
                            )
                            if not df_peer.empty:
                                df_peer['name'] = peer_name
                                peer_data_list.append(df_peer)
                        except:
                            pass
                        time.sleep(0.2)

                # 获取申万行业指数的PE数据（L3无数据时降级L2/L1，最后AKShare兜底）
                try:
                    print(f" 正在获取申万行业指数PE数据: {target_index_code}")
                    df_index = None
                    # 收集L3/L2/L1候选代码（部分L3如农业综合Ⅲ无sw_daily数据）
                    candidate_codes = []
                    for lv in ('l3_code', 'l2_code', 'l1_code'):
                        code = latest_industry.get(lv)
                        if code:
                            candidate_codes.append((lv, code))
                    if not candidate_codes:
                        candidate_codes = [('l3_code', target_index_code)]

                    end_date = datetime.now().strftime('%Y%m%d')
                    start_date = (datetime.now() - timedelta(days=10)).strftime('%Y%m%d')

                    # 逐级尝试 DB缓存 → sw_daily
                    for lv, cand_code in candidate_codes:
                        # 1. DB缓存
                        try:
                            from utils.db_manager import ValuationDB
                            db = ValuationDB()
                            df_index = db.load_industry_daily(cand_code, start_date, end_date)
                            if df_index is not None and not df_index.empty:
                                if cand_code != target_index_code:
                                    print(f"  L3无数据，降级使用 {lv.replace('_code','').upper()} {cand_code}")
                                else:
                                    print(f"  使用DB缓存的行业PE数据({cand_code})")
                                break
                        except Exception:
                            pass
                        # 2. sw_daily
                        try:
                            df_index = pro.sw_daily(ts_code=cand_code, start_date=start_date, end_date=end_date)
                            if isinstance(df_index, str):
                                df_index = None
                            if df_index is not None and not df_index.empty:
                                if cand_code != target_index_code:
                                    print(f"  L3无数据，降级使用 {lv.replace('_code','').upper()} {cand_code}")
                                break
                        except Exception:
                            df_index = None

                    # 从sw_daily/DB结果提取PE/PB/PS
                    if df_index is not None and not df_index.empty:
                        latest = df_index.iloc[-1]
                        sw_index_pe = latest.get('pe', None)
                        sw_index_pb = latest.get('pb', None)
                        sw_index_ps = latest.get('ps_ttm', None)
                        if sw_index_pe:
                            print(f"  sw_daily PE数据获取成功: PE={sw_index_pe:.2f}, PB={sw_index_pb:.2f}")
                    # 3. 全部失败 → AKShare申万三级PE兜底
                    if not sw_index_pe:
                        try:
                            import akshare as ak
                            ths_info = ak.sw_index_third_info()
                            matched = ths_info[ths_info['行业代码'] == target_index_code]
                            if not matched.empty:
                                row = matched.iloc[0]
                                sw_index_pe = row.get('TTM(滚动)市盈率', None)
                                sw_index_pb = row.get('市净率', None)
                                sw_index_ps = None  # AKShare不提供PS
                                if sw_index_pe:
                                    print(f"  使用AKShare申万三级PE数据(兜底): PE={sw_index_pe:.2f}, PB={sw_index_pb:.2f}")
                        except Exception as e:
                            print(f"  AKShare PE获取失败: {e}")
                    if not sw_index_pe:
                        print(f"  申万行业指数PE数据为空，将跳过行业PE对比")
                except Exception as e:
                    print(f" 获取申万行业指数PE数据失败: {e}")

                if peer_data_list:
                    peer_companies_val = pd.concat(peer_data_list, ignore_index=True)

                    # 过滤异常数据
                    peer_filtered = peer_companies_val[
                        (peer_companies_val['pe_ttm'] > 0) &
                        (peer_companies_val['pe_ttm'] < 500) &
                        (peer_companies_val['pb'] > 0) &
                        (peer_companies_val['pb'] < 20)
                    ]
                    # 如果过滤后为空（小行业成分股少），使用原始数据
                    if peer_filtered.empty and not peer_companies_val.empty:
                        print(f"  ⚠️ 行业成分股较少({len(peer_companies_val)}家)，使用未过滤数据")
                        peer_filtered = peer_companies_val
                    peer_companies_val = peer_filtered

                    # 重命名列并进行单位转换
                    peer_companies_val['market_cap'] = peer_companies_val['total_mv'] / 10000

                    peer_companies_val = peer_companies_val.rename(columns={
                        'pe_ttm': 'pe',
                        'ps_ttm': 'ps',
                        'ts_code': 'code'
                    })

                    peer_companies_val = peer_companies_val[['name', 'code', 'pe', 'ps', 'pb', 'market_cap']]
                else:
                    # 无同行数据，用目标公司自身数据作为 fallback
                    print(f"  ⚠️ 未获取到同行公司数据，使用目标公司自身数据")
                    context['industry_data_warning'] = '行业成分股不足，无法进行同行对比分析，以下行业统计数据仅基于目标公司自身数据，仅供参考。'
                    peer_companies_val = pd.DataFrame([{
                        'name': stock_name,
                        'code': stock_code,
                        'pe': current_metrics_val.get('pe', 0),
                        'ps': current_metrics_val.get('ps', 0),
                        'pb': current_metrics_val.get('pb', 0),
                        'market_cap': current_metrics_val.get('market_cap', 0),
                    }])

            else:
                raise ValueError("TUSHARE_TOKEN 未设置")

            # API 成功，保存缓存到DB和文件
            try:
                cache_data = {
                    'cache_date': today_str,
                    'trade_date': trade_date,
                    'current_metrics': current_metrics_val,
                    'peer_companies': peer_companies_val.to_dict('records'),
                    'sw_index_pe': sw_index_pe,
                    'sw_index_pb': sw_index_pb,
                    'sw_index_ps': sw_index_ps,
                    'target_index_code': target_index_code,
                    'target_industry_l3': target_industry_l3,
                }
                # 保存到DB
                try:
                    from utils.db_manager import ValuationDB
                    db = ValuationDB()
                    db.save_relative_valuation(stock_code, cache_data)
                    print(f"  已缓存相对估值数据到DB")
                except Exception:
                    pass
                _cache_loaded = True
            except Exception:
                pass

        except Exception as e:
            print(f"获取相对估值数据失败: {e}，尝试使用缓存")

        # 尝试使用DB缓存（API不可用时的降级方案）
        if not _cache_loaded:
            try:
                from utils.db_manager import ValuationDB
                db = ValuationDB()
                cached = db.load_relative_valuation(stock_code)
                if cached:
                    current_metrics_val = cached.get('current_metrics', {})
                    peer_data = cached.get('peer_companies', [])
                    peer_companies_val = pd.DataFrame(peer_data) if peer_data else pd.DataFrame()
                    sw_index_pe = cached.get('sw_index_pe')
                    sw_index_pb = cached.get('sw_index_pb')
                    sw_index_ps = cached.get('sw_index_ps')
                    target_index_code = cached.get('target_index_code')
                    target_industry_l3 = cached.get('target_industry_l3')
                    trade_date = cached.get('trade_date')
                    print(f"  使用DB缓存的相对估值数据（交易日期: {trade_date}）")
                    _cache_loaded = True
            except Exception:
                pass

        # 回退到JSON文件缓存
        if not _cache_loaded and os.path.exists(cache_file):
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cached = json.load(f)
                current_metrics_val = cached['current_metrics']
                peer_companies_val = pd.DataFrame(cached['peer_companies'])
                sw_index_pe = cached.get('sw_index_pe')
                sw_index_pb = cached.get('sw_index_pb')
                sw_index_ps = cached.get('sw_index_ps')
                target_index_code = cached.get('target_index_code')
                target_industry_l3 = cached.get('target_industry_l3')
                trade_date = cached.get('trade_date')
                print(f"  使用文件缓存的相对估值数据（交易日期: {trade_date}）")
                _cache_loaded = True
            except (json.JSONDecodeError, KeyError):
                pass

        if not _cache_loaded:
            raise ValueError(f"无法获取相对估值数据且无缓存，请检查网络连接")

    # 计算行业统计指标（填充NaN，避免成分股过少时崩溃）
    _pe = peer_companies_val['pe'].fillna(0)
    _ps = peer_companies_val['ps'].fillna(0)
    _pb = peer_companies_val['pb'].fillna(0)
    industry_stats_val = {
        'pe': {
            'mean': float(_pe.mean()),
            'median': float(_pe.median()),
            'q1': float(_pe.quantile(0.25)),
            'q3': float(_pe.quantile(0.75)),
            'min': float(_pe.min()),
            'max': float(_pe.max()),
            'std': float(_pe.std()) if len(_pe) > 1 else 0.0
        },
        'ps': {
            'mean': float(_ps.mean()),
            'median': float(_ps.median()),
            'q1': float(_ps.quantile(0.25)),
            'q3': float(_ps.quantile(0.75)),
            'min': float(_ps.min()),
            'max': float(_ps.max()),
            'std': float(_ps.std()) if len(_ps) > 1 else 0.0
        },
        'pb': {
            'mean': float(_pb.mean()),
            'median': float(_pb.median()),
            'q1': float(_pb.quantile(0.25)),
            'q3': float(_pb.quantile(0.75)),
            'min': float(_pb.min()),
            'max': float(_pb.max()),
            'std': float(_pb.std()) if len(_pb) > 1 else 0.0
        }
    }

    # 剔除3倍标准差异常值后的行业平均
    industry_avg_val = {
        'pe': industry_stats_val['pe']['mean'],
        'ps': industry_stats_val['ps']['mean'],
        'pb': industry_stats_val['pb']['mean']
    }

    # 估值指标对比表
    def _fmt_val(v, suffix='倍'):
        return f"{v:.2f}{suffix}" if v else 'N/A'

    def _fmt_dev(val, ref):
        if val and ref and ref != 0:
            return f"{(val - ref) / ref * 100:+.1f}%"
        return 'N/A'

    valuation_headers = ['指标', stock_name, '行业平均', '中位数', 'Q1(25分位)', 'Q3(75分位)', '最小值', '最大值', '偏离度']
    valuation_data = [
        ['PE (TTM)',
         _fmt_val(current_metrics_val['pe']),
         _fmt_val(industry_stats_val['pe']['mean']),
         _fmt_val(industry_stats_val['pe']['median']),
         _fmt_val(industry_stats_val['pe']['q1']),
         _fmt_val(industry_stats_val['pe']['q3']),
         _fmt_val(industry_stats_val['pe']['min']),
         _fmt_val(industry_stats_val['pe']['max']),
         _fmt_dev(current_metrics_val['pe'], industry_stats_val['pe']['mean'])],
        ['PB',
         _fmt_val(current_metrics_val['pb']),
         _fmt_val(industry_stats_val['pb']['mean']),
         _fmt_val(industry_stats_val['pb']['median']),
         _fmt_val(industry_stats_val['pb']['q1']),
         _fmt_val(industry_stats_val['pb']['q3']),
         _fmt_val(industry_stats_val['pb']['min']),
         _fmt_val(industry_stats_val['pb']['max']),
         _fmt_dev(current_metrics_val['pb'], industry_stats_val['pb']['mean'])],
        ['PS (TTM)',
         _fmt_val(current_metrics_val['ps']),
         _fmt_val(industry_stats_val['ps']['mean']),
         _fmt_val(industry_stats_val['ps']['median']),
         _fmt_val(industry_stats_val['ps']['q1']),
         _fmt_val(industry_stats_val['ps']['q3']),
         _fmt_val(industry_stats_val['ps']['min']),
         _fmt_val(industry_stats_val['ps']['max']),
         _fmt_dev(current_metrics_val['ps'], industry_stats_val['ps']['mean'])]
    ]
    add_table_data(document, valuation_headers, valuation_data)

    # 行业数据不足警告
    if context.get('industry_data_warning'):
        add_paragraph(document, '')
        add_paragraph(document, context['industry_data_warning'], bold=True)

    # 添加统计分析说明
    add_paragraph(document, '')
    add_paragraph(document, '行业估值统计说明：')
    add_paragraph(document, '• 行业平均：所有同行公司的算术平均值（受极端值影响）')
    add_paragraph(document, '• 中位数：行业50%的公司估值低于此水平，抗极端值干扰')
    add_paragraph(document, '• Q1(25分位)：行业25%的公司估值低于此水平')
    add_paragraph(document, '• Q3(75分位)：行业75%的公司估值低于此水平（即25%的公司高于此水平）')
    add_paragraph(document, '• 最小/最大值：行业中的估值极值')
    add_paragraph(document, '• 数据已过滤异常值（PE<500, PB<20）以避免极端情况影响统计')
    add_paragraph(document, f'• 样本量：共{len(peer_companies_val)}家同行公司')

    # 添加同行公司名单
    add_paragraph(document, '')
    add_title(document, '2.1.1 同行公司名单', level=3)
    add_paragraph(document, f'基于申万三级行业分类"{sw_l3_name}"筛选的同行公司：')

    # 按市值排序的同行公司表格
    peer_companies_sorted = peer_companies_val.sort_values('market_cap', ascending=False)
    peer_headers = ['公司名称', '股票代码', 'PE (TTM)', 'PB', 'PS (TTM)', '市值(亿元)']
    peer_rows = []
    for _, row in peer_companies_sorted.iterrows():
        peer_rows.append([
            row['name'],
            row['code'],
            f"{row['pe']:.2f}",
            f"{row['pb']:.2f}",
            f"{row['ps']:.2f}",
            f"{row['market_cap']:.2f}"
        ])
    add_table_data(document, peer_headers, peer_rows)

    # 添加行业统计汇总
    add_paragraph(document, '')
    add_paragraph(document, '行业估值统计汇总：')
    add_paragraph(document, f'• PE: 平均{industry_stats_val["pe"]["mean"]:.2f}倍，中位数{peer_companies_val["pe"].median():.2f}倍，标准差{industry_stats_val["pe"]["std"]:.2f}倍')
    add_paragraph(document, f'  • Q1-Q3区间: [{industry_stats_val["pe"]["q1"]:.2f}, {industry_stats_val["pe"]["q3"]:.2f}]倍，极值范围: [{industry_stats_val["pe"]["min"]:.2f}, {industry_stats_val["pe"]["max"]:.2f}]倍')
    add_paragraph(document, f'• PB: 平均{industry_stats_val["pb"]["mean"]:.2f}倍，中位数{peer_companies_val["pb"].median():.2f}倍，标准差{industry_stats_val["pb"]["std"]:.2f}倍')
    add_paragraph(document, f'  • Q1-Q3区间: [{industry_stats_val["pb"]["q1"]:.2f}, {industry_stats_val["pb"]["q3"]:.2f}]倍，极值范围: [{industry_stats_val["pb"]["min"]:.2f}, {industry_stats_val["pb"]["max"]:.2f}]倍')
    add_paragraph(document, f'• PS: 平均{industry_stats_val["ps"]["mean"]:.2f}倍，中位数{peer_companies_val["ps"].median():.2f}倍，标准差{industry_stats_val["ps"]["std"]:.2f}倍')
    add_paragraph(document, f'  • Q1-Q3区间: [{industry_stats_val["ps"]["q1"]:.2f}, {industry_stats_val["ps"]["q3"]:.2f}]倍，极值范围: [{industry_stats_val["ps"]["min"]:.2f}, {industry_stats_val["ps"]["max"]:.2f}]倍')

    add_paragraph(document, '图表 2.0: 相对估值对比分析 - 估值指标对比')
    chart_paths, df_scenarios = generate_relative_valuation_charts_split(
        current_metrics_val, industry_avg_val, peer_companies_val, IMAGES_DIR, stock_name=stock_name
    )
    add_image(document, chart_paths[0])  # 估值指标对比

    add_paragraph(document, '图表 2.1: 相对估值对比分析 - PE倍数对比')
    add_image(document, chart_paths[1])

    add_paragraph(document, '图表 2.2: 相对估值对比分析 - PB倍数对比')
    add_image(document, chart_paths[2])

    add_paragraph(document, '图表 2.3: 相对估值对比分析 - PS倍数对比')
    add_image(document, chart_paths[3])

    # ==================== 2.1.2 申万行业指数估值 ====================
    add_title(document, '2.1.2 申万行业指数估值', level=3)
    add_paragraph(document, '')

    if sw_index_pe is not None:
        add_paragraph(document, f'本节展示申万三级行业指数"{target_industry_l3}"的估值数据，提供官方行业指数视角的估值基准。')
        add_paragraph(document, f'申万行业指数代码: {target_index_code}')
        add_paragraph(document, '')

        # 申万行业指数估值表格
        sw_index_headers = ['指标', '申万行业指数', stock_name, '差异', '说明']
        sw_index_data = [
            ['PE (TTM)',
             _fmt_val(sw_index_pe),
             _fmt_val(current_metrics_val['pe']),
             _fmt_dev(current_metrics_val['pe'], sw_index_pe),
             '行业指数PE反映行业整体估值水平' if current_metrics_val['pe'] and current_metrics_val['pe'] > sw_index_pe else '个股PE低于行业指数，相对低估'],
            ['PB',
             _fmt_val(sw_index_pb),
             _fmt_val(current_metrics_val['pb']),
             _fmt_dev(current_metrics_val['pb'], sw_index_pb),
             '市净率反映行业整体账面价值溢价'],
            ['PS (TTM)',
             _fmt_val(sw_index_ps),
             _fmt_val(current_metrics_val['ps']),
             _fmt_dev(current_metrics_val['ps'], sw_index_ps),
             '市销率反映行业整体营收能力']
        ]
        add_table_data(document, sw_index_headers, sw_index_data)

        add_paragraph(document, '')
        add_paragraph(document, '申万行业指数估值说明：')
        add_paragraph(document, f'• 申万行业指数是基于该行业所有成分股按市值加权计算的指数')
        add_paragraph(document, f'• 指数PE/PB/PS反映行业整体的估值水平，不同于同行公司平均（简单平均）')
        add_paragraph(document, f'• 指数估值受大盘股权重影响，更能代表行业龙头公司的估值水平')
        add_paragraph(document, f'• 同行公司平均反映行业内典型公司的估值，受小盘股影响较大')
        add_paragraph(document, '')

        # 对比分析
        add_paragraph(document, ' 对比分析：', bold=True)
        if current_metrics_val['pe'] and sw_index_pe and sw_index_pe > 0:
            if abs(current_metrics_val['pe'] - sw_index_pe) / sw_index_pe < 0.1:
                add_paragraph(document, f' 个股PE({_fmt_val(current_metrics_val["pe"])})与申万行业指数PE({_fmt_val(sw_index_pe)})基本一致，估值合理')
            elif current_metrics_val['pe'] < sw_index_pe:
                add_paragraph(document, f' 个股PE({_fmt_val(current_metrics_val["pe"])})低于申万行业指数PE({_fmt_val(sw_index_pe)})，相对行业指数低估')
            else:
                add_paragraph(document, f' 个股PE({_fmt_val(current_metrics_val["pe"])})高于申万行业指数PE({_fmt_val(sw_index_pe)})，相对行业指数高估')
    else:
        add_paragraph(document, f' 申万行业指数"{target_industry_l3}"的估值数据暂时无法获取')

    add_title(document, '2.2 估值偏离度分析', level=2)
    add_paragraph(document, '本节分析标的公司与同行公司和申万行业指数的估值偏离情况，评估估值相对位置。')
    add_paragraph(document, '')

    # 计算PE在行业中的分位数位置
    n_peers = max(len(peer_companies_val), 1)
    pe_position = (peer_companies_val['pe'] < current_metrics_val['pe']).sum() / n_peers * 100 if current_metrics_val['pe'] else 0
    pb_position = (peer_companies_val['pb'] < current_metrics_val['pb']).sum() / n_peers * 100 if current_metrics_val['pb'] else 0
    ps_position = (peer_companies_val['ps'] < current_metrics_val['ps']).sum() / n_peers * 100 if current_metrics_val['ps'] else 0

    # 2.2.1 与同行公司对比
    add_title(document, '2.2.1 与同行公司对比', level=3)
    add_paragraph(document, '')

    add_paragraph(document, f"• PE偏离度: {_fmt_dev(current_metrics_val['pe'], industry_avg_val['pe'])}，位于行业{pe_position:.1f}%分位")
    add_paragraph(document, f"• PB偏离度: {_fmt_dev(current_metrics_val['pb'], industry_avg_val['pb'])}，位于行业{pb_position:.1f}%分位")
    add_paragraph(document, f"• PS偏离度: {_fmt_dev(current_metrics_val['ps'], industry_avg_val['ps'])}，位于行业{ps_position:.1f}%分位")

    add_paragraph(document, '')

    # PE分位数分析
    cur_pe = current_metrics_val['pe']
    if cur_pe is not None and industry_stats_val['pe']['q3'] is not None:
        if cur_pe > industry_stats_val['pe']['q3']:
            add_paragraph(document, f' PE({_fmt_val(cur_pe)})高于行业Q3({_fmt_val(industry_stats_val["pe"]["q3"])})，处于行业高位，估值偏高')
        elif cur_pe < industry_stats_val['pe']['q1']:
            add_paragraph(document, f' PE({_fmt_val(cur_pe)})低于行业Q1({_fmt_val(industry_stats_val["pe"]["q1"])})，处于行业低位，估值偏低')
        else:
            add_paragraph(document, f' PE({_fmt_val(cur_pe)})介于行业Q1({_fmt_val(industry_stats_val["pe"]["q1"])})和Q3({_fmt_val(industry_stats_val["pe"]["q3"])})之间，估值合理')
    elif cur_pe is not None:
        add_paragraph(document, f' PE({_fmt_val(cur_pe)})，行业分位数据不完整，无法进行分位分析')

    # PB分位数分析
    cur_pb = current_metrics_val['pb']
    if cur_pb is not None and industry_stats_val['pb']['q3'] is not None:
        if cur_pb > industry_stats_val['pb']['q3']:
            add_paragraph(document, f' PB({_fmt_val(cur_pb)})高于行业Q3({_fmt_val(industry_stats_val["pb"]["q3"])})，市净率偏高')
        elif cur_pb < industry_stats_val['pb']['q1']:
            add_paragraph(document, f' PB({_fmt_val(cur_pb)})低于行业Q1({_fmt_val(industry_stats_val["pb"]["q1"])})，市净率偏低')

    # PS分位数分析
    cur_ps = current_metrics_val['ps']
    if cur_ps is not None and industry_stats_val['ps']['q3'] is not None:
        if cur_ps > industry_stats_val['ps']['q3']:
            add_paragraph(document, f' PS({_fmt_val(cur_ps)})高于行业Q3({_fmt_val(industry_stats_val["ps"]["q3"])})，市销率偏高')
        elif cur_ps < industry_stats_val['ps']['q1']:
            add_paragraph(document, f' PS({_fmt_val(cur_ps)})低于行业Q1({_fmt_val(industry_stats_val["ps"]["q1"])})，市销率偏低')

    # 2.2.2 与申万行业指数对比
    if sw_index_pe is not None and cur_pe is not None:
        add_paragraph(document, '')
        add_title(document, '2.2.2 与申万行业指数对比', level=3)
        add_paragraph(document, '')

        # 计算与申万行业指数的偏离度
        pe_dev_sw_val = (cur_pe - sw_index_pe) / sw_index_pe * 100 if sw_index_pe else None
        pe_dev_sw = _fmt_dev(cur_pe, sw_index_pe)
        add_paragraph(document, f"• PE偏离度: {pe_dev_sw}（标的{_fmt_val(cur_pe)} vs 申万{_fmt_val(sw_index_pe)}）")
        pb_dev_sw_val = None
        if cur_pb is not None and sw_index_pb:
            pb_dev_sw_val = (cur_pb - sw_index_pb) / sw_index_pb * 100 if sw_index_pb else None
            pb_dev_sw = _fmt_dev(cur_pb, sw_index_pb)
            add_paragraph(document, f"• PB偏离度: {pb_dev_sw}（标的{_fmt_val(cur_pb)} vs 申万{_fmt_val(sw_index_pb)}）")
        ps_dev_sw_val = None
        if cur_ps is not None and sw_index_ps:
            ps_dev_sw_val = (cur_ps - sw_index_ps) / sw_index_ps * 100 if sw_index_ps else None
            ps_dev_sw = _fmt_dev(cur_ps, sw_index_ps)
            add_paragraph(document, f"• PS偏离度: {ps_dev_sw}（标的{_fmt_val(cur_ps)} vs 申万{_fmt_val(sw_index_ps)}）")

        add_paragraph(document, '')

        # PE申万指数对比分析
        if pe_dev_sw_val is not None and abs(pe_dev_sw_val) < 10:
            add_paragraph(document, f' PE({current_metrics_val["pe"]:.2f}倍)与申万行业指数PE({sw_index_pe:.2f}倍)基本一致，偏离度{pe_dev_sw_val:+.1f}%')
        elif pe_dev_sw_val is not None and pe_dev_sw_val > 0:
            add_paragraph(document, f' PE({current_metrics_val["pe"]:.2f}倍)高于申万行业指数PE({sw_index_pe:.2f}倍)，溢价{pe_dev_sw_val:+.1f}%')
        elif pe_dev_sw_val is not None:
            add_paragraph(document, f' PE({current_metrics_val["pe"]:.2f}倍)低于申万行业指数PE({sw_index_pe:.2f}倍)，折价{pe_dev_sw_val:+.1f}%')

        # PB申万指数对比分析
        if pb_dev_sw_val is not None:
            if abs(pb_dev_sw_val) < 10:
                add_paragraph(document, f' PB({current_metrics_val["pb"]:.2f}倍)与申万行业指数PB({sw_index_pb:.2f}倍)基本一致')
            elif pb_dev_sw_val > 0:
                add_paragraph(document, f' PB({current_metrics_val["pb"]:.2f}倍)高于申万行业指数PB({sw_index_pb:.2f}倍)，溢价{pb_dev_sw_val:+.1f}%')
            else:
                add_paragraph(document, f' PB({current_metrics_val["pb"]:.2f}倍)低于申万行业指数PB({sw_index_pb:.2f}倍)，折价{pb_dev_sw_val:+.1f}%')

        add_paragraph(document, '')
        add_paragraph(document, '申万行业指数说明：')
        add_paragraph(document, f'• 申万行业指数基于所有成分股市值加权，反映行业整体估值水平')
        add_paragraph(document, f'• 与申万指数对比可判断个股相对行业整体的估值位置')
        add_paragraph(document, f'• 正偏离表示估值高于行业平均，负偏离表示估值低于行业平均')

    
    # ==================== 2.3 PE历史分位数趋势分析 ====================
    add_title(document, '2.3 PE历史分位数趋势分析', level=2)

    add_paragraph(document, '本节通过分析标的股票和所属行业的PE历史走势及分位数变化，从时间维度评估估值的合理性。')
    add_paragraph(document, '基于最近5年的历史数据（约1250个交易日），通过对比个股与行业的PE历史分位数趋势，可以更清晰地判断当前估值处于历史哪个水平。')
    add_paragraph(document, '5年的历史周期能够覆盖完整的牛熊周期，提供更可靠的估值基准。')
    add_paragraph(document, '')

    # 尝试从tushare获取历史PE数据并生成趋势图
    try:
        from utils.pe_history_analyzer import PEHistoryAnalyzer

        print("\n=== 开始PE历史分位数趋势分析 ===")

        # 创建PE历史分析器
        pe_analyzer = PEHistoryAnalyzer()

        # 获取个股历史PE数据（最近5年）
        print(" 正在获取个股历史PE数据...")
        stock_pe_data = pe_analyzer.get_stock_pe_history(stock_code, days=1825)

        if stock_pe_data is not None:
            print(f" 个股PE数据获取成功: {len(stock_pe_data)}条记录")
        else:
            print(f" 个股PE数据获取失败")

        # 获取行业历史PE数据
        print(" 正在获取行业历史PE数据...")
        industry_name, industry_code, industry_pe_data = pe_analyzer.get_industry_pe_history(stock_code, days=1825)

        if industry_pe_data is not None:
            print(f" 行业PE数据获取成功: {len(industry_pe_data)}条记录, 行业: {industry_name}({industry_code})")
        else:
            print(f" 行业PE数据获取失败, 行业: {industry_name}({industry_code})")

        # 保存可用的PE数据到context（即使部分数据失败也保存可用的数据）
        pe_data_saved = False

        if stock_pe_data is not None:
            context['stock_pe_data'] = stock_pe_data
            print(f" 已保存个股PE数据到context，供第六章情景分析使用")
            pe_data_saved = True

        if industry_pe_data is not None:
            context['industry_pe_data'] = industry_pe_data
            context['industry_name'] = industry_name
            context['industry_code'] = industry_code
            print(f" 已保存行业PE数据到context，供第六章情景分析使用")
            pe_data_saved = True

        if stock_pe_data is not None and industry_pe_data is not None:
            # 完整的PE分析（需要个股和行业数据都可用）
            print(f" 成功获取完整历史PE数据，进行详细分析")
            print(f"   个股数据: {len(stock_pe_data)}条")
            print(f"   行业数据: {len(industry_pe_data)}条")
            print(f"   行业: {industry_name} ({industry_code})")

            # 计算个股历史分位数统计
            stock_pe_current = stock_pe_data.iloc[-1]['pe_ttm']
            stock_pe_min = stock_pe_data['pe_ttm'].min()
            stock_pe_max = stock_pe_data['pe_ttm'].max()
            stock_pe_median = stock_pe_data['pe_ttm'].median()
            stock_pe_25 = stock_pe_data['pe_ttm'].quantile(0.25)
            stock_pe_75 = stock_pe_data['pe_ttm'].quantile(0.75)
            stock_pe_percentile = (stock_pe_data['pe_ttm'] < stock_pe_current).sum() / len(stock_pe_data) * 100

            # 计算申万行业指数历史分位数统计
            sw_index_pe_current = industry_pe_data.iloc[-1]['pe_ttm']
            sw_index_pe_min = industry_pe_data['pe_ttm'].min()
            sw_index_pe_max = industry_pe_data['pe_ttm'].max()
            sw_index_pe_median = industry_pe_data['pe_ttm'].median()
            sw_index_pe_25 = industry_pe_data['pe_ttm'].quantile(0.25)
            sw_index_pe_75 = industry_pe_data['pe_ttm'].quantile(0.75)
            sw_index_pe_percentile = (industry_pe_data['pe_ttm'] < sw_index_pe_current).sum() / len(industry_pe_data) * 100

            # 获取同行公司的历史PE数据
            print("\n获取同行公司历史PE数据...")
            custom_peer_pe_current = None
            custom_peer_pe_min = None
            custom_peer_pe_max = None
            custom_peer_pe_median = None
            custom_peer_pe_percentile = None

            try:
                # 提取同行公司代码列表
                peer_codes = peer_companies_val['code'].tolist()[:20]  # 取前20个
                print(f"  同行公司数量: {len(peer_codes)}")

                # 获取同行公司历史PE（并行查询，避免逐家串行20次API调用）
                from concurrent.futures import ThreadPoolExecutor

                def _fetch_one_peer(code):
                    try:
                        h = pe_analyzer.get_stock_pe_history(code, days=1825)
                        if h is not None and not h.empty:
                            return h.rename(columns={'pe_ttm': 'pe'})
                    except Exception:
                        pass
                    return None

                peer_pe_histories = []
                with ThreadPoolExecutor(max_workers=5) as executor:
                    for h in executor.map(_fetch_one_peer, peer_codes):
                        if h is not None:
                            peer_pe_histories.append(h)
                print(f"  并行获取 {len(peer_pe_histories)}/{len(peer_codes)} 家同行历史PE完成")

                if peer_pe_histories:
                    # 合并所有同行公司的历史PE
                    peer_pe_df = pd.concat(peer_pe_histories, ignore_index=True)

                    # 按日期分组计算平均值
                    custom_peer_pe_data = peer_pe_df.groupby('trade_date')['pe'].mean().reset_index()
                    custom_peer_pe_data = custom_peer_pe_data.sort_values('trade_date').reset_index(drop=True)

                    # 计算统计指标
                    # 安全提取标量值，处理各种数据类型
                    try:
                        # 获取最后一个PE值
                        pe_last = custom_peer_pe_data.iloc[-1]['pe']
                        if isinstance(pe_last, (pd.Series, list)):
                            custom_peer_pe_current = float(pe_last.iloc[0] if isinstance(pe_last, pd.Series) else pe_last[0])
                        else:
                            custom_peer_pe_current = float(pe_last)

                        # 其他聚合函数，添加类型检查
                        pe_min = custom_peer_pe_data['pe'].min()
                        custom_peer_pe_min = float(pe_min.iloc[0] if isinstance(pe_min, pd.Series) else pe_min)

                        pe_max = custom_peer_pe_data['pe'].max()
                        custom_peer_pe_max = float(pe_max.iloc[0] if isinstance(pe_max, pd.Series) else pe_max)

                        pe_median = custom_peer_pe_data['pe'].median()
                        custom_peer_pe_median = float(pe_median.iloc[0] if isinstance(pe_median, pd.Series) else pe_median)

                        # 计算25%和75%分位数
                        pe_25 = custom_peer_pe_data['pe'].quantile(0.25)
                        custom_peer_pe_25 = float(pe_25.iloc[0] if isinstance(pe_25, pd.Series) else pe_25)

                        pe_75 = custom_peer_pe_data['pe'].quantile(0.75)
                        custom_peer_pe_75 = float(pe_75.iloc[0] if isinstance(pe_75, pd.Series) else pe_75)

                        # 计算百分位
                        pe_percentile_count = (custom_peer_pe_data['pe'] < custom_peer_pe_current).sum()
                        custom_peer_pe_percentile = float(pe_percentile_count.iloc[0] if isinstance(pe_percentile_count, pd.Series) else pe_percentile_count) / len(custom_peer_pe_data) * 100

                    except Exception as e:
                        print(f"   计算统计指标时出错: {e}")
                        print(f"   custom_peer_pe_data形状: {custom_peer_pe_data.shape}")
                        print(f"   custom_peer_pe_data列: {custom_peer_pe_data.columns.tolist()}")
                        raise

                    print(f"   同行公司历史PE计算成功:")
                    print(f"     当前PE: {custom_peer_pe_current:.2f}倍")
                    print(f"     数据点数: {len(custom_peer_pe_data)}")
                else:
                    print(f"   未获取到同行公司历史PE数据")

            except Exception as e:
                print(f"   获取同行公司历史PE失败: {e}")

            # 添加历史分位数统计表格
            if custom_peer_pe_current is not None:
                # 有自定义同行数据的完整表格
                pe_history_headers = ['指标', '标的股票', '行业指数(申万)', '行业指数(自定义)', '差异(vs申万)', '差异(vs自定义)']
                pe_history_data = [
                    ['当前PE-TTM',
                     f'{stock_pe_current:.2f}倍',
                     f'{sw_index_pe_current:.2f}倍',
                     f'{custom_peer_pe_current:.2f}倍',
                     f'{(stock_pe_current/sw_index_pe_current-1)*100:+.1f}%',
                     f'{(stock_pe_current/custom_peer_pe_current-1)*100:+.1f}%'],
                    ['历史最小PE',
                     f'{stock_pe_min:.2f}倍',
                     f'{sw_index_pe_min:.2f}倍',
                     f'{custom_peer_pe_min:.2f}倍',
                     f'{stock_pe_min-sw_index_pe_min:+.2f}倍',
                     f'{stock_pe_min-custom_peer_pe_min:+.2f}倍'],
                    ['25%分位数PE',
                     f'{stock_pe_25:.2f}倍',
                     f'{sw_index_pe_25:.2f}倍',
                     f'{custom_peer_pe_25:.2f}倍',
                     f'{stock_pe_25-sw_index_pe_25:+.2f}倍',
                     f'{stock_pe_25-custom_peer_pe_25:+.2f}倍'],
                    ['历史中位数PE',
                     f'{stock_pe_median:.2f}倍',
                     f'{sw_index_pe_median:.2f}倍',
                     f'{custom_peer_pe_median:.2f}倍',
                     f'{stock_pe_median-sw_index_pe_median:+.2f}倍',
                     f'{stock_pe_median-custom_peer_pe_median:+.2f}倍'],
                    ['75%分位数PE',
                     f'{stock_pe_75:.2f}倍',
                     f'{sw_index_pe_75:.2f}倍',
                     f'{custom_peer_pe_75:.2f}倍',
                     f'{stock_pe_75-sw_index_pe_75:+.2f}倍',
                     f'{stock_pe_75-custom_peer_pe_75:+.2f}倍'],
                    ['历史最大PE',
                     f'{stock_pe_max:.2f}倍',
                     f'{sw_index_pe_max:.2f}倍',
                     f'{custom_peer_pe_max:.2f}倍',
                     f'{stock_pe_max-sw_index_pe_max:+.2f}倍',
                     f'{stock_pe_max-custom_peer_pe_max:+.2f}倍'],
                    ['当前分位数',
                     f'{stock_pe_percentile:.1f}%',
                     f'{sw_index_pe_percentile:.1f}%',
                     f'{custom_peer_pe_percentile:.1f}%',
                     f'{stock_pe_percentile-sw_index_pe_percentile:+.1f}%',
                     f'{stock_pe_percentile-custom_peer_pe_percentile:+.1f}%']
                ]
            else:
                # 只有申万数据的简化表格
                pe_history_headers = ['指标', '标的股票', '行业指数(申万)', '差异']
                pe_history_data = [
                    ['当前PE-TTM', f'{stock_pe_current:.2f}倍', f'{sw_index_pe_current:.2f}倍', f'{(stock_pe_current/sw_index_pe_current-1)*100:+.1f}%'],
                    ['历史最小PE', f'{stock_pe_min:.2f}倍', f'{sw_index_pe_min:.2f}倍', f'{stock_pe_min-sw_index_pe_min:+.2f}倍'],
                    ['25%分位数PE', f'{stock_pe_25:.2f}倍', f'{sw_index_pe_25:.2f}倍', f'{stock_pe_25-sw_index_pe_25:+.2f}倍'],
                    ['历史中位数PE', f'{stock_pe_median:.2f}倍', f'{sw_index_pe_median:.2f}倍', f'{stock_pe_median-sw_index_pe_median:+.2f}倍'],
                    ['75%分位数PE', f'{stock_pe_75:.2f}倍', f'{sw_index_pe_75:.2f}倍', f'{stock_pe_75-sw_index_pe_75:+.2f}倍'],
                    ['历史最大PE', f'{stock_pe_max:.2f}倍', f'{sw_index_pe_max:.2f}倍', f'{stock_pe_max-sw_index_pe_max:+.2f}倍'],
                    ['当前分位数', f'{stock_pe_percentile:.1f}%', f'{sw_index_pe_percentile:.1f}%', f'{stock_pe_percentile-sw_index_pe_percentile:+.1f}%']
                ]

            add_table_data(document, pe_history_headers, pe_history_data)

            add_paragraph(document, '')
            add_paragraph(document, '历史分位数说明：')
            add_paragraph(document, f'• 当前分位数表示当前PE在历史数据中的相对位置')
            add_paragraph(document, f'• 例如：{stock_pe_percentile:.1f}%分位数表示历史上只有{stock_pe_percentile:.1f}%的时间PE低于当前值')
            add_paragraph(document, f'• 50%分位数即为中位数，代表历史平均水平')
            if custom_peer_pe_current is not None:
                add_paragraph(document, f'• 行业指数（申万）：申万行业指数的PE，基于所有成分股按市值加权计算')
                add_paragraph(document, f'• 行业指数（自定义）：2.1.1节中同行公司PE的简单平均，反映所选同行公司的平均估值水平')
            add_paragraph(document, '')

            # 生成PE趋势图
            pe_trend_chart_path = os.path.join(IMAGES_DIR, '02_4_pe_trend_analysis.png')
            chart_path = pe_analyzer.generate_pe_trend_chart(
                stock_code, stock_pe_data,
                industry_name, industry_pe_data,
                pe_trend_chart_path
            )

            # 添加图表到文档
            if chart_path and os.path.exists(chart_path):
                add_paragraph(document, '图表 2.4: PE历史分位数趋势分析')
                add_image(document, chart_path, width=Inches(6.5))

                add_paragraph(document, '')
                add_paragraph(document, '图表解读：', bold=True)
                add_paragraph(document, f'左上-PE走势对比：')
                add_paragraph(document, f'  • 蓝线：{stock_code}的PE-TTM走势')
                add_paragraph(document, f'  • 红线：{industry_name}的PE-TTM走势')
                add_paragraph(document, f'  • 两条线的相对位置反映个股相对行业的估值水平')
                add_paragraph(document, '')

                add_paragraph(document, f'右上-PE相对位置（个股/行业）：')
                add_paragraph(document, f'  • 比值>1：个股PE高于行业，溢价')
                add_paragraph(document, f'  • 比值<1：个股PE低于行业，折价')
                add_paragraph(document, f'  • 比值=1：与行业持平')
                add_paragraph(document, '')

                add_paragraph(document, f'左下-个股PE历史分位数：')
                add_paragraph(document, f'  • 显示{stock_code}的PE在历史中的位置变化')
                add_paragraph(document, f'  • 当前分位数：{stock_pe_percentile:.1f}%')
                add_paragraph(document, f'  • 分位数上升表示估值相对历史提升')
                add_paragraph(document, '')

                add_paragraph(document, f'右下-行业PE历史分位数：')
                add_paragraph(document, f'  • 显示{industry_name}的PE在历史中的位置变化')
                add_paragraph(document, f'  • 当前分位数：{sw_index_pe_percentile:.1f}%')
                add_paragraph(document, f'  • 可用于判断行业整体估值水平')
                add_paragraph(document, '')

            # 添加分析结论
            add_paragraph(document, '')
            add_paragraph(document, 'PE历史分位数趋势分析结论：', bold=True)

            # 估值水平判断
            if stock_pe_percentile >= 80:
                stock_valuation_level = "历史高位"
                stock_emoji = ""
                stock_comment = f"当前PE处于历史{stock_pe_percentile:.1f}%分位数，属于历史高位，估值偏高，需警惕回调风险"
            elif stock_pe_percentile >= 60:
                stock_valuation_level = "历史中高位"
                stock_emoji = "🟠"
                stock_comment = f"当前PE处于历史{stock_pe_percentile:.1f}%分位数，属于历史中高位，估值相对偏高"
            elif stock_pe_percentile >= 40:
                stock_valuation_level = "历史中位数"
                stock_emoji = "🟡"
                stock_comment = f"当前PE处于历史{stock_pe_percentile:.1f}%分位数，接近历史中位数，估值合理"
            elif stock_pe_percentile >= 20:
                stock_valuation_level = "历史中低位"
                stock_emoji = "🟢"
                stock_comment = f"当前PE处于历史{stock_pe_percentile:.1f}%分位数，属于历史中低位，估值相对偏低"
            else:
                stock_valuation_level = "历史低位"
                stock_emoji = ""
                stock_comment = f"当前PE处于历史{stock_pe_percentile:.1f}%分位数，属于历史低位，估值偏低，安全边际较高"

            add_paragraph(document, f'{stock_emoji} 个股估值水平：{stock_valuation_level}')
            add_paragraph(document, f'   {stock_comment}')
            add_paragraph(document, '')

            # 与行业对比
            if stock_pe_percentile > sw_index_pe_percentile + 20:
                relative_comment = f"个股分位数({stock_pe_percentile:.1f}%)显著高于行业({sw_index_pe_percentile:.1f}%)，相对行业估值偏高"
                relative_emoji = ""
            elif stock_pe_percentile < sw_index_pe_percentile - 20:
                relative_comment = f"个股分位数({stock_pe_percentile:.1f}%)显著低于行业({sw_index_pe_percentile:.1f}%)，相对行业估值偏低，安全边际较高"
                relative_emoji = ""
            else:
                relative_comment = f"个股分位数({stock_pe_percentile:.1f}%)与行业({sw_index_pe_percentile:.1f}%)基本持平"
                relative_emoji = "ℹ️"

            add_paragraph(document, f'{relative_emoji} 相对行业估值：{relative_comment}')
            add_paragraph(document, '')

            # 投资建议
            add_paragraph(document, '历史分位数投资启示：')
            if stock_pe_percentile <= 25:
                add_paragraph(document, f'• 当前PE处于历史{stock_pe_percentile:.1f}%分位数（低位），历史上仅{stock_pe_percentile:.1f}%的时间估值更低')
                add_paragraph(document, f'• 从历史角度看，当前估值具备较好的安全边际')
                add_paragraph(document, f'• 建议积极关注，估值修复空间较大')
            elif stock_pe_percentile <= 50:
                add_paragraph(document, f'• 当前PE处于历史{stock_pe_percentile:.1f}%分位数（中低位），估值相对合理或偏低')
                add_paragraph(document, f'• 从历史角度看，当前估值风险可控')
                add_paragraph(document, f'• 建议适度配置，关注基本面变化')
            elif stock_pe_percentile <= 75:
                add_paragraph(document, f'• 当前PE处于历史{stock_pe_percentile:.1f}%分位数（中高位），估值相对偏高')
                add_paragraph(document, f'• 从历史角度看，当前估值风险上升')
                add_paragraph(document, f'• 建议谨慎参与，等待更好的买入时机')
            else:
                add_paragraph(document, f'• 当前PE处于历史{stock_pe_percentile:.1f}%分位数（高位），估值处于历史高位')
                add_paragraph(document, f'• 从历史角度看，当前估值风险较大')
                add_paragraph(document, f'• 建议等待估值回落至历史中低位再考虑参与')

        else:
            print(" PE历史数据获取不完整，跳过趋势图生成")
            add_paragraph(document, ' PE历史分位数趋势图暂时无法生成，可能原因：')
            add_paragraph(document, '   • tushare数据缺失或API调用限制')
            add_paragraph(document, '   • 股票或行业历史数据不足')

    except ImportError as e:
        print(f" PE历史分析器导入失败: {e}")
        add_paragraph(document, ' PE历史分位数趋势分析功能暂不可用')

    except Exception as e:
        print(f" PE历史分位数趋势分析失败: {e}")
        add_paragraph(document, f' PE历史分位数趋势分析执行失败: {e}')

    add_paragraph(document, '')

    # ==================== 2.4 PE估值分析（绝对估值与修正估值） ====================
    add_title(document, '2.4 PE估值分析', level=2)
    add_paragraph(document, '基于行业净利率、净利润增长率等参数，对PE进行绝对估值和修正估值分析。')

    try:
        import sys
        _PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        _ROOT_DIR = os.path.dirname(_PROJECT_DIR)
        if _ROOT_DIR not in sys.path:
            sys.path.insert(0, _ROOT_DIR)

        from industry_dcf.utils.pe_estimator import PEEstimator

        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if ts_token and stock_code:
            import tushare as ts
            from industry_dcf.utils.industry_data_fetcher import IndustryDataFetcher
            from industry_dcf.utils.rate_limiter import RateLimiter
            from industry_dcf.utils.shenwan_lookup import find_l3_by_code
            from industry_dcf.utils.industry_dcf_calculator import IndustryDCFCalculator

            pro = ts.pro_api(ts_token)
            rl = RateLimiter()
            fetcher = IndustryDataFetcher(pro, rate_limiter=rl)
            calculator = IndustryDCFCalculator()

            # 行业识别
            ind_info = find_l3_by_code(stock_code, pro)
            if ind_info:
                l3_code = ind_info['l3_code']
                ind_name = ind_info.get('l3_name', '')

                # 获取行业数据（优先从context复用，避免Ch2/Ch3重复获取）
                print("  [2.4] 获取行业数据用于PE估值...")
                ind_financials = context.get('_industry_financials')
                ind_pe_data = context.get('_industry_pe_data')
                if not ind_financials or not ind_pe_data:
                    ind_financials = fetcher.get_industry_financials(l3_code)
                    ind_pe_data = fetcher.get_industry_daily_basics(l3_code)
                    # 存入context供Ch3复用
                    context['_industry_financials'] = ind_financials
                    context['_industry_pe_data'] = ind_pe_data
                    context['_industry_l3_code'] = l3_code
                benchmark = calculator.calculate_industry_benchmark(ind_financials, industry_pe_data=ind_pe_data)

                # 获取个股数据
                company_data = fetcher.get_company_financials(stock_code)

                # 市值数据
                current_price = project_params.get('current_price', 0)
                total_shares = project_params.get('total_shares', 0)
                market_cap = current_price * total_shares if current_price and total_shares else 0

                # 尝试从daily_basic获取更准确的市值
                if ts_token:
                    try:
                        for db_back in range(0, 10):
                            td = (datetime.now() - timedelta(days=db_back)).strftime('%Y%m%d')
                            db_df = pro.daily_basic(ts_code=stock_code, trade_date=td,
                                                    fields='ts_code,total_mv,total_share,close')
                            if db_df is not None and not db_df.empty:
                                row = db_df.iloc[0]
                                market_cap = float(row.get('total_mv', 0) or 0) * 10000
                                total_shares = float(row.get('total_share', 0) or 0)
                                current_price = float(row.get('close', 0) or 0)
                                break
                    except Exception:
                        pass

                market_data = {
                    'market_cap': market_cap,
                    'current_price': current_price,
                    'total_shares': total_shares,
                }

                # 执行PE估值
                estimator = PEEstimator()
                pe_result = estimator.estimate_normalized_pe(
                    ts_code=stock_code,
                    industry_financials=ind_financials,
                    industry_benchmark=benchmark,
                    company_financials=company_data,
                    market_data=market_data,
                )

                # ====== 2.4.1 绝对估值 ======
                add_title(document, '2.4.1 绝对PE估值', level=3)
                add_paragraph(document, '基于当前市值和实际归母净利润计算的PE值。')
                add_paragraph(document, '')

                actual_pe = pe_result.get('actual_pe')
                ind_pe_med = pe_result.get('industry_pe_median', 0)
                company_rev = pe_result.get('company_revenue_latest', 0)
                actual_ni = pe_result.get('actual_net_income', 0)
                co_margin = pe_result.get('company_net_margin', {})

                pe_headers = ['指标', '数值']
                pe_data_rows = [
                    ['当前股价', f'{current_price:.2f} 元'],
                    ['总市值', f'{market_cap/100000000:.2f} 亿元' if market_cap else 'N/A'],
                    ['最新营收', f'{company_rev/10000:.2f} 亿元' if company_rev else 'N/A'],
                    ['归母净利润', f'{actual_ni/10000:.2f} 亿元' if actual_ni else 'N/A'],
                    ['公司净利率(最新)', f'{co_margin.get("latest", 0)*100:.2f}%'],
                    ['公司净利率(3年均)', f'{co_margin.get("avg_3y", 0)*100:.2f}%'],
                    ['实际PE(静态)', f'{actual_pe:.1f} 倍' if actual_pe else 'N/A(亏损)'],
                    ['行业PE中位数', f'{ind_pe_med:.1f} 倍'],
                ]
                add_table_data(document, pe_headers, pe_data_rows)
                add_paragraph(document, '')

                if actual_pe and ind_pe_med > 0:
                    if actual_pe > ind_pe_med * 1.5:
                        pe_comment = f'实际PE({actual_pe:.1f}倍)显著高于行业PE中位数({ind_pe_med:.1f}倍)，估值偏高'
                    elif actual_pe > ind_pe_med:
                        pe_comment = f'实际PE({actual_pe:.1f}倍)略高于行业PE中位数({ind_pe_med:.1f}倍)，估值略高'
                    elif actual_pe > ind_pe_med * 0.8:
                        pe_comment = f'实际PE({actual_pe:.1f}倍)接近行业PE中位数({ind_pe_med:.1f}倍)，估值合理'
                    else:
                        pe_comment = f'实际PE({actual_pe:.1f}倍)低于行业PE中位数({ind_pe_med:.1f}倍)，估值偏低'
                    add_paragraph(document, f'• {pe_comment}')

                # ====== 2.4.2 修正估值 ======
                add_title(document, '2.4.2 修正PE估值（行业修正）', level=3)
                add_paragraph(document, '当企业净利率低于行业水平时，使用行业净利率修正净利润，推算"正常"PE和对应股价。')
                add_paragraph(document, '')

                ind_margin = pe_result.get('industry_net_margin', {})
                ind_eg = pe_result.get('industry_earnings_growth', {})
                norm_ni = pe_result.get('normalized_net_income', 0)
                norm_pe = pe_result.get('normalized_pe')
                norm_src = pe_result.get('normalized_margin_source', 'industry')
                norm_margin = pe_result.get('normalized_margin_used', 0)
                margin_gap = pe_result.get('margin_gap', 0)

                mod_headers = ['指标', '数值']
                mod_rows = [
                    ['行业净利率中位数', f'{ind_margin.get("median", 0)*100:.2f}%'],
                    ['行业净利润增速中位数', f'{ind_eg.get("median", 0)*100:.2f}%'],
                    ['净利率差距(公司-行业)', f'{margin_gap*100:+.2f}%'],
                    ['采用的净利率', f'{norm_margin*100:.2f}% ({("公司实际" if norm_src == "company" else "行业中位数")})'],
                    ['修正后净利润', f'{norm_ni/10000:.2f} 亿元' if norm_ni else 'N/A'],
                    ['修正后PE', f'{norm_pe:.1f} 倍' if norm_pe else 'N/A'],
                ]
                add_table_data(document, mod_headers, mod_rows)
                add_paragraph(document, '')

                if norm_src == 'industry':
                    add_paragraph(document, f'• 公司净利率低于行业水平，使用行业净利率({ind_margin.get("median",0)*100:.1f}%)修正盈利')
                else:
                    add_paragraph(document, f'• 公司净利率高于或等于行业水平，使用公司实际净利率({co_margin.get("latest",0)*100:.1f}%)')

                # ====== 修正估值反推股价 ======
                add_title(document, '2.4.3 修正估值目标股价', level=3)
                add_paragraph(document, '基于修正后净利润和行业PE中位数，反推目标股价及增长空间。')
                add_paragraph(document, '')

                total_shares_wan = total_shares  # Tushare total_share 单位为万股
                if norm_ni > 0 and ind_pe_med > 0 and total_shares_wan > 0 and current_price > 0:
                    target_price = norm_ni * ind_pe_med / total_shares_wan
                    upside = (target_price / current_price - 1) * 100

                    target_headers = ['指标', '数值']
                    target_rows = [
                        ['当前股价', f'{current_price:.2f} 元'],
                        ['修正后净利润', f'{norm_ni/10000:.2f} 亿元'],
                        ['行业PE中位数', f'{ind_pe_med:.1f} 倍'],
                        ['目标股价', f'{target_price:.2f} 元'],
                        ['增长空间', f'{upside:+.1f}%'],
                    ]
                    add_table_data(document, target_headers, target_rows)
                    add_paragraph(document, '')

                    if upside > 20:
                        target_comment = f'修正估值目标价({target_price:.2f}元)较当前价({current_price:.2f}元)有{upside:.1f}%的增长空间，股价具有上行潜力'
                    elif upside > 0:
                        target_comment = f'修正估值目标价({target_price:.2f}元)较当前价({current_price:.2f}元)有{upside:.1f}%的增长空间'
                    elif upside > -20:
                        target_comment = f'修正估值目标价({target_price:.2f}元)略低于当前价({current_price:.2f}元)，估值基本合理'
                    else:
                        target_comment = f'修正估值目标价({target_price:.2f}元)较当前价({current_price:.2f}元)低{abs(upside):.1f}%，存在高估风险'
                    add_paragraph(document, f'• {target_comment}')
                    add_paragraph(document, '')
                    add_paragraph(document, f'计算公式: 目标股价 = 修正后净利润 × 行业PE中位数 ÷ 总股本')
                    add_paragraph(document, f'         = {norm_ni/10000:.2f}亿 × {ind_pe_med:.1f} ÷ {total_shares_wan:.0f}万股')
                    add_paragraph(document, f'         = {target_price:.2f} 元')
                else:
                    add_paragraph(document, '• 数据不足，无法计算修正估值目标股价')

                # ====== 前瞻PE预测 ======
                projections = pe_result.get('forward_pe_projections', [])
                if projections and current_price > 0:
                    add_title(document, '2.4.4 前瞻PE预测', level=3)
                    eg = pe_result.get('earnings_growth_used', 0)
                    add_paragraph(document, f'基于行业净利润增速({eg*100:.1f}%)，预测未来PE变化及对应目标股价：')
                    add_paragraph(document, '')

                    fwd_headers = ['年份', '预测净利润(亿元)', '前瞻PE(倍)', '行业PE对应价(元)', '增长空间']
                    fwd_rows = []
                    for p in projections:
                        yr = p['year']
                        if yr in (1, 2, 3, 5, 7, 10):
                            fwd_rows.append([
                                f'第{yr}年',
                                f'{p["projected_ni"]/10000:.2f}',
                                f'{p["forward_pe"]:.1f}',
                                f'{p["projected_price"]:.2f}',
                                f'{p["upside_pct"]:+.1f}%',
                            ])
                    add_table_data(document, fwd_headers, fwd_rows)
                    add_paragraph(document, '')
                    add_paragraph(document, '注: 前瞻PE = 当前市值 ÷ 预测净利润; 行业PE对应价 = 预测净利润 × 行业PE中位数 ÷ 总股本')

                # 保存到context
                context['pe_estimation'] = pe_result

                print(f"  [2.4] PE估值完成: 实际PE={actual_pe}, 修正PE={norm_pe}")
            else:
                add_paragraph(document, '• 无法识别行业分类，PE估值分析暂不可用')
        else:
            add_paragraph(document, '• 未配置Tushare Token，PE估值分析暂不可用')

    except ImportError as e:
        print(f"  [2.4] PE估值模块导入失败: {e}")
        add_paragraph(document, f'• PE估值分析功能暂不可用 (导入失败: {e})')
    except Exception as e:
        print(f"  [2.4] PE估值分析失败: {e}")
        import traceback
        traceback.print_exc()
        add_paragraph(document, f'• PE估值分析执行失败: {e}')

    add_section_break(document)

    # 保存数据到context，供后续章节使用（第七章等需要）
    context['current_metrics_val'] = current_metrics_val
    context['industry_stats_val'] = industry_stats_val
    context['industry_avg_val'] = industry_avg_val
    context['peer_companies_val'] = peer_companies_val

    # 返回更新后的context
    return context
