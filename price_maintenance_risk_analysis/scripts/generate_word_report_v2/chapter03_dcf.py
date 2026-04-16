#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
第三章：DCF估值分析

本章节实现DCF（现金流折现）估值分析，包括：
- DCF估值方法说明
- 估值敏感性分析
- 估值参数与过程
- 历史FCF数据分析
- 逐年FCF预测与折现计算
- 终值计算
- 企业价值与股权价值计算
- 估值结论
"""

import os
import sys
from datetime import datetime, timedelta
from docx.shared import Inches

import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # 使用非交互式后端

import numpy as np
import pandas as pd
import seaborn as sns

# 添加父目录到路径以导入工具模块
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(os.path.dirname(SCRIPT_DIR))
sys.path.insert(0, PROJECT_DIR)

from utils.font_manager import get_font_prop
from utils.wacc_calculator import WACCCalculator, calculate_wacc_simple
from module_utils import (
    add_title,
    add_paragraph,
    add_table_data,
    add_image,
    add_section_break
)

# 获取中文字体
font_prop = get_font_prop()

# 配置路径
DATA_DIR = os.path.join(PROJECT_DIR, 'data')
REPORTS_DIR = os.path.join(PROJECT_DIR, 'reports')
IMAGES_DIR = os.path.join(REPORTS_DIR, 'images')


def generate_chapter(context):
    """
    生成第三章：DCF估值分析

    参数:
        context (dict): 包含以下键的上下文字典:
            - document: Word文档对象
            - stock_code: 股票代码
            - project_params: 项目参数字典
            - placement_params: 定增参数字典

    返回:
        dict: 更新后的上下文字典
    """
    document = context['document']
    stock_code = context['stock_code']
    project_params = context['project_params']
    placement_params = context.get('placement_params', {})

    # ==================== 三、DCF估值分析 ====================
    add_title(document, '三、DCF估值分析', level=1)

    add_paragraph(document, '本章节使用现金流折现模型评估公司的内在价值。')

    add_title(document, '3.1 DCF估值方法说明', level=2)

    add_paragraph(document, 'DCF（Discounted Cash Flow）估值法通过预测公司未来自由现金流，并以加权平均资本成本（WACC）折现至现值，得出公司的内在价值。')

    add_paragraph(document, '')
    add_paragraph(document, '估值步骤：')
    add_paragraph(document, '1. 预测未来10年自由现金流（FCF）')
    add_paragraph(document, '2. 计算终值（Terminal Value）')
    add_paragraph(document, '3. 以WACC折现至现值')
    add_paragraph(document, '4. 减去净债务，得到股权价值')
    add_paragraph(document, '5. 除以总股数，得到每股价值')

    # 从placement_params中获取财务数据
    net_income = project_params.get('net_income', 253532329.85)  # 净利润

    # 获取真实总股本（从 Tushare）
    total_shares = 767460689  # 光弘科技真实总股本（约7.67亿股）
    try:
        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if ts_token:
            import tushare as ts

            pro = ts.pro_api(ts_token)
            trade_date = (datetime.now() - timedelta(days=1)).strftime('%Y%m%d')

            df_basic = pro.daily_basic(
                ts_code=stock_code,
                trade_date=trade_date,
                fields='ts_code,total_share'
            )

            if not df_basic.empty:
                total_shares = float(df_basic.iloc[0]['total_share']) * 10000  # 万股转股
    except Exception as e:
        print(f"获取总股本失败，使用估算值: {e}")

    net_assets = project_params.get('net_assets', 4939639031.34)  # 净资产（元）
    total_debt = project_params.get('total_debt', 4507009002.2)  # 总债务（元）
    net_income = project_params.get('net_income', 253532329.85)  # 净利润（元）

    # 尝试从Tushare获取资产负债表和利润表数据
    net_debt = None
    cash_assets = 0.0
    total_liab_value = 0.0

    # 初始化CAGR计算变量
    net_income_cagr = None  # 净利润CAGR（备选）
    fcf_cagr = None  # FCF的CAGR（优先使用）
    historical_incomes = None

    try:
        from update_market_data import TushareFinancialData

        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if ts_token:
            financial = TushareFinancialData(stock_code)

            # 获取资产负债表数据
            balancesheet = financial.get_latest_balancesheet()
            if balancesheet:
                total_liab_value = balancesheet['total_liabilities'] / 100000000  # 转亿元
                cash_equivalents = balancesheet['cash_equivalents'] / 100000000  # 转亿元

                # 更新财务数据
                net_assets = balancesheet['total_equity']
                total_debt = balancesheet['total_liabilities']

                # 净债务 = 总负债 - 货币资金
                net_debt = total_liab_value - cash_equivalents
                cash_assets = cash_equivalents

                print(f" 从Tushare获取资产负债表数据成功，期间: {balancesheet['report_date']}")
                print(f"   总负债: {total_liab_value:.2f} 亿元")
                print(f"   货币资金: {cash_equivalents:.2f} 亿元")
                if net_debt > 0:
                    print(f"   净债务: {net_debt:.2f} 亿元")
                else:
                    print(f"   净现金: {-net_debt:.2f} 亿元")

            # 获取利润表数据（净利润）
            income = financial.get_latest_income()
            if income:
                net_income = income['net_income']
                print(f" 从Tushare获取利润表数据成功，期间: {income['report_date']}")
                print(f"   净利润: {net_income/100000000:.2f} 亿元")

                # 更新项目参数
                project_params['net_income'] = net_income

            # 获取历史净利润数据以计算CAGR（作为备选方案）
            historical_incomes = financial.get_historical_net_income(years=5)
            net_income_cagr = None  # 净利润CAGR（备选）
            if historical_incomes and len(historical_incomes) >= 2:
                # 计算CAGR (复合年增长率)
                # CAGR = (终值/起始值)^(1/年数) - 1
                start_income = historical_incomes[-1]
                end_income = historical_incomes[0]
                n_years = len(historical_incomes) - 1

                if start_income > 0 and end_income > 0:
                    net_income_cagr = (end_income / start_income) ** (1 / n_years) - 1
                    print(f" 计算净利润历史CAGR（备选）:")
                    print(f"   起始净利润: {start_income/100000000:.2f} 亿元")
                    print(f"   最新净利润: {end_income/100000000:.2f} 亿元")
                    print(f"   历史年数: {n_years} 年")
                    print(f"   净利润CAGR: {net_income_cagr*100:+.2f}%")
                    print(f"   注：将优先使用FCF的CAGR，如果不可用则使用净利润CAGR")

            # 获取完整历史FCF数据（用于DCF估值）
            historical_fcf_data = financial.get_historical_fcf_for_dcf(max_years=15)
            if historical_fcf_data:
                project_params['historical_fcf_data'] = historical_fcf_data
                print(f" 已保存历史FCF数据到项目参数")

    except Exception as e:
        print(f"从Tushare获取财务数据失败: {e}")

    # 如果无法获取实际数据，使用项目参数中的估计值
    if net_debt is None:
        # 计算净债务 = 总负债 - 货币资金
        cash_equivalents = project_params.get('cash_equivalents', 0)
        net_debt = (total_debt - cash_equivalents) / 100000000  # 转亿元
        cash_assets = cash_equivalents / 100000000
        print(f" 使用参数文件中的财务数据")
        print(f"   净债务: {net_debt:.2f} 亿元")

    # 生成并插入 DCF 热力图
    add_title(document, '3.2 估值敏感性分析', level=2)

    dcf_chart_path = os.path.join(IMAGES_DIR, '04_dcf_valuation_heatmap.png')
    generate_dcf_valuation_heatmap(dcf_chart_path, project_params['current_price'], net_income, total_shares, net_debt)

    add_paragraph(document, '图表 3: DCF估值敏感性分析矩阵')
    add_image(document, dcf_chart_path, width=Inches(6))

    add_paragraph(document, '分析说明：')
    add_paragraph(document, '• 横轴：WACC（加权平均资本成本），代表折现率')
    add_paragraph(document, '• 纵轴：增长率，代表预测期FCF增长率')
    add_paragraph(document, '• 颜色：绿色表示估值高于当前价，红色表示估值低于当前价')
    add_paragraph(document, f'• 当前股价: {project_params["current_price"]:.2f} 元/股')

    add_title(document, '3.3 估值参数与过程', level=2)

    # 计算基础FCF
    base_fcf = net_income  # FCF = 净利润（保守估计）
    using_net_income_as_fcf = True  # 标记是否使用净利润作为FCF基准

    # 获取历史FCF数据用于DCF估值
    historical_fcf_years = 0
    base_fcf = net_income  # 默认使用净利润作为FCF基准
    fcf_data_description = "使用净利润近似FCF（无历史FCF数据）"
    fcf_cagr = None  # FCF复合增长率（优先使用）

    if 'historical_fcf_data' in project_params and project_params['historical_fcf_data']:
        historical_fcf = project_params['historical_fcf_data']
        historical_fcf_years = historical_fcf['years']

        print(f"\n 历史FCF数据分析:")
        print(f"   总年数: {historical_fcf_years}")
        print(f"{'年份':<6} {'营收(亿)':<12} {'NOPAT(亿)':<12} {'FCF(亿)':<12}")
        print("-"*50)

        # 显示所有历史FCF数据
        for year_data in historical_fcf['data']:
            print(f"{year_data['year']:<6} {year_data['revenue']:>10.2f}     "
                  f"{year_data['nopat']:>10.2f}     {year_data['fcf']:>10.2f}")

        # 使用最新一年的真实FCF作为基准
        latest_fcf = historical_fcf['data'][-1]['fcf'] * 100000000  # 转回元

        # 检查FCF是否为正
        print(f"\n   最新FCF: {latest_fcf/100000000:.2f} 亿元")
        if latest_fcf > 0:
            base_fcf = latest_fcf
            using_net_income_as_fcf = False
            print(f"    使用最新FCF作为基准")
        else:
            base_fcf = net_income  # 使用净利润作为基准
            using_net_income_as_fcf = True
            print(f"    最新FCF为负，使用净利润作为基准")

        # 计算历史FCF复合增长率（CAGR）- 优先使用
        # 只计算正FCF年份
        positive_fcf_data = [d for d in historical_fcf['data'] if d['fcf'] > 0]
        print(f"   正值FCF年数: {len(positive_fcf_data)}")

        if len(positive_fcf_data) >= 2:
            # 使用最早和最新的正FCF年计算CAGR
            first_year_data = positive_fcf_data[0]
            last_year_data = positive_fcf_data[-1]

            first_fcf_val = first_year_data['fcf']
            last_fcf_val = last_year_data['fcf']
            years_count = last_year_data['year'] - first_year_data['year']

            # CAGR = (最终值 / 初始值)^(1/年数) - 1
            fcf_cagr = (last_fcf_val / first_fcf_val) ** (1 / years_count) - 1

            # 限制增长率在合理范围（-5%到30%）
            fcf_cagr = max(-0.05, min(0.30, fcf_cagr))

            print(f"    FCF历史CAGR: {fcf_cagr*100:.2f}%（从{first_year_data['year']}年{first_fcf_val:.2f}亿到{last_year_data['year']}年{last_fcf_val:.2f}亿，{years_count}年）")
            print(f"    将使用FCF的CAGR作为增长率基准")
        else:
            fcf_cagr = None
            print(f"    正值FCF年数少于2年，无法计算FCF的CAGR")
            if net_income_cagr is not None:
                print(f"   ℹ️  将使用净利润的CAGR作为备选: {net_income_cagr*100:.2f}%")

        fcf_data_description = f"使用真实历史FCF数据（{historical_fcf_years}年）"

    # 使用历史CAGR作为FCF增长率，优先级：FCF CAGR > 净利润CAGR > 默认值
    # 初始化FCF增长率变量（需要在表格前初始化）
    fcf_growth_example = 0.10  # 默认10%
    fcf_growth_source = "默认假设（无历史数据）"

    # 优先使用FCF的CAGR
    if fcf_cagr is not None:
        fcf_growth_example = max(fcf_cagr, 0.05)  # 至少5%，避免过于保守
        if fcf_growth_example > 0.30:  # 限制最高30%
            fcf_growth_example = 0.30
        fcf_growth_source = f"基于历史{historical_fcf_years}年FCF数据（CAGR: {fcf_cagr*100:.1f}%）"
        print(f"\n 最终采用FCF增长率: {fcf_growth_example*100:.1f}%（数据源: FCF历史CAGR）")
    # 如果FCF的CAGR不可用，使用净利润的CAGR作为备选
    elif net_income_cagr is not None:
        fcf_growth_example = max(net_income_cagr, 0.05)  # 至少5%
        if fcf_growth_example > 0.30:  # 限制最高30%
            fcf_growth_example = 0.30
        fcf_growth_source = f"基于历史净利润CAGR（CAGR: {net_income_cagr*100:.1f}%）"
        print(f"\n 最终采用净利润增长率: {fcf_growth_example*100:.1f}%（数据源: 净利润历史CAGR，因无足够正FCF数据）")

    # 计算真实数据年数和预测年数
    actual_data_years = min(historical_fcf_years, 10) if historical_fcf_years > 0 else 0
    projection_years = max(0, 10 - actual_data_years)

    # 获取真实的最新FCF（如果有的话）
    if 'historical_fcf_data' in project_params and project_params['historical_fcf_data']:
        real_latest_fcf = project_params['historical_fcf_data']['data'][-1]['fcf'] * 100000000
        real_latest_fcf_yuan = real_latest_fcf
    else:
        real_latest_fcf_yuan = None

    # 构建DCF参数表格
    dcf_process_data = [
        ['财务数据（来源：Tushare）', ''],
        ['净利润', f'{net_income/100000000:.2f} 亿元'],
    ]

    # 显示真实的最新FCF和基准FCF
    if real_latest_fcf_yuan is not None:
        real_fcf_display = f'{real_latest_fcf_yuan/100000000:.2f} 亿元'
        if using_net_income_as_fcf:
            real_fcf_display += ' （负值）'
        dcf_process_data.append(['真实最新FCF', real_fcf_display])

    # 显示基准FCF（用于估值计算）
    if using_net_income_as_fcf:
        dcf_process_data.append(['基准FCF（用于计算）', f'{base_fcf/100000000:.2f} 亿元（使用净利润）'])
    else:
        dcf_process_data.append(['基准FCF（用于计算）', f'{base_fcf/100000000:.2f} 亿元'])

    # 继续添加其他财务数据
    dcf_process_data.extend([
        ['净资产', f'{net_assets/100000000:.2f} 亿元'],
        ['总负债', f'{total_debt/100000000:.2f} 亿元'],
        ['货币资金', f'{cash_assets:.2f} 亿元'],
        ['净债务', f'{net_debt:.2f} 亿元'],
        ['总股本', f'{total_shares:,} 股'],
        [''],
        ['DCF假设参数', ''],
        ['数据来源', fcf_data_description],
        ['真实数据年数', f'{actual_data_years} 年'],
        ['预测数据年数', f'{projection_years} 年'],
        ['预测期总计', '10年'],
        ['FCF增长率（预测期）', f'{fcf_growth_example*100:.1f}% ({fcf_growth_source})'],
        ['WACC范围', '8.0% - 13.0%'],
        ['永续增长率', '3.0%']
    ])
    add_table_data(document, ['参数', '值'], dcf_process_data)

    add_paragraph(document, '')
    add_paragraph(document, '估值计算过程（示例）：')

    # 添加FCF计算方法说明
    add_title(document, '3.3.1 自由现金流（FCF）计算方法', level=3)
    add_paragraph(document, '自由现金流（Free Cash Flow, FCF）代表公司在支付了所有运营费用、资本支出后，可分配给股东和债权人的现金。')
    add_paragraph(document, '')
    add_paragraph(document, '本报告采用现金流量表法计算FCF：')
    add_paragraph(document, 'FCF = 经营活动现金流 - 资本支出')
    add_paragraph(document, '')
    add_paragraph(document, '其中：')
    add_paragraph(document, '• 经营活动现金流：来自现金流量表，反映公司主营业务的现金生成能力')
    add_paragraph(document, '• 资本支出：直接使用投资活动现金流，反映购建固定资产、无形资产和其他长期资产支付的现金')
    add_paragraph(document, '')
    add_paragraph(document, '该方法的优势：')
    add_paragraph(document, '1. 直接使用现金流量表数据，避免复杂的营运资本计算')
    add_paragraph(document, '2. 经营活动现金流已自动考虑了营运资本变化')
    add_paragraph(document, '3. 更真实地反映公司的现金生成能力')

    # 添加历史FCF数据分析
    if 'historical_fcf_data' in project_params and project_params['historical_fcf_data']:
        add_title(document, '3.3.3 历史FCF数据分析', level=3)

        historical_fcf = project_params['historical_fcf_data']
        fcf_data_list = historical_fcf['data']

        # 显示最近10年数据
        recent_fcf = fcf_data_list[-10:] if len(fcf_data_list) >= 10 else fcf_data_list

        fcf_history_headers = ['年份', '营业收入（亿）', 'NOPAT（亿）', 'FCF（亿）', 'FCF/营收']
        fcf_history_data = []
        total_fcf_to_revenue = 0
        valid_years = 0

        for year_data in recent_fcf:
            year = year_data['year']
            revenue = year_data['revenue']
            nopat = year_data['nopat']
            fcf = year_data['fcf']

            if revenue > 0:
                fcf_ratio = (fcf / revenue) * 100
                total_fcf_to_revenue += fcf_ratio
                valid_years += 1
            else:
                fcf_ratio = 0

            fcf_history_data.append([
                str(year),
                f'{revenue:.2f}',
                f'{nopat:.2f}',
                f'{fcf:.2f}',
                f'{fcf_ratio:.1f}%'
            ])

        add_table_data(document, fcf_history_headers, fcf_history_data)

        add_paragraph(document, '')
        if valid_years > 0:
            avg_fcf_ratio = total_fcf_to_revenue / valid_years
            add_paragraph(document, f'历史平均FCF/营收比率：{avg_fcf_ratio:.1f}%')
            add_paragraph(document, '该指标反映公司每100元营业收入能产生多少自由现金流。')

        # FCF趋势分析
        add_paragraph(document, '')
        add_paragraph(document, 'FCF趋势分析：')
        if len(recent_fcf) >= 3:
            recent_3_years = recent_fcf[-3:]
            avg_fcf_3y = sum(d['fcf'] for d in recent_3_years) / 3
            add_paragraph(document, f'• 最近3年平均FCF：{avg_fcf_3y:.2f} 亿元')

        latest_fcf = recent_fcf[-1]['fcf']
        add_paragraph(document, f'• 最新年度FCF（{recent_fcf[-1]["year"]}年）：{latest_fcf:.2f} 亿元')

        if latest_fcf > 0:
            add_paragraph(document, '•  FCF为正值，说明公司具备良好的现金生成能力')
        else:
            add_paragraph(document, '•  FCF为负值，可能由于大额资本支出或营运资金占用')

    # ==================== WACC计算 ====================
    add_title(document, '3.3.2 WACC（加权平均资本成本）计算', level=3)

    add_paragraph(document, 'WACC是DCF估值中的关键参数，代表投资者要求的必要收益率。本节使用CAPM模型计算WACC。')
    add_paragraph(document, '')
    add_paragraph(document, '【参数说明】')
    add_paragraph(document, '• 无风险利率（Rf）：1.85%（十年期国债收益率，2025年末数据）')
    add_paragraph(document, '• 市场收益率（Rm）：8.0%（沪深300预期收益率）')
    add_paragraph(document, '• 市场风险溢价（Rm-Rf）：6.15%')
    add_paragraph(document, '• 债务成本：无风险利率 × (1 + 50%) = 1.85% × 1.5 = 2.78%')
    add_paragraph(document, '• Beta系数：使用行业Beta（行业内成分股相对沪深300的Beta中位数）')
    add_paragraph(document, '')

    # 尝试使用tushare计算WACC
    wacc_result = None
    industry_code = None
    try:
        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if ts_token:
            import tushare as ts
            pro = ts.pro_api(ts_token)
            calculator = WACCCalculator(pro_api=pro)

            # 获取申万行业代码
            try:
                # 使用stock_basic接口获取行业信息
                df_stock_basic = pro.stock_basic(
                    ts_code=stock_code,
                    fields='ts_code,name,area,industry,list_date'
                )
                if not df_stock_basic.empty and 'industry' in df_stock_basic.columns:
                    industry = df_stock_basic['industry'].iloc[0]
                    print(f"   股票所属行业：{industry}")
                else:
                    print(f"   获取行业信息为空或缺少industry字段")
                    print(f"   将使用默认行业代码（申万电子）")
            except Exception as e:
                print(f"   获取行业代码失败: {e}")
                print(f"   将使用默认行业代码（申万电子）")

            # 计算WACC（使用行业Beta）
            # 获取第二章的同行公司数据
            peer_companies_val = context.get('peer_companies_val', None)

            wacc_result = calculator.calculate_wacc(
                stock_code=stock_code,
                market_data=project_params,  # 传递市场数据
                industry_code='801010.SI',   # 申万一级行业代码（电子）
                peer_companies=peer_companies_val,  # 传入同行公司数据
                use_industry_beta=True       # 使用行业Beta
            )

            # 显示WACC计算详情
            add_paragraph(document, '【WACC计算结果】')

            # 构建详细表格
            cap_struct = wacc_result['capital_structure']
            params = wacc_result['parameters']

            # 第一部分：参数说明
            capital_structure_data = [
                ['【参数说明】', '', ''],
                ['无风险利率（Rf）', f'{params["risk_free_rate"]*100:.2f}%', '十年期国债收益率（2025年末数据）'],
                ['市场收益率（Rm）', f'{params.get("market_return", (params["risk_free_rate"] + params["market_premium"]))*100:.1f}%', '沪深300预期收益率'],
                ['市场风险溢价（Rm-Rf）', f'{params["market_premium"]*100:.2f}%', '= 8.0% - 1.85%'],
                ['企业所得税率', f'{params["tax_rate"]:.0%}', '法定税率'],
                ['', '', ''],
                ['【Beta系数】', '', ''],
                ['个股Beta', f'{wacc_result.get("beta_stock", 1.0):.3f}', f'{stock_code} 相对沪深300（{params["beta_window"]}天窗口）'],
                ['行业Beta', f'{wacc_result.get("beta_industry", 1.0):.3f}', '行业内成分股Beta中位数'],
                ['采用Beta', f'{wacc_result["beta"]:.3f}', '使用行业Beta（更稳健）'],
                ['', '', ''],
                ['【资本结构】', '数值（亿元）', '占比'],
                ['股权市值', f'{cap_struct["market_cap"]/1e8:.2f}', f'{cap_struct["equity_ratio"]*100:.1f}%'],
                ['有息负债', f'{cap_struct.get("interest_bearing_debt", cap_struct.get("total_debt", 0))/1e8:.2f}', f'{cap_struct["debt_ratio"]*100:.1f}%'],
            ]

            # 如果有有息负债明细，添加到表格
            if cap_struct.get('debt_breakdown'):
                debt_breakdown = cap_struct['debt_breakdown']
                capital_structure_data.extend([
                    ['  - 短期借款', f'{debt_breakdown.get("短期借款", 0)/1e8:.2f}', ''],
                    ['  - 长期借款', f'{debt_breakdown.get("长期借款", 0)/1e8:.2f}', ''],
                    ['  - 应付债券', f'{debt_breakdown.get("应付债券", 0)/1e8:.2f}', ''],
                    ['  - 一年内到期非流动负债', f'{debt_breakdown.get("一年内到期非流动负债", 0)/1e8:.2f}', ''],
                ])

            capital_structure_data.extend([
                ['', '', ''],
                ['【资本成本】', '', ''],
                ['股权成本（Re）', f'{wacc_result["cost_of_equity"]*100:.2f}%', 'CAPM计算：Rf + β × (Rm-Rf)'],
                ['债务成本（税前）', f'{wacc_result["cost_of_debt"]*100:.2f}%', f'无风险利率 × (1 + {params["debt_premium"]*100:.0f}%)'],
                ['债务成本（税后）', f'{wacc_result["cost_of_debt_after_tax"]*100:.2f}%', f'扣除{params["tax_rate"]:.0%}所得税'],
                ['', '', ''],
                ['【WACC结果】', '', ''],
                ['WACC', f'{wacc_result["wacc"]*100:.2f}%', '加权平均资本成本'],
            ])

            add_table_data(document, ['', '', ''], capital_structure_data)

            # 使用计算得到的WACC
            wacc_example = wacc_result['wacc']

            add_paragraph(document, '')
            add_paragraph(document, '【WACC计算公式】')
            add_paragraph(document, f'WACC = 股权占比 × 股权成本 + 债务占比 × 税后债务成本')
            add_paragraph(document, f'     = {cap_struct["equity_ratio"]:.1%} × {wacc_result["cost_of_equity"]*100:.2f}% + {cap_struct["debt_ratio"]:.1%} × {wacc_result["cost_of_debt_after_tax"]*100:.2f}%')
            add_paragraph(document, f'     = {wacc_example*100:.2f}%')

        else:
            raise ValueError("未设置TUSHARE_TOKEN环境变量")
    except Exception as e:
        print(f" WACC自动计算失败: {e}，使用简化计算")
        add_paragraph(document, ' WACC自动计算失败，使用简化计算方法（默认参数）')
        add_paragraph(document, '')

        # 使用简化计算
        beta_from_config = project_params.get('beta', 1.0)
        wacc_example = calculate_wacc_simple(beta=beta_from_config)

        wacc_result = {
            'wacc': wacc_example,
            'cost_of_equity': 0.0185 + beta_from_config * 0.0615,
            'cost_of_debt': 0.0185 * 1.5,
            'beta': beta_from_config,
            'capital_structure': {
                'market_cap': 10_000_000_000,
                'total_debt': 3_000_000_000,
                'equity_ratio': 0.7,
                'debt_ratio': 0.3
            },
            'parameters': {
                'risk_free_rate': 0.0185,
                'market_premium': 0.0615,
                'tax_rate': 0.25
            }
        }

        # 显示简化计算的参数
        add_paragraph(document, ' 简化计算参数：')
        add_paragraph(document, f'• Beta系数：{beta_from_config:.2f}（来自配置文件）')
        add_paragraph(document, f'• 无风险利率：1.85%（十年期国债收益率）')
        add_paragraph(document, f'• 市场收益率：8.0%（沪深300预期收益率）')
        add_paragraph(document, f'• 市场风险溢价：6.15%（= 8.0% - 1.85%）')
        add_paragraph(document, f'• 股权占比：70%（默认值）')
        add_paragraph(document, f'• 债务占比：30%（默认值）')
        add_paragraph(document, f'• 企业所得税率：25%（默认值）')
        add_paragraph(document, f'• WACC结果：{wacc_example*100:.2f}%')

    add_paragraph(document, '')

    growth_example = 0.025  # 2.5% 永续增长率

    # 计算示例 - 使用真实FCF数据（如果有）
    fcfs = []
    fcf_sources = []  # 记录每个FCF的来源
    if actual_data_years > 0 and 'historical_fcf_data' in project_params:
        # 使用最新历史FCF数据作为基准，正向预测未来10年
        historical_fcf = project_params['historical_fcf_data']['data']
        latest_fcf = historical_fcf[-1]['fcf'] * 100000000  # 最新一年（2025年）FCF，转回元
        latest_year = historical_fcf[-1]['year']

        # 计算历史FCF复利增长率（CAGR）- 仅用于报告显示
        # 使用最早的正FCF年份和最新的正FCF年份计算
        positive_fcf_years = [d for d in historical_fcf if d['fcf'] > 0]

        if len(positive_fcf_years) >= 2:
            first_fcf_year = positive_fcf_years[0]
            last_fcf_year = positive_fcf_years[-1]

            first_fcf = first_fcf_year['fcf']
            last_fcf = last_fcf_year['fcf']
            years_span = last_fcf_year['year'] - first_fcf_year['year']

            # CAGR = (最终值 / 初始值)^(1/年数) - 1
            cagr_fcf_display = (last_fcf / first_fcf) ** (1 / years_span) - 1
        else:
            cagr_fcf_display = 0.05  # 默认5%（仅用于显示）
            first_fcf_year = None
            last_fcf_year = None
            first_fcf = 0
            last_fcf = 0
            years_span = 0

        # 注意：实际使用 fcf_growth_example（已在前面计算，优先使用FCF CAGR）
        # 这里不重新计算，保持与前面的一致性

        add_paragraph(document, f' 使用最新历史数据（{latest_year}年FCF {latest_fcf/100000000:.2f}亿元）作为基准，正向预测未来10年')
        add_paragraph(document, '')
        add_paragraph(document, f' 历史FCF增长率分析：')

        # 显示最近5年FCF数据
        recent_5 = historical_fcf[-5:]
        fcf_table_data = []
        for i, item in enumerate(recent_5):
            fcf_table_data.append([f"{item['year']}年", f"{item['fcf']:.2f}"])
        add_table_data(document, ['年份', 'FCF（亿）'], fcf_table_data)

        if first_fcf_year is not None:
            add_paragraph(document, f'• 历史CAGR: {cagr_fcf_display*100:.2f}%（从{first_fcf_year["year"]}年{first_fcf:.2f}亿到{last_fcf_year["year"]}年{last_fcf:.2f}亿，跨度{years_span}年）')
        else:
            add_paragraph(document, f'• 历史CAGR: {cagr_fcf_display*100:.2f}%（正FCF数据不足，使用默认增长率）')
        add_paragraph(document, f'• 采用预测增长率: {fcf_growth_example*100:.1f}%（基于历史CAGR，优先使用FCF数据计算）')

        # 基于最新FCF，预测未来10年（使用历史增长率）
        for i in range(10):
            fcf_proj = latest_fcf * ((1 + fcf_growth_example) ** (i + 1))
            fcfs.append(fcf_proj)
            if i == 0:
                fcf_sources.append(f"基准年（{latest_year}年实际）")
            else:
                fcf_sources.append(f"预测（增长{fcf_growth_example*100:.1f}%）")
    else:
        # 全部使用预测
        fcfs = [base_fcf * ((1 + fcf_growth_example) ** i) for i in range(10)]
        fcf_sources = [f"预测（增长{fcf_growth_example*100:.1f}%）" for _ in range(10)]

    # 添加逐年FCF预测和折现计算
    add_title(document, '3.3.4 逐年FCF预测与折现计算', level=3)

    add_paragraph(document, f'基于WACC={wacc_example*100:.1f}%，预测期FCF增长率={fcf_growth_example*100:.1f}%，永续增长率={growth_example*100:.1f}%：')
    add_paragraph(document, '')

    # 构建逐年计算表格
    dcf_year_headers = ['年份', 'FCF（亿）', '折现因子', '现值（亿）', '数据来源']
    dcf_year_data = []

    pv_fcfs_detail = 0
    for i, fcf in enumerate(fcfs):
        discount_factor = 1 / ((1 + wacc_example) ** (i+1))
        pv = fcf * discount_factor
        pv_fcfs_detail += pv
        year_num = i + 1

        dcf_year_data.append([
            str(year_num),
            f'{fcf/100000000:.2f}',
            f'{discount_factor:.4f}',
            f'{pv/100000000:.2f}',
            fcf_sources[i] if i < len(fcf_sources) else '预测'
        ])

    add_table_data(document, dcf_year_headers, dcf_year_data)

    add_paragraph(document, '')
    add_paragraph(document, f'预测期FCF现值合计：{pv_fcfs_detail/100000000:.2f} 亿元')

    # 添加终值计算详细过程
    add_title(document, '3.3.5 终值（Terminal Value）计算', level=3)

    terminal_fcf = fcfs[-1] * (1 + growth_example)
    terminal_value = terminal_fcf / (wacc_example - growth_example)
    pv_terminal = terminal_value / ((1 + wacc_example) ** 10)

    add_paragraph(document, '终值代表预测期后公司持续经营的价值，采用永续增长模型计算：')
    add_paragraph(document, '')
    add_paragraph(document, '计算公式：')
    add_paragraph(document, '终值 = 第10年末FCF × (1 + 永续增长率) / (WACC - 永续增长率)')
    add_paragraph(document, '')
    add_paragraph(document, '计算步骤：')
    add_paragraph(document, f'1. 第10年末FCF = {fcfs[-1]/100000000:.2f} 亿元')
    add_paragraph(document, f'2. 终值FCF = 第10年末FCF × (1 + {growth_example*100:.1f}%)')
    add_paragraph(document, f'          = {fcfs[-1]/100000000:.2f} × {1 + growth_example:.4f}')
    add_paragraph(document, f'          = {terminal_fcf/100000000:.2f} 亿元')
    add_paragraph(document, f'3. 终值 = {terminal_fcf/100000000:.2f} / ({wacc_example*100:.1f}% - {growth_example*100:.1f}%)')
    add_paragraph(document, f'        = {terminal_value/100000000:.2f} 亿元')
    add_paragraph(document, f'4. 终值现值 = {terminal_value/100000000:.2f} / (1 + {wacc_example*100:.1f}%)^10')
    add_paragraph(document, f'          = {pv_terminal/100000000:.2f} 亿元')

    # 添加企业价值和股权价值计算
    add_title(document, '3.3.6 企业价值与股权价值计算', level=3)

    enterprise_value = pv_fcfs_detail + pv_terminal  # 使用详细计算的现值

    add_paragraph(document, '企业价值（Enterprise Value）计算：')
    add_paragraph(document, '')
    add_paragraph(document, f'企业价值 = 预测期FCF现值 + 终值现值')
    add_paragraph(document, f'        = {pv_fcfs_detail/100000000:.2f} + {pv_terminal/100000000:.2f}')
    add_paragraph(document, f'        = {enterprise_value/100000000:.2f} 亿元')
    add_paragraph(document, '')
    add_paragraph(document, '股权价值（Equity Value）计算：')
    add_paragraph(document, '')
    add_paragraph(document, f'股权价值 = 企业价值 - 净债务')
    add_paragraph(document, f'        = {enterprise_value/100000000:.2f} - {net_debt:.2f}')

    # 修正：股权价值 = 企业价值 - 净债务（注意单位转换）
    equity_value = enterprise_value - (net_debt * 100000000)  # 净债务从亿元转为元

    add_paragraph(document, f'        = {equity_value/100000000:.2f} 亿元')
    add_paragraph(document, '')
    add_paragraph(document, '内在价值（Intrinsic Value）计算：')
    add_paragraph(document, '')
    add_paragraph(document, f'内在价值 = 股权价值 / 总股本')
    add_paragraph(document, f'        = {equity_value/100000000:.2f} 亿元 / {total_shares/100000000:.2f} 亿股')

    intrinsic_value = equity_value / total_shares  # 单位：元/股

    add_paragraph(document, f'        = {intrinsic_value:.2f} 元/股')

    # 添加计算汇总表格
    add_paragraph(document, '')
    add_paragraph(document, '估值计算汇总：')

    example_data = [
        ['WACC', f'{wacc_example*100:.1f}%'],
        ['FCF增长率（预测期）', f'{fcf_growth_example*100:.1f}%'],
        ['（数据来源）', fcf_growth_source],
        ['真实数据年数', f'{actual_data_years} 年'],
        ['预测数据年数', f'{projection_years} 年'],
        ['永续增长率', f'{growth_example*100:.1f}%'],
        ['预测期FCF现值', f'{pv_fcfs_detail/100000000:.2f} 亿元'],
        ['终值', f'{terminal_value/100000000:.2f} 亿元'],
        ['终值现值', f'{pv_terminal/100000000:.2f} 亿元'],
        ['企业价值', f'{enterprise_value/100000000:.2f} 亿元'],
        ['净债务', f'{net_debt:.2f} 亿元'],
        ['股权价值', f'{equity_value/100000000:.2f} 亿元'],
        ['每股价值', f'{intrinsic_value:.2f} 元/股']
    ]
    add_table_data(document, ['项目', '计算结果'], example_data)

    add_title(document, '3.4 估值结论', level=2)

    margin_market = (intrinsic_value - project_params['current_price']) / project_params['current_price'] * 100
    margin_issue = (intrinsic_value - project_params['issue_price']) / project_params['issue_price'] * 100

    add_paragraph(document, f'• DCF内在价值: {intrinsic_value:.2f} 元/股')
    add_paragraph(document, f'• 当前价格: {project_params["current_price"]:.2f} 元/股')
    add_paragraph(document, f'• 发行价格: {project_params["issue_price"]:.2f} 元/股')
    add_paragraph(document, f'• 相对市价安全边际: {margin_market:+.1f}%')
    add_paragraph(document, f'• 相对发行价安全边际: {margin_issue:+.1f}%')

    if margin_issue > 15:
        conclusion = "DCF估值显示，相比发行价有显著安全边际，估值合理偏低。"
    elif margin_issue > 0:
        conclusion = "DCF估值显示，相比发行价有一定安全边际，估值相对合理。"
    else:
        conclusion = "DCF估值显示，发行价高于内在价值，需谨慎。"

    add_paragraph(document, '')
    add_paragraph(document, f'估值结论: {conclusion}')

    add_section_break(document)

    # 更新上下文
    context['intrinsic_value'] = intrinsic_value
    context['net_debt'] = net_debt
    context['total_shares'] = total_shares
    context['enterprise_value'] = enterprise_value
    context['equity_value'] = equity_value

    return context


if __name__ == '__main__':
    # 测试代码
    print("第三章：DCF估值分析模块")
    print("本模块需要通过 generate_chapter(context) 函数调用")
def generate_dcf_valuation_heatmap(save_path, current_price, net_income, total_shares, net_debt=20.0):
    """生成DCF估值热力图

    参数:
        save_path: 保存路径
        current_price: 当前股价
        net_income: 净利润（元）
        total_shares: 总股数
        net_debt: 净债务（亿元），默认20亿元
    """
    fig, ax = plt.subplots(figsize=(14, 9))

    # WACC 和永续增长率范围
    wacc_range = np.linspace(0.08, 0.13, 7)   # 8%-13%
    growth_range = np.linspace(0.15, 0.30, 7) # 15%-30% 增长率

    # 估值参数
    # FCF = 净利润 × 100%
    base_fcf = net_income / 100000000  # 转换为亿元
    # 股数（亿股）
    shares_billion = total_shares / 100000000

    # 计算每股价值矩阵
    pivot_data = []
    for g in growth_range:
        row = []
        for w in wacc_range:
            # DCF计算：FCF按g增长
            fcfs = [base_fcf * ((1 + g) ** i) for i in range(10)]
            pv_fcfs = sum([fcf / ((1 + w) ** (i+1)) for i, fcf in enumerate(fcfs)])

            # 终值：永续增长率设为3%
            terminal_growth = 0.03
            terminal_fcf = fcfs[-1] * (1 + terminal_growth)
            terminal_value = terminal_fcf / (w - terminal_growth) if w > terminal_growth else 0
            pv_terminal = terminal_value / ((1 + w) ** 10)

            enterprise_value = pv_fcfs + pv_terminal  # 亿元
            equity_value = enterprise_value - net_debt  # 减净债务（亿元）
            value_per_share = equity_value / shares_billion  # 每股价值（元）
            row.append(value_per_share)
        pivot_data.append(row)

    pivot_table = pd.DataFrame(
        pivot_data,
        index=[f'{g*100:.1f}%' for g in growth_range],
        columns=[f'{w*100:.1f}%' for w in wacc_range]
    )

    # 绘制热力图
    heatmap = sns.heatmap(pivot_table, annot=True, fmt='.2f', cmap='RdYlGn',
                          center=current_price, vmin=current_price*0.5, vmax=current_price*1.5,
                          cbar_kws={'label': '每股价值（元）'}, ax=ax)

    # 修复 colorbar 中文
    cbar = heatmap.collections[0].colorbar
    cbar.set_label('每股价值（元）', fontproperties=font_prop, fontsize=12)

    # 设置轴标签
    ax.set_xlabel('WACC（加权平均资本成本）', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('预测期FCF增长率', fontproperties=font_prop, fontsize=12)
    ax.set_title('DCF估值敏感性分析矩阵（每股价值：元）',
                fontproperties=font_prop, fontsize=14, fontweight='bold', pad=15)

    # 设置刻度标签字体
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(10)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(10)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_var_chart(var_95, var_99, cvar_95, save_path):
    """生成VaR风险度量图表"""
    plt.figure(figsize=(10, 6))
    metrics = ['95% VaR', '99% VaR', '95% CVaR']
