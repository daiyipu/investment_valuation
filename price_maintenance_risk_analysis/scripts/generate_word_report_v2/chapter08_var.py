"""
第八章：VaR风险度量

本模块提供VaR（Value at Risk）和CVaR（Conditional Value at Risk）的计算和展示功能，
包括多窗口期VaR对比、不同置信水平下的VaR、CVaR分析、回撤分析和潜在损失金额估算。
"""

import numpy as np
import matplotlib.pyplot as plt
from docx.shared import Inches
import sys
import os

# 添加路径
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, PROJECT_DIR)

from module_utils import (
    add_title, add_paragraph, add_table_data,
    add_image, add_section_break, generate_var_chart
)


def generate_chapter(context):
    """
    生成第八章：VaR风险度量

    Parameters:
    -----------
    context : dict
        包含所有必要数据的上下文字典

    Returns:
    --------
    dict
        更新后的上下文字典
    """
    # 从context中提取变量
    document = context['document']
    project_params = context['project_params']
    market_data = context['market_data']
    IMAGES_DIR = context['IMAGES_DIR']

    # 从context['results']中获取第五章的蒙特卡洛参数
    mc_results = context['results'].get('mc_results', {})

    # 使用默认值或从第五章结果中获取
    mc_volatility_60d = mc_results.get('volatility_60d', market_data.get('volatility_60d', 0.3))
    mc_drift_60d = mc_results.get('drift_60d', market_data.get('annual_return_60d', 0.0))
    mc_volatility_120d = mc_results.get('volatility_120d', market_data.get('volatility_120d', 0.3))
    mc_drift_120d = mc_results.get('drift_120d', market_data.get('annual_return_120d', 0.0))
    mc_volatility_250d = mc_results.get('volatility_250d', market_data.get('volatility_250d', 0.3))
    mc_drift_250d = mc_results.get('drift_250d', market_data.get('annual_return_250d', 0.0))

    # ==================== 八、VaR风险度量 ====================
    add_title(document, '八、VaR风险度量', level=1)

    add_paragraph(document, '本章节使用多种方法计算风险价值（VaR）和条件风险价值（CVaR），全面评估定增项目的下行风险。')
    add_paragraph(document, '通过多窗口期VaR对比、增量风险分析等方法，提供更全面的风险度量视角。')

    # ==================== 8.1 VaR计算参数说明 ====================
    add_title(document, '8.1 VaR计算参数说明', level=2)

    add_paragraph(document, 'VaR（Value at Risk）风险价值表示在给定置信水平下，投资组合可能遭受的最大损失。')
    add_paragraph(document, '本报告使用蒙特卡洛模拟法计算VaR，基于10,000次模拟路径，分别计算60日、120日、250日三个窗口期的VaR。')
    add_paragraph(document, '')

    add_paragraph(document, 'VaR计算参数说明（多窗口期对比）：', bold=True)

    # 多窗口期参数
    var_windows = {
        '60日': {
            'vol': mc_volatility_60d,
            'drift': mc_drift_60d,
            'days': 60
        },
        '120日': {
            'vol': mc_volatility_120d,
            'drift': mc_drift_120d,
            'days': 120
        },
        '250日': {
            'vol': mc_volatility_250d,
            'drift': mc_drift_250d,
            'days': 250
        }
    }

    # 计算每个窗口期的VaR
    var_results = {}
    lockup_months = project_params['lockup_period']
    n_sim_var = 5000  # 模拟次数

    for window_name, params in var_windows.items():
        window_vol = params['vol']
        window_drift = params['drift']
        time_steps = params['days']

        # 计算锁定期的参数
        lockup_drift = window_drift * (lockup_months / 12)
        lockup_vol = window_vol * np.sqrt(lockup_months / 12)

        # 蒙特卡洛模拟
        np.random.seed(42)
        returns_var = np.random.normal(lockup_drift, lockup_vol, n_sim_var)
        final_prices_var = project_params['current_price'] * np.exp(returns_var)
        profit_losses_var = (final_prices_var - project_params['issue_price']) / project_params['issue_price']

        # 限制最大损失为-100%（本金全部亏完）
        profit_losses_var = np.maximum(profit_losses_var, -1.0)

        # 计算不同置信水平的VaR和CVaR
        # VaR取绝对值，表示损失的幅度
        var_90 = abs(np.percentile(profit_losses_var, 10))
        var_95 = abs(np.percentile(profit_losses_var, 5))
        var_99 = abs(np.percentile(profit_losses_var, 1))
        cvar_95 = abs(profit_losses_var[profit_losses_var <= np.percentile(profit_losses_var, 5)].mean())
        cvar_99 = abs(profit_losses_var[profit_losses_var <= np.percentile(profit_losses_var, 1)].mean())

        var_results[window_name] = {
            'vol': window_vol,
            'drift': window_drift,
            'var_90': var_90,
            'var_95': var_95,
            'var_99': var_99,
            'cvar_95': cvar_95,
            'cvar_99': cvar_99,
            'profit_losses': profit_losses_var
        }

    # 创建参数说明表格
    var_params_headers = ['参数', '60日窗口', '120日窗口', '250日窗口']
    var_params_data = [
        ['数据窗口', '60日（季度）', '120日（半年线）', '250日（年线）'],
        ['年化波动率', f'{var_results["60日"]["vol"]*100:.2f}%', f'{var_results["120日"]["vol"]*100:.2f}%', f'{var_results["250日"]["vol"]*100:.2f}%'],
        ['年化漂移率', f'{var_results["60日"]["drift"]*100:+.2f}%', f'{var_results["120日"]["drift"]*100:+.2f}%', f'{var_results["250日"]["drift"]*100:+.2f}%'],
        ['锁定期', f'{lockup_months}个月', f'{lockup_months}个月', f'{lockup_months}个月'],
        ['模拟次数', '5,000次', '5,000次', '5,000次']
    ]

    # 添加项目参数
    ma120_var = market_data.get('ma_120', 0)
    if ma120_var > 0:
        discount_var = (project_params["issue_price"] - ma120_var) / ma120_var * 100
        discount_note = f"（相对MA120: {ma120_var:.2f}元）"
    else:
        discount_var = (project_params["issue_price"] - project_params["current_price"]) / project_params["current_price"] * 100
        discount_note = f"（相对当前价: {project_params['current_price']:.2f}元）"

    var_params_data.extend([
        ['发行价格', f'{project_params["issue_price"]:.2f}元/股', f'{project_params["issue_price"]:.2f}元/股', f'{project_params["issue_price"]:.2f}元/股'],
        ['当前价格', f'{project_params["current_price"]:.2f}元/股', f'{project_params["current_price"]:.2f}元/股', f'{project_params["current_price"]:.2f}元/股'],
        ['折价/溢价率', f'{discount_var:+.2f}% {discount_note}', f'{discount_var:+.2f}% {discount_note}', f'{discount_var:+.2f}% {discount_note}']
    ])

    add_table_data(document, var_params_headers, var_params_data)

    add_paragraph(document, '')
    add_paragraph(document, '参数说明：')
    add_paragraph(document, '• 60日窗口：反映短期波动特征，适合评估短期风险')
    add_paragraph(document, '• 120日窗口：反映中期波动特征，平衡稳定性和时效性（推荐）')
    add_paragraph(document, '• 250日窗口：反映长期波动特征，数据最稳定，适合长期风险评估')
    add_paragraph(document, '• VaR值表示定增项目在到期时的收益率损失风险，正值表示亏损幅度')
    
    # ==================== 8.2 多窗口期VaR对比 ====================
    add_title(document, '8.2 多窗口期VaR对比', level=2)

    add_paragraph(document, '本节对比不同时间窗口下的VaR计算结果，分析窗口期选择对风险度量的影响。')
    add_paragraph(document, '')

    # 创建VaR对比表格
    var_comparison_headers = ['窗口期', '95% VaR', '99% VaR', '95% CVaR', '99% CVaR', '风险等级']

    # 定义风险等级
    def get_var_risk_level(var_value):
        if var_value <= 0.15:
            return "🟢 低风险"
        elif var_value <= 0.25:
            return "🟡 中等风险"
        elif var_value <= 0.40:
            return "🟠 较高风险"
        else:
            return "高风险"

    var_comparison_data = []
    for window in ['60日', '120日', '250日']:
        result = var_results[window]
        var_comparison_data.append([
            window,
            f'{result["var_95"]*100:.2f}%',
            f'{result["var_99"]*100:.2f}%',
            f'{result["cvar_95"]*100:.2f}%',
            f'{result["cvar_99"]*100:.2f}%',
            get_var_risk_level(result["var_95"])
        ])

    add_table_data(document, var_comparison_headers, var_comparison_data)

    add_paragraph(document, '')
    add_paragraph(document, '多窗口期VaR对比分析：', bold=True)

    # 分析窗口期差异
    var_95_values = [var_results[w]['var_95'] for w in ['60日', '120日', '250日']]
    var_95_min = min(var_95_values)
    var_95_max = max(var_95_values)
    var_95_range = var_95_max - var_95_min
    add_paragraph(document, f'• VaR数值范围：95% VaR在{var_95_min*100:.2f}%至{var_95_max*100:.2f}%之间，差异{var_95_range*100:.2f}%')
    add_paragraph(document, f'• 60日窗口VaR：{var_results["60日"]["var_95"]*100:.2f}%（{"短期波动较大" if var_results["60日"]["var_95"] > var_results["120日"]["var_95"] else "短期风险相对较低"}）')
    add_paragraph(document, f'• 120日窗口VaR：{var_results["120日"]["var_95"]*100:.2f}%（中期基准，推荐使用）')
    add_paragraph(document, f'• 250日窗口VaR：{var_results["250日"]["var_95"]*100:.2f}%（长期趋势，{"数据更稳定" if var_results["250日"]["vol"] < var_results["60日"]["vol"] else "波动特征"}）')

    # 推荐窗口期
    recommended_window = '120日'
    recommended_var = var_results[recommended_window]['var_95']
    recommended_cvar = var_results[recommended_window]['cvar_95']

    add_paragraph(document, '')
    add_paragraph(document, f'推荐窗口期：{recommended_window}')
    add_paragraph(document, f'   理由：平衡了数据稳定性和时效性，95% VaR为{recommended_var*100:.2f}%，95% CVaR为{recommended_cvar*100:.2f}%')
    add_paragraph(document, '')

    # 使用推荐窗口期的数据作为后续分析基准
    var_90 = var_results[recommended_window]['var_90']
    var_95 = recommended_var
    var_99 = var_results[recommended_window]['var_99']
    cvar_95 = recommended_cvar
    cvar_99 = var_results[recommended_window]['cvar_99']
    profit_losses = var_results[recommended_window]['profit_losses']
    mc_volatility = var_results[recommended_window]['vol']
    mc_drift = var_results[recommended_window]['drift']

    # 保存图表
    var_chart_path = f'{IMAGES_DIR}/05_var_analysis.png'
    generate_var_chart(var_95, var_99, cvar_95, var_chart_path)

    # ==================== 8.3 不同置信水平下的VaR ====================
    add_title(document, '8.3 不同置信水平下的VaR', level=2)

    # VaR表格（使用推荐窗口期的数据）
    var_table_data = [
        ['90%', f'{abs(np.percentile(profit_losses, 10))*100:.1f}%', f'{abs(profit_losses[profit_losses <= np.percentile(profit_losses, 10)].mean())*100:.1f}%', '有10%的概率损失超过此值'],
        ['95%', f'{var_95*100:.1f}%', f'{cvar_95*100:.1f}%', '有5%的概率损失超过此值'],
        ['99%', f'{var_99*100:.1f}%', f'{cvar_99*100:.1f}%', '有1%的概率损失超过此值']
    ]
    add_table_data(document, ['置信水平', 'VaR', 'CVaR', '说明'], var_table_data)

    add_paragraph(document, '图表 8.1: VaR风险度量')
    add_image(document, var_chart_path, width=Inches(6))

    # ==================== 8.3.2 蒙特卡洛模拟收益率分布图 ====================
    add_paragraph(document, '图表 8.2: 蒙特卡洛模拟收益率分布（VaR位置标注）')

    # 生成收益率分布图
    try:
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))

        # 1. 收益率分布直方图
        ax1.hist(profit_losses * 100, bins=50, color='steelblue', alpha=0.7, edgecolor='black')

        # 标注VaR位置（在负收益侧）
        ax1.axvline(x=-var_95*100, color='red', linestyle='--',
                   label=f'95% VaR (损失{var_95*100:.1f}%)', linewidth=2)
        ax1.axvline(x=-var_99*100, color='darkred', linestyle='--',
                   label=f'99% VaR (损失{var_99*100:.1f}%)', linewidth=2)
        ax1.axvline(x=0, color='green', linestyle='-', linewidth=2, label='盈亏平衡')

        ax1.set_xlabel('收益率 (%)', fontsize=12)
        ax1.set_ylabel('频数', fontsize=12)
        ax1.set_title('蒙特卡洛模拟 - 收益率分布（VaR位置标注）', fontsize=14, fontweight='bold')
        ax1.legend(loc='upper right')
        ax1.grid(True, alpha=0.3)

        # 2. VaR柱状图
        cl_labels = ['90%', '95%', '99%']
        var_values = [var_90 * 100, var_95 * 100, var_99 * 100]
        colors = ['#f39c12', '#e74c3c', '#8b0000']
        bars = ax2.bar(cl_labels, var_values, color=colors, alpha=0.7, edgecolor='black', linewidth=1.5)

        ax2.set_xlabel('置信水平', fontsize=12)
        ax2.set_ylabel('VaR (%损失)', fontsize=12)
        ax2.set_title('不同置信水平下的VaR对比', fontsize=14, fontweight='bold')
        ax2.grid(True, axis='y', alpha=0.3)

        # 添加数值标签
        for bar, value in zip(bars, var_values):
            ax2.text(bar.get_x() + bar.get_width()/2., bar.get_height(),
                    f'{value:.1f}%', ha='center', va='bottom', fontsize=11, fontweight='bold')

        plt.tight_layout()

        # 保存图表
        var_distribution_path = f'{IMAGES_DIR}/08_var_distribution_with_var.png'
        plt.savefig(var_distribution_path, dpi=300, bbox_inches='tight')
        plt.close()

        add_image(document, var_distribution_path, width=Inches(6.5))

        add_paragraph(document, '')
        add_paragraph(document, '分布图解读：', bold=True)
        add_paragraph(document, f'• 左图：收益率分布直方图，红色虚线标注95%和99% VaR位置')
        add_paragraph(document, f'• 右图：不同置信水平的VaR对比柱状图')
        add_paragraph(document, f'• 95% VaR = {var_95*100:.1f}%：有95%的概率亏损不超过此值')
        add_paragraph(document, f'• 99% VaR = {var_99*100:.1f}%：有99%的概率亏损不超过此值')
        add_paragraph(document, f'• 绿色实线：盈亏平衡点（0%）')

    except Exception as e:
        print(f" 生成收益率分布图失败: {e}")

    # ==================== 8.4 CVaR（条件风险价值）分析 ====================
    add_title(document, '8.4 CVaR（条件风险价值）分析', level=2)

    add_paragraph(document, 'CVaR（Conditional Value at Risk）又称期望短缺（Expected Shortfall），表示在损失超过VaR时的平均损失。')
    add_paragraph(document, 'CVaR比VaR更保守，能更好地反映极端情况下的风险。')
    add_paragraph(document, '')

    cvar_comparison_data = [
        ['VaR vs CVaR (95%)', f'{var_95*100:.1f}%', f'{cvar_95*100:.1f}%', f'{(cvar_95-var_95)*100:.1f}%', f'{(cvar_95/var_95-1)*100:.1f}%'],
        ['VaR vs CVaR (99%)', f'{var_99*100:.1f}%', f'{cvar_99*100:.1f}%', f'{(cvar_99-var_99)*100:.1f}%', f'{(cvar_99/var_99-1)*100:.1f}%']
    ]
    add_table_data(document, ['项目', 'VaR', 'CVaR', '差值', 'CVaR/VaR'], cvar_comparison_data)
    add_paragraph(document, 'CVaR分析结论：')
    add_paragraph(document, f'• 95%置信水平下，CVaR比VaR高出 {(cvar_95/var_95-1)*100:.1f}%')
    add_paragraph(document, f'• 99%置信水平下，CVaR比VaR高出 {(cvar_99/var_99-1)*100:.1f}%')
    add_paragraph(document, '• 说明极端情况下的损失比VaR估计的更严重')

    # 增量风险分析
    add_paragraph(document, 'VaR vs CVaR 的区别：')
    add_paragraph(document, '• VaR：回答"有5%的概率损失至少多少"')
    add_paragraph(document, '• CVaR：回答"在最差5%的情况下，平均损失多少"')
    add_paragraph(document, '• CVaR > VaR，说明极端损失的"尾部"比VaR预测的更严重')

    # ==================== 8.5 潜在损失金额估算 ====================
    add_title(document, '8.5 潜在损失金额估算', level=2)
    add_paragraph(document, '本节将VaR和CVaR转换为具体金额，直观展示潜在损失规模。')

    # 计算损失金额和剩余本金
    investment_amount = project_params['issue_price'] * project_params['issue_shares']
    loss_var_95 = investment_amount * var_95
    loss_cvar_95 = investment_amount * cvar_95
    loss_var_99 = investment_amount * var_99

    # 计算剩余本金
    remaining_var_95 = investment_amount - loss_var_95
    remaining_cvar_95 = investment_amount - loss_cvar_95
    remaining_var_99 = investment_amount - loss_var_99

    loss_amount_data = [
        ['投资总额', f'{investment_amount/10000:.2f}', '-', '-'],
        ['95% VaR潜在损失', f'{loss_var_95/10000:.2f}', f'{var_95*100:.2f}%', f'{remaining_var_95/10000:.2f}'],
        ['95% CVaR潜在损失', f'{loss_cvar_95/10000:.2f}', f'{cvar_95*100:.2f}%', f'{remaining_cvar_95/10000:.2f}'],
        ['99% VaR潜在损失', f'{loss_var_99/10000:.2f}', f'{var_99*100:.2f}%', f'{remaining_var_99/10000:.2f}']
    ]
    add_table_data(document, ['项目', '金额（万元）', '占投资比例', '剩余本金（万元）'], loss_amount_data)

    add_paragraph(document, '')
    add_paragraph(document, '金额说明：')
    add_paragraph(document, f'• 投资总额：{investment_amount/10000:.2f}万元（发行价×发行股数）')
    add_paragraph(document, '• 占投资比例：潜在损失占投资总额的百分比（即亏损比例）')
    add_paragraph(document, '• 剩余本金：投资总额减去潜在损失后的余额')
    add_paragraph(document, f'• 95% VaR：95%概率亏损不超过{loss_var_95/10000:.2f}万元（{var_95*100:.2f}%），剩余本金{remaining_var_95/10000:.2f}万元')
    add_paragraph(document, f'• 95% CVaR：最差5%情况下平均损失{loss_cvar_95/10000:.2f}万元（{cvar_95*100:.2f}%），剩余本金{remaining_cvar_95/10000:.2f}万元')
    add_paragraph(document, '')

    add_section_break(document)

    # 返回VaR相关变量供后续章节使用
    context['results']['var_90'] = var_90
    context['results']['var_95'] = var_95
    context['results']['var_99'] = var_99
    context['results']['cvar_95'] = cvar_95
    context['results']['cvar_99'] = cvar_99
    context['results']['profit_losses'] = profit_losses
    context['results']['mc_volatility'] = mc_volatility
    context['results']['mc_drift'] = mc_drift
    context['results']['var_results'] = var_results

    return context
