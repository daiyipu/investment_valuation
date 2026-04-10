# -*- coding: utf-8 -*-
"""
第六章 - 情景分析

功能：生成情景分析章节内容
"""

import sys
import os
import json
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

# 添加路径
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, PROJECT_DIR)

from utils.font_manager import get_font_prop
from utils.analysis_tools import PrivatePlacementRiskAnalyzer

# 获取中文字体
font_prop = get_font_prop()


def generate_scenario_matrix(drift_levels, vol_levels, premium_levels, scenario_name_prefix, drift_source="", project_params=None, start_idx=1):
    """
    生成情景矩阵（3×3×5=45个情景或1×3×5=15个情景）

    参数:
        drift_levels: 漂移率列表，如[高, 中, 低]
        vol_levels: 波动率列表，如[高, 中, 低]
        premium_levels: 溢价率列表，如[-0.20, -0.15, -0.10, -0.05, 0.0]
        scenario_name_prefix: 情景名称前缀，如"市场指数"
        drift_source: 漂移率来源说明，如"基于各大指数的120日年化收益率"
        project_params: 项目参数，用于计算发行价
        start_idx: 起始编号，默认为1

    返回:
        scenarios: 情景配置列表
    """
    scenarios = []

    drift_labels = ['高', '中', '低']
    vol_labels = ['高', '中', '低']
    premium_labels = ['-20%', '-15%', '-10%', '-5%', '0%']

    scenario_idx = start_idx

    for i, drift in enumerate(drift_levels):
        for j, vol in enumerate(vol_levels):
            for k, premium in enumerate(premium_levels):
                # 计算发行价（基于当前价格和溢价率）
                if project_params and 'current_price' in project_params:
                    issue_price = project_params['current_price'] * (1 + premium)
                else:
                    issue_price = 20.0  # 默认值

                scenario = {
                    'name': f'{scenario_name_prefix}-{scenario_idx:02d}',
                    'drift_level': drift_labels[i] if len(drift_levels) == 3 else '',
                    'vol_level': vol_labels[j],
                    'premium_level': premium_labels[k],
                    'drift': drift,
                    'volatility': vol,
                    'premium_rate': premium,
                    'discount': premium,  # 兼容附件要求的字段名
                    'issue_price': issue_price,  # 添加发行价字段
                    'description': f'漂移率{drift_labels[i] if len(drift_levels) == 3 else ""}({drift*100:+.2f}%)、波动率{vol_labels[j]}({vol*100:.2f}%)、溢价率{premium_labels[k]}'
                }
                scenarios.append(scenario)
                scenario_idx += 1

    return scenarios


def add_title(document, text, level=1):
    """添加标题"""
    if level == 1:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = '方正公文小标宋_GBK'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
            run.font.size = Pt(16)
    elif level == 2:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = '方正公文小标宋_GBK'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
            run.font.size = Pt(15)
    else:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(14)
    return heading


def add_paragraph(document, text, bold=False, font_size=14):
    """添加段落"""
    para = document.add_paragraph(text)
    for run in para.runs:
        run.font.bold = bold
        run.font.size = Pt(font_size)
        if '<b>' in text and '</b>' in text:
            # 处理HTML格式的粗体标记
            run.text = text.replace('<b>', '').replace('</b>', '')
            run.font.bold = True
    return para


def add_table_data(document, headers, data, font_size=12):
    """添加表格"""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 设置表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        para = header_cells[i].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.size = Pt(font_size)
        para.runs[0].font.name = '宋体'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            para = row_cells[i].paragraphs[0]
            para.runs[0].font.size = Pt(font_size)
            para.runs[0].font.name = '宋体'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if isinstance(cell_data, (int, float)):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def add_image(document, image_path, width=Inches(5)):
    """添加图片到文档"""
    if os.path.exists(image_path):
        document.add_picture(image_path, width=width)
        # 设置图片居中
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    else:
        print(f" 图片不存在: {image_path}")
        return False


def add_section_break(document):
    """添加分页符"""
    document.add_page_break()


def generate_multi_dimension_scenario_charts_split(current_price, base_price, volatility, drift, lockup_period, save_dir):
    """生成多维度情景图表 - 拆分版本"""
    from scipy import stats

    os.makedirs(save_dir, exist_ok=True)
    chart_paths = []

    # 定义参数范围
    drift_range = np.arange(-0.30, 0.35, 0.05)
    vol_range = np.arange(0.10, 0.55, 0.10)
    discount_range = np.arange(-0.20, 0.25, 0.05)

    all_results = []
    lockup_years = lockup_period / 12
    n_sim = 1000

    for d in drift_range:
        for v in vol_range:
            for disc in discount_range:
                issue_price = base_price * (1 + disc)
                lockup_drift = d * lockup_years
                lockup_vol = v * np.sqrt(lockup_years)

                np.random.seed(42)
                # sim_returns是对数收益率（连续复利）
                sim_returns = np.random.normal(lockup_drift, lockup_vol, n_sim)
                final_prices = current_price * np.exp(sim_returns)
                # 使用对数收益率计算总收益
                total_log_returns = np.log(final_prices / issue_price)
                # 年化：总对数收益率 / 年数
                annualized_log_returns = total_log_returns / lockup_years

                profit_prob = (total_log_returns > 0).mean() * 100
                mean_return = annualized_log_returns.mean() * 100

                all_results.append({
                    'drift': d,
                    'volatility': v,
                    'discount': disc,
                    'profit_prob': profit_prob,
                    'mean_return': mean_return
                })

    df = pd.DataFrame(all_results)

    # 图1: 不同波动率下的情景对比
    fig, ax = plt.subplots(figsize=(14, 7))
    vol_groups = df.groupby('volatility')

    for vol, group in vol_groups:
        ax.plot(group['drift'] * 100, group['profit_prob'],
               marker='o', label=f'波动率 {vol*100:.0f}%', linewidth=2, markersize=6)

    ax.set_xlabel('漂移率 (%)', fontproperties=font_prop, fontsize=13)
    ax.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=13)
    ax.set_title('不同波动率下的盈利概率对比', fontproperties=font_prop, fontsize=15, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=11)
    ax.grid(True, alpha=0.3)
    ax.axhline(y=50, color='red', linestyle='--', alpha=0.5, label='盈亏平衡线')

    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(11)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(11)

    plt.tight_layout()
    chart1_path = os.path.join(save_dir, '06_06_volatility_comparison.png')
    plt.savefig(chart1_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart1_path)

    # 图2: 优质情景TOP10
    top_scenarios = df.nlargest(10, 'profit_prob')

    fig, ax = plt.subplots(figsize=(14, 7))
    colors = ['green' if p >= 70 else 'orange' if p >= 50 else 'red' for p in top_scenarios['profit_prob']]
    bars = ax.bar(range(len(top_scenarios)), top_scenarios['profit_prob'], color=colors, alpha=0.7, edgecolor='black')

    labels = [f"D:{row['drift']*100:+.0f}%, V:{row['volatility']*100:.0f}%, Dsc:{row['discount']*100:+.0f}%"
              for _, row in top_scenarios.iterrows()]
    ax.set_xticks(range(len(top_scenarios)))
    ax.set_xticklabels(labels, fontproperties=font_prop, fontsize=9, rotation=45, ha='right')
    ax.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=13)
    ax.set_title('优质情景TOP10 (盈利概率排序)', fontproperties=font_prop, fontsize=15, fontweight='bold')
    ax.set_ylim(0, 100)
    ax.grid(True, alpha=0.3, axis='y')

    for i, (bar, prob) in enumerate(zip(bars, top_scenarios['profit_prob'])):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
               f'{prob:.1f}%', ha='center', va='bottom', fontproperties=font_prop, fontsize=10)

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(11)

    plt.tight_layout()
    chart2_path = os.path.join(save_dir, '06_07_top_scenarios.png')
    plt.savefig(chart2_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart2_path)

    return chart_paths


def generate_chapter(context):
    """
    生成第六章内容：情景分析

    参数:
        context: 包含所有必要数据的字典
            - document: Word文档对象
            - project_params: 项目参数
            - market_data: 市场数据
            - IMAGES_DIR: 图片保存目录
            - DATA_DIR: 数据目录
            - analyzer: 分析器对象
            - risk_params: 风险参数
            - stock_code: 股票代码
            - ma20: MA20价格
            - stock_pe_data: 个股PE数据（可选）
            - industry_pe_data: 行业PE数据（可选）
            - intrinsic_value: DCF内在价值（可选）
    """
    document = context['document']
    project_params = context['project_params']
    market_data = context['market_data']
    IMAGES_DIR = context['IMAGES_DIR']
    DATA_DIR = context.get('DATA_DIR', os.path.join(os.path.dirname(PROJECT_DIR), 'data'))
    analyzer = context['analyzer']
    risk_params = context.get('risk_params', {})
    stock_code = context.get('stock_code', '300735.SZ')
    ma20 = context.get('ma20', market_data.get('ma_20', project_params['current_price']))
    stock_pe_data = context.get('stock_pe_data', None)
    industry_pe_data = context.get('industry_pe_data', None)
    intrinsic_value = context.get('intrinsic_value', None)

    # ==================== 六、情景分析 ====================
    add_title(document, '六、情景分析', level=1)

    add_paragraph(document, '本章节分析定增项目在不同预设情景下的风险表现，包括多维度情景分析（预期期间收益率、净利润率、波动率、发行价折价）以及破发概率情景分析。')

    # ==================== 6.1 多参数情景分析 ====================
    add_paragraph(document, '')
    add_title(document, '6.1 多参数情景分析', level=2)

    add_paragraph(document, '本节通过穷举漂移率、波动率、溢价率三个参数的组合，全面分析不同市场环境下的定增项目收益预期。')
    
    # ====== 参数空间定义 ======
    add_paragraph(document, '6.1.1 参数空间定义', bold=True, font_size=14)
    add_paragraph(document, '通过以下参数的穷举组合，模拟585种不同市场情景：')

    # 定义参数范围
    drift_rates = np.arange(-0.30, 0.35, 0.05)  # -30% 到 +30%，每档5%
    volatilities = np.arange(0.10, 0.55, 0.10)  # 10% 到 50%，每档10%
    discounts = np.arange(-0.20, 0.25, 0.05)  # -20% 到 +20%，每档5%

    # 参数说明表格
    param_data = [
        ['参数', '范围', '步长', '组合数'],
        ['漂移率（年化收益率）', '-30% ~ +30%', '5%', len(drift_rates)],
        ['波动率（年化）', '10% ~ 50%', '10%', len(volatilities)],
        ['溢价率', '-20% ~ +20%', '5%', len(discounts)],
        ['总组合数', '-', '-', f'{len(drift_rates) * len(volatilities) * len(discounts)}']
    ]
    add_table_data(document, ['参数', '范围', '步长', '组合数'], param_data)

    add_paragraph(document, '')
    add_paragraph(document, '参数说明：')
    add_paragraph(document, f'• 当前价格: {project_params["current_price"]:.2f} 元/股')
    add_paragraph(document, f'• MA20价格: {market_data.get("ma_20", 0):.2f} 元/股（作为发行定价基准）')
    add_paragraph(document, f'• 锁定期: {project_params["lockup_period"]} 个月')
    add_paragraph(document, f'• 漂移率: 反映股价的预期趋势（负值=下跌，正值=上涨）')
    add_paragraph(document, f'• 波动率: 反映股价的不确定性（越高=风险越大）')
    add_paragraph(document, f'• 溢价率: 发行价相对MA20的溢价（负值=折价，正值=溢价，与配置的premium_rate一致）')
    
    # ====== 模拟所有组合 ======
    print("\n运行多参数组合模拟...")

    all_scenarios = []
    n_sim = 2000  # 每个组合模拟2000次
    lockup_years = project_params['lockup_period'] / 12
    current_price = project_params['current_price']
    ma20_price = market_data.get('ma_20', current_price)  # 使用MA20作为发行定价基准

    scenario_count = 0
    for drift in drift_rates:
        for vol in volatilities:
            for discount in discounts:
                scenario_count += 1

                # 计算发行价（使用MA20作为基准，与配置保持一致）
                issue_price = ma20_price * (1 + discount)

                # 蒙特卡洛模拟
                lockup_drift = drift * lockup_years
                lockup_vol = vol * np.sqrt(lockup_years)

                np.random.seed(42)  # 固定种子以确保可重复性
                # sim_returns是对数收益率（连续复利）
                sim_returns = np.random.normal(lockup_drift, lockup_vol, n_sim)
                final_prices = current_price * np.exp(sim_returns)

                # 计算收益：使用对数收益率（连续复利）
                total_log_returns = np.log(final_prices / issue_price)
                # 年化对数收益率
                annualized_log_returns = total_log_returns / lockup_years

                # 统计指标
                mean_return = annualized_log_returns.mean()
                median_return = np.median(annualized_log_returns)
                profit_prob = (total_log_returns > 0).mean() * 100
                var_5 = np.percentile(annualized_log_returns, 5)
                var_95 = np.percentile(annualized_log_returns, 95)

                all_scenarios.append({
                    'drift': drift,
                    'volatility': vol,
                    'discount': discount,
                    'issue_price': issue_price,
                    'mean_return': mean_return,
                    'median_return': median_return,
                    'profit_prob': profit_prob,
                    'var_5': var_5,
                    'var_95': var_95
                })

                if scenario_count % 100 == 0:
                    print(f"  已完成 {scenario_count}/{len(drift_rates) * len(volatilities) * len(discounts)} 个情景...")

    print(f"  完成！共模拟 {len(all_scenarios)} 个情景")

    # ====== 按收益和概率排序 ======
    add_paragraph(document, '')
    add_paragraph(document, '6.1.2 情景分析结果', bold=True, font_size=14)

    # 按预期年化收益排序（Top 20）
    add_paragraph(document, '表1: 按预期年化收益率排序的Top 20情景')
    add_paragraph(document, '')

    top_by_return = sorted(all_scenarios, key=lambda x: x['mean_return'], reverse=True)[:20]

    top_return_data = []
    for i, s in enumerate(top_by_return, 1):
        top_return_data.append([
            f"{i}",
            f"{s['drift']*100:+.0f}%",
            f"{s['volatility']*100:.0f}%",
            f"{s['discount']*100:+.0f}%",
            f"{s['issue_price']:.2f} 元",
            f"{s['mean_return']*100:+.2f}%",
            f"{s['profit_prob']:.1f}%"
        ])

    return_headers = ['排名', '漂移率', '波动率', '溢价率', '发行价', '预期年化收益', '盈利概率']
    add_table_data(document, return_headers, top_return_data)

    # 按盈利概率排序（Top 20）
    add_paragraph(document, '')
    add_paragraph(document, '表2: 按盈利概率排序的Top 20情景')
    add_paragraph(document, '')

    top_by_prob = sorted(all_scenarios, key=lambda x: x['profit_prob'], reverse=True)[:20]

    top_prob_data = []
    for i, s in enumerate(top_by_prob, 1):
        top_prob_data.append([
            f"{i}",
            f"{s['drift']*100:+.0f}%",
            f"{s['volatility']*100:.0f}%",
            f"{s['discount']*100:+.0f}%",
            f"{s['issue_price']:.2f} 元",
            f"{s['mean_return']*100:+.2f}%",
            f"{s['profit_prob']:.1f}%"
        ])

    add_table_data(document, return_headers, top_prob_data)

    # ====== 综合最优情景 ======
    add_paragraph(document, '')
    add_paragraph(document, '6.1.3 综合最优情景分析', bold=True, font_size=14)

    # 找出同时满足高收益和高概率的情景
    high_return = [s for s in all_scenarios if s['mean_return'] > 0.20]  # 年化收益>20%
    high_prob = [s for s in all_scenarios if s['profit_prob'] > 80]  # 盈利概率>80%
    best_scenarios = [s for s in high_return if s in high_prob]

    if best_scenarios:
        add_paragraph(document, f'找到 {len(best_scenarios)} 个同时满足"年化收益>20%"和"盈利概率>80%"的最优情景：')
        add_paragraph(document, '')

        best_data = []
        for i, s in enumerate(sorted(best_scenarios, key=lambda x: x['mean_return'], reverse=True)[:10], 1):
            best_data.append([
                f"{i}",
                f"{s['drift']*100:+.0f}%",
                f"{s['volatility']*100:.0f}%",
                f"{s['discount']*100:+.0f}%",
                f"{s['mean_return']*100:+.2f}%",
                f"{s['profit_prob']:.1f}%"
            ])
        add_table_data(document, ['排名', '漂移率', '波动率', '溢价率', '预期年化收益', '盈利概率'], best_data)
    else:
        add_paragraph(document, '未找到同时满足"年化收益>20%"和"盈利概率>80%"的情景。')

    # ====== 当前项目定位 ======
    add_paragraph(document, '')
    add_paragraph(document, '6.1.4 当前项目在所有情景中的定位', bold=True, font_size=14)

    current_drift = market_data.get('annual_return_120d', risk_params.get('drift', 0.08))
    current_vol = market_data.get('volatility_120d', risk_params.get('volatility', 0.30))
    # 使用MA20作为基准计算折价率（与配置保持一致）
    ma20_for_discount = market_data.get('ma_20', project_params['current_price'])
    current_discount = (project_params['issue_price'] / ma20_for_discount) - 1

    # 找到最接近当前参数的情景
    for s in all_scenarios:
        s['distance'] = (
            abs(s['drift'] - current_drift) +
            abs(s['volatility'] - current_vol) +
            abs(s['discount'] - current_discount)
        )

    closest_scenario = min(all_scenarios, key=lambda x: x['distance'])

    # 计算当前项目的排名
    return_rank = sorted(all_scenarios, key=lambda x: x['mean_return'], reverse=True).index(closest_scenario) + 1
    prob_rank = sorted(all_scenarios, key=lambda x: x['profit_prob'], reverse=True).index(closest_scenario) + 1

    add_paragraph(document, '基于当前市场参数的最接近情景：')
    add_paragraph(document, f'• 漂移率: {current_drift*100:+.2f}%')
    add_paragraph(document, f'• 波动率: {current_vol*100:.2f}%')
    add_paragraph(document, f'• 溢价率: {current_discount*100:+.2f}% （相对MA20: {ma20_for_discount:.2f}元）')
    add_paragraph(document, f'• 发行价: {project_params["issue_price"]:.2f} 元/股')
    add_paragraph(document, f'• 预期年化收益（算术平均）: {closest_scenario["mean_return"]*100:+.2f}%')
    add_paragraph(document, f'• 收益率中位数: {closest_scenario["median_return"]*100:+.2f}%')
    add_paragraph(document, f'• 盈利概率: {closest_scenario["profit_prob"]:.1f}%')
    add_paragraph(document, f'• 收益率排名: 第 {return_rank} 名 / 共 {len(all_scenarios)} 个情景')
    add_paragraph(document, f'• 盈利概率排名: 第 {prob_rank} 名 / 共 {len(all_scenarios)} 个情景')

    # 保存情景数据供附件使用（分两种类型）
    multi_param_scenarios = all_scenarios.copy()  # 6.1节的585种多参数构造情景
    # 注意：不初始化historical_scenarios_for_appendix，避免6.2-6.5节情景混入585种情景表中
    historical_scenarios_for_appendix = []  # 6.2-6.5节的历史数据情景

    # ==================== 6.2.1 市场指数情景分析 ====================
    add_paragraph(document, '')
    add_title(document, '6.2.1 市场指数情景分析', level=2)

    add_paragraph(document, '本节基于各大市场指数（沪深300、中证500、创业板指、科创50）的历史数据，构建全面的情景矩阵进行分析。')
    add_paragraph(document, '通过组合不同漂移率、波动率和溢价率水平，生成3×3×5=45个情景，全面评估市场环境对定增收益的影响。')
    add_paragraph(document, '')

    # 加载指数数据
    indices_data_available = False
    index_vol_values = []
    index_return_values = []

    try:
        indices_data_file = os.path.join(DATA_DIR, 'market_indices_scenario_data_v2.json')
        if os.path.exists(indices_data_file):
            with open(indices_data_file, 'r', encoding='utf-8') as f:
                indices_data_loaded = json.load(f)

            # 收集主要指数的波动率和收益率
            major_indices = ['沪深300', '中证500', '创业板指', '科创50']

            for index_name in major_indices:
                if index_name in indices_data_loaded:
                    index_vol_values.append(indices_data_loaded[index_name]['volatility_120d'])
                    index_return_values.append(indices_data_loaded[index_name]['return_120d'])

            if index_vol_values and index_return_values:
                indices_data_available = True
                print(f" 已加载指数数据: {indices_data_file}")
                print(f"   涵盖指数: {', '.join(major_indices)}")
    except Exception as e:
        print(f" 无法加载指数数据，使用默认值: {e}")

    # 如果指数数据可用，进行情景分析
    if indices_data_available:
        # 计算指数波动率和收益率的高、中、低分位
        index_vol_high = max(index_vol_values) * 1.1  # 高波动率
        index_vol_mid = np.mean(index_vol_values)      # 中波动率
        index_vol_low = min(index_vol_values) * 0.9    # 低波动率

        index_drift_high = max(index_return_values) * 1.1  # 高漂移率
        index_drift_mid = np.mean(index_return_values)      # 中漂移率
        index_drift_low = min(index_return_values) * 0.9    # 低漂移率

        # 溢价率档次：-20%, -15%, -10%, -5%, 0%
        premium_levels = [-0.20, -0.15, -0.10, -0.05, 0.0]

        # 生成市场指数情景矩阵（3×3×5=45个情景）
        index_scenarios = generate_scenario_matrix(
            drift_levels=[index_drift_high, index_drift_mid, index_drift_low],
            vol_levels=[index_vol_high, index_vol_mid, index_vol_low],
            premium_levels=premium_levels,
            scenario_name_prefix='市场指数',
            drift_source='基于各大指数的120日年化收益率',
            project_params=project_params
        )

        print(f" 生成市场指数情景矩阵: {len(index_scenarios)}个情景")

        # 运行市场指数情景模拟
        index_scenario_results = []
        for idx, scenario in enumerate(index_scenarios):
            try:
                sim_result = analyzer.monte_carlo_simulation(
                    n_simulations=5000,
                    time_steps=120,
                    volatility=scenario['volatility'],
                    drift=scenario['drift'],
                    seed=42+idx  # 使用不同的seed确保每个情景的随机性不同
                )

                final_prices = sim_result.iloc[:, -1].values
                # 使用情景相关的发行价（考虑了溢价率）
                scenario_issue_price = scenario['issue_price']
                # 对数收益率（连续复利）
                log_returns = np.log(final_prices / scenario_issue_price)
                # 年化：120日 = 120个交易日
                annualized_log_returns = log_returns * (252.0 / 120)

                # 计算VaR（使用年化对数收益率）
                var_5 = np.percentile(annualized_log_returns, 5)   # 5%分位数（最坏5%）
                var_95 = np.percentile(annualized_log_returns, 95) # 95%分位数（最好5%）

                # 包装成兼容旧代码的格式
                scenario_result = {
                    'scenario': scenario,  # 将原始情景对象嵌套在scenario键中
                    'profit_prob': (log_returns > 0).mean() * 100,
                    'median_return': np.median(annualized_log_returns),
                    'mean_return': np.mean(annualized_log_returns),
                    'var_5': var_5,
                    'var_95': var_95,
                    'actual_premium_rate': scenario['premium_rate']
                }

                index_scenario_results.append(scenario_result)
            except Exception as e:
                print(f" 情景{scenario['name']}模拟失败: {e}")

        # 添加到汇总列表
        all_scenarios.extend(index_scenario_results)

        # 为附件展开情景对象（附件期望扁平结构）
        for scenario_result in index_scenario_results:
            scenario_flat = scenario_result['scenario'].copy()
            scenario_flat.update({
                'profit_prob': scenario_result['profit_prob'],
                'median_return': scenario_result['median_return'],
                'mean_return': scenario_result['mean_return'],
                'var_5': scenario_result['var_5'],
                'var_95': scenario_result['var_95'],
                'issue_price': scenario_result['scenario']['issue_price']
            })
            historical_scenarios_for_appendix.append(scenario_flat)

        # 生成市场指数情景对比表格
        if index_scenario_results:
            add_paragraph(document, '市场指数情景参数表（全部45个情景）：')
            add_paragraph(document, '')

            # 按盈利概率排序，展示全部情景
            sorted_scenarios = sorted(index_scenario_results, key=lambda x: x['profit_prob'], reverse=True)

            index_table_data = []
            for scenario_result in sorted_scenarios:
                scenario = scenario_result['scenario']  # 获取嵌套的情景对象
                index_table_data.append([
                    scenario['name'],
                    f"{scenario['drift_level']}({scenario['drift']*100:+.1f}%)",
                    f"{scenario['vol_level']}({scenario['volatility']*100:.1f}%)",
                    scenario['premium_level'],
                    f"{scenario_result['profit_prob']:.1f}%",
                    f"{scenario_result['median_return']*100:+.1f}%"
                ])

            add_table_data(document, ['情景', '漂移率', '波动率', '溢价率', '盈利概率', '中位数收益'], index_table_data)

            add_paragraph(document, '')
            add_paragraph(document, '市场指数情景分析结论：')
            best_scenario = max(index_scenario_results, key=lambda x: x['profit_prob'])
            worst_scenario = min(index_scenario_results, key=lambda x: x['profit_prob'])
            add_paragraph(document, f'• 最优情景：{best_scenario["scenario"]["name"]}，盈利概率{best_scenario["profit_prob"]:.1f}%，中位数收益{best_scenario["median_return"]*100:+.1f}%')
            add_paragraph(document, f'• 最差情景：{worst_scenario["scenario"]["name"]}，盈利概率{worst_scenario["profit_prob"]:.1f}%，中位数收益{worst_scenario["median_return"]*100:+.1f}%')

        
    # ==================== 6.2.2 行业指数情景分析 ====================
    add_paragraph(document, '')
    add_title(document, '6.2.2 行业指数情景分析', level=2)

    add_paragraph(document, '本节基于标的股票所属行业指数的历史数据，构建全面的情景矩阵进行分析。')
    add_paragraph(document, '通过组合不同漂移率、波动率和溢价率水平，生成3×3×5=45个情景，评估行业环境对定增收益的影响。')
    add_paragraph(document, '')

    # 加载行业数据
    industry_data_available = False
    industry_vol_values = []
    industry_return_values = []

    try:
        industry_data_file = os.path.join(DATA_DIR, f'{stock_code.replace(".", "_")}_industry_data.json')
        if os.path.exists(industry_data_file):
            with open(industry_data_file, 'r', encoding='utf-8') as f:
                industry_data_loaded = json.load(f)

            # 行业数据文件直接包含波动率和收益率字段
            if 'volatility_120d' in industry_data_loaded:
                industry_vol_values.append(industry_data_loaded.get('volatility_120d', 0.35))
                industry_return_values.append(industry_data_loaded.get('annual_return_120d', 0.08))

            if industry_vol_values and industry_return_values:
                industry_data_available = True
                print(f" 已加载行业数据: {industry_data_file}")
                print(f"   行业波动率: {industry_vol_values[0]*100:.2f}%")
                print(f"   行业收益率: {industry_return_values[0]*100:+.2f}%")
    except Exception as e:
        print(f" 无法加载行业数据: {e}")

    # 如果行业数据可用，进行情景分析
    if industry_data_available:
        # 计算行业波动率和收益率的高、中、低分位
        industry_vol_high = max(industry_vol_values) * 1.1  # 高波动率
        industry_vol_mid = np.mean(industry_vol_values)      # 中波动率
        industry_vol_low = min(industry_vol_values) * 0.9    # 低波动率

        industry_drift_high = max(industry_return_values) * 1.1  # 高漂移率
        industry_drift_mid = np.mean(industry_return_values)      # 中漂移率
        industry_drift_low = min(industry_return_values) * 0.9    # 低漂移率

        # 溢价率档次：-20%, -15%, -10%, -5%, 0%
        premium_levels = [-0.20, -0.15, -0.10, -0.05, 0.0]

        # 生成行业指数情景矩阵（3×3×5=45个情景）
        industry_scenarios = generate_scenario_matrix(
            drift_levels=[industry_drift_high, industry_drift_mid, industry_drift_low],
            vol_levels=[industry_vol_high, industry_vol_mid, industry_vol_low],
            premium_levels=premium_levels,
            scenario_name_prefix='行业指数',
            drift_source='基于行业指数的120日年化收益率',
            project_params=project_params
        )

        print(f" 生成行业指数情景矩阵: {len(industry_scenarios)}个情景")

        # 运行行业指数情景模拟
        industry_scenario_results = []
        for idx, scenario in enumerate(industry_scenarios):
            try:
                sim_result = analyzer.monte_carlo_simulation(
                    n_simulations=5000,
                    time_steps=120,
                    volatility=scenario['volatility'],
                    drift=scenario['drift'],
                    seed=42+idx  # 使用不同的seed确保每个情景的随机性不同
                )

                final_prices = sim_result.iloc[:, -1].values
                # 使用情景相关的发行价（考虑了溢价率）
                scenario_issue_price = scenario['issue_price']
                # 对数收益率（连续复利）
                log_returns = np.log(final_prices / scenario_issue_price)
                # 年化：120日 = 120个交易日
                annualized_log_returns = log_returns * (252.0 / 120)

                # 计算VaR（使用年化对数收益率）
                var_5 = np.percentile(annualized_log_returns, 5)   # 5%分位数（最坏5%）
                var_95 = np.percentile(annualized_log_returns, 95) # 95%分位数（最好5%）

                # 包装成兼容旧代码的格式
                scenario_result = {
                    'scenario': scenario,  # 将原始情景对象嵌套在scenario键中
                    'profit_prob': (log_returns > 0).mean() * 100,
                    'median_return': np.median(annualized_log_returns),
                    'mean_return': np.mean(annualized_log_returns),
                    'var_5': var_5,
                    'var_95': var_95,
                    'actual_premium_rate': scenario['premium_rate']
                }

                industry_scenario_results.append(scenario_result)
            except Exception as e:
                print(f" 情景{scenario['name']}模拟失败: {e}")

        # 添加到汇总列表
        all_scenarios.extend(industry_scenario_results)

        # 为附件展开情景对象（附件期望扁平结构）
        for scenario_result in industry_scenario_results:
            scenario_flat = scenario_result['scenario'].copy()
            scenario_flat.update({
                'profit_prob': scenario_result['profit_prob'],
                'median_return': scenario_result['median_return'],
                'mean_return': scenario_result['mean_return'],
                'var_5': scenario_result['var_5'],
                'var_95': scenario_result['var_95'],
                'issue_price': scenario_result['scenario']['issue_price']
            })
            historical_scenarios_for_appendix.append(scenario_flat)

        # 生成行业指数情景对比表格
        if industry_scenario_results:
            add_paragraph(document, '行业指数情景参数表（全部45个情景）：')
            add_paragraph(document, '')

            # 按盈利概率排序，展示全部情景
            sorted_scenarios = sorted(industry_scenario_results, key=lambda x: x['profit_prob'], reverse=True)

            industry_table_data = []
            for scenario_result in sorted_scenarios:
                scenario = scenario_result['scenario']  # 获取嵌套的情景对象
                industry_table_data.append([
                    scenario['name'],
                    f"{scenario['drift_level']}({scenario['drift']*100:+.1f}%)",
                    f"{scenario['vol_level']}({scenario['volatility']*100:.1f}%)",
                    scenario['premium_level'],
                    f"{scenario_result['profit_prob']:.1f}%",
                    f"{scenario_result['median_return']*100:+.1f}%"
                ])

            add_table_data(document, ['情景', '漂移率', '波动率', '溢价率', '盈利概率', '中位数收益'], industry_table_data)

            add_paragraph(document, '')
            add_paragraph(document, '行业指数情景分析结论：')
            best_scenario = max(industry_scenario_results, key=lambda x: x['profit_prob'])
            worst_scenario = min(industry_scenario_results, key=lambda x: x['profit_prob'])
            add_paragraph(document, f'• 最优情景：{best_scenario["scenario"]["name"]}，盈利概率{best_scenario["profit_prob"]:.1f}%，中位数收益{best_scenario["median_return"]*100:+.1f}%')
            add_paragraph(document, f'• 最差情景：{worst_scenario["scenario"]["name"]}，盈利概率{worst_scenario["profit_prob"]:.1f}%，中位数收益{worst_scenario["median_return"]*100:+.1f}%')

        
        # ==================== 6.3 基于行业PE分位数的情景分析 ====================
        add_paragraph(document, '')
        add_title(document, '6.3 基于行业PE分位数的情景分析', level=2)

        add_paragraph(document, '本节基于行业PE分位数的估值水平，构建全面的情景矩阵进行分析。')
        add_paragraph(document, '通过组合不同PE分位数（高/中/低）、波动率和溢价率，生成3×3×5=45个情景，评估估值回归对定增收益的影响。')
        add_paragraph(document, '')

        add_paragraph(document, '行业PE分位数计算逻辑：')
        add_paragraph(document, '• 收集行业历史PE数据（通常为5年）')
        add_paragraph(document, '• 计算75%、50%、25%分位数，分别代表高估、中性、低估水平')
        add_paragraph(document, '• 漂移率 = (行业PE分位数 / 个股当前PE - 1)，反映估值回归潜力')
        add_paragraph(document, '• 如果个股PE低于行业PE分位数，估值便宜，未来收益率应该更高')
        add_paragraph(document, '')

        # 检查行业PE数据
        if industry_pe_data is not None and stock_pe_data is not None:
            try:
                # 计算行业PE分位数
                industry_pe_values = industry_pe_data['pe_ttm'].dropna()
                pe_75 = industry_pe_values.quantile(0.75)
                pe_50 = industry_pe_values.quantile(0.50)
                pe_25 = industry_pe_values.quantile(0.25)

                # 获取标的个股当前PE值
                stock_pe_current = stock_pe_data['pe_ttm'].iloc[-1]

                # 根据行业PE分位数与标的个股PE的比值计算漂移率
                ratio_75 = pe_75 / stock_pe_current if stock_pe_current > 0 else 1.0
                ratio_50 = pe_50 / stock_pe_current if stock_pe_current > 0 else 1.0
                ratio_25 = pe_25 / stock_pe_current if stock_pe_current > 0 else 1.0

                # 漂移率 = 比值 - 1（直接作为增长率假设）
                drift_pe_75 = ratio_75 - 1.0
                drift_pe_50 = ratio_50 - 1.0
                drift_pe_25 = ratio_25 - 1.0

                # 限制漂移率在合理范围内（-50%到+100%）
                drift_pe_75 = max(-0.50, min(1.00, drift_pe_75))
                drift_pe_50 = max(-0.50, min(1.00, drift_pe_50))
                drift_pe_25 = max(-0.50, min(1.00, drift_pe_25))

                print(f" 行业PE分位数漂移率计算完成:")
                print(f"   标的个股当前PE: {stock_pe_current:.2f}倍")
                print(f"   行业PE 75%分位数: {pe_75:.1f}倍, 比值={ratio_75:.2f}, 漂移率={drift_pe_75*100:+.2f}%")
                print(f"   行业PE 50%分位数: {pe_50:.1f}倍, 比值={ratio_50:.2f}, 漂移率={drift_pe_50*100:+.2f}%")
                print(f"   行业PE 25%分位数: {pe_25:.1f}倍, 比值={ratio_25:.2f}, 漂移率={drift_pe_25*100:+.2f}%")

                # 使用市场指数的波动率（高、中、低三档）
                if index_vol_values:
                    index_vol_high = max(index_vol_values) * 1.1
                    index_vol_mid = np.mean(index_vol_values)
                    index_vol_low = min(index_vol_values) * 0.9
                else:
                    # 如果没有指数数据，使用市场数据
                    index_vol_high = market_data.get('volatility_120d', 0.35) * 1.1
                    index_vol_mid = market_data.get('volatility_120d', 0.35)
                    index_vol_low = market_data.get('volatility_120d', 0.35) * 0.9

                # 溢价率档次：-20%, -15%, -10%, -5%, 0%
                premium_levels = [-0.20, -0.15, -0.10, -0.05, 0.0]

                # 生成行业PE情景矩阵（3×3×5=45个情景）
                industry_pe_scenarios = generate_scenario_matrix(
                    drift_levels=[drift_pe_75, drift_pe_50, drift_pe_25],
                    vol_levels=[index_vol_high, index_vol_mid, index_vol_low],
                    premium_levels=premium_levels,
                    scenario_name_prefix='行业PE',
                    drift_source='基于行业PE分位数与个股当前PE的比值',
                    project_params=project_params
                )

                print(f" 生成行业PE情景矩阵: {len(industry_pe_scenarios)}个情景")

                # 运行行业PE情景模拟
                industry_pe_scenario_results = []
                for idx, scenario in enumerate(industry_pe_scenarios):
                    try:
                        sim_result = analyzer.monte_carlo_simulation(
                            n_simulations=5000,
                            time_steps=120,
                            volatility=scenario['volatility'],
                            drift=scenario['drift'],
                            seed=42+idx  # 使用不同的seed确保每个情景的随机性不同
                        )

                        final_prices = sim_result.iloc[:, -1].values
                        # 使用情景相关的发行价（考虑了溢价率）
                        scenario_issue_price = scenario['issue_price']
                        # 对数收益率（连续复利）
                        log_returns = np.log(final_prices / scenario_issue_price)
                        # 年化：120日 = 120个交易日
                        annualized_log_returns = log_returns * (252.0 / 120)

                        # 计算VaR（使用年化对数收益率）
                        var_5 = np.percentile(annualized_log_returns, 5)
                        var_95 = np.percentile(annualized_log_returns, 95)

                        # 包装成兼容旧代码的格式
                        scenario_result = {
                            'scenario': scenario,
                            'profit_prob': (log_returns > 0).mean() * 100,
                            'median_return': np.median(annualized_log_returns),
                            'mean_return': np.mean(annualized_log_returns),
                            'var_5': var_5,
                            'var_95': var_95,
                            'actual_premium_rate': scenario['premium_rate']
                        }

                        industry_pe_scenario_results.append(scenario_result)
                    except Exception as e:
                        print(f" 情景{scenario['name']}模拟失败: {e}")

                # 添加到汇总列表
                all_scenarios.extend(industry_pe_scenario_results)

                # 为附件展开情景对象
                for scenario_result in industry_pe_scenario_results:
                    scenario_flat = scenario_result['scenario'].copy()
                    scenario_flat.update({
                        'profit_prob': scenario_result['profit_prob'],
                        'median_return': scenario_result['median_return'],
                        'mean_return': scenario_result['mean_return'],
                        'var_5': scenario_result['var_5'],
                        'var_95': scenario_result['var_95'],
                        'issue_price': scenario_result['scenario']['issue_price']
                    })
                    historical_scenarios_for_appendix.append(scenario_flat)

                # 生成行业PE情景对比表格
                if industry_pe_scenario_results:
                    add_paragraph(document, '行业PE情景参数表（全部45个情景）：')
                    add_paragraph(document, '')

                    # 按盈利概率排序，展示全部情景
                    sorted_scenarios = sorted(industry_pe_scenario_results, key=lambda x: x['profit_prob'], reverse=True)

                    industry_pe_table_data = []
                    for scenario_result in sorted_scenarios:
                        scenario = scenario_result['scenario']
                        industry_pe_table_data.append([
                            scenario['name'],
                            f"{scenario['drift_level']}({scenario['drift']*100:+.1f}%)",
                            f"{scenario['vol_level']}({scenario['volatility']*100:.1f}%)",
                            scenario['premium_level'],
                            f"{scenario_result['profit_prob']:.1f}%",
                            f"{scenario_result['median_return']*100:+.1f}%"
                        ])

                    add_table_data(document, ['情景', 'PE分位-漂移率', '波动率', '溢价率', '盈利概率', '中位数收益'], industry_pe_table_data, font_size=10.5)

                    add_paragraph(document, '')
                    add_paragraph(document, ' 行业PE情景分析结论：')
                    best_scenario = max(industry_pe_scenario_results, key=lambda x: x['profit_prob'])
                    worst_scenario = min(industry_pe_scenario_results, key=lambda x: x['profit_prob'])
                    add_paragraph(document, f'• 最优情景：{best_scenario["scenario"]["name"]}，盈利概率{best_scenario["profit_prob"]:.1f}%，中位数收益{best_scenario["median_return"]*100:+.1f}%')
                    add_paragraph(document, f'• 最差情景：{worst_scenario["scenario"]["name"]}，盈利概率{worst_scenario["profit_prob"]:.1f}%，中位数收益{worst_scenario["median_return"]*100:+.1f}%')

                add_paragraph(document, '')
            except Exception as e:
                print(f" 行业PE数据计算失败: {e}")
                add_paragraph(document, ' 行业PE数据不可用，跳过本节分析')
        else:
            print(" 行业PE数据不可用，跳过6.3节")
            add_paragraph(document, ' 行业PE数据不可用，跳过本节分析')

        # ==================== 6.4 基于个股PE分位数的情景分析 ====================
        add_paragraph(document, '')
        add_title(document, '6.4 基于个股PE分位数的情景分析', level=2)

        add_paragraph(document, '本节基于个股PE分位数的估值水平，构建全面的情景矩阵进行分析。')
        add_paragraph(document, '通过组合不同PE分位数（高/中/低）、波动率和溢价率，生成3×3×5=45个情景，评估个股历史估值水平对定增收益的影响。')
        add_paragraph(document, '')

        add_paragraph(document, '个股PE分位数计算逻辑：')
        add_paragraph(document, '• 收集个股历史PE数据（通常为5年）')
        add_paragraph(document, '• 计算75%、50%、25%分位数，分别代表历史高估、中性、低估水平')
        add_paragraph(document, '• 漂移率 = (个股PE分位数 / 个股当前PE - 1)，反映估值回归潜力')
        add_paragraph(document, '• 如果当前PE低于历史分位数，估值便宜，未来收益率应该更高')
        add_paragraph(document, '')

        # 检查个股PE数据
        if stock_pe_data is not None:
            try:
                # 计算个股PE分位数
                stock_pe_values = stock_pe_data['pe_ttm'].dropna()
                stock_pe_75 = stock_pe_values.quantile(0.75)
                stock_pe_50 = stock_pe_values.quantile(0.50)
                stock_pe_25 = stock_pe_values.quantile(0.25)

                # 获取标的个股当前PE值
                stock_pe_current = stock_pe_data['pe_ttm'].iloc[-1]

                # 根据个股PE分位数与个股当前PE的比值计算漂移率
                stock_ratio_75 = stock_pe_75 / stock_pe_current if stock_pe_current > 0 else 1.0
                stock_ratio_50 = stock_pe_50 / stock_pe_current if stock_pe_current > 0 else 1.0
                stock_ratio_25 = stock_pe_25 / stock_pe_current if stock_pe_current > 0 else 1.0

                stock_drift_75 = stock_ratio_75 - 1.0
                stock_drift_50 = stock_ratio_50 - 1.0
                stock_drift_25 = stock_ratio_25 - 1.0

                # 限制漂移率在合理范围内（-50%到+100%）
                stock_drift_75 = max(-0.50, min(1.00, stock_drift_75))
                stock_drift_50 = max(-0.50, min(1.00, stock_drift_50))
                stock_drift_25 = max(-0.50, min(1.00, stock_drift_25))

                print(f" 个股PE分位数漂移率计算完成:")
                print(f"   个股当前PE: {stock_pe_current:.2f}倍")
                print(f"   个股PE 75%分位数: {stock_pe_75:.1f}倍, 比值={stock_ratio_75:.2f}, 漂移率={stock_drift_75*100:+.2f}%")
                print(f"   个股PE 50%分位数: {stock_pe_50:.1f}倍, 比值={stock_ratio_50:.2f}, 漂移率={stock_drift_50*100:+.2f}%")
                print(f"   个股PE 25%分位数: {stock_pe_25:.1f}倍, 比值={stock_ratio_25:.2f}, 漂移率={stock_drift_25*100:+.2f}%")

                # 使用市场指数的波动率（高、中、低三档）
                if index_vol_values:
                    index_vol_high = max(index_vol_values) * 1.1
                    index_vol_mid = np.mean(index_vol_values)
                    index_vol_low = min(index_vol_values) * 0.9
                else:
                    # 如果没有指数数据，使用市场数据
                    index_vol_high = market_data.get('volatility_120d', 0.35) * 1.1
                    index_vol_mid = market_data.get('volatility_120d', 0.35)
                    index_vol_low = market_data.get('volatility_120d', 0.35) * 0.9

                # 溢价率档次：-20%, -15%, -10%, -5%, 0%
                premium_levels = [-0.20, -0.15, -0.10, -0.05, 0.0]

                # 生成个股PE情景矩阵（3×3×5=45个情景）
                stock_pe_scenarios = generate_scenario_matrix(
                    drift_levels=[stock_drift_75, stock_drift_50, stock_drift_25],
                    vol_levels=[index_vol_high, index_vol_mid, index_vol_low],
                    premium_levels=premium_levels,
                    scenario_name_prefix='个股PE',
                    drift_source='基于个股PE历史分位数与当前PE的比值',
                    project_params=project_params
                )

                print(f" 生成个股PE情景矩阵: {len(stock_pe_scenarios)}个情景")

                # 运行个股PE情景模拟
                stock_pe_scenario_results = []
                for idx, scenario in enumerate(stock_pe_scenarios):
                    try:
                        sim_result = analyzer.monte_carlo_simulation(
                            n_simulations=5000,
                            time_steps=120,
                            volatility=scenario['volatility'],
                            drift=scenario['drift'],
                            seed=42+idx  # 使用不同的seed确保每个情景的随机性不同
                        )

                        final_prices = sim_result.iloc[:, -1].values
                        # 使用情景相关的发行价（考虑了溢价率）
                        scenario_issue_price = scenario['issue_price']
                        # 对数收益率（连续复利）
                        log_returns = np.log(final_prices / scenario_issue_price)
                        # 年化：120日 = 120个交易日
                        annualized_log_returns = log_returns * (252.0 / 120)

                        # 计算VaR（使用年化对数收益率）
                        var_5 = np.percentile(annualized_log_returns, 5)
                        var_95 = np.percentile(annualized_log_returns, 95)

                        # 包装成兼容旧代码的格式
                        scenario_result = {
                            'scenario': scenario,
                            'profit_prob': (log_returns > 0).mean() * 100,
                            'median_return': np.median(annualized_log_returns),
                            'mean_return': np.mean(annualized_log_returns),
                            'var_5': var_5,
                            'var_95': var_95,
                            'actual_premium_rate': scenario['premium_rate']
                        }

                        stock_pe_scenario_results.append(scenario_result)
                    except Exception as e:
                        print(f" 情景{scenario['name']}模拟失败: {e}")

                # 添加到汇总列表
                all_scenarios.extend(stock_pe_scenario_results)

                # 为附件展开情景对象
                for scenario_result in stock_pe_scenario_results:
                    scenario_flat = scenario_result['scenario'].copy()
                    scenario_flat.update({
                        'profit_prob': scenario_result['profit_prob'],
                        'median_return': scenario_result['median_return'],
                        'mean_return': scenario_result['mean_return'],
                        'var_5': scenario_result['var_5'],
                        'var_95': scenario_result['var_95'],
                        'issue_price': scenario_result['scenario']['issue_price']
                    })
                    historical_scenarios_for_appendix.append(scenario_flat)

                # 生成个股PE情景对比表格
                if stock_pe_scenario_results:
                    add_paragraph(document, '个股PE情景参数表（全部45个情景）：')
                    add_paragraph(document, '')

                    # 按盈利概率排序，展示全部情景
                    sorted_scenarios = sorted(stock_pe_scenario_results, key=lambda x: x['profit_prob'], reverse=True)

                    stock_pe_table_data = []
                    for scenario_result in sorted_scenarios:
                        scenario = scenario_result['scenario']
                        stock_pe_table_data.append([
                            scenario['name'],
                            f"{scenario['drift_level']}({scenario['drift']*100:+.1f}%)",
                            f"{scenario['vol_level']}({scenario['volatility']*100:.1f}%)",
                            scenario['premium_level'],
                            f"{scenario_result['profit_prob']:.1f}%",
                            f"{scenario_result['median_return']*100:+.1f}%"
                        ])

                    add_table_data(document, ['情景', 'PE分位-漂移率', '波动率', '溢价率', '盈利概率', '中位数收益'], stock_pe_table_data, font_size=10.5)

                    add_paragraph(document, '')
                    add_paragraph(document, ' 个股PE情景分析结论：')
                    best_scenario = max(stock_pe_scenario_results, key=lambda x: x['profit_prob'])
                    worst_scenario = min(stock_pe_scenario_results, key=lambda x: x['profit_prob'])
                    add_paragraph(document, f'• 最优情景：{best_scenario["scenario"]["name"]}，盈利概率{best_scenario["profit_prob"]:.1f}%，中位数收益{best_scenario["median_return"]*100:+.1f}%')
                    add_paragraph(document, f'• 最差情景：{worst_scenario["scenario"]["name"]}，盈利概率{worst_scenario["profit_prob"]:.1f}%，中位数收益{worst_scenario["median_return"]*100:+.1f}%')

                add_paragraph(document, '')
            except Exception as e:
                print(f" 个股PE数据计算失败: {e}")
                add_paragraph(document, ' 个股PE数据不可用，跳过本节分析')
        else:
            print(" 个股PE数据不可用，跳过6.4节")
            add_paragraph(document, ' 个股PE数据不可用，跳过本节分析')


        # ==================== 6.5 基于DCF估值的情景分析 ====================
        add_paragraph(document, '')
        add_title(document, '6.5 基于DCF估值的情景分析', level=2)

        add_paragraph(document, '本节基于DCF绝对估值方法，构建全面的情景矩阵进行分析。')
        add_paragraph(document, '通过组合不同波动率（高/中/低）和溢价率，生成1×3×5=15个情景，评估公司内在价值对定增收益的影响。')
        add_paragraph(document, '')

        add_paragraph(document, 'DCF估值方法说明：')
        add_paragraph(document, '• DCF（现金流折现模型）通过预测未来自由现金流并折现计算内在价值')
        add_paragraph(document, '• 漂移率 = (DCF内在价值 / 当前价格 - 1)，反映内在价值与市场价格的偏离')
        add_paragraph(document, '• 如果DCF内在价值高于当前价格，股票被低估，未来收益率应该为正')
        add_paragraph(document, '• DCF估值提供绝对价值判断，独立于市场相对估值')
        add_paragraph(document, '')

        # 检查DCF估值数据
        if intrinsic_value is not None and intrinsic_value > 0:
            try:
                current_price_dcf = project_params['current_price']
                dcf_ratio = intrinsic_value / current_price_dcf if current_price_dcf > 0 else 1.0
                dcf_drift = dcf_ratio - 1.0

                # 限制漂移率在合理范围内（-50%到+100%）
                dcf_drift = max(-0.50, min(1.00, dcf_drift))

                print(f" DCF估值漂移率计算完成:")
                print(f"   DCF内在价值: {intrinsic_value:.2f}元")
                print(f"   当前价格: {current_price_dcf:.2f}元")
                print(f"   比值: {dcf_ratio:.2f}")
                print(f"   漂移率: {dcf_drift*100:+.2f}%")

                # 使用市场指数的波动率（高、中、低三档）
                if index_vol_values:
                    index_vol_high = max(index_vol_values) * 1.1
                    index_vol_mid = np.mean(index_vol_values)
                    index_vol_low = min(index_vol_values) * 0.9
                else:
                    # 如果没有指数数据，使用市场数据
                    index_vol_high = market_data.get('volatility_120d', 0.35) * 1.1
                    index_vol_mid = market_data.get('volatility_120d', 0.35)
                    index_vol_low = market_data.get('volatility_120d', 0.35) * 0.9

                # 溢价率档次：-20%, -15%, -10%, -5%, 0%
                premium_levels = [-0.20, -0.15, -0.10, -0.05, 0.0]

                # 生成DCF估值情景矩阵（1×3×5=15个情景）
                # 注意：DCF只有1个漂移率（基于内在价值），所以是1×3×5
                dcf_scenarios = generate_scenario_matrix(
                    drift_levels=[dcf_drift],  # 只有1个漂移率
                    vol_levels=[index_vol_high, index_vol_mid, index_vol_low],
                    premium_levels=premium_levels,
                    scenario_name_prefix='DCF估值',
                    drift_source='基于DCF内在价值与当前价格的比值',
                    project_params=project_params
                )

                print(f" 生成DCF估值情景矩阵: {len(dcf_scenarios)}个情景")

                # 运行DCF估值情景模拟
                dcf_scenario_results = []
                for idx, scenario in enumerate(dcf_scenarios):
                    try:
                        sim_result = analyzer.monte_carlo_simulation(
                            n_simulations=5000,
                            time_steps=120,
                            volatility=scenario['volatility'],
                            drift=scenario['drift'],
                            seed=42+idx  # 使用不同的seed确保每个情景的随机性不同
                        )

                        final_prices = sim_result.iloc[:, -1].values
                        # 使用情景相关的发行价（考虑了溢价率）
                        scenario_issue_price = scenario['issue_price']
                        # 对数收益率（连续复利）
                        log_returns = np.log(final_prices / scenario_issue_price)
                        # 年化：120日 = 120个交易日
                        annualized_log_returns = log_returns * (252.0 / 120)

                        # 计算VaR（使用年化对数收益率）
                        var_5 = np.percentile(annualized_log_returns, 5)
                        var_95 = np.percentile(annualized_log_returns, 95)

                        # 包装成兼容旧代码的格式
                        scenario_result = {
                            'scenario': scenario,
                            'profit_prob': (log_returns > 0).mean() * 100,
                            'median_return': np.median(annualized_log_returns),
                            'mean_return': np.mean(annualized_log_returns),
                            'var_5': var_5,
                            'var_95': var_95,
                            'actual_premium_rate': scenario['premium_rate']
                        }

                        dcf_scenario_results.append(scenario_result)
                    except Exception as e:
                        print(f" 情景{scenario['name']}模拟失败: {e}")

                # 添加到汇总列表
                all_scenarios.extend(dcf_scenario_results)

                # 为附件展开情景对象
                for scenario_result in dcf_scenario_results:
                    scenario_flat = scenario_result['scenario'].copy()
                    scenario_flat.update({
                        'profit_prob': scenario_result['profit_prob'],
                        'median_return': scenario_result['median_return'],
                        'mean_return': scenario_result['mean_return'],
                        'var_5': scenario_result['var_5'],
                        'var_95': scenario_result['var_95'],
                        'issue_price': scenario_result['scenario']['issue_price']
                    })
                    historical_scenarios_for_appendix.append(scenario_flat)

                # 生成DCF估值情景对比表格
                if dcf_scenario_results:
                    add_paragraph(document, 'DCF估值情景参数表（全部15个情景）：')
                    add_paragraph(document, '')

                    # 按盈利概率排序，展示所有情景
                    sorted_scenarios = sorted(dcf_scenario_results, key=lambda x: x['profit_prob'], reverse=True)

                    dcf_table_data = []
                    for scenario_result in sorted_scenarios:
                        scenario = scenario_result['scenario']
                        dcf_table_data.append([
                            scenario['name'],
                            f"DCF({intrinsic_value:.2f}元)",
                            f"{scenario['vol_level']}({scenario['volatility']*100:.1f}%)",
                            scenario['premium_level'],
                            f"{scenario_result['profit_prob']:.1f}%",
                            f"{scenario_result['median_return']*100:+.1f}%"
                        ])

                    add_table_data(document, ['情景', 'DCF估值', '波动率', '溢价率', '盈利概率', '中位数收益'], dcf_table_data, font_size=10.5)

                    add_paragraph(document, '')
                    add_paragraph(document, ' DCF估值情景分析结论：')
                    best_scenario = max(dcf_scenario_results, key=lambda x: x['profit_prob'])
                    worst_scenario = min(dcf_scenario_results, key=lambda x: x['profit_prob'])
                    add_paragraph(document, f'• 最优情景：{best_scenario["scenario"]["name"]}，盈利概率{best_scenario["profit_prob"]:.1f}%，中位数收益{best_scenario["median_return"]*100:+.1f}%')
                    add_paragraph(document, f'• 最差情景：{worst_scenario["scenario"]["name"]}，盈利概率{worst_scenario["profit_prob"]:.1f}%，中位数收益{worst_scenario["median_return"]*100:+.1f}%')

                add_paragraph(document, '')
            except Exception as e:
                print(f" DCF估值计算失败: {e}")
                add_paragraph(document, ' DCF估值数据不可用，跳过本节分析')
        else:
            print(" DCF内在价值数据不可用，跳过6.5节")
            add_paragraph(document, ' DCF估值数据不可用，跳过本节分析')
        # ==================== 6.6 情景综合分析汇总表 ====================
        add_paragraph(document, '')
        add_title(document, '6.6 情景综合分析汇总表', level=2)

        add_paragraph(document, '本节对所有情景进行统计分析，汇总各类情景的数量分布和关键指标。')
        add_paragraph(document, '详细的情景数据请参见报告附件《情景数据表》。')
        add_paragraph(document, '')

        # 定义comprehensive_results（兼容旧代码）
        comprehensive_results = all_scenarios

        # 统计各类情景数量
        scenario_type_counts = {}
        for result in comprehensive_results:
            scenario_obj = result.get('scenario', result)
            scenario_name = scenario_obj.get('name', '未知')

            # 提取情景类型（市场指数、行业指数、行业PE、个股PE、DCF估值）
            if '市场指数' in scenario_name:
                scenario_type = '市场指数情景'
            elif '行业指数' in scenario_name:
                scenario_type = '行业指数情景'
            elif '行业PE' in scenario_name:
                scenario_type = '行业PE情景'
            elif '个股PE' in scenario_name:
                scenario_type = '个股PE情景'
            elif 'DCF估值' in scenario_name:
                scenario_type = 'DCF估值情景'
            else:
                scenario_type = '多参数情景'

            scenario_type_counts[scenario_type] = scenario_type_counts.get(scenario_type, 0) + 1

        # 计算整体统计指标
        all_profit_probs = [r.get('profit_prob', 0) for r in comprehensive_results]
        all_median_returns = [r.get('median_return', 0) for r in comprehensive_results]

        avg_profit_prob = np.mean(all_profit_probs)
        max_profit_prob = np.max(all_profit_probs)
        min_profit_prob = np.min(all_profit_probs)

        avg_median_return = np.mean(all_median_returns)
        max_median_return = np.max(all_median_returns)
        min_median_return = np.min(all_median_returns)

        # 找出最优和最差情景
        best_scenario = max(comprehensive_results, key=lambda x: x.get('profit_prob', 0))
        worst_scenario = min(comprehensive_results, key=lambda x: x.get('profit_prob', 0))

        # 获取情景名称，优先从scenario对象获取，其次从顶层获取，最后生成描述
        def get_scenario_name(scenario):
            if 'scenario' in scenario and isinstance(scenario['scenario'], dict):
                return scenario['scenario'].get('name', '未知情景')
            elif 'name' in scenario:
                return scenario['name']
            else:
                # 根据情景参数生成描述性名称
                drift = scenario.get('drift', 0)
                vol = scenario.get('volatility', 0)
                discount = scenario.get('discount', 0)
                return f'漂移率{drift*100:.0f}%_波动率{vol*100:.0f}%_折价{discount*100:.0f}%'

        best_scenario_name = get_scenario_name(best_scenario)
        worst_scenario_name = get_scenario_name(worst_scenario)

        # 生成汇总统计表格
        summary_data = [
            ['统计维度', '数值/说明', '备注'],
            ['情景总数', f'{len(comprehensive_results)}个', '包含所有类型的情景'],
            ['', '', ''],
            ['情景类型分布', '', ''],
        ]

        # 添加各类型情景数量
        for scenario_type, count in scenario_type_counts.items():
            summary_data.append([f'• {scenario_type}', f'{count}个', ''])

        # 添加整体统计
        summary_data.extend([
            ['', '', ''],
            ['盈利概率统计', '', ''],
            [f'• 平均盈利概率', f'{avg_profit_prob:.1f}%', '所有情景的平均值'],
            [f'• 最高盈利概率', f'{max_profit_prob:.1f}%', f'最优情景：{best_scenario_name}'],
            [f'• 最低盈利概率', f'{min_profit_prob:.1f}%', f'最差情景：{worst_scenario_name}'],
            ['', '', ''],
            ['收益率中位数统计', '', ''],
            [f'• 平均收益率', f'{avg_median_return*100:+.1f}%', '年化收益率'],
            [f'• 最高收益率', f'{max_median_return*100:+.1f}%', f'情景：{best_scenario_name}'],
            [f'• 最低收益率', f'{min_median_return*100:+.1f}%', f'情景：{worst_scenario_name}'],
        ])

        add_table_data(document, ['统计项目', '数值', '说明'], summary_data)

        # 添加说明（表格与正文之间不需要空行）
        add_paragraph(document, ' 情景说明：')
        add_paragraph(document, f'• 6.1多参数情景：{scenario_type_counts.get("多参数情景", 0)}个，涵盖漂移率、波动率、溢价率的全部组合')
        add_paragraph(document, f'• 6.2.1市场指数情景：{scenario_type_counts.get("市场指数情景", 0)}个（3×3×5矩阵）')
        add_paragraph(document, f'• 6.2.2行业指数情景：{scenario_type_counts.get("行业指数情景", 0)}个（3×3×5矩阵）')
        add_paragraph(document, f'• 6.3行业PE情景：{scenario_type_counts.get("行业PE情景", 0)}个（1×3×5矩阵）')
        add_paragraph(document, f'• 6.4个股PE情景：{scenario_type_counts.get("个股PE情景", 0)}个（1×3×5矩阵）')
        add_paragraph(document, f'• 6.5 DCF估值情景：{scenario_type_counts.get("DCF估值情景", 0)}个（1×3×5矩阵）')
        add_paragraph(document, '')

        add_paragraph(document, ' 评级说明：')
        add_paragraph(document, f'• 平均盈利概率：{avg_profit_prob:.1f}%，{"较高" if avg_profit_prob >= 70 else "中等" if avg_profit_prob >= 50 else "较低"}')
        add_paragraph(document, f'• 平均收益率：{avg_median_return*100:+.1f}%，{"表现良好" if avg_median_return >= 0.10 else "表现一般" if avg_median_return >= 0 else "表现不佳"}')
        add_paragraph(document, f'• 最优情景：{best_scenario_name}，盈利概率{best_scenario.get("profit_prob", 0):.1f}%，收益率{best_scenario.get("median_return", 0)*100:+.1f}%')
        add_paragraph(document, f'• 最差情景：{worst_scenario_name}，盈利概率{worst_scenario.get("profit_prob", 0):.1f}%，收益率{worst_scenario.get("median_return", 0)*100:+.1f}%')

    # 生成多维度情景图表（补充分析）
    # 注：波动率×折价率热力图已在6.3章节展示，此处不再重复
    try:
        multi_dim_chart_paths = generate_multi_dimension_scenario_charts_split(
            project_params['current_price'], ma20, risk_params['volatility'],
            risk_params['drift'], project_params['lockup_period'], IMAGES_DIR)

        # 该函数只返回1个图表（3D分析图）
        if len(multi_dim_chart_paths) > 0:
            add_paragraph(document, '图表 6.6: 多维度情景分析（波动率 × 时间窗口 × 折扣率）')
            add_image(document, multi_dim_chart_paths[0], width=Inches(6.5))
            add_paragraph(document, '')

    except Exception as e:
        print(f" 生成多维度图表失败: {e}")
        import traceback
        traceback.print_exc()

    add_section_break(document)

    # 分别保存两种情景数据供附件使用
    if 'multi_param_scenarios' in locals():
        context['results']['multi_param_scenarios_585'] = multi_param_scenarios
        print(f" 已保存6.1节{len(multi_param_scenarios)}种多参数构造情景到context")
    else:
        print(f" 警告：6.1节多参数构造情景未生成")

    if 'historical_scenarios_for_appendix' in locals() and len(historical_scenarios_for_appendix) > 0:
        context['results']['historical_scenarios_195'] = historical_scenarios_for_appendix
        print(f" 已保存6.2-6.5节{len(historical_scenarios_for_appendix)}种历史数据情景到context")
    else:
        print(f" 警告：6.2-6.5节历史数据情景未生成或为空")
        print(f"   可能原因：")
        print(f"   - 市场指数数据文件不可用")
        print(f"   - PE数据获取失败（已在第二章尝试获取）")
        print(f"   - DCF估值计算失败")
        print(f"   - 情景模拟过程中出现异常")

    return context
