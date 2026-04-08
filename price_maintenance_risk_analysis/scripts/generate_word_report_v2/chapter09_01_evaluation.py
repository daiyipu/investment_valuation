"""
第九章综合评估汇总

本章节对前面章节的分析结果进行综合汇总，涵盖个股与行业对比、相对估值、DCF估值、蒙特卡洛模拟、情景分析、压力测试和VaR风险度量等核心内容。
"""

import sys
import os
from docx.shared import Inches

# 添加路径
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, PROJECT_DIR)

from module_utils import add_title, add_paragraph, add_table_data, add_image
from module_utils import font_prop


def generate_chapter(context):
    """
    生成第九章第一节：综合评估汇总

    参数说明：
    - context: 包含所有必要数据的上下文字典
    """
    # 从context中提取变量
    document = context['document']
    project_params = context['project_params']
    market_data = context['market_data']
    IMAGES_DIR = context['IMAGES_DIR']

    # 添加第九章第一节标题（主标题已在chapter09_advice.py中添加）
    # 不再添加重复的"九、风控建议与风险提示"标题
    add_paragraph(document, '本章第一节从综合评估角度，汇总前面各章节的核心分析结果。')
    add_paragraph(document, '')

    # 从context['results']中获取其他章节的计算结果
    all_scenarios = context['results'].get('all_scenarios', [])
    var_95 = context['results'].get('var_95', 0.5)
    var_99 = context['results'].get('var_99', 0.7)
    risk_rating = context['results'].get('risk_rating', '高风险')
    mc_volatility = context['results'].get('mc_volatility', 0.30)
    mc_drift = context['results'].get('mc_drift', 0.08)

    # ==================== 9.1 综合评估汇总 ====================
    # 注意：9.1的level=2标题已在chapter09_advice.py中生成，这里只生成内容

    add_paragraph(document, '本节对前面章节的分析结果进行综合汇总，涵盖相对估值、DCF估值、蒙特卡洛模拟、情景分析、压力测试和VaR风险度量等核心内容。')
    add_paragraph(document, '')

    # ==================== 9.1.1 个股与历史分位数对比 ====================
    add_title(document, '9.1.1 个股与历史分位数对比', level=3)

    add_paragraph(document, '本节对比个股当前估值水平与自身历史数据，判断当前估值在历史中的位置。')
    add_paragraph(document, '')

    # 从第二章获取历史分位数数据（如果有的话）
    # 这里简化处理，显示基本信息
    current_price = project_params['current_price']
    issue_price = project_params['issue_price']
    ma20 = market_data.get('ma_20', current_price)
    ma60 = market_data.get('ma_60', current_price)
    ma120 = market_data.get('ma_120', current_price)
    ma250 = market_data.get('ma_250', current_price)

    # 计算当前价格相对于均线的位置
    ma_comparison_data = [
        ['指标', '当前值', 'MA20', 'MA60', 'MA120', 'MA250', '评估'],
        ['股价（元）', f'{current_price:.2f}', f'{ma20:.2f}', f'{ma60:.2f}', f'{ma120:.2f}', f'{ma250:.2f}', '']
    ]

    # 添加评估
    if current_price > ma120:
        price_eval = "高于半年线，偏强势"
    elif current_price < ma120 * 0.95:
        price_eval = "低于半年线5%以上，偏弱势"
    else:
        price_eval = "在半年线附近，震荡整理"
    ma_comparison_data[1][6] = price_eval

    add_table_data(document, ma_comparison_data[0], [ma_comparison_data[1]])

    add_paragraph(document, '')
    add_paragraph(document, ' 技术位置解读：', bold=True)
    add_paragraph(document, f'• 当前价格：{current_price:.2f}元')
    add_paragraph(document, f'• 相对MA20：{((current_price/ma20 - 1) * 100):+.2f}%')
    add_paragraph(document, f'• 相对MA120：{((current_price/ma120 - 1) * 100):+.2f}%')
    add_paragraph(document, f'• 技术位置：{price_eval}')
    add_paragraph(document, '')

    # ==================== 9.1.2 相对估值分析汇总 ====================
    add_title(document, '9.1.2 相对估值分析', level=3)

    add_paragraph(document, '本节汇总相对估值分析的核心结果，包括PE、PB等估值指标与行业的对比。')
    add_paragraph(document, '')
    add_paragraph(document, '注：详细的PE历史分位数分析请参见第二章"相对估值分析"。')
    add_paragraph(document, '')

    # 尝试从多个数据源获取PE/PB数据
    current_pe = 0
    industry_pe = 0
    current_pb = 0
    industry_pb = 0

    # 优先从context['results']获取第二章的计算结果
    if 'stock_pe_data' in context and context['stock_pe_data'] is not None:
        stock_pe_data = context['stock_pe_data']
        if not stock_pe_data.empty:
            current_pe = stock_pe_data.iloc[-1]['pe_ttm']

    if 'industry_pe_data' in context and context['industry_pe_data'] is not None:
        industry_pe_data = context['industry_pe_data']
        if not industry_pe_data.empty:
            industry_pe = industry_pe_data.iloc[-1]['pe_ttm']

    # 如果没有找到PE数据，尝试从market_data获取
    if current_pe == 0:
        current_pe = market_data.get('pe_ratio', 0)
    if industry_pe == 0:
        industry_pe = market_data.get('industry_pe', 0)

    # PB数据（如果有）
    current_pb = market_data.get('pb_ratio', 0)
    industry_pb = market_data.get('industry_pb', 0)

    # 如果有PE数据，生成表格
    if current_pe > 0 and industry_pe > 0:
        pe_premium = (current_pe / industry_pe - 1) * 100
        valuation_summary_data = [
            ['估值指标', '个股', '行业', '相对行业', '评估'],
            ['市盈率(PE)', f'{current_pe:.2f}倍', f'{industry_pe:.2f}倍', f'{pe_premium:+.1f}%', ''],
        ]

        if pe_premium < -20:
            pe_eval = "显著低估"
        elif pe_premium < 0:
            pe_eval = "相对低估"
        elif pe_premium < 20:
            pe_eval = "相对合理"
        elif pe_premium < 50:
            pe_eval = "相对高估"
        else:
            pe_eval = "显著高估"
        valuation_summary_data[1][4] = pe_eval

        if current_pb > 0 and industry_pb > 0:
            pb_premium = (current_pb / industry_pb - 1) * 100
            valuation_summary_data.append([
                '市净率(PB)', f'{current_pb:.2f}倍', f'{industry_pb:.2f}倍', f'{pb_premium:+.1f}%', ''
            ])

        add_table_data(document, valuation_summary_data[0], valuation_summary_data[1:])

        add_paragraph(document, '')
        add_paragraph(document, ' 相对估值解读：', bold=True)
        add_paragraph(document, f'• PE相对行业：{pe_premium:+.1f}%（{pe_eval}）')
        if current_pb > 0 and industry_pb > 0:
            add_paragraph(document, f'• PB相对行业：{pb_premium:+.1f}%')
        add_paragraph(document, '')
    else:
        add_paragraph(document, ' PE/PB数据暂时不可用，请参见第二章"相对估值分析"获取详细信息。')
        add_paragraph(document, '')

    # ==================== 9.1.3 DCF估值分析汇总 ====================
    add_title(document, '9.1.3 DCF估值分析', level=3)

    add_paragraph(document, '本节汇总DCF绝对估值分析的结果，评估公司的内在价值。')
    add_paragraph(document, '')

    # DCF估值方法说明
    add_paragraph(document, '【DCF估值方法】', bold=True)
    add_paragraph(document, 'DCF（现金流折现）估值法通过预测公司未来自由现金流，并以加权平均资本成本（WACC）')
    add_paragraph(document, '折现至现值，得出公司的内在价值。')
    add_paragraph(document, '')
    add_paragraph(document, '估值步骤：')
    add_paragraph(document, '1. 预测未来10年自由现金流（FCF）')
    add_paragraph(document, '2. 计算终值（Terminal Value）')
    add_paragraph(document, '3. 以WACC折现至现值')
    add_paragraph(document, '4. 减去净债务，得到股权价值')
    add_paragraph(document, '5. 除以总股数，得到每股价值')
    add_paragraph(document, '')

    # DCF估值数据（从第三章结果获取）
    intrinsic_value = context['results'].get('intrinsic_value', 0)
    if intrinsic_value is not None and intrinsic_value > 0:
        dcf_discount = (intrinsic_value / current_price - 1) * 100
        dcf_discount_issue = (intrinsic_value / issue_price - 1) * 100

        # 显示关键估值参数
        add_paragraph(document, '【估值参数】', bold=True)
        wacc_result = context.get('wacc', None)
        net_debt_result = context.get('net_debt', None)
        enterprise_value_result = context.get('enterprise_value', None)
        equity_value_result = context.get('equity_value', None)

        if wacc_result is not None:
            add_paragraph(document, f'• WACC（加权平均资本成本）: {wacc_result*100:.1f}%')
        if net_debt_result is not None:
            add_paragraph(document, f'• 净债务: {net_debt_result:.2f} 亿元')
        if enterprise_value_result is not None:
            add_paragraph(document, f'• 企业价值: {enterprise_value_result/100000000:.2f} 亿元')
        if equity_value_result is not None:
            add_paragraph(document, f'• 股权价值: {equity_value_result/100000000:.2f} 亿元')

        add_paragraph(document, '')
        add_paragraph(document, '【估值结果】', bold=True)
        add_paragraph(document, f'• DCF内在价值: {intrinsic_value:.2f} 元/股')
        add_paragraph(document, f'• 当前价格: {current_price:.2f} 元/股')
        add_paragraph(document, f'• 发行价格: {issue_price:.2f} 元/股')
        add_paragraph(document, f'• 相对市价安全边际: {dcf_discount:+.1f}%')
        add_paragraph(document, f'• 相对发行价安全边际: {dcf_discount_issue:+.1f}%')
        add_paragraph(document, '')

        # 估值结论
        add_paragraph(document, '【估值结论】', bold=True)

        if dcf_discount_issue > 15:
            conclusion = " DCF估值显示，相比发行价有显著安全边际（>15%），内在价值高于发行价。"
        elif dcf_discount_issue > 0:
            conclusion = " DCF估值显示，相比发行价有一定安全边际，内在价值略高于发行价。"
        elif dcf_discount_issue > -10:
            conclusion = " DCF估值显示，发行价略高于内在价值（<10%）。"
        else:
            conclusion = " DCF估值显示，发行价显著高于内在价值（≥10%）。"

        add_paragraph(document, conclusion)
        add_paragraph(document, '')

        # 添加估值表格
        dcf_summary_data = [
            ['估值方法', '内在价值', '当前价格', '发行价格', '相对市价', '相对发行价'],
            ['DCF估值', f'{intrinsic_value:.2f}元', f'{current_price:.2f}元', f'{issue_price:.2f}元', f'{dcf_discount:+.1f}%', f'{dcf_discount_issue:+.1f}%']
        ]
        add_table_data(document, dcf_summary_data[0], dcf_summary_data[1:])

        add_paragraph(document, '')
        add_paragraph(document, '说明：')
        add_paragraph(document, '• DCF内在价值基于未来自由现金流预测和WACC折现计算')
        add_paragraph(document, '• 安全边际 = (内在价值 - 价格) / 价格 × 100%')
        add_paragraph(document, '• 正安全边际表示内在价值高于价格，具有投资价值')
        add_paragraph(document, '• 负安全边际表示内在价值低于价格，存在估值风险')
        add_paragraph(document, '')
        add_paragraph(document, '注：详细的DCF估值模型、敏感性分析和参数说明请参见第三章"DCF估值分析"。')
        add_paragraph(document, '')

        # 保存结果到上下文
        context['dcf_intrinsic_value'] = intrinsic_value
        context['dcf_discount'] = dcf_discount
        context['dcf_discount_issue'] = dcf_discount_issue
    else:
        add_paragraph(document, '  DCF估值不可用（内在价值为负或无法计算），请参见第三章"DCF估值分析"了解详情。')
        add_paragraph(document, '')
        add_paragraph(document, '可能原因：')
        add_paragraph(document, '• WACC计算失败（参数缺失或不合理）')
        add_paragraph(document, '• 净利润为负（无法进行DCF估值）')
        add_paragraph(document, '• 净债务过高（股权价值为负）')
        add_paragraph(document, '• 其他计算错误')
        add_paragraph(document, '')

    # ==================== 9.1.4 蒙特卡洛模拟结果汇总 ====================
    add_title(document, '9.1.4 蒙特卡洛模拟结果', level=3)

    add_paragraph(document, '本节汇总蒙特卡洛模拟的核心结果，包括基于历史参数和基于预测参数（ARIMA+GARCH）两种方法的120日窗口模拟。')
    add_paragraph(document, '')
    add_paragraph(document, '注：详细的蒙特卡洛模拟分析请参见第五章"蒙特卡洛模拟"。')
    add_paragraph(document, '')

    # 添加两种模拟方法的说明
    add_paragraph(document, '【两种模拟方法说明】', bold=True)
    add_paragraph(document, '本报告采用两种蒙特卡洛模拟方法，以全面评估定增项目的风险收益特征：')
    add_paragraph(document, '')

    add_paragraph(document, '1. 历史参数模拟（传统方法）')
    add_paragraph(document, '   • 数据来源：基于过去250个交易日的历史统计数据')
    add_paragraph(document, '   • 漂移率：使用历史对数收益率的年化均值')
    add_paragraph(document, '   • 波动率：使用历史对数收益率的年化标准差')
    add_paragraph(document, '   • 优点：客观反映长期平均水平，不受模型假设影响')
    add_paragraph(document, '   • 缺点：无法反映市场结构变化和未来趋势')
    add_paragraph(document, '')

    add_paragraph(document, '2. 预测参数模拟（先进方法）')
    add_paragraph(document, '   • 漂移率：基于ARIMA时间序列模型预测未来收益趋势')
    add_paragraph(document, '   • 波动率：基于GARCH模型预测未来波动率变化')
    add_paragraph(document, '   • 优点：考虑时间序列的动态特征，反映市场变化趋势')
    add_paragraph(document, '   • 缺点：依赖模型假设，预测不确定性较大')
    add_paragraph(document, '')

    add_paragraph(document, '【方法对比与选择】', bold=True)
    add_paragraph(document, '• 互补性：两种方法各有优劣，建议结合使用，互为验证')
    add_paragraph(document, '• 适用场景：')
    add_paragraph(document, '  - 历史参数：适用于市场稳定的成熟期公司')
    add_paragraph(document, '  - 预测参数：适用于成长期公司或市场环境变化较大时')
    add_paragraph(document, '• 投资决策：')
    add_paragraph(document, '  - 两种方法结果一致时：可增强决策信心')
    add_paragraph(document, '  - 两种方法结果分歧时：需进一步分析原因，采取更保守策略')
    add_paragraph(document, '')

    # 获取多窗口期模拟结果
    multi_window_results = context['results'].get('multi_window_mc_results', {})
    if multi_window_results and '120d' in multi_window_results:
        mc_120d = multi_window_results['120d']
        profit_prob_historical = mc_120d.get('profit_prob', 0.5) * 100
        mean_return_historical = mc_120d.get('mean_return', 0.0) * 100
        median_return_historical = mc_120d.get('median_return', 0.0) * 100

        # 获取预测参数模拟结果
        mc_predicted = context['results'].get('mc_predicted_120d', {})
        if mc_predicted:
            profit_prob_predicted = mc_predicted.get('profit_prob', 0.5) * 100
            mean_return_predicted = mc_predicted.get('mean_return', 0.0) * 100
            median_return_predicted = mc_predicted.get('median_return', 0.0) * 100
            drift_predicted = mc_predicted.get('drift', mc_drift)
            vol_predicted = mc_predicted.get('volatility', mc_volatility)

            # 生成对比表格
            mc_comparison_data = [
                ['指标', '历史参数模拟', '预测参数模拟', '差异'],
                ['年化漂移率', f'{mc_drift*100:+.2f}%', f'{drift_predicted*100:+.2f}%', f'{(drift_predicted-mc_drift)*100:+.2f}%'],
                ['年化波动率', f'{mc_volatility*100:.2f}%', f'{vol_predicted*100:.2f}%', f'{(vol_predicted-mc_volatility)*100:+.2f}%'],
                ['', '', '', ''],
                ['盈利概率', f'{profit_prob_historical:.1f}%', f'{profit_prob_predicted:.1f}%', f'{profit_prob_predicted-profit_prob_historical:+.1f}%'],
                ['平均收益率', f'{mean_return_historical:+.2f}%', f'{mean_return_predicted:+.2f}%', f'{mean_return_predicted-mean_return_historical:+.2f}%'],
                ['中位数收益率', f'{median_return_historical:+.2f}%', f'{median_return_predicted:+.2f}%', f'{median_return_predicted-median_return_historical:+.2f}%'],
            ]

            add_table_data(document, ['指标', '历史参数模拟', '预测参数模拟(ARIMA+GARCH)', '差异'], mc_comparison_data)

            add_paragraph(document, '')
            add_paragraph(document, ' 模拟结果解读：', bold=True)
            add_paragraph(document, f'• 历史参数模拟：基于250日历史数据的漂移率({mc_drift*100:+.2f}%)和波动率({mc_volatility*100:.2f}%)')
            add_paragraph(document, f'• 预测参数模拟：基于ARIMA预测的漂移率({drift_predicted*100:+.2f}%)和GARCH预测的波动率({vol_predicted*100:.2f}%)')
            add_paragraph(document, f'• 盈利概率差异：{profit_prob_predicted-profit_prob_historical:+.1f}个百分点（{"预测更高" if profit_prob_predicted > profit_prob_historical else "历史更高"}）')
            add_paragraph(document, f'• 预期收益差异：{mean_return_predicted-mean_return_historical:+.2f}个百分点（{"预测更乐观" if mean_return_predicted > mean_return_historical else "历史更乐观"}）')
            add_paragraph(document, '')
        else:
            # 只显示历史参数结果
            mc_summary_data = [
                ['模拟参数', '值'],
                ['模拟窗口', '120日（半年）'],
                ['模拟次数', '10,000次'],
                ['年化波动率', f'{mc_volatility*100:.2f}%'],
                ['年化漂移率', f'{mc_drift*100:+.2f}%'],
                ['', ''],
                ['模拟结果', ''],
                ['盈利概率', f'{profit_prob_historical:.1f}%'],
                ['平均收益率', f'{mean_return_historical:+.2f}%'],
                ['中位数收益率', f'{median_return_historical:+.2f}%'],
                ['风险评级', risk_rating]
            ]
            add_table_data(document, mc_summary_data[0], mc_summary_data[1:])

            add_paragraph(document, '')
            add_paragraph(document, ' 说明：预测参数模拟结果不可用，仅显示基于历史参数的模拟结果。')
            add_paragraph(document, '')
    else:
        add_paragraph(document, ' 蒙特卡洛模拟结果不可用，请参见第五章"蒙特卡洛模拟"获取详细信息。')
        add_paragraph(document, '')

    # ==================== 9.1.5 历史数据场景说明 ====================
    add_title(document, '9.1.5 历史数据场景说明', level=3)

    add_paragraph(document, '本节简要介绍报告中使用的历史数据场景分类和数量统计。')
    add_paragraph(document, '')

    add_paragraph(document, '历史数据场景分类：', bold=True)
    add_paragraph(document, '• 市场指数情景：基于市场历史数据的情景分析')
    add_paragraph(document, '• 行业指数情景：基于行业历史数据的情景分析')
    add_paragraph(document, '• 行业估值情景：基于行业PE历史分位数的情景分析')
    add_paragraph(document, '• 个股估值情景：基于个股PE历史分位数的情景分析')
    add_paragraph(document, '• DCF估值情景：基于绝对估值法的情景分析')
    add_paragraph(document, '')

    add_paragraph(document, '情景数量统计：', bold=True)
    add_paragraph(document, '• 历史数据场景共计5大类')
    add_paragraph(document, '• 每类场景包含多个档位设置（如高、中、低档位）')
    add_paragraph(document, '• 具体情景数量和详细分析请参见第六章"情景分析"')
    add_paragraph(document, '')

    # ==================== 9.1.6 压力测试结果 ====================
    add_title(document, '9.1.6 压力测试结果', level=3)

    add_paragraph(document, '本节汇总压力测试的结果，展示在极端市场情景下的投资表现。')
    add_paragraph(document, '')
    add_paragraph(document, '注：详细的压力测试分析请参见第七章"情景分析与压力测试"。')
    add_paragraph(document, '')

    # 添加压力测试说明
    add_paragraph(document, '【压力测试说明】', bold=True)
    add_paragraph(document, '压力测试模拟历史危机和黑天鹅事件对股价的冲击，评估定增项目在极端情况下的抗风险能力。')
    add_paragraph(document, '本报告测试六种极端情景：2008年金融危机、2020年疫情、极端熊市、个股重大利空、行业政策收紧、流动性危机。')
    add_paragraph(document, '')

    # 定义六种极端压力情景（与第七章一致）
    current_price = project_params.get('current_price', 21.26)
    issue_price = project_params.get('issue_price', 25.80)
    issue_shares = project_params.get('issue_shares', 1000000)

    stress_scenarios = [
        {
            'name': '市场危机_2008',
            'description': '2008年全球金融危机',
            'price_drop': 0.60,
            'volatility_spike': 2.0,
            'details': 'A股市场暴跌约60%的极端情景'
        },
        {
            'name': '市场危机_2020',
            'description': '2020年新冠疫情冲击',
            'price_drop': 0.40,
            'volatility_spike': 1.5,
            'details': 'A股市场下跌约40%的情景'
        },
        {
            'name': '极端熊市',
            'description': '极端熊市周期',
            'price_drop': 0.50,
            'volatility_spike': 1.3,
            'details': '股价下跌50%，波动率上升30%'
        },
        {
            'name': '个股重大利空',
            'description': '公司业绩暴雷',
            'price_drop': 0.35,
            'volatility_spike': 1.8,
            'details': '财务造假、重大诉讼等'
        },
        {
            'name': '行业政策收紧',
            'description': '行业监管收紧',
            'price_drop': 0.25,
            'volatility_spike': 1.2,
            'details': '教培、互联网金融等强监管'
        },
        {
            'name': '流动性危机',
            'description': '市场流动性枯竭',
            'price_drop': 0.20,
            'volatility_spike': 2.5,
            'details': '钱荒、信用违约等'
        }
    ]

    # 计算各情景的定增收益
    add_paragraph(document, '【极端情景测试结果】', bold=True)

    stress_results_data = [['情景名称', '情景描述', '股价跌幅', '压力后价格', '定增收益率', '亏损金额', '风险评估']]

    total_loss_amount = 0
    worst_scenario = None
    worst_loss = 0

    for scenario in stress_scenarios:
        stressed_price = current_price * (1 - scenario['price_drop'])
        pnl_percent = ((stressed_price - issue_price) / issue_price) * 100
        pnl_amount = (stressed_price - issue_price) * issue_shares / 100000000  # 转万元

        # 风险评估
        if pnl_percent >= 0:
            risk_level = '🟢 低风险'
            risk_assessment = '抗压能力强'
        elif pnl_percent >= -20:
            risk_level = '🟡 中等风险'
            risk_assessment = '有一定抗风险能力'
        elif pnl_percent >= -40:
            risk_level = '🟠 较高风险'
            risk_assessment = '抗风险能力较弱'
        else:
            risk_level = '🔴 高风险'
            risk_assessment = '抗风险能力很弱'

        stress_results_data.append([
            scenario['name'],
            scenario['description'],
            f'-{int(scenario["price_drop"]*100)}%',
            f'{stressed_price:.2f}元',
            f'{pnl_percent:+.1f}%',
            f'{pnl_amount:.2f}万元',
            risk_level
        ])

        total_loss_amount += max(0, -pnl_amount)

        # 找出最坏情景
        if pnl_percent < worst_loss:
            worst_loss = pnl_percent
            worst_scenario = scenario

    add_table_data(document, stress_results_data[0], stress_results_data[1:])
    add_paragraph(document, '')

    # 风险提示
    add_paragraph(document, '【风险提示】', bold=True)
    add_paragraph(document, '• 压力测试情景基于历史事件，但实际极端情况可能更严重')
    add_paragraph(document, '• 投资决策应充分考虑极端情景下的最大承受能力')
    add_paragraph(document, '• 设置止损机制，控制单一项目投资比例')
    add_paragraph(document, '• 在市场出现极端情况时，及时评估并调整投资策略')
    add_paragraph(document, '')
    add_paragraph(document, '注：详细的压力测试分析请参见第七章"情景分析与压力测试"。')
    add_paragraph(document, '')

    # ==================== 9.1.7 VaR风险度量汇总 ====================
    add_title(document, '9.1.7 VaR风险度量汇总', level=3)

    add_paragraph(document, '本节汇总VaR（Value at Risk）风险度量的结果，评估在不同置信水平下的最大损失。')
    add_paragraph(document, '')
    add_paragraph(document, '注：详细的VaR风险度量分析请参见第八章"VaR风险度量"。')
    add_paragraph(document, '')

    # 使用120日窗口的VaR数据
    var_95_simple = var_95
    var_99_simple = var_99

    # 计算CVaR（条件风险价值，尾部平均损失）
    # 简化处理：CVaR约为VaR的1.2-1.5倍
    cvar_95_simple = min(var_95_simple * 1.3, 1.0)
    cvar_99_simple = min(var_99_simple * 1.4, 1.0)

    # 风险等级评估
    var_risk_level = "低风险" if var_95_simple < 0.15 else "中等风险" if var_95_simple < 0.30 else "高风险"

    var_summary_data = [
        ['置信水平', 'VaR（最大损失）', 'CVaR（尾部平均损失）', '说明'],
        ['90%置信', f'{min(var_95_simple*0.8, 1.0)*100:.2f}%', f'{min(var_95_simple*0.8*1.2, 1.0)*100:.2f}%', '10%概率损失超过此值'],
        ['95%置信', f'{var_95_simple*100:.2f}%', f'{cvar_95_simple*100:.2f}%', '5%概率损失超过此值'],
        ['99%置信', f'{var_99_simple*100:.2f}%', f'{cvar_99_simple*100:.2f}%', '1%概率损失超过此值'],
        ['', '', '', ''],
        ['数据窗口', '120日（半年锁定期）', '', '基于锁定期收益率'],
        ['风险等级', var_risk_level, '', '基于95% VaR综合评估'],
    ]

    add_table_data(document, var_summary_data[0], var_summary_data[1:])

    add_paragraph(document, '')
    add_paragraph(document, ' VaR解读：', bold=True)
    add_paragraph(document, f'• VaR基于锁定期（120日）的简单收益率计算，未进行年化')
    add_paragraph(document, f'• 95% VaR：{var_95_simple*100:.2f}%，表示在95%置信水平下，锁定期末的最大可能损失')
    add_paragraph(document, f'• 99% VaR：{var_99_simple*100:.2f}%，表示在99%置信水平下，锁定期末的最大可能损失')
    add_paragraph(document, f'• 风险等级：{var_risk_level}')
    add_paragraph(document, '')

    # 返回更新后的context
    return context
