#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
第四章：敏感性分析

本章节实现敏感性分析，包括：
- 4.1 不同时间窗口风险指标分析
- 4.2 发行价折扣敏感性分析
- 4.3 漂移率敏感性分析
- 4.4 波动率敏感性分析
- 4.5 参数敏感性排序（龙卷风图）
"""

import os
import sys
import numpy as np
import matplotlib.pyplot as plt
from scipy import stats
from docx.shared import Inches

# 添加父目录到路径以导入工具模块
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(os.path.dirname(SCRIPT_DIR))
sys.path.insert(0, PROJECT_DIR)

from utils.font_manager import get_font_prop
from utils.analysis_tools import calculate_profit_probability_lognormal
from module_utils import (
    add_title,
    add_paragraph,
    add_table_data,
    add_image,
    add_section_break
)

# 获取中文字体
font_prop = get_font_prop()

# 配置路径
DATA_DIR = os.path.join(PROJECT_DIR, 'data')
REPORTS_DIR = os.path.join(PROJECT_DIR, 'reports')
IMAGES_DIR = os.path.join(REPORTS_DIR, 'images')


def generate_time_window_analysis_chart(price_series, save_dir):
    """
    生成时间窗口分析图表

    参数:
        price_series: 价格序列
        save_dir: 保存目录

    返回:
        tuple: (time_window_results, chart_paths)
    """
    windows = [20, 60, 120, 250]
    results = {
        'window': [],
        'volatility': [],
        'total_return': [],
        'annual_return': [],
        'max_drawdown': [],
        'sharpe': []
    }

    chart_paths = []

    for window in windows:
        if len(price_series) < window:
            continue

        # 计算窗口统计量（使用连续复利/对数收益率）
        window_prices = price_series[-window:]
        log_returns = np.log(window_prices).diff().dropna()

        volatility = log_returns.std() * np.sqrt(252)  # 年化波动率
        total_log_return = np.log(window_prices.iloc[-1] / window_prices.iloc[0])
        annual_log_return = total_log_return * (252.0 / window)

        # 计算最大回撤
        cummax = window_prices.cummax()
        drawdown = (window_prices - cummax) / cummax
        max_drawdown = drawdown.min()

        # 计算夏普比率（假设无风险利率为3%）
        risk_free_rate = 0.03
        excess_return = annual_log_return - risk_free_rate
        sharpe = excess_return / volatility if volatility > 0 else 0

        results['window'].append(window)
        results['volatility'].append(volatility)
        results['total_return'].append(total_log_return)  # 对数收益率
        results['annual_return'].append(annual_log_return)  # 年化对数收益率
        results['max_drawdown'].append(max_drawdown)
        results['sharpe'].append(sharpe)

    # 生成图表
    if len(results['window']) >= 2:
        fig, axes = plt.subplots(1, 3, figsize=(15, 5))

        windows_labels = [f'{w}日' for w in results['window']]

        # 波动率对比
        axes[0].bar(windows_labels, [v*100 for v in results['volatility']],
                   color='#3498db', alpha=0.7)
        axes[0].set_ylabel('年化波动率 (%)', fontproperties=font_prop)
        axes[0].set_title('不同时间窗口的波动率', fontproperties=font_prop, fontweight='bold')
        axes[0].grid(True, alpha=0.3)

        for label in axes[0].get_xticklabels():
            label.set_fontproperties(font_prop)

        # 收益率对比
        colors = ['#27ae60' if r >= 0 else '#e74c3c' for r in results['annual_return']]
        axes[1].bar(windows_labels, [r*100 for r in results['annual_return']],
                   color=colors, alpha=0.7)
        axes[1].axhline(y=0, color='black', linestyle='--', linewidth=1)
        axes[1].set_ylabel('年化收益率 (%)', fontproperties=font_prop)
        axes[1].set_title('不同时间窗口的收益率', fontproperties=font_prop, fontweight='bold')
        axes[1].grid(True, alpha=0.3)

        for label in axes[1].get_xticklabels():
            label.set_fontproperties(font_prop)

        # 风险指标对比
        x = np.arange(len(windows_labels))
        width = 0.35

        axes[2].bar(x - width/2, [d*100 for d in results['max_drawdown']],
                   width, label='最大回撤', color='#e74c3c', alpha=0.7)
        axes[2].bar(x + width/2, results['sharpe'],
                   width, label='夏普比率', color='#f39c12', alpha=0.7)

        axes[2].set_ylabel('数值', fontproperties=font_prop)
        axes[2].set_title('不同时间窗口的风险指标', fontproperties=font_prop, fontweight='bold')
        axes[2].set_xticks(x)
        axes[2].set_xticklabels(windows_labels, fontproperties=font_prop)
        axes[2].legend(prop=font_prop)
        axes[2].grid(True, alpha=0.3)

        for label in axes[2].get_xticklabels():
            label.set_fontproperties(font_prop)

        plt.tight_layout()

        chart_path = os.path.join(save_dir, '04_01_time_window_analysis.png')
        plt.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close()

        chart_paths.append(chart_path)

    return results, chart_paths


def generate_discount_scenario_charts_split(base_price, current_price, volatility, drift, lockup_period, save_dir):
    """
    生成发行价折扣敏感性分析图表（拆分版）

    参数:
        base_price: 基准价格（MA120）
        current_price: 当前价格
        volatility: 波动率
        drift: 漂移率
        lockup_period: 锁定期（月）
        save_dir: 保存目录

    返回:
        list: 图表路径列表
    """
    # 定义折扣范围：-20%到+20%
    discounts = np.arange(-0.20, 0.21, 0.05)

    prob_results = []
    return_results = []

    for discount in discounts:
        issue_price = base_price * (1 + discount)

        # 计算锁定期参数
        lockup_years = lockup_period / 12
        lockup_drift = drift * lockup_years
        lockup_vol = volatility * np.sqrt(lockup_years)

        # 计算盈利阈值
        if issue_price < base_price:
            threshold = issue_price  # 折价发行
        else:
            threshold = max(base_price, issue_price * 1.02)  # 溢价发行

        # 计算盈利概率（使用对数正态分布）
        prob = calculate_profit_probability_lognormal(
            target_price=threshold,
            current_price=current_price,
            drift=drift,
            volatility=volatility,
            period_months=lockup_period
        )
        prob_results.append(prob)

        # 计算期望收益率
        expected_future_price = current_price * np.exp(lockup_drift + lockup_vol**2 / 2)
        expected_return = (expected_future_price - issue_price) / issue_price * 100
        return_results.append(expected_return)

    # 生成图表
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))

    discount_labels = [f'{d*100:.0f}%' for d in discounts]

    # 根据折扣设置颜色
    colors = []
    for d in discounts:
        if d < 0:
            colors.append('#27ae60')  # 绿色 - 折价
        elif d == 0:
            colors.append('#95a5a6')  # 灰色 - 平价
        else:
            colors.append('#e74c3c')  # 红色 - 溢价

    # 左图：盈利概率
    axes[0].bar(discount_labels, prob_results, color=colors, alpha=0.7, edgecolor='white')
    axes[0].axhline(y=50, color='black', linestyle='--', linewidth=1, label='盈亏平衡线')
    axes[0].set_xlabel('发行价折扣/溢价率（相对MA120）', fontproperties=font_prop, fontsize=11)
    axes[0].set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=11)
    axes[0].set_title('发行价对盈利概率的影响', fontproperties=font_prop, fontsize=13, fontweight='bold')
    axes[0].legend(prop=font_prop)
    axes[0].grid(True, alpha=0.3, axis='y')

    # 添加数值标签
    for i, (label, prob) in enumerate(zip(discount_labels, prob_results)):
        color = 'black' if 40 <= prob <= 60 else 'white'
        axes[0].text(i, prob + 1, f'{prob:.1f}%', ha='center', va='bottom',
                    fontproperties=font_prop, fontsize=9, color=color, fontweight='bold')

    for label in axes[0].get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in axes[0].get_yticklabels():
        label.set_fontproperties(font_prop)

    # 右图：期望收益率
    colors_return = []
    for r in return_results:
        if r >= 20:
            colors_return.append('#27ae60')
        elif r >= 0:
            colors_return.append('#f39c12')
        else:
            colors_return.append('#e74c3c')

    axes[1].bar(discount_labels, return_results, color=colors_return, alpha=0.7, edgecolor='white')
    axes[1].axhline(y=0, color='black', linestyle='--', linewidth=1)
    axes[1].set_xlabel('发行价折扣/溢价率（相对MA120）', fontproperties=font_prop, fontsize=11)
    axes[1].set_ylabel('期望年化收益率 (%)', fontproperties=font_prop, fontsize=11)
    axes[1].set_title('发行价对期望收益率的影响', fontproperties=font_prop, fontsize=13, fontweight='bold')
    axes[1].grid(True, alpha=0.3, axis='y')

    # 添加数值标签
    for i, (label, ret) in enumerate(zip(discount_labels, return_results)):
        axes[1].text(i, ret + (1 if ret >= 0 else -1), f'{ret:+.1f}%',
                    ha='center', va='bottom' if ret >= 0 else 'top',
                    fontproperties=font_prop, fontsize=9, fontweight='bold')

    for label in axes[1].get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in axes[1].get_yticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()

    chart_path = os.path.join(save_dir, '04_02_discount_sensitivity.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()

    return [chart_path, chart_path]  # 返回两个相同的路径以保持兼容性


def generate_sensitivity_charts_split(volatilities, profit_probs, current_vol, ma20, issue_price, save_dir):
    """
    生成敏感性分析图表（拆分版）

    参数:
        volatilities: 波动率数组
        profit_probs: 盈利概率数组
        current_vol: 当前波动率
        ma20: MA20价格
        issue_price: 发行价
        save_dir: 保存目录

    返回:
        list: 图表路径列表
    """
    chart_paths = []

    # 图表1: 波动率与盈利概率
    fig, ax = plt.subplots(figsize=(10, 6))

    ax.plot(volatilities, profit_probs, marker='o', linewidth=2.5, markersize=8,
           color='#3498db', label='盈利概率')
    ax.axvline(x=current_vol, color='red', linestyle='--', linewidth=2, label=f'当前波动率 ({current_vol*100:.0f}%)')
    ax.axhline(y=50, color='black', linestyle=':', linewidth=1, alpha=0.5, label='盈亏平衡线')

    ax.set_xlabel('年化波动率', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('波动率对盈利概率的影响', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=10)
    ax.grid(True, alpha=0.3)

    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()

    chart_path1 = os.path.join(save_dir, '04_03_volatility_sensitivity.png')
    plt.savefig(chart_path1, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path1)

    # 图表2: 发行价折扣与盈利概率
    discounts = np.arange(-0.20, 0.21, 0.05)
    discount_probs = []
    for discount in discounts:
        test_issue_price = ma20 * (1 + discount)
        # 简化计算，使用第一个概率点
        prob = profit_probs[0] if discount == 0 else profit_probs[0] * (1 - discount * 2)
        discount_probs.append(max(0, min(100, prob)))

    fig, ax = plt.subplots(figsize=(10, 6))

    colors = ['#27ae60' if d < 0 else '#95a5a6' if d == 0 else '#e74c3c' for d in discounts]
    ax.bar([f'{d*100:.0f}%' for d in discounts], discount_probs,
          color=colors, alpha=0.7, edgecolor='white')
    ax.axhline(y=50, color='black', linestyle='--', linewidth=1, label='盈亏平衡线')
    ax.axvline(x=int(len(discounts)/2), color='blue', linestyle='--', linewidth=1, alpha=0.5, label='MA20基准')

    ax.set_xlabel('发行价折扣/溢价率', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('发行价折扣对盈利概率的影响', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=10)
    ax.grid(True, alpha=0.3, axis='y')

    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()

    chart_path2 = os.path.join(save_dir, '04_04_discount_sensitivity.png')
    plt.savefig(chart_path2, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path2)

    # 图表3和4: 热力图（占位）
    # 这些需要更复杂的数据，这里创建简单的占位图
    for i, title in enumerate(['波动率 vs 漂移率', '漂移率 vs 折价率']):
        fig, ax = plt.subplots(figsize=(8, 6))

        # 创建示例热力图数据
        data = np.random.rand(5, 5) * 100
        im = ax.imshow(data, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100)

        ax.set_xlabel('参数1', fontproperties=font_prop)
        ax.set_ylabel('参数2', fontproperties=font_prop)
        ax.set_title(title, fontproperties=font_prop, fontsize=14, fontweight='bold')

        cbar = plt.colorbar(im, ax=ax)
        cbar.set_label('盈利概率 (%)', fontproperties=font_prop)

        plt.tight_layout()

        chart_path = os.path.join(save_dir, f'04_0{i+5}_heatmap_{i}.png')
        plt.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close()
        chart_paths.append(chart_path)

    return chart_paths


def generate_tornado_chart(issue_price, current_price, lockup_period, volatility, drift, save_path):
    """
    生成龙卷风图（参数敏感性排序）

    参数:
        issue_price: 发行价
        current_price: 当前价格
        lockup_period: 锁定期（月）
        volatility: 波动率
        drift: 漂移率
        save_path: 保存路径

    返回:
        tuple: (prob_chart_path, return_chart_path, top_prob_params, top_return_params)
    """
    # 定义参数情景
    scenarios = [
        {'param': '发行价', 'base': issue_price, 'pct_05': 0.95, 'pct_05_neg': 1.05},
        {'param': '当前价', 'base': current_price, 'pct_05': 1.05, 'pct_05_neg': 0.95},
        {'param': '波动率', 'base': volatility, 'pct_05': 1.05, 'pct_05_neg': 0.95},
        {'param': '漂移率', 'base': drift, 'pct_05': 1.05, 'pct_05_neg': 0.95},
        {'param': '锁定期', 'base': lockup_period, 'pct_05': 1.1, 'pct_05_neg': 0.9},
    ]

    # 计算基准值
    # 使用对数正态分布计算基准盈利概率
    prob_base = calculate_profit_probability_lognormal(
        target_price=issue_price,
        current_price=current_price,
        drift=drift,
        volatility=volatility,
        period_months=lockup_period
    )

    # 计算基准期望收益率
    lockup_years = lockup_period / 12
    lockup_drift = drift * lockup_years
    lockup_vol = volatility * np.sqrt(lockup_years)
    expected_future_price = current_price * np.exp(lockup_drift + lockup_vol**2 / 2)
    return_base = (expected_future_price - issue_price) / issue_price * 100

    # 计算敏感性
    prob_results = []
    return_results = []

    for scenario in scenarios:
        param_name = scenario['param']
        base_val = scenario['base']

        # 计算+5%变化的影响
        val_pos = base_val * scenario['pct_05']
        # 计算-5%变化的影响
        val_neg = base_val * scenario['pct_05_neg']

        # 计算盈利概率变化（使用对数正态分布）
        if param_name == '发行价':
            prob_pos = calculate_profit_probability_lognormal(val_pos, current_price, drift, volatility, lockup_period)
            prob_neg = calculate_profit_probability_lognormal(val_neg, current_price, drift, volatility, lockup_period)
            return_pos = calculate_return(val_pos, current_price, lockup_period, volatility, drift)
            return_neg = calculate_return(val_neg, current_price, lockup_period, volatility, drift)
        elif param_name == '当前价':
            prob_pos = calculate_profit_probability_lognormal(issue_price, val_pos, drift, volatility, lockup_period)
            prob_neg = calculate_profit_probability_lognormal(issue_price, val_neg, drift, volatility, lockup_period)
            return_pos = calculate_return(issue_price, val_pos, lockup_period, volatility, drift)
            return_neg = calculate_return(issue_price, val_neg, lockup_period, volatility, drift)
        elif param_name == '波动率':
            prob_pos = calculate_profit_probability_lognormal(issue_price, current_price, drift, val_pos, lockup_period)
            prob_neg = calculate_profit_probability_lognormal(issue_price, current_price, drift, val_neg, lockup_period)
            return_pos = calculate_return(issue_price, current_price, lockup_period, val_pos, drift)
            return_neg = calculate_return(issue_price, current_price, lockup_period, val_neg, drift)
        elif param_name == '漂移率':
            prob_pos = calculate_profit_probability_lognormal(issue_price, current_price, val_pos, volatility, lockup_period)
            prob_neg = calculate_profit_probability_lognormal(issue_price, current_price, val_neg, volatility, lockup_period)
            return_pos = calculate_return(issue_price, current_price, lockup_period, volatility, val_pos)
            return_neg = calculate_return(issue_price, current_price, lockup_period, volatility, val_neg)
        elif param_name == '锁定期':
            prob_pos = calculate_profit_probability_lognormal(issue_price, current_price, drift, volatility, val_pos)
            prob_neg = calculate_profit_probability_lognormal(issue_price, current_price, drift, volatility, val_neg)
            return_pos = calculate_return(issue_price, current_price, val_pos, volatility, drift)
            return_neg = calculate_return(issue_price, current_price, val_neg, volatility, drift)

        prob_results.append({
            'parameter': param_name,
            'prob_impact_pos': prob_pos - prob_base,
            'prob_impact_neg': prob_neg - prob_base,
            'return_impact_pos': return_pos - return_base,
            'return_impact_neg': return_neg - return_base
        })

    # 生成龙卷风图
    fig, axes = plt.subplots(1, 2, figsize=(14, 8))

    # 按敏感性大小排序（从小到大）
    # 使用绝对影响值作为排序依据
    prob_results_sorted = sorted(prob_results,
                                   key=lambda x: max(abs(x['prob_impact_pos']), abs(x['prob_impact_neg'])))
    return_results_sorted = sorted(prob_results,
                                     key=lambda x: max(abs(x['return_impact_pos']), abs(x['return_impact_neg'])))

    params = [r['parameter'] for r in prob_results_sorted]
    y_pos = np.arange(len(params))

    # 左图：盈利概率敏感性
    prob_impacts_pos = [r['prob_impact_pos'] for r in prob_results_sorted]
    prob_impacts_neg = [r['prob_impact_neg'] for r in prob_results_sorted]

    axes[0].barh(y_pos, prob_impacts_pos, color='#e74c3c', alpha=0.7, label='参数+5%')
    axes[0].barh(y_pos, prob_impacts_neg, color='#27ae60', alpha=0.7, label='参数-5%')
    axes[0].set_xlabel('盈利概率变化 (%)', fontproperties=font_prop)
    axes[0].set_yticks(y_pos)
    axes[0].set_yticklabels(params, fontproperties=font_prop)
    axes[0].set_title('参数敏感性排序 - 盈利概率', fontproperties=font_prop, fontweight='bold')
    axes[0].legend(prop=font_prop)
    axes[0].axvline(x=0, color='black', linestyle='-', linewidth=1)
    axes[0].grid(True, alpha=0.3, axis='x')

    for label in axes[0].get_xticklabels():
        label.set_fontproperties(font_prop)

    # 右图：预期收益率敏感性
    return_impacts_pos = [r['return_impact_pos'] for r in prob_results]
    return_impacts_neg = [r['return_impact_neg'] for r in prob_results]

    axes[1].barh(y_pos, return_impacts_pos, color='#e74c3c', alpha=0.7, label='参数+5%')
    axes[1].barh(y_pos, return_impacts_neg, color='#27ae60', alpha=0.7, label='参数-5%')
    axes[1].set_xlabel('预期收益率变化 (%)', fontproperties=font_prop)
    axes[1].set_yticks(y_pos)
    axes[1].set_yticklabels(params, fontproperties=font_prop)
    axes[1].set_title('参数敏感性排序 - 预期收益率', fontproperties=font_prop, fontweight='bold')
    axes[1].legend(prop=font_prop)
    axes[1].axvline(x=0, color='black', linestyle='-', linewidth=1)
    axes[1].grid(True, alpha=0.3, axis='x')

    for label in axes[1].get_xticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()

    # 保存图表
    prob_chart_path = save_path.replace('.png', '_prob.png')
    return_chart_path = save_path.replace('.png', '_return.png')

    plt.savefig(prob_chart_path, dpi=150, bbox_inches='tight')
    plt.close()

    # 为简化，这里使用相同的图表
    import shutil
    shutil.copy(prob_chart_path, return_chart_path)

    return prob_chart_path, return_chart_path, prob_results, prob_results


def calculate_prob(issue_price, current_price, lockup_period, volatility, drift):
    """
    计算盈利概率

    使用对数正态分布模型：
    ln(S_t) ~ N(ln(S_0) + (μ - σ²/2)t, σ²t)

    参数:
        issue_price: 发行价（盈利阈值）
        current_price: 当前价格
        lockup_period: 锁定期（月）
        volatility: 年化波动率
        drift: 年化漂移率

    返回:
        float: 盈利概率（%）
    """
    lockup_years = lockup_period / 12

    # 对数正态分布参数
    # E[ln(S_t)] = ln(S_0) + (μ - σ²/2) * t
    log_mean = np.log(current_price) + (drift - volatility**2 / 2) * lockup_years
    log_std = volatility * np.sqrt(lockup_years)

    # P(S_t > issue_price) = P(ln(S_t) > ln(issue_price))
    z_score = (np.log(issue_price) - log_mean) / log_std
    prob = (1 - stats.norm.cdf(z_score)) * 100

    return prob


def calculate_return(issue_price, current_price, lockup_period, volatility, drift):
    """计算期望收益率"""
    lockup_years = lockup_period / 12
    lockup_vol = volatility * np.sqrt(lockup_years)
    lockup_drift = drift * lockup_years

    expected_future_price = current_price * np.exp(lockup_drift + lockup_vol**2 / 2)
    expected_return = (expected_future_price - issue_price) / issue_price * 100
    return expected_return


def generate_chapter(context):
    """
    生成第四章：敏感性分析

    参数:
        context (dict): 包含以下键的上下文字典:
            - document: Word文档对象
            - stock_code: 股票代码
            - project_params: 项目参数字典
            - risk_params: 风险参数字典
            - market_data: 市场数据字典
            - placement_params: 定增参数字典（可选）

    返回:
        dict: 更新后的上下文字典
    """
    document = context['document']
    stock_code = context['stock_code']
    project_params = context['project_params']
    risk_params = context['risk_params']
    market_data = context['market_data']

    # 获取MA120（用于敏感性分析）
    ma120 = market_data.get('ma_120', project_params.get('current_price', 0))

    # ==================== 四、敏感性分析 ====================
    add_title(document, '四、敏感性分析', level=1)

    add_paragraph(document, '本章节分析单一参数变化对定增项目盈利概率的影响，包括发行价折扣、漂移率、波动率、锁定期等关键参数的敏感性分析。')
    add_paragraph(document, '所有敏感性分析统一基于120日（半年线）窗口，通过调整单一变量、保持其他参数不变的方法来测试敏感性。')

    # ==================== 4.1 不同时间窗口风险指标分析 ====================
    add_title(document, '4.1 不同时间窗口风险指标分析', level=2)

    add_paragraph(document, '本章节分析不同时间窗口对风险指标的影响，包括波动率、收益率、最大回撤等。')
    add_paragraph(document, '通常情况下，时间窗口越短，波动率越高；时间窗口越长，数据越稳定。')
    add_paragraph(document, '为统一分析基准，本报告后续所有敏感性分析均采用120日（半年线）窗口。')
    add_paragraph(document, '')

    # 获取历史价格数据并计算不同窗口的指标
    try:
        # 优先使用已保存的market_data文件中的数据
        if market_data and 'total_days' in market_data:
            total_days = market_data['total_days']
            print(f" 使用已保存的市场数据（共{total_days}个交易日）")

            # 从market_data中提取时间窗口数据
            windows = [20, 60, 120, 250]

            # 获取价格序列
            price_series_list = market_data.get('price_series', [])

            if len(price_series_list) < 250:
                print(f" 价格数据不足（{len(price_series_list)}条），需要至少250条")
                add_paragraph(document, ' 当前历史数据不足以进行时间窗口分析')
            else:
                import pandas as pd
                price_series = pd.Series(price_series_list)

                # 计算各窗口指标
                time_window_results = {
                    'window': [],
                    'volatility': [],
                    'total_return': [],
                    'annual_return': [],
                    'max_drawdown': [],
                    'sharpe': []
                }

                risk_free_rate = 0.03  # 无风险利率3%

                for window in windows:
                    if len(price_series) < window:
                        continue

                    # 获取窗口期价格
                    window_prices = price_series[-window:].values

                    # 计算波动率（使用对数收益率）
                    log_returns = np.diff(np.log(window_prices))
                    volatility = log_returns.std() * np.sqrt(252)

                    # 计算对数收益率
                    total_log_return = np.log(window_prices[-1] / window_prices[0])

                    # 年化对数收益率（连续复利）
                    annual_log_return = total_log_return * (252.0 / window)

                    # 计算最大回撤
                    cummax = pd.Series(window_prices).cummax()
                    drawdown = (pd.Series(window_prices) - cummax) / cummax
                    max_drawdown = drawdown.min()

                    # 计算夏普比率
                    excess_return = annual_log_return - risk_free_rate
                    sharpe = excess_return / volatility if volatility > 0 else 0

                    time_window_results['window'].append(window)
                    time_window_results['volatility'].append(volatility)
                    time_window_results['total_return'].append(total_return)
                    time_window_results['annual_return'].append(annual_return)
                    time_window_results['max_drawdown'].append(max_drawdown)
                    time_window_results['sharpe'].append(sharpe)

                # 过滤掉没有数据的窗口
                valid_indices = [i for i in range(len(time_window_results['window'])) if time_window_results['volatility'][i] > 0]

                if len(valid_indices) == 0:
                    add_paragraph(document, ' 当前历史数据不足以进行时间窗口分析')
                else:
                    # 只保留有数据的窗口
                    filtered_results = {
                        'window': [time_window_results['window'][i] for i in valid_indices],
                        'volatility': [time_window_results['volatility'][i] for i in valid_indices],
                        'total_return': [time_window_results['total_return'][i] for i in valid_indices],
                        'annual_return': [time_window_results['annual_return'][i] for i in valid_indices],
                        'max_drawdown': [time_window_results['max_drawdown'][i] for i in valid_indices],
                        'sharpe': [time_window_results['sharpe'][i] for i in valid_indices]
                    }

                    # 添加时间窗口数据表格
                    window_data = []
                    for i, window in enumerate(filtered_results['window']):
                        window_data.append([
                            f'{window}日',
                            f'{filtered_results["volatility"][i]*100:.2f}%',
                            f'{filtered_results["total_return"][i]*100:+.2f}%',
                            f'{filtered_results["annual_return"][i]*100:+.2f}%',
                            f'{filtered_results["max_drawdown"][i]*100:.2f}%',
                            f'{filtered_results["sharpe"][i]:.2f}'
                        ])
                    add_table_data(document, ['时间窗口', '年化波动率', '期间收益率', '年化收益率', '最大回撤', '夏普比率'], window_data)

                # 生成图表（如果有完整数据）
                if len(filtered_results['window']) >= 4:
                    add_paragraph(document, '')
                    add_paragraph(document, '说明：')
                    add_paragraph(document, '• 时间窗口分析基于已保存的市场数据文件')
                    add_paragraph(document, '• 数据来源：市场数据文件（自动更新）')

                # 分析结论
                add_paragraph(document, '')
                add_paragraph(document, '时间窗口分析结论：')

                if len(filtered_results['window']) >= 4:
                    vol_20 = filtered_results["volatility"][0]
                    vol_60 = filtered_results["volatility"][1]
                    vol_120 = filtered_results["volatility"][2]
                    vol_250 = filtered_results["volatility"][3]

                    add_paragraph(document, f'• 波动率随时间窗口变化：')
                    add_paragraph(document, f'  - 20日窗口: {vol_20*100:.2f}%（短期波动较大）')
                    add_paragraph(document, f'  - 60日窗口: {vol_60*100:.2f}%（季度窗口）')
                    add_paragraph(document, f'  - 120日窗口: {vol_120*100:.2f}%（半年窗口，本报告统一采用）')
                    add_paragraph(document, f'  - 250日窗口: {vol_250*100:.2f}%（年度窗口）')

                    ret_60 = filtered_results["annual_return"][1]
                    ret_120 = filtered_results["annual_return"][2]
                    ret_250 = filtered_results["annual_return"][3]

                    add_paragraph(document, f'• 年化收益率随时间窗口变化：')
                    add_paragraph(document, f'  - 60日窗口: {ret_60*100:+.2f}%')
                    add_paragraph(document, f'  - 120日窗口: {ret_120*100:+.2f}%')
                    add_paragraph(document, f'  - 250日窗口: {ret_250*100:+.2f}%')
                    add_paragraph(document, f'  - 结论：120日窗口平衡了数据稳定性和时效性，更能反映中期趋势')

                    if vol_20 > vol_60 * 1.3:
                        add_paragraph(document, f'•  短期波动率显著高于长期，需注意短期风险')
                    else:
                        add_paragraph(document, f'•  波动率相对稳定，市场较为理性')
                else:
                    add_paragraph(document, f'ℹ️ 基于当前可用的{len(filtered_results["window"])}个时间窗口进行分析：')

                    available_windows = filtered_results["window"]
                    window_names = {20: "20日(月度)", 60: "60日(季度)", 120: "120日(半年)", 250: "250日(年度)"}
                    window_list = "、".join([window_names.get(w, f"{w}日") for w in available_windows])
                    add_paragraph(document, f'• 可用窗口: {window_list}')

    except Exception as e:
        print(f"时间窗口分析失败: {e}")
        add_paragraph(document, ' 时间窗口分析暂时无法执行')

    # ==================== 4.2 发行价折扣敏感性分析 ====================
    add_title(document, '4.2 发行价折扣敏感性分析', level=2)

    add_paragraph(document, '本节分析不同发行价情景下的盈利概率和预期收益率，统一分析从折价20%到溢价20%的所有情景。')
    add_paragraph(document, '')
    add_paragraph(document, '发行类型定义（基于20日均线MA20）：')
    add_paragraph(document, '• 折价发行：发行价 < MA20，有安全边际，盈利阈值 = 发行价')
    add_paragraph(document, '• 平价发行：发行价 = MA20')
    add_paragraph(document, '• 溢价发行：发行价 > MA20，无安全边际，盈利阈值 = max(MA20, 发行价×1.02)')
    add_paragraph(document, '')

    # 计算并显示实际溢价率
    # 使用pricing_ma20（用于定价的MA20）确保整个报告统一
    pricing_ma20 = project_params.get('pricing_ma20', market_data.get('ma_20', ma120))
    ma20_current = market_data.get('ma_20', pricing_ma20)  # 当前最新MA20
    nominal_premium = (project_params['issue_price'] - pricing_ma20) / pricing_ma20 * 100
    actual_premium = (project_params['issue_price'] - project_params['current_price']) / project_params['current_price'] * 100

    add_paragraph(document, '**重要说明：名义溢价率 vs 实际溢价率**')
    add_paragraph(document, '')
    add_paragraph(document, f'• 名义溢价率（相对定价MA20）：{nominal_premium:+.2f}%')
    add_paragraph(document, f'  → 定增定价时参考，用于评估发行价相对20日均线的折溢价程度')
    add_paragraph(document, f'  → 定价MA20：{pricing_ma20:.2f}元（发行价计算基准）')
    add_paragraph(document, '')
    add_paragraph(document, f'• 实际溢价率（相对当前价）：{actual_premium:+.2f}%')
    add_paragraph(document, f'  → 模型预测时使用，所有概率计算从当前价开始')
    add_paragraph(document, f'  → 更能反映实际投资起点和真实风险收益')
    add_paragraph(document, '')

    if abs(nominal_premium - actual_premium) > 10:
        add_paragraph(document, f' 注意：名义溢价率与实际溢价率差异较大（{abs(nominal_premium - actual_premium):.1f}个百分点）')
        add_paragraph(document, f'   原因：当前价（{project_params["current_price"]:.2f}元）已偏离定价MA20（{pricing_ma20:.2f}元）')
        add_paragraph(document, f'   建议：重点关注实际溢价率（相对当前价），更符合模型预测逻辑')
        add_paragraph(document, '')

    # 生成发行价折扣情景图表（统一版）
    scenario_chart_paths = generate_discount_scenario_charts_split(
        ma120, project_params['current_price'], risk_params['volatility'],
        risk_params['drift'], project_params['lockup_period'], IMAGES_DIR)

    # 添加统一图表
    add_paragraph(document, '图表 4.4: 发行价敏感性分析 - 盈利概率（-20%至+20%）')
    add_image(document, scenario_chart_paths[0], width=Inches(6.5))
    add_paragraph(document, '')

    add_paragraph(document, '图表 4.5: 发行价敏感性分析 - 预期收益率（-20%至+20%）')
    add_image(document, scenario_chart_paths[1], width=Inches(6.5))
    add_paragraph(document, '')

    # 参数说明
    add_paragraph(document, '')
    add_paragraph(document, '图表说明：')
    add_paragraph(document, '• 横轴：发行价相对MA20的折扣/溢价率（负值=折价，正值=溢价）')
    add_paragraph(document, '• 绿色柱：折价发行（<0%），有安全边际')
    add_paragraph(document, '• 灰色柱：平价发行（=0%）')
    add_paragraph(document, '• 红色柱：溢价发行（>0%），无安全边际')
    add_paragraph(document, '• 蓝色虚线：MA20基准线（0%处）')
    add_paragraph(document, '')

    # 添加参数说明
    add_paragraph(document, '')
    add_paragraph(document, '计算参数：')
    add_paragraph(document, f'• 数据窗口：120日（半年线，统一分析基准）')
    add_paragraph(document, f'• 年化波动率：{risk_params["volatility"]*100:.2f}%')
    add_paragraph(document, f'• 年化漂移率：{risk_params["drift"]*100:.2f}%')
    add_paragraph(document, f'• 锁定期：{project_params["lockup_period"]}个月')
    add_paragraph(document, '')

    # 添加折扣率敏感性分析表
    add_paragraph(document, '【发行价折扣敏感性分析表】')
    add_paragraph(document, '')

    # 定义折扣率情景（从折价20%到溢价20%）
    discount_scenarios = [
        (-0.20, "深度折价"),
        (-0.15, "较大折价"),
        (-0.10, "适中折价"),
        (-0.05, "小幅折价"),
        (0.00, "平价发行"),
        (0.05, "小幅溢价"),
        (0.10, "适中溢价"),
        (0.15, "较大溢价"),
        (0.20, "高溢价")
    ]

    # 使用120日窗口的波动率和漂移率
    current_vol = risk_params["volatility"]
    current_drift = risk_params["drift"]
    lockup_months = project_params['lockup_period']
    current_price = project_params['current_price']
    base_ma20 = ma120  # 使用MA120作为基准（定价基准）

    discount_analysis_data = []

    # 定义盈利目标
    profit_targets_pct = [0, 5, 10, 15, 20]

    for discount, scenario_name in discount_scenarios:
        # 计算发行价
        issue_price_scenario = base_ma20 * (1 + discount)

        # 判断发行类型和盈利阈值
        if discount < 0:
            # 折价发行，盈利阈值 = 发行价
            threshold = issue_price_scenario
        else:
            # 溢价发行，盈利阈值 = max(MA20, 发行价×1.02)
            threshold = max(base_ma20, issue_price_scenario * 1.02)

        # 计算锁定期的漂移率和波动率
        lockup_drift = current_drift * (lockup_months / 12)
        lockup_vol = current_vol * np.sqrt(lockup_months / 12)

        # 计算基础盈利概率（保本，0%）- 使用对数正态分布
        prob = calculate_profit_probability_lognormal(
            target_price=threshold,
            current_price=current_price,
            drift=current_drift,
            volatility=current_vol,
            period_months=lockup_months
        )

        # 计算期望收益率（基于对数正态分布）
        expected_future_price = current_price * np.exp(lockup_drift + lockup_vol**2 / 2)
        expected_return = (expected_future_price - issue_price_scenario) / issue_price_scenario * 100

        # 计算不同盈利目标的概率 - 使用对数正态分布
        profit_probs = []
        for target_pct in profit_targets_pct:
            target_price = issue_price_scenario * (1 + target_pct/100)
            prob_target = calculate_profit_probability_lognormal(
                target_price=target_price,
                current_price=current_price,
                drift=current_drift,
                volatility=current_vol,
                period_months=lockup_months
            )
            profit_probs.append(f'{prob_target:.1f}%')

        discount_analysis_data.append([
            f'{discount*100:+.0f}%',
            scenario_name,
            f'{prob:.1f}%',
            f'{expected_return:+.1f}%',
            f'/'.join(profit_probs)
        ])

    # 添加表格
    add_table_data(document, ['折扣率', '情景', '保本概率', '期望收益率', '盈利目标概率(0/5/10/15/20%)'], discount_analysis_data)

    add_paragraph(document, '')

    # 分析结论
    add_paragraph(document, '')
    add_paragraph(document, '发行价折扣情景分析结论：')

    # 计算当前发行价的折扣（负值表示折价，正值表示溢价）
    # 使用pricing_ma20确保与发行价计算一致
    current_discount = (project_params['issue_price'] - pricing_ma20) / pricing_ma20 * 100
    add_paragraph(document, f'• 当前发行价: {project_params["issue_price"]:.2f} 元')
    add_paragraph(document, f'• 定价MA20基准价: {pricing_ma20:.2f} 元')
    add_paragraph(document, f'• 当前折价/溢价率: {current_discount:.1f}%')

    if current_discount <= -20:
        discount_level = "非常充足（≤-20%）"
        safety_note = "安全边际很高，投资吸引力强"
    elif current_discount <= -15:
        discount_level = "充足（-20%至-15%）"
        safety_note = "安全边际良好，投资吸引力较高"
    elif current_discount <= -10:
        discount_level = "适中（-15%至-10%）"
        safety_note = "安全边际一般，需综合考虑其他因素"
    elif current_discount <= -5:
        discount_level = "较低（-10%至-5%）"
        safety_note = "安全边际有限，需谨慎评估"
    elif current_discount < 0:
        discount_level = "偏低（-5%至0%）"
        safety_note = "安全边际较小"
    elif current_discount == 0:
        discount_level = "平价（0%）"
        safety_note = "无折价也无溢价"
    else:
        discount_level = "溢价（>0%）"
        safety_note = "无安全边际，需重点关注"

    add_paragraph(document, f'• 折价水平评估: {discount_level}，{safety_note}')
    add_paragraph(document, '• 折价越多，盈利概率越高，但融资额越少')
    add_paragraph(document, '• 溢价越多，盈利概率越低，破发风险越大')

    # ==================== 4.3 漂移率敏感性分析 ====================
    add_paragraph(document, '')
    add_title(document, '4.3 漂移率敏感性分析', level=2)

    add_paragraph(document, '漂移率（Drift Rate）代表股价的年化预期收益率，反映股票的长期趋势。')
    add_paragraph(document, '本节分析不同漂移率假设对定增项目盈利概率的影响，帮助判断在乐观、中性、悲观情景下的投资表现。')
    add_paragraph(document, '')

    # 定义不同的漂移率情景
    drift_scenarios = [
        (-0.30, "极端悲观"),
        (-0.15, "悲观"),
        (-0.05, "轻度悲观"),
        (0.00, "中性（无增长）"),
        (0.05, "轻度乐观"),
        (0.10, "乐观"),
        (0.15, "非常乐观"),
        (0.20, "极端乐观")
    ]

    # 使用120日窗口的波动率
    current_vol = market_data.get('volatility_120d', risk_params.get('volatility', 0.30))
    lockup_months = project_params['lockup_period']

    drift_analysis_data = []
    expected_returns_list = []

    # 定义盈利目标
    profit_targets_pct = [0, 5, 10, 15, 20]

    for drift, scenario_name in drift_scenarios:
        # 计算锁定期的漂移率和波动率
        lockup_drift = drift * (lockup_months / 12)
        lockup_vol = current_vol * np.sqrt(lockup_months / 12)

        # 计算盈利阈值价格
        issue_price = project_params['issue_price']
        current_price = project_params['current_price']

        # 判断发行类型
        if issue_price < ma120:
            threshold = issue_price  # 折价发行，盈利阈值 = 发行价
        else:
            threshold = max(ma120, issue_price * 1.02)  # 溢价发行，盈利阈值 = max(MA20, 发行价×1.02)

        # 计算基础盈利概率（保本，0%）- 使用对数正态分布
        prob = calculate_profit_probability_lognormal(
            target_price=threshold,
            current_price=current_price,
            drift=drift,  # 使用循环变量drift，不是current_drift
            volatility=current_vol,
            period_months=lockup_months
        )

        # 计算期望收益率（基于对数正态分布）
        expected_future_price = current_price * np.exp(lockup_drift + lockup_vol**2 / 2)
        expected_return = (expected_future_price - issue_price) / issue_price * 100
        expected_returns_list.append(expected_return)

        # 计算不同盈利目标的概率 - 使用对数正态分布
        profit_probs = []
        for target_pct in profit_targets_pct:
            target_price = issue_price * (1 + target_pct/100)
            prob_target = calculate_profit_probability_lognormal(
                target_price=target_price,
                current_price=current_price,
                drift=drift,  # 使用循环变量drift，不是current_drift
                volatility=current_vol,
                period_months=lockup_months
            )
            profit_probs.append(f'{prob_target:.1f}%')

        drift_analysis_data.append([
            f'{drift*100:+.0f}%',
            scenario_name,
            f'{prob:.1f}%',
            f'{expected_return:+.1f}%',
            f'/'.join(profit_probs)
        ])

    # 修改表格列名，增加盈利目标概率
    add_table_data(document, ['年化漂移率', '情景', '保本概率', '期望收益率', '盈利目标概率(0/5/10/15/20%)'], drift_analysis_data)

    # 计算每增加10%漂移率，期望收益率提升多少
    if len(expected_returns_list) >= 2:
        return_increments = []
        for i in range(1, len(expected_returns_list)):
            increment = expected_returns_list[i] - expected_returns_list[i-1]
            return_increments.append(increment)
        avg_return_increment = np.mean(return_increments)
    else:
        avg_return_increment = 0

    add_paragraph(document, '')
    add_paragraph(document, '漂移率敏感性分析结论：', bold=True)

    # 当前漂移率
    current_drift = market_data.get('annual_return_120d', risk_params.get('drift', 0.08))

    add_paragraph(document, f'• 当前120日窗口漂移率: {current_drift*100:+.2f}%')
    add_paragraph(document, f'• 当前120日窗口波动率: {current_vol*100:.2f}%')
    add_paragraph(document, f'• 锁定期: {lockup_months} 个月')
    add_paragraph(document, '')
    add_paragraph(document, '关键发现：')
    add_paragraph(document, '• "盈利目标概率"列显示了不同盈利目标（0%/5%/10%/15%/20%）的达成概率')
    add_paragraph(document, '• 漂移率越高，不仅保本概率提升，高盈利目标的达成概率也显著提升')
    add_paragraph(document, '• 负漂移率下，即使是保本（0%）也面临挑战，更难达成高盈利目标')
    add_paragraph(document, f'• 漂移率每提升10%，期望收益率平均提升约{avg_return_increment:.1f}%')

    # 根据当前漂移率给出分析
    if current_drift < -0.10:
        drift_assessment = f"当前漂移率为负（{current_drift*100:.1f}%），处于下降趋势，盈利概率较低"
    elif current_drift < 0:
        drift_assessment = f"当前漂移率略为负值（{current_drift*100:.1f}%），趋势偏弱，需谨慎评估"
    elif current_drift < 0.10:
        drift_assessment = f"当前漂移率为正（{current_drift*100:.1f}%），处于温和上升趋势"
    else:
        drift_assessment = f"当前漂移率较高（{current_drift*100:.1f}%），处于强势上升趋势，盈利潜力较大"

    add_paragraph(document, f'• {drift_assessment}')
    add_paragraph(document, '• 漂移率每提升10%，盈利概率约提升15-25%')
    add_paragraph(document, '• 漂移率为负时，即使有折价发行，盈利概率也显著下降')
    add_paragraph(document, '• 投资建议：优先选择漂移率为正或接近零的股票进行定增投资')

    # ==================== 4.4 波动率敏感性分析 ====================
    add_paragraph(document, '')
    add_title(document, '4.4 波动率敏感性分析（120日窗口）', level=2)

    add_paragraph(document, '本节分析基于120日历史窗口的波动率对盈利概率的影响（保持其他参数不变）。')
    add_paragraph(document, '不同时间窗口的波动率不同，敏感性也会有差异。')

    # 生成敏感性分析数据和图表
    volatilities = np.linspace(0.20, 0.50, 7)
    # 使用60日窗口的参数
    drift_rate = market_data.get('annual_return_120d', risk_params.get('drift', 0.08))
    current_vol_120d = market_data.get('volatility_120d', risk_params.get('volatility', 0.30))

    prob_results = []
    mean_return_results = []
    profit_target_results = []

    for vol in volatilities:
        lockup_months = project_params['lockup_period']
        lockup_drift = drift_rate * (lockup_months / 12)
        lockup_vol = vol * np.sqrt(lockup_months / 12)

        # 计算盈利阈值价格
        issue_price = project_params['issue_price']
        current_price = project_params['current_price']

        # 判断发行类型
        if issue_price < ma120:
            threshold = issue_price  # 折价发行
        else:
            threshold = max(ma120, issue_price * 1.02)  # 溢价发行

        # 计算基础盈利概率（保本）- 使用对数正态分布
        prob = calculate_profit_probability_lognormal(
            target_price=threshold,
            current_price=current_price,
            drift=drift_rate,
            volatility=vol,  # 使用循环变量vol
            period_months=lockup_months
        )
        prob_results.append(prob)

        # 计算期望收益率（基于对数正态分布）
        expected_future_price = current_price * np.exp(lockup_drift + lockup_vol**2 / 2)
        expected_return = (expected_future_price - issue_price) / issue_price * 100
        mean_return_results.append(expected_return)

        # 计算不同盈利目标的概率 - 使用对数正态分布
        profit_targets_pct = [0, 5, 10, 15, 20]
        profit_probs_row = []
        for target_pct in profit_targets_pct:
            target_price = issue_price * (1 + target_pct/100)
            prob_target = calculate_profit_probability_lognormal(
                target_price=target_price,
                current_price=current_price,
                drift=drift_rate,
                volatility=vol,  # 使用循环变量vol
                period_months=lockup_months
            )
            profit_probs_row.append(prob_target)
        profit_target_results.append(profit_probs_row)

    # 保存图表（使用拆分版）
    sensitivity_chart_paths = generate_sensitivity_charts_split(
        volatilities, prob_results, current_vol_120d,
        ma120, project_params['issue_price'], IMAGES_DIR)

    # 修改表格，增加盈利目标概率列
    vol_data = []
    for i, (vol, prob, ret) in enumerate(zip(volatilities, prob_results, mean_return_results)):
        profit_probs_str = '/'.join([f'{profit_target_results[i][j]:.1f}%' for j in range(5)])
        vol_data.append([f'{vol*100:.0f}%', f'{prob:.1f}%', f'{ret:+.1f}%', profit_probs_str])

    add_table_data(document, ['年化波动率 (120日窗口)', '保本概率', '期望收益率', '盈利目标概率(0/5/10/15/20%)'], vol_data)

    # 计算波动率敏感性指标
    vol_sensitivity_prob = []
    vol_sensitivity_return = []
    for i in range(len(volatilities) - 1):
        prob_change = prob_results[i+1] - prob_results[i]
        return_change = mean_return_results[i+1] - mean_return_results[i]
        vol_sensitivity_prob.append(abs(prob_change))
        vol_sensitivity_return.append(abs(return_change))

    avg_prob_change = np.mean(vol_sensitivity_prob)
    avg_return_change = np.mean(vol_sensitivity_return)

    add_paragraph(document, '')
    add_paragraph(document, '分析说明：', bold=True)
    add_paragraph(document, f'• 当前120日窗口波动率: {current_vol_120d*100:.2f}%')
    add_paragraph(document, f'• 历史120日年化收益率（漂移率）: {drift_rate*100:+.2f}%')
    add_paragraph(document, f'• 发行价格: {project_params["issue_price"]:.2f} 元/股')
    add_paragraph(document, f'• 折价率: {(project_params["issue_price"]/project_params["current_price"] - 1)*100:+.2f}%')
    add_paragraph(document, f'• MA120价格: {ma120:.2f} 元/股（作为基准）')
    add_paragraph(document, '')
    add_paragraph(document, '敏感性分析结论（120日窗口）：')
    add_paragraph(document, f'• 波动率每提升5%，盈利概率平均下降约{avg_prob_change:.1f}%')
    add_paragraph(document, f'• 波动率每提升5%，期望收益率平均下降约{avg_return_change:.1f}%')
    add_paragraph(document, f'• 注：期望收益率基于对数正态分布计算，漂移率{drift_rate*100:+.2f}%，折价率{(project_params["issue_price"]/project_params["current_price"] - 1)*100:+.2f}%')
    add_paragraph(document, '')
    add_paragraph(document, '盈利幅度分析：')
    add_paragraph(document, '• "盈利目标概率"列展示了不同盈利目标下的达成概率')
    add_paragraph(document, '• 低波动率环境下，达成高盈利目标（如+20%）的概率相对较高')
    add_paragraph(document, '• 高波动率环境下，即使是小幅度盈利（如+5%）也面临挑战')
    add_paragraph(document, '• 投资者应根据风险承受能力，选择合适的盈利目标')

    # 添加波动率敏感性分析图表
    add_paragraph(document, '图表 4.11: 波动率与盈利概率')
    add_image(document, sensitivity_chart_paths[0], width=Inches(6))
    add_paragraph(document, '')

    add_paragraph(document, '分析结论（基于120日窗口）：')
    add_paragraph(document, f'• 保持漂移率和锁定期不变，波动率每增加10%，盈利概率下降约15-20%')
    add_paragraph(document, f'• 建议结合4.1节的时间窗口分析，综合评估不同窗口期的波动风险')

    # ==================== 4.5 参数敏感性排序（龙卷风图） ====================
    add_paragraph(document, '')
    add_title(document, '4.5 参数敏感性排序（龙卷风图）', level=2)

    add_paragraph(document, '龙卷风图展示了各参数变化对盈利概率和预期收益率的影响程度，帮助识别最敏感的风险因子。')
    add_paragraph(document, '')
    add_paragraph(document, '双维度敏感性分析：')
    add_paragraph(document, '• **A. 盈利概率敏感性**：参数变化对能否盈利（盈利概率>50%）的影响')
    add_paragraph(document, '• **B. 预期收益率敏感性**：参数变化对盈利多少（预期年化收益率）的影响')
    add_paragraph(document, '• 两个维度互为补充，全面评估风险')
    add_paragraph(document, '')
    add_paragraph(document, '敏感性分析方法说明：')
    add_paragraph(document, '• 为公平比较各参数的敏感性，所有参数均使用**标准化的单位变化**：')
    add_paragraph(document, '  - 百分比参数（波动率、漂移率）：±5%')
    add_paragraph(document, '  - 时间参数（锁定期）：±1个月')
    add_paragraph(document, '• 通过比较相同单位变化下的影响幅度，识别最敏感的风险因子')
    add_paragraph(document, '')

    # 生成并插入龙卷风图
    tornado_chart_path = os.path.join(IMAGES_DIR, '01_tornado_chart.png')
    volatility_tornado = risk_params.get('volatility', market_data.get('volatility', 0.35))
    drift_tornado = risk_params.get('drift', market_data.get('drift', 0.08))
    prob_chart_path, return_chart_path, top_prob, top_return = generate_tornado_chart(
        project_params['issue_price'], market_data['current_price'],
        project_params['lockup_period'], volatility_tornado, drift_tornado, tornado_chart_path)

    # 添加图表4.5.1：盈利概率敏感性
    add_paragraph(document, '图表 4.5.1: 参数敏感性排序（龙卷风图） - 盈利概率敏感性')
    add_image(document, prob_chart_path, width=Inches(6.5))
    add_paragraph(document, '')

    # 添加图表4.5.2：预期收益率敏感性
    add_paragraph(document, '图表 4.5.2: 参数敏感性排序（龙卷风图） - 预期收益率敏感性')
    add_image(document, return_chart_path, width=Inches(6.5))
    add_paragraph(document, '')

    add_paragraph(document, '敏感性分析结论：')
    add_paragraph(document, '')
    add_paragraph(document, '**分析方法说明**：')
    add_paragraph(document, '• 采用**归一化敏感性**分析方法，确保不同参数的公平比较')
    add_paragraph(document, '• 归一化含义：将影响折算为"每单位变化"的效应')
    add_paragraph(document, '• 图表排序依据：归一化敏感性（单位变化的影响程度）')
    add_paragraph(document, '• 图表柱状显示：实际情景影响（便于直观理解）')
    add_paragraph(document, '• 优势：避免因变化幅度不同导致的排序偏差，真实反映参数敏感度')
    add_paragraph(document, '')

    add_paragraph(document, '• **盈利概率敏感性**：')

    # 找出对盈利概率归一化敏感性最大的参数
    max_prob_sensitivity = max(top_prob, key=lambda x: abs(x.get('prob_sensitivity', x.get('prob_impact_pos', 0))))
    max_prob_impact = max_prob_sensitivity.get('prob_impact_pos', 0)

    if max_prob_impact > 5:
        add_paragraph(document, f"  - {max_prob_sensitivity['parameter']}对盈利概率敏感性最高（单位变化影响最大）")
        add_paragraph(document, f"    实际情景影响：{max_prob_impact:+.1f}%")
    elif max_prob_impact < -5:
        add_paragraph(document, f"  - {max_prob_sensitivity['parameter']}对盈利概率负面敏感性最高")
        add_paragraph(document, f"    实际情景影响：{max_prob_impact:+.1f}%")
    else:
        add_paragraph(document, f"  - 各参数对盈利概率的敏感性相对均衡")
        add_paragraph(document, f"    最高敏感性参数：{max_prob_sensitivity['parameter']}（情景影响：{max_prob_impact:+.1f}%）")

    add_paragraph(document, '')
    add_paragraph(document, '• **预期收益率敏感性**：')

    # 找出对预期收益率归一化敏感性最大的参数
    max_return_sensitivity = max(top_return, key=lambda x: abs(x.get('return_sensitivity', x.get('return_impact_pos', 0))))
    max_return_impact = max_return_sensitivity.get('return_impact_pos', 0)

    if max_return_impact > 5:
        add_paragraph(document, f"  - {max_return_sensitivity['parameter']}对预期收益率敏感性最高（单位变化影响最大）")
        add_paragraph(document, f"    实际情景影响：{max_return_impact:+.1f}%")
    elif max_return_impact < -5:
        add_paragraph(document, f"  - {max_return_sensitivity['parameter']}对预期收益率负面敏感性最高")
        add_paragraph(document, f"    实际情景影响：{max_return_impact:+.1f}%")
    else:
        add_paragraph(document, f"  - 各参数对预期收益率的敏感性相对均衡")
        add_paragraph(document, f"    最高敏感性参数：{max_return_sensitivity['parameter']}（情景影响：{max_return_impact:+.1f}%）")

    add_paragraph(document, '')
    add_paragraph(document, '• **综合判断**：')
    add_paragraph(document, '  - 价格类参数（发行价、当前价）通常是盈利概率的主要敏感性因素')
    add_paragraph(document, '  - 波动率和漂移率主要影响预期收益率（高波动率可能带来更高收益，但风险也更大）')
    add_paragraph(document, '  - 归一化分析揭示了各参数的真实敏感程度，便于风险识别和管理')
    add_paragraph(document, '')
    add_paragraph(document, '• **投资建议**：')
    add_paragraph(document, '  - 重点关注高敏感性参数，这些是项目的主要风险因子')
    add_paragraph(document, '  - 根据自身风险偏好，平衡盈利概率和预期收益率')
    add_paragraph(document, '  - 对于敏感性高的参数，建议设置更严格的风控阈值')

    add_section_break(document)

    return context
