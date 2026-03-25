# -*- coding: utf-8 -*-
"""
通用工具函数模块

包含所有Word文档格式化和图表生成的工具函数
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats

# 获取中文字体（在导入时初始化）
try:
    import sys
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
    sys.path.insert(0, PROJECT_DIR)
    from utils.font_manager import get_font_prop
    from utils.analysis_tools import calculate_profit_probability_lognormal
    font_prop = get_font_prop()
except Exception as e:
    print(f"⚠️ 无法加载中文字体: {e}")
    font_prop = None


# ==================== Word文档格式化函数 ====================

def setup_chinese_font(document):
    """设置中文字体 - V2版本"""
    # 正文：仿宋-GB2312，四号（14pt）
    document.styles['Normal'].font.name = '仿宋_GB2312'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    document.styles['Normal'].font.size = Pt(14)


def add_title(document, text, level=1):
    """添加标题 - V2版本"""
    if level == 0:  # 报告标题，使用段落而不是heading
        para = document.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = '方正公文小标宋_GBK'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
            run.font.size = Pt(22)  # 二号（22pt）
            run.font.bold = True
        return para
    else:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            # 根据级别设置字体和大小
            if level == 1:  # 一级标题
                run.font.name = '方正公文小标宋_GBK'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
                run.font.size = Pt(16)  # 三号（16pt）
            elif level == 2:  # 二级标题
                run.font.name = '方正公文小标宋_GBK'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
                run.font.size = Pt(15)  # 小三（15pt）
            else:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(14)
        return heading


def add_paragraph(document, text, bold=False, font_size=14):
    """添加段落 - V2版本（正文：仿宋-GB2312，四号）"""
    para = document.add_paragraph(text)
    for run in para.runs:
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(font_size)
        if bold:
            run.font.bold = True
    return para


def add_image(document, image_path, width=Inches(5)):
    """添加图片到文档"""
    if os.path.exists(image_path):
        document.add_picture(image_path, width=width)
        # 设置图片居中
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    else:
        print(f"⚠️ 图片不存在: {image_path}")
        return False


def add_table_data(document, headers, data, font_size=12):
    """添加表格 - V2版本（表格：宋体，默认小四号12pt）

    Args:
        document: Word文档对象
        headers: 表头列表
        data: 数据列表
        font_size: 字体大小（单位：磅），默认12pt（小四号），10.5pt为五号
    """
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 设置表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        para = header_cells[i].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.size = Pt(font_size)
        para.runs[0].font.name = '宋体'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            para = row_cells[i].paragraphs[0]
            para.runs[0].font.size = Pt(font_size)
            para.runs[0].font.name = '宋体'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if isinstance(cell_data, (int, float)):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def add_section_break(document):
    """添加分页符"""
    document.add_page_break()


# ==================== 图表生成函数 ====================

def generate_break_even_chart(issue_price, current_price, lockup_period, save_path):
    """生成盈亏平衡价格敏感性曲线"""
    # 不同期望收益率下的盈亏平衡价格
    target_returns = np.linspace(0.05, 0.50, 10)  # 5%到50%年化收益率
    break_even_prices = [issue_price * (1 + r * (lockup_period / 12)) for r in target_returns]

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(target_returns * 100, break_even_prices, 'b-', linewidth=2, label='盈亏平衡价格')
    ax.axhline(y=issue_price, color='r', linestyle='--', label='发行价格')
    ax.axhline(y=current_price, color='g', linestyle='--', label='当前价格')

    ax.set_xlabel('期望年化收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('盈亏平衡价格 (元)', fontproperties=font_prop, fontsize=12)
    ax.set_title('不同收益率要求下的盈亏平衡价格', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    ax.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_lockup_sensitivity_charts_split(issue_price, current_price, save_dir):
    """生成锁定期敏感性分析图表 - 拆分版（V2：生成2个单独图片）"""
    lockup_periods = [6, 9, 12, 18, 24, 36]  # 月
    target_return = 0.20  # 20%年化收益率

    lockup_analysis = []
    for period in lockup_periods:
        period_years = period / 12
        be_price = issue_price * (1 + target_return * period_years)
        required_increase = (be_price - issue_price) / issue_price * 100
        lockup_analysis.append({'period': period, 'be_price': be_price, 'increase': required_increase})

    # 图1：盈亏平衡价格
    fig1, ax1 = plt.subplots(figsize=(12, 6))
    ax1.bar([p['period'] for p in lockup_analysis], [p['be_price'] for p in lockup_analysis],
            color='steelblue', alpha=0.7)
    ax1.axhline(y=current_price, color='r', linestyle='--', label='当前价格')
    ax1.set_xlabel('锁定期 (月)', fontproperties=font_prop, fontsize=14)
    ax1.set_ylabel('盈亏平衡价格 (元)', fontproperties=font_prop, fontsize=14)
    ax1.set_title('锁定期对盈亏平衡价格的影响', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax1.legend(prop=font_prop, fontsize=12)
    for label in ax1.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax1.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    ax1.grid(True, alpha=0.3)

    plt.tight_layout()
    chart1_path = os.path.join(save_dir, 'lockup_sensitivity_be_price.png')
    plt.savefig(chart1_path, dpi=150, bbox_inches='tight')
    plt.close()

    # 图2：需要涨幅
    fig2, ax2 = plt.subplots(figsize=(12, 6))
    ax2.plot([p['period'] for p in lockup_analysis], [p['increase'] for p in lockup_analysis],
            'ro-', linewidth=3, markersize=10)
    ax2.set_xlabel('锁定期 (月)', fontproperties=font_prop, fontsize=14)
    ax2.set_ylabel('需要涨幅 (%)', fontproperties=font_prop, fontsize=14)
    ax2.set_title('锁定期与需要涨幅的关系', fontproperties=font_prop, fontsize=16, fontweight='bold')
    for label in ax2.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax2.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    ax2.grid(True, alpha=0.3)

    plt.tight_layout()
    chart2_path = os.path.join(save_dir, 'lockup_sensitivity_increase.png')
    plt.savefig(chart2_path, dpi=150, bbox_inches='tight')
    plt.close()

    return chart1_path, chart2_path


def generate_time_window_analysis_chart(price_series, save_dir):
    """生成不同时间窗口的风险指标对比图表"""
    chart_paths = []

    # 定义不同时间窗口（天数）
    # 月度(20日)、季度(60日)、半年(120日)、年度(250日)
    windows = [20, 60, 120, 250]

    # 计算每个窗口的统计指标
    results = {
        'window': [],
        'volatility': [],
        'total_return': [],     # 期间收益率
        'annual_return': [],    # 年化收益率
        'mean': [],
        'std': [],
        'max_drawdown': [],
        'sharpe': []
    }

    risk_free_rate = 0.03  # 无风险利率

    for window in windows:
        if len(price_series) < window:
            continue

        # 取最近window天的数据
        recent_prices = price_series[-window:]

        # 计算收益率
        returns = recent_prices.pct_change().dropna()

        # 年化波动率
        volatility = returns.std() * np.sqrt(250)

        # 期间收益率
        total_return = (recent_prices.iloc[-1] / recent_prices.iloc[0]) - 1

        # 年化收益率（使用单利计算，按交易日比例）
        # 年化收益率 = 期间收益率 × (250 / 窗口期天数)
        annual_return = total_return * (250.0 / window)

        # 均值
        mean = recent_prices.mean()

        # 标准差
        std = recent_prices.std()

        # 最大回撤
        cummax = recent_prices.cummax()
        drawdown = (recent_prices - cummax) / cummax
        max_drawdown = drawdown.min()

        # 夏普比率
        sharpe = (annual_return - risk_free_rate) / volatility if volatility > 0 else 0

        results['window'].append(window)
        results['volatility'].append(volatility)
        results['total_return'].append(total_return)     # 期间收益率
        results['annual_return'].append(annual_return)    # 年化收益率
        results['mean'].append(mean)
        results['std'].append(std)
        results['max_drawdown'].append(max_drawdown)
        results['sharpe'].append(sharpe)

    # 图1：波动率对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    bars = ax.bar(range(len(results['window'])), [v*100 for v in results['volatility']],
                   color='steelblue', alpha=0.7)
    ax.set_xlabel('时间窗口（交易日）', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('年化波动率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同时间窗口的波动率对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(results['window'])))
    ax.set_xticklabels([f'{w}日' for w in results['window']], fontproperties=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, axis='y')

    # 添加数值标签
    for bar, vol in zip(bars, results['volatility']):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                f'{vol*100:.1f}%', ha='center', va='bottom', fontsize=11, fontproperties=font_prop)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'time_window_volatility_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2：年化收益率对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    colors_return = ['green' if r > 0 else 'red' for r in results['annual_return']]
    bars = ax.bar(range(len(results['window'])), [r*100 for r in results['annual_return']],
                   color=colors_return, alpha=0.7)
    ax.set_xlabel('时间窗口（交易日）', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('年化收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同时间窗口的年化收益率对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(results['window'])))
    ax.set_xticklabels([f'{w}日' for w in results['window']], fontproperties=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, axis='y')
    ax.axhline(y=0, color='black', linestyle='-', linewidth=1)

    # 添加数值标签
    for bar, ret in zip(bars, results['annual_return']):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3 if ret > 0 else bar.get_height() - 0.3,
                f'{ret*100:+.1f}%', ha='center', va='bottom' if ret > 0 else 'top',
                fontsize=11, fontproperties=font_prop, fontweight='bold')

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'time_window_return_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3：风险指标汇总对比（单独大图）
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 7))

    # 左图：最大回撤
    bars1 = ax1.bar(range(len(results['window'])), [d*100 for d in results['max_drawdown']],
                    color='#e74c3c', alpha=0.7)
    ax1.set_xlabel('时间窗口（交易日）', fontproperties=font_prop, fontsize=12)
    ax1.set_ylabel('最大回撤 (%)', fontproperties=font_prop, fontsize=12)
    ax1.set_title('不同时间窗口的最大回撤', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax1.set_xticks(range(len(results['window'])))
    ax1.set_xticklabels([f'{w}日' for w in results['window']], fontproperties=font_prop, fontsize=10)
    ax1.grid(True, alpha=0.3, axis='y')

    for bar, dd in zip(bars1, results['max_drawdown']):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() - 0.5,
                f'{dd*100:.1f}%', ha='center', va='top', fontsize=9, fontproperties=font_prop)

    # 右图：夏普比率
    colors_sharpe = ['green' if s > 0.5 else 'orange' if s > 0 else 'red' for s in results['sharpe']]
    bars2 = ax2.bar(range(len(results['window'])), results['sharpe'],
                    color=colors_sharpe, alpha=0.7)
    ax2.set_xlabel('时间窗口（交易日）', fontproperties=font_prop, fontsize=12)
    ax2.set_ylabel('夏普比率', fontproperties=font_prop, fontsize=12)
    ax2.set_title('不同时间窗口的夏普比率', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax2.set_xticks(range(len(results['window'])))
    ax2.set_xticklabels([f'{w}日' for w in results['window']], fontproperties=font_prop, fontsize=10)
    ax2.grid(True, alpha=0.3, axis='y')
    ax2.axhline(y=1, color='black', linestyle='--', alpha=0.5, label='基准线')

    for bar, sr in zip(bars2, results['sharpe']):
        ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.02,
                f'{sr:.2f}', ha='center', va='bottom', fontsize=9, fontproperties=font_prop)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'time_window_risk_metrics_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return results, chart_paths


def generate_tornado_chart_split(issue_price, current_price, lockup_period, volatility, drift, save_dir):
    """生成龙卷风图 - 参数敏感性排序（拆分版）

    敏感性分析原则：
    - 所有参数使用标准化的单位变化
    - 价格类参数：变化±10%
    - 百分比参数（波动率、漂移率）：变化±10%
    - 这样可以公平比较各参数对盈利概率的影响程度

    双维度分析：
    - 维度1：盈利概率敏感性（能否盈利）
    - 维度2：预期收益率敏感性（盈利多少）

    生成两个独立的图表，分别展示盈利概率敏感性和预期收益率敏感性
    """
    # 基准参数
    base_params = {
        'issue_price': issue_price,
        'current_price': current_price,
        'volatility': volatility,
        'drift': drift
    }

    # 计算基准盈利概率（使用对数正态分布）
    def calculate_profit_prob(params):
        period_years = lockup_period / 12  # 使用外部变量lockup_period
        vol_period = params['volatility'] * np.sqrt(period_years)
        drift_period = params['drift'] * period_years

        # 使用对数正态分布计算盈利概率
        # 发行价作为盈利阈值
        prob = calculate_profit_probability_lognormal(
            target_price=params['issue_price'],
            current_price=params['current_price'],
            drift=params['drift'],
            volatility=params['volatility'],
            period_months=lockup_period
        )
        return prob

    # 计算预期收益率（年化）
    def calculate_expected_return(params):
        period_years = lockup_period / 12  # 使用外部变量lockup_period
        # 使用GBM模型计算预期终值
        # E[S_T] = S_0 * exp(μ * T)
        expected_final_price = params['current_price'] * np.exp(params['drift'] * period_years)
        expected_return = (expected_final_price - params['issue_price']) / params['issue_price']
        return expected_return * 100  # 转为百分比

    base_prob = calculate_profit_prob(base_params)
    base_return = calculate_expected_return(base_params)
    print(f"基准盈利概率: {base_prob:.2f}%")
    print(f"基准预期收益率: {base_return:.2f}%")

    # 参数变化分析 - 使用标准化的单位变化
    param_changes = []

    # 定义每个参数的变化方式和幅度
    param_definitions = [
        {
            'name': 'issue_price',
            'name_cn': '发行价',
            'changes': [-0.10, 0.10],  # ±10%
            'change_type': 'percent',
            'apply_change': lambda base_val, change: base_val * (1 + change)
        },
        {
            'name': 'current_price',
            'name_cn': '当前价',
            'changes': [-0.10, 0.10],  # ±10%
            'change_type': 'percent',
            'apply_change': lambda base_val, change: base_val * (1 + change)
        },
        {
            'name': 'volatility',
            'name_cn': '波动率',
            'changes': [-0.10, 0.10],  # ±10%
            'change_type': 'percent',
            'apply_change': lambda base_val, change: base_val * (1 + change)
        },
        {
            'name': 'drift',
            'name_cn': '漂移率',
            'changes': [-0.10, 0.10],  # ±10%
            'change_type': 'percent',
            'apply_change': lambda base_val, change: base_val * (1 + change)
        }
    ]

    # 计算各参数变化的敏感性（双维度）
    for param_def in param_definitions:
        param_name = param_def['name']
        param_name_cn = param_def['name_cn']

        for change in param_def['changes']:
            params = base_params.copy()
            params[param_name] = param_def['apply_change'](base_params[param_name], change)

            new_prob = calculate_profit_prob(params)
            new_return = calculate_expected_return(params)

            prob_impact = new_prob - base_prob
            return_impact = new_return - base_return

            # 计算归一化的敏感性（用于排序）
            # 盈利概率敏感性：每单位变化的影响
            # 预期收益率敏感性：每单位变化的影响
            if param_def['change_type'] == 'percent':
                # 百分比变化：每1%的影响
                prob_sensitivity = prob_impact / abs(change * 100)
                return_sensitivity = return_impact / abs(change * 100)
                change_unit = f'每1%变化'
            else:  # absolute
                # 绝对值变化：每1单位的影响
                prob_sensitivity = prob_impact / abs(change)
                return_sensitivity = return_impact / abs(change)
                change_unit = f'每1个月变化'

            # 生成变化描述
            if change < 0:
                change_desc = f"下降{abs(int(change*100))}%"
            else:
                change_desc = f"上升{int(change*100)}%"

            param_changes.append({
                'parameter': param_name_cn,
                'change_desc': change_desc,
                'prob_impact': prob_impact,
                'return_impact': return_impact,
                'prob_sensitivity': prob_sensitivity,  # 归一化的敏感性（用于排序）
                'return_sensitivity': return_sensitivity,  # 归一化的敏感性（用于排序）
                'change_unit': change_unit,  # 变化单位说明
                'full_label_prob': f'{param_name_cn}{change_desc}',
                'full_label_return': f'{param_name_cn}{change_desc}'
            })

            print(f"{param_name_cn}{change_desc}: 盈利概率 {base_prob:.2f}% → {new_prob:.2f}% ({prob_impact:+.2f}%, 敏感性: {prob_sensitivity:+.3f}), 预期收益率 {base_return:.2f}% → {new_return:.2f}% ({return_impact:+.2f}%, 敏感性: {return_sensitivity:+.3f})")

    # 按归一化的敏感性排序（找出真正最敏感的参数）
    param_changes_by_prob = sorted(param_changes, key=lambda x: abs(x['prob_sensitivity']), reverse=True)
    param_changes_by_return = sorted(param_changes, key=lambda x: abs(x['return_sensitivity']), reverse=True)

    # 只保留影响最大的前8个
    top_prob = param_changes_by_prob[:8]
    top_return = param_changes_by_return[:8]

    # 生成两个独立的图表
    # 图表1：盈利概率敏感性
    fig1, ax1 = plt.subplots(figsize=(12, 8))
    colors1 = ['red' if p['prob_impact'] < 0 else 'green' for p in top_prob]
    bars1 = ax1.barh(range(len(top_prob)), [p['prob_impact'] for p in top_prob],
                    color=colors1, alpha=0.7)

    ax1.set_xlabel('对盈利概率的影响 (百分点)', fontproperties=font_prop, fontsize=12)
    ax1.set_title('图表 4.5.1: 参数敏感性排序（龙卷风图） - 盈利概率敏感性\n标准化单位变化：价格±10%、波动率±10%、漂移率±10%',
                 fontproperties=font_prop, fontsize=13, fontweight='bold')
    ax1.set_yticks(range(len(top_prob)))
    ax1.set_yticklabels([p['full_label_prob'] for p in top_prob], fontproperties=font_prop, fontsize=10)
    ax1.axvline(x=0, color='black', linestyle='-', linewidth=1)
    ax1.grid(True, axis='x', alpha=0.3)

    # 添加数值标签
    for i, (bar, p) in enumerate(zip(bars1, top_prob)):
        width = bar.get_width()
        xpos = width + 0.3 if width >= 0 else width - 0.3
        ax1.text(xpos, i, f'{p["prob_impact"]:+.1f}%',
                va='center', fontsize=9, fontproperties=font_prop,
                color='green' if width >= 0 else 'red')

    # 设置字体
    for label in ax1.get_xticklabels() + ax1.get_yticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()
    prob_chart_path = os.path.join(save_dir, 'tornado_prob_sensitivity.png')
    plt.savefig(prob_chart_path, dpi=150, bbox_inches='tight')
    plt.close()

    # 图表2：预期收益率敏感性
    fig2, ax2 = plt.subplots(figsize=(12, 8))
    colors2 = ['red' if p['return_impact'] < 0 else 'green' for p in top_return]
    bars2 = ax2.barh(range(len(top_return)), [p['return_impact'] for p in top_return],
                    color=colors2, alpha=0.7)

    ax2.set_xlabel('对预期收益率的影响 (百分点)', fontproperties=font_prop, fontsize=12)
    ax2.set_title('图表 4.5.2: 参数敏感性排序（龙卷风图） - 预期收益率敏感性\n标准化单位变化：价格±10%、波动率±10%、漂移率±10%',
                 fontproperties=font_prop, fontsize=13, fontweight='bold')
    ax2.set_yticks(range(len(top_return)))
    ax2.set_yticklabels([p['full_label_return'] for p in top_return], fontproperties=font_prop, fontsize=10)
    ax2.axvline(x=0, color='black', linestyle='-', linewidth=1)
    ax2.grid(True, axis='x', alpha=0.3)

    # 添加数值标签
    for i, (bar, p) in enumerate(zip(bars2, top_return)):
        width = bar.get_width()
        xpos = width + 0.3 if width >= 0 else width - 0.3
        ax2.text(xpos, i, f'{p["return_impact"]:+.1f}%',
                va='center', fontsize=9, fontproperties=font_prop,
                color='green' if width >= 0 else 'red')

    # 设置字体
    for label in ax2.get_xticklabels() + ax2.get_yticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()
    return_chart_path = os.path.join(save_dir, 'tornado_return_sensitivity.png')
    plt.savefig(return_chart_path, dpi=150, bbox_inches='tight')
    plt.close()

    return prob_chart_path, return_chart_path, top_prob, top_return


def generate_sensitivity_charts_split(volatilities, profit_probs, current_vol, ma20, issue_price, save_dir):
    """生成敏感性分析图表 - 拆分版本（V2：生成4个单独图片）"""
    chart_paths = []

    # 图1：波动率 vs 盈利概率柱状图（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    colors = ['#27ae60' if p >= 60 else '#f39c12' if p >= 40 else '#e74c3c' for p in profit_probs]
    bars = ax.bar(range(len(volatilities)), profit_probs, color=colors, alpha=0.8, edgecolor='black', linewidth=1.5)
    ax.set_xlabel('年化波动率', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同波动率下的盈利概率', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(volatilities)))
    ax.set_xticklabels([f'{v*100:.0f}%' for v in volatilities], fontproperties=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.set_ylim(0, 100)
    ax.axhline(y=50, color='gray', linestyle='--', alpha=0.5)

    # 标记当前波动率
    current_idx = np.argmin(np.abs(np.array(volatilities) - current_vol))
    ax.axvline(x=current_idx, color='red', linestyle='-', linewidth=2, alpha=0.7)
    ax.text(current_idx, 95, f'当前\n{current_vol*100:.1f}%', ha='center', va='top',
            fontsize=12, fontproperties=font_prop, color='red', fontweight='bold')
    for bar, prob in zip(bars, profit_probs):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                f'{prob:.1f}%', ha='center', va='bottom', fontsize=11, fontweight='bold')

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'sensitivity_volatility_profit_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2：发行价折扣分析（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    discounts = np.array([5, 10, 15, 20])  # 折价率
    issue_prices_discounted = ma20 * (1 - discounts/100)

    # 计算不同发行价下的盈利概率
    discount_probs = []
    for discounted_price in issue_prices_discounted:
        lockup_months = 6
        drift_rate = -0.1875
        lockup_drift = drift_rate * (lockup_months / 12)
        lockup_vol = 0.3063 * np.sqrt(lockup_months / 12)
        z = (np.log(ma20 / discounted_price) - lockup_drift) / lockup_vol
        prob = stats.norm.cdf(z) * 100
        discount_probs.append(prob)

    colors_discount = ['#27ae60', '#2ecc71', '#f39c12', '#e74c3c']
    bars = ax.bar(range(len(discounts)), discount_probs, color=colors_discount,
                   alpha=0.8, edgecolor='black', linewidth=1.5)
    ax.set_xlabel('发行价相对MA20折价率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同发行价折扣下的盈利概率', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(discounts)))
    ax.set_xticklabels([f'{d}%' for d in discounts], fontproperties=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.set_ylim(0, 100)
    ax.axhline(y=50, color='gray', linestyle='--', alpha=0.5, label='盈亏平衡线')
    ax.axvline(x=0, color='blue', linestyle='--', linewidth=2, alpha=0.7, label='当前发行价位置')
    ax.legend(prop=font_prop, fontsize=12, loc='lower right')

    for bar, prob, price in zip(bars, discount_probs, issue_prices_discounted):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                f'{prob:.1f}%', ha='center', va='bottom', fontsize=11, fontweight='bold')
        ax.text(bar.get_x() + bar.get_width()/2, -3,
                f'{price:.2f}元', ha='center', va='top', fontsize=10, fontproperties=font_prop)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'sensitivity_discount_profit_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3：热力图（波动率 vs 漂移率）- 单独大图
    fig, ax = plt.subplots(figsize=(14, 10))
    drift_range = np.linspace(-0.3, 0.15, 8)
    vol_range_heatmap = np.linspace(0.20, 0.50, 8)
    heatmap_data = []

    for d in drift_range:
        row = []
        for v in vol_range_heatmap:
            lockup_drift = d * (6/12)
            lockup_vol = v * np.sqrt(6/12)
            z = (np.log(ma20/issue_price) - lockup_drift) / lockup_vol
            row.append(stats.norm.cdf(z) * 100)
        heatmap_data.append(row)

    im = ax.imshow(heatmap_data, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100,
                   extent=[vol_range_heatmap[0]*100, vol_range_heatmap[-1]*100, drift_range[0]*100, drift_range[-1]*100])
    ax.set_xlabel('年化波动率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('年化漂移率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('盈利概率敏感性热力图（波动率 vs 漂移率）', fontproperties=font_prop, fontsize=16, fontweight='bold')
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    cbar = plt.colorbar(im, ax=ax)
    cbar.set_label('盈利概率 (%)', fontproperties=font_prop, fontsize=12)

    # 标记当前位置
    ax.scatter([current_vol*100], [drift_rate*100], color='red', s=300, marker='*',
               edgecolors='white', linewidths=2, label='当前位置', zorder=5)
    ax.legend(prop=font_prop, loc='upper right', fontsize=12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'sensitivity_heatmap_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图4：热力图（漂移率 vs 折价率）- 单独大图
    fig, ax = plt.subplots(figsize=(16, 10))
    drift_range_new = np.linspace(-0.30, 0.30, 7)  # -30%到+30%，7个档位
    discount_range = np.array([-30, -20, -10, -5, 0, 5, 10, 20, 30])  # 折价/溢价率-30%到+30%，更全面的范围
    heatmap_data_drift_discount = []

    for d in drift_range_new:
        row = []
        for discount in discount_range:
            issue_price_hm = ma20 * (1 + discount/100)  # 负值表示折价，正值表示溢价
            lockup_drift = d * (6/12)
            lockup_vol = 0.3063 * np.sqrt(6/12)  # 使用当前波动率
            z = (np.log(ma20/issue_price_hm) - lockup_drift) / lockup_vol
            row.append(stats.norm.cdf(z) * 100)
        heatmap_data_drift_discount.append(row)

    im = ax.imshow(heatmap_data_drift_discount, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100)

    # 设置刻度
    ax.set_xticks(np.arange(len(discount_range)))
    ax.set_yticks(np.arange(len(drift_range_new)))
    ax.set_xticklabels([f'{d}%' for d in discount_range], fontproperties=font_prop, fontsize=12)
    ax.set_yticklabels([f'{d*100:+.0f}%' for d in drift_range_new], fontproperties=font_prop, fontsize=12)

    # 添加数值标注
    for i in range(len(drift_range_new)):
        for j in range(len(discount_range)):
            text = ax.text(j, i, f'{heatmap_data_drift_discount[i][j]:.0f}%',
                         ha="center", va="center", color="black", fontproperties=font_prop, fontsize=8,
                         fontweight='bold' if heatmap_data_drift_discount[i][j] >= 50 else 'normal')

    ax.set_xlabel('发行价折价/溢价率（负值表示折价，正值表示溢价）(%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('年化漂移率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('盈利概率敏感性热力图（漂移率 vs 折价率）', fontproperties=font_prop, fontsize=16, fontweight='bold')

    cbar = plt.colorbar(im, ax=ax)
    cbar.set_label('盈利概率 (%)', fontproperties=font_prop, fontsize=12)

    # 标记当前位置（需要传入当前折价/溢价率）
    current_vol_for_hm = 0.3063  # 从外部获取当前波动率
    current_drift_for_hm = -0.1875  # 从外部获取当前漂移率
    current_discount_for_hm = ((issue_price - ma20) / ma20) * 100  # 计算当前折价/溢价率（负值表示折价）

    # 找到最接近的刻度位置
    drift_idx = np.argmin(np.abs(drift_range_new - current_drift_for_hm))
    discount_idx = np.argmin(np.abs(discount_range - current_discount_for_hm))

    ax.scatter([discount_idx], [drift_idx], color='red', s=300, marker='*',
               edgecolors='white', linewidths=2, label='当前位置', zorder=5)
    ax.legend(prop=font_prop, loc='upper right', fontsize=12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'sensitivity_heatmap_drift_discount_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths


def generate_discount_scenario_charts_split(base_price, current_price, volatility, drift, lockup_period, save_dir):
    """生成发行价折扣情景图表 - 统一版本（-20%至+20%）

    参数:
        base_price: 基准价格（MA20）
        current_price: 当前价格
        volatility: 波动率
        drift: 漂移率
        lockup_period: 锁定期（月）
        save_dir: 保存目录

    返回:
        图表路径列表 [盈利概率图, 预期收益率图]
    """
    chart_paths = []
    # 统一范围：从-20%（折价）到+20%（溢价），包括0%
    all_rates = [-20, -15, -10, -5, 0, 5, 10, 15, 20]

    # 计算所有情景的数据
    all_data = []
    lockup_vol = volatility * np.sqrt(lockup_period / 12)
    lockup_drift = drift * (lockup_period / 12)
    expected_future_price = current_price * np.exp(lockup_drift)

    for rate in all_rates:
        if rate < 0:
            # 折价发行
            issue_price = base_price * (1 + rate/100)  # rate是负数
            threshold = issue_price
        elif rate == 0:
            # 平价发行
            issue_price = base_price
            threshold = issue_price
        else:
            # 溢价发行
            issue_price = base_price * (1 + rate/100)
            threshold = max(base_price, issue_price * 1.02)

        # 计算盈利概率（使用对数正态分布）
        profit_prob = calculate_profit_probability_lognormal(
            target_price=threshold,
            current_price=current_price,
            drift=drift,
            volatility=volatility,
            period_months=lockup_period
        )

        # 计算预期收益率
        expected_return = (expected_future_price - issue_price) / issue_price * 100

        all_data.append({
            'rate': rate,
            'profit_prob': profit_prob,
            'expected_return': expected_return,
            'issue_price': issue_price
        })

    # 图1：盈利概率对比（统一图）
    fig, ax = plt.subplots(figsize=(14, 8))

    # 根据折价/溢价设置颜色
    colors = []
    for d in all_data:
        if d['rate'] < 0:
            # 折价发行：绿色
            colors.append('#27ae60')
        elif d['rate'] == 0:
            # 平价发行：灰色
            colors.append('#7f8c8d')
        else:
            # 溢价发行：红色
            colors.append('#e74c3c')

    bars = ax.bar([d['rate'] for d in all_data], [d['profit_prob'] for d in all_data],
                 color=colors, alpha=0.8, edgecolor='black', linewidth=2)

    ax.set_xlabel('发行价相对MA20的折扣/溢价率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('发行价敏感性分析 - 盈利概率', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks([d['rate'] for d in all_data])
    ax.set_xticklabels([f"{d['rate']}%" for d in all_data], fontproperties=font_prop, fontsize=12)
    ax.set_ylim(0, 100)
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.axhline(y=50, color='gray', linestyle='--', alpha=0.5, label='盈亏平衡线')
    ax.axvline(x=0, color='blue', linestyle=':', linewidth=2, label='MA20基准线')

    for bar, d in zip(bars, all_data):
        y_pos = d['profit_prob'] + 1 if d['profit_prob'] < 95 else d['profit_prob'] - 3
        ax.text(bar.get_x() + bar.get_width()/2, y_pos,
                f"{d['profit_prob']:.1f}%", ha='center', va='bottom' if d['profit_prob'] < 95 else 'top',
                fontsize=11, fontweight='bold')

    ax.legend(prop=font_prop, fontsize=12)
    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'scenario_unified_profit_prob.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2：预期收益率对比（统一图）
    fig, ax = plt.subplots(figsize=(14, 8))

    colors = []
    for d in all_data:
        if d['rate'] < 0:
            colors.append('#27ae60')  # 折价：绿色
        elif d['rate'] == 0:
            colors.append('#7f8c8d')  # 平价：灰色
        else:
            colors.append('#e74c3c')  # 溢价：红色

    bars = ax.bar([d['rate'] for d in all_data], [d['expected_return'] for d in all_data],
                 color=colors, alpha=0.8, edgecolor='black', linewidth=2)

    ax.set_xlabel('发行价相对MA20的折扣/溢价率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('预期收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('发行价敏感性分析 - 预期收益率', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks([d['rate'] for d in all_data])
    ax.set_xticklabels([f"{d['rate']}%" for d in all_data], fontproperties=font_prop, fontsize=12)
    ax.axhline(y=0, color='gray', linestyle='--', alpha=0.7, linewidth=2)
    ax.axvline(x=0, color='blue', linestyle=':', linewidth=2, label='MA20基准线')
    ax.grid(True, alpha=0.3, linestyle='--')

    for bar, d in zip(bars, all_data):
        height = d['expected_return']
        y_pos = height + 1 if height >= 0 else height - 3
        ax.text(bar.get_x() + bar.get_width()/2, y_pos,
                f'{height:.1f}%', ha='center', va='bottom' if height >= 0 else 'top',
                fontsize=11, fontweight='bold')

    ax.legend(prop=font_prop, fontsize=12)
    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'scenario_unified_expected_return.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths


def generate_multi_dimension_scenario_charts_split(current_price, base_price, volatility, drift, lockup_period, save_dir):
    """生成多维度情景分析图表（波动率×折扣率×时间窗口） - 拆分版本（3个单独图片）

    参数:
        current_price: 当前价格
        base_price: 基准价格（MA20）
        volatility: 波动率
        drift: 漂移率
        lockup_period: 锁定期（月）
        save_dir: 保存目录

    返回:
        图表路径列表
    """
    import os
    chart_paths = []

    # 定义情景参数
    volatility_scenarios = [0.20, 0.30, 0.40, 0.50]
    discount_scenarios = [-10, -5, 0, 5, 10, 15, 20]
    window_scenarios = [60, 120, 180]

    # 计算所有情景
    multi_dim_results = []
    for vol in volatility_scenarios:
        for discount in discount_scenarios:
            for window in window_scenarios:
                if discount >= 0:
                    issue_price = base_price * (1 - discount/100)
                    threshold = issue_price
                else:
                    issue_price = base_price * (1 + abs(discount)/100)
                    threshold = issue_price

                # 计算盈利概率（使用对数正态分布）
                profit_prob = calculate_profit_probability_lognormal(
                    target_price=threshold,
                    current_price=current_price,
                    drift=drift,
                    volatility=vol,
                    period_months=window
                )

                multi_dim_results.append({
                    'volatility': vol,
                    'discount': discount,
                    'window': window,
                    'profit_prob': profit_prob
                })

    # 转换为DataFrame方便处理
    df_scenarios = pd.DataFrame(multi_dim_results)

    # 图1：波动率 vs 折扣率热力图（使用60日窗口）
    fig, ax = plt.subplots(figsize=(14, 10))
    pivot_data = df_scenarios[df_scenarios['window'] == 60].pivot(
        index='volatility', columns='discount', values='profit_prob'
    )

    im = ax.imshow(pivot_data.values, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100,
                   extent=[pivot_data.columns.min(), pivot_data.columns.max(),
                          pivot_data.index.min(), pivot_data.index.max()])

    ax.set_xlabel('发行价折扣率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('年化波动率', fontproperties=font_prop, fontsize=14)
    ax.set_title('盈利概率热力图：波动率 × 折扣率（60日窗口）', fontproperties=font_prop, fontsize=16, fontweight='bold')
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    cbar = plt.colorbar(im, ax=ax)
    cbar.set_label('盈利概率 (%)', fontproperties=font_prop, fontsize=12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'multi_dim_vol_discount_heatmap.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2：不同窗口下的盈利概率对比
    fig, ax = plt.subplots(figsize=(14, 8))

    for window in window_scenarios:
        data_window = df_scenarios[df_scenarios['window'] == window]
        # 取中等波动率(0.30)的数据
        data_median = data_window[data_window['volatility'] == 0.30]
        ax.plot(data_median['discount'], data_median['profit_prob'],
               marker='o', linewidth=2, markersize=8, label=f'{window}日窗口')

    ax.set_xlabel('发行价折扣率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同时间窗口下的盈利概率对比（波动率=30%）', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'multi_dim_window_comparison.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3：3D柱状图 - 波动率×窗口×折扣率
    fig = plt.figure(figsize=(16, 10))
    ax = fig.add_subplot(111, projection='3d')

    # 选择几个代表性的折扣率
    selected_discounts = [-10, 0, 10, 20]

    for discount in selected_discounts:
        data_disc = df_scenarios[df_scenarios['discount'] == discount]

        # 创建网格
        vol_mesh = []
        win_mesh = []
        prob_mesh = []

        for vol in volatility_scenarios:
            for window in window_scenarios:
                vol_mesh.append(vol)
                win_mesh.append(window)
                prob_val = data_disc[(data_disc['volatility'] == vol) &
                                    (data_disc['window'] == window)]['profit_prob'].values
                if len(prob_val) > 0:
                    prob_mesh.append(prob_val[0])
                else:
                    prob_mesh.append(0)

        # 绘制3D柱状图
        ax.bar3d(vol_mesh, win_mesh, [0]*len(vol_mesh),
                dx=0.03, dy=15, dz=prob_mesh,
                alpha=0.6, label=f'折扣{discount}%')

    ax.set_xlabel('年化波动率', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('时间窗口（日）', fontproperties=font_prop, fontsize=12)
    ax.set_zlabel('盈利概率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('多维度情景分析：波动率 × 时间窗口 × 折扣率', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=10)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'multi_dim_3d_analysis.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths


def generate_stress_test_charts_split(stress_test_results, save_dir):
    """生成压力测试图表 - 拆分版本

    参数:
        stress_test_results: 压力测试结果字典
        save_dir: 保存目录

    返回:
        图表路径列表
    """
    import os
    chart_paths = []

    # 提取数据
    scenarios = list(stress_test_results.keys())
    returns = [stress_test_results[s]['return'] * 100 for s in scenarios]

    # 图1：压力测试收益率对比
    fig, ax = plt.subplots(figsize=(14, 8))
    colors = ['#e74c3c' if r < 0 else '#2ecc71' for r in returns]
    bars = ax.barh(range(len(scenarios)), returns, color=colors, alpha=0.7)
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('情景', fontproperties=font_prop, fontsize=14)
    ax.set_title('压力测试情景分析', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_yticks(range(len(scenarios)))
    ax.set_yticklabels(scenarios, fontproperties=font_prop, fontsize=12)
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1)
    ax.grid(True, alpha=0.3, axis='x')

    for bar, ret in zip(bars, returns):
        ax.text(bar.get_width() + (1 if ret > 0 else -1),
                bar.get_y() + bar.get_height()/2,
                f'{ret:.1f}%', ha='left' if ret > 0 else 'right',
                va='center', fontsize=11, fontproperties=font_prop)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'stress_test_returns.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths


def generate_monte_carlo_charts_split(final_prices, issue_price, current_price, save_dir):
    """生成蒙特卡洛模拟图表 - 拆分版本（6个单独大图）"""
    returns = (final_prices - issue_price) / issue_price * 100
    chart_paths = []

    # 图1: 价格分布直方图（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.hist(final_prices, bins=50, alpha=0.7, color='steelblue', edgecolor='black')
    ax.axvline(x=issue_price, color='red', linestyle='--', linewidth=2, label=f'发行价: {issue_price:.2f}')
    ax.axvline(x=current_price, color='green', linestyle='--', linewidth=2, label=f'当前价: {current_price:.2f}')
    ax.set_xlabel('未来价格 (元)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('频数', fontproperties=font_prop, fontsize=14)
    ax.set_title('锁定期后价格分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_price_distribution_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2: 收益率分布（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.hist(returns, bins=50, alpha=0.7, color='coral', edgecolor='black')
    ax.axvline(x=0, color='red', linestyle='--', linewidth=2, label='盈亏平衡')
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('频数', fontproperties=font_prop, fontsize=14)
    ax.set_title('收益率分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_return_distribution_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3: 累积分布函数（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    sorted_returns = np.sort(returns)
    cumulative = np.arange(1, len(sorted_returns) + 1) / len(sorted_returns) * 100
    ax.plot(sorted_returns, cumulative, linewidth=2, color='#8e44ad')
    ax.axvline(x=0, color='red', linestyle='--', linewidth=1.5, label='盈亏平衡')
    ax.axhline(y=50, color='gray', linestyle='--', linewidth=1, alpha=0.5)
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('累积概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('累积分布函数', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_cdf_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图4: 盈亏概率饼图（单独大图）
    fig, ax = plt.subplots(figsize=(12, 10))
    profit_prob = (returns > 0).mean() * 100
    loss_prob = 100 - profit_prob
    colors_pie = ['#27ae60', '#e74c3c']
    wedges, texts, autotexts = ax.pie([profit_prob, loss_prob], labels=['盈利', '亏损'],
                                         autopct='%1.1f%%', colors=colors_pie, startangle=90,
                                         textprops={'fontproperties': font_prop, 'fontsize': 14})
    ax.set_title('盈亏概率分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    for autotext in autotexts:
        autotext.set_fontproperties(font_prop)
        autotext.set_fontsize(16)
        autotext.set_fontweight('bold')

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_profit_loss_pie_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图5: 关键统计指标（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.axis('off')
    stats_text = f"""
    关键统计指标
    ════════════════

    盈利概率: {profit_prob:.1f}%

    预期收益率: {returns.mean():.1f}%

    收益率中位数: {np.median(returns):.1f}%

    95% VaR: {abs(np.percentile(returns, 5)):.1f}%

    99% VaR: {abs(np.percentile(returns, 1)):.1f}%

    最大回撤: {abs(returns.min()):.1f}%

    收益率标准差: {returns.std():.1f}%

    盈亏平衡价: {issue_price:.2f}元

    当前价格: {current_price:.2f}元
    """
    ax.text(0.1, 0.5, stats_text, transform=ax.transAxes,
            fontsize=16, verticalalignment='center',
            fontproperties=font_prop,
            family='monospace',
            bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.3))

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_stats_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图6: 概率分布曲线（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    from scipy.stats import norm
    mu, std = returns.mean(), returns.std()
    x = np.linspace(returns.min(), returns.max(), 100)
    pdf = norm.pdf(x, mu, std)
    ax.plot(x, pdf, linewidth=2, color='#2980b9', label='正态分布拟合')
    ax.fill_between(x, pdf, alpha=0.3, color='#3498db')
    ax.axvline(x=0, color='red', linestyle='--', linewidth=2, label='盈亏平衡')
    ax.axvline(x=mu, color='green', linestyle='--', linewidth=2, label=f'均值: {mu:.1f}%')
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('概率密度', fontproperties=font_prop, fontsize=14)
    ax.set_title('收益率概率分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_pdf_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths


def generate_var_chart(var_95, var_99, cvar_95, save_path):
    """生成VaR风险度量图表"""
    plt.figure(figsize=(10, 6))
    metrics = ['95% VaR', '99% VaR', '95% CVaR']
    values = [var_95, var_99, cvar_95]
    colors = ['#e74c3c', '#c0392b', '#922b21']

    bars = plt.bar(metrics, [v*100 for v in values], color=colors, alpha=0.7)
    plt.ylabel('潜在损失 (%)', fontproperties=font_prop, fontsize=12)
    plt.title('风险度量 (VaR/CVaR)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    plt.xticks(fontproperties=font_prop)
    plt.yticks(fontproperties=font_prop)
    plt.grid(True, alpha=0.3, axis='y')

    for bar, value in zip(bars, values):
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                f'{value*100:.1f}%', ha='center', va='bottom',
                fontsize=11, fontweight='bold', fontproperties=font_prop)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_relative_valuation_charts_split(current_metrics, industry_avg, peer_companies_data, save_dir):
    """生成相对估值分析图表 - 拆分版本（单独大图）"""
    chart_paths = []

    # 图1: 估值指标对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    metrics = ['PE', 'PS', 'PB']
    current_vals = [current_metrics['pe'], current_metrics['ps'], current_metrics['pb']]
    industry_vals = [industry_avg['pe'], industry_avg['ps'], industry_avg['pb']]

    x = np.arange(len(metrics))
    width = 0.35

    bars1 = ax.bar(x - width/2, current_vals, width, label='光弘科技', color='#3498db', alpha=0.8)
    bars2 = ax.bar(x + width/2, industry_vals, width, label='行业平均', color='#e74c3c', alpha=0.8)

    ax.set_xlabel('估值指标', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('倍数', fontproperties=font_prop, fontsize=14)
    ax.set_title('相对估值对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(metrics, fontproperties=font_prop, fontsize=12)
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, axis='y')

    for bars in [bars1, bars2]:
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                    f'{height:.1f}', ha='center', va='bottom', fontsize=12, fontproperties=font_prop)

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_metrics_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2: PE对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    peer_names = list(peer_companies_data['name'])
    peer_pes = list(peer_companies_data['pe'])
    peer_pes.append(industry_avg['pe'])
    peer_names.append('行业平均')
    peer_pes.append(current_metrics['pe'])
    peer_names.append('光弘科技')

    colors_pe = ['#95a5a6'] * len(peer_companies_data) + ['#e74c3c', '#3498db']
    bars = ax.bar(range(len(peer_names)), peer_pes, color=colors_pe, alpha=0.7)
    ax.axhline(y=industry_avg['pe'], color='red', linestyle='--', alpha=0.5)
    ax.set_xlabel('公司', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('PE (倍)', fontproperties=font_prop, fontsize=14)
    ax.set_title('PE倍数对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(peer_names)))
    ax.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=11, rotation=45)
    ax.grid(True, alpha=0.3, axis='y')

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_pe_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3: PB对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    peer_pbs = list(peer_companies_data['pb'])
    peer_pbs.append(industry_avg['pb'])
    peer_pbs.append(current_metrics['pb'])

    bars = ax.bar(range(len(peer_names)), peer_pbs, color=colors_pe, alpha=0.7)
    ax.axhline(y=industry_avg['pb'], color='red', linestyle='--', alpha=0.5)
    ax.set_xlabel('公司', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('PB (倍)', fontproperties=font_prop, fontsize=14)
    ax.set_title('PB倍数对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(peer_names)))
    ax.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=11, rotation=45)
    ax.grid(True, alpha=0.3, axis='y')

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_pb_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图4: PS对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    peer_pss = list(peer_companies_data['ps'])
    peer_pss.append(industry_avg['ps'])
    peer_pss.append(current_metrics['ps'])

    bars = ax.bar(range(len(peer_names)), peer_pss, color=colors_pe, alpha=0.7)
    ax.axhline(y=industry_avg['ps'], color='red', linestyle='--', alpha=0.5)
    ax.set_xlabel('公司', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('PS (倍)', fontproperties=font_prop, fontsize=14)
    ax.set_title('PS倍数对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(peer_names)))
    ax.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=11, rotation=45)
    ax.grid(True, alpha=0.3, axis='y')

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_ps_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图5: 目标价格情景（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    net_income = 253532329.85
    net_assets = 4939639031.34
    total_shares = 767460689
    current_price = 23.88

    eps = net_income / total_shares
    bps = net_assets / total_shares

    scenarios_data = []
    scenarios_data.append(['当前估值', current_metrics['pe'], current_metrics['pb'], current_metrics['ps'], current_price, 0])
    target_price_pe = eps * industry_avg['pe']
    return_pe = (target_price_pe - current_price) / current_price * 100
    scenarios_data.append(['PE→行业平均', industry_avg['pe'], current_metrics['pb'], current_metrics['ps'], target_price_pe, return_pe])
    target_price_pb = bps * industry_avg['pb']
    return_pb = (target_price_pb - current_price) / current_price * 100
    scenarios_data.append(['PB→行业平均', current_metrics['pe'], industry_avg['pb'], current_metrics['ps'], target_price_pb, return_pb])
    target_price_avg = (target_price_pe + target_price_pb) / 2
    return_avg = (target_price_avg - current_price) / current_price * 100
    scenarios_data.append(['全面回归', industry_avg['pe'], industry_avg['pb'], industry_avg['ps'], target_price_avg, return_avg])
    min_pe = peer_companies_data['pe'].min()
    min_pb = peer_companies_data['pb'].min()
    target_price_pessimistic = (eps * min_pe + bps * min_pb) / 2
    return_pessimistic = (target_price_pessimistic - current_price) / current_price * 100
    scenarios_data.append(['行业最低估值', min_pe, min_pb, industry_avg['ps'], target_price_pessimistic, return_pessimistic])
    max_pe = peer_companies_data['pe'].max()
    max_pb = peer_companies_data['pb'].max()
    target_price_optimistic = (eps * max_pe + bps * max_pb) / 2
    return_optimistic = (target_price_optimistic - current_price) / current_price * 100
    scenarios_data.append(['行业最高估值', max_pe, max_pb, industry_avg['ps'], target_price_optimistic, return_optimistic])

    df_scenarios = pd.DataFrame(scenarios_data, columns=['情景', 'PE', 'PB', 'PS', '目标价格(元)', '预期收益率(%)'])

    colors_scenario = ['gray', 'blue', 'green', 'orange', 'red', 'green']
    bars = ax.barh(df_scenarios['情景'], df_scenarios['目标价格(元)'], color=colors_scenario, alpha=0.7)
    ax.axvline(x=current_price, color='black', linestyle='--', label='当前价格')
    ax.set_xlabel('目标价格 (元)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('情景', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同估值情景下的目标价格', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, axis='x')
    # 设置y轴刻度位置和标签
    ax.set_yticks(range(len(df_scenarios)))
    ax.set_yticklabels(df_scenarios['情景'], fontproperties=font_prop, fontsize=11)

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_target_price_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图6: 预期收益率情景（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    colors_return = ['green' if r > 0 else 'red' for r in df_scenarios['预期收益率(%)']]
    bars = ax.barh(df_scenarios['情景'], df_scenarios['预期收益率(%)'], color=colors_return, alpha=0.7)
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1)
    ax.set_xlabel('预期收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('情景', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同估值情景下的预期收益率', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.grid(True, alpha=0.3, axis='x')
    # 设置y轴刻度位置和标签
    ax.set_yticks(range(len(df_scenarios)))
    ax.set_yticklabels(df_scenarios['情景'], fontproperties=font_prop, fontsize=11)

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_expected_return_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths, df_scenarios


def generate_stock_market_data_charts_split(market_data, price_data, volatility_data, save_dir):
    """
    生成个股市场数据图表（拆分版）

    参数:
        market_data: 市场数据字典
        price_data: 价格数据（包含close序列和均线）
        volatility_data: 波动率数据
        save_dir: 保存目录

    返回:
        图片路径列表 [股价走势图, 波动率图, 收益率分布图]
    """
    paths = []

    # 1. 股价走势与均线图
    fig, ax = plt.subplots(figsize=(14, 6))

    x_range = range(len(price_data['close']))
    ax.plot(x_range, price_data['close'], label='收盘价', linewidth=1.5, alpha=0.8, color='#2c3e50')

    colors = ['#e74c3c', '#e67e22', '#3498db', '#27ae60', '#9b59b6']
    windows = [20, 60, 120, 250]  # 与调用者提供的数据一致
    for i, window in enumerate(windows):
        ma_key = f'ma_{window}'
        if ma_key in price_data:
            ax.plot(x_range, price_data[ma_key], label=f'MA{window}', color=colors[i], linewidth=1.2, alpha=0.7)

    ax.set_xlabel('交易日', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('价格 (元)', fontproperties=font_prop, fontsize=12)
    ax.set_title(f'{market_data.get("stock_name", "N/A")} - 股价走势与均线',
                 fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, framealpha=0.9)
    ax.grid(True, alpha=0.3)

    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    path1 = os.path.join(save_dir, 'market_data_price_trend.png')
    plt.tight_layout()
    plt.savefig(path1, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path1)

    # 2. 滚动波动率图
    fig, ax = plt.subplots(figsize=(14, 5))

    colors = ['#e67e22', '#3498db', '#27ae60', '#9b59b6']
    windows = [20, 60, 120, 250]  # 与调用者提供的数据一致
    for i, window in enumerate(windows):
        vol_key = f'volatility_{window}d_series'
        if vol_key in volatility_data:
            ax.plot(x_range, volatility_data[vol_key] * 100,
                   label=f'{window}日波动率', color=colors[i], linewidth=1.2, alpha=0.7)

    ax.set_xlabel('交易日', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('滚动波动率走势', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, framealpha=0.9)
    ax.grid(True, alpha=0.3)

    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    path2 = os.path.join(save_dir, 'market_data_volatility.png')
    plt.tight_layout()
    plt.savefig(path2, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path2)

    # 3. 日收益率分布图
    fig, ax = plt.subplots(figsize=(12, 5))

    returns_pct = price_data.get('pct_chg', [])
    if len(returns_pct) > 0:
        # pct_chg 已经是百分比形式（如 2.5 表示 2.5%），不需要再乘以 100
        ax.hist(returns_pct, bins=50, color='#3498db', alpha=0.7, edgecolor='black')

        mean_return = np.mean(returns_pct)
        ax.axvline(x=mean_return, color='#e74c3c', linestyle='--',
                  label=f'平均收益率 {mean_return:.3f}%', linewidth=2)
        ax.axvline(x=0, color='#27ae60', linestyle='-', label='盈亏平衡', linewidth=1.5)

    ax.set_xlabel('日收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('频数', fontproperties=font_prop, fontsize=12)
    ax.set_title('日收益率分布', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, framealpha=0.9)
    ax.grid(True, alpha=0.3, axis='y')

    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    path3 = os.path.join(save_dir, 'market_data_returns_dist.png')
    plt.tight_layout()
    plt.savefig(path3, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path3)

    return paths


def generate_industry_index_charts(industry_data, save_dir):
    """
    生成行业指数图表（走势、波动率、收益率分布）

    参数:
        industry_data: 行业数据字典
        save_dir: 保存目录

    返回:
        图片路径列表 [指数走势图, 波动率图, 收益率分布图]
    """
    import tushare as ts
    from datetime import datetime, timedelta

    paths = []

    try:
        # 重新获取行业指数历史数据用于绘图
        index_code = industry_data.get('index_code')
        if not index_code:
            print("⚠️ 缺少行业指数代码，无法生成图表")
            return paths

        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if not ts_token:
            print("⚠️ 未设置TUSHARE_TOKEN，无法获取行业指数数据")
            return paths

        pro = ts.pro_api(ts_token)

        # 获取最近500天的数据
        end_date = datetime.now().strftime('%Y%m%d')
        start_date = (datetime.now() - timedelta(days=500*2)).strftime('%Y%m%d')

        df = pro.sw_daily(ts_code=index_code, start_date=start_date, end_date=end_date)

        if df is None or df.empty:
            print(f"⚠️ 未获取到{index_code}的历史数据")
            return paths

        # 按日期排序
        df = df.sort_values('trade_date')
        df.reset_index(drop=True, inplace=True)

        # 准备价格数据
        price_data = {'close': df['close'].values}
        pct_col = 'pct_chg' if 'pct_chg' in df.columns else 'pct_change'
        price_data['pct_chg'] = df[pct_col].values

        # 计算均线
        for window in [20, 60, 120, 250]:
            ma = df['close'].rolling(window=window).mean()
            price_data[f'ma_{window}'] = ma.values

        # 准备波动率数据
        volatility_data = {}
        for window in [20, 60, 120, 250]:
            pct_decimal = df[pct_col] / 100.0
            rolling_std = pct_decimal.rolling(window=window).std()
            rolling_vol = rolling_std * np.sqrt(250)
            volatility_data[f'volatility_{window}d_series'] = rolling_vol.values

        industry_name = industry_data.get('industry_name', industry_data.get('sw_l3_name', '行业指数'))

        # 1. 指数走势与均线图
        fig, ax = plt.subplots(figsize=(14, 6))

        x_range = range(len(price_data['close']))
        ax.plot(x_range, price_data['close'], label='收盘点位', linewidth=1.5, alpha=0.8, color='#2c3e50')

        colors = ['#e74c3c', '#e67e22', '#3498db', '#27ae60']
        windows = [20, 60, 120, 250]
        for i, window in enumerate(windows):
            ma_key = f'ma_{window}'
            if ma_key in price_data:
                ax.plot(x_range, price_data[ma_key], label=f'MA{window}', color=colors[i], linewidth=1.2, alpha=0.7)

        ax.set_xlabel('交易日', fontproperties=font_prop, fontsize=12)
        ax.set_ylabel('指数点位', fontproperties=font_prop, fontsize=12)
        ax.set_title(f'{industry_name} - 指数走势与均线',
                     fontproperties=font_prop, fontsize=14, fontweight='bold')
        ax.legend(prop=font_prop, framealpha=0.9)
        ax.grid(True, alpha=0.3)

        for label in ax.get_xticklabels() + ax.get_yticklabels():
            label.set_fontproperties(font_prop)

        path1 = os.path.join(save_dir, 'industry_index_trend.png')
        plt.tight_layout()
        plt.savefig(path1, dpi=150, bbox_inches='tight')
        plt.close()
        paths.append(path1)

        # 2. 滚动波动率图
        fig, ax = plt.subplots(figsize=(14, 5))

        colors = ['#e67e22', '#3498db', '#27ae60', '#9b59b6']
        windows = [20, 60, 120, 250]
        for i, window in enumerate(windows):
            vol_key = f'volatility_{window}d_series'
            if vol_key in volatility_data:
                ax.plot(x_range, volatility_data[vol_key] * 100,
                       label=f'{window}日波动率', color=colors[i], linewidth=1.2, alpha=0.7)

        ax.set_xlabel('交易日', fontproperties=font_prop, fontsize=12)
        ax.set_ylabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
        ax.set_title('行业指数滚动波动率走势', fontproperties=font_prop, fontsize=14, fontweight='bold')
        ax.legend(prop=font_prop, framealpha=0.9)
        ax.grid(True, alpha=0.3)

        for label in ax.get_xticklabels() + ax.get_yticklabels():
            label.set_fontproperties(font_prop)

        path2 = os.path.join(save_dir, 'industry_index_volatility.png')
        plt.tight_layout()
        plt.savefig(path2, dpi=150, bbox_inches='tight')
        plt.close()
        paths.append(path2)

        # 3. 日收益率分布图
        fig, ax = plt.subplots(figsize=(12, 5))

        returns_pct = price_data.get('pct_chg', [])
        if len(returns_pct) > 0:
            ax.hist(returns_pct, bins=50, color='#3498db', alpha=0.7, edgecolor='black')

            mean_return = np.mean(returns_pct)
            ax.axvline(x=mean_return, color='#e74c3c', linestyle='--',
                      label=f'平均收益率 {mean_return:.3f}%', linewidth=2)
            ax.axvline(x=0, color='#27ae60', linestyle='-', label='盈亏平衡', linewidth=1.5)

        ax.set_xlabel('日收益率 (%)', fontproperties=font_prop, fontsize=12)
        ax.set_ylabel('频数', fontproperties=font_prop, fontsize=12)
        ax.set_title('行业指数日收益率分布', fontproperties=font_prop, fontsize=14, fontweight='bold')
        ax.legend(prop=font_prop, framealpha=0.9)
        ax.grid(True, alpha=0.3, axis='y')

        for label in ax.get_xticklabels() + ax.get_yticklabels():
            label.set_fontproperties(font_prop)

        path3 = os.path.join(save_dir, 'industry_index_returns_dist.png')
        plt.tight_layout()
        plt.savefig(path3, dpi=150, bbox_inches='tight')
        plt.close()
        paths.append(path3)

        print(f"✅ 成功生成 {len(paths)} 个行业指数图表")

    except Exception as e:
        print(f"⚠️ 生成行业指数图表失败: {e}")
        import traceback
        traceback.print_exc()

    return paths


def generate_index_data_charts_split(indices_data, save_dir):
    """
    生成市场指数数据图表（拆分版）

    参数:
        indices_data: 指数数据字典
        save_dir: 保存目录

    返回:
        图片路径列表 [60日波动率, 60日收益率, 60日胜率, 60日技术位置,
                     120日波动率, 120日收益率, 250日波动率, 250日收益率]
    """
    paths = []

    index_names = list(indices_data.keys())

    # 提取60日数据
    volatilities_60d = [indices_data[name].get('volatility_60d', 0) * 100 for name in index_names]
    returns_60d = [indices_data[name].get('return_60d', 0) * 100 for name in index_names]
    win_rates_60d = [indices_data[name].get('win_rate_60d', 0) * 100 for name in index_names]
    ma_distances_60 = [(indices_data[name]['current_level'] /
                     indices_data[name].get('ma_60', 1) - 1) * 100 for name in index_names]

    # 提取120日数据
    volatilities_120d = [indices_data[name].get('volatility_120d', 0) * 100 for name in index_names]
    returns_120d = [indices_data[name].get('return_120d', 0) * 100 for name in index_names]
    win_rates_120d = [indices_data[name].get('win_rate_120d', 0) * 100 for name in index_names]

    # 提取250日数据
    volatilities_250d = [indices_data[name].get('volatility_250d', 0) * 100 for name in index_names]
    returns_250d = [indices_data[name].get('return_250d', 0) * 100 for name in index_names]
    win_rates_250d = [indices_data[name].get('win_rate_250d', 0) * 100 for name in index_names]

    # ====== 60日窗口图表 ======
    # 1. 波动率对比图（60日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors1 = plt.cm.Reds(np.linspace(0.4, 0.9, len(index_names)))
    bars1 = ax.barh(index_names, volatilities_60d, color=colors1, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数波动率对比 (60日窗口)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, vol) in enumerate(zip(bars1, volatilities_60d)):
        ax.text(vol + 0.5, i, f'{vol:.1f}%', va='center', fontsize=10,
               fontproperties=font_prop, fontweight='bold')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path1 = os.path.join(save_dir, 'index_volatility_60d.png')
    plt.tight_layout()
    plt.savefig(path1, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path1)

    # 2. 年化收益率对比图（60日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors2 = ['#27ae60' if r > 0 else '#e74c3c' for r in returns_60d]
    bars2 = ax.barh(index_names, returns_60d, color=colors2, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数年化收益率对比 (60日窗口)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1.2)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, ret) in enumerate(zip(bars2, returns_60d)):
        pos = ret + 1.5 if ret > 0 else ret - 2
        ha = 'left' if ret > 0 else 'right'
        ax.text(pos, i, f'{ret:+.1f}%', va='center', ha=ha, fontsize=10,
               fontproperties=font_prop, fontweight='bold', color='white')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path2 = os.path.join(save_dir, 'index_returns_60d.png')
    plt.tight_layout()
    plt.savefig(path2, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path2)

    # 3. 胜率对比图（60日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors3 = plt.cm.Blues(np.linspace(0.4, 0.9, len(index_names)))
    bars3 = ax.barh(index_names, win_rates_60d, color=colors3, alpha=0.85, edgecolor='white')

    ax.set_xlabel('胜率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数胜率对比 (60日窗口)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.set_xlim(0, 100)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, wr) in enumerate(zip(bars3, win_rates_60d)):
        ax.text(wr + 1, i, f'{wr:.1f}%', va='center', fontsize=10,
               fontproperties=font_prop, fontweight='bold')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path3 = os.path.join(save_dir, 'index_winrate_60d.png')
    plt.tight_layout()
    plt.savefig(path3, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path3)

    # 4. 技术位置对比图（60日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors4 = ['#27ae60' if d > 0 else '#e74c3c' for d in ma_distances_60]
    bars4 = ax.barh(index_names, ma_distances_60, color=colors4, alpha=0.85, edgecolor='white')

    ax.set_xlabel('相对MA60的偏离 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数技术位置对比 (60日窗口)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1.2)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, dist) in enumerate(zip(bars4, ma_distances_60)):
        pos = dist + 0.8 if dist > 0 else dist - 0.8
        ha = 'left' if dist > 0 else 'right'
        ax.text(pos, i, f'{dist:+.1f}%', va='center', ha=ha, fontsize=10,
               fontproperties=font_prop, fontweight='bold', color='white')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path4 = os.path.join(save_dir, 'index_technical_60d.png')
    plt.tight_layout()
    plt.savefig(path4, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path4)

    # ====== 120日窗口图表 ======
    # 5. 波动率对比图（120日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors5 = plt.cm.Reds(np.linspace(0.4, 0.9, len(index_names)))
    bars5 = ax.barh(index_names, volatilities_120d, color=colors5, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数波动率对比 (120日窗口/半年线)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, vol) in enumerate(zip(bars5, volatilities_120d)):
        ax.text(vol + 0.5, i, f'{vol:.1f}%', va='center', fontsize=10,
               fontproperties=font_prop, fontweight='bold')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path5 = os.path.join(save_dir, 'index_volatility_120d.png')
    plt.tight_layout()
    plt.savefig(path5, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path5)

    # 6. 年化收益率对比图（120日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors6 = ['#27ae60' if r > 0 else '#e74c3c' for r in returns_120d]
    bars6 = ax.barh(index_names, returns_120d, color=colors6, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数年化收益率对比 (120日窗口/半年线)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1.2)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, ret) in enumerate(zip(bars6, returns_120d)):
        pos = ret + 1.5 if ret > 0 else ret - 2
        ha = 'left' if ret > 0 else 'right'
        ax.text(pos, i, f'{ret:+.1f}%', va='center', ha=ha, fontsize=10,
               fontproperties=font_prop, fontweight='bold', color='white')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path6 = os.path.join(save_dir, 'index_returns_120d.png')
    plt.tight_layout()
    plt.savefig(path6, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path6)

    # ====== 250日窗口图表 ======
    # 7. 波动率对比图（250日）
    fig, ax = plt.subplots(figsize=(20, 10))
    colors7 = plt.cm.Reds(np.linspace(0.4, 0.9, len(index_names)))
    bars7 = ax.barh(index_names, volatilities_250d, color=colors7, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化波动率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('各指数波动率对比 (250日窗口/年线)', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, vol) in enumerate(zip(bars7, volatilities_250d)):
        ax.text(vol + 0.5, i, f'{vol:.1f}%', va='center', fontsize=12,
               fontproperties=font_prop, fontweight='bold')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(13)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(13)

    path7 = os.path.join(save_dir, 'index_volatility_250d.png')
    plt.tight_layout()
    plt.savefig(path7, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path7)

    # 8. 年化收益率对比图（250日）
    fig, ax = plt.subplots(figsize=(20, 10))
    colors8 = ['#27ae60' if r > 0 else '#e74c3c' for r in returns_250d]
    bars8 = ax.barh(index_names, returns_250d, color=colors8, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('各指数年化收益率对比 (250日窗口/年线)', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1.2)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, ret) in enumerate(zip(bars8, returns_250d)):
        pos = ret + 1.5 if ret > 0 else ret - 2
        ha = 'left' if ret > 0 else 'right'
        ax.text(pos, i, f'{ret:+.1f}%', va='center', ha=ha, fontsize=12,
               fontproperties=font_prop, fontweight='bold', color='white')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(13)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(13)

    path8 = os.path.join(save_dir, 'index_returns_250d.png')
    plt.tight_layout()
    plt.savefig(path8, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path8)

    return paths


def generate_radar_chart(scores, save_path):
    """生成风险评分雷达图"""
    categories = list(scores.keys())
    values = list(scores.values())

    # 闭合雷达图
    angles_plot = np.linspace(0, 2*np.pi, len(categories), endpoint=False).tolist()
    angles_plot += angles_plot[:1]
    values_plot = values + values[:1]

    fig, ax = plt.subplots(figsize=(10, 10), subplot_kw=dict(projection='polar'))
    ax.plot(angles_plot, values_plot, 'o-', linewidth=2, color='#3498db')
    ax.fill(angles_plot, values_plot, alpha=0.25, color='#3498db')
    ax.set_xticks(angles_plot[:-1])
    ax.set_xticklabels(categories, fontproperties=font_prop, fontsize=11)
    ax.set_ylim(0, 30)
    ax.set_title('定增项目风险雷达图', fontsize=14, fontweight='bold',
                pad=20, fontproperties=font_prop)
    ax.grid(True)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()
