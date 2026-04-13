"""
第九章：风控建议与风险提示

本章节从风险控制角度，给出盈亏平衡分析、核心指标汇总、最终报价建议和全面的风险提示。
基于保守原则，确保投资决策在合理风险可控范围内进行。
"""

import sys
import os
from datetime import datetime
from docx.shared import Inches

# 添加路径
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, PROJECT_DIR)

from module_utils import add_title, add_paragraph, add_table_data, add_image, add_section_break
from module_utils import generate_break_even_chart, generate_market_turnover_chart
import chapter09_01_evaluation


def calculate_macro_environment_assessment(market_data, document, context=None, industry_cycle_override=None):
    """
    计算市场环境评估得分

    评估维度：
    1. 货币政策（15%权重）：紧缩(40分)/稳健(60分)/适度宽松(80分)/宽松(100分)
    2. 财政政策（15%权重）：紧缩(50分)/稳健(70分)/积极(90分)
    3. 行业发展周期（30%权重）：成长(100分)/成熟(80分)/幼稚(60分)/衰退(40分) - 支持参数指定
    4. 二级市场活跃度（40%权重）：基于历史分位数

    Args:
        market_data: 市场数据字典
        document: Word文档对象（用于添加表格）
        industry_cycle_override: 行业发展周期覆盖参数，可选值："成长"(100)/"成熟"(80)/"幼稚"(60)/"衰退"(40)
                                  如果不指定则默认为"幼稚"(60分)

    Returns:
        dict: 包含total_score, assessment_level, assessment_description等
    """
    import numpy as np

    # ==================== 1. 货币政策评估（15%权重）====================
    # 这里简化处理：基于市场利率和M2增速评估
    # 实际应用中可以根据最新宏观经济数据调整
    current_rdr = market_data.get('risk_free_rate', 0.03)  # 无风险利率作为代理指标

    # 简化的货币政策评分逻辑
    if current_rdr < 0.025:
        monetary_policy_score = 100  # 宽松
        monetary_policy_status = "宽松"
    elif current_rdr < 0.035:
        monetary_policy_score = 80  # 适度宽松
        monetary_policy_status = "适度宽松"
    elif current_rdr < 0.045:
        monetary_policy_score = 60  # 稳健
        monetary_policy_status = "稳健"
    else:
        monetary_policy_score = 40  # 紧缩
        monetary_policy_status = "紧缩"

    # ==================== 2. 财政政策评估（15%权重）====================
    # 基于财政赤字率和基建投资等指标
    # 这里简化处理，默认为"积极"
    fiscal_policy_score = 90  # 积极
    fiscal_policy_status = "积极"

    # ==================== 3. 行业发展周期评估（30%权重）====================
    # 支持参数指定行业发展周期，否则默认为幼稚期（60分）
    if industry_cycle_override is not None:
        # 使用参数指定的行业发展周期
        cycle_mapping = {
            "成长": 100,
            "成熟": 80,
            "幼稚": 60,
            "衰退": 40
        }
        industry_cycle_score = cycle_mapping.get(industry_cycle_override, 60)
        industry_cycle_status = industry_cycle_override
        # 添加说明：使用参数指定
        industry_cycle_note = f"（参数指定）"
    else:
        # 基于行业PE和行业估值水平评估
        industry_pe = market_data.get('industry_pe', 30)
        current_pe = market_data.get('pe_ratio', 30)

        # 行业相对估值：当前PE / 行业PE
        if industry_pe > 0:
            industry_relative_valuation = current_pe / industry_pe
        else:
            industry_relative_valuation = 1.0

        # 基于相对估值判断行业发展周期
        if industry_relative_valuation < 0.7:
            industry_cycle_score = 100  # 成长期（低估）
            industry_cycle_status = "成长"
        elif industry_relative_valuation < 0.9:
            industry_cycle_score = 80  # 成熟期（合理）
            industry_cycle_status = "成熟"
        elif industry_relative_valuation < 1.2:
            industry_cycle_score = 60  # 幼稚期（略高估）
            industry_cycle_status = "幼稚"
        else:
            industry_cycle_score = 40  # 衰退期（高估）
            industry_cycle_status = "衰退"

        # 如果无法判断，给予幼稚期默认分数60分
        if industry_pe == 0 or current_pe == 0:
            industry_cycle_score = 60
            industry_cycle_status = "幼稚"
            industry_cycle_note = f"（默认）"
        else:
            industry_cycle_note = f"（自动判断）"

    # ==================== 4. 二级市场活跃度评估（40%权重）====================
    # 基于历史换手率分位数（最近5年）

    # 获取市场换手率数据
    current_turnover_percentile = None
    if 'market_turnover' in market_data and market_data['market_turnover'] is not None:
        turnover_info = market_data['market_turnover']
        current_turnover_percentile = turnover_info.get('historical_percentile', 50)

    # 根据历史分位数确定活跃度等级和得分
    if current_turnover_percentile is not None:
        if current_turnover_percentile < 20:
            market_activity_score = 20
            market_activity_status = "不活跃"
        elif current_turnover_percentile < 40:
            market_activity_score = 40
            market_activity_status = "较不活跃"
        elif current_turnover_percentile < 60:
            market_activity_score = 60
            market_activity_status = "平稳"
        elif current_turnover_percentile < 80:
            market_activity_score = 80
            market_activity_status = "较活跃"
        else:
            market_activity_score = 100
            market_activity_status = "活跃"

        market_activity_note = f"（历史分位数{current_turnover_percentile:.1f}%）"
    else:
        # 如果没有换手率数据，使用默认值
        market_activity_score = 60
        market_activity_status = "平稳"
        market_activity_note = "（默认值，换手率数据不可用）"

    # ==================== 计算加权总分 ====================
    weights = {
        'monetary_policy': 0.15,
        'fiscal_policy': 0.15,
        'industry_cycle': 0.30,
        'market_activity': 0.40
    }

    weighted_scores = {
        'monetary_policy': monetary_policy_score * weights['monetary_policy'],
        'fiscal_policy': fiscal_policy_score * weights['fiscal_policy'],
        'industry_cycle': industry_cycle_score * weights['industry_cycle'],
        'market_activity': market_activity_score * weights['market_activity']
    }

    total_score = sum(weighted_scores.values())

    # ==================== 评估等级 ====================
    if total_score >= 90:
        assessment_level = "积极"
        assessment_description = "宏观环境良好，市场情绪乐观"
    elif total_score >= 80:
        assessment_level = "适度"
        assessment_description = "宏观环境较好，市场情绪稳定"
    elif total_score >= 70:
        assessment_level = "稳健"
        assessment_description = "宏观环境中性，市场情绪谨慎"
    elif total_score >= 60:
        assessment_level = "偏悲观"
        assessment_description = "宏观环境偏弱，市场情绪低迷"
    else:
        assessment_level = "悲观"
        assessment_description = "宏观环境较差，市场情绪悲观"

    # ==================== 生成评估表格 ====================
    assessment_data = [
        ['评估维度', '当前状态', '得分', '权重', '加权得分', '说明'],
        [
            '货币政策',
            monetary_policy_status,
            f'{monetary_policy_score:.0f}',
            '15%',
            f'{weighted_scores["monetary_policy"]:.1f}',
            '紧缩(40分)/稳健(60分)/适度宽松(80分)/宽松(100分)'
        ],
        [
            '财政政策',
            fiscal_policy_status,
            f'{fiscal_policy_score:.0f}',
            '15%',
            f'{weighted_scores["fiscal_policy"]:.1f}',
            '紧缩(50分)/稳健(70分)/积极(90分)'
        ],
        [
            '行业发展周期',
            f'{industry_cycle_status}{industry_cycle_note if "industry_cycle_note" in locals() else ""}',
            f'{industry_cycle_score:.0f}',
            '30%',
            f'{weighted_scores["industry_cycle"]:.1f}',
            '成长(100分)/成熟(80分)/幼稚(60分)/衰退(40分) - 默认幼稚期，可参数指定'
        ],
        [
            '二级市场活跃度',
            market_activity_status,
            f'{market_activity_score:.1f}',
            '40%',
            f'{weighted_scores["market_activity"]:.1f}',
            f'基于市场换手率历史分位数{market_activity_note}'
        ],
        [
            '合计',
            '',
            '',
            '100%',
            f'{total_score:.1f}',
            assessment_level
        ]
    ]

    add_table_data(document, assessment_data[0], assessment_data[1:])

    # 返回评估结果
    return {
        'total_score': total_score,
        'assessment_level': assessment_level,
        'assessment_description': assessment_description,
        'weighted_scores': weighted_scores,
        'individual_scores': {
            'monetary_policy': monetary_policy_score,
            'fiscal_policy': fiscal_policy_score,
            'industry_cycle': industry_cycle_score,
            'market_activity': market_activity_score
        },
        'individual_statuses': {
            'monetary_policy': monetary_policy_status,
            'fiscal_policy': fiscal_policy_status,
            'industry_cycle': industry_cycle_status,
            'market_activity': market_activity_status
        }
    }


def generate_chapter(context):
    """
    生成第九章和第十章：综合评估 + 投资建议与风险提示

    参数说明：
    - context: 包含所有必要数据的上下文字典
    """
    # 从context中提取变量
    document = context['document']
    project_params = context['project_params']
    market_data = context['market_data']
    IMAGES_DIR = context['IMAGES_DIR']

    # 从context['results']中获取其他章节的计算结果
    all_scenarios_for_appendix = context['results'].get('all_scenarios', [])
    intrinsic_value = context.get('intrinsic_value', 25.0)
    discount_premium = context['results'].get('discount_premium', 0.0)
    issue_type = project_params.get('issue_type', '竞价')

    # 从第五章获取蒙特卡洛结果
    mc_results = context['results'].get('mc_results', {})
    profit_prob = mc_results.get('profit_prob_120d', 0.5)
    mean_return = mc_results.get('mean_return_120d', 0.0)

    # 从第八章获取VaR结果
    var_95 = context['results'].get('var_95', 0.5)
    var_99 = context['results'].get('var_99', 0.7)

    # 其他参数（使用默认值或从context获取）
    total_score = context['results'].get('total_score', 50)
    ma20_mc = market_data.get('ma_20', project_params['current_price'])

    # 蒙特卡洛模拟参数（用于风险评估）
    mc_volatility_120d = market_data.get('volatility_120d', market_data.get('volatility', 0.30))
    mc_drift_120d = market_data.get('annual_return_120d', market_data.get('drift', 0.08))

    # 配置常量（使用context中的IMAGES_DIR）
    # IMAGES_DIR已在context中定义

    # ==================== 第九章标题 ====================
    add_title(document, '九、风控建议与风险提示', level=1)

    add_paragraph(document, '本章节从风险控制角度，给出盈亏平衡分析、核心指标汇总、最终报价建议和全面的风险提示。')
    add_paragraph(document, '基于保守原则，确保投资决策在合理风险可控范围内进行。')

    # ==================== 9.1 综合评估汇总 ====================
    # 生成9.1节的内容（通过调用chapter09_01_evaluation模块）
    print("\n 生成第九章第一节：综合评估汇总...")
    context = chapter09_01_evaluation.generate_chapter(context)

    # ==================== 9.2 盈亏平衡分析 ====================
    add_title(document, '9.2 盈亏平衡分析', level=2)
    add_paragraph(document, '本节通过盈亏平衡分析，评估在不同目标收益率和不同退出周期下的安全边际。')

    # ==================== 9.2.1 基于当前锁定期的盈亏平衡分析 ====================
    add_title(document, '9.2.1 基于当前锁定期的盈亏平衡分析', level=3)

    add_paragraph(document, '本节基于当前锁定期，分析不同目标收益率下的盈亏平衡价格。')

    # 计算不同收益率下的盈亏平衡价格
    import numpy as np

    # 定义要显示的收益率档位：10%, 15%, 20%, 25%, 30%, 35%, 40%, 45%
    display_returns = np.array([0.10, 0.15, 0.20, 0.25, 0.30, 0.35, 0.40, 0.45])

    issue_price = project_params['issue_price']
    issue_price_eval = issue_price  # 添加此变量供后续使用
    lockup_years = project_params['lockup_period'] / 12
    current_price_eval = project_params['current_price']

    # 检查发行日是否确定，用于区分"发行日价格"和"发行日价格（拟）"
    is_fixed_issue_date = project_params.get('invitation_date_fixed', False)
    if is_fixed_issue_date:
        issue_date = project_params.get('issue_date', '')
        price_label = '发行日价格'
        price_description = f'发行日价格（{issue_date}）'
    else:
        price_label = '发行日价格（拟）'
        price_description = '发行日价格（拟）'

    # 为每个收益率档位计算盈亏平衡价格
    break_even_prices = []
    for target_return in display_returns:
        # 使用复利计算：发行价 × (1 + 年化收益率)^锁定期年数
        be_price = issue_price * (1 + target_return) ** lockup_years
        break_even_prices.append(be_price)

    # 生成表格数据
    be_data = []
    for ret, price in zip(display_returns, break_even_prices):
        # 计算实际需要的收益率（锁定期内）
        actual_return = (price - issue_price) / issue_price
        # 计算离目标的差距：(盈亏平衡价 / 发行日价格) - 1，以发行日价格为基准
        gap_to_target = (price / current_price_eval - 1) * 100
        status = "" if gap_to_target > 0 else ""
        be_data.append([f'{ret*100:.0f}%', f'{actual_return*100:.1f}%', f'{price:.2f}', f'{gap_to_target:+.1f}%', status])

    add_table_data(document, ['期望年化收益率', '实际需要收益率', '盈亏平衡价格(元)', '离目标差距', '状态'], be_data)

    add_paragraph(document, '说明：年化收益率是基于全年计算的收益率，实际需要收益率是锁定期内需要的价格涨幅。')
    add_paragraph(document, f'离目标差距：基于{price_description}计算，正值表示需要上涨才能达到目标，负值表示{price_label}已超过目标有安全边际。')
    add_paragraph(document, '')

    add_paragraph(document, '盈亏平衡分析结论：')
    add_paragraph(document, f'• 发行价格: {issue_price:.2f} 元/股')
    add_paragraph(document, f'• {price_label}: {current_price_eval:.2f} 元/股')
    add_paragraph(document, f'• 锁定期: {project_params["lockup_period"]}个月（{lockup_years:.2f}年）')

    # 20%收益率对应的索引是2（display_returns中：10%, 15%, 20%, ...）
    be_20 = break_even_prices[2]
    add_paragraph(document, f'• 20%年化收益率要求下盈亏平衡价: {be_20:.2f} 元')

    # 以发行日价格为基准计算差距
    gap_to_target = (be_20 / current_price_eval - 1) * 100
    if gap_to_target > 0:
        add_paragraph(document, f'•  {price_label}需要上涨{gap_to_target:.1f}%才能达到20%收益率目标')
    else:
        add_paragraph(document, f'•  {price_label}已超过20%收益率目标{abs(gap_to_target):.1f}%，有安全边际')

    
    # ==================== 9.2.2 考虑退出周期的盈亏平衡分析 ====================
    add_title(document, '9.2.2 考虑退出周期的盈亏平衡分析', level=3)

    add_paragraph(document, '本节考虑不同的退出周期（锁定期），计算在不同周期下的盈亏平衡价格。')
    add_paragraph(document, '分析考虑资金的时间成本，年化资金成本统一按4%计算。')
    add_paragraph(document, '注：盈亏平衡价使用复利计算：发行价 × (1 + 年化收益率)^锁定期年数')

    # 期望收益率和资金成本
    net_return = 0.08  # 净收益率（期望收益率）年化8%，已扣除资金成本
    cost_of_capital = 0.04  # 资金成本年化4%
    total_return = net_return + cost_of_capital  # 总收益率要求12%

    add_paragraph(document, ' 计算参数：', bold=True)
    add_paragraph(document, f'• 期望收益率：{net_return*100:.0f}%（年化，已扣除资金成本）')
    add_paragraph(document, f'• 资金成本：{cost_of_capital*100:.0f}%（年化）')
    add_paragraph(document, f'• 总收益率要求：{total_return*100:.0f}%（年化）')
    add_paragraph(document, f'• 当前发行价：{issue_price:.2f} 元/股')
    add_paragraph(document, f'• 当前价格：{current_price_eval:.2f} 元/股')
    add_paragraph(document, '')

    # 按照3个月一期，从6个月到24个月
    lockup_periods = [6, 9, 12, 15, 18, 21, 24]  # 单位：月
    period_analysis = []

    for months in lockup_periods:
        years = months / 12

        # 计算考虑资金成本后的盈亏平衡价
        # 公式：盈亏平衡价 = 发行价 × (1 + 总收益率要求)^年数
        # 总收益率要求 = 净收益率（期望收益率）+ 资金成本

        # 方案1：考虑资金成本和时间价值
        be_price_with_cost = issue_price * (1 + total_return) ** years

        # 方案2：传统方法（仅考虑净收益率的时间价值）
        # 传统方法下，盈亏平衡价应该考虑锁定期内的净收益增长
        be_price_traditional = issue_price * (1 + net_return) ** years

        # 距离当前价的差距
        distance_with_cost = (current_price_eval - be_price_with_cost) / current_price_eval * 100
        distance_traditional = (current_price_eval - be_price_traditional) / current_price_eval * 100

        period_analysis.append({
            'months': months,
            'years': years,
            'be_price_with_cost': be_price_with_cost,
            'be_price_traditional': be_price_traditional,
            'distance_with_cost': distance_with_cost,
            'distance_traditional': distance_traditional
        })

    # 生成表格数据
    period_data = [
        ['退出周期', '盈亏平衡价(考虑资金成本)', '盈亏平衡价(传统)', '距离当前价(考虑成本)', '距离当前价(传统)', '评估'],
    ]

    for item in period_analysis:
        months = item['months']
        distance_with_cost = item['distance_with_cost']
        distance_traditional = item['distance_traditional']

        # 评估状态
        if distance_with_cost > 0:
            eval_text = "安全边际充足"
        elif distance_with_cost > -5:
            eval_text = "安全边际有限"
        else:
            eval_text = "安全边际不足"

        period_data.append([
            f'{months}个月',
            f'{item["be_price_with_cost"]:.2f}元',
            f'{item["be_price_traditional"]:.2f}元',
            f'{distance_with_cost:+.1f}%',
            f'{distance_traditional:+.1f}%',
            eval_text
        ])

    add_table_data(document, period_data[0], period_data[1:])

    add_paragraph(document, ' 退出周期分析结论：', bold=True)
    add_paragraph(document, '说明：')
    add_paragraph(document, f'• 考虑资金成本方法：盈亏平衡价 = 发行价 × (1 + {total_return*100:.0f}% × 年数)')
    add_paragraph(document, f'• 传统方法：盈亏平衡价 = 发行价 × (1 + {net_return*100:.0f}% × 年数)（仅考虑期望收益率的时间价值）')
    add_paragraph(document, f'• 总收益率要求{total_return*100:.0f}% = 期望收益率{net_return*100:.0f}% + 资金成本{cost_of_capital*100:.0f}%')

    # 分析不同退出周期的影响
    add_paragraph(document, ' 退出周期影响分析：', bold=True)

    # 找出最佳退出周期
    best_period = max(period_analysis, key=lambda x: x['distance_with_cost'])
    add_paragraph(document, f'• 最佳退出周期：{best_period["months"]}个月（{best_period["years"]:.2f}年）')
    add_paragraph(document, f'  盈亏平衡价：{best_period["be_price_with_cost"]:.2f}元（考虑资金成本）')
    add_paragraph(document, f'  距离当前价：{best_period["distance_with_cost"]:+.1f}%')

    if best_period['distance_with_cost'] > 0:
        add_paragraph(document, f'   在{best_period["months"]}个月退出时，当前价格{distance_with_cost:+.1f}%高于盈亏平衡价')
    else:
        add_paragraph(document, f'   在{best_period["months"]}个月退出时，当前价格{distance_with_cost:.1f}%低于盈亏平衡价')

    add_paragraph(document, '')

    # 对比当前锁定期
    current_lockup_months = project_params['lockup_period']
    current_result = next((item for item in period_analysis if item['months'] == current_lockup_months), None)

    if current_result:
        add_paragraph(document, f'• 当前锁定期（{current_lockup_months}个月）：')
        add_paragraph(document, f'  盈亏平衡价（考虑资金成本）：{current_result["be_price_with_cost"]:.2f}元')
        add_paragraph(document, f'  距离当前价：{current_result["distance_with_cost"]:+.1f}%')

        if current_result['distance_with_cost'] > 0:
            add_paragraph(document, f'   当前价格高于盈亏平衡价{current_result["distance_with_cost"]:.1f}%，具有安全边际')
        else:
            add_paragraph(document, f'   当前价格低于盈亏平衡价{abs(current_result["distance_with_cost"]):.1f}%，安全边际不足')
    add_paragraph(document, ' 投资建议：', bold=True)
    add_paragraph(document, '• 退出周期越长，资金成本越高，要求的盈亏平衡价也越高')
    add_paragraph(document, '• 建议选择退出周期较短（6-12个月）的项目，以降低资金成本压力')
    add_paragraph(document, '• 如果当前价格低于考虑资金成本的盈亏平衡价，建议要求更高的折价率')

    # 生成并插入盈亏平衡价格敏感性图表
    break_even_chart_path = os.path.join(IMAGES_DIR, '09_break_even_analysis.png')
    generate_break_even_chart(issue_price, current_price_eval, project_params['lockup_period'], break_even_chart_path)
    add_paragraph(document, '图表 9.1: 盈亏平衡价格敏感性曲线')
    add_image(document, break_even_chart_path, width=Inches(6))

    # ==================== 9.3 报价方案建议 ====================
    add_title(document, '9.3 报价方案建议', level=2)
    add_paragraph(document, '本节提供不同目标收益率下的报价建议，帮助投资者根据风险偏好选择合适的报价方案。')

    # ==================== 市场环境评估 ====================
    add_paragraph(document, '【市场环境评估】', bold=True)
    add_paragraph(document, '在进行报价方案建议之前，首先评估当前宏观环境，为情景选择提供参考依据。')
    add_paragraph(document, '')

    # 计算宏观环境评分
    macro_assessment = calculate_macro_environment_assessment(market_data, document, context)

    # 保存市场环境评估结果到context，供后续情景选择使用
    context['macro_environment_assessment'] = macro_assessment

    # 显示市场环境评估结论
    total_score = macro_assessment['total_score']
    assessment_level = macro_assessment['assessment_level']

    add_paragraph(document, f'综合评分：{total_score:.1f}分 | 评估等级：{assessment_level}')

    # 根据评估等级给出情景选择建议
    if total_score >= 90:
        add_paragraph(document, f' 宏观环境"积极"（{total_score:.1f}分）：建议溢价率和波动率选择第一档（乐观），漂移率结合公司基本面选择')
    elif total_score >= 80:
        add_paragraph(document, f' 宏观环境"适度"（{total_score:.1f}分）：建议溢价率和波动率选择第二档（中性偏乐观），漂移率需结合公司基本面判断')
    elif total_score >= 70:
        add_paragraph(document, f' 宏观环境"稳健"（{total_score:.1f}分）：建议溢价率和波动率选择第三档（中性），漂移率结合公司基本面选择偏保守一档')
    elif total_score >= 60:
        add_paragraph(document, f'⚠ 宏观环境"偏悲观"（{total_score:.1f}分）：建议溢价率和波动率选择第四档（谨慎），漂移率结合公司基本面选择保守档')
    else:
        add_paragraph(document, f' 宏观环境"悲观"（{total_score:.1f}分）：建议溢价率和波动率选择第五档（极度保守），漂移率结合公司基本面选择最保守档或谨慎参与')

    # 添加公司基本面判断说明
    add_paragraph(document, '')
    add_paragraph(document, '【公司基本面判断说明】', bold=True)
    add_paragraph(document, '漂移率选择需结合公司基本面判断（经营周期）：')
    add_paragraph(document, '• 优质期（业绩增长稳定、行业地位领先）：漂移率范围 +10% 至 +30%')
    add_paragraph(document, '• 成熟期（业绩稳定、行业地位稳固）：漂移率范围 0% 至 +10%')
    add_paragraph(document, '• 衰退期（业绩下滑、行业地位下降）：漂移率范围 -10% 至 0%')
    add_paragraph(document, '• 极度保守期：漂移率范围 -30% 至 -10%')
    add_paragraph(document, '后续将根据经营周期对应的漂移率范围，筛选符合条件的情景。')
    # 参数档位说明
    add_paragraph(document, '【参数档位说明】', bold=True)
    add_paragraph(document, '溢价率档位（从优到差）：')
    add_paragraph(document, '• 第一档：0%（无折价，最高报价）')
    add_paragraph(document, '• 第二档：-5%（小幅折价）')
    add_paragraph(document, '• 第三档：-10%（中等折价）')
    add_paragraph(document, '• 第四档：-15%（较大折价）')
    add_paragraph(document, '• 第五档：-20%（最大折价，最低报价）')

    add_paragraph(document, '波动率档位（从高到低）：')
    add_paragraph(document, '• 第一档：40%-50%（高波动，风险较高）')
    add_paragraph(document, '• 第二档：30%-40%（中高波动）')
    add_paragraph(document, '• 第三档：20%-30%（中低波动）')
    add_paragraph(document, '• 第四档：10%-20%（低波动，风险较低）')
    add_paragraph(document, '')
    add_paragraph(document, '注：上述溢价率建议基于市场环境评估，具体报价还需结合项目基本面和市场情况综合判断。')
    # 9.3.1 市场环境评估
    add_title(document, '9.3.1 市场环境评估', level=3)
    add_paragraph(document, '在制定报价方案前，先评估当前的宏观环境，包括货币政策与财政政策、行业发展周期、二级市场活跃度三个维度。')

    # 计算二级市场活跃度（使用换手率）
    print(f"正在读取市场换手率数据...")

    # 尝试从market_data中读取
    if 'market_turnover' in market_data and market_data['market_turnover'] is not None:
        turnover_info = market_data['market_turnover']
        weighted_turnover = turnover_info.get('current_turnover', 1.99)
        current_turnover_percentile = turnover_info.get('historical_percentile', 50)
        historical_count = turnover_info.get('historical_count', 0)

        print(f" 成功读取市场换手率数据：")
        print(f"   当前换手率={weighted_turnover:.2f}%")
        print(f"   历史分位数={current_turnover_percentile:.1f}%（基于{historical_count}个交易日）")
    else:
        print(f" 市场换手率数据不可用，使用默认值")
        weighted_turnover = 1.99
        current_turnover_percentile = 50
        historical_count = 0
        turnover_info = {}

    # 根据历史分位数确定活跃度等级和得分
    if current_turnover_percentile < 20:
        market_activity_level = "不活跃"
        market_activity_score = 20
        market_activity_desc = "历史分位数<20%（5年）"
    elif current_turnover_percentile < 40:
        market_activity_level = "较不活跃"
        market_activity_score = 40
        market_activity_desc = f"历史分位数20%-40%（5年，当前{current_turnover_percentile:.1f}%）"
    elif current_turnover_percentile < 60:
        market_activity_level = "平稳"
        market_activity_score = 60
        market_activity_desc = f"历史分位数40%-60%（5年，当前{current_turnover_percentile:.1f}%）"
    elif current_turnover_percentile < 80:
        market_activity_level = "较活跃"
        market_activity_score = 80
        market_activity_desc = f"历史分位数60%-80%（5年，当前{current_turnover_percentile:.1f}%）"
    else:
        market_activity_level = "活跃"
        market_activity_score = 100
        market_activity_desc = f"历史分位数>80%（5年，当前{current_turnover_percentile:.1f}%）"

    current_percentile = current_turnover_percentile

    # 市场环境评估表格
    add_paragraph(document, ' 宏观环境三维度评估：', bold=True)

    # 货币政策和财政政策（固定值）
    monetary_policy = '适度宽松'
    monetary_policy_score = 80
    monetary_policy_weight = 0.15  # 15%
    monetary_policy_weighted = monetary_policy_score * monetary_policy_weight

    fiscal_policy = '积极'
    fiscal_policy_score = 90
    fiscal_policy_weight = 0.15  # 15%
    fiscal_policy_weighted = fiscal_policy_score * fiscal_policy_weight

    # 行业发展周期（初始值：幼稚期60分，可根据实际情况调整）
    industry_cycle = '幼稚期'
    industry_cycle_score = 60
    industry_cycle_weight = 0.30  # 30%
    industry_cycle_weighted = industry_cycle_score * industry_cycle_weight

    macro_env_data = [
        ['评估维度', '当前状态', '得分', '权重', '加权得分', '说明'],
        # 1. 货币政策（权重15%）
        ['货币政策', monetary_policy, f'{monetary_policy_score}', '15%', f'{monetary_policy_weighted:.1f}',
         '紧缩(40分)/稳健(60分)/适度宽松(80分)/宽松(100分)'],
        # 2. 财政政策（权重15%）
        ['财政政策', fiscal_policy, f'{fiscal_policy_score}', '15%', f'{fiscal_policy_weighted:.1f}',
         '紧缩(50分)/稳健(70分)/积极(90分)'],
        # 3. 行业发展周期（权重30%，初始值幼稚期60分）
        ['行业发展周期', industry_cycle, f'{industry_cycle_score}', '30%', f'{industry_cycle_weighted:.1f}',
         '成长(100分)/成熟(80分)/幼稚(60分)/衰退(40分) - 可根据实际情况调整'],
        # 4. 二级市场活跃度（权重40%，自动计算）
        ['二级市场活跃度', market_activity_level, f'{market_activity_score}', '40%', f'{market_activity_score * 0.4:.1f}',
         f'{market_activity_desc}，基于最近5年历史换手率计算'],
    ]

    # 分离表头和数据
    macro_env_headers = macro_env_data[0]
    macro_env_table_data = macro_env_data[1:]
    add_table_data(document, macro_env_headers, macro_env_table_data)

    # 计算宏观环境总分
    macro_total_score = monetary_policy_weighted + fiscal_policy_weighted + industry_cycle_weighted + (market_activity_score * 0.4)

    # 将总分保存到context中，供后续筛选使用
    context['macro_score'] = macro_total_score

    add_paragraph(document, f'• 宏观环境总分：{macro_total_score:.1f}分（满分100分）')

    # 评估说明
    add_paragraph(document, '评估说明：', bold=True)
    add_paragraph(document, f'• 宏观环境总分：{macro_total_score:.1f}分（满分100分）')
    add_paragraph(document, '• 评估周期：基于最近5年历史数据计算分位数')
    add_paragraph(document, '• 权重分配：货币政策(15%) + 财政政策(15%) + 行业发展周期(30%) + 二级市场活跃度(40%)')
    add_paragraph(document, '• 得分标准：')
    add_paragraph(document, '  - 货币政策：紧缩(40分)、稳健(60分)、适度宽松(80分)、宽松(100分)')
    add_paragraph(document, '  - 财政政策：紧缩(50分)、稳健(70分)、积极(90分)')
    add_paragraph(document, '  - 行业发展周期：成长(100分)、成熟(80分)、幼稚(60分)、衰退(40分)')
    add_paragraph(document, '  - 二级市场活跃度：根据历史分位数自动计算（20-100分）')

    # 二级市场活跃度详细说明
    add_paragraph(document, ' 二级市场活跃度详情：', bold=True)
    add_paragraph(document, f'• 当前状态：{market_activity_level}（得分{market_activity_score}分）')
    add_paragraph(document, f'• 当前换手率：{weighted_turnover:.2f}%（120日中位数）')
    add_paragraph(document, f'• 历史分位数：{current_percentile:.1f}%（最近{historical_count}个交易日）')
    add_paragraph(document, '• 数据来源：深圳市场换手率（代表全市场，tushare daily_info接口）')
    add_paragraph(document, '')

    # 生成并添加换手率曲线图
    if 'market_turnover' in market_data and market_data['market_turnover'] is not None:
        turnover_chart_path = os.path.join(IMAGES_DIR, '09_market_turnover_history.png')
        print(f"正在生成市场换手率曲线图...")
        chart_path = generate_market_turnover_chart(market_data['market_turnover'], turnover_chart_path)
        if chart_path and os.path.exists(chart_path):
            add_paragraph(document, '图表 9.3: 市场换手率历史走势')
            add_image(document, chart_path, width=Inches(6.0))
            add_paragraph(document, '')
            add_paragraph(document, '图表说明：')
            add_paragraph(document, f'• 蓝色曲线：历史换手率走势（共{historical_count}个交易日）')
            add_paragraph(document, f'• 红色虚线：当前120日中位数（{weighted_turnover:.2f}%），代表近期市场活跃度水平')
            add_paragraph(document, f'• 绿色点线：最新一日换手率（{turnover_info.get("latest_turnover", 0):.2f}%），反映当日市场情绪')
            add_paragraph(document, f'• 当前分位数：{current_percentile:.1f}%，表示当前活跃度在历史中的相对位置')

    # 活跃度对定增的影响
    add_paragraph(document, ' 活跃度对定增的影响：', bold=True)
    if market_activity_score >= 80:
        add_paragraph(document, f'•  市场较活跃，流动性充足，有利于定增项目发行和退出')
        add_paragraph(document, f'• 投资者情绪较高，可适当提高报价预期')
    elif market_activity_score >= 60:
        add_paragraph(document, f'•  市场平稳，流动性适中，定增项目正常发行')
        add_paragraph(document, f'• 建议按照标准折价率报价')
    else:
        add_paragraph(document, f'•  市场不活跃，流动性偏紧，定增发行难度增加')
        add_paragraph(document, f'• 建议要求更高折价率或谨慎参与')

    add_paragraph(document, '')

    # 市场环境评分参数映射表（用于参数构造场景585种）
    add_paragraph(document, '市场环境评分参数映射表（参数构造场景）：', bold=True)
    add_paragraph(document, '本表定义了不同宏观评分下对应的漂移率和波动率参数范围，用于585种参数构造场景的筛选。')
    add_paragraph(document, '')

    # 创建映射表数据
    mapping_table_data = [
        ['80-100分', '市场环境较好', '25% - 35%', '40% - 50%', '高分环境：高漂移率，高波动率'],
        ['75-80分', '市场环境良好', '20% - 30%', '35% - 45%', '较高分环境：较高漂移率和波动率'],
        ['70-75分', '市场环境中上', '15% - 25%', '30% - 40%', '中上分环境：中等偏上参数'],
        ['65-70分', '市场环境中等', '10% - 20%', '25% - 35%', '中等分环境：中等参数'],
        ['60-65分', '市场环境中下', '5% - 15%', '20% - 30%', '中下分环境：中等偏下参数'],
        ['50-60分', '市场环境偏差', '0% - 10%', '18% - 28%', '较低分环境：较低参数'],
        ['40-50分', '市场环境较差', '-5% - 5%', '15% - 25%', '低分环境：低参数'],
        ['<40分', '市场环境很差', '-10% - 0%', '12% - 22%', '极低分环境：极低参数']
    ]

    add_table_data(document, ['宏观评分范围', '市场环境', '漂移率范围', '波动率范围', '说明'], mapping_table_data)
    add_paragraph(document, '')

    # 9.3.2 报价方案建议
    add_title(document, '9.3.2 报价方案建议', level=3)
    add_paragraph(document, '本节基于历史收益率和情景分析，提供多角度的报价建议。')

    # 定义市场环境评分映射函数（供整个9.3.2节使用）
    def get_drift_volatility_mapping(macro_score):
        """
        根据宏观评分获取对应的漂移率和波动率范围
        返回：(drift_min, drift_max, volatility_min, volatility_max)
        """
        if macro_score >= 80:
            # 高分环境：高漂移率，高波动率
            return (0.25, 0.35, 0.40, 0.50)
        elif macro_score >= 75:
            # 较高分环境
            return (0.20, 0.30, 0.35, 0.45)
        elif macro_score >= 70:
            return (0.15, 0.25, 0.30, 0.40)
        elif macro_score >= 65:
            return (0.10, 0.20, 0.25, 0.35)
        elif macro_score >= 60:
            return (0.05, 0.15, 0.20, 0.30)
        elif macro_score >= 50:
            return (0.00, 0.10, 0.18, 0.28)
        elif macro_score >= 40:
            return (-0.05, 0.05, 0.15, 0.25)
        else:
            # 低分环境：低漂移率，低波动率
            return (-0.10, 0.00, 0.12, 0.22)

    add_title(document, '9.3.2.1 反向推算最高报价', level=4)

    # 处理报价日信息
    # 优先使用项目参数中的具体报价日，如果没有则使用当前日（遇节假日往前推）
    if 'issue_date' in project_params and project_params['issue_date']:
        # 如果项目中有具体的报价日
        issue_date_input = project_params['issue_date']
        if isinstance(issue_date_input, str):
            issue_date_display = issue_date_input  # 已经是字符串格式
        else:
            issue_date_display = issue_date_input.strftime('%Y年%m月%d日')
        date_source = "指定报价日"
    else:
        # 如果没有具体报价日，使用当前日
        today = datetime.now()
        issue_date_display = today.strftime('%Y年%m月%d日')
        date_source = "当前日（遇节假日往前推）"

    add_paragraph(document, f'• 报价日期：{issue_date_display}（{date_source}）')

    # 计算参数
    target_return = 0.08  # 目标收益率8%
    lockup_years = project_params['lockup_period'] / 12

    # 获取市场环境评分
    macro_score = context.get('macro_score', 83.5)

    # 使用已有的映射函数获取波动率范围
    drift_min, drift_max, vol_min, vol_max = get_drift_volatility_mapping(macro_score)
    # 使用波动率范围的中值
    environment_volatility = (vol_min + vol_max) / 2

    current_price_eval = project_params['current_price']
    ma20_price = market_data.get('ma_20', current_price_eval)

    # 方法1：简单反推（保守）- 不考虑波动率
    # 使用复利：最高发行价 = 当前价格 / (1 + 目标收益率)^锁定期年数
    max_issue_price_simple = current_price_eval / (1 + target_return) ** lockup_years
    premium_to_current_simple = (max_issue_price_simple - current_price_eval) / current_price_eval * 100
    premium_to_ma20_simple = (max_issue_price_simple - ma20_price) / ma20_price * 100

    # 方法2：考虑市场环境波动率的反推（使用安全边际）
    safety_margin = 0.5 * environment_volatility ** 2
    adjusted_return = target_return + safety_margin
    # 使用复利：最高发行价 = 当前价格 / (1 + 调整后收益率)^锁定期年数
    max_issue_price_adjusted = current_price_eval / (1 + adjusted_return) ** lockup_years
    premium_to_current_adjusted = (max_issue_price_adjusted - current_price_eval) / current_price_eval * 100
    premium_to_ma20_adjusted = (max_issue_price_adjusted - ma20_price) / ma20_price * 100

    # 执行逻辑说明
    add_paragraph(document, '执行逻辑说明：', bold=True)
    add_paragraph(document, '基于目标收益率8%，使用以下逻辑反推最高可接受溢价率：')
    add_paragraph(document, f'1. 设定目标收益率：8%（年化）')
    add_paragraph(document, f'2. 市场环境评分：{macro_score:.1f}分，对应波动率范围：{vol_min*100:.0f}%-{vol_max*100:.0f}%（中值：{environment_volatility*100:.1f}%）')
    add_paragraph(document, f'3. 锁定期：{project_params["lockup_period"]}个月（{lockup_years:.2f}年）')
    add_paragraph(document, f'4. 报价日期：{issue_date_display}')
    add_paragraph(document, f'5. 参考价格：当前价格{current_price_eval:.2f}元（报价日价格），MA20价格{ma20_price:.2f}元')

    add_paragraph(document, '计算方法：', bold=True)
    add_paragraph(document, '• 方法1（保守）：简单反推，不考虑波动率')
    add_paragraph(document, f'  当前价格 = {current_price_eval:.2f}元（报价日价格）')
    add_paragraph(document, '  公式：最高发行价 = 当前价格 ÷ (1 + 目标收益率)^锁定期年数')
    add_paragraph(document, f'  最高发行价 = {current_price_eval:.2f} ÷ (1 + 8%)^{lockup_years:.2f} = {max_issue_price_simple:.2f}元')
    add_paragraph(document, f'  对应名义溢价率（相对MA20）：{premium_to_ma20_simple:+.2f}%（MA20={ma20_price:.2f}元）')
    add_paragraph(document, f'  对应实际溢价率（相对当前价）：{premium_to_current_simple:+.2f}%')
    add_paragraph(document, '')

    add_paragraph(document, '• 方法2（更保守）：考虑市场环境波动率安全边际')
    add_paragraph(document, f'  当前价格 = {current_price_eval:.2f}元（报价日价格）')
    add_paragraph(document, f'  市场环境波动率 = {environment_volatility*100:.1f}%（基于评分{macro_score:.1f}分的映射值）')
    add_paragraph(document, f'  安全边际 = 0.5 × 波动率² = 0.5 × ({environment_volatility*100:.1f}%)² = {safety_margin*100:.2f}%')
    add_paragraph(document, f'  调整后收益率 = 目标收益率 + 安全边际 = 8% + {safety_margin*100:.2f}% = {adjusted_return*100:.2f}%')
    add_paragraph(document, f'  最高发行价 = {current_price_eval:.2f} ÷ (1 + {adjusted_return*100:.2f}%)^{lockup_years:.2f} = {max_issue_price_adjusted:.2f}元')
    add_paragraph(document, f'  对应名义溢价率（相对MA20）：{premium_to_ma20_adjusted:+.2f}%（MA20={ma20_price:.2f}元）')
    add_paragraph(document, f'  对应实际溢价率（相对当前价）：{premium_to_current_adjusted:+.2f}%')
    add_paragraph(document, '')

    # 推荐方案对比表
    add_paragraph(document, '推荐方案对比：', bold=True)

    pricing_options = [
        ['计算方法', '最高发行价', '实际溢价率（相对报价日价格）', '名义溢价率（相对MA20）', '特点'],
        [
            '方法1：简单反推',
            f'{max_issue_price_simple:.2f}元',
            f'{premium_to_current_simple:+.2f}%',
            f'{premium_to_ma20_simple:+.2f}%',
            '不考虑波动率，相对激进'
        ],
        [
            '方法2：考虑波动率安全边际',
            f'{max_issue_price_adjusted:.2f}元',
            f'{premium_to_current_adjusted:+.2f}%',
            f'{premium_to_ma20_adjusted:+.2f}%',
            '更保守，推荐使用'
        ],
    ]

    add_table_data(document, pricing_options[0], pricing_options[1:])

    # 推荐结论
    add_paragraph(document, '')
    add_paragraph(document, '推荐结论：', bold=True)

    # 选择更保守的方法2
    recommended_price = max_issue_price_adjusted
    premium_to_ma20 = premium_to_ma20_adjusted  # 名义溢价率（相对MA20）
    premium_to_current = premium_to_current_adjusted  # 实际溢价率（相对报价日价格）

    # 保存反向推算结果到context，供后续章节使用
    if 'results' not in context:
        context['results'] = {}
    context['results']['reverse_calculation'] = {
        'max_issue_price_adjusted': max_issue_price_adjusted,
        'premium_to_ma20_adjusted': premium_to_ma20_adjusted,
        'premium_to_current_adjusted': premium_to_current_adjusted,
        'current_price_eval': current_price_eval,
        'recommended_price': recommended_price,
        'premium_to_ma20': premium_to_ma20,
        'premium_to_current': premium_to_current
    }
    print(f"调试：反向推算结果已保存到context - 溢价率: {premium_to_ma20_adjusted:.2f}%")

    add_paragraph(document, '• 推荐使用方法2（考虑市场环境波动率安全边际）')
    add_paragraph(document, f'• 建议最高报价：不高于{recommended_price:.2f}元/股')
    add_paragraph(document, f'• 名义溢价率（相对MA20）：{premium_to_ma20:+.2f}%（溢价率上限）')
    # 根据是否有确定报价日显示不同的描述
    if date_source == "指定报价日":
        add_paragraph(document, f'• 实际溢价率（相对报价日价格{current_price_eval:.2f}元）：{premium_to_current:+.2f}%')
    else:
        add_paragraph(document, f'• 实际溢价率（相对当前价格{current_price_eval:.2f}元，默认为报价日价格）：{premium_to_current:+.2f}%')
    add_paragraph(document, f'• 该溢价率已考虑市场环境评分{macro_score:.1f}分对应波动率的安全边际')
    add_paragraph(document, f'• 说明：定增名义溢价率最低限为MA20的-20%，反向推算确定上限为{premium_to_ma20:+.2f}%，建议区间为[-20%, {premium_to_ma20:+.2f}%]')
    add_paragraph(document, '')

    # 投资建议
    add_paragraph(document, ' 投资建议：', bold=True)

    # 根据是否有确定报价日显示不同的建议描述
    price_desc = f"报价日价格{current_price_eval:.2f}元" if date_source == "指定报价日" else f"当前价格{current_price_eval:.2f}元（默认为报价日价格）"

    if premium_to_ma20 <= -5:
        add_paragraph(document, f'•  建议报价 - 名义溢价率{premium_to_ma20:+.2f}%，实际溢价率{premium_to_current:+.2f}%（相对{price_desc}），安全边际充足，符合保守投资原则')
    elif premium_to_ma20 <= 0:
        add_paragraph(document, f'•  谨慎报价 - 名义溢价率{premium_to_ma20:+.2f}%，实际溢价率{premium_to_current:+.2f}%（相对{price_desc}），安全边际有限，建议折价至少5%')
    else:
        add_paragraph(document, f'•  不建议报价 - 名义溢价率{premium_to_ma20:+.2f}%为溢价发行，实际溢价率{premium_to_current:+.2f}%（相对{price_desc}），缺乏安全边际')
        add_paragraph(document, '  建议：等待市场时机或选择其他投资机会')

    add_paragraph(document, ' 注意事项：')
    add_paragraph(document, '• 本计算基于目标收益率8%，实际收益可能因市场波动而偏离')
    add_paragraph(document, '• 建议结合其他分析方法（如DCF估值、相对估值）综合判断')
    add_paragraph(document, '• 如果实际发行价高于推荐报价，建议谨慎参与或放弃')

    # ==================== 9.3.2.2 基于蒙特卡洛模拟的情景筛选 ====================
    add_title(document, '9.3.2.2 基于蒙特卡洛模拟的情景筛选', level=4)

    add_paragraph(document, '本节基于第五章5.5节的蒙特卡洛模拟结果，分析不同溢价率水平下的风险收益特征，找出符合严格筛选条件的溢价率区间，为报价决策提供量化依据。')
    add_paragraph(document, '')

    # 从第五章获取溢价率模拟结果
    premium_simulation_results = context['results'].get('premium_simulation_results', [])
    premium_simulation_params = context['results'].get('premium_simulation_params', {})

    if not premium_simulation_results:
        add_paragraph(document, '⚠️ 未找到第五章的溢价率模拟结果，无法进行筛选分析。')
        add_paragraph(document, '   请确认第五章5.5节已正常生成溢价率模拟数据。')
    else:
        # 获取模拟参数信息
        param_source = premium_simulation_params.get('param_source', '未知参数来源')
        drift_source = premium_simulation_params.get('drift', 0)
        vol_source = premium_simulation_params.get('volatility', 0.3)

        # 筛选条件说明
        add_paragraph(document, '蒙特卡洛模拟筛选分析：', bold=True)
        add_paragraph(document, f'本分析基于第五章的{param_source}（漂移率{drift_source*100:.2f}%，波动率{vol_source*100:.2f}%），对21种溢价率（-20%至0%，1%一档）进行筛选，找出符合以下条件的方案：')
        add_paragraph(document, '')

        screening_conditions = [
            '1. 预测中位数收益率 > 8%（年化收益达标）',
            '2. 95% VaR > -30%（最大损失控制在30%以内）',
            '3. 盈利概率 > 70%（盈利可能性较高）',
            '4. 溢价率不设下限（根据市场情况灵活定价）'
        ]

        for condition in screening_conditions:
            add_paragraph(document, condition)

        add_paragraph(document, '')

        # 分析现有结果，检查符合条件的情况
        qualified_mc_rates = []
        mc_analysis_results = []

        print(f"\n从第五章获取蒙特卡洛模拟筛选分析...")
        print(f"  参数来源: {param_source}")
        print(f"  溢价率模拟结果: {len(premium_simulation_results)}个情景")

        for result in premium_simulation_results:
            premium_rate = result['premium_rate']
            # 第五章存储的是年化对数收益率（小数形式）
            median_return_annual = result['median_return'] * 100  # 年化收益率
            var_5_annual = result['percentile_5'] * 100  # 年化VaR
            profit_prob_mc = result['profit_prob']  # 盈利概率

            # 转换为锁定期收益率用于筛选（避免年化VaR超过-100%的问题）
            # 锁定期收益率 = 年化收益率 × (6/12) = 年化收益率 × 0.5
            median_return_lockup = median_return_annual * 0.5
            var_5_lockup = var_5_annual * 0.5

            # 检查是否符合筛选条件（使用锁定期收益率）
            condition_1 = median_return_lockup > 8  # 中位数收益率 > 8%（锁定期）
            condition_2 = var_5_lockup > -30  # VaR损失控制在30%以内（锁定期）
            condition_3 = profit_prob_mc > 70  # 盈利概率 > 70%
            condition_4 = True  # 溢价率不设下限

            is_qualified = condition_1 and condition_2 and condition_3 and condition_4

            mc_analysis_results.append({
                'premium_rate': premium_rate,
                'issue_price': result['issue_price'],
                'profit_prob': profit_prob_mc,
                'median_return_annual': median_return_annual,  # 年化收益率
                'median_return_lockup': median_return_lockup,  # 锁定期收益率
                'var_5_annual': var_5_annual,  # 年化VaR
                'var_5_lockup': var_5_lockup,  # 锁定期VaR
                'is_qualified': is_qualified,
                'conditions': {
                    'return_check': condition_1,
                    'var_check': condition_2,
                    'prob_check': condition_3
                }
            })

            if is_qualified:
                qualified_mc_rates.append(premium_rate)

            # 输出调试信息（仅显示关键节点）
            if premium_rate in [-20, -15, -10, -5, 0]:
                status = "✅符合" if is_qualified else "❌不符合"
                print(f"  溢价率{premium_rate:+3d}%: {status} - 盈利概率{profit_prob_mc:.1f}%, 锁定期中位数收益{median_return_lockup:.2f}%, 锁定期VaR{var_5_lockup:.2f}%")

        # 创建筛选结果表格
        add_paragraph(document, '不同溢价率筛选结果（基于锁定期收益率）：', bold=True)

        mc_screening_headers = [
            '名义溢价率',
            '发行价格(元)',
            '盈利概率(%)',
            '中位数收益(%)',
            '5% VaR(%)',
            '收益率>8%',
            'VaR>-30%',
            '盈利概率>70%',
            '符合条件'
        ]

        mc_screening_data = []
        for result in mc_analysis_results:
            premium_rate = result['premium_rate']
            # 对关键节点进行标记
            if premium_rate in [-20, -15, -10, -5, 0]:
                rate_label = f"★ {premium_rate:+d}%"
            else:
                rate_label = f"  {premium_rate:+d}%"

            # 使用锁定期收益率进行显示和筛选（避免年化VaR超过-100%的问题）
            mc_screening_data.append([
                rate_label,
                f"{result['issue_price']:.2f}",
                f"{result['profit_prob']:.1f}",
                f"{result['median_return_lockup']:.2f}",  # 锁定期收益率
                f"{result['var_5_lockup']:.2f}",  # 锁定期VaR
                "✅" if result['conditions']['return_check'] else "❌",
                "✅" if result['conditions']['var_check'] else "❌",
                "✅" if result['conditions']['prob_check'] else "❌",
                "✅符合" if result['is_qualified'] else "❌不符合"
            ])

        add_table_data(document, mc_screening_headers, mc_screening_data)

        # 添加表格说明
        add_paragraph(document, '')
        add_paragraph(document, '表格说明：', bold=True)
        add_paragraph(document, '• 表格中的收益率和VaR均基于锁定期（6个月）计算，避免年化值超过-100%的问题')
        add_paragraph(document, '• 锁定期收益率 = 年化收益率 × 0.5，更直观反映定增实际投资回报')
        add_paragraph(document, '• 第六章情景分析使用年化收益率，但本节使用锁定期收益率以保持数值合理性')
        add_paragraph(document, '• 筛选条件（收益率>8%、VaR>-30%）均基于锁定期收益率进行判断')
        add_paragraph(document, '• 盈利概率基于锁定期（6个月）的模拟结果计算')

        # 筛选结果分析
        add_paragraph(document, '')
        add_paragraph(document, '筛选结果分析：', bold=True)

        if qualified_mc_rates:
            best_mc_rate = max(qualified_mc_rates, key=lambda x: next(r['median_return_lockup'] for r in mc_analysis_results if r['premium_rate'] == x))
            worst_mc_rate = min(qualified_mc_rates, key=lambda x: next(r['median_return_lockup'] for r in mc_analysis_results if r['premium_rate'] == x))

            add_paragraph(document, f'• 符合条件的溢价率范围：{min(qualified_mc_rates):+d}% 至 {max(qualified_mc_rates):+d}%')
            add_paragraph(document, f'• 推荐溢价率区间：{worst_mc_rate:+d}% 至 {best_mc_rate:+d}%')

            # 获取对应的结果数据
            best_result = next(r for r in mc_analysis_results if r['premium_rate'] == best_mc_rate)
            worst_result = next(r for r in mc_analysis_results if r['premium_rate'] == worst_mc_rate)

            add_paragraph(document, f'  对应发行价区间：{worst_result["issue_price"]:.2f}元 至 {best_result["issue_price"]:.2f}元')
            add_paragraph(document, f'  盈利概率区间：{worst_result["profit_prob"]:.1f}% 至 {best_result["profit_prob"]:.1f}%')
            add_paragraph(document, f'  预期收益区间：{worst_result["median_return_lockup"]:.2f}% 至 {best_result["median_return_lockup"]:.2f}%（锁定期）')
            add_paragraph(document, f'• 共{len(qualified_mc_rates)}个溢价率档次符合严格的筛选条件')
        else:
            # 如果没有完全符合条件的，找出最接近的
            if mc_analysis_results:
                # 按综合得分排序（盈利概率 * 锁定期收益率 * VaR安全性）
                best_overall = max(mc_analysis_results, key=lambda x: x['profit_prob'] * x['median_return_lockup'] * (1 if x['var_5_lockup'] > -30 else 0.5))
                add_paragraph(document, f'• 无溢价率能完全符合所有筛选条件')
                add_paragraph(document, f'• 最接近条件的溢价率：{best_overall["premium_rate"]:+d}%')
                add_paragraph(document, f'  对应发行价：{best_overall["issue_price"]:.2f}元')
                add_paragraph(document, f'  盈利概率：{best_overall["profit_prob"]:.1f}%')
                add_paragraph(document, f'  预期收益：{best_overall["median_return_lockup"]:.2f}%（锁定期）')
                add_paragraph(document, f'  VaR损失：{best_overall["var_5_lockup"]:.2f}%（锁定期）')

                # 分析最接近条件的溢价率不满足哪些条件
                failed_conditions = []
                if not best_overall['conditions']['return_check']:
                    failed_conditions.append(f"中位数收益{best_overall['median_return_lockup']:.2f}%未达到8%要求")
                if not best_overall['conditions']['var_check']:
                    failed_conditions.append(f"VaR损失{best_overall['var_5_lockup']:.2f}%超过-30%限制")
                if not best_overall['conditions']['prob_check']:
                    failed_conditions.append(f"盈利概率{best_overall['profit_prob']:.1f}%未达到70%要求")

                if failed_conditions:
                    add_paragraph(document, f'• 未满足条件：{"；".join(failed_conditions)}')

        add_paragraph(document, '')
        add_paragraph(document, '投资建议：', bold=True)

        if qualified_mc_rates:
            add_paragraph(document, f'• 优先选择符合条件区间内折价率较大的方案（{min(qualified_mc_rates):+d}%至{max(qualified_mc_rates):+d}%），确保安全边际')
            add_paragraph(document, f'• 推荐重点关注{best_mc_rate:+d}%折价率，预期收益最高且符合所有风控要求')
        else:
            if mc_analysis_results:
                best_overall = max(mc_analysis_results, key=lambda x: x['profit_prob'] * x['median_return_lockup'])
                add_paragraph(document, f'• 建议谨慎参与，最优溢价率{best_overall["premium_rate"]:+d}%仍无法完全满足风控要求')
                add_paragraph(document, f'• 如果必须参与，建议要求更高的折价率或等待更好的市场时机')

        add_paragraph(document, f'• 蒙特卡洛模拟基于{param_source}，考虑了市场随机性和波动特征')
        add_paragraph(document, '• 建议结合历史数据场景分析和参数构造场景分析，综合判断最优报价方案')

        print(f" 蒙特卡洛模拟筛选完成：共{len(mc_analysis_results)}个情景，{len(qualified_mc_rates)}个符合条件")

    # ==================== 9.3.2.3 基于历史数据场景的筛选 ====================
    add_title(document, '9.3.2.3 基于历史数据场景的筛选', level=4)

    add_paragraph(document, '本节基于第六章的历史数据场景（市场指数、行业PE、个股PE），通过严格筛选条件，从风险收益平衡角度推荐最优报价方案。')

    # 市场环境条件筛选
    add_paragraph(document, ' 市场环境条件筛选：', bold=True)

    # 获取宏观评分
    macro_score = context.get('macro_score', 60)

    # 根据宏观评分确定档位选择策略
    if macro_score >= 80:
        tier_level = "高档位"
        tier_selection = "选择各历史数据场景的高档位（如75%分位数、高估水平）"
        market_env_desc = f"市场宏观评分{macro_score}分（≥80分），市场环境较好"
    elif macro_score >= 60:
        tier_level = "中档位"
        tier_selection = "选择各历史数据场景的中档位（如50%分位数、中性水平）"
        market_env_desc = f"市场宏观评分{macro_score}分（60-80分），市场环境一般"
    else:
        tier_level = "低档位"
        tier_selection = "选择各历史数据场景的低档位（如25%分位数、低估水平）"
        market_env_desc = f"市场宏观评分{macro_score}分（<60分），市场环境较差"

    add_paragraph(document, f'• {market_env_desc}')
    add_paragraph(document, f'• 档位选择：{tier_selection}')
    add_paragraph(document, f'• 筛选逻辑：根据市场环境评分，从每个历史数据场景的多个档位中选择对应{tier_level}进行分析')
    add_paragraph(document, '')

    # 基础筛选条件
    add_paragraph(document, ' 基础筛选条件（必须全部满足）：', bold=True)
    add_paragraph(document, '1. 预测中位数收益率 > 8%（年化收益达标）')
    add_paragraph(document, '2. 95% VaR < 30%（最大损失控制在30%以内）')
    add_paragraph(document, '3. 盈利概率 > 70%（盈利可能性较高）')
    add_paragraph(document, '4. 溢价率不设下限（根据市场情况灵活定价）')
    add_paragraph(document, '')

    # 选择逻辑
    add_paragraph(document, ' 选择逻辑（从符合条件中优选）：', bold=True)
    add_paragraph(document, '• 优先选择最接近8%收益率的方案（保守原则）')
    add_paragraph(document, '• 同等收益率下选择报价更低的方案')
    add_paragraph(document, '• 同等情况下选择保守程度更高的方案')
    add_paragraph(document, '')

    # 从context获取历史数据情景（第六章6.2-6.5节的情景）
    comprehensive_results = context['results'].get('historical_scenarios_195', [])
    print(f" 第九章9.3.2.3节：获取到{len(comprehensive_results)}个历史数据场景")

    # 确定当前应该选择的档位
    if macro_score >= 80:
        target_tier = "高"
        tier_keywords = ["75%", "高", "upper", "high"]
    elif macro_score >= 60:
        target_tier = "中"
        tier_keywords = ["50%", "中", "middle", "mid"]
    else:
        target_tier = "低"
        tier_keywords = ["25%", "低", "lower", "low"]

    # 收集情景方案（仅筛选6.2-6.5的情景：市场指数、行业PE、个股PE、DCF估值）
    scenario_options = []

    if comprehensive_results:
        # 从comprehensive_results提取所有情景
        for result in comprehensive_results:
            # 兼容两种数据结构：
            # 1. 嵌套结构：{'scenario': {...}, 'median_return': ..., 'profit_prob': ...}
            # 2. 扁平结构：{'name': ..., 'drift_level': ..., 'profit_prob': ..., 'median_return': ...}
            if 'scenario' in result:
                # 嵌套结构
                scenario_obj = result['scenario']
                scenario_name = scenario_obj['name']
                profit_prob = result['profit_prob']
                median_return = result['median_return']
                var_5 = result.get('var_5', result.get('var_95', 0))
                issue_price = result.get('issue_price', scenario_obj.get('issue_price', 0))

                # 只筛选6.2-6.5的情景（市场指数、行业PE、个股PE、DCF估值）
                if not any(keyword in scenario_name for keyword in ['市场指数', '行业PE', '个股PE', 'DCF估值']):
                    continue

                # 根据市场环境评分筛选对应档位的场景
                # 档位信息存储在drift_level和vol_level字段中，不是在名称里
                is_target_tier = False
                drift_level = scenario_obj.get('drift_level', '')
                vol_level = scenario_obj.get('vol_level', '')

                # DCF估值特殊处理：只有1个漂移率，所以没有drift_level档位，只需检查波动率档位
                if scenario_name.startswith('DCF估值'):
                    is_target_tier = (vol_level == target_tier)
                else:
                    # 其他历史数据场景：需要漂移率和波动率档位都匹配
                    if target_tier == "高":
                        # 检查漂移率和波动率是否都为高档位
                        is_target_tier = (drift_level == "高" and vol_level == "高")
                    elif target_tier == "中":
                        # 检查漂移率和波动率是否都为中档位
                        is_target_tier = (drift_level == "中" and vol_level == "中")
                    else:  # 低档位
                        # 检查漂移率和波动率是否都为低档位
                        is_target_tier = (drift_level == "低" and vol_level == "低")

                # 如果不符合目标档位，跳过该场景
                if not is_target_tier:
                    continue

                # 数据类型转换（确保格式一致）
                median_return = median_return * 100  # 从小数转换为百分比
                # 修正：使用var_5（5%分位数）作为95% VaR，表示95%置信水平下的最大损失
                var_95 = var_5 * 100  # 优先使用var_5

                scenario_options.append({
                    'name': scenario_name,
                    'description': scenario_obj.get('description', ''),
                    'median_return': median_return,
                    'profit_prob': profit_prob,
                    'premium_rate': scenario_obj.get('premium_rate', scenario_obj.get('discount', 0)) * 100,  # 转为百分比
                    'var_95': var_95,
                    'issue_price': issue_price,
                    'drift': scenario_obj.get('drift', 0),
                    'volatility': scenario_obj.get('volatility', 0)
                })

            elif 'name' in result and 'median_return' in result and 'profit_prob' in result:
                # 扁平结构（第六章存储的格式）
                scenario_obj = result
                scenario_name = result['name']
                profit_prob = result['profit_prob']
                median_return = result['median_return']
                var_5 = result.get('var_5', result.get('var_95', 0))
                issue_price = result.get('issue_price', 0)

                # 只筛选6.2-6.5的情景（市场指数、行业PE、个股PE、DCF估值）
                if not any(keyword in scenario_name for keyword in ['市场指数', '行业PE', '个股PE', 'DCF估值']):
                    continue

                # 根据市场环境评分筛选对应档位的场景
                # 档位信息存储在drift_level和vol_level字段中，不是在名称里
                is_target_tier = False
                drift_level = scenario_obj.get('drift_level', '')
                vol_level = scenario_obj.get('vol_level', '')

                # DCF估值特殊处理：只有1个漂移率，所以没有drift_level档位，只需检查波动率档位
                if scenario_name.startswith('DCF估值'):
                    is_target_tier = (vol_level == target_tier)
                else:
                    # 其他历史数据场景：需要漂移率和波动率档位都匹配
                    if target_tier == "高":
                        # 检查漂移率和波动率是否都为高档位
                        is_target_tier = (drift_level == "高" and vol_level == "高")
                    elif target_tier == "中":
                        # 检查漂移率和波动率是否都为中档位
                        is_target_tier = (drift_level == "中" and vol_level == "中")
                    else:  # 低档位
                        # 检查漂移率和波动率是否都为低档位
                        is_target_tier = (drift_level == "低" and vol_level == "低")

                # 如果不符合目标档位，跳过该场景
                if not is_target_tier:
                    continue

                # 数据类型转换（确保格式一致）
                median_return = median_return * 100  # 从小数转换为百分比
                # 修正：使用var_5（5%分位数）作为95% VaR，表示95%置信水平下的最大损失
                var_95 = var_5 * 100  # 优先使用var_5

                scenario_options.append({
                    'name': scenario_name,
                    'description': scenario_obj.get('description', ''),
                    'median_return': median_return,
                    'profit_prob': profit_prob,
                    'premium_rate': scenario_obj.get('premium_rate', scenario_obj.get('discount', 0)) * 100,  # 转为百分比
                    'var_95': var_95,
                    'issue_price': issue_price,
                    'drift': scenario_obj.get('drift', 0),
                    'volatility': scenario_obj.get('volatility', 0)
                })

            else:
                # 数据格式不匹配，跳过
                continue

    # 如果有情景方案，进行筛选和统计
    if scenario_options:
        # 按大场景分组（3个历史数据场景：市场指数、行业PE、个股PE）
        # 注意：DCF估值场景在9.3.2.4节单独处理
        scenario_categories = {
            '市场指数': [],
            '行业PE': [],
            '个股PE': []
        }

        # 将情景按大场景分类
        for scenario in scenario_options:
            for category in scenario_categories.keys():
                if category in scenario['name']:
                    scenario_categories[category].append(scenario)
                    break

        # 对每个大场景应用基础筛选条件（优化版）
        # 优化说明：
        # 1. 放开溢价率下限：不设最低折价要求，允许根据市场情况灵活定价
        # 2. 中位数收益率 > 8%：保持年化收益达标要求
        # 3. 损失率控制在30%：95% VaR > -30%（即最大损失不超过30%）
        # 4. 盈利概率 > 70%：提高盈利概率要求，降低投资风险
        qualified_by_category = {}
        for category, scenarios in scenario_categories.items():
            qualified = []
            for scenario in scenarios:
                # 检查四个基础条件（优化后）
                # 注意：premium_rate和median_return已经是百分比形式，直接比较
                condition_1 = True  # 放开溢价率下限限制
                condition_2 = scenario['median_return'] > 8  # 中位数收益率 > 8%
                condition_3 = scenario['var_95'] > -30  # 95% VaR > -30%（损失率控制在30%以内）
                condition_4 = scenario['profit_prob'] > 70  # 盈利概率 > 70%

                # 显示前几个场景的条件检查结果（调试用）
                if len(qualified) == 0:
                    print(f"{category} - 场景: {scenario['name']}")
                    print(f"  条件检查: c1={condition_1}(无溢价限制), c2={condition_2}(中位数收益{scenario['median_return']:.1f}%), c3={condition_3}(VaR{scenario['var_95']:.1f}%), c4={condition_4}(盈利概率{scenario['profit_prob']:.1f}%)")

                if condition_1 and condition_2 and condition_3 and condition_4:
                    qualified.append(scenario)
            qualified_by_category[category] = qualified

        print(f"筛选结果:")
        for category, scenarios in qualified_by_category.items():
            print(f"  {category}: {len(scenarios)}个场景符合条件")

        # 统计结果
        qualified_categories = {cat: scenarios for cat, scenarios in qualified_by_category.items() if scenarios}
        total_qualified_categories = len(qualified_categories)
        total_qualified_scenarios = sum(len(scenarios) for scenarios in qualified_by_category.values())

        # 显示统计结果
        add_paragraph(document, '历史数据场景筛选统计：', bold=True)
        add_paragraph(document, f'• 大场景总数：3个（市场指数、行业PE、个股PE）')
        add_paragraph(document, f'• 符合条件的大场景数：{total_qualified_categories}个')
        add_paragraph(document, f'• 符合条件的总场景数：{total_qualified_scenarios}个')
        add_paragraph(document, '')

        # 显示各场景筛选结果
        for category, scenarios in qualified_by_category.items():
            category_count = len(scenarios)

            if category_count > 0:
                add_paragraph(document, f'{category}：{category_count}个场景符合条件', bold=True)

                # 创建该类别符合条件场景表格
                category_scenarios_data = []
                for scenario in scenarios:
                    category_scenarios_data.append([
                        scenario['name'],
                        f"{scenario.get('drift', 0)*100:.0f}%",
                        f"{scenario.get('volatility', 0)*100:.0f}%",
                        f"{scenario['premium_rate']:+.1f}%",
                        f"{scenario['median_return']:+.1f}%",
                        f"{scenario['var_95']:+.1f}%",
                        f"{scenario['profit_prob']:.1f}%"
                    ])

                add_table_data(document, ['情景方案', '漂移率', '波动率', '溢价率', '中位数收益率', '95% VaR', '盈利概率'], category_scenarios_data)
                add_paragraph(document, '')
            # 删除了"无场景符合条件"的显示

        # 投资建议判断
        add_paragraph(document, '投资建议：', bold=True)
        if total_qualified_categories >= 2:
            # 至少两个大场景符合条件，支持报价
            # 计算各场景溢价率范围的交集（保守策略）
            category_premium_ranges = {}
            for category, scenarios in qualified_by_category.items():
                if scenarios:
                    premiums = [s['premium_rate'] for s in scenarios]
                    category_premium_ranges[category] = {
                        'min': min(premiums),
                        'max': max(premiums),
                        'count': len(premiums)
                    }

            # 计算交集：取各场景范围的最大值作为下限，最小值作为上限
            if len(category_premium_ranges) >= 2:
                # 交集的下限是各范围下限的最大值（最保守的下限）
                intersection_min = max(r['min'] for r in category_premium_ranges.values())
                # 交集的上限是各范围上限的最小值（最保守的上限）
                intersection_max = min(r['max'] for r in category_premium_ranges.values())

                # 如果交集为空（下限>上限），则取并集
                if intersection_min > intersection_max:
                    # 无交集，取并集作为备选方案
                    all_premiums = [s['premium_rate'] for s in all_qualified]
                    min_premium = min(all_premiums)
                    max_premium = max(all_premiums)
                    range_type = "并集"
                    range_note = f"（各场景无交集，取{len(all_qualified)}个场景的并集）"
                else:
                    # 有交集，使用交集
                    min_premium = intersection_min
                    max_premium = intersection_max
                    range_type = "交集"
                    range_note = f"（{total_qualified_categories}个大场景的共同范围）"
            else:
                # 只有一个场景，直接使用其范围
                min_premium = list(category_premium_ranges.values())[0]['min']
                max_premium = list(category_premium_ranges.values())[0]['max']
                range_type = "单场景"
                range_note = f"（仅1个大场景符合条件）"

            total_scenarios = sum(r['count'] for r in category_premium_ranges.values())
            add_paragraph(document, f'• 至少{total_qualified_categories}个大场景符合条件，支持参与报价')
            add_paragraph(document, f'• 指导溢价率：{min_premium:+.2f}% 至 {max_premium:+.2f}%（{range_type}{range_note}）')
            add_paragraph(document, f'• 建议报价范围：溢价率在[{min_premium:+.2f}%, {max_premium:+.2f}%]区间内（折价{abs(min_premium):.1f}%至{abs(max_premium):.1f}%）')
        else:
            # 少于两个大场景符合条件，建议不参与
            add_paragraph(document, f'• 仅{total_qualified_categories}个大场景符合条件（少于2个），建议本项目不参与')
            add_paragraph(document, '• 风险提示：历史数据场景支持不足，建议谨慎考虑或放弃本次定增')

        add_paragraph(document, '')
    else:
        add_paragraph(document, '历史数据场景筛选统计：', bold=True)
        add_paragraph(document, '• 无历史数据场景可供分析')
        add_paragraph(document, '• 原因：第六章情景分析未生成或数据缺失')
        add_paragraph(document, '• 建议：请检查第六章情景分析是否正常生成')

    # ==================== 9.3.2.4 基于DCF估值场景的筛选 ====================
    add_title(document, '9.3.2.4 基于DCF估值场景的筛选', level=4)

    add_paragraph(document, '本节基于DCF内在估值模型的情景分析，从绝对估值角度筛选符合风险收益要求的报价方案。')
    add_paragraph(document, 'DCF估值基于公司基本面分析，通过预测未来现金流并折现计算内在价值，为定增定价提供理论支撑。')
    add_paragraph(document, '')

    # 基础筛选条件
    add_paragraph(document, ' 基础筛选条件（必须全部满足）：', bold=True)
    add_paragraph(document, '1. 预测中位数收益率 > 8%（年化收益达标）')
    add_paragraph(document, '2. 95% VaR < 30%（最大损失控制在30%以内）')
    add_paragraph(document, '3. 盈利概率 > 70%（盈利可能性较高）')
    add_paragraph(document, '4. 溢价率不设下限（根据市场情况灵活定价）')
    add_paragraph(document, '')

    # 从context获取历史数据情景，筛选DCF估值场景
    comprehensive_results = context['results'].get('historical_scenarios_195', [])

    # 筛选DCF估值场景
    dcf_scenario_options = []
    for result in comprehensive_results:
        # 兼容两种数据结构
        if 'scenario' in result:
            scenario_obj = result['scenario']
        else:
            scenario_obj = result

        scenario_name = scenario_obj.get('name', '')

        # 只处理DCF估值场景
        if 'DCF估值' not in scenario_name:
            continue

        # 提取关键指标
        median_return = result.get('median_return', 0)
        profit_prob = result.get('profit_prob', 0)
        var_5 = result.get('var_5', result.get('var_95', 0))
        issue_price = scenario_obj.get('issue_price', 0)

        # 数据类型转换
        median_return = median_return * 100  # 从小数转换为百分比
        var_95 = var_5 * 100  # 95% VaR

        dcf_scenario_options.append({
            'name': scenario_name,
            'description': scenario_obj.get('description', ''),
            'median_return': median_return,
            'profit_prob': profit_prob,
            'premium_rate': scenario_obj.get('premium_rate', scenario_obj.get('discount', 0)) * 100,
            'var_95': var_95,
            'issue_price': issue_price
        })

    if dcf_scenario_options:
        # 对DCF估值场景应用基础筛选条件
        qualified_dcf_scenarios = []
        for scenario in dcf_scenario_options:
            condition_1 = True  # 放开溢价率下限限制
            condition_2 = scenario['median_return'] > 8  # 中位数收益率 > 8%
            condition_3 = scenario['var_95'] > -30  # 95% VaR > -30%
            condition_4 = scenario['profit_prob'] > 70  # 盈利概率 > 70%

            if condition_1 and condition_2 and condition_3 and condition_4:
                qualified_dcf_scenarios.append(scenario)

        # 显示筛选结果
        add_paragraph(document, 'DCF估值场景筛选统计：', bold=True)
        add_paragraph(document, f'• DCF估值场景总数：{len(dcf_scenario_options)}个')
        add_paragraph(document, f'• 符合条件的场景数：{len(qualified_dcf_scenarios)}个')
        add_paragraph(document, '')

        if qualified_dcf_scenarios:
            add_paragraph(document, '符合条件的DCF估值场景：', bold=True)

            # 创建DCF估值场景表格
            dcf_scenarios_data = []
            for scenario in qualified_dcf_scenarios:
                dcf_scenarios_data.append([
                    scenario['name'],
                    f"{scenario.get('drift', 0)*100:.0f}%",
                    f"{scenario.get('volatility', 0)*100:.0f}%",
                    f"{scenario['premium_rate']:+.1f}%",
                    f"{scenario['median_return']:+.1f}%",
                    f"{scenario['var_95']:+.1f}%",
                    f"{scenario['profit_prob']:.1f}%"
                ])

            add_table_data(document, ['情景方案', '漂移率', '波动率', '溢价率', '中位数收益率', '95% VaR', '盈利概率'], dcf_scenarios_data)
            add_paragraph(document, '')

            # 计算DCF估值场景的溢价率区间
            dcf_premiums = [s['premium_rate'] for s in qualified_dcf_scenarios]
            min_dcf_premium = min(dcf_premiums)
            max_dcf_premium = max(dcf_premiums)

            add_paragraph(document, 'DCF估值投资建议：', bold=True)
            add_paragraph(document, f'• DCF估值支持报价：共{len(qualified_dcf_scenarios)}个场景符合条件')
            add_paragraph(document, f'• DCF估值指导溢价率：{min_dcf_premium:+.2f}% 至 {max_dcf_premium:+.2f}%')
            add_paragraph(document, f'• 说明：DCF估值基于公司内在价值，为定增定价提供理论参考')
        # 删除了"无DCF估值场景符合条件"的显示

        add_paragraph(document, '')
    # 删除了无DCF估值场景的提示信息

    add_paragraph(document, '')

    # ==================== 9.3.2.5 基于参数构造场景的筛选 ====================
    add_title(document, '9.3.2.5 基于参数构造场景的筛选', level=4)

    add_paragraph(document, '本节从585种参数构造场景（漂移率13档 × 波动率5档 × 溢价率9档）中筛选符合条件的方案，提供更全面的报价参考。')

    # 基础筛选条件（与历史数据场景保持一致）
    add_paragraph(document, ' 基础筛选条件（必须全部满足）：', bold=True)
    add_paragraph(document, '1. 预测中位数收益率 > 8%（年化收益达标）')
    add_paragraph(document, '2. 95% VaR < 30%（最大损失控制在30%以内）')
    add_paragraph(document, '3. 盈利概率 > 70%（盈利可能性较高）')
    add_paragraph(document, '4. 溢价率不设下限（根据市场情况灵活定价）')
    add_paragraph(document, '')

    # 从context获取多参数构造情景（第六章6.1节的情景）
    comprehensive_results = context['results'].get('multi_param_scenarios_585', [])

    # 第一步：筛选符合基础条件的585种参数构造场景
    qualified_585_scenarios = []
    checked_count = 0

    for result in comprehensive_results:
        # 兼容新旧格式
        scenario_obj = result.get('scenario', result)

        # 只处理多参数情景（名称不带"市场指数"、"行业PE"等前缀的）
        scenario_name = scenario_obj.get('name', '')
        if any(prefix in scenario_name for prefix in ['市场指数', '行业PE', '个股PE', 'DCF估值']):
            continue

        checked_count += 1

        # 获取实际溢价率（百分比形式）
        premium_rate_pct = scenario_obj.get('premium_rate', scenario_obj.get('discount', 0)) * 100
        median_return_pct = result.get('median_return', 0) * 100
        var_5_raw = result.get('var_5', result.get('var_95', 0)) * 100  # 原始值，可能是负值
        var_95_pct = abs(var_5_raw)  # 绝对值，用于筛选条件
        profit_prob_pct = result.get('profit_prob', 0)

        # 检查四个基础条件（与历史数据场景保持一致）
        condition_1 = True  # 放开溢价率下限限制，根据市场情况灵活定价
        condition_2 = median_return_pct > 8  # 中位数收益率 > 8%
        condition_3 = var_5_raw > -30  # VaR损失控制在30%以内
        condition_4 = profit_prob_pct > 70  # 盈利概率 > 70%

        if condition_1 and condition_2 and condition_3 and condition_4:
            qualified_585_scenarios.append({
                'name': scenario_name,
                'drift': scenario_obj['drift'],
                'volatility': scenario_obj['volatility'],
                'premium_rate': premium_rate_pct,
                'issue_price': scenario_obj.get('issue_price', 0),
                'median_return': median_return_pct,
                'var_95': var_5_raw,
                'profit_prob': profit_prob_pct
            })

    if qualified_585_scenarios:
        add_paragraph(document, f'参数构造场景筛选结果：', bold=True)
        add_paragraph(document, f'• 从585种参数构造场景中找到{len(qualified_585_scenarios)}个符合基础条件的方案')

        # 第二步：根据市场环境评分进行参数映射
        add_paragraph(document, '根据市场环境评分进行参数映射：', bold=True)
        add_paragraph(document, f'• 参考9.3.1节的市场环境评分参数映射表')
        add_paragraph(document, '')

        add_paragraph(document, '当前项目参数映射结果：', bold=True)

        # 使用前面定义的映射函数获取当前项目的参数范围
        drift_min, drift_max, vol_min, vol_max = get_drift_volatility_mapping(macro_score)

        add_paragraph(document, f'• 当前宏观评分：{macro_score}分')
        add_paragraph(document, f'• 对应漂移率范围：{drift_min*100:.0f}% - {drift_max*100:.0f}%')
        add_paragraph(document, f'• 对应波动率范围：{vol_min*100:.0f}% - {vol_max*100:.0f}%')
        add_paragraph(document, '')

        # 第三步：从所有符合基础条件的场景中筛选符合市场环境参数的场景
        environment_matched_scenarios = []
        for scenario in qualified_585_scenarios:
            drift = scenario['drift']
            volatility = scenario['volatility']

            # 检查是否在对应范围内
            if drift_min <= drift <= drift_max and vol_min <= volatility <= vol_max:
                environment_matched_scenarios.append(scenario)

        add_paragraph(document, f'环境匹配筛选结果：', bold=True)
        if environment_matched_scenarios:
            add_paragraph(document, f'• 从{len(qualified_585_scenarios)}个符合基础条件的场景中找到{len(environment_matched_scenarios)}个符合当前市场环境的方案')

            # 按溢价率从高到低排序
            environment_matched_scenarios.sort(key=lambda x: x['premium_rate'], reverse=True)

            add_paragraph(document, '')
            add_paragraph(document, '参数构造场景环境匹配方案详情：', bold=True)
            add_paragraph(document, f'以下{len(environment_matched_scenarios)}个方案同时符合基础条件和当前市场环境（宏观评分{macro_score}分）：')
            add_paragraph(document, '')

            # 构建表格数据
            matched_scenarios_data = []
            for scenario in environment_matched_scenarios:
                matched_scenarios_data.append([
                    f"{scenario['premium_rate']:+.1f}%",
                    f"{scenario['issue_price']:.2f}",
                    f"{scenario['median_return']:+.1f}%",
                    f"{scenario['profit_prob']:.1f}%",
                    f"{scenario['var_95']:+.1f}%",
                    f"{scenario['drift']*100:+.0f}%",
                    f"{scenario['volatility']*100:.0f}%"
                ])

            # 添加表格
            matched_headers = ['溢价率', '发行价(元)', '收益率(%)', '盈利概率(%)', '95% VaR(%)', '漂移率(%)', '波动率(%)']
            add_table_data(document, matched_headers, matched_scenarios_data)

            # 添加表格说明
            add_paragraph(document, '')
            add_paragraph(document, '表格说明：', bold=True)
            add_paragraph(document, '• 溢价率：相对MA20的溢价率（负值表示折价）')
            add_paragraph(document, '• 发行价：基于MA20和溢价率计算的发行价格')
            add_paragraph(document, '• 收益率：预期中位数年化收益率')
            add_paragraph(document, '• 盈利概率：蒙特卡洛模拟的盈利概率')
            add_paragraph(document, '• 95% VaR：95%置信水平下的最大损失（负值表示损失）')
            add_paragraph(document, '• 漂移率：年化漂移率参数（基于市场环境映射）')
            add_paragraph(document, '• 波动率：年化波动率参数（基于市场环境映射）')
            add_paragraph(document, '')

            # 取符合环境场景中能接受的最大溢价率（即溢价率最高的）作为推荐
            best_matched_scenario = environment_matched_scenarios[0]
            environment_premium_threshold = best_matched_scenario['premium_rate']

            add_paragraph(document, '推荐方案：', bold=True)
            add_paragraph(document, f'• 推荐溢价率：{environment_premium_threshold:+.1f}%（最高溢价率，确保最大安全边际）')
            add_paragraph(document, f'• 推荐发行价：{best_matched_scenario["issue_price"]:.2f}元/股')
            add_paragraph(document, f'• 推荐参数：漂移率{best_matched_scenario["drift"]*100:+.0f}%, 波动率{best_matched_scenario["volatility"]*100:.0f}%')
        else:
            add_paragraph(document, f'• 前10个场景中没有完全符合当前市场环境（漂移率{drift_min*100:.0f}-{drift_max*100:.0f}%, 波动率{vol_min*100:.0f}-{vol_max*100:.0f}%）的方案')
            add_paragraph(document, '• 建议适当放宽市场环境要求或调整参数范围')
    else:
        add_paragraph(document, '参数构造场景筛选结果：', bold=True)
        add_paragraph(document, '• 从585种参数构造场景中未找到符合基础条件的方案')
        add_paragraph(document, '• 建议调整筛选条件或等待更好的市场时机')

    
    # ==================== 9.3.2.6 综合溢价率区间建议 ====================
    add_title(document, '9.3.2.6 综合溢价率区间建议', level=4)

    add_paragraph(document, '综合以上五种情景筛选结果（反向推算、蒙特卡洛模拟、历史数据、DCF估值、参数构造），给出最终的溢价率区间建议。')
    add_paragraph(document, '')

    # 9.3.2.6.1 溢价率情景对照表
    add_title(document, '9.3.2.6.1 溢价率情景对照表', level=5)

    add_paragraph(document, '为明确推导过程，以下是各筛选方法中每一档溢价率对应的情景统计：')
    add_paragraph(document, '本表综合了五种场景类型的分析结果：')
    add_paragraph(document, '• 反向推算：基于目标收益率8%的反向推算结果')
    add_paragraph(document, '• 蒙特卡洛模拟：基于ARIMA+GARCH预测参数的585种情景模拟')
    add_paragraph(document, '• 历史数据：基于历史数据情景（市场指数、行业PE、个股PE）')
    add_paragraph(document, '• DCF估值：基于DCF内在估值模型的绝对估值分析')
    add_paragraph(document, '• 参数构造：基于585种参数构造情景（13种漂移率×5种波动率×9种溢价率）')
    add_paragraph(document, '')

    # 创建综合的溢价率情景对照表 - 包含五种场景类型
    # 收集所有场景的详细信息
    all_scenarios_details = []

    # 1. 反向推算场景 - 从context中获取9.3.2.1节的计算结果
    # 反向推算结果要在多个档次中显示：-20%, -15%, -10%
    # 因为反向推算确定的是上限（-11.56%），更低溢价率更安全也符合条件
    reverse_calc = context['results'].get('reverse_calculation', {})
    if reverse_calc:
        reverse_premium = reverse_calc.get('premium_to_ma20_adjusted', 0)
        reverse_price = reverse_calc.get('max_issue_price_adjusted', 0)
        reverse_current_price = reverse_calc.get('current_price_eval', 0)

        print(f"调试：从context读取反向推算结果 - 溢价率: {reverse_premium:.2f}%, 发行价: {reverse_price:.2f}元")

        # 为多个档次添加反向推算场景
        reverse_levels = [-20.0, -15.0, -10.0]  # 在这些档次中都显示
        for level in reverse_levels:
            all_scenarios_details.append({
                'premium_rate': level,  # 使用档次溢价率而不是实际计算值
                'scenario_type': '反向推算',
                'scenario_name': '目标收益率法',
                'params': f"目标收益8%, 当前价{reverse_current_price:.2f}元, 上限溢价率{reverse_premium:+.1f}%",
                'results': f"上限发行价{reverse_price:.2f}元",
                'is_qualified': True,
                'is_reverse_calc': True  # 标记为反向推算场景
            })
        print(f"调试：反向推算场景已添加到all_scenarios_details，总场景数: {len(all_scenarios_details)}")
    else:
        print(f"调试：context中未找到反向推算结果")

    # 2. 蒙特卡洛模拟场景
    if 'mc_analysis_results' in locals() and mc_analysis_results:
        for result in mc_analysis_results:
            if result.get('is_qualified', False):
                all_scenarios_details.append({
                    'premium_rate': result['premium_rate'],
                    'scenario_type': '蒙特卡洛模拟',
                    'scenario_name': '预测参数模拟',
                    'params': f"漂移率{predicted_drift*100:.0f}%, 波动率{predicted_vol*100:.0f}%, 溢价率{result['premium_rate']:+.0f}%",
                    'results': f"概率{result['profit_prob']:.1f}%, 收益{result['median_return_lockup']:.1f}%",
                    'is_qualified': True
                })

    # 3. 历史数据场景
    if 'qualified_by_category' in locals() and qualified_by_category:
        for category, scenarios in qualified_by_category.items():
            for scenario in scenarios:
                scenario_name = scenario['name']

                # 从场景对象中提取漂移率和波动率信息
                drift_val = scenario.get('drift', 0)
                vol_val = scenario.get('volatility', 0)
                premium_rate = scenario['premium_rate']

                # 显示场景对应的实际漂移率和波动率值
                params = f"漂移率{drift_val*100:.0f}%, 波动率{vol_val*100:.0f}%"

                # 构建结果指标：盈利概率 + 中位数收益率
                results = f"概率{scenario['profit_prob']:.1f}%, 收益{scenario['median_return']:.1f}%"

                all_scenarios_details.append({
                    'premium_rate': scenario['premium_rate'],
                    'scenario_type': '历史数据',
                    'scenario_name': f"{category}",
                    'params': params,
                    'results': results,
                    'is_qualified': True
                })

    # 4. DCF估值场景
    if 'qualified_dcf_scenarios' in locals() and qualified_dcf_scenarios:
        for scenario in qualified_dcf_scenarios:
            scenario_name = scenario['name']

            # 从场景对象中提取漂移率和波动率信息
            drift_val = scenario.get('drift', 0)
            vol_val = scenario.get('volatility', 0)
            premium_rate = scenario['premium_rate']

            # 显示场景对应的实际漂移率和波动率值
            params = f"漂移率{drift_val*100:.0f}%, 波动率{vol_val*100:.0f}%"

            # 构建结果指标：盈利概率 + 中位数收益率
            results = f"概率{scenario['profit_prob']:.1f}%, 收益{scenario['median_return']:.1f}%"

            all_scenarios_details.append({
                'premium_rate': scenario['premium_rate'],
                'scenario_type': 'DCF估值',
                'scenario_name': 'DCF绝对估值',
                'params': params,
                'results': results,
                'is_qualified': True
            })

    # 5. 参数构造场景
    if 'environment_matched_scenarios' in locals() and environment_matched_scenarios:
        for scenario in environment_matched_scenarios:
            all_scenarios_details.append({
                'premium_rate': scenario['premium_rate'],
                'scenario_type': '参数构造',
                'scenario_name': '多参数组合',
                'params': f"漂移率{scenario['drift']*100:.0f}%, 波动率{scenario['volatility']*100:.0f}%, 溢价率{scenario['premium_rate']:+.0f}%",
                'results': f"概率{scenario['profit_prob']:.1f}%, 收益{scenario['median_return']:.1f}%",
                'is_qualified': True
            })

    # 按溢价率档次排序（从低到高：-20.0% → +0.0%）
    premium_levels = [-20.0, -15.0, -10.0, -5.0, +0.0]

    # 构建表格数据 - 每个场景单独一行
    comprehensive_table_data = []

    for premium_level in premium_levels:
        # 找到该溢价率档次的所有场景
        scenarios_in_level = []
        for scenario in all_scenarios_details:
            premium_rate = scenario['premium_rate']
            # 检查是否在当前档次范围内（±2.5%）
            if premium_level - 2.5 <= premium_rate < premium_level + 2.5:
                scenarios_in_level.append(scenario)

        if scenarios_in_level:
            # 添加溢价率档次标题行
            comprehensive_table_data.append([
                f"{premium_level:+.1f}%",
                '', '', '', '', ''
            ])

            # 为该档位的每个场景添加一行
            for scenario in scenarios_in_level:
                comprehensive_table_data.append([
                    '',
                    scenario['scenario_type'],
                    scenario['scenario_name'],
                    f"{scenario['premium_rate']:+.1f}%",  # 具体溢价率数值
                    scenario['params'],
                    scenario['results']
                ])
        else:
            # 没有符合该档位的场景，显示空行
            comprehensive_table_data.append([
                f"{premium_level:+.1f}%",
                '无符合条件场景', '', '', '', ''
            ])

    # 添加综合表格
    comprehensive_headers = ['溢价率档次', '场景类型', '场景名称', '具体溢价率', '参数条件', '结果指标']
    add_table_data(document, comprehensive_headers, comprehensive_table_data)

    # 添加表格说明
    add_paragraph(document, '')
    add_paragraph(document, '表格说明：', bold=True)
    add_paragraph(document, '• 溢价率档次：按名义溢价率（相对MA20）从低到高排序，-20.0%表示最大折价，+0.0%表示平价')
    add_paragraph(document, '• 场景类型：包含反向推算、蒙特卡洛模拟、历史数据、参数构造四种类型')
    add_paragraph(document, '• 场景名称：具体情景的类别名称')
    add_paragraph(document, '• 具体溢价率：该情景使用的精确溢价率数值')
    add_paragraph(document, '• 参数条件：该情景使用的关键参数（如漂移率、波动率等）')
    add_paragraph(document, '• 结果指标：盈利概率和中位数收益率')
    add_paragraph(document, '')

    # 9.3.2.6.2 综合分析结果
    add_title(document, '9.3.2.6.2 综合分析结果', level=5)
    add_paragraph(document, '')

    # 这里需要汇总前面三个子节的阈值结果
    # 由于前面结果是在不同作用域中计算的，这里需要重新计算或存储
    # 为了简化，我们在这里重新计算关键的阈值

    # 1. 历史数据场景阈值（如果有符合条件的）
    historical_threshold = None
    historical_valid = False  # 标记历史场景是否能形成有效交集
    historical_min_premium = None
    historical_max_premium = None

    if 'scenario_options' in locals():
        historical_qualified = []
        for scenario in scenario_options:
            condition_1 = True  # 放开溢价率下限限制
            condition_2 = scenario['median_return'] > 8
            condition_3 = scenario['var_95'] > -30  # 95% VaR > -30%（损失率控制在30%以内）
            condition_4 = scenario['profit_prob'] > 70  # 盈利概率 > 70%

            if condition_1 and condition_2 and condition_3 and condition_4:
                historical_qualified.append(scenario)

        if historical_qualified:
            # 重新按场景类型分组，确保在当前作用域中可用
            scenario_categories = {
                '市场指数': [],
                '行业PE': [],
                '个股PE': []
            }

            # 将情景按大场景分类
            for scenario in historical_qualified:
                for category in scenario_categories.keys():
                    if category in scenario['name']:
                        scenario_categories[category].append(scenario)
                        break

            # 对每个大场景应用基础筛选条件
            qualified_by_category = {}
            for category, scenarios in scenario_categories.items():
                qualified = []
                for scenario in scenarios:
                    # 检查四个基础条件（优化后）
                    condition_1 = True  # 放开溢价率下限限制
                    condition_2 = scenario['median_return'] > 8  # 中位数收益率 > 8%
                    condition_3 = scenario['var_95'] > -30  # 95% VaR > -30%
                    condition_4 = scenario['profit_prob'] > 70  # 盈利概率 > 70%

                    if condition_1 and condition_2 and condition_3 and condition_4:
                        qualified.append(scenario)
                qualified_by_category[category] = qualified

            # 计算各场景类型的溢价率范围，然后取交集（保守策略）
            # 关键：至少需要2个场景类型有符合条件的小场景才能形成有效交集
            category_premium_ranges = {}
            for category, scenarios in qualified_by_category.items():
                if scenarios:
                    premiums = [s['premium_rate'] for s in scenarios]
                    category_premium_ranges[category] = {
                        'min': min(premiums),
                        'max': max(premiums),
                        'count': len(premiums)
                    }

            # 只有当至少2个场景类型有符合条件的小场景时，才计算有效交集
            if len(category_premium_ranges) >= 2:
                # 交集的下限是各范围下限的最大值（最保守的下限）
                historical_min_premium = max(r['min'] for r in category_premium_ranges.values())
                # 交集的上限是各范围上限的最小值（最保守的上限）
                historical_max_premium = min(r['max'] for r in category_premium_ranges.values())

                # 如果交集为空（下限>上限），则历史场景无效
                if historical_min_premium > historical_max_premium:
                    # 无交集，历史场景无法形成有效报价空间
                    historical_valid = False
                    historical_threshold = None
                    historical_min_premium = None
                    historical_max_premium = None
                else:
                    # 有效交集
                    historical_valid = True
                    historical_threshold = (historical_min_premium + historical_max_premium) / 2
            else:
                # 少于2个场景类型有符合条件的小场景，历史场景无效
                historical_valid = False
                historical_threshold = None
                historical_min_premium = None
                historical_max_premium = None

    # 2. 参数构造场景阈值（如果有符合条件的）
    param_threshold = None
    param_threshold_actual = None
    param_valid = False  # 默认无效
    if 'environment_matched_scenarios' in locals() and environment_matched_scenarios:
        param_nominal_premium = environment_matched_scenarios[0]['premium_rate']  # 名义溢价率（相对MA20）
        # 参数构造确定上限，下限为定增规则最低限-20%
        param_min_premium = -20.0  # 定增名义溢价率最低限
        param_max_premium = param_nominal_premium  # 参数构造确定的上限
        # 基于名义溢价率计算发行价：发行价 = MA20 × (1 + 名义溢价率)
        issue_price_from_nominal = ma20_price * (1 + param_nominal_premium / 100)
        # 计算实际溢价率（相对当前价格）：实际溢价率 = 发行价 / 当前价格 - 1
        param_threshold_actual = (issue_price_from_nominal / current_price_eval - 1) * 100
        # 名义溢价率（保持不变）
        param_threshold = param_nominal_premium
        param_valid = True  # 有符合条件的环境匹配场景，有效

    # 3. 反向推算阈值（使用名义溢价率相对MA20）
    reverse_nominal_premium = premium_to_ma20_adjusted  # 从9.3.2.1节获取名义溢价率（相对MA20）
    reverse_valid = False  # 默认无效
    # 反向推算确定上限，下限为定增规则最低限-20%
    reverse_min_premium = -20.0  # 定增名义溢价率最低限
    reverse_max_premium = reverse_nominal_premium  # 反向推算确定的上限
    # 只有当反向推算上限 > -20%时才有效（即有溢价空间）
    if reverse_max_premium > reverse_min_premium:
        reverse_valid = True  # 反向推算有效
    else:
        reverse_valid = False  # 反向推算无效（上限等于或低于下限）
    # 基于名义溢价率计算发行价：发行价 = MA20 × (1 + 名义溢价率)
    issue_price_from_nominal = ma20_price * (1 + reverse_nominal_premium / 100)
    # 计算实际溢价率（相对当前价格）：实际溢价率 = 发行价 / 当前价格 - 1
    reverse_threshold_actual = (issue_price_from_nominal / current_price_eval - 1) * 100
    # 名义溢价率（保持不变）
    reverse_threshold = reverse_nominal_premium

    # 汇总结果（统一使用区间形式，只包含有效的大类）
    thresholds = []

    # 只有当历史场景形成有效交集时才加入
    if historical_valid and historical_threshold is not None:
        total_count = sum(r['count'] for r in category_premium_ranges.values())
        thresholds.append(('历史数据场景', {
            'min_premium': historical_min_premium,
            'max_premium': historical_max_premium,
            'avg_premium': historical_threshold,
            'count': total_count
        }))

    # 参数构造场景（只有当有符合条件的环境匹配场景时才有效）
    if param_valid and param_threshold is not None:
        thresholds.append(('参数构造场景', {
            'min_premium': param_min_premium,
            'max_premium': param_max_premium,
            'avg_premium': param_threshold,
            'count': len(environment_matched_scenarios) if 'environment_matched_scenarios' in locals() else 0
        }))

    # 反向推算（只有当上限 > -20%时才有效）
    if reverse_valid and 'premium_to_ma20_adjusted' in locals():
        thresholds.append(('反向推算', {
            'min_premium': reverse_min_premium,
            'max_premium': reverse_max_premium,
            'avg_premium': reverse_nominal_premium,
            'count': 1  # 反向推算只有一个结果
        }))

    if len(thresholds) >= 2:
        # 提取所有阈值的最小值和最大值
        all_min_premiums = []
        all_max_premiums = []

        for name, threshold in thresholds:
            # 所有阈值现在都是区间形式（dict）
            all_min_premiums.append(threshold['min_premium'])
            all_max_premiums.append(threshold['max_premium'])

        min_premium = min(all_min_premiums)
        max_premium = max(all_max_premiums)
        premium_range = max_premium - min_premium

        add_paragraph(document, f'• 有效阈值数量：{len(thresholds)}个')
        for name, threshold in thresholds:
            # 所有阈值都显示为区间形式
            count_val = threshold['count']
            count_str = f"，共{count_val}个符合情景" if count_val > 1 else ""
            min_val = threshold['min_premium']
            max_val = threshold['max_premium']
            add_paragraph(document, f'• {name}阈值：[{min_val:+.1f}%, {max_val:+.1f}%]（名义溢价率，相对MA20{count_str}）')

        add_paragraph(document, '')
        add_paragraph(document, '最终溢价率区间建议：', bold=True)
        add_paragraph(document, f'• 名义溢价率区间（相对MA20）：[{min_premium:+.2f}%, {max_premium:+.2f}%]')
        add_paragraph(document, f'• 区间宽度：{premium_range:.2f}%')

        add_paragraph(document, '')
        add_paragraph(document, '最终建议：', bold=True)
        add_paragraph(document, f'• 名义溢价率上限：≤{max_premium:.2f}%（相对MA20，硬约束）')
        add_paragraph(document, f'• 名义溢价率下限：{min_premium:+.2f}%（相对MA20，最优方案）')
        add_paragraph(document, f'• 区间宽度：{premium_range:.2f}%')
    else:
        add_paragraph(document, '• 有效阈值数量不足（少于2个）')
        add_paragraph(document, '• 无法形成可靠的溢价率区间，建议谨慎参与本次定增')
        add_paragraph(document, '• 可能原因：')
        add_paragraph(document, '  - 历史数据场景：未能形成有效交集（需要至少2个场景类型有符合条件的小场景）')
        add_paragraph(document, '  - 参数构造场景：未能筛选出符合条件的环境匹配场景')
        add_paragraph(document, '  - 反向推算：上限值≤-20%，无溢价空间（无法在定增规则下达到8%收益目标）')
        add_paragraph(document, '• 建议：参考单一可用的有效阈值（如有）或谨慎参与本次定增')

        # ==================== 9.4 主要风险提示 ====================
    add_title(document, '9.4 主要风险提示', level=2)
    add_paragraph(document, '基于多维度风险分析，提示以下主要风险（详见前文各章节详细分析）：')
    add_paragraph(document, '')

    # 1. 市场风险
    add_paragraph(document, '1. 市场风险')
    add_paragraph(document, f'   • 波动率风险：当前120日窗口年化波动率为{mc_volatility_120d*100:.1f}%，市场波动可能导致实际收益偏离预期')
    add_paragraph(document, f'   • 趋势风险：当前120日窗口年化漂移率为{mc_drift_120d*100:+.2f}%，{"上升趋势" if mc_drift_120d > 0 else "下降趋势" if mc_drift_120d < 0 else "震荡趋势"}可能影响解禁时收益')

    # 2. 流动性风险
    add_paragraph(document, '2. 流动性风险')
    add_paragraph(document, f'   • 锁定期风险：{project_params["lockup_period"]}个月锁定期内无法交易，需承担期间价格波动')
    add_paragraph(document, '   • 解禁冲击：解禁后可能面临抛压，导致实际变现价格低于理论价格')

    # 3. VaR在险价值风险
    add_paragraph(document, '3. VaR在险价值风险')
    # 使用基于期收益率（锁定期收益率）的VaR，而不是年化VaR
    # 从context中获取var_results，然后获取120日窗口的var_95
    var_results = context.get('results', {}).get('var_results', {})
    var_95_period = var_results.get('120日', {}).get('var_95', 0.5)
    var_95_display = var_95_period * 100  # 转换为百分比
    add_paragraph(document, f'   • 120日窗口：95%置信水平下最大可能亏损{var_95_display:.1f}%')
    add_paragraph(document, '   • 尾部风险：历史数据显示，小概率极端事件（黑天鹅）可能导致损失超过VaR预测值')

    # 4. 估值风险
    add_paragraph(document, '4. 估值风险')
    add_paragraph(document, f'   • DCF估值风险：DCF内在价值{intrinsic_value:.2f}元/股基于多个假设，实际业绩可能偏离预测')
    add_paragraph(document, '   • 相对估值风险：PE/PS/PB相对估值基于行业平均水平，行业景气度变化可能导致估值体系重构')

    # 5. 发行定价风险
    add_paragraph(document, '5. 发行定价风险')
    add_paragraph(document, '   • 溢价风险：如发行价过高，影响安全边际和整体抗风险能力。')
    add_paragraph(document, '   • 溢价发行无安全边际，需重点关注公司成长性')
    add_paragraph(document, '   • 定价偏离：若发行价显著高于盈亏平衡价格，将大幅降低盈利概率和预期收益')

    # 6. 其他风险
    add_paragraph(document, '6. 其他风险')
    add_paragraph(document, '   • 行业政策风险：需关注行业监管政策变化')
    add_paragraph(document, '   • 业绩波动风险：需关注公司业绩预告、审计报告等')
    add_paragraph(document, '   • 竞争格局风险：行业竞争加剧可能影响盈利能力')

    # ==================== 9.5 免责声明 ====================
    add_title(document, '9.5 免责声明', level=2)

    disclaimer = f'''
    本报告数据取至tushare平台，报告系半自动化生成，仅供参考使用，不构成投资建议。

    1. 本报告基于历史数据和公开信息进行分析，不能保证tushare数据不会出错；
    2. 市场有风险，投资需谨慎。本报告中的任何分析观点不代表未来表现；
    3. 本报告提到的任何证券或投资标的，仅为分析示例，不构成推荐；
    4. 本报告的知识产权归分析团队所有，未经许可不得转载或使用。

    报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    '''

    add_paragraph(document, disclaimer, font_size=9)

    add_section_break(document)

    return context
