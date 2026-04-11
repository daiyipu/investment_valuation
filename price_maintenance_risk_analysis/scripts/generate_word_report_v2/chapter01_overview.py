# -*- coding: utf-8 -*-
"""
第一章：项目概况

本章节生成报告的项目概况部分，包括：
- 报告封面
- 目录生成
- 1.1 基本信息
- 1.2 个股数据分析
- 1.3 市场指数分析
- 1.4 行业数据分析
"""

import os
import json
import numpy as np
from datetime import datetime, timedelta
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 导入工具函数
from module_utils import (
    add_title, add_paragraph, add_table_data, add_image, add_section_break,
    generate_stock_market_data_charts_split, generate_industry_index_charts,
    generate_index_data_charts_split
)


def generate_chapter(context):
    """
    生成第一章：项目概况

    Args:
        context: 包含以下键的字典:
            - document: Word文档对象
            - project_params: 项目参数
            - market_data: 市场数据
            - industry_data: 行业数据（可选）
            - IMAGES_DIR: 图片目录
            - DATA_DIR: 数据目录

    Returns:
        更新后的context字典
    """
    # 从context中提取数据
    document = context['document']
    project_params = context['project_params']
    market_data = context['market_data']
    industry_data = context.get('industry_data', {})
    IMAGES_DIR = context['IMAGES_DIR']
    DATA_DIR = context['DATA_DIR']

    # 提取常用变量
    stock_code = project_params.get('stock_code', '300735.SZ')
    issue_price = project_params['issue_price']
    current_price = project_params['current_price']
    lockup_period = project_params['lockup_period']
    ma20 = market_data.get('ma_20', 0)

    # 计算发行类型和溢价率
    # 名义溢价率（相对MA20）：用于定价说明
    nominal_premium = (issue_price - ma20) / ma20 * 100

    # 实际溢价率（相对发行日价格）：用于投资分析
    actual_premium = (issue_price - current_price) / current_price * 100

    # 发行类型固定为询价发行
    issue_type = "询价发行"

    # 获取数据日期（用于显示当前价格的实际日期）
    data_date = market_data.get('analysis_date', '')
    if data_date:
        try:
            from datetime import datetime
            data_date_obj = datetime.strptime(data_date, '%Y%m%d')
            data_date_formatted = data_date_obj.strftime('%Y年%m月%d日')
            current_price_desc = f'{market_data["current_price"]:.2f} 元/股（数据日期{data_date_formatted}）'
        except:
            current_price_desc = f'{market_data["current_price"]:.2f} 元/股（最近交易日数据）'
    else:
        current_price_desc = f'{market_data["current_price"]:.2f} 元/股（最近交易日数据）'

    current_date = datetime.now().strftime('%Y年%m月%d日')

    # ==================== 报告封面 ====================
    add_title(document, '定增项目风险分析报告', level=0)
    add_paragraph(document, '')

    # 报告基本信息
    info_table = document.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'

    # 根据发行日是否确定，显示不同的价格标签
    is_fixed_issue_date = project_params.get('invitation_date_fixed', False)
    if is_fixed_issue_date:
        price_label_cover = '发行日价格'
        price_value_cover = f'{current_price:.2f} 元/股（发行日{project_params.get("issue_date", "")}）'
    else:
        price_label_cover = '发行日价格（拟）'
        price_value_cover = f'{current_price:.2f} 元/股'

    info_data = [
        ['公司名称', f'{project_params.get("company_name", "光弘科技")} ({stock_code})'],
        ['报告日期', current_date],
        ['发行价格', f'{issue_price:.2f} 元/股'],
        [price_label_cover, price_value_cover],
        ['MA20价格', f'{ma20:.2f} 元/股'],
        ['发行类型', issue_type]
    ]

    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    # 设置表格样式
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if len(paragraph.runs) > 0:
                    paragraph.runs[0].font.size = 14
                    for run in paragraph.runs:
                        run.font.name = '仿宋_GB2312'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    add_paragraph(document, '')

    # ==================== 目录 ====================
    add_title(document, '目 录', level=1)
    add_paragraph(document, '一、项目概况')
    add_paragraph(document, '二、相对估值分析')
    add_paragraph(document, '三、DCF估值分析')
    add_paragraph(document, '四、敏感性分析')
    add_paragraph(document, '五、蒙特卡洛模拟')
    add_paragraph(document, '六、情景分析')
    add_paragraph(document, '七、压力测试')
    add_paragraph(document, '八、VaR风险度量')
    add_paragraph(document, '九、风控建议与风险提示')

    add_section_break(document)

    # ==================== 一、项目概况 ====================
    add_title(document, '一、项目概况', level=1)

    # 1.1 基本信息
    add_title(document, '1.1 基本信息', level=2)

    # 根据发行日是否确定，决定价格字段的显示内容
    is_fixed_issue_date = project_params.get('invitation_date_fixed', False)

    # 设置各个价格字段的显示内容
    if is_fixed_issue_date:
        # 发行日确定的情况
        issue_date_price_note = f'{current_price:.2f} 元/股（发行日{project_params.get("issue_date", "")}）'
        planned_price_note = f'{current_price:.2f} 元/股'
    else:
        # 发行日不确定的情况
        issue_date_price_note = ''  # "发行日价格"字段显示空值
        planned_price_note = f'{current_price:.2f} 元/股'

    project_headers = ['指标', '具体内容']
    project_data = [
        ['股票代码', stock_code],
        ['公司名称', project_params.get('company_name', '光弘科技')],
        ['发行日/报价日', project_params.get('issue_date', 'N/A')],
        ['询价邀请日', project_params.get('invitation_date', 'N/A')],
        ['发行价格', f'{issue_price:.2f} 元/股'],
        ['发行数量', f'{project_params["issue_shares"]:,} 股'],
        ['锁定期', f'{lockup_period} 个月'],
        ['投资金额（模拟）', f'{project_params["financing_amount"] / 100000000:.2f} 亿元'],
        ['当前价格', current_price_desc],
        ['发行日价格（拟）', planned_price_note],
        ['发行日价格', issue_date_price_note],
        ['MA20价格', f'{ma20:.2f} 元/股（基于发行日的成交量加权均价）'],
        ['发行类型', issue_type],
        ['实际溢价率', f'{actual_premium:+.2f}%']
    ]
    add_table_data(document, project_headers, project_data)

    # 1.2 个股数据分析
    add_title(document, '1.2 个股数据分析', level=2)

    # 获取最新的交易日价格（如果有保存的话）
    latest_trading_price = project_params.get('latest_trading_date_price', market_data.get('current_price'))

    # 根据发行日是否确定，显示不同的价格标签
    if is_fixed_issue_date:
        price_label_market = '最新交易日价格'
        price_note_market = f'（数据日期{market_data.get("analysis_date", "")}）'
    else:
        price_label_market = '最新交易日价格'
        price_note_market = f'（数据日期{market_data.get("analysis_date", "")}）'

    market_headers = ['指标', '数值']
    market_table_data = [
        [price_label_market, f'{latest_trading_price:.2f} 元/股{price_note_market}'],
        ['平均价格', f'{market_data.get("avg_price_all", 0):.2f} 元/股（{market_data.get("total_days", 0)}个交易日）'],
        ['价格标准差', f'{market_data.get("price_std", 0):.2f}'],
        ['月度波动率(20日)', f'{market_data.get("volatility_20d", 0)*100:.2f}%'],
        ['季度波动率(60日)', f'{market_data.get("volatility_60d", 0)*100:.2f}%'],
        ['半年波动率(120日)', f'{market_data.get("volatility_120d", 0)*100:.2f}%'],
        ['年度波动率(250日)', f'{market_data.get("volatility_250d", 0)*100:.2f}%'],
        ['月度区间收益率(20日)', f'{market_data.get("period_return_20d", 0)*100:+.2f}%'],
        ['月度年化收益率(20日)', f'{market_data.get("annual_return_20d", 0)*100:+.2f}%'],
        ['季度区间收益率(60日)', f'{market_data.get("period_return_60d", 0)*100:+.2f}%'],
        ['季度年化收益率(60日)', f'{market_data.get("annual_return_60d", 0)*100:+.2f}%'],
        ['半年区间收益率(120日)', f'{market_data.get("period_return_120d", 0)*100:+.2f}%'],
        ['半年年化收益率(120日)', f'{market_data.get("annual_return_120d", 0)*100:+.2f}%'],
        ['年度区间收益率(250日)', f'{market_data.get("period_return_250d", 0)*100:+.2f}%'],
        ['年度年化收益率(250日)', f'{market_data.get("annual_return_250d", 0)*100:+.2f}%'],
        ['MA20', f'{market_data.get("ma_20", 0):.2f} 元/股'],
        ['MA60', f'{market_data.get("ma_60", 0):.2f} 元/股'],
        ['MA120', f'{market_data.get("ma_120", 0):.2f} 元/股'],
        ['MA250', f'{market_data.get("ma_250", 0):.2f} 元/股'],
        ['数据来源', 'Tushare API']
    ]
    add_table_data(document, market_headers, market_table_data)

    # 1.2.1 生成个股市场数据图表
    print("生成个股市场数据图表...")
    stock_market_data_file = os.path.join(DATA_DIR, f"{stock_code.replace('.', '_')}_market_data.json")

    # 尝试从Tushare获取历史数据生成图表
    stock_charts_paths = []
    try:
        import tushare as ts

        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if ts_token:
            pro = ts.pro_api(ts_token)

            # 获取历史价格数据
            end_date = datetime.now().strftime('%Y%m%d')
            start_date = (datetime.now() - timedelta(days=500)).strftime('%Y%m%d')

            df_daily = pro.daily(ts_code=stock_code, start_date=start_date, end_date=end_date)
            if not df_daily.empty:
                df_daily = df_daily.sort_values('trade_date').reset_index(drop=True)

                # 准备价格数据
                price_data = {'close': df_daily['close'].values, 'pct_chg': df_daily['pct_chg'].values}
                for window in [20, 60, 120, 250]:
                    ma = df_daily['close'].rolling(window=window).mean()
                    price_data[f'ma_{window}'] = ma.values

                # 准备波动率数据
                volatility_data = {}
                for window in [20, 60, 120, 250]:
                    pct_decimal = df_daily['pct_chg'] / 100.0
                    rolling_std = pct_decimal.rolling(window=window).std()
                    rolling_vol = rolling_std * np.sqrt(250)
                    volatility_data[f'volatility_{window}d_series'] = rolling_vol.values

                # 生成图表
                stock_charts_paths = generate_stock_market_data_charts_split(
                    market_data, price_data, volatility_data, IMAGES_DIR
                )

                # 添加图表到文档
                add_paragraph(document, '图表 1.1: 股价走势与均线分析')
                if len(stock_charts_paths) > 0 and os.path.exists(stock_charts_paths[0]):
                    add_image(document, stock_charts_paths[0], width=Inches(6.5))
                    add_paragraph(document, '')

                add_paragraph(document, '图表 1.2: 滚动波动率分析')
                if len(stock_charts_paths) > 1 and os.path.exists(stock_charts_paths[1]):
                    add_image(document, stock_charts_paths[1], width=Inches(6.5))
                    add_paragraph(document, '')

                add_paragraph(document, '图表 1.3: 日收益率分布')
                if len(stock_charts_paths) > 2 and os.path.exists(stock_charts_paths[2]):
                    add_image(document, stock_charts_paths[2], width=Inches(6.5))
                    add_paragraph(document, '')

                add_paragraph(document, '分析结论：')
                add_paragraph(document, f'• 当前股价{market_data["current_price"]:.2f}元，{issue_type}于MA20({ma20:.2f}元)')
                add_paragraph(document, f'• 60日年化波动率为{market_data.get("volatility_60d", 0)*100:.2f}%，{"高于" if market_data.get("volatility_60d", 0) > 0.3 else "低于"}市场平均水平')
                add_paragraph(document, f'• 60日年化收益率为{market_data.get("annual_return_60d", 0)*100:.2f}%，{"表现良好" if market_data.get("annual_return_60d", 0) > 0 else "表现不佳"}')
    except Exception as e:
        print(f" 生成个股市场数据图表失败: {e}")

    # 1.3 市场指数分析
    add_paragraph(document, '')
    add_title(document, '1.3 市场指数分析', level=2)

    add_paragraph(document, '本章节分析主要市场指数的表现，包括波动率、收益率、胜率等指标，为项目风险评估提供市场环境参考。')

    # 加载指数数据（从data目录）
    # 优先使用新方法计算的指数数据（使用修复后的年化收益率公式）
    indices_data_file_v2 = os.path.join(DATA_DIR, 'market_indices_scenario_data_v2.json')
    indices_data_file = os.path.join(DATA_DIR, 'market_indices_scenario_data.json')
    indices_charts_paths = []

    # 优先使用新数据
    if os.path.exists(indices_data_file_v2):
        print(" 使用新方法计算的指数数据（修复后的年化收益率公式）")
        indices_data_file = indices_data_file_v2

    if os.path.exists(indices_data_file):
        print("加载市场指数数据...")
        try:
            with open(indices_data_file, 'r', encoding='utf-8') as f:
                indices_data = json.load(f)

            # 生成指数对比图表
            indices_charts_paths = generate_index_data_charts_split(indices_data, IMAGES_DIR)

            # 添加指数对比表格 - 60日指标
            add_paragraph(document, '')
            add_paragraph(document, '主要市场指数60日指标对比：')

            index_headers_60d = ['指数', '当前点位', '波动率', '区间收益率', '年化收益率', '胜率']
            index_table_data_60d = []
            for name, data in indices_data.items():
                # 从年化收益率反推区间收益率（使用简单年化方法）
                # 年化收益率 = 区间收益率 × (250/60)
                # 反推：区间收益率 = 年化收益率 × (60/250)
                annual_ret = data.get('return_60d', 0)
                period_ret_60d = annual_ret * (60.0 / 250.0)

                index_table_data_60d.append([
                    name,
                    f"{data.get('current_level', 0):.2f}",
                    f"{data.get('volatility_60d', 0)*100:.2f}%",
                    f"{period_ret_60d*100:+.2f}%",
                    f"{data.get('return_60d', 0)*100:+.2f}%",
                    f"{data.get('win_rate_60d', 0)*100:.1f}%"
                ])
            add_table_data(document, index_headers_60d, index_table_data_60d)

            # 添加指数对比表格 - 120日指标
            add_paragraph(document, '')
            add_paragraph(document, '主要市场指数120日指标对比（半年线）：')

            index_headers_120d = ['指数', '当前点位', '波动率', '区间收益率', '年化收益率', '胜率']
            index_table_data_120d = []
            for name, data in indices_data.items():
                # 从年化收益率反推区间收益率（使用简单年化方法）
                annual_ret = data.get('return_120d', 0)
                period_ret_120d = annual_ret * (120.0 / 250.0)

                # 如果没有win_rate_120d，使用win_rate_60d作为近似值
                win_rate_120d = data.get('win_rate_120d', 0)

                index_table_data_120d.append([
                    name,
                    f"{data.get('current_level', 0):.2f}",
                    f"{data.get('volatility_120d', 0)*100:.2f}%",
                    f"{period_ret_120d*100:+.2f}%",
                    f"{data.get('return_120d', 0)*100:+.2f}%",
                    f"{win_rate_120d*100:.1f}%"
                ])
            add_table_data(document, index_headers_120d, index_table_data_120d)

            # 添加指数对比表格 - 250日指标
            add_paragraph(document, '')
            add_paragraph(document, '主要市场指数250日指标对比（年线）：')

            index_headers_250d = ['指数', '当前点位', '波动率', '区间收益率', '年化收益率', '胜率']
            index_table_data_250d = []
            for name, data in indices_data.items():
                # 对于250日窗口，年化收益率≈区间收益率（250个交易日≈1年）
                annual_ret = data.get('return_250d', 0)
                period_ret = annual_ret  # 250日≈1年，区间收益率≈年化收益率

                # 如果没有win_rate_250d，使用win_rate_60d作为近似值
                win_rate_250d = data.get('win_rate_250d', 0)

                index_table_data_250d.append([
                    name,
                    f"{data.get('current_level', 0):.2f}",
                    f"{data.get('volatility_250d', 0)*100:.2f}%",
                    f"{period_ret*100:+.2f}%",
                    f"{annual_ret*100:+.2f}%",
                    f"{win_rate_250d*100:.1f}%"
                ])
            add_table_data(document, index_headers_250d, index_table_data_250d)

            add_paragraph(document, '')
            add_paragraph(document, '说明：')
            add_paragraph(document, '• 250日（年线）包含约250个交易日，接近一年的交易日数量（约252天）')
            add_paragraph(document, '• 年化收益率≈区间收益率（因250日≈1年，无需额外年化处理）')
            add_paragraph(document, '• 胜率基于实际历史数据计算，反映各时间窗口的上涨天数占比')
            add_paragraph(document, '• 建议以沪深300为基准参考，其年化收益率反映了市场整体表现')
            add_paragraph(document, '• 个股投资应结合行业特征和市场环境综合判断')

            # 添加图表 - 60日窗口
            add_paragraph(document, '图表 1.4: 各指数波动率对比 (60日窗口)')
            if len(indices_charts_paths) > 0 and os.path.exists(indices_charts_paths[0]):
                add_image(document, indices_charts_paths[0], width=Inches(6))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.5: 各指数年化收益率对比 (60日窗口)')
            if len(indices_charts_paths) > 1 and os.path.exists(indices_charts_paths[1]):
                add_image(document, indices_charts_paths[1], width=Inches(6))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.6: 各指数胜率对比 (60日窗口)')
            if len(indices_charts_paths) > 2 and os.path.exists(indices_charts_paths[2]):
                add_image(document, indices_charts_paths[2], width=Inches(6))
                add_paragraph(document, '')

            # 添加图表 - 120日窗口（半年线）
            add_paragraph(document, '图表 1.7: 各指数波动率对比 (120日窗口/半年线)')
            if len(indices_charts_paths) > 4 and os.path.exists(indices_charts_paths[4]):
                add_image(document, indices_charts_paths[4], width=Inches(6))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.8: 各指数年化收益率对比 (120日窗口/半年线)')
            if len(indices_charts_paths) > 5 and os.path.exists(indices_charts_paths[5]):
                add_image(document, indices_charts_paths[5], width=Inches(6))
                add_paragraph(document, '')

            # 添加图表 - 250日窗口
            add_paragraph(document, '图表 1.9: 各指数波动率对比 (250日年线窗口)')
            if len(indices_charts_paths) > 6 and os.path.exists(indices_charts_paths[6]):
                add_image(document, indices_charts_paths[6], width=Inches(6.3))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.10: 各指数年化收益率对比 (250日年线窗口)')
            if len(indices_charts_paths) > 7 and os.path.exists(indices_charts_paths[7]):
                add_image(document, indices_charts_paths[7], width=Inches(6.3))
                add_paragraph(document, '')

            # 添加分析结论
            add_paragraph(document, '分析结论：')
            add_paragraph(document, '• 沪深300作为市场基准，波动率相对稳定，适合作为风险参照')
            add_paragraph(document, '• 创业板指和科创50波动率较高，体现成长股的高风险高收益特征')
            add_paragraph(document, '• 60日（短趋势）反映市场近期情绪，120日（半年线）反映中期趋势，250日（年线）反映长期趋势，适合作为战略参考')
            add_paragraph(document, '• 当前市场环境下，建议关注指数技术位置对个股表现的影响')

        except Exception as e:
            print(f" 加载指数数据失败: {e}")
            add_paragraph(document, ' 指数数据暂未加载，请先运行07_market_data_analysis.ipynb生成数据')
    else:
        add_paragraph(document, ' 指数数据文件不存在，请先运行07_market_data_analysis.ipynb生成数据')

    # 1.4 行业数据分析
    add_paragraph(document, '')
    add_title(document, '1.4 行业数据分析', level=2)

    add_paragraph(document, '本章节分析所属行业的申万三级行业指数表现，为项目风险评估提供行业环境参考。')

    # 尝试加载行业数据
    try:
        industry_data_file = os.path.join(DATA_DIR, f'{stock_code.replace(".", "_")}_industry_data.json')

        if os.path.exists(industry_data_file):
            with open(industry_data_file, 'r', encoding='utf-8') as f:
                industry_data = json.load(f)

            print(f" 已加载行业数据: {industry_data_file}")

            # 1.4.1 行业基本信息
            add_paragraph(document, '')
            add_title(document, '1.4.1 行业基本信息', level=3)

            industry_headers = ['指标', '数值']
            industry_table_data = [
                ['申万一级行业', f"{industry_data.get('sw_l1_name', 'N/A')}"],
                ['申万二级行业', f"{industry_data.get('sw_l2_name', 'N/A')}"],
                ['申万三级行业', f"{industry_data.get('sw_l3_name', 'N/A')}"],
                ['行业指数代码', f"{industry_data.get('index_code', 'N/A')}"],
                ['当前点位', f"{industry_data.get('current_level', 0):.2f}"],
                ['分析日期', f"{industry_data.get('analysis_date', 'N/A')}"],
            ]
            add_table_data(document, industry_headers, industry_table_data)

            # 1.4.2 行业风险与技术指标综合分析
            add_paragraph(document, '')
            add_title(document, '1.4.2 行业风险与技术指标综合分析', level=3)

            add_paragraph(document, '本节从风险水平、收益率表现、技术位置和胜率四个维度综合分析行业指数特征。')
            add_paragraph(document, '')

            # 综合指标表格（风险 + 收益率）
            risk_return_data = [
                ['时间窗口', '波动率', '区间收益率', '年化收益率'],
                ['月度(20日)',
                 f"{industry_data.get('volatility_20d', 0)*100:.2f}%",
                 f"{industry_data.get('period_return_20d', 0)*100:+.2f}%",
                 f"{industry_data.get('annual_return_20d', 0)*100:+.2f}%"],
                ['季度(60日)',
                 f"{industry_data.get('volatility_60d', 0)*100:.2f}%",
                 f"{industry_data.get('period_return_60d', 0)*100:+.2f}%",
                 f"{industry_data.get('annual_return_60d', 0)*100:+.2f}%"],
                ['半年(120日)',
                 f"{industry_data.get('volatility_120d', 0)*100:.2f}%",
                 f"{industry_data.get('period_return_120d', 0)*100:+.2f}%",
                 f"{industry_data.get('annual_return_120d', 0)*100:+.2f}%"],
                ['年度(250日)',
                 f"{industry_data.get('volatility_250d', 0)*100:.2f}%",
                 f"{industry_data.get('period_return_250d', 0)*100:+.2f}%",
                 f"{industry_data.get('annual_return_250d', 0)*100:+.2f}%"],
            ]
            add_table_data(document, ['时间窗口', '波动率(风险)', '区间收益率', '年化收益率'], risk_return_data)

            add_paragraph(document, '')

            # 技术位置与胜率表格
            add_paragraph(document, '')

            # 获取胜率数据
            win_rate_20d = industry_data.get('win_rate_20d', 0)
            win_rate_60d = industry_data.get('win_rate_60d', 0)
            win_rate_120d = industry_data.get('win_rate_120d', 0)
            win_rate_250d = industry_data.get('win_rate_250d', 0)

            tech_position_data = [
                ['时间窗口', '移动平均线', '胜率(上涨天数占比)'],
                ['月度(20日)',
                 f"{industry_data.get('ma_20', 0):.2f} 点",
                 f"{win_rate_20d*100:.1f}%"],
                ['季度(60日)',
                 f"{industry_data.get('ma_60', 0):.2f} 点",
                 f"{win_rate_60d*100:.1f}%"],
                ['半年(120日)',
                 f"{industry_data.get('ma_120', 0):.2f} 点",
                 f"{win_rate_120d*100:.1f}%"],
                ['年度(250日)',
                 f"{industry_data.get('ma_250', 0):.2f} 点",
                 f"{win_rate_250d*100:.1f}%"],
            ]
            add_table_data(document, ['时间窗口', '移动平均线(技术位置)', '胜率'], tech_position_data)

            add_paragraph(document, '')
            add_paragraph(document, '指标说明：')
            add_paragraph(document, '• 波动率：反映行业指数的不确定性程度，越高风险越大')
            add_paragraph(document, '• 区间收益率：时间窗口内的累计涨跌幅')
            add_paragraph(document, '• 年化收益率：将区间收益率年化后的收益率，便于不同窗口期对比')
            add_paragraph(document, '• 移动平均线：行业指数在相应时间窗口内的平均价格，反映趋势')
            add_paragraph(document, '• 胜率：上涨天数占总交易日的比例，反映趋势强度')

            # 1.4.3 行业指数图表分析
            add_paragraph(document, '')
            add_title(document, '1.4.3 行业指数图表分析', level=3)

            # 生成行业指数图表
            add_paragraph(document, '')
            industry_charts_paths = generate_industry_index_charts(industry_data, IMAGES_DIR)

            # 添加图表到文档
            add_paragraph(document, '图表 1.12: 行业指数走势与均线分析')
            if len(industry_charts_paths) > 0 and os.path.exists(industry_charts_paths[0]):
                add_image(document, industry_charts_paths[0], width=Inches(6.5))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.13: 行业指数滚动波动率分析')
            if len(industry_charts_paths) > 1 and os.path.exists(industry_charts_paths[1]):
                add_image(document, industry_charts_paths[1], width=Inches(6.5))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.14: 行业指数日收益率分布')
            if len(industry_charts_paths) > 2 and os.path.exists(industry_charts_paths[2]):
                add_image(document, industry_charts_paths[2], width=Inches(6.5))
                add_paragraph(document, '')

            # 1.4.4 个股与行业对比分析
            add_paragraph(document, '')
            add_title(document, '1.4.4 个股与行业对比分析', level=3)

            add_paragraph(document, '')
            add_paragraph(document, ' 个股与行业表现对比（多窗口期）：')

            # 定义窗口期
            windows = [
                ('20日窗口', '20d'),
                ('60日窗口', '60d'),
                ('120日窗口', '120d'),
                ('250日窗口', '250d')
            ]

            # 构建多窗口期对比表格
            comparison_headers = ['窗口期', '波动率（个股）', '波动率（行业）', '差异', '年化收益率（个股）', '年化收益率（行业）', '差异']
            comparison_data = []

            for window_name, window_suffix in windows:
                stock_vol = market_data.get(f'volatility_{window_suffix}', 0)
                industry_vol = industry_data.get(f'volatility_{window_suffix}', 0)
                stock_ret = market_data.get(f'annual_return_{window_suffix}', 0)
                industry_ret = industry_data.get(f'annual_return_{window_suffix}', 0)

                vol_diff = (stock_vol - industry_vol) * 100
                ret_diff = (stock_ret - industry_ret) * 100

                comparison_data.append([
                    window_name,
                    f"{stock_vol*100:.2f}%" if stock_vol > 0 else "N/A",
                    f"{industry_vol*100:.2f}%" if industry_vol > 0 else "N/A",
                    f"{vol_diff:+.2f}%" if stock_vol > 0 and industry_vol > 0 else "N/A",
                    f"{stock_ret*100:+.2f}%" if stock_ret != 0 else "N/A",
                    f"{industry_ret*100:+.2f}%" if industry_ret != 0 else "N/A",
                    f"{ret_diff:+.2f}%" if stock_ret != 0 and industry_ret != 0 else "N/A"
                ])

            add_table_data(document, comparison_headers, comparison_data)

            # 分析结论（以60日窗口为主）
            add_paragraph(document, '')
            add_paragraph(document, '行业分析结论（以60日窗口为主）：')

            stock_volatility = market_data.get('volatility_60d', 0)
            industry_volatility = industry_data.get('volatility_60d', 0)
            stock_return = market_data.get('annual_return_60d', 0)
            industry_return = industry_data.get('annual_return_60d', 0)

            if stock_volatility > 0 and industry_volatility > 0:
                if stock_volatility > industry_volatility:
                    add_paragraph(document, f'• 个股波动率({stock_volatility*100:.2f}%)高于行业({industry_volatility*100:.2f}%)，表明个股风险水平较高')
                else:
                    add_paragraph(document, f'• 个股波动率({stock_volatility*100:.2f}%)低于行业({industry_volatility*100:.2f}%)，表明个股相对稳定')

            if stock_return != 0 and industry_return != 0:
                if stock_return > industry_return:
                    add_paragraph(document, f'• 个股收益率({stock_return*100:+.2f}%)跑赢行业({industry_return*100:+.2f}%)，表现优异')
                else:
                    add_paragraph(document, f'• 个股收益率({stock_return*100:+.2f}%)跑输行业({industry_return*100:+.2f}%)，需关注行业竞争压力')

            add_paragraph(document, '• 行业数据为申万三级行业指数，反映细分行业的整体表现')
            add_paragraph(document, '• 建议结合行业景气度、公司基本面等因素综合评估投资价值')
            add_paragraph(document, '• 不同窗口期反映不同时间维度的市场表现，短期波动较大，长期趋势相对稳定')

        else:
            print(f" 行业数据文件不存在: {industry_data_file}")
            add_paragraph(document, ' 行业数据暂未加载，请先运行update_market_data.py生成数据')

    except Exception as e:
        print(f" 加载行业数据失败: {e}")
        add_paragraph(document, f' 行业数据加载失败: {e}')

    add_section_break(document)

    # 更新context并返回
    return context
