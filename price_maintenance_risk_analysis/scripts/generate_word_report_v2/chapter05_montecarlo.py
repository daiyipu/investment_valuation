# -*- coding: utf-8 -*-
"""
第五章 - 蒙特卡洛模拟

功能：生成蒙特卡洛模拟章节内容
"""

import sys
import os
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import random

# 添加路径
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, PROJECT_DIR)

from utils.font_manager import get_font_prop
from utils.time_series_forecaster import TimeSeriesForecaster
from utils.analysis_tools import PrivatePlacementRiskAnalyzer

# 获取中文字体
font_prop = get_font_prop()


def add_title(document, text, level=1):
    """添加标题"""
    if level == 1:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = '方正公文小标宋_GBK'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
            run.font.size = Pt(16)
    elif level == 2:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = '方正公文小标宋_GBK'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
            run.font.size = Pt(15)
    else:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(14)
    return heading


def add_paragraph(document, text, bold=False):
    """添加段落"""
    para = document.add_paragraph(text)
    for run in para.runs:
        run.font.bold = bold
        if '<b>' in text and '</b>' in text:
            # 处理HTML格式的粗体标记
            run.text = text.replace('<b>', '').replace('</b>', '')
            run.font.bold = True
    return para


def add_table_data(document, headers, data, font_size=12):
    """添加表格"""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 设置表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        para = header_cells[i].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.size = Pt(font_size)
        para.runs[0].font.name = '宋体'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            para = row_cells[i].paragraphs[0]
            para.runs[0].font.size = Pt(font_size)
            para.runs[0].font.name = '宋体'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if isinstance(cell_data, (int, float)):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def add_image(document, image_path, width=Inches(5)):
    """添加图片到文档"""
    if os.path.exists(image_path):
        document.add_picture(image_path, width=width)
        # 设置图片居中
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    else:
        print(f" 图片不存在: {image_path}")
        return False


def add_section_break(document):
    """添加分页符"""
    document.add_page_break()


def generate_monte_carlo_charts_split(final_prices, issue_price, current_price, save_dir):
    """生成蒙特卡洛模拟图表 - 拆分版本"""
    import os

    returns = (final_prices - issue_price) / issue_price * 100
    chart_paths = []

    # 图1: 价格分布直方图（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.hist(final_prices, bins=50, alpha=0.7, color='steelblue', edgecolor='black')
    ax.axvline(x=issue_price, color='red', linestyle='--', linewidth=2, label=f'发行价: {issue_price:.2f}')
    ax.axvline(x=current_price, color='green', linestyle='--', linewidth=2, label=f'当前价: {current_price:.2f}')
    ax.set_xlabel('未来价格 (元)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('频数', fontproperties=font_prop, fontsize=14)
    ax.set_title('锁定期后价格分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_price_distribution_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2: 收益率分布（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.hist(returns, bins=50, alpha=0.7, color='coral', edgecolor='black')
    ax.axvline(x=0, color='red', linestyle='--', linewidth=2, label='盈亏平衡')
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('频数', fontproperties=font_prop, fontsize=14)
    ax.set_title('收益率分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_return_distribution_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3: 累积分布函数（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    sorted_returns = np.sort(returns)
    cumulative = np.arange(1, len(sorted_returns) + 1) / len(sorted_returns) * 100
    ax.plot(sorted_returns, cumulative, linewidth=2, color='#8e44ad')
    ax.axvline(x=0, color='red', linestyle='--', linewidth=1.5, label='盈亏平衡')
    ax.axhline(y=50, color='gray', linestyle='--', linewidth=1, alpha=0.5)
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('累积概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('累积分布函数', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_cdf_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图4: 盈亏概率饼图（单独大图）
    fig, ax = plt.subplots(figsize=(12, 10))
    profit_prob = (returns > 0).mean() * 100
    loss_prob = 100 - profit_prob
    colors_pie = ['#27ae60', '#e74c3c']
    wedges, texts, autotexts = ax.pie([profit_prob, loss_prob], labels=['盈利', '亏损'],
                                         autopct='%1.1f%%', colors=colors_pie, startangle=90,
                                         textprops={'fontproperties': font_prop, 'fontsize': 14})
    ax.set_title('盈亏概率分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    for autotext in autotexts:
        autotext.set_fontproperties(font_prop)
        autotext.set_fontsize(16)
        autotext.set_fontweight('bold')

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_profit_loss_pie_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths


def generate_chapter(context):
    """
    生成第五章内容：蒙特卡洛模拟

    参数:
        context: 包含所有必要数据的字典
            - document: Word文档对象
            - project_params: 项目参数
            - market_data: 市场数据
            - IMAGES_DIR: 图片保存目录
            - analyzer: 分析器对象
    """
    document = context['document']
    project_params = context['project_params']
    market_data = context['market_data']
    IMAGES_DIR = context['IMAGES_DIR']
    analyzer = context['analyzer']
    risk_params = context.get('risk_params', {})

    # ==================== 五、蒙特卡洛模拟 ====================
    add_title(document, '五、蒙特卡洛模拟', level=1)

    add_paragraph(document, '本章节使用蒙特卡洛方法模拟未来股价路径，评估收益分布。')

    add_title(document, '5.1 蒙特卡洛模拟基础分析', level=2)
    add_paragraph(document, '本节介绍蒙特卡洛模拟的理论基础、参数设置和基准结果。')
    add_paragraph(document, '')

    add_paragraph(document, '蒙特卡洛模拟是一种基于随机采样的数值计算方法，通过大量重复随机试验来估计复杂系统的行为。')
    add_paragraph(document, '')

    add_paragraph(document, ' 本项目采用的模型：几何布朗运动（Geometric Brownian Motion, GBM）')
    add_paragraph(document, '')

    add_paragraph(document, '模型假设：')
    add_paragraph(document, '• **股价的连续复利收益率服从正态分布**')
    add_paragraph(document, '  - 收益率 r = ln(S(t)/S(0)) ~ N(μ, σ²)')
    add_paragraph(document, '• **股价本身服从对数正态分布**')
    add_paragraph(document, '  - 股价 S(t) = S(0) × exp(r)，其中r为正态分布')
    add_paragraph(document, '• 股价变化具有连续性，不会出现跳跃')
    add_paragraph(document, '• 波动率（σ）和漂移率（μ）在模拟期间保持常数')
    add_paragraph(document, '• 市场有效，历史数据可反映未来特征')
    add_paragraph(document, '')

    add_paragraph(document, '📐 数学模型：')
    add_paragraph(document, '连续时间形式：dS = μSdt + σSdW')
    add_paragraph(document, '离散时间形式：S(t+Δt) = S(t) × exp((μ - σ²/2)Δt + σ√Δt × ε)')
    add_paragraph(document, '')

    add_paragraph(document, '其中：')
    add_paragraph(document, '• S(t)：时刻t的股价')
    add_paragraph(document, '• μ：漂移率（年化期望收益率）- 来自历史数据')
    add_paragraph(document, '• σ：波动率（年化标准差）- 来自历史数据')
    add_paragraph(document, '• Δt：时间步长')
    add_paragraph(document, '• ε：标准正态分布随机变量，ε ~ N(0,1)')
    add_paragraph(document, '')

    add_paragraph(document, ' 模拟步骤：')
    add_paragraph(document, '1. **参数估计**：根据历史数据计算年化漂移率（μ）和波动率（σ）')
    add_paragraph(document, '2. **时间离散化**：将锁定期T分割为n个时间步长（Δt = T/n）')
    add_paragraph(document, '3. **生成随机数**：对每个时间步长生成标准正态分布随机数 ε ~ N(0,1)')
    add_paragraph(document, '4. **路径模拟**：使用离散化公式计算下期价格：')
    add_paragraph(document, '   S(t+Δt) = S(t) × exp((μ - σ²/2)Δt + σ√Δt × ε)')
    add_paragraph(document, '5. **重复模拟**：重复步骤3-4共N次（N=10,000），得到N条可能的价格路径')
    add_paragraph(document, '6. **统计分析**：基于N条路径的终值，计算收益分布、盈利概率等统计量')
    add_paragraph(document, '')

    add_paragraph(document, '🎲 随机抽样特性：')
    add_paragraph(document, '• 每次模拟独立进行，使用固定随机种子（seed=42）确保结果可复现')
    add_paragraph(document, '• 通过大量模拟（10,000次），样本统计量收敛到真实分布特征')
    add_paragraph(document, '• 中心极限定理保证样本均值近似服从正态分布')
    add_paragraph(document, '')

    add_paragraph(document, ' 收益率计算说明：')
    add_paragraph(document, '• **锁定期收益率**（单利）：R = (S(T) - 发行价) / 发行价')
    add_paragraph(document, '• **年化收益率**（单利）：年化R = R × (12 / 锁定期月数)')
    add_paragraph(document, '• **预期收益率**：所有模拟路径年化收益率的算术平均值')
    add_paragraph(document, '• **收益率中位数**：所有模拟路径年化收益率的中位数（更稳健，不受极端值影响）')
    add_paragraph(document, '')

    # ==================== 模拟参数 ====================
    add_paragraph(document, '1. 模拟参数', bold=True)

    # 计算定增溢价率（相对于MA20，与项目概况保持一致）
    ma20_mc = market_data.get('ma_20', 0)
    if ma20_mc > 0:
        # 使用MA20作为基准（与项目概况一致）
        discount_rate = (project_params['issue_price'] - ma20_mc) / ma20_mc * 100
        discount_basis = f"（相对MA20: {ma20_mc:.2f}元）"
    else:
        # 如果没有MA20数据，使用当前价格作为基准
        discount_rate = (project_params['issue_price'] - project_params['current_price']) / project_params['current_price'] * 100
        discount_basis = f"（相对当前价: {project_params['current_price']:.2f}元）"

    # 使用250日窗口参数（对应年度解禁期，年线指标）
    # 保存为全局变量供后续章节（如VaR）使用
    mc_volatility_250d = market_data.get('volatility_250d', market_data.get('volatility', 0.30))
    mc_drift_250d = market_data.get('annual_return_250d', market_data.get('drift', 0.08))

    # 同时获取60日和120日窗口参数
    mc_volatility_60d = market_data.get('volatility_60d', market_data.get('volatility', 0.30))
    mc_drift_60d = market_data.get('annual_return_60d', market_data.get('drift', 0.08))

    mc_volatility_120d = market_data.get('volatility_120d', market_data.get('volatility', 0.30))
    mc_drift_120d = market_data.get('annual_return_120d', market_data.get('drift', 0.08))

    mc_params = [
        ['模拟次数', '10,000次'],
        ['锁定期', f'{project_params["lockup_period"]} 个月'],
        ['初始价格', f'{project_params["current_price"]:.2f} 元/股'],
        ['发行价格', f'{project_params["issue_price"]:.2f} 元/股'],
        ['定增溢价率（premium_rate）', f'{discount_rate:+.2f}% {discount_basis}'],
        [''],
        ['60日窗口（季度）参数', ''],
        ['  - 年化波动率', f'{mc_volatility_60d*100:.1f}%'],
        ['  - 年化漂移率', f'{mc_drift_60d*100:.1f}%'],
        [''],
        ['120日窗口（半年线）参数', ''],
        ['  - 年化波动率', f'{mc_volatility_120d*100:.1f}%'],
        ['  - 年化漂移率', f'{mc_drift_120d*100:.1f}%'],
        [''],
        ['250日窗口（年线）参数', ''],
        ['  - 年化波动率', f'{mc_volatility_250d*100:.1f}%'],
        ['  - 年化漂移率', f'{mc_drift_250d*100:.1f}%']
    ]
    add_table_data(document, ['参数', '值'], mc_params)

    add_paragraph(document, '')
    add_paragraph(document, '参数说明：')
    add_paragraph(document, '• 60日窗口：反映季度波动和收益特征，适合短期趋势分析')
    add_paragraph(document, '• 120日窗口：反映半年期波动和收益特征，平衡稳定性和时效性')
    add_paragraph(document, '• 250日窗口：反映年度波动和收益特征，数据最稳定，适合长期趋势分析')
    add_paragraph(document, '• 不同窗口期的参数差异会影响蒙特卡洛模拟结果，5.2节将对比分析')
    add_paragraph(document, '')

    # 默认使用250日窗口参数
    mc_volatility = mc_volatility_250d
    mc_drift = mc_drift_250d

    # ==================== 模拟结果 ====================
    add_paragraph(document, '2. 模拟结果（250日窗口基准）', bold=True)

    # 运行简化模拟
    lockup_days = project_params['lockup_period'] * 30
    n_simulations = 10000

    lockup_drift = mc_drift * (lockup_days / 365)
    lockup_vol = mc_volatility * np.sqrt(lockup_days / 365)

    np.random.seed(42)
    returns = np.random.normal(lockup_drift, lockup_vol, n_simulations)
    final_prices = project_params['current_price'] * np.exp(returns)
    profit_losses = (final_prices - project_params['issue_price']) / project_params['issue_price']

    # 统计
    profit_prob = (profit_losses > 0).mean()
    mean_return = profit_losses.mean()
    median_return = np.median(profit_losses)
    percent_5 = np.percentile(profit_losses, 5)
    percent_95 = np.percentile(profit_losses, 95)

    # ==================== 生成模拟路径可视化 ====================
    # 生成部分路径用于可视化（减少计算量）
    n_paths_visualize = 100  # 可视化100条路径
    n_steps = lockup_days  # 锁定期的天数

    # 每日漂移率和波动率
    daily_drift = mc_drift / 365
    daily_vol = mc_volatility / np.sqrt(365)

    # 生成路径
    np.random.seed(42)
    path_returns = np.random.normal(daily_drift, daily_vol, (n_paths_visualize, n_steps))

    # 计算价格路径（几何布朗运动）
    price_paths = np.zeros((n_paths_visualize, n_steps + 1))
    price_paths[:, 0] = project_params['current_price']

    for t in range(1, n_steps + 1):
        price_paths[:, t] = price_paths[:, t-1] * np.exp(path_returns[:, t-1])

    # 生成路径可视化图表
    fig, ax = plt.subplots(figsize=(14, 6))

    # 随机选择50条路径进行绘制
    n_paths_to_plot = min(50, n_paths_visualize)
    random_indices = random.sample(range(n_paths_visualize), n_paths_to_plot)

    # 绘制模拟路径
    for idx in random_indices:
        path = price_paths[idx, :]
        color = 'green' if path[-1] >= project_params['issue_price'] else 'red'
        ax.plot(range(n_steps + 1), path, color=color, alpha=0.15, linewidth=0.8)

    # 绘制关键水平线
    ax.axhline(y=project_params['issue_price'], color='blue', linestyle='--',
              linewidth=2, label=f"发行价格 ({project_params['issue_price']:.2f}元)")
    ax.axhline(y=project_params['current_price'], color='orange', linestyle='--',
              linewidth=2, label=f"当前价格 ({project_params['current_price']:.2f}元)")

    ax.set_xlabel('天数', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('股价 (元)', fontproperties=font_prop, fontsize=12)
    ax.set_title(f'蒙特卡洛模拟 - 股价路径可视化（{n_paths_to_plot}条样本路径）',
                fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, loc='upper left')
    ax.grid(True, alpha=0.3)

    plt.tight_layout()

    # 保存图表
    path_chart_path = os.path.join(IMAGES_DIR, '05_monte_carlo_paths.png')
    plt.savefig(path_chart_path, dpi=300, bbox_inches='tight')
    plt.close()

    print(f" 已生成蒙特卡洛路径可视化图表: {path_chart_path}")

    # 保存图表（拆分版）
    mc_chart_paths = generate_monte_carlo_charts_split(
        final_prices, project_params['issue_price'], project_params['current_price'], IMAGES_DIR)

    # 计算折价率
    discount_rate = (project_params["issue_price"] / project_params["current_price"] - 1) * 100

    # 创建合并的表格：模拟参数 + 结果
    mc_results = [
        ['模拟次数', f'{n_simulations:,}次'],
        ['锁定期', f'{project_params["lockup_period"]}个月（{lockup_days}天）'],
        ['发行价格', f'{project_params["issue_price"]:.2f}元/股'],
        ['当前价格', f'{project_params["current_price"]:.2f}元/股'],
        ['定价方式', f'{project_params.get("pricing_method", "N/A")}（折价率：{discount_rate:+.1f}%）'],
        ['年化漂移率', f'{mc_drift*100:+.2f}%'],
        ['年化波动率', f'{mc_volatility*100:.2f}%'],
        ['', ''],  # 空行分隔
        ['盈利概率', f'{profit_prob*100:.1f}%'],
        ['预期收益率', f'{mean_return*100:.1f}%'],
        ['收益率中位数', f'{median_return*100:.1f}%'],
        ['5%分位数（最差情况）', f'{percent_5*100:.1f}%'],
        ['95%分位数（较好情况）', f'{percent_95*100:.1f}%']
    ]
    add_table_data(document, ['指标', '值'], mc_results)

    add_paragraph(document, '')
    add_paragraph(document, '参数说明：')
    add_paragraph(document, '• 年化漂移率反映股价的长期趋势，正值表示上升趋势（基于250日窗口历史数据）')
    add_paragraph(document, '• 年化波动率反映股价的不确定性，值越大风险越高（基于250日窗口历史数据）')
    add_paragraph(document, '• 模拟基于几何布朗运动（GBM）模型，假设股价服从对数正态分布')
    add_paragraph(document, '• 锁定期内的收益率 = 漂移率×时间 + 波动率×√时间×随机因子')

    add_paragraph(document, '图表 5.1: 蒙特卡洛模拟 - 股价路径可视化')
    add_image(document, path_chart_path, width=Inches(6.5))
    add_paragraph(document, '')
    add_paragraph(document, '路径可视化说明：')
    add_paragraph(document, '• 绿色路径：锁定期末价格高于发行价（盈利）')
    add_paragraph(document, '• 红色路径：锁定期末价格低于发行价（亏损）')
    add_paragraph(document, '• 蓝色虚线：发行价格（盈利/亏损分界线）')
    add_paragraph(document, '• 橙色虚线：当前价格（模拟起点）')
    add_paragraph(document, f'• 显示{n_paths_to_plot}条随机样本路径，实际模拟{n_simulations:,}次')
    add_paragraph(document, '')

    add_paragraph(document, '图表 5.2: 蒙特卡洛模拟结果 - 价格分布')
    add_image(document, mc_chart_paths[0])

    add_paragraph(document, '图表 5.3: 蒙特卡洛模拟结果 - 收益率分布')
    add_image(document, mc_chart_paths[1])

    add_paragraph(document, '图表 5.4: 蒙特卡洛模拟结果 - 累积分布函数')
    add_image(document, mc_chart_paths[2])

    add_paragraph(document, '图表 5.5: 蒙特卡洛模拟结果 - 盈亏概率')
    add_image(document, mc_chart_paths[3])

    add_paragraph(document, '蒙特卡洛模拟结论：')
    add_paragraph(document, f'• 在 {n_simulations:,} 次模拟中，约 {profit_prob*100:.1f}% 的场景实现盈利')
    add_paragraph(document, f'• 预期收益率约 {mean_return*100:.1f}%')
    add_paragraph(document, f'• 95%的置信区间内，亏损可能不超过 {abs(percent_5)*100:.1f}%')

    # ==================== 5.2 多窗口期对比分析 ====================
    add_title(document, '5.2 多窗口期对比分析', level=2)

    add_paragraph(document, '本节分析不同时间窗口（60日/120日/250日）下的蒙特卡洛模拟结果，对比不同窗口期的波动率和收益率对定增收益的影响。')

    # 多窗口期配置
    windows_mc = {
        '60日': {'days': 60, 'vol_key': 'volatility_60d', 'return_key': 'annual_return_60d'},
        '120日': {'days': 120, 'vol_key': 'volatility_120d', 'return_key': 'annual_return_120d'},
        '250日': {'days': 250, 'vol_key': 'volatility_250d', 'return_key': 'annual_return_250d'}
    }

    # 为每个窗口期运行模拟
    multi_window_mc_results = {}
    mc_n_sim = 5000  # 使用5000次模拟以加快速度

    print("\n运行多窗口期蒙特卡洛模拟...")

    # 使用250日年化收益率作为基准（更稳定，避免极端值）
    base_annual_drift = market_data.get('annual_return_250d', risk_params.get('drift', 0.08))

    # 限制漂移率在合理范围内：-30%到+30%（避免极端值导致模拟失真）
    base_annual_drift = max(-0.30, min(0.30, base_annual_drift))

    print(f"  使用基准年化漂移率: {base_annual_drift*100:.2f}%（已限制在±30%范围内）")

    for window_name, config in windows_mc.items():
        print(f"  模拟 {window_name}...")

        # 获取该窗口期的波动率
        window_vol = market_data.get(config['vol_key'], risk_params.get('volatility', 0.30))

        # 所有窗口期使用统一的基准漂移率（避免极端值）
        window_drift = base_annual_drift

        # 计算时间步数（交易日转换为天，假设1月=21交易日）
        time_steps = config['days']

        # 运行模拟
        try:
            sim_window = analyzer.monte_carlo_simulation(
                n_simulations=mc_n_sim,
                time_steps=time_steps,
                volatility=window_vol,
                drift=window_drift,
                seed=42
            )

            # 提取锁定期末价格
            final_prices = sim_window.iloc[:, -1].values
            returns = (final_prices - project_params['issue_price']) / project_params['issue_price']

            # 年化收益率计算：使用固定系数（单利）
            # 60日=1/4(季度), 120日=1/2(半年), 250日=1(年度)
            if time_steps == 60:
                coefficient = 0.25  # 季度
            elif time_steps == 120:
                coefficient = 0.5   # 半年
            elif time_steps == 250:
                coefficient = 1.0   # 年度
            else:
                coefficient = time_steps / 252.0  # 其他情况用实际计算
            annualized_returns = returns / coefficient if coefficient > 0 else returns

            # 调试信息
            print(f"    调试: 当前价={project_params['current_price']:.2f}, 发行价={project_params['issue_price']:.2f}")
            print(f"    调试: 模拟期数={time_steps}日(系数={coefficient}), 锁定期={project_params['lockup_period']}月")
            print(f"    调试: 漂移率={window_drift*100:.2f}%, 波动率={window_vol*100:.2f}%")
            print(f"    调试: 最终价格均值={final_prices.mean():.2f}, 中位数={np.median(final_prices):.2f}")
            print(f"    调试: 期收益率均值={returns.mean()*100:.2f}%, 年化收益率均值={annualized_returns.mean()*100:.2f}%")
            print(f"    调试: 盈利次数={(returns>0).sum()}/{len(returns)}, 盈利概率={(returns>0).mean()*100:.1f}%")

            # 统计分析
            multi_window_mc_results[window_name] = {
                'volatility': window_vol,
                'drift': window_drift,
                'mean_return': annualized_returns.mean(),
                'median_return': np.median(annualized_returns),
                'std_return': annualized_returns.std(),
                'profit_prob': (returns > 0).mean() * 100,
                'percentile_5': np.percentile(annualized_returns, 5),
                'percentile_25': np.percentile(annualized_returns, 25),
                'percentile_75': np.percentile(annualized_returns, 75),
                'percentile_95': np.percentile(annualized_returns, 95),
            }

            print(f"    完成: 预期收益={multi_window_mc_results[window_name]['mean_return']*100:.2f}%, 盈利概率={multi_window_mc_results[window_name]['profit_prob']:.1f}%")

        except Exception as e:
            print(f"    错误: {e}")
            continue

    # 添加多窗口期对比表格
    add_paragraph(document, '')
    add_paragraph(document, '不同窗口期蒙特卡洛模拟对比：')
    add_paragraph(document, '下表展示在不同折价/溢价率情景下的模拟结果，涵盖深度折价到深度溢价的完整区间。')

    # 计算当前项目的折价/溢价率（相对于当前价格）
    current_premium_discount = (project_params['issue_price'] - project_params['current_price']) / project_params['current_price'] * 100

    # 修正表头，明确区分历史数据和模拟结果，删除折价/溢价率列（已在窗口期中体现）
    window_headers = ['窗口期', '波动率', '历史年化收益率', '模拟预期年化收益', '中位数收益', '盈利概率', '5% VaR', '95% VaR']
    window_table_data = []

    # 为每个窗口期，计算不同溢价率下的模拟结果
    # 溢价率：-20%, -15%, -10%, -5%（负值表示折价）
    # 溢价率：0%, 5%, 10%, 15%, 20%（正值表示溢价）
    discount_scenarios = [-20, -15, -10, -5, 0, 5, 10, 15, 20]

    for window_name in ['60日', '120日', '250日']:
        if window_name in multi_window_mc_results:
            r = multi_window_mc_results[window_name]

            # 显示当前项目的折价/溢价率（在窗口期列显示）
            discount_label = f"{window_name} ({current_premium_discount:+.1f}%)"
            window_table_data.append([
                discount_label,
                f"{r['volatility']*100:.2f}%",
                f"{r['drift']*100:.2f}%",  # 历史年化收益率（蒙特卡洛输入参数）
                f"{r['mean_return']*100:.2f}%",  # 模拟预期年化收益（蒙特卡洛输出）
                f"{r['median_return']*100:.2f}%",
                f"{r['profit_prob']:.1f}%",
                f"{r['percentile_5']*100:.2f}%",
                f"{r['percentile_95']*100:.2f}%"
            ])

            # 为该窗口期添加不同折价/溢价率情景的分析
            window_vol = r['volatility']
            window_drift = r['drift']
            time_steps = {'60日': 60, '120日': 120, '250日': 250}[window_name]

            # 对每个折价/溢价率运行模拟
            for rate in discount_scenarios:
                # 计算该折价/溢价率下的发行价
                # rate为负数表示折价（发行价低于市场价），正数表示溢价（发行价高于市场价）
                issue_price_with_rate = project_params['current_price'] * (1 + rate/100)

                sim_rate = analyzer.monte_carlo_simulation(
                    n_simulations=2000,  # 使用较少次数加快速度
                    time_steps=time_steps,
                    volatility=window_vol,
                    drift=window_drift,
                    seed=42
                )

                final_prices = sim_rate.iloc[:, -1].values
                returns = (final_prices - issue_price_with_rate) / issue_price_with_rate
                annualized_returns = returns * (12 / project_params['lockup_period'])

                # 根据折价/溢价率显示不同的标签
                if rate < 0:
                    rate_label = f"  └ 折价{abs(rate):.0f}%"
                elif rate == 0:
                    rate_label = "  └ 平价"
                else:
                    rate_label = f"  └ 溢价{rate:.0f}%"

                window_table_data.append([
                    rate_label,
                    "",  # 波动率留空
                    "",  # 历史收益率留空
                    f"{annualized_returns.mean()*100:.2f}%",  # 该折价/溢价率下的预期收益
                    f"{np.median(annualized_returns)*100:.2f}%",
                    f"{(returns > 0).mean()*100:.1f}%",
                    f"{np.percentile(annualized_returns, 5)*100:.2f}%",
                    f"{np.percentile(annualized_returns, 95)*100:.2f}%"
                ])

    add_table_data(document, window_headers, window_table_data)

    # 添加说明，解释各项指标含义
    add_paragraph(document, '')
    add_paragraph(document, '指标说明：')
    add_paragraph(document, '• 历史年化收益率：该窗口期内股票的实际历史表现，作为蒙特卡洛模拟的输入参数（drift，漂移率）')
    add_paragraph(document, '• 模拟预期年化收益率：')
    add_paragraph(document, '  - 基于历史波动率和漂移率，通过蒙特卡洛模拟5000次得出锁定期收益')
    add_paragraph(document, '  - 采用单利年化：年化收益 = 锁定期收益 × (12 ÷ 锁定期月数)')
    add_paragraph(document, '  - 例如：6个月锁定期收益10%，年化收益 = 10% × (12÷6) = 20%')
    add_paragraph(document, '  - 正值表示预期盈利，负值表示预期亏损')
    add_paragraph(document, '• 中位数收益：模拟结果的50%分位数，比平均值更稳健（不受极端值影响）')
    add_paragraph(document, '• 盈利概率：模拟中收益率为正的概率，反映投资获利的可能性')
    add_paragraph(document, '• 5% VaR：95%置信度下的最大损失，即只有5%的概率损失会超过此值')
    add_paragraph(document, '• 95% VaR：5%置信度下的最大收益，即只有5%的概率收益会超过此值（乐观情景）')
    add_paragraph(document, '• 窗口期标签中的百分比为当前定增项目的折价/溢价率（如-2.92%表示折价2.92%）')
    add_paragraph(document, '• 由于定增通常折价发行，即使历史收益率为负，模拟预期收益仍可能为正')

    # 生成多窗口期对比图表
    add_paragraph(document, '图表 5.7: 多窗口期蒙特卡洛模拟对比')

    try:
        # 检查是否有可用的窗口期数据
        if not multi_window_mc_results:
            raise ValueError("没有可用的窗口期数据，所有窗口期的蒙特卡洛模拟都失败了")

        # 创建对比图表
        fig, axes = plt.subplots(2, 2, figsize=(14, 10))

        window_names_ordered = ['60日', '120日', '250日']
        vols_mc = [multi_window_mc_results[w]['volatility']*100 for w in window_names_ordered if w in multi_window_mc_results]
        drifts_mc = [multi_window_mc_results[w]['drift']*100 for w in window_names_ordered if w in multi_window_mc_results]
        mean_returns_mc = [multi_window_mc_results[w]['mean_return']*100 for w in window_names_ordered if w in multi_window_mc_results]
        profit_probs_mc = [multi_window_mc_results[w]['profit_prob'] for w in window_names_ordered if w in multi_window_mc_results]

        # 更新窗口名称列表为实际有数据的窗口
        window_names_actual = [w for w in window_names_ordered if w in multi_window_mc_results]

        # 1. 波动率对比
        ax1 = axes[0, 0]
        bars1 = ax1.bar(window_names_actual, vols_mc, color='steelblue', alpha=0.7)
        ax1.set_ylabel('波动率 (%)', fontproperties=font_prop, fontsize=12)
        ax1.set_title('不同窗口期的波动率对比', fontproperties=font_prop, fontsize=13, fontweight='bold')
        ax1.grid(True, alpha=0.3, axis='y')
        for bar, vol in zip(bars1, vols_mc):
            ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                    f'{vol:.1f}%', ha='center', va='bottom', fontproperties=font_prop, fontsize=11)
        for label in ax1.get_xticklabels():
            label.set_fontproperties(font_prop)
        for label in ax1.get_yticklabels():
            label.set_fontproperties(font_prop)

        # 2. 收益率对比
        ax2 = axes[0, 1]
        bars2 = ax2.bar(window_names_actual, drifts_mc, color='coral', alpha=0.7)
        ax2.set_ylabel('收益率 (%)', fontproperties=font_prop, fontsize=12)
        ax2.set_title('不同窗口期的收益率对比', fontproperties=font_prop, fontsize=13, fontweight='bold')
        ax2.grid(True, alpha=0.3, axis='y')
        ax2.axhline(y=0, color='black', linestyle='--', linewidth=1)
        for bar, drift in zip(bars2, drifts_mc):
            y_pos = drift + 3 if drift > 0 else drift - 4
            ax2.text(bar.get_x() + bar.get_width()/2, y_pos,
                    f'{drift:.1f}%', ha='center', va='bottom' if drift > 0 else 'top',
                    fontproperties=font_prop, fontsize=11)
        for label in ax2.get_xticklabels():
            label.set_fontproperties(font_prop)
        for label in ax2.get_yticklabels():
            label.set_fontproperties(font_prop)

        # 3. 预期年化收益率对比
        ax3 = axes[1, 0]
        bars3 = ax3.bar(window_names_actual, mean_returns_mc,
                     color=['green' if r > 0 else 'red' for r in mean_returns_mc], alpha=0.7)
        ax3.set_ylabel('预期年化收益率 (%)', fontproperties=font_prop, fontsize=12)
        ax3.set_title('不同窗口期的预期收益率对比', fontproperties=font_prop, fontsize=13, fontweight='bold')
        ax3.grid(True, alpha=0.3, axis='y')
        ax3.axhline(y=0, color='black', linestyle='--', linewidth=1)
        for bar, ret in zip(bars3, mean_returns_mc):
            y_pos = ret + 3 if ret > 0 else ret - 4
            ax3.text(bar.get_x() + bar.get_width()/2, y_pos,
                    f'{ret:.1f}%', ha='center', va='bottom' if ret > 0 else 'top',
                    fontproperties=font_prop, fontsize=11)
        for label in ax3.get_xticklabels():
            label.set_fontproperties(font_prop)
        for label in ax3.get_yticklabels():
            label.set_fontproperties(font_prop)

        # 4. 盈利概率对比
        ax4 = axes[1, 1]
        bars4 = ax4.bar(window_names_actual, profit_probs_mc, color='purple', alpha=0.7)
        ax4.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=12)
        ax4.set_title('不同窗口期的盈利概率对比', fontproperties=font_prop, fontsize=13, fontweight='bold')
        ax4.set_ylim(0, 100)
        ax4.grid(True, alpha=0.3, axis='y')
        for bar, prob in zip(bars4, profit_probs_mc):
            ax4.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                    f'{prob:.1f}%', ha='center', va='bottom', fontproperties=font_prop, fontsize=11)
        for label in ax4.get_xticklabels():
            label.set_fontproperties(font_prop)
        for label in ax4.get_yticklabels():
            label.set_fontproperties(font_prop)

        plt.suptitle('蒙特卡洛模拟 - 多窗口期对比分析', fontproperties=font_prop, fontsize=15, fontweight='bold')
        plt.tight_layout()

        # 保存图表
        mc_multi_window_chart_path = os.path.join(IMAGES_DIR, 'monte_carlo_multi_window_comparison.png')
        plt.savefig(mc_multi_window_chart_path, dpi=150, bbox_inches='tight')
        plt.close()

        # 添加图表到文档
        add_image(document, mc_multi_window_chart_path, width=Inches(6))
        add_paragraph(document, '')

        # 添加分析结论
        add_paragraph(document, '多窗口期分析结论：')

        # 根据实际有数据的窗口期显示结论
        if '60日' in window_names_actual:
            idx_60 = window_names_actual.index('60日')
            add_paragraph(document, f'• 短期（60日）波动率约{vols_mc[idx_60]:.1f}%，更能反映当前市场情绪')
        if '120日' in window_names_actual:
            idx_120 = window_names_actual.index('120日')
            add_paragraph(document, f'• 中期（120日）波动率约{vols_mc[idx_120]:.1f}%，平衡了短期波动和长期趋势')
        if '250日' in window_names_actual:
            idx_250 = window_names_actual.index('250日')
            add_paragraph(document, f'• 长期（250日）波动率约{vols_mc[idx_250]:.1f}%，反映长期基本面特征')

        # 找出最佳和最差窗口
        best_window = window_names_actual[np.argmax(mean_returns_mc)]
        worst_window = window_names_actual[np.argmin(mean_returns_mc)]
        best_prob = max(profit_probs_mc)
        best_prob_window = window_names_actual[np.argmax(profit_probs_mc)]

        add_paragraph(document, f'• 预期收益率最高的是{best_window}窗口，约{max(mean_returns_mc):.1f}%')
        add_paragraph(document, f'• 盈利概率最高的是{best_prob_window}窗口，约{best_prob:.1f}%')
        add_paragraph(document, '')

        # 添加折价/溢价率的影响分析
        add_paragraph(document, '折价/溢价率影响分析（基于120日窗口）：')
        add_paragraph(document, f'• 当前项目折价/溢价率：{current_premium_discount:+.2f}%（{"折价发行" if current_premium_discount < 0 else "平价发行" if current_premium_discount == 0 else "溢价发行"}）')

        # 从表格数据中提取不同折价/溢价率的结果
        if '120日' in multi_window_mc_results:
            r_120 = multi_window_mc_results['120日']

            # 分析不同折价/溢价率对盈利概率的影响
            # 深度折价（-20%）：通常盈利概率最高
            # 深度溢价（+20%）：通常盈利概率最低

            add_paragraph(document, '• 深度折价（-20%）情景：安全边际充足，即使在悲观市场环境下也有较高盈利概率')
            add_paragraph(document, '• 中度折价（-10%至-15%）情景：平衡安全边际和投资机会，盈利概率显著提升')
            add_paragraph(document, '• 平价（0%）情景：无安全边际，盈利概率取决于市场趋势')
            add_paragraph(document, '• 溢价（+5%至+20%）情景：无安全边际且发行成本高，盈利概率较低，风险较大')

        add_paragraph(document, '')
        add_paragraph(document, '投资建议：')
        add_paragraph(document, '• 优先选择溢价率不高于-10%的定增项目（即折价≥10%），确保足够安全边际')
        add_paragraph(document, '• 对于溢价发行的项目，要求更高的预期收益率和更低的市场波动率')
        add_paragraph(document, '• 综合参考多个窗口期的结果，避免单一窗口期的偏差')

    except Exception as e:
        print(f" 生成多窗口期图表失败: {e}")
        add_paragraph(document, f' 多窗口期图表生成失败: {e}')

    # ==================== 5.3 ARIMA预测漂移率 ====================
    add_title(document, '5.3 ARIMA时间序列预测漂移率', level=2)

    add_paragraph(document, '本节使用ARIMA（自回归积分滑动平均）模型预测未来120日的漂移率（年化收益率）。')
    add_paragraph(document, '')
    add_paragraph(document, 'ARIMA模型原理：')
    add_paragraph(document, '• ARIMA(p,d,q)模型是一种经典的时间序列预测方法')
    add_paragraph(document, '• p：自回归阶数，利用历史值的线性组合预测当前值')
    add_paragraph(document, '• d：差分阶数，通过差分实现序列平稳化')
    add_paragraph(document, '• q：滑动平均阶数，利用历史预测误差的线性组合')
    add_paragraph(document, '• 适用于捕捉趋势性和周期性模式')
    add_paragraph(document, '')

    # ==================== 获取价格序列数据 ====================
    # 直接从market_data获取（由update_market_data.py生成）
    prices_list = market_data.get('price_series', [])

    if not prices_list or len(prices_list) < 100:
        print("\n" + "="*70)
        print(" 错误：price_series数据不足或缺失")
        print("="*70)
        print(f"   当前数据量: {len(prices_list) if prices_list else 0} 条")
        print(f"   需要数据量: 至少 100 条")
        print("")
        print("解决方案：")
        print("   请先运行数据更新脚本生成完整的市场数据")
        print("="*70)

        # 跳过时间序列预测部分
        arima_result = {'forecast_drift': mc_drift_120d, 'model_fitted': False}
        garch_result = {'forecast_volatility': mc_volatility_120d, 'model_fitted': False}
        skip_time_series = True
    else:
        skip_time_series = False
        print(f" 从market_data加载price_series: {len(prices_list)}个交易日")

    # 创建时间序列预测器
    if not skip_time_series:
        try:
            if prices_list and len(prices_list) >= 100:
                prices_series = pd.Series(prices_list)
                forecaster = TimeSeriesForecaster(prices_series)

                # ARIMA预测
                print("\n运行ARIMA模型预测漂移率...")
                arima_result = forecaster.forecast_drift_with_arima(horizon=120)

                # 添加模型设置说明
                add_paragraph(document, '模型设置：', bold=True)
                arima_params = [
                    ['模型类型', 'ARIMA(1,1,1)'],
                    ['预测期数', '120日（半年）'],
                    ['数据量', f'{len(prices_series)}个交易日'],
                    ['拟合状态', ' 成功' if arima_result['model_fitted'] else ' 失败（使用历史平均）']
                ]
                if arima_result['model_fitted']:
                    arima_params.append(['AIC准则', f"{arima_result['aic']:.2f}"])
                add_table_data(document, ['参数', '值'], arima_params)

                # 添加预测结果
                add_paragraph(document, '')
                add_paragraph(document, '预测结果：', bold=True)
                add_paragraph(document, f'• 预测漂移率（年化，对数收益率）：{arima_result["forecast_drift"]*100:.2f}%')
                add_paragraph(document, f'• 预测期间：未来120个交易日（约半年）')
                add_paragraph(document, f'• 含义：如果模型预测准确，未来半年股价预期年化对数收益率为{arima_result["forecast_drift"]*100:.2f}%（连续复利）')

                if not arima_result['model_fitted']:
                    add_paragraph(document, '')
                    add_paragraph(document, ' 注：ARIMA模型拟合失败，使用历史平均收益率作为预测值。这可能是因为：')
                    add_paragraph(document, '   - 数据量不足（建议至少250个交易日）')
                    add_paragraph(document, '   - 序列平稳性不足')
                    add_paragraph(document, '   - 存在异常值或结构断点')

                # 对比历史漂移率
                add_paragraph(document, '')
                add_paragraph(document, '与历史漂移率对比：', bold=True)
                historical_drifts = [
                    ['60日窗口', f'{mc_drift_60d*100:.2f}%'],
                    ['120日窗口', f'{mc_drift_120d*100:.2f}%'],
                    ['250日窗口', f'{mc_drift_250d*100:.2f}%'],
                    ['ARIMA预测', f'{arima_result["forecast_drift"]*100:.2f}%']
                ]
                add_table_data(document, ['数据来源', '年化漂移率'], historical_drifts)

                print(f" ARIMA预测完成：{arima_result['forecast_drift']*100:.2f}%")

            else:
                add_paragraph(document, ' 价格序列数据不足，无法进行ARIMA预测（需要至少100个交易日数据）')
                print(" 价格序列数据不足，跳过ARIMA预测")
                arima_result = {'forecast_drift': mc_drift_120d, 'model_fitted': False}

        except Exception as e:
            add_paragraph(document, f' ARIMA预测失败: {e}')
            print(f" ARIMA预测失败: {e}")
            arima_result = {'forecast_drift': mc_drift_120d, 'model_fitted': False}
    else:
        # 跳过5.3-5.5节（数据不足）
        print(" 数据不足，跳过5.3-5.5节时间序列预测")
        arima_result = {'forecast_drift': mc_drift_120d, 'model_fitted': False}
        garch_result = {'forecast_volatility': mc_volatility_120d, 'model_fitted': False}

    add_section_break(document)

    # ==================== 5.4 GARCH预测波动率 ====================
    add_title(document, '5.4 GARCH模型预测波动率', level=2)

    add_paragraph(document, '本节使用GARCH（广义自回归条件异方差）模型预测未来120日的波动率。')
    add_paragraph(document, '')
    add_paragraph(document, 'GARCH模型原理：')
    add_paragraph(document, '• GARCH(p,q)模型专门用于建模和预测波动率（金融时间序列的二阶矩）')
    add_paragraph(document, '• 能捕捉"波动率聚集"现象：高波动期后往往还是高波动，低波动后还是低波动')
    add_paragraph(document, '• ω：长期平均方差（波动率的基准水平）')
    add_paragraph(document, '• α（ARCH项）：过去冲击对当前波动的影响（新信息的冲击）')
    add_paragraph(document, '• β（GARCH项）：过去波动对当前波动的影响（波动的持续性）')
    add_paragraph(document, '• α+β < 1：模型平稳，波动率会回归长期均值')
    add_paragraph(document, '')

    # GARCH预测
    if not skip_time_series:
        try:
            if 'forecaster' in locals():
                print("\n运行GARCH模型预测波动率...")
                garch_result = forecaster.forecast_volatility_with_garch(horizon=120)

                # 添加模型设置说明
                add_paragraph(document, '模型设置：', bold=True)
                garch_params = [
                    ['模型类型', 'GARCH(1,1)'],
                    ['预测期数', '120日（半年）'],
                    ['数据量', f'{len(prices_series)}个交易日'],
                    ['拟合状态', ' 成功' if garch_result['model_fitted'] else ' 失败（使用历史波动率）']
                ]
                add_table_data(document, ['参数', '值'], garch_params)

                # 添加模型参数（如果拟合成功）
                if garch_result['model_fitted']:
                    add_paragraph(document, '')
                    add_paragraph(document, '模型参数：', bold=True)
                    model_params = [
                        ['ω（长期平均方差）', f"{garch_result['omega']:.6f}"],
                        ['α（ARCH系数）', f"{garch_result['alpha']:.4f}"],
                        ['β（GARCH系数）', f"{garch_result['beta']:.4f}"],
                        ['α+β（持续性）', f"{garch_result['alpha'] + garch_result['beta']:.4f}"]
                    ]
                    add_table_data(document, ['参数', '值'], model_params)
                    add_paragraph(document, '')
                    add_paragraph(document, '参数解读：')
                    add_paragraph(document, f'• α+β = {garch_result["alpha"] + garch_result["beta"]:.4f} {"< 1，模型平稳，波动率会逐渐回归长期均值" if garch_result["alpha"] + garch_result["beta"] < 1 else "≥ 1，模型非平稳，波动率可能发散"}')
                    if garch_result['alpha'] + garch_result['beta'] < 1:
                        half_life = -np.log(0.5) / (1 - (garch_result['alpha'] + garch_result['beta']))
                        add_paragraph(document, f'• 波动率半衰期：约{half_life:.0f}日（即冲击影响衰减一半所需时间）')

                # 添加预测结果
                add_paragraph(document, '')
                add_paragraph(document, '预测结果：', bold=True)
                add_paragraph(document, f'• 预测波动率（年化）：{garch_result["forecast_volatility"]*100:.2f}%')
                add_paragraph(document, f'• 预测期间：未来120个交易日')
                add_paragraph(document, f'• 含义：未来半年股价的预期年化波动率为{garch_result["forecast_volatility"]*100:.2f}%')

                if not garch_result['model_fitted']:
                    add_paragraph(document, '')
                    add_paragraph(document, ' 注：GARCH模型拟合失败，使用历史波动率作为预测值。')

                # 对比历史波动率
                add_paragraph(document, '')
                add_paragraph(document, '与历史波动率对比：', bold=True)
                historical_vols = [
                    ['60日窗口', f'{mc_volatility_60d*100:.2f}%'],
                    ['120日窗口', f'{mc_volatility_120d*100:.2f}%'],
                    ['250日窗口', f'{mc_volatility_250d*100:.2f}%'],
                    ['GARCH预测', f'{garch_result["forecast_volatility"]*100:.2f}%']
                ]
                add_table_data(document, ['数据来源', '年化波动率'], historical_vols)

                print(f" GARCH预测完成：{garch_result['forecast_volatility']*100:.2f}%")

            else:
                add_paragraph(document, ' 无法创建预测器，跳过GARCH预测')
                print(" 无法创建预测器，跳过GARCH预测")
                garch_result = {'forecast_volatility': mc_volatility_120d, 'model_fitted': False}

        except Exception as e:
            add_paragraph(document, f' GARCH预测失败: {e}')
            print(f" GARCH预测失败: {e}")
            garch_result = {'forecast_volatility': mc_volatility_120d, 'model_fitted': False}

    add_section_break(document)

    # ==================== 5.5 基于预测参数的蒙特卡洛模拟 ====================
    if not skip_time_series:
        add_title(document, '5.5 基于预测参数的蒙特卡洛模拟（120日）', level=2)

        add_paragraph(document, '本节使用5.3和5.4预测的漂移率和波动率，进行120日蒙特卡洛模拟，评估基于预测参数的定增收益风险。')
        add_paragraph(document, '')

        # 使用预测参数
        predicted_drift = arima_result['forecast_drift']
        predicted_vol = garch_result['forecast_volatility']

        # 运行蒙特卡洛模拟
        n_simulations = 10000
        time_steps = 120  # 120日

        print(f"\n运行基于预测参数的蒙特卡洛模拟（120日）...")
        print(f"  预测漂移率: {predicted_drift*100:.2f}%")
        print(f"  预测波动率: {predicted_vol*100:.2f}%")

        sim_predicted = analyzer.monte_carlo_simulation(
            n_simulations=n_simulations,
            time_steps=time_steps,
            volatility=predicted_vol,
            drift=predicted_drift,
            seed=42
        )

        # 计算结果
        final_prices = sim_predicted.iloc[:, -1].values
        returns = (final_prices - project_params['issue_price']) / project_params['issue_price']

        # 年化收益率计算：使用固定系数（单利）
        # 120日=1/2(半年)
        coefficient = 0.5
        annualized_returns = returns / coefficient

        # 统计
        profit_prob = (returns > 0).mean() * 100
        mean_return = annualized_returns.mean()
        median_return = np.median(annualized_returns)
        percentile_5 = np.percentile(annualized_returns, 5)
        percentile_95 = np.percentile(annualized_returns, 95)

        # 保存基于预测参数的模拟结果到context，供第九章使用
        context['results']['mc_predicted_120d'] = {
            'profit_prob': profit_prob / 100,  # 转换为小数
            'mean_return': mean_return / 100,   # 转换为小数
            'median_return': median_return / 100,
            'var_5': percentile_5 / 100,
            'var_95': percentile_95 / 100,
            'drift': predicted_drift,
            'volatility': predicted_vol
        }
        print(f" 已保存预测参数模拟结果到context：盈利概率{profit_prob:.1f}%，预期收益{mean_return:.2f}%")

        # 添加模拟参数说明
        add_paragraph(document, '模拟参数：', bold=True)
        mc_5_5_params = [
            ['当前价格', f'{project_params["current_price"]:.2f}元'],
            ['发行价格', f'{project_params["issue_price"]:.2f}元'],
            ['溢价率', f'{(project_params["issue_price"]/project_params["current_price"]-1)*100:+.2f}%'],
            ['预测漂移率', f'{predicted_drift*100:.2f}%（来自5.3节ARIMA预测）'],
            ['预测波动率', f'{predicted_vol*100:.2f}%（来自5.4节GARCH预测）'],
            ['模拟期数', f'{time_steps}日（120个交易日，约半年）'],
            ['模拟次数', f'{n_simulations:,}次'],
            ['锁定期', f'{project_params["lockup_period"]}个月']
        ]
        add_table_data(document, ['参数', '值'], mc_5_5_params)

        # 添加模拟结果
        add_paragraph(document, '')
        add_paragraph(document, '模拟结果：', bold=True)
        mc_5_5_results = [
            ['盈利概率', f'{profit_prob:.1f}%'],
            ['预期年化收益', f'{mean_return*100:.2f}%'],
            ['中位数收益', f'{median_return*100:.2f}%'],
            ['5% VaR', f'{percentile_5*100:.2f}%'],
            ['95% VaR', f'{percentile_95*100:.2f}%']
        ]
        add_table_data(document, ['指标', '值'], mc_5_5_results)

        # 添加结论说明
        add_paragraph(document, '')
        add_paragraph(document, '结果解读：', bold=True)
        add_paragraph(document, f'• 在{n_simulations:,}次模拟中，约{profit_prob:.1f}%的场景实现盈利')
        add_paragraph(document, f'• 预期年化收益率为{mean_return*100:.2f}%，{"高于" if mean_return > 0 else "低于"}无风险收益')
        add_paragraph(document, f'• 95%置信区间：年化收益在{percentile_5*100:.2f}%至{percentile_95*100:.2f}%之间')

        # 生成图表
        try:
            # 使用现有的图表生成函数
            mc_5_5_charts = generate_monte_carlo_charts_split(
                final_prices,
                project_params['issue_price'],
                project_params['current_price'],
                IMAGES_DIR
            )

            add_paragraph(document, '图表 5.8: 基于预测参数的蒙特卡洛模拟 - 价格分布')
            add_image(document, mc_5_5_charts[0])

            add_paragraph(document, '图表 5.9: 基于预测参数的蒙特卡洛模拟 - 收益率分布')
            add_image(document, mc_5_5_charts[1])

            add_paragraph(document, '图表 5.10: 基于预测参数的蒙特卡洛模拟 - 累积分布函数')
            add_image(document, mc_5_5_charts[2])

            add_paragraph(document, '图表 5.11: 基于预测参数的蒙特卡洛模拟 - 盈亏概率')
            add_image(document, mc_5_5_charts[3])

        except Exception as e:
            print(f" 生成图表失败: {e}")

        # 与历史参数方法对比
        add_paragraph(document, '')
        add_paragraph(document, '与历史参数方法对比：', bold=True)
        add_paragraph(document, '为了验证预测参数的有效性，下表对比基于预测参数和历史参数的模拟结果：')
        add_paragraph(document, '')

        # 使用历史250日参数模拟（对比基准）
        historical_drift = mc_drift_250d
        historical_vol = mc_volatility_250d

        sim_historical = analyzer.monte_carlo_simulation(
            n_simulations=n_simulations,
            time_steps=time_steps,
            volatility=historical_vol,
            drift=historical_drift,
            seed=42
        )

        final_prices_hist = sim_historical.iloc[:, -1].values
        returns_hist = (final_prices_hist - project_params['issue_price']) / project_params['issue_price']
        annualized_returns_hist = returns_hist * (12 / project_params['lockup_period'])

        # 统计
        profit_prob_hist = (returns_hist > 0).mean() * 100
        mean_return_hist = annualized_returns_hist.mean()

        # 对比表格
        comparison_table = [
            ['参数来源', '漂移率', '波动率', '盈利概率', '预期年化收益'],
            ['预测参数（ARIMA+GARCH）',
             f'{predicted_drift*100:.2f}%',
             f'{predicted_vol*100:.2f}%',
             f'{profit_prob:.1f}%',
             f'{mean_return*100:.2f}%'],
            ['历史参数（250日窗口）',
             f'{historical_drift*100:.2f}%',
             f'{historical_vol*100:.2f}%',
             f'{profit_prob_hist:.1f}%',
             f'{mean_return_hist*100:.2f}%'],
            ['', '', '', '', ''],
            ['差异',
             f'{(predicted_drift - historical_drift)*100:+.2f}%',
             f'{(predicted_vol - historical_vol)*100:+.2f}%',
             f'{profit_prob - profit_prob_hist:+.1f}%',
             f'{(mean_return - mean_return_hist)*100:+.2f}%']
        ]
        add_table_data(document, ['对比项', '漂移率', '波动率', '盈利概率', '预期收益'], comparison_table)

        # 对比分析
        add_paragraph(document, '')
        add_paragraph(document, '对比分析：', bold=True)
        add_paragraph(document, f'• 漂移率：预测参数{"高于" if predicted_drift > historical_drift else "低于"}历史参数，差{abs(predicted_drift - historical_drift)*100:.2f}个百分点')
        add_paragraph(document, f'• 波动率：预测参数{"高于" if predicted_vol > historical_vol else "低于"}历史参数，差{abs(predicted_vol - historical_vol)*100:.2f}个百分点')
        add_paragraph(document, f'• 盈利概率：预测参数{"更高" if profit_prob > profit_prob_hist else "更低"}，差{abs(profit_prob - profit_prob_hist):.1f}个百分点')
        add_paragraph(document, f'• 预期收益：预测参数{"更高" if mean_return > mean_return_hist else "更低"}，差{abs(mean_return - mean_return_hist)*100:.2f}个百分点')

        # 投资建议
        add_paragraph(document, '')
        add_paragraph(document, '投资建议：', bold=True)

        if mean_return > mean_return_hist and profit_prob > profit_prob_hist:
            add_paragraph(document, f' 基于预测参数的模拟结果优于历史参数，预期收益更高且盈利概率更大，建议积极考虑投资。')
        elif mean_return > mean_return_hist:
            add_paragraph(document, f' 基于预测参数的预期收益高于历史参数，但盈利概率{"略高" if profit_prob > profit_prob_hist else "略低"}，需结合其他因素综合判断。')
        elif profit_prob > profit_prob_hist:
            add_paragraph(document, f' 基于预测参数的盈利概率高于历史参数，但预期收益{"略高" if mean_return > mean_return_hist else "略低"}，安全性较好但收益空间有限。')
        else:
            add_paragraph(document, f' 基于预测参数的模拟结果全面劣于历史参数，建议谨慎投资或要求更高的折价率。')

        add_paragraph(document, '')
        add_paragraph(document, '说明：')
        add_paragraph(document, '• 预测参数基于ARIMA和GARCH模型，考虑了时间序列的动态特征')
        add_paragraph(document, '• 历史参数基于过去250个交易日的统计，反映长期平均水平')
        add_paragraph(document, '• 两种方法各有优劣，建议结合使用，互为验证')

        print(f" 5.5节完成：盈利概率{profit_prob:.1f}%，预期收益{mean_return*100:.2f}%")

        add_section_break(document)

    # 保存数据到context供后续章节使用
    context['results']['arima_result'] = arima_result
    context['results']['garch_result'] = garch_result
    context['results']['multi_window_mc_results'] = multi_window_mc_results
    print(f" 已保存multi_window_mc_results到context，共{len(multi_window_mc_results)}个窗口期")

    return context
