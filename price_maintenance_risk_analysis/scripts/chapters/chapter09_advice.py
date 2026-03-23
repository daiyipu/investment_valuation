"""
第九章：风控建议与风险提示

本章节从风险控制角度，给出盈亏平衡分析、核心指标汇总、最终报价建议和全面的风险提示。
基于保守原则，确保投资决策在合理风险可控范围内进行。
"""

def generate_chapter09(document, project_params, market_data, all_scenarios_for_appendix,
                       intrinsic_value, discount_premium, issue_type,
                       profit_prob, mean_return, var_95, var_99,
                       total_score, ma20_mc):
    """
    生成第九章：风控建议与风险提示

    参数说明：
    - document: Word文档对象
    - project_params: 项目参数字典
    - market_data: 市场数据字典
    - all_scenarios_for_appendix: 所有情景数据（用于附件）
    - intrinsic_value: DCF内在价值
    - discount_premium: 溢价率（相对MA20）
    - issue_type: 发行类型
    - profit_prob: 盈利概率
    - mean_return: 预期收益率
    - var_95: 95% VaR
    - var_99: 99% VaR
    - total_score: 风险评分
    - ma20_mc: MA20价格
    """
    from docx.shared import Inches
    from .utils import add_title, add_paragraph, add_table_data, add_image, add_section_break
    from .charts import generate_break_even_chart
    import os
    from datetime import datetime

    # 配置常量
    IMAGES_DIR = os.path.join(os.path.dirname(__file__), '..', 'outputs', 'images')

    # ==================== 第九章标题 ====================
    add_title(document, '九、风控建议与风险提示', level=1)

    add_paragraph(document, '本章节从风险控制角度，给出盈亏平衡分析、核心指标汇总、最终报价建议和全面的风险提示。')
    add_paragraph(document, '基于保守原则，确保投资决策在合理风险可控范围内进行。')
    add_paragraph(document, '')

    # ==================== 9.1 核心风险指标汇总 ====================
    add_title(document, '9.1 核心风险指标汇总', level=2)

    add_paragraph(document, '本节汇总核心风险指标，为投资决策提供全面的量化评估依据。')
    add_paragraph(document, '')

    # 处理盈利概率和预期收益率的格式
    profit_prob_display = profit_prob * 100 if profit_prob <= 1 else profit_prob
    mean_return_display = mean_return * 100 if abs(mean_return) <= 1 else mean_return

    # 统一使用120日窗口参数
    mc_volatility_120d = market_data.get('volatility_120d', market_data.get('volatility', 0.30))
    mc_drift_120d = market_data.get('annual_return_120d', market_data.get('drift', 0.08))

    # 计算评估值（避免在f-string中使用嵌套三元运算符）
    risk_level_eval = "低风险" if total_score >= 80 else "中等风险" if total_score >= 60 else "高风险"
    profit_level_eval = "较高" if profit_prob >= 0.7 else "中等" if profit_prob >= 0.5 else "较低"
    issue_type_eval = "折价发行" if discount_premium < 0 else "平价发行" if discount_premium < 5 else "溢价发行"

    summary_data = [
        ['风险评分', f'{total_score}/100', risk_level_eval],
        ['盈利概率', f'{profit_prob_display:.1f}%', profit_level_eval],
        ['发行类型', f'{issue_type}', issue_type_eval],
        ['溢价率（相对MA20）', f'{discount_premium:+.2f}%', ''],
        ['预期收益率', f'{mean_return_display:.1f}%', ''],
        ['95% VaR', f'{var_95*100:.1f}%', ''],
        ['波动率(120日)', f'{mc_volatility_120d*100:.1f}%', ''],
        ['年化漂移率(120日)', f'{mc_drift_120d*100:+.2f}%', ''],
        ['DCF内在价值', f'{intrinsic_value:.2f} 元/股', '']
    ]
    add_table_data(document, ['指标', '值', '评估'], summary_data)

    add_paragraph(document, '')

    # ==================== 9.2 盈亏平衡分析 ====================
    add_title(document, '9.2 盈亏平衡分析', level=2)
    add_paragraph(document, '本节通过盈亏平衡分析，评估在不同目标收益率下的安全边际。')
    add_paragraph(document, '')

    # 计算不同收益率下的盈亏平衡价格
    import numpy as np
    target_returns = np.linspace(0.05, 0.50, 10)
    break_even_prices = []
    issue_price = project_params['issue_price']
    lockup_years = project_params['lockup_period'] / 12
    current_price_eval = project_params['current_price']

    for target_return in target_returns:
        be_price = issue_price * (1 + target_return * lockup_years)
        break_even_prices.append(be_price)

    # 生成表格数据
    be_data = []
    for ret, price in zip(target_returns[::2], break_even_prices[::2]):
        distance = (current_price_eval - price) / current_price_eval * 100
        status = "✅" if distance > 0 else "⚠️"
        be_data.append([f'{ret*100:.0f}%', f'{price:.2f}', f'{distance:+.1f}%', status])

    add_table_data(document, ['期望年化收益率', '盈亏平衡价格(元)', '距离当前价', '状态'], be_data)

    add_paragraph(document, '')
    add_paragraph(document, '盈亏平衡分析结论：')
    add_paragraph(document, f'• 当前价格: {current_price_eval:.2f} 元')
    add_paragraph(document, f'• 发行价格: {issue_price:.2f} 元')
    add_paragraph(document, f'• 20%年化收益率要求下盈亏平衡价: {break_even_prices[3]:.2f} 元')

    be_20 = break_even_prices[3]
    if current_price_eval > be_20:
        margin = (current_price_eval - be_20) / current_price_eval * 100
        add_paragraph(document, f'• ✅ 当前价格高于20%收益率盈亏平衡价{margin:.1f}%，具有较好安全边际')
    else:
        gap = (be_20 - current_price_eval) / current_price_eval * 100
        add_paragraph(document, f'• ⚠️ 当前价格低于20%收益率盈亏平衡价{gap:.1f}%，安全边际不足')

    add_paragraph(document, '')

    # 生成并插入盈亏平衡价格敏感性图表
    break_even_chart_path = os.path.join(IMAGES_DIR, '10_01_break_even_analysis.png')
    generate_break_even_chart(issue_price, current_price_eval, project_params['lockup_period'], break_even_chart_path)
    add_paragraph(document, '图表 9.1: 盈亏平衡价格敏感性曲线')
    add_image(document, break_even_chart_path, width=Inches(6))

    add_paragraph(document, '')

    # ==================== 9.3 报价方案建议 ====================
    add_title(document, '9.3 报价方案建议', level=2)
    add_paragraph(document, '本节提供不同目标收益率下的报价建议，帮助投资者根据风险偏好选择合适的报价方案。')
    add_paragraph(document, '')

    # 显示计算参数
    current_price_eval = project_params['current_price']
    issue_price_eval = project_params['issue_price']
    lockup_years = project_params['lockup_period'] / 12
    historical_return = market_data.get('annual_return_120d', 0.08)
    expected_price = current_price_eval * (1 + historical_return * lockup_years)

    # 添加计算参数说明
    add_paragraph(document, '💡 计算参数（前置假设）：')
    add_paragraph(document, f'• 当前价格：{current_price_eval:.2f} 元/股')
    add_paragraph(document, f'• 锁定期：{project_params["lockup_period"]} 个月（{lockup_years:.2f} 年）')
    add_paragraph(document, f'• 历史年化收益率：{historical_return*100:+.2f}%（基于120日窗口数据）')
    add_paragraph(document, f'• 预期价格（锁定期末）：{expected_price:.2f} 元/股')
    add_paragraph(document, f'  计算公式：预期价格 = 当前价格 × (1 + 历史年化收益率 × 锁定期年数)')
    add_paragraph(document, f'          = {current_price_eval:.2f} × (1 + {historical_return*100:+.2f}% × {lockup_years:.2f})')
    add_paragraph(document, f'          = {expected_price:.2f} 元')
    add_paragraph(document, '')
    add_paragraph(document, '📊 报价计算逻辑：')
    add_paragraph(document, '• 目标：实现年化收益率 R（如15%/20%/25%）')
    add_paragraph(document, '• 公式：最高报价 = 预期价格 ÷ (1 + R × 锁定期年数)')
    add_paragraph(document, '• 溢价率：(最高报价 - 当前价格) ÷ 当前价格 × 100%')
    add_paragraph(document, '  （溢价率为负表示折价，为正表示溢价）')
    add_paragraph(document, '')

    # 计算不同目标收益率下的建议发行价
    target_returns_pricing = [0.15, 0.20, 0.25]
    pricing_recommendations = []

    for target_ret in target_returns_pricing:
        max_issue_price = expected_price / (1 + target_ret * lockup_years)
        premium_rate = (max_issue_price - current_price_eval) / current_price_eval * 100
        pricing_recommendations.append([
            f'{target_ret*100:.0f}%',
            f'{max_issue_price:.2f}',
            f'{premium_rate:+.2f}%'
        ])

    add_table_data(document, ['目标年化收益率', '最高报价(元)', '溢价率'], pricing_recommendations)
    add_paragraph(document, '')

    # 详细说明三个方案
    add_paragraph(document, '1. 保守型方案（15%目标收益）')
    add_paragraph(document, f'   • 建议报价：不高于{pricing_recommendations[0][1]}元（溢价率{pricing_recommendations[0][2]}）')
    add_paragraph(document, '   • 优势：折价较深，安全边际充足，盈利概率较高')
    add_paragraph(document, '   • 风险：可能因报价过低而错配机会，或无法获得足额配售')
    add_paragraph(document, '')

    add_paragraph(document, '2. 平衡型方案（20%目标收益）')
    add_paragraph(document, f'   • 建议报价：不高于{pricing_recommendations[1][1]}元（溢价率{pricing_recommendations[1][2]}）')
    add_paragraph(document, '   • 优势：收益与风险相对平衡，符合一般投资要求')
    add_paragraph(document, '   • 风险：对市场波动敏感，需密切关注120日窗口指标变化')
    add_paragraph(document, '')

    add_paragraph(document, '3. 积极型方案（25%目标收益）')
    add_paragraph(document, f'   • 建议报价：不高于{pricing_recommendations[2][1]}元（溢价率{pricing_recommendations[2][2]}）')
    add_paragraph(document, '   • 优势：可参与更多项目，提高资金使用效率')
    add_paragraph(document, '   • 风险：溢价率较高或溢价发行，安全边际有限，需承受较高波动风险')
    add_paragraph(document, '')

    # ==================== 9.4 当前发行价评估 ====================
    add_title(document, '9.4 当前发行价评估', level=2)
    add_paragraph(document, '本节评估当前发行价的合理性，并给出投资建议。')
    add_paragraph(document, '')

    # 评估当前发行价
    current_premium = (issue_price_eval - current_price_eval) / current_price_eval * 100
    add_paragraph(document, f'• 当前发行价: {issue_price_eval:.2f} 元')
    add_paragraph(document, f'• 当前价格: {current_price_eval:.2f} 元')
    add_paragraph(document, f'• 溢价率（相对当前价）: {current_premium:+.2f}%')

    # 说明：溢价率为负表示折价，为正表示溢价
    if current_premium > 0:
        add_paragraph(document, f'  （溢价{current_premium:.2f}%，发行价高于当前价）')
    elif current_premium < 0:
        add_paragraph(document, f'  （折价{abs(current_premium):.2f}%，发行价低于当前价）')
    else:
        add_paragraph(document, f'  （平价发行，发行价等于当前价）')

    add_paragraph(document, f'• 溢价率（相对MA20: {ma20_mc:.2f}元）: {discount_premium:+.2f}%')
    add_paragraph(document, '')

    # 判断当前发行价是否合理
    max_price_20 = expected_price / (1 + 0.20 * lockup_years)

    if issue_price_eval <= max_price_20:
        add_paragraph(document, f'• ✅ 当前发行价低于20%目标收益对应的最高报价({max_price_20:.2f}元)，具有投资价值')
        investment_advice = "建议积极参与"
    elif issue_price_eval <= max_price_20 * 1.05:
        add_paragraph(document, f'• ⚠️ 当前发行价接近20%目标收益对应的最高报价({max_price_20:.2f}元)，需谨慎评估')
        investment_advice = "建议谨慎参与或要求更高折价"
    else:
        add_paragraph(document, f'• ❌ 当前发行价高于20%目标收益对应的最高报价({max_price_20:.2f}元)，安全边际不足')
        investment_advice = "建议规避或等待更好时机"

    add_paragraph(document, '')

    # ==================== 9.5 主要风险提示 ====================
    add_title(document, '9.5 主要风险提示', level=2)
    add_paragraph(document, '基于多维度风险分析，提示以下主要风险（详见前文各章节详细分析）：')
    add_paragraph(document, '')

    # 1. 市场风险
    add_paragraph(document, '1. 市场风险')
    add_paragraph(document, f'   • 波动率风险：当前120日窗口年化波动率为{mc_volatility_120d*100:.1f}%，市场波动可能导致实际收益偏离预期')
    add_paragraph(document, f'   • 趋势风险：当前120日窗口年化漂移率为{mc_drift_120d*100:+.2f}%，{"上升趋势" if mc_drift_120d > 0 else "下降趋势" if mc_drift_120d < 0 else "震荡趋势"}可能影响解禁时收益')
    add_paragraph(document, '')

    # 2. 流动性风险
    add_paragraph(document, '2. 流动性风险')
    add_paragraph(document, f'   • 锁定期风险：{project_params["lockup_period"]}个月锁定期内无法交易，需承担期间价格波动')
    add_paragraph(document, '   • 解禁冲击：解禁后可能面临抛压，导致实际变现价格低于理论价格')
    add_paragraph(document, '')

    # 3. VaR在险价值风险
    add_paragraph(document, '3. VaR在险价值风险')
    var_95_display = var_95 * 100
    add_paragraph(document, f'   • 120日窗口：95%置信水平下最大可能亏损{var_95_display:.1f}%')
    add_paragraph(document, '   • 尾部风险：历史数据显示，小概率极端事件（黑天鹅）可能导致损失超过VaR预测值')
    add_paragraph(document, '')

    # 4. 估值风险
    add_paragraph(document, '4. 估值风险')
    add_paragraph(document, f'   • DCF估值风险：DCF内在价值{intrinsic_value:.2f}元/股基于多个假设，实际业绩可能偏离预测')
    add_paragraph(document, '   • 相对估值风险：PE/PS/PB相对估值基于行业平均水平，行业景气度变化可能导致估值体系重构')
    add_paragraph(document, '')

    # 5. 发行定价风险
    add_paragraph(document, '5. 发行定价风险')
    if current_premium < 0:
        discount_amount = abs(current_premium)
        add_paragraph(document, f'   • 折价情况：当前溢价率为{current_premium:+.2f}%（折价{discount_amount:.2f}%）')
        if discount_amount < 10:
            add_paragraph(document, f'   • 折价不足风险：折价率低于10%，安全边际有限')
        else:
            add_paragraph(document, f'   • 折价合理：折价{discount_amount:.2f}%提供了一定的安全边际')
    else:
        add_paragraph(document, f'   • 溢价风险：当前溢价率为{current_premium:+.2f}%（溢价发行）')
        add_paragraph(document, f'   • 溢价发行无安全边际，需重点关注公司成长性')
    add_paragraph(document, '   • 定价偏离：若发行价显著高于盈亏平衡价格，将大幅降低盈利概率和预期收益')
    add_paragraph(document, '')

    # 6. 其他风险
    add_paragraph(document, '6. 其他风险')
    add_paragraph(document, '   • 行业政策风险：需关注行业监管政策变化')
    add_paragraph(document, '   • 业绩波动风险：需关注公司业绩预告、审计报告等')
    add_paragraph(document, '   • 竞争格局风险：行业竞争加剧可能影响盈利能力')
    add_paragraph(document, '')

    add_paragraph(document, '💡 综合建议：')
    risk_advice_text = "风险较低，可积极参与" if total_score >= 80 else "风险中等，需谨慎评估" if total_score >= 60 else "风险较高，建议规避"
    add_paragraph(document, f'• 当前项目风险评分{total_score}/100分，{risk_advice_text}')
    add_paragraph(document, f'• 推荐方案：{investment_advice}')
    add_paragraph(document, '• 建议结合个人风险承受能力、资金成本和市场环境选择合适的报价方案')
    add_paragraph(document, '')
    if total_score >= 80 and issue_price_eval <= max_price_20:
        final_recommendation = f"🟢 建议积极参与"
        reason = f"风险评分{total_score}/100（优秀），发行价具有较好安全边际，符合20%目标收益率要求。"
    elif total_score >= 60:
        final_recommendation = f"🟡 谨慎参与"
        reason = f"风险评分{total_score}/100（中等），需结合溢价率和增长预期综合评估。"
    else:
        final_recommendation = f"🔴 建议规避"
        reason = f"风险评分{total_score}/100（较低），安全边际不足，风险较高。"

    add_paragraph(document, f'{final_recommendation}')
    add_paragraph(document, reason)

    add_paragraph(document, '')

    # ==================== 9.6 风控策略建议 ====================
    add_title(document, '9.6 风控策略建议', level=2)
    add_paragraph(document, '根据当前溢价率和盈利概率，提供风控策略建议。')
    add_paragraph(document, '')

    # 使用current_premium（溢价率），为负表示折价
    if current_premium <= -15:
        add_paragraph(document, '• 当前溢价率≤-15%（折价≥15%），可按原计划参与')
    elif current_premium <= -10:
        add_paragraph(document, '• 当前溢价率-10%~-15%（折价10%~15%），建议关注公司基本面和行业前景')
    else:
        add_paragraph(document, f'• 当前溢价率{current_premium:+.2f}%（折价不足或溢价），建议要求更高折价或等待更好时机')

    # 根据盈利概率给出建议（统一处理百分比格式）
    profit_prob_pct = profit_prob * 100 if profit_prob <= 1 else profit_prob
    if profit_prob_pct >= 70:
        add_paragraph(document, f'• 盈利概率{profit_prob_pct:.1f}%（≥70%），投资安全边际充足')
    elif profit_prob_pct >= 50:
        add_paragraph(document, f'• 盈利概率{profit_prob_pct:.1f}%（50%-70%），建议分批参与或控制仓位')
    else:
        add_paragraph(document, f'• 盈利概率{profit_prob_pct:.1f}%（<50%），风险较大，建议谨慎')

    add_paragraph(document, '')

    # ==================== 9.7 压力测试与VaR风险提示 ====================
    add_title(document, '9.7 压力测试与VaR风险提示', level=2)
    add_paragraph(document, '本节汇总压力测试和VaR分析的关键风险指标。')
    add_paragraph(document, '')

    # 从多窗口期VaR分析中提取关键风险指标
    var_95_pct = var_95 * 100
    var_99_pct = var_99 * 100 if 'var_99' in locals() else None

    add_paragraph(document, f'• 120日窗口VaR风险：95%置信水平下最大可能亏损为{var_95_pct:.1f}%')
    if var_99_pct:
        add_paragraph(document, f'• 极端情况VaR：99%置信水平下最大可能亏损为{var_99_pct:.1f}%')

    # 压力测试风险提示
    add_paragraph(document, '• 压力测试情景：需关注PE回归、市场危机、三重打击等极端情景下的潜在损失')
    add_paragraph(document, '• 波动率放大风险：当实际波动率超过120日窗口统计值时，风险敞口将显著增加')

    add_paragraph(document, '')

    # ==================== 9.8 多重方案选项建议 ====================
    add_title(document, '9.8 多重方案选项建议', level=2)
    add_paragraph(document, '根据不同的风险偏好，提供以下参与方案：')
    add_paragraph(document, '')

    add_paragraph(document, '根据不同的风险偏好，提供以下参与方案：')
    add_paragraph(document, '')

    # 方案一：保守型（低风险偏好）
    add_paragraph(document, '🛡️ 方案一：保守型（低风险偏好）')
    add_paragraph(document, f'   • 适用场景：追求确定性，可接受15%年化收益率')
    max_price_conservative = expected_price / (1 + 0.15 * lockup_years)
    discount_conservative = (current_price_eval - max_price_conservative) / current_price_eval * 100
    add_paragraph(document, f'   • 建议报价：不高于{max_price_conservative:.2f}元（折价率{discount_conservative:+.2f}%）')
    add_paragraph(document, f'   • 风险控制：要求较高折价，确保足够安全边际')
    add_paragraph(document, '')

    # 方案二：平衡型（中等风险偏好）
    add_paragraph(document, '⚖️ 方案二：平衡型（中等风险偏好）')
    add_paragraph(document, f'   • 适用场景：平衡收益与风险，目标20%年化收益率')
    max_price_balanced = expected_price / (1 + 0.20 * lockup_years)
    discount_balanced = (current_price_eval - max_price_balanced) / current_price_eval * 100
    add_paragraph(document, f'   • 建议报价：不高于{max_price_balanced:.2f}元（折价率{discount_balanced:+.2f}%）')
    add_paragraph(document, f'   • 风险控制：适度折价，关注盈利概率和VaR指标')
    add_paragraph(document, '')

    # 方案三：积极型（高风险偏好）
    add_paragraph(document, '🚀 方案三：积极型（高风险偏好）')
    add_paragraph(document, f'   • 适用场景：看好公司长期成长，可承受较大波动，目标25%年化收益率')
    max_price_aggressive = expected_price / (1 + 0.25 * lockup_years)
    discount_aggressive = (current_price_eval - max_price_aggressive) / current_price_eval * 100
    add_paragraph(document, f'   • 建议报价：不高于{max_price_aggressive:.2f}元（折价率{discount_aggressive:+.2f}%）')
    add_paragraph(document, f'   • 风险控制：需重点关注公司基本面和行业前景，建议控制仓位')
    add_paragraph(document, '')

    # 风险提示
    add_paragraph(document, '⚠️ 重要提示：')
    add_paragraph(document, '• 本建议基于120日窗口历史数据，实际收益可能偏离预期')
    add_paragraph(document, '• 市场环境变化、公司业绩波动等风险因素可能影响最终收益')
    add_paragraph(document, '• 建议结合最新市场情况和公司公告动态调整报价策略')
    add_paragraph(document, '• 投资有风险，决策需谨慎')

    add_paragraph(document, '')

    # ==================== 9.9 免责声明 ====================
    add_title(document, '9.9 免责声明', level=2)

    disclaimer = f'''
    本报告由自动化分析系统生成，仅供参考使用，不构成投资建议。

    1. 本报告基于历史数据和公开信息进行分析，不保证数据的准确性和完整性；
    2. 市场有风险，投资需谨慎。本报告中的任何分析观点不代表未来表现；
    3. 本报告提到的任何证券或投资标的，仅为分析示例，不构成推荐；
    4. 投资者应根据自身情况独立判断，自行承担投资风险；
    5. 本报告的知识产权归分析团队所有，未经许可不得转载或使用。

    报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    '''

    add_paragraph(document, disclaimer, font_size=9)
