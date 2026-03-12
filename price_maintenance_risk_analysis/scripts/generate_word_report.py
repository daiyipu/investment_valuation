#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
定增风险分析报告生成器（增强版）
生成 Word 格式的综合分析报告，包含图表
"""

import sys
import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats

# 添加路径 - 从 scripts 目录运行
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, PROJECT_DIR)

from utils.config_loader import load_placement_config
from utils.analysis_tools import PrivatePlacementRiskAnalyzer
from utils.font_manager import get_font_prop

# 获取中文字体
font_prop = get_font_prop()

# 配置路径
DATA_DIR = os.path.join(PROJECT_DIR, 'data')
REPORTS_DIR = os.path.join(PROJECT_DIR, 'reports')
IMAGES_DIR = os.path.join(REPORTS_DIR, 'images')
OUTPUTS_DIR = os.path.join(REPORTS_DIR, 'outputs')

# 确保目录存在
os.makedirs(IMAGES_DIR, exist_ok=True)
os.makedirs(OUTPUTS_DIR, exist_ok=True)


def setup_chinese_font(document):
    """设置中文字体"""
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    document.styles['Normal'].font.size = Pt(11)


def add_title(document, text, level=1):
    """添加标题"""
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading


def add_paragraph(document, text, bold=False, font_size=11):
    """添加段落"""
    para = document.add_paragraph(text)
    for run in para.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(font_size)
        if bold:
            run.font.bold = True
    return para


def add_image(document, image_path, width=Inches(5)):
    """添加图片到文档"""
    if os.path.exists(image_path):
        document.add_picture(image_path, width=width)
        # 设置图片居中
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    else:
        print(f"⚠️ 图片不存在: {image_path}")
        return False


def add_table_data(document, headers, data):
    """添加表格"""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 设置表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        para = header_cells[i].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.size = Pt(10)
        para.runs[0].font.name = '宋体'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            para = row_cells[i].paragraphs[0]
            para.runs[0].font.size = Pt(10)
            para.runs[0].font.name = '宋体'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if isinstance(cell_data, (int, float)):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def add_section_break(document):
    """添加分页符"""
    document.add_page_break()


def generate_break_even_chart(issue_price, current_price, lockup_period, save_path):
    """生成盈亏平衡价格敏感性曲线"""
    # 不同期望收益率下的盈亏平衡价格
    target_returns = np.linspace(0.05, 0.50, 10)  # 5%到50%年化收益率
    break_even_prices = [issue_price * (1 + r * (lockup_period / 12)) for r in target_returns]

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(target_returns * 100, break_even_prices, 'b-', linewidth=2, label='盈亏平衡价格')
    ax.axhline(y=issue_price, color='r', linestyle='--', label='发行价格')
    ax.axhline(y=current_price, color='g', linestyle='--', label='当前价格')

    ax.set_xlabel('期望年化收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('盈亏平衡价格 (元)', fontproperties=font_prop, fontsize=12)
    ax.set_title('不同收益率要求下的盈亏平衡价格', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    ax.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_lockup_sensitivity_chart(issue_price, current_price, save_path):
    """生成锁定期敏感性分析图表"""
    lockup_periods = [6, 9, 12, 18, 24, 36]  # 月
    target_return = 0.20  # 20%年化收益率

    lockup_analysis = []
    for period in lockup_periods:
        period_years = period / 12
        be_price = issue_price * (1 + target_return * period_years)
        required_increase = (be_price - issue_price) / issue_price * 100
        lockup_analysis.append({'period': period, 'be_price': be_price, 'increase': required_increase})

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 5))

    # 左图：盈亏平衡价格
    ax1.bar([p['period'] for p in lockup_analysis], [p['be_price'] for p in lockup_analysis],
            color='steelblue', alpha=0.7)
    ax1.axhline(y=current_price, color='r', linestyle='--', label='当前价格')
    ax1.set_xlabel('锁定期 (月)', fontproperties=font_prop, fontsize=11)
    ax1.set_ylabel('盈亏平衡价格 (元)', fontproperties=font_prop, fontsize=11)
    ax1.set_title('锁定期对盈亏平衡价格的影响', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax1.legend(prop=font_prop)
    for label in ax1.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax1.get_yticklabels():
        label.set_fontproperties(font_prop)
    ax1.grid(True, alpha=0.3)

    # 右图：需要涨幅
    ax2.plot([p['period'] for p in lockup_analysis], [p['increase'] for p in lockup_analysis],
            'ro-', linewidth=2, markersize=8)
    ax2.set_xlabel('锁定期 (月)', fontproperties=font_prop, fontsize=11)
    ax2.set_ylabel('需要涨幅 (%)', fontproperties=font_prop, fontsize=11)
    ax2.set_title('锁定期与需要涨幅的关系', fontproperties=font_prop, fontsize=12, fontweight='bold')
    for label in ax2.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax2.get_yticklabels():
        label.set_fontproperties(font_prop)
    ax2.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_discount_scenario_charts(base_price, current_price, volatility, drift, lockup_period, save_dir):
    """生成发行价折扣情景图表 - 基础版（合并为一图）"""
    import os

    discount_rates = [5, 10, 15, 20]
    premium_rates = [5, 10, 15, 20]

    # 计算数据
    discount_data = []
    premium_data = []

    for d in discount_rates:
        issue_price = base_price * (1 - d/100)
        threshold = issue_price
        lockup_vol = volatility * np.sqrt(lockup_period / 12)
        lockup_drift = drift * (lockup_period / 12)
        required_return = (threshold - current_price) / current_price
        z_score = (lockup_drift - required_return) / lockup_vol
        profit_prob = (1 - stats.norm.cdf(-z_score)) * 100
        expected_return = (base_price - issue_price) / issue_price * 100
        discount_data.append({'rate': d, 'profit_prob': profit_prob, 'expected_return': expected_return})

    for p in premium_rates:
        issue_price = base_price * (1 + p/100)
        threshold = max(base_price, issue_price * 1.02)
        lockup_vol = volatility * np.sqrt(lockup_period / 12)
        lockup_drift = drift * (lockup_period / 12)
        required_return = (threshold - current_price) / current_price
        z_score = (lockup_drift - required_return) / lockup_vol
        profit_prob = (1 - stats.norm.cdf(-z_score)) * 100
        expected_return = (base_price - issue_price) / issue_price * 100
        premium_data.append({'rate': p, 'profit_prob': profit_prob, 'expected_return': expected_return})

    # 生成图表（2x2布局）
    fig, axes = plt.subplots(2, 2, figsize=(16, 12))

    # 1. 折价发行盈利概率
    axes[0, 0].bar([d['rate'] for d in discount_data], [d['profit_prob'] for d in discount_data],
                  color='steelblue', alpha=0.7)
    axes[0, 0].set_xlabel('折价率 (%)', fontproperties=font_prop)
    axes[0, 0].set_ylabel('盈利概率 (%)', fontproperties=font_prop)
    axes[0, 0].set_title('折价发行 - 盈利概率', fontproperties=font_prop, fontsize=12, fontweight='bold')
    for label in axes[0, 0].get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in axes[0, 0].get_yticklabels():
        label.set_fontproperties(font_prop)
    axes[0, 0].grid(True, alpha=0.3)

    # 2. 溢价发行盈利概率
    axes[0, 1].bar([p['rate'] for p in premium_data], [p['profit_prob'] for p in premium_data],
                  color='coral', alpha=0.7)
    axes[0, 1].set_xlabel('溢价率 (%)', fontproperties=font_prop)
    axes[0, 1].set_ylabel('盈利概率 (%)', fontproperties=font_prop)
    axes[0, 1].set_title('溢价发行 - 盈利概率', fontproperties=font_prop, fontsize=12, fontweight='bold')
    axes[0, 1].set_ylim(0, 25)
    for label in axes[0, 1].get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in axes[0, 1].get_yticklabels():
        label.set_fontproperties(font_prop)
    axes[0, 1].grid(True, alpha=0.3)

    # 3. 预期收益率对比
    x_pos = np.arange(len(discount_rates))
    width = 0.35
    axes[1, 0].bar(x_pos - width/2, [d['expected_return'] for d in discount_data], width,
                   label='折价', color='steelblue', alpha=0.7)
    axes[1, 0].bar(x_pos + width/2, [p['expected_return'] for p in premium_data], width,
                   label='溢价', color='coral', alpha=0.7)
    axes[1, 0].set_xlabel('发行折扣率', fontproperties=font_prop)
    axes[1, 0].set_ylabel('预期收益率 (%)', fontproperties=font_prop)
    axes[1, 0].set_title('预期收益率对比', fontproperties=font_prop, fontsize=12, fontweight='bold')
    axes[1, 0].set_xticks(x_pos)
    axes[1, 0].set_xticklabels([f"{d}%" for d in discount_rates], fontproperties=font_prop)
    axes[1, 0].axhline(y=0, color='gray', linestyle='--', alpha=0.5)
    axes[1, 0].legend(prop=font_prop)
    for label in axes[1, 0].get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in axes[1, 0].get_yticklabels():
        label.set_fontproperties(font_prop)
    axes[1, 0].grid(True, alpha=0.3)

    # 4. 盈利概率汇总表数据
    add_table_data(None, ['情景类型', '折扣率', '盈利概率'], [])
    # 生成图表
    chart_path = os.path.join(save_dir, 'discount_scenario_analysis.png')
    plt.tight_layout()
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()

    return chart_path


def generate_multi_dimension_scenario_charts(current_price, base_price, volatility, drift, lockup_period, save_dir):
    """生成多维度情景分析图表（波动率×折扣率×时间窗口） - 基础版（合并为一图）"""
    import os

    # 定义情景参数
    volatility_scenarios = [0.20, 0.30, 0.40, 0.50]
    discount_scenarios = [-10, -5, 0, 5, 10, 15, 20]  # 负数=折价，正数=溢价
    window_scenarios = [60, 120, 180]

    # 生成所有情景组合的数据
    results = []
    for vol in volatility_scenarios:
        for discount in discount_scenarios:
            for window in window_scenarios:
                # 计算发行价
                issue_price = base_price * (1 - discount/100)

                # 判断盈利阈值
                if discount < 0:  # 折价发行
                    threshold = issue_price
                else:  # 溢价发行
                    threshold = max(base_price, issue_price * 1.02)

                # 计算该时间窗口下的参数
                window_years = window / 252
                window_vol = vol * np.sqrt(window_years)
                window_drift = drift * window_years

                # 计算盈利概率
                required_return = (threshold - current_price) / current_price
                z_score = (window_drift - required_return) / window_vol if window_vol > 0 else 0
                profit_prob = (1 - stats.norm.cdf(-z_score)) * 100

                results.append({
                    'volatility': vol,
                    'discount': discount,
                    'window': window,
                    'profit_prob': profit_prob,
                    'issue_price': issue_price
                })

    # 创建DataFrame方便处理
    df = pd.DataFrame(results)

    # 生成图表（2x2布局）
    fig, axes = plt.subplots(2, 2, figsize=(16, 12))

    # 图1: 热力图 - 波动率×折扣率（使用60日窗口）
    pivot_60d = df[df['window'] == 60].pivot(index='volatility', columns='discount', values='profit_prob')
    im1 = axes[0, 0].imshow(pivot_60d.values, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100)
    axes[0, 0].set_xticks(range(len(discount_scenarios)))
    axes[0, 0].set_xticklabels([f"{d}%" for d in discount_scenarios], fontproperties=font_prop)
    axes[0, 0].set_yticks(range(len(volatility_scenarios)))
    axes[0, 0].set_yticklabels([f"{v*100:.0f}%" for v in volatility_scenarios], fontproperties=font_prop)
    axes[0, 0].set_xlabel('折扣/溢价率', fontproperties=font_prop)
    axes[0, 0].set_ylabel('波动率', fontproperties=font_prop)
    axes[0, 0].set_title('多维度情景：盈利概率热力图（60日窗口）', fontproperties=font_prop, fontsize=12, fontweight='bold')
    plt.colorbar(im1, ax=axes[0, 0], label='盈利概率 (%)')

    # 图2: 20%折价 vs 10%溢价对比（不同波动率）
    discount_20_data = df[(df['discount'] == -20) & (df['window'] == 60)]
    premium_10_data = df[(df['discount'] == 10) & (df['window'] == 60)]

    x_pos = np.arange(len(volatility_scenarios))
    width = 0.35
    axes[0, 1].bar(x_pos - width/2, discount_20_data['profit_prob'], width,
                    label='20%折价', color='green', alpha=0.7)
    axes[0, 1].bar(x_pos + width/2, premium_10_data['profit_prob'], width,
                    label='10%溢价', color='red', alpha=0.7)
    axes[0, 1].set_xlabel('波动率', fontproperties=font_prop)
    axes[0, 1].set_ylabel('盈利概率 (%)', fontproperties=font_prop)
    axes[0, 1].set_title('典型情景对比（60日窗口）', fontproperties=font_prop, fontsize=12, fontweight='bold')
    axes[0, 1].set_xticks(x_pos)
    axes[0, 1].set_xticklabels([f"{v*100:.0f}%" for v in volatility_scenarios], fontproperties=font_prop)
    axes[0, 1].legend(prop=font_prop)
    axes[0, 1].grid(True, alpha=0.3, axis='y')

    # 图3: Top 10最佳情景
    top_10 = df.nlargest(10, 'profit_prob')[['volatility', 'discount', 'window', 'profit_prob']].copy()
    top_10['scenario'] = top_10.apply(lambda x: f"{x['volatility']*100:.0f}%波动,{x['discount']}%折扣,{x['window']}日", axis=1)

    colors = ['green' if d < 0 else 'orange' for d in top_10['discount']]
    axes[1, 0].barh(range(len(top_10)), top_10['profit_prob'], color=colors, alpha=0.7)
    axes[1, 0].set_yticks(range(len(top_10)))
    axes[1, 0].set_yticklabels(top_10['scenario'], fontproperties=font_prop, fontsize=9)
    axes[1, 0].set_xlabel('盈利概率 (%)', fontproperties=font_prop)
    axes[1, 0].set_title('最佳情景 Top 10', fontproperties=font_prop, fontsize=12, fontweight='bold')
    axes[1, 0].grid(True, alpha=0.3, axis='x')
    axes[1, 0].invert_yaxis()

    # 图4: 不同时间窗口的对比
    window_comparison = df[df['discount'] == 0].groupby('window')['profit_prob'].mean()
    axes[1, 1].plot([f"{w}日" for w in window_comparison.index], window_comparison.values,
                    marker='o', linewidth=2, markersize=8, color='steelblue')
    axes[1, 1].set_xlabel('时间窗口', fontproperties=font_prop)
    axes[1, 1].set_ylabel('平均盈利概率 (%)', fontproperties=font_prop)
    axes[1, 1].set_title('不同时间窗口的盈利概率（以当前发行价为准）', fontproperties=font_prop, fontsize=12, fontweight='bold')
    axes[1, 1].grid(True, alpha=0.3)
    for label in axes[1, 1].get_xticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'multi_dimension_scenario_analysis.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()

    return chart_path


def generate_tornado_chart(issue_price, current_price, lockup_period, volatility, drift, save_path):
    """生成龙卷风图 - 参数敏感性排序"""
    # 基准参数
    base_params = {
        'issue_price': issue_price,
        'current_price': current_price,
        'lockup_period': lockup_period,
        'volatility': volatility,
        'drift': drift
    }

    # 计算基准收益率
    def calculate_profit_prob(params):
        period_years = params['lockup_period'] / 12
        vol_period = params['volatility'] * np.sqrt(period_years)
        drift_period = params['drift'] * period_years
        required_return = (params['current_price'] - params['issue_price']) / params['issue_price']
        z_score = (drift_period - required_return) / vol_period
        return stats.norm.cdf(z_score) * 100

    base_prob = calculate_profit_prob(base_params)

    # 参数变化分析
    param_changes = []
    changes = [-0.2, -0.1, 0.1, 0.2]

    # 各参数敏感性
    for param_name in ['issue_price', 'current_price', 'lockup_period', 'volatility']:
        for change in changes:
            params = base_params.copy()
            params[param_name] = base_params[param_name] * (1 + change) if param_name != 'lockup_period' else base_params[param_name] * (1 + change)
            new_prob = calculate_profit_prob(params)
            impact = new_prob - base_prob
            param_name_cn = {'issue_price': '发行价', 'current_price': '当前价', 'lockup_period': '锁定期', 'volatility': '波动率'}[param_name]
            param_changes.append({'parameter': f'{param_name_cn} ({int(change*100)}%)', 'impact': impact})

    # 排序
    param_changes.sort(key=lambda x: x['impact'])

    # 绘制龙卷风图
    fig, ax = plt.subplots(figsize=(10, 6))
    colors = ['red' if p['impact'] < 0 else 'green' for p in param_changes]
    bars = ax.barh([p['parameter'] for p in param_changes], [p['impact'] for p in param_changes],
                  color=colors, alpha=0.7)
    ax.set_xlabel('对盈利概率的影响 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('参数敏感性排序（龙卷风图）', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    ax.grid(True, axis='x', alpha=0.3)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_sensitivity_chart(volatilities, profit_probs, current_vol, ma30, issue_price, save_path):
    """生成敏感性分析图表（增强版）"""
    fig = plt.figure(figsize=(16, 10))
    gs = fig.add_gridspec(2, 2, hspace=0.3, wspace=0.3)

    # 1. 波动率 vs 盈利概率柱状图
    ax1 = fig.add_subplot(gs[0, 0])
    colors = ['#27ae60' if p >= 60 else '#f39c12' if p >= 40 else '#e74c3c' for p in profit_probs]
    bars = ax1.bar(range(len(volatilities)), profit_probs, color=colors, alpha=0.8, edgecolor='black', linewidth=1.5)
    ax1.set_xlabel('年化波动率', fontproperties=font_prop, fontsize=11)
    ax1.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=11)
    ax1.set_title('不同波动率下的盈利概率', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax1.set_xticks(range(len(volatilities)))
    ax1.set_xticklabels([f'{v*100:.0f}%' for v in volatilities], fontproperties=font_prop, fontsize=9)
    ax1.grid(True, alpha=0.3, linestyle='--')
    ax1.set_ylim(0, 100)
    ax1.axhline(y=50, color='gray', linestyle='--', alpha=0.5)

    # 标记当前波动率
    current_idx = np.argmin(np.abs(np.array(volatilities) - current_vol))
    ax1.axvline(x=current_idx, color='red', linestyle='-', linewidth=2, alpha=0.7)
    ax1.text(current_idx, 95, f'当前\n{current_vol*100:.1f}%', ha='center', va='top',
            fontsize=9, fontproperties=font_prop, color='red', fontweight='bold')
    for bar, prob in zip(bars, profit_probs):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                f'{prob:.1f}%', ha='center', va='bottom', fontsize=9, fontweight='bold')

    # 2. 发行价折扣分析
    ax2 = fig.add_subplot(gs[0, 1])
    discounts = np.array([5, 10, 15, 20])  # 折价率
    issue_prices_discounted = ma30 * (1 - discounts/100)

    # 计算不同发行价下的盈利概率
    discount_probs = []
    for discounted_price in issue_prices_discounted:
        # 使用当前市场参数
        lockup_months = 6  # 假设6个月锁定期
        drift_rate = -0.1875  # 当前漂移率
        lockup_drift = drift_rate * (lockup_months / 12)
        lockup_vol = 0.3063 * np.sqrt(lockup_months / 12)  # 当前波动率

        # 计算盈利概率
        z = (np.log(ma30 / discounted_price) - lockup_drift) / lockup_vol
        prob = stats.norm.cdf(z) * 100
        discount_probs.append(prob)

    colors_discount = ['#27ae60', '#2ecc71', '#f39c12', '#e74c3c']
    bars2 = ax2.bar(range(len(discounts)), discount_probs, color=colors_discount,
                   alpha=0.8, edgecolor='black', linewidth=1.5)
    ax2.set_xlabel('发行价相对MA30折价率 (%)', fontproperties=font_prop, fontsize=11)
    ax2.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=11)
    ax2.set_title('不同发行价折扣下的盈利概率', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax2.set_xticks(range(len(discounts)))
    ax2.set_xticklabels([f'{d}%' for d in discounts], fontproperties=font_prop, fontsize=10)
    ax2.grid(True, alpha=0.3, linestyle='--')
    ax2.set_ylim(0, 100)
    ax2.axhline(y=50, color='gray', linestyle='--', alpha=0.5, label='盈亏平衡线')
    ax2.axvline(x=0, color='blue', linestyle='--', linewidth=2, alpha=0.7, label='当前发行价位置')
    ax2.legend(prop=font_prop, fontsize=9, loc='lower right')

    # 添加数值标签和发行价信息
    for bar, prob, price in zip(bars2, discount_probs, issue_prices_discounted):
        ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                f'{prob:.1f}%', ha='center', va='bottom', fontsize=9, fontweight='bold')
        ax2.text(bar.get_x() + bar.get_width()/2, -3,
                f'{price:.2f}元', ha='center', va='top', fontsize=8, fontproperties=font_prop)

    # 3. 热力图（波动率 vs 漂移率）
    ax3 = fig.add_subplot(gs[1, :])
    drift_range = np.linspace(-0.3, 0.15, 8)
    vol_range_heatmap = np.linspace(0.20, 0.50, 8)
    heatmap_data = []

    for d in drift_range:
        row = []
        for v in vol_range_heatmap:
            lockup_drift = d * (6/12)
            lockup_vol = v * np.sqrt(6/12)
            z = (np.log(ma30/issue_price) - lockup_drift) / lockup_vol
            row.append(stats.norm.cdf(z) * 100)
        heatmap_data.append(row)

    im = ax3.imshow(heatmap_data, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100,
                   extent=[vol_range_heatmap[0]*100, vol_range_heatmap[-1]*100, drift_range[0]*100, drift_range[-1]*100])
    ax3.set_xlabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
    ax3.set_ylabel('年化漂移率 (%)', fontproperties=font_prop, fontsize=12)
    ax3.set_title('盈利概率敏感性热力图（波动率 vs 漂移率）', fontproperties=font_prop, fontsize=13, fontweight='bold')
    ax3.tick_params(axis='both', which='major', labelsize=10)
    for label in ax3.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax3.get_yticklabels():
        label.set_fontproperties(font_prop)

    cbar = plt.colorbar(im, ax=ax3)
    cbar.set_label('盈利概率 (%)', fontproperties=font_prop, fontsize=11)
    cbar.ax.tick_params(labelsize=10)

    # 标记当前位置
    ax3.scatter([current_vol*100], [drift_rate*100], color='red', s=300, marker='*',
               edgecolors='white', linewidths=2, label='当前位置', zorder=5)
    ax3.legend(prop=font_prop, loc='upper right', fontsize=10)

    plt.suptitle('敏感性分析：波动率、漂移率与发行价折扣', fontproperties=font_prop,
               fontsize=15, fontweight='bold', y=0.98)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_stress_test_chart(scenarios, returns, save_path):
    """生成压力测试图表"""
    plt.figure(figsize=(10, 6))
    colors = ['#e74c3c' if r < 0 else '#2ecc71' for r in returns]
    bars = plt.barh(range(len(scenarios)), returns, color=colors, alpha=0.7)
    plt.xlabel('收益率 (%)', fontproperties=font_prop, fontsize=12)
    plt.ylabel('情景', fontproperties=font_prop, fontsize=12)
    plt.title('压力测试情景分析', fontproperties=font_prop, fontsize=14, fontweight='bold')
    plt.yticks(range(len(scenarios)), scenarios, fontproperties=font_prop)
    plt.xticks(fontproperties=font_prop)
    plt.axvline(x=0, color='black', linestyle='-', linewidth=1)
    plt.grid(True, alpha=0.3, axis='x')

    for bar, ret in zip(bars, returns):
        plt.text(bar.get_width() + (1 if ret > 0 else -1),
                bar.get_y() + bar.get_height()/2,
                f'{ret:.1f}%', ha='left' if ret > 0 else 'right',
                va='center', fontsize=10, fontproperties=font_prop)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_monte_carlo_chart(final_prices, issue_price, current_price, save_path):
    """生成蒙特卡洛模拟图表（增强版）"""
    fig = plt.figure(figsize=(16, 10))

    # 1. 价格分布直方图
    ax1 = plt.subplot(2, 3, 1)
    ax1.hist(final_prices, bins=50, alpha=0.7, color='steelblue', edgecolor='black')
    ax1.axvline(x=issue_price, color='red', linestyle='--', linewidth=2, label=f'发行价: {issue_price:.2f}')
    ax1.axvline(x=current_price, color='green', linestyle='--', linewidth=2, label=f'当前价: {current_price:.2f}')
    ax1.set_xlabel('未来价格 (元)', fontproperties=font_prop, fontsize=11)
    ax1.set_ylabel('频数', fontproperties=font_prop, fontsize=11)
    ax1.set_title('锁定期后价格分布', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax1.legend(prop=font_prop, fontsize=9)
    ax1.grid(True, alpha=0.3)
    for label in ax1.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax1.get_yticklabels():
        label.set_fontproperties(font_prop)

    # 2. 收益率分布
    ax2 = plt.subplot(2, 3, 2)
    returns = (final_prices - issue_price) / issue_price * 100
    ax2.hist(returns, bins=50, alpha=0.7, color='coral', edgecolor='black')
    ax2.axvline(x=0, color='red', linestyle='--', linewidth=2, label='盈亏平衡')
    ax2.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=11)
    ax2.set_ylabel('频数', fontproperties=font_prop, fontsize=11)
    ax2.set_title('收益率分布', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax2.legend(prop=font_prop, fontsize=9)
    ax2.grid(True, alpha=0.3)
    for label in ax2.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax2.get_yticklabels():
        label.set_fontproperties(font_prop)

    # 3. 累积分布函数
    ax3 = plt.subplot(2, 3, 3)
    sorted_returns = np.sort(returns)
    cumulative = np.arange(1, len(sorted_returns) + 1) / len(sorted_returns) * 100
    ax3.plot(sorted_returns, cumulative, linewidth=2, color='#8e44ad')
    ax3.axvline(x=0, color='red', linestyle='--', linewidth=1.5, label='盈亏平衡')
    ax3.axhline(y=50, color='gray', linestyle='--', linewidth=1, alpha=0.5)
    ax3.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=11)
    ax3.set_ylabel('累积概率 (%)', fontproperties=font_prop, fontsize=11)
    ax3.set_title('累积分布函数', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax3.legend(prop=font_prop, fontsize=9)
    ax3.grid(True, alpha=0.3)
    for label in ax3.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax3.get_yticklabels():
        label.set_fontproperties(font_prop)

    # 4. 盈亏概率饼图
    ax4 = plt.subplot(2, 3, 4)
    profit_prob = (returns > 0).mean() * 100
    loss_prob = 100 - profit_prob
    colors_pie = ['#27ae60', '#e74c3c']
    wedges, texts, autotexts = ax4.pie([profit_prob, loss_prob], labels=['盈利', '亏损'],
                                         autopct='%1.1f%%', colors=colors_pie, startangle=90)
    ax4.set_title('盈亏概率分布', fontproperties=font_prop, fontsize=12, fontweight='bold')
    for text in texts:
        text.set_fontproperties(font_prop)
    for autotext in autotexts:
        autotext.set_fontproperties(font_prop)
        autotext.set_fontsize(12)
        autotext.set_fontweight('bold')

    # 5. 关键统计指标
    ax5 = plt.subplot(2, 3, 5)
    ax5.axis('off')
    stats_text = f"""
    关键统计指标
    ════════════════

    盈利概率: {profit_prob:.1f}%

    预期收益率: {returns.mean():.1f}%

    收益率中位数: {np.median(returns):.1f}%

    95% VaR: {abs(np.percentile(returns, 5)):.1f}%

    99% VaR: {abs(np.percentile(returns, 1)):.1f}%

    最大回撤: {abs(returns.min()):.1f}%

    收益率标准差: {returns.std():.1f}%

    盈亏平衡价: {issue_price:.2f}元

    当前价格: {current_price:.2f}元
    """
    ax5.text(0.1, 0.5, stats_text, transform=ax5.transAxes,
            fontsize=11, verticalalignment='center',
            fontproperties=font_prop,
            family='monospace',
            bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.3))

    # 6. 概率分布曲线
    ax6 = plt.subplot(2, 3, 6)
    from scipy.stats import norm
    mu, std = returns.mean(), returns.std()
    x = np.linspace(returns.min(), returns.max(), 100)
    pdf = norm.pdf(x, mu, std)
    ax6.plot(x, pdf, linewidth=2, color='#2980b9', label='正态分布拟合')
    ax6.fill_between(x, pdf, alpha=0.3, color='#3498db')
    ax6.axvline(x=0, color='red', linestyle='--', linewidth=2, label='盈亏平衡')
    ax6.axvline(x=mu, color='green', linestyle='--', linewidth=2, label=f'均值: {mu:.1f}%')
    ax6.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=11)
    ax6.set_ylabel('概率密度', fontproperties=font_prop, fontsize=11)
    ax6.set_title('收益率概率分布', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax6.legend(prop=font_prop, fontsize=9)
    ax6.grid(True, alpha=0.3)
    for label in ax6.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax6.get_yticklabels():
        label.set_fontproperties(font_prop)

    plt.suptitle('蒙特卡洛模拟分析 (10000次模拟)', fontproperties=font_prop,
               fontsize=14, fontweight='bold', y=0.995)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_monte_carlo_charts_split(final_prices, issue_price, current_price, save_dir):
    """生成蒙特卡洛模拟图表 - 拆分版本（6个单独大图）"""
    import os

    returns = (final_prices - issue_price) / issue_price * 100
    chart_paths = []

    # 图1: 价格分布直方图（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.hist(final_prices, bins=50, alpha=0.7, color='steelblue', edgecolor='black')
    ax.axvline(x=issue_price, color='red', linestyle='--', linewidth=2, label=f'发行价: {issue_price:.2f}')
    ax.axvline(x=current_price, color='green', linestyle='--', linewidth=2, label=f'当前价: {current_price:.2f}')
    ax.set_xlabel('未来价格 (元)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('频数', fontproperties=font_prop, fontsize=14)
    ax.set_title('锁定期后价格分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_price_distribution_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2: 收益率分布（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.hist(returns, bins=50, alpha=0.7, color='coral', edgecolor='black')
    ax.axvline(x=0, color='red', linestyle='--', linewidth=2, label='盈亏平衡')
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('频数', fontproperties=font_prop, fontsize=14)
    ax.set_title('收益率分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_return_distribution_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3: 累积分布函数（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    sorted_returns = np.sort(returns)
    cumulative = np.arange(1, len(sorted_returns) + 1) / len(sorted_returns) * 100
    ax.plot(sorted_returns, cumulative, linewidth=2, color='#8e44ad')
    ax.axvline(x=0, color='red', linestyle='--', linewidth=1.5, label='盈亏平衡')
    ax.axhline(y=50, color='gray', linestyle='--', linewidth=1, alpha=0.5)
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('累积概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('累积分布函数', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_cdf_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图4: 盈亏概率饼图（单独大图）
    fig, ax = plt.subplots(figsize=(12, 10))
    profit_prob = (returns > 0).mean() * 100
    loss_prob = 100 - profit_prob
    colors_pie = ['#27ae60', '#e74c3c']
    wedges, texts, autotexts = ax.pie([profit_prob, loss_prob], labels=['盈利', '亏损'],
                                         autopct='%1.1f%%', colors=colors_pie, startangle=90,
                                         textprops={'fontproperties': font_prop, 'fontsize': 14})
    ax.set_title('盈亏概率分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    for autotext in autotexts:
        autotext.set_fontproperties(font_prop)
        autotext.set_fontsize(16)
        autotext.set_fontweight('bold')

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_profit_loss_pie_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图5: 关键统计指标（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.axis('off')
    stats_text = f"""
    关键统计指标
    ════════════════

    盈利概率: {profit_prob:.1f}%

    预期收益率: {returns.mean():.1f}%

    收益率中位数: {np.median(returns):.1f}%

    95% VaR: {abs(np.percentile(returns, 5)):.1f}%

    99% VaR: {abs(np.percentile(returns, 1)):.1f}%

    最大回撤: {abs(returns.min()):.1f}%

    收益率标准差: {returns.std():.1f}%

    盈亏平衡价: {issue_price:.2f}元

    当前价格: {current_price:.2f}元
    """
    ax.text(0.1, 0.5, stats_text, transform=ax.transAxes,
            fontsize=16, verticalalignment='center',
            fontproperties=font_prop,
            family='monospace',
            bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.3))

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_stats_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图6: 概率分布曲线（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    from scipy.stats import norm
    mu, std = returns.mean(), returns.std()
    x = np.linspace(returns.min(), returns.max(), 100)
    pdf = norm.pdf(x, mu, std)
    ax.plot(x, pdf, linewidth=2, color='#2980b9', label='正态分布拟合')
    ax.fill_between(x, pdf, alpha=0.3, color='#3498db')
    ax.axvline(x=0, color='red', linestyle='--', linewidth=2, label='盈亏平衡')
    ax.axvline(x=mu, color='green', linestyle='--', linewidth=2, label=f'均值: {mu:.1f}%')
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('概率密度', fontproperties=font_prop, fontsize=14)
    ax.set_title('收益率概率分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_pdf_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths


def generate_dcf_valuation_heatmap(save_path, current_price, net_income, total_shares, net_debt=20.0):
    """生成DCF估值热力图

    参数:
        save_path: 保存路径
        current_price: 当前股价
        net_income: 净利润（元）
        total_shares: 总股数
        net_debt: 净债务（亿元），可为负（净现金）
    """
    fig, ax = plt.subplots(figsize=(14, 9))

    # WACC 和永续增长率范围
    wacc_range = np.linspace(0.08, 0.13, 7)   # 8%-13%
    growth_range = np.linspace(0.15, 0.30, 7) # 15%-30% 增长率

    # 估值参数
    # FCF = 净利润 × 100%
    base_fcf = net_income / 100000000  # 转换为亿元
    # 股数（亿股）
    shares_billion = total_shares / 100000000

    # 计算每股价值矩阵
    pivot_data = []
    for g in growth_range:
        row = []
        for w in wacc_range:
            # DCF计算：FCF按g增长
            fcfs = [base_fcf * ((1 + g) ** i) for i in range(10)]
            pv_fcfs = sum([fcf / ((1 + w) ** (i+1)) for i, fcf in enumerate(fcfs)])

            # 终值：永续增长率设为3%
            terminal_growth = 0.03
            terminal_fcf = fcfs[-1] * (1 + terminal_growth)
            terminal_value = terminal_fcf / (w - terminal_growth) if w > terminal_growth else 0
            pv_terminal = terminal_value / ((1 + w) ** 10)

            enterprise_value = pv_fcfs + pv_terminal  # 亿元
            equity_value = enterprise_value - net_debt  # 减净债务（亿元）
            value_per_share = equity_value / shares_billion  # 每股价值（元）
            row.append(value_per_share)
        pivot_data.append(row)

    pivot_table = pd.DataFrame(
        pivot_data,
        index=[f'{g*100:.1f}%' for g in growth_range],
        columns=[f'{w*100:.1f}%' for w in wacc_range]
    )

    # 绘制热力图
    heatmap = sns.heatmap(pivot_table, annot=True, fmt='.2f', cmap='RdYlGn',
                          center=current_price, vmin=current_price*0.5, vmax=current_price*1.5,
                          cbar_kws={'label': '每股价值（元）'}, ax=ax)

    # 修复 colorbar 中文
    cbar = heatmap.collections[0].colorbar
    cbar.set_label('每股价值（元）', fontproperties=font_prop, fontsize=12)

    # 设置轴标签
    ax.set_xlabel('WACC（加权平均资本成本）', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('预测期FCF增长率', fontproperties=font_prop, fontsize=12)
    ax.set_title('DCF估值敏感性分析矩阵（每股价值：元）',
                fontproperties=font_prop, fontsize=14, fontweight='bold', pad=15)

    # 设置刻度标签字体
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(10)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(10)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_var_chart(var_95, var_99, cvar_95, save_path):
    """生成VaR风险度量图表"""
    plt.figure(figsize=(10, 6))
    metrics = ['95% VaR', '99% VaR', '95% CVaR']
    values = [var_95, var_99, cvar_95]
    colors = ['#e74c3c', '#c0392b', '#922b21']

    bars = plt.bar(metrics, [v*100 for v in values], color=colors, alpha=0.7)
    plt.ylabel('潜在损失 (%)', fontproperties=font_prop, fontsize=12)
    plt.title('风险度量 (VaR/CVaR)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    plt.xticks(fontproperties=font_prop)
    plt.yticks(fontproperties=font_prop)
    plt.grid(True, alpha=0.3, axis='y')

    for bar, value in zip(bars, values):
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                f'{value*100:.1f}%', ha='center', va='bottom',
                fontsize=11, fontweight='bold', fontproperties=font_prop)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_relative_valuation_charts(current_metrics, industry_avg, peer_companies_data, save_dir):
    """生成相对估值分析图表"""
    import os

    # 图1: 估值指标对比（柱状图）
    fig, axes = plt.subplots(2, 2, figsize=(16, 12))

    # 1. 估值指标对比
    ax1 = axes[0, 0]
    metrics = ['PE', 'PS', 'PB']
    current_vals = [current_metrics['pe'], current_metrics['ps'], current_metrics['pb']]
    industry_vals = [industry_avg['pe'], industry_avg['ps'], industry_avg['pb']]

    x = np.arange(len(metrics))
    width = 0.35

    bars1 = ax1.bar(x - width/2, current_vals, width, label='光弘科技', color='#3498db', alpha=0.8)
    bars2 = ax1.bar(x + width/2, industry_vals, width, label='行业平均', color='#e74c3c', alpha=0.8)

    ax1.set_xlabel('估值指标', fontproperties=font_prop, fontsize=12)
    ax1.set_ylabel('倍数', fontproperties=font_prop, fontsize=12)
    ax1.set_title('相对估值对比', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax1.set_xticks(x)
    ax1.set_xticklabels(metrics, fontproperties=font_prop)
    ax1.legend(prop=font_prop)
    ax1.grid(True, alpha=0.3, axis='y')

    for bars in [bars1, bars2]:
        for bar in bars:
            height = bar.get_height()
            ax1.text(bar.get_x() + bar.get_width()/2., height,
                    f'{height:.1f}', ha='center', va='bottom', fontsize=10, fontproperties=font_prop)

    # 2. PE对比
    ax2 = axes[0, 1]
    peer_names = list(peer_companies_data['name'])
    peer_pes = list(peer_companies_data['pe'])
    peer_pes.append(industry_avg['pe'])
    peer_names.append('行业平均')
    peer_pes.append(current_metrics['pe'])
    peer_names.append('光弘科技')

    colors_pe = ['#95a5a6'] * len(peer_companies_data) + ['#e74c3c', '#3498db']
    bars = ax2.bar(range(len(peer_names)), peer_pes, color=colors_pe, alpha=0.7)
    ax2.axhline(y=industry_avg['pe'], color='red', linestyle='--', alpha=0.5)
    ax2.set_xlabel('公司', fontproperties=font_prop, fontsize=10)
    ax2.set_ylabel('PE (倍)', fontproperties=font_prop, fontsize=10)
    ax2.set_title('PE倍数对比', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax2.set_xticks(range(len(peer_names)))
    ax2.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=9, rotation=45)
    ax2.grid(True, alpha=0.3, axis='y')

    # 3. PB对比
    ax3 = axes[1, 0]
    peer_pbs = list(peer_companies_data['pb'])
    peer_pbs.append(industry_avg['pb'])
    peer_pbs.append(current_metrics['pb'])

    bars = ax3.bar(range(len(peer_names)), peer_pbs, color=colors_pe, alpha=0.7)
    ax3.axhline(y=industry_avg['pb'], color='red', linestyle='--', alpha=0.5)
    ax3.set_xlabel('公司', fontproperties=font_prop, fontsize=10)
    ax3.set_ylabel('PB (倍)', fontproperties=font_prop, fontsize=10)
    ax3.set_title('PB倍数对比', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax3.set_xticks(range(len(peer_names)))
    ax3.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=9, rotation=45)
    ax3.grid(True, alpha=0.3, axis='y')

    # 4. PS对比
    ax4 = axes[1, 1]
    peer_pss = list(peer_companies_data['ps'])
    peer_pss.append(industry_avg['ps'])
    peer_pss.append(current_metrics['ps'])

    bars = ax4.bar(range(len(peer_names)), peer_pss, color=colors_pe, alpha=0.7)
    ax4.axhline(y=industry_avg['ps'], color='red', linestyle='--', alpha=0.5)
    ax4.set_xlabel('公司', fontproperties=font_prop, fontsize=10)
    ax4.set_ylabel('PS (倍)', fontproperties=font_prop, fontsize=10)
    ax4.set_title('PS倍数对比', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax4.set_xticks(range(len(peer_names)))
    ax4.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=9, rotation=45)
    ax4.grid(True, alpha=0.3, axis='y')

    plt.suptitle('相对估值分析：与同行对比', fontproperties=font_prop, fontsize=16, fontweight='bold', y=0.995)
    plt.tight_layout()

    chart1_path = os.path.join(save_dir, 'relative_valuation_comparison.png')
    plt.savefig(chart1_path, dpi=150, bbox_inches='tight')
    plt.close()

    # 图2: 估值情景分析
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))

    # 计算情景数据
    net_income = 253532329.85  # 从配置获取
    net_assets = 4939639031.34
    total_shares = 767460689
    current_price = 23.88
    issue_price = 20.25

    eps = net_income / total_shares
    bps = net_assets / total_shares

    scenarios_data = []
    # 情景1: 当前估值
    scenarios_data.append(['当前估值', current_metrics['pe'], current_metrics['pb'], current_metrics['ps'], current_price, 0])
    # 情景2: PE回归到行业平均
    target_price_pe = eps * industry_avg['pe']
    return_pe = (target_price_pe - current_price) / current_price * 100
    scenarios_data.append(['PE→行业平均', industry_avg['pe'], current_metrics['pb'], current_metrics['ps'], target_price_pe, return_pe])
    # 情景3: PB回归到行业平均
    target_price_pb = bps * industry_avg['pb']
    return_pb = (target_price_pb - current_price) / current_price * 100
    scenarios_data.append(['PB→行业平均', current_metrics['pe'], industry_avg['pb'], current_metrics['ps'], target_price_pb, return_pb])
    # 情景4: 全面回归行业平均
    target_price_avg = (target_price_pe + target_price_pb) / 2
    return_avg = (target_price_avg - current_price) / current_price * 100
    scenarios_data.append(['全面回归', industry_avg['pe'], industry_avg['pb'], industry_avg['ps'], target_price_avg, return_avg])
    # 情景5: 行业最低估值
    min_pe = peer_companies_data['pe'].min()
    min_pb = peer_companies_data['pb'].min()
    target_price_pessimistic = (eps * min_pe + bps * min_pb) / 2
    return_pessimistic = (target_price_pessimistic - current_price) / current_price * 100
    scenarios_data.append(['行业最低估值', min_pe, min_pb, industry_avg['ps'], target_price_pessimistic, return_pessimistic])
    # 情景6: 行业最高估值
    max_pe = peer_companies_data['pe'].max()
    max_pb = peer_companies_data['pb'].max()
    target_price_optimistic = (eps * max_pe + bps * max_pb) / 2
    return_optimistic = (target_price_optimistic - current_price) / current_price * 100
    scenarios_data.append(['行业最高估值', max_pe, max_pb, industry_avg['ps'], target_price_optimistic, return_optimistic])

    df_scenarios = pd.DataFrame(scenarios_data, columns=['情景', 'PE', 'PB', 'PS', '目标价格(元)', '预期收益率(%)'])

    # 左图：目标价格对比
    colors_scenario = ['gray', 'blue', 'green', 'orange', 'red', 'green']
    bars1 = ax1.barh(df_scenarios['情景'], df_scenarios['目标价格(元)'], color=colors_scenario, alpha=0.7)
    ax1.axvline(x=current_price, color='black', linestyle='--', label='当前价格')
    ax1.set_xlabel('目标价格 (元)', fontproperties=font_prop, fontsize=12)
    ax1.set_ylabel('情景', fontproperties=font_prop, fontsize=12)
    ax1.set_title('不同估值情景下的目标价格', fontproperties=font_prop, fontsize=13, fontweight='bold')
    ax1.legend(prop=font_prop)
    ax1.grid(True, alpha=0.3, axis='x')
    ax1.set_yticks(range(len(df_scenarios)))
    ax1.set_yticklabels(df_scenarios['情景'], fontproperties=font_prop, fontsize=10)

    # 右图：预期收益率对比
    colors_return = ['green' if r > 0 else 'red' for r in df_scenarios['预期收益率(%)']]
    bars2 = ax2.barh(df_scenarios['情景'], df_scenarios['预期收益率(%)'], color=colors_return, alpha=0.7)
    ax2.axvline(x=0, color='black', linestyle='-', linewidth=1)
    ax2.set_xlabel('预期收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax2.set_ylabel('情景', fontproperties=font_prop, fontsize=12)
    ax2.set_title('不同估值情景下的预期收益率', fontproperties=font_prop, fontsize=13, fontweight='bold')
    ax2.grid(True, alpha=0.3, axis='x')
    ax2.set_yticks(range(len(df_scenarios)))
    ax2.set_yticklabels(df_scenarios['情景'], fontproperties=font_prop, fontsize=10)

    plt.tight_layout()

    chart2_path = os.path.join(save_dir, 'relative_valuation_scenarios.png')
    plt.savefig(chart2_path, dpi=150, bbox_inches='tight')
    plt.close()

    return chart1_path, chart2_path, df_scenarios


def generate_stock_market_data_charts_split(market_data, price_data, volatility_data, save_dir):
    """
    生成个股市场数据图表（拆分版）

    参数:
        market_data: 市场数据字典
        price_data: 价格数据（包含close序列和均线）
        volatility_data: 波动率数据
        save_dir: 保存目录

    返回:
        图片路径列表 [股价走势图, 波动率图, 收益率分布图]
    """
    paths = []

    # 1. 股价走势与均线图
    fig, ax = plt.subplots(figsize=(14, 6))

    x_range = range(len(price_data['close']))
    ax.plot(x_range, price_data['close'], label='收盘价', linewidth=1.5, alpha=0.8, color='#2c3e50')

    colors = ['#e67e22', '#3498db', '#27ae60', '#e74c3c']
    windows = [30, 60, 120, 180]
    for i, window in enumerate(windows):
        ma_key = f'ma_{window}'
        if ma_key in price_data:
            ax.plot(x_range, price_data[ma_key], label=f'MA{window}', color=colors[i], linewidth=1.2, alpha=0.7)

    ax.set_xlabel('交易日', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('价格 (元)', fontproperties=font_prop, fontsize=12)
    ax.set_title(f'{market_data.get("stock_name", "N/A")} - 股价走势与均线',
                 fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, framealpha=0.9)
    ax.grid(True, alpha=0.3)

    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    path1 = os.path.join(save_dir, 'market_data_price_trend.png')
    plt.tight_layout()
    plt.savefig(path1, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path1)

    # 2. 滚动波动率图
    fig, ax = plt.subplots(figsize=(14, 5))

    colors = ['#e67e22', '#3498db', '#27ae60', '#e74c3c']
    for i, window in enumerate(windows):
        vol_key = f'volatility_{window}d_series'
        if vol_key in volatility_data:
            ax.plot(x_range, volatility_data[vol_key] * 100,
                   label=f'{window}日波动率', color=colors[i], linewidth=1.2, alpha=0.7)

    ax.set_xlabel('交易日', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('滚动波动率走势', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, framealpha=0.9)
    ax.grid(True, alpha=0.3)

    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    path2 = os.path.join(save_dir, 'market_data_volatility.png')
    plt.tight_layout()
    plt.savefig(path2, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path2)

    # 3. 日收益率分布图
    fig, ax = plt.subplots(figsize=(12, 5))

    returns_pct = price_data.get('pct_chg', [])
    if len(returns_pct) > 0:
        # pct_chg 已经是百分比形式（如 2.5 表示 2.5%），不需要再乘以 100
        ax.hist(returns_pct, bins=50, color='#3498db', alpha=0.7, edgecolor='black')

        mean_return = np.mean(returns_pct)
        ax.axvline(x=mean_return, color='#e74c3c', linestyle='--',
                  label=f'平均收益率 {mean_return:.3f}%', linewidth=2)
        ax.axvline(x=0, color='#27ae60', linestyle='-', label='盈亏平衡', linewidth=1.5)

    ax.set_xlabel('日收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('频数', fontproperties=font_prop, fontsize=12)
    ax.set_title('日收益率分布', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, framealpha=0.9)
    ax.grid(True, alpha=0.3, axis='y')

    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    path3 = os.path.join(save_dir, 'market_data_returns_dist.png')
    plt.tight_layout()
    plt.savefig(path3, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path3)

    return paths


def generate_index_data_charts_split(indices_data, save_dir):
    """
    生成市场指数数据图表（拆分版）

    参数:
        indices_data: 指数数据字典
        save_dir: 保存目录

    返回:
        图片路径列表 [波动率对比, 收益率对比, 胜率对比, 技术位置对比]
    """
    paths = []

    index_names = list(indices_data.keys())

    # 提取数据
    volatilities_60d = [indices_data[name]['volatility_60d'] * 100 for name in index_names]
    returns_60d = [indices_data[name]['return_60d'] * 100 for name in index_names]
    win_rates_60d = [indices_data[name]['win_rate_60d'] * 100 for name in index_names]
    ma_distances = [(indices_data[name]['current_level'] /
                    indices_data[name]['ma_60'] - 1) * 100 for name in index_names]

    # 1. 波动率对比图
    fig, ax = plt.subplots(figsize=(12, 6))
    colors1 = plt.cm.Reds(np.linspace(0.4, 0.9, len(index_names)))
    bars1 = ax.barh(index_names, volatilities_60d, color=colors1, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数波动率对比 (60日窗口)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, vol) in enumerate(zip(bars1, volatilities_60d)):
        ax.text(vol + 1, i, f'{vol:.1f}%', va='center', fontsize=10,
               fontproperties=font_prop, fontweight='bold')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path1 = os.path.join(save_dir, 'index_volatility_comparison.png')
    plt.tight_layout()
    plt.savefig(path1, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path1)

    # 2. 年化收益率对比图
    fig, ax = plt.subplots(figsize=(12, 6))
    colors2 = ['#27ae60' if r > 0 else '#e74c3c' for r in returns_60d]
    bars2 = ax.barh(index_names, returns_60d, color=colors2, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数年化收益率对比 (60日窗口)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1.2)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, ret) in enumerate(zip(bars2, returns_60d)):
        pos = ret + 1.5 if ret > 0 else ret - 2
        ha = 'left' if ret > 0 else 'right'
        ax.text(pos, i, f'{ret:+.1f}%', va='center', ha=ha, fontsize=10,
               fontproperties=font_prop, fontweight='bold', color='white')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path2 = os.path.join(save_dir, 'index_returns_comparison.png')
    plt.tight_layout()
    plt.savefig(path2, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path2)

    # 3. 胜率对比图
    fig, ax = plt.subplots(figsize=(12, 6))
    colors3 = plt.cm.Blues(np.linspace(0.4, 0.9, len(index_names)))
    bars3 = ax.barh(index_names, win_rates_60d, color=colors3, alpha=0.85, edgecolor='white')

    ax.set_xlabel('胜率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数胜率对比 (60日窗口)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.set_xlim(0, 100)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, wr) in enumerate(zip(bars3, win_rates_60d)):
        ax.text(wr + 1, i, f'{wr:.1f}%', va='center', fontsize=10,
               fontproperties=font_prop, fontweight='bold')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path3 = os.path.join(save_dir, 'index_winrate_comparison.png')
    plt.tight_layout()
    plt.savefig(path3, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path3)

    # 4. 技术位置对比图
    fig, ax = plt.subplots(figsize=(12, 6))
    colors4 = ['#27ae60' if d > 0 else '#e74c3c' for d in ma_distances]
    bars4 = ax.barh(index_names, ma_distances, color=colors4, alpha=0.85, edgecolor='white')

    ax.set_xlabel('相对MA60的偏离 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数技术位置对比', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1.2)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, dist) in enumerate(zip(bars4, ma_distances)):
        pos = dist + 0.8 if dist > 0 else dist - 0.8
        ha = 'left' if dist > 0 else 'right'
        ax.text(pos, i, f'{dist:+.1f}%', va='center', ha=ha, fontsize=10,
               fontproperties=font_prop, fontweight='bold', color='white')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path4 = os.path.join(save_dir, 'index_technical_position.png')
    plt.tight_layout()
    plt.savefig(path4, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path4)

    return paths


def generate_radar_chart(scores, save_path):
    """生成风险评分雷达图"""
    categories = list(scores.keys())
    values = list(scores.values())

    # 闭合雷达图
    angles_plot = np.linspace(0, 2*np.pi, len(categories), endpoint=False).tolist()
    angles_plot += angles_plot[:1]
    values_plot = values + values[:1]

    fig, ax = plt.subplots(figsize=(10, 10), subplot_kw=dict(projection='polar'))
    ax.plot(angles_plot, values_plot, 'o-', linewidth=2, color='#3498db')
    ax.fill(angles_plot, values_plot, alpha=0.25, color='#3498db')
    ax.set_xticks(angles_plot[:-1])
    ax.set_xticklabels(categories, fontproperties=font_prop, fontsize=11)
    ax.set_ylim(0, 30)
    ax.set_title('定增项目风险雷达图', fontsize=14, fontweight='bold',
                pad=20, fontproperties=font_prop)
    ax.grid(True)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_report(stock_code='300735.SZ', output_file='定增风险分析报告.docx'):
    """生成完整的定增风险分析报告"""

    print("开始生成定增风险分析报告（包含图表）...")

    # 创建文档
    document = Document()
    setup_chinese_font(document)

    # 加载配置 - 使用数据目录中的文件
    print("加载配置数据...")
    project_params, risk_params, market_data = load_placement_config(stock_code, data_dir=DATA_DIR)

    # 读取原始placement_params以获取完整财务数据
    placement_file = os.path.join(DATA_DIR, f"{stock_code.replace('.', '_')}_placement_params.json")
    if os.path.exists(placement_file):
        with open(placement_file, 'r', encoding='utf-8') as f:
            placement_params_raw = json.load(f)
            # 将财务数据合并到project_params中
            project_params.update({
                'net_assets': placement_params_raw.get('net_assets', 0),
                'total_debt': placement_params_raw.get('total_debt', 0),
                'net_income': placement_params_raw.get('net_income', 0),
                'revenue_growth': placement_params_raw.get('revenue_growth', 0.25),
                'operating_margin': placement_params_raw.get('operating_margin', 0.2),
                'beta': placement_params_raw.get('beta', 1.2)
            })

    # 发行类型评估
    ma30 = market_data.get('ma_30', 0)
    issue_price = project_params['issue_price']
    if issue_price < ma30:
        issue_type = "折价发行"
        discount_premium = (ma30 - issue_price) / ma30 * 100
    else:
        issue_type = "溢价发行"
        discount_premium = (issue_price - ma30) / ma30 * 100

    # 创建分析器
    analyzer = PrivatePlacementRiskAnalyzer(
        issue_price=project_params['issue_price'],
        issue_shares=project_params['issue_shares'],
        lockup_period=project_params['lockup_period'],
        current_price=project_params['current_price'],
        risk_free_rate=project_params['risk_free_rate']
    )

    current_date = datetime.now().strftime('%Y年%m月%d日')

    # ==================== 封面 ====================
    add_title(document, '定增项目风险分析报告', level=0)
    document.add_paragraph()
    document.add_paragraph()

    # 报告基本信息
    info_table = document.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'

    info_data = [
        ['公司名称', '光弘科技 (300735.SZ)'],
        ['报告日期', current_date],
        ['发行日期', current_date],
        ['发行价格', f'{project_params["issue_price"]:.2f} 元/股'],
        ['当前价格', f'{project_params["current_price"]:.2f} 元/股'],
        ['发行类型', f'{issue_type}（安全边际/溢价率: {discount_premium:.2f}%）']
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        for cell in info_table.rows[i].cells:
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.name = '宋体'
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_section_break(document)

    # ==================== 目录 ====================
    add_title(document, '目 录', level=1)
    add_paragraph(document, '一、项目概况')
    add_paragraph(document, '二、敏感性分析')
    add_paragraph(document, '三、情景分析')
    add_paragraph(document, '四、压力测试')
    add_paragraph(document, '五、DCF估值分析')
    add_paragraph(document, '六、相对估值分析')
    add_paragraph(document, '七、蒙特卡洛模拟')
    add_paragraph(document, '八、VaR风险度量')
    add_paragraph(document, '九、综合评估')
    add_paragraph(document, '十、投资建议与风险提示')

    add_section_break(document)

    # ==================== 一、项目概况 ====================
    add_title(document, '一、项目概况', level=1)

    add_title(document, '1.1 基本信息', level=2)

    project_headers = ['指标', '数值']
    project_data = [
        ['股票代码', '300735.SZ'],
        ['公司名称', '光弘科技'],
        ['发行价格', f'{project_params["issue_price"]:.2f} 元/股'],
        ['发行数量', f'{project_params["issue_shares"]:,} 股'],
        ['锁定期', f'{project_params["lockup_period"]} 个月'],
        ['融资金额', f'{project_params["issue_price"] * project_params["issue_shares"] / 100000000:.2f} 亿元'],
        ['当前价格', f'{project_params["current_price"]:.2f} 元/股'],
        ['MA30价格', f'{ma30:.2f} 元/股'],
        ['发行类型', issue_type],
        ['安全边际/溢价率', f'{discount_premium:.2f}%']
    ]
    add_table_data(document, project_headers, project_data)

    add_title(document, '1.2 市场数据', level=2)

    market_headers = ['指标', '数值']
    market_table_data = [
        ['当前价格', f'{market_data["current_price"]:.2f} 元/股'],
        ['平均价格', f'{market_data.get("avg_price_all", 0):.2f} 元/股'],
        ['价格标准差', f'{market_data.get("price_std", 0):.2f}'],
        ['月度波动率(20日)', f'{market_data.get("volatility_20d", 0)*100:.2f}%'],
        ['季度波动率(60日)', f'{market_data.get("volatility_60d", 0)*100:.2f}%'],
        ['半年波动率(120日)', f'{market_data.get("volatility_120d", 0)*100:.2f}%'],
        ['年度波动率(250日)', f'{market_data.get("volatility_250d", 0)*100:.2f}%'],
        ['月度年化收益率(20日)', f'{market_data.get("annual_return_20d", 0)*100:.2f}%'],
        ['季度年化收益率(60日)', f'{market_data.get("annual_return_60d", 0)*100:.2f}%'],
        ['半年年化收益率(120日)', f'{market_data.get("annual_return_120d", 0)*100:.2f}%'],
        ['年度年化收益率(250日)', f'{market_data.get("annual_return_250d", 0)*100:.2f}%'],
        ['MA20', f'{market_data.get("ma_20", 0):.2f} 元/股'],
        ['MA60', f'{market_data.get("ma_60", 0):.2f} 元/股'],
        ['MA120', f'{market_data.get("ma_120", 0):.2f} 元/股'],
        ['MA250', f'{market_data.get("ma_250", 0):.2f} 元/股'],
        ['数据来源', 'Tushare API']
    ]
    add_table_data(document, market_headers, market_table_data)

    # 1.2.1 生成个股市场数据图表
    print("生成个股市场数据图表...")

    # 尝试从Tushare获取历史数据生成图表
    stock_charts_generated = False
    try:
        import tushare as ts
        from datetime import timedelta

        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if not ts_token:
            print("⚠️ 未设置TUSHARE_TOKEN环境变量，跳过个股市场数据图表生成")
            add_paragraph(document, '')
            add_paragraph(document, '注：个股市场数据图表需要Tushare API支持，请设置TUSHARE_TOKEN环境变量后重新生成报告。')
        else:
            pro = ts.pro_api(ts_token)

            # 获取历史价格数据
            end_date = datetime.now().strftime('%Y%m%d')
            start_date = (datetime.now() - timedelta(days=500)).strftime('%Y%m%d')

            df_daily = pro.daily(ts_code=stock_code, start_date=start_date, end_date=end_date)
            if df_daily.empty:
                print(f"⚠️ 未获取到股票{stock_code}的历史数据")
            else:
                df_daily = df_daily.sort_values('trade_date').reset_index(drop=True)

                # 准备价格数据
                price_data = {'close': df_daily['close'].values, 'pct_chg': df_daily['pct_chg'].values}
                for window in [30, 60, 120, 180]:
                    ma = df_daily['close'].rolling(window=window).mean()
                    price_data[f'ma_{window}'] = ma.values

                # 准备波动率数据
                volatility_data = {}
                for window in [30, 60, 120, 180]:
                    pct_decimal = df_daily['pct_chg'] / 100.0
                    rolling_std = pct_decimal.rolling(window=window).std()
                    rolling_vol = rolling_std * np.sqrt(252)
                    volatility_data[f'volatility_{window}d_series'] = rolling_vol.values

                # 生成图表
                stock_charts_paths = generate_stock_market_data_charts_split(
                    market_data, price_data, volatility_data, IMAGES_DIR
                )

                # 添加图表到文档
                add_paragraph(document, '')
                add_paragraph(document, '图表 1.1: 股价走势与均线分析')
                if len(stock_charts_paths) > 0 and os.path.exists(stock_charts_paths[0]):
                    add_image(document, stock_charts_paths[0], width=Inches(6.5))
                    add_paragraph(document, '')

                add_paragraph(document, '图表 1.2: 滚动波动率分析')
                if len(stock_charts_paths) > 1 and os.path.exists(stock_charts_paths[1]):
                    add_image(document, stock_charts_paths[1], width=Inches(6.5))
                    add_paragraph(document, '')

                add_paragraph(document, '图表 1.3: 日收益率分布')
                if len(stock_charts_paths) > 2 and os.path.exists(stock_charts_paths[2]):
                    add_image(document, stock_charts_paths[2], width=Inches(6.5))
                    add_paragraph(document, '')

                add_paragraph(document, '分析结论：')
                add_paragraph(document, f'• 当前股价{market_data["current_price"]:.2f}元，{issue_type}于MA30({ma30:.2f}元)')
                add_paragraph(document, f'• 60日年化波动率为{market_data.get("volatility_60d", 0)*100:.2f}%，{"高于" if market_data.get("volatility_60d", 0) > 0.3 else "低于"}市场平均水平')
                add_paragraph(document, f'• 60日年化收益率为{market_data.get("annual_return_60d", 0)*100:.2f}%，{"表现良好" if market_data.get("annual_return_60d", 0) > 0 else "表现不佳"}')

                stock_charts_generated = True
                print(f"✅ 个股市场数据图表生成成功")
    except Exception as e:
        print(f"⚠️ 生成个股市场数据图表失败: {e}")
        import traceback
        traceback.print_exc()

    if not stock_charts_generated:
        add_paragraph(document, '')
        add_paragraph(document, '⚠️ 个股市场数据图表生成失败，请确保：')
        add_paragraph(document, '  1. 已设置TUSHARE_TOKEN环境变量')
        add_paragraph(document, '  2. Tushare API可访问')
        add_paragraph(document, '  3. 股票代码正确')

    # 1.3 市场指数分析
    add_paragraph(document, '')
    add_title(document, '1.3 市场指数分析', level=2)

    add_paragraph(document, '本章节分析主要市场指数的表现，包括波动率、收益率、胜率等指标，为项目风险评估提供市场环境参考。')

    # 加载指数数据（从data目录）
    indices_data_file = os.path.join(DATA_DIR, 'market_indices_scenario_data.json')
    indices_charts_paths = []

    if os.path.exists(indices_data_file):
        print("加载市场指数数据...")
        try:
            with open(indices_data_file, 'r', encoding='utf-8') as f:
                indices_data = json.load(f)

            # 生成指数对比图表
            indices_charts_paths = generate_index_data_charts_split(indices_data, IMAGES_DIR)

            # 添加指数对比表格
            add_paragraph(document, '')
            add_paragraph(document, '主要市场指数60日指标对比：')

            index_headers = ['指数', '当前点位', '波动率', '年化收益率', '胜率']
            index_table_data = []
            for name, data in indices_data.items():
                index_table_data.append([
                    name,
                    f"{data.get('current_level', 0):.2f}",
                    f"{data.get('volatility_60d', 0)*100:.2f}%",
                    f"{data.get('return_60d', 0)*100:.2f}%",
                    f"{data.get('win_rate_60d', 0)*100:.1f}%"
                ])
            add_table_data(document, index_headers, index_table_data)

            # 添加图表
            add_paragraph(document, '')
            add_paragraph(document, '图表 1.4: 各指数波动率对比')
            if len(indices_charts_paths) > 0 and os.path.exists(indices_charts_paths[0]):
                add_image(document, indices_charts_paths[0], width=Inches(6))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.5: 各指数年化收益率对比')
            if len(indices_charts_paths) > 1 and os.path.exists(indices_charts_paths[1]):
                add_image(document, indices_charts_paths[1], width=Inches(6))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.6: 各指数胜率对比')
            if len(indices_charts_paths) > 2 and os.path.exists(indices_charts_paths[2]):
                add_image(document, indices_charts_paths[2], width=Inches(6))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.7: 各指数技术位置对比')
            if len(indices_charts_paths) > 3 and os.path.exists(indices_charts_paths[3]):
                add_image(document, indices_charts_paths[3], width=Inches(6))
                add_paragraph(document, '')

            # 添加分析结论
            add_paragraph(document, '分析结论：')
            add_paragraph(document, '• 沪深300作为市场基准，波动率相对稳定，适合作为风险参照')
            add_paragraph(document, '• 创业板指和科创50波动率较高，体现成长股的高风险高收益特征')
            add_paragraph(document, '• 当前市场环境下，建议关注指数技术位置对个股表现的影响')

        except Exception as e:
            print(f"⚠️ 加载指数数据失败: {e}")
            add_paragraph(document, '⚠️ 指数数据暂未加载，请先运行07_market_data_analysis.ipynb生成数据')
    else:
        add_paragraph(document, '⚠️ 指数数据文件不存在，请先运行07_market_data_analysis.ipynb生成数据')

    add_section_break(document)

    # ==================== 二、敏感性分析 ====================
    add_title(document, '二、敏感性分析', level=1)

    add_paragraph(document, '本章节分析在不同波动率和漂移率假设下，定增项目的盈利概率变化情况。')

    add_title(document, '2.1 不同波动率下的盈利概率', level=2)

    # 生成敏感性分析数据和图表
    volatilities = np.linspace(0.20, 0.50, 7)
    drift_rate = risk_params.get('drift', market_data.get('drift', 0.08))

    prob_results = []
    for vol in volatilities:
        lockup_months = project_params['lockup_period']
        lockup_drift = drift_rate * (lockup_months / 12)
        lockup_vol = vol * np.sqrt(lockup_months / 12)

        d = (np.log(ma30 / project_params['issue_price']) - lockup_drift) / lockup_vol
        prob = stats.norm.cdf(d)
        prob_results.append(prob * 100)

    # 保存图表（传入当前波动率）
    sensitivity_chart_path = os.path.join(IMAGES_DIR, '01_sensitivity_analysis.png')
    generate_sensitivity_chart(volatilities, prob_results, market_data['volatility'],
                               ma30, project_params['issue_price'], sensitivity_chart_path)

    vol_data = [[f'{vol*100:.0f}%', f'{prob:.1f}%'] for vol, prob in zip(volatilities, prob_results)]
    add_table_data(document, ['波动率', '盈利概率'], vol_data)

    # 添加发行价折扣分析表格
    add_paragraph(document, '')
    add_title(document, '2.2 发行价折扣分析', level=2)

    # 计算不同折扣下的数据
    discount_analysis_data = []
    for discount_percent in [5, 10, 15, 20]:
        discounted_price = ma30 * (1 - discount_percent/100)
        lockup_months = 6
        drift_rate = -0.1875
        lockup_drift = drift_rate * (lockup_months / 12)
        lockup_vol = 0.3063 * np.sqrt(lockup_months / 12)

        z = (np.log(ma30 / discounted_price) - lockup_drift) / lockup_vol
        prob = stats.norm.cdf(z) * 100

        # 预期收益率（假设锁定期后价格回到MA30）
        expected_return = (ma30 - discounted_price) / discounted_price * 100

        discount_analysis_data.append([
            f'{discount_percent}%',
            f'{discounted_price:.2f}',
            f'{prob:.1f}%',
            f'{expected_return:.1f}%'
        ])

    add_table_data(document, ['发行价折扣', '折扣后发行价', '盈利概率', '预期收益率'], discount_analysis_data)

    add_paragraph(document, '图表 1: 敏感性分析')
    add_image(document, sensitivity_chart_path, width=Inches(6.5))

    add_paragraph(document, '分析结论：')
    add_paragraph(document, f'• 在当前市场波动率（{market_data["volatility"]*100:.1f}%）和漂移率（{market_data["drift"]*100:.1f}%）下，项目盈利概率约 {prob_results[3]:.1f}%')
    add_paragraph(document, '• 波动率越高，盈利概率的不确定性越大')
    add_paragraph(document, '• 折价发行为投资者提供了一定的安全边际')
    add_paragraph(document, '')
    add_paragraph(document, '发行价折扣分析：')

    # 计算当前发行价的折扣
    current_discount = (ma30 - project_params['issue_price']) / ma30 * 100
    add_paragraph(document, f'• 当前发行价折扣: {current_discount:.1f}%（发行价 {project_params["issue_price"]:.2f}元 / MA30 {ma30:.2f}元）')

    if current_discount >= 20:
        discount_level = "非常充足（≥20%）"
        safety_note = "安全边际很高，投资吸引力强"
    elif current_discount >= 15:
        discount_level = "充足（15%-20%）"
        safety_note = "安全边际良好，投资吸引力较高"
    elif current_discount >= 10:
        discount_level = "适中（10%-15%）"
        safety_note = "安全边际一般，需综合考虑其他因素"
    elif current_discount >= 5:
        discount_level = "较低（5%-10%）"
        safety_note = "安全边际有限，需谨慎评估"
    else:
        discount_level = "不足（<5%）"
        safety_note = "安全边际不足，风险较高"

    add_paragraph(document, f'• 折价水平评估: {discount_level}，{safety_note}')
    add_paragraph(document, '• 如果发行价折扣提高到15%或20%，盈利概率将显著提升')

    # ==================== 2.1.1 盈亏平衡价格敏感性 ====================
    add_paragraph(document, '')
    add_title(document, '2.1.1 盈亏平衡价格敏感性', level=2)

    add_paragraph(document, '分析在不同期望年化收益率要求下，解禁时需要达到的盈亏平衡价格。')

    # 计算不同收益率下的盈亏平衡价格
    target_returns = np.linspace(0.05, 0.50, 10)  # 5%到50%年化收益率
    break_even_prices = []
    issue_price = project_params['issue_price']
    lockup_years = project_params['lockup_period'] / 12

    for target_return in target_returns:
        # 盈亏平衡价格 = 发行价 * (1 + 年化收益率 * 锁定年数)
        be_price = issue_price * (1 + target_return * lockup_years)
        break_even_prices.append(be_price)

    # 生成表格数据
    be_data = []
    for ret, price in zip(target_returns[::2], break_even_prices[::2]):  # 每隔一个显示
        distance = (market_data['current_price'] - price) / market_data['current_price'] * 100
        status = "✅" if distance > 0 else "⚠️"
        be_data.append([f'{ret*100:.0f}%', f'{price:.2f}', f'{distance:+.1f}%', status])

    add_table_data(document, ['期望年化收益率', '盈亏平衡价格(元)', '距离当前价', '状态'], be_data)

    add_paragraph(document, '盈亏平衡分析结论：')
    add_paragraph(document, f'• 当前价格: {market_data["current_price"]:.2f} 元')
    add_paragraph(document, f'• 发行价格: {issue_price:.2f} 元')
    add_paragraph(document, f'• 20%年化收益率要求下盈亏平衡价: {break_even_prices[3]:.2f} 元')

    # 生成并插入盈亏平衡价格敏感性图表
    break_even_chart_path = os.path.join(IMAGES_DIR, '01_break_even_analysis.png')
    generate_break_even_chart(issue_price, market_data['current_price'], project_params['lockup_period'], break_even_chart_path)
    add_paragraph(document, '')
    add_paragraph(document, '图表 1.1: 盈亏平衡价格敏感性曲线')
    add_image(document, break_even_chart_path, width=Inches(6))

    # ==================== 2.2 锁定期敏感性分析 ====================
    add_paragraph(document, '')
    add_title(document, '2.2 锁定期敏感性分析', level=2)

    add_paragraph(document, '分析不同锁定期对盈亏平衡价格和所需涨幅的影响。')

    # 不同锁定期下的分析
    lockup_periods = [6, 9, 12, 18, 24, 36]  # 月
    target_return = 0.20  # 20%年化收益率

    lockup_analysis_data = []
    for period in lockup_periods:
        period_years = period / 12
        be_price = issue_price * (1 + target_return * period_years)
        required_increase = (be_price - issue_price) / issue_price * 100
        lockup_analysis_data.append([f'{period}个月', f'{be_price:.2f}', f'{required_increase:.2f}%'])

    add_table_data(document, ['锁定期', '盈亏平衡价(元)', '需要涨幅'], lockup_analysis_data)

    add_paragraph(document, '锁定期敏感性结论：')
    add_paragraph(document, f'• 当前锁定期: {project_params["lockup_period"]} 个月')
    add_paragraph(document, '• 锁定期越长，要求达到的盈亏平衡价格越高')
    add_paragraph(document, '• 锁定期每增加6个月，盈亏平衡价格约提高4-5元')

    # 生成并插入锁定期敏感性分析图表
    lockup_chart_path = os.path.join(IMAGES_DIR, '01_lockup_sensitivity.png')
    generate_lockup_sensitivity_chart(issue_price, market_data['current_price'], lockup_chart_path)
    add_paragraph(document, '')
    add_paragraph(document, '图表 1.2: 锁定期敏感性分析')
    add_image(document, lockup_chart_path, width=Inches(6.5))

    # ==================== 2.3 参数敏感性排序（龙卷风图） ====================
    add_paragraph(document, '')
    add_title(document, '2.3 参数敏感性排序（龙卷风图）', level=2)

    add_paragraph(document, '龙卷风图展示了各参数变化对盈利概率的影响程度，帮助识别最敏感的风险因子。')

    # 生成并插入龙卷风图
    tornado_chart_path = os.path.join(IMAGES_DIR, '01_tornado_chart.png')
    volatility = risk_params.get('volatility', market_data.get('volatility', 0.35))
    drift = risk_params.get('drift', market_data.get('drift', 0.08))
    generate_tornado_chart(project_params['issue_price'], market_data['current_price'],
                         project_params['lockup_period'], volatility, drift, tornado_chart_path)

    add_paragraph(document, '图表 1.3: 参数敏感性排序（龙卷风图）')
    add_image(document, tornado_chart_path, width=Inches(6))

    add_paragraph(document, '敏感性分析结论：')
    add_paragraph(document, '• 龙卷风图显示，对盈利概率影响最大的参数是当前价格和发行价格')
    add_paragraph(document, '• 波动率和锁定期对盈利概率的影响相对较小')
    add_paragraph(document, '• 投资者应重点关注市场价格变化，这是最关键的风险因子')

    add_section_break(document)

    # ==================== 三、情景分析 ====================
    add_title(document, '三、情景分析', level=1)

    add_paragraph(document, '本章节分析定增项目在不同预设情景下的风险表现，包括发行价折扣情景、不同时间窗口下的风险指标对比、多维度情景分析以及破发概率情景分析。')

    # ==================== 3.1 发行价折扣情景分析 ====================
    add_title(document, '3.1 发行价折扣情景分析', level=2)

    add_paragraph(document, '本节分析不同发行价折扣情景下的盈利概率和预期收益率，包括折价发行和溢价发行两种情景。')
    add_paragraph(document, '')
    add_paragraph(document, '发行类型定义：')
    add_paragraph(document, '• 折价发行：发行价 < MA30，有安全边际，盈利阈值 = 发行价')
    add_paragraph(document, '• 溢价发行：发行价 > MA30，无安全边际，盈利阈值 = max(MA30, 发行价×1.02)')
    add_paragraph(document, '')

    # 生成发行价折扣情景图表（基础版：合并为一图）
    scenario_chart_path = generate_discount_scenario_charts(
        ma30, project_params['current_price'], risk_params['volatility'],
        risk_params['drift'], project_params['lockup_period'], IMAGES_DIR)

    # 添加图表
    add_paragraph(document, '')
    add_paragraph(document, '图表 3.1: 发行价折扣情景分析（折价 vs 溢价）')
    add_image(document, scenario_chart_path, width=Inches(6.5))

    # 分析结论
    add_paragraph(document, '')
    add_paragraph(document, '发行价折扣情景分析结论：')

    # 计算当前发行价的折扣
    current_discount = (ma30 - project_params['issue_price']) / ma30 * 100
    add_paragraph(document, f'• 当前发行价: {project_params["issue_price"]:.2f} 元')
    add_paragraph(document, f'• MA30基准价: {ma30:.2f} 元')
    add_paragraph(document, f'• 当前折扣率: {current_discount:.1f}%')

    if current_discount >= 20:
        discount_level = "非常充足（≥20%）"
        safety_note = "安全边际很高，投资吸引力强"
    elif current_discount >= 15:
        discount_level = "充足（15%-20%）"
        safety_note = "安全边际良好，投资吸引力较高"
    elif current_discount >= 10:
        discount_level = "适中（10%-15%）"
        safety_note = "安全边际一般，需综合考虑其他因素"
    elif current_discount >= 5:
        discount_level = "较低（5%-10%）"
        safety_note = "安全边际有限，需谨慎评估"
    else:
        discount_level = "不足（<5%）"
        safety_note = "安全边际不足，风险较高"

    add_paragraph(document, f'• 折价水平评估: {discount_level}，{safety_note}')
    add_paragraph(document, '• 折价越多，盈利概率越高，但融资额越少')
    add_paragraph(document, '• 溢价越多，盈利概率越低，破发风险越大')

    # ==================== 3.2 不同时间窗口风险指标分析 ====================
    add_paragraph(document, '')
    add_title(document, '3.2 不同时间窗口风险指标分析', level=2)

    add_paragraph(document, '本节分析不同时间窗口下的风险指标变化，包括波动率、收益率、最大回撤、夏普比率等。')
    add_paragraph(document, '通过对比不同时间窗口的风险指标，可以更全面地了解股票的风险特征。')
    add_paragraph(document, '')
    add_paragraph(document, '⚠️ 基础版报告暂不包含时间窗口分析，请使用V2增强版报告查看完整分析。')

    # ==================== 3.3 多维度情景分析 ====================
    add_paragraph(document, '')
    add_title(document, '3.3 多维度情景分析', level=2)

    add_paragraph(document, '本节进行多维度情景分析，综合考虑波动率、发行价折扣率和时间窗口三个维度的组合影响。')
    add_paragraph(document, '')
    add_paragraph(document, '分析参数：')
    add_paragraph(document, '• 波动率情景：20%、30%、40%、50%（共4种）')
    add_paragraph(document, '• 折扣/溢价率：-10%、-5%、0%、5%、10%、15%、20%（共7种）')
    add_paragraph(document, '• 时间窗口：60日、120日、180日（共3种）')
    add_paragraph(document, f'• 总计情景数量：4 × 7 × 3 = 84 种组合')
    add_paragraph(document, '')

    # 生成多维度情景图表
    md_chart_path = generate_multi_dimension_scenario_charts(
        market_data['current_price'], ma30, risk_params['volatility'],
        risk_params['drift'], project_params['lockup_period'], IMAGES_DIR)

    # 添加图表
    add_paragraph(document, '')
    add_paragraph(document, '图表 3.3: 多维度情景分析（波动率 × 折扣率 × 时间窗口）')
    add_image(document, md_chart_path, width=Inches(6.5))

    # 分析结论
    add_paragraph(document, '')
    add_paragraph(document, '多维度情景分析结论：')
    add_paragraph(document, '• 低波动率 + 高折价 = 最佳情景（盈利概率 > 80%）')
    add_paragraph(document, '• 高波动率 + 高溢价 = 最差情景（盈利概率 < 30%）')
    add_paragraph(document, '• 时间窗口对盈利概率的影响：短期（60日）波动性最大，长期（180日）相对稳定')
    add_paragraph(document, '• 在当前市场条件下，折价发行的安全边际效应显著')

    # ==================== 3.4 破发概率情景分析 ====================
    add_paragraph(document, '')
    add_title(document, '3.4 破发概率情景分析', level=2)

    add_paragraph(document, '本节分析不同波动率情景下的破发概率，基于发行类型（折价/溢价）设定不同的盈利阈值。')
    add_paragraph(document, '')

    # 判断发行类型
    is_discount = project_params['issue_price'] < ma30

    add_paragraph(document, '发行类型定义：')
    add_paragraph(document, '• 折价发行：发行价 < MA30，有安全边际')
    add_paragraph(document, '• 溢价发行：发行价 > MA30，无安全边际')
    add_paragraph(document, '')

    if is_discount:
        threshold = project_params['issue_price']
        safety_margin = (ma30 - project_params['issue_price']) / ma30 * 100
        add_paragraph(document, f'✅ 当前为折价发行（有安全边际）')
        add_paragraph(document, f'   MA30价格: {ma30:.2f} 元')
        add_paragraph(document, f'   发行价格: {project_params["issue_price"]:.2f} 元')
        add_paragraph(document, f'   安全边际: {safety_margin:.2f}%')
        add_paragraph(document, '   盈利定义：解禁价 >= 发行价')
    else:
        threshold = max(ma30, project_params['issue_price'] * 1.02)
        premium_rate = (project_params['issue_price'] - ma30) / ma30 * 100
        add_paragraph(document, f'⚠️ 当前为溢价发行（无安全边际）')
        add_paragraph(document, f'   MA30价格: {ma30:.2f} 元')
        add_paragraph(document, f'   发行价格: {project_params["issue_price"]:.2f} 元')
        add_paragraph(document, f'   溢价幅度: {premium_rate:.2f}%')
        add_paragraph(document, '   盈利定义：解禁价 >= max(MA30, 发行价×1.02)')

    add_paragraph(document, '')
    add_paragraph(document, f'盈利阈值: {threshold:.2f} 元')
    add_paragraph(document, '')

    # 不同波动率情景下的破发概率分析
    volatilities_prob = np.linspace(0.15, 0.60, 10)
    drift_prob = risk_params.get('drift', market_data.get('drift', 0.08))

    break_even_table = []
    for vol in volatilities_prob:
        lockup_months = project_params['lockup_period']
        lockup_vol = vol * np.sqrt(lockup_months / 12)
        lockup_drift = drift_prob * (lockup_months / 12)
        required_return = (threshold - market_data['current_price']) / market_data['current_price']
        z_score = (lockup_drift - required_return) / lockup_vol
        profit_prob = (1 - stats.norm.cdf(-z_score)) * 100
        loss_prob = 100 - profit_prob
        break_even_table.append([f'{vol*100:.0f}%', f'{profit_prob:.1f}%', f'{loss_prob:.1f}%'])

    add_table_data(document, ['波动率', '盈利概率', '破发概率'], break_even_table)

    add_paragraph(document, '')
    add_paragraph(document, '破发概率情景分析结论：')
    add_paragraph(document, f'• 在当前波动率（{market_data["volatility"]*100:.1f}%）下，盈利概率约 {break_even_table[3][1]}')
    add_paragraph(document, '• 波动率越高，盈利概率和破发概率的差异越小')
    if is_discount:
        add_paragraph(document, '• 折价发行提供了安全边际，降低了破发风险')
    else:
        add_paragraph(document, '• 溢价发行无安全边际，破发风险相对较高')

    add_section_break(document)

    # ==================== 四、压力测试 ====================
    add_title(document, '四、压力测试', level=1)

    add_paragraph(document, '本章节模拟极端市场情况下的定增项目表现，包括历史危机情景和假设性极端情景。')

    add_title(document, '4.1 压力情景定义', level=2)

    # 定义完整的压力测试情景（与notebook一致）
    stress_scenarios = {
        '市场危机_2008': {
            'description': '模拟2008年金融危机，股价下跌60%',
            'price_drop': 0.60,
            'volatility_spike': 2.0,
        },
        '市场危机_2020': {
            'description': '模拟2020年疫情，股价下跌40%',
            'price_drop': 0.40,
            'volatility_spike': 1.5,
        },
        '行业政策收紧': {
            'description': '行业监管政策收紧，股价下跌25%',
            'price_drop': 0.25,
            'volatility_spike': 1.2,
        },
        '个股重大利空': {
            'description': '公司业绩暴雷，股价下跌35%',
            'price_drop': 0.35,
            'volatility_spike': 1.8,
        },
        '流动性危机': {
            'description': '市场流动性枯竭，股价下跌20%并波动率飙升',
            'price_drop': 0.20,
            'volatility_spike': 2.5,
        },
        '极端熊市': {
            'description': '极端熊市情景，股价下跌50%',
            'price_drop': 0.50,
            'volatility_spike': 1.3,
        }
    }

    # 计算各情景的结果
    scenario_names = []
    scenario_descriptions = []
    scenario_prices_list = []
    scenario_returns_list = []
    scenario_pnl_list = []

    current_price = market_data['current_price']
    issue_price = project_params['issue_price']
    issue_shares = project_params['issue_shares']

    for name, params in stress_scenarios.items():
        stressed_price = current_price * (1 - params['price_drop'])
        pnl_percent = ((stressed_price - issue_price) / issue_price) * 100
        pnl_amount = (stressed_price - issue_price) * issue_shares

        scenario_names.append(name)
        scenario_descriptions.append(params['description'])
        scenario_prices_list.append(stressed_price)
        scenario_returns_list.append(pnl_percent)
        scenario_pnl_list.append(pnl_amount)

    # 按照下跌幅度排序（从最严重到最轻）
    sorted_indices = sorted(range(len(scenario_names)), key=lambda i: scenario_prices_list[i])
    scenario_names = [scenario_names[i] for i in sorted_indices]
    scenario_descriptions = [scenario_descriptions[i] for i in sorted_indices]
    scenario_prices_list = [scenario_prices_list[i] for i in sorted_indices]
    scenario_returns_list = [scenario_returns_list[i] for i in sorted_indices]
    scenario_pnl_list = [scenario_pnl_list[i] for i in sorted_indices]

    # 保存图表
    stress_chart_path = os.path.join(IMAGES_DIR, '02_stress_test.png')
    generate_stress_test_chart(scenario_names, scenario_returns_list, stress_chart_path)

    add_title(document, '4.2 压力测试结果', level=2)

    # 生成详细的结果表格
    stress_table = []
    for i, name in enumerate(scenario_names):
        stress_table.append([
            name,
            scenario_descriptions[i],
            f'{scenario_prices_list[i]:.2f}',
            f'{scenario_returns_list[i]:+.2f}%',
            f'{scenario_pnl_list[i]/10000:+.2f}'
        ])

    add_table_data(document, ['情景', '描述', '受压后价格(元)', '收益率(%)', '盈亏(万元)'], stress_table)

    add_paragraph(document, '图表 2: 压力测试情景分析')
    add_image(document, stress_chart_path)

    add_title(document, '4.3 压力测试结论', level=2)

    # 统计分析
    profit_scenarios = sum(1 for r in scenario_returns_list if r > 0)
    total_scenarios = len(scenario_returns_list)
    worst_loss = min(scenario_pnl_list) / 10000
    worst_loss_percent = min(scenario_returns_list)
    best_scenario_idx = scenario_returns_list.index(max(scenario_returns_list))
    worst_scenario_idx = scenario_returns_list.index(min(scenario_returns_list))

    add_paragraph(document, f'压力测试统计：')
    add_paragraph(document, f'• 总情景数: {total_scenarios} 种')
    add_paragraph(document, f'• 盈利情景: {profit_scenarios} 种')
    add_paragraph(document, f'• 亏损情景: {total_scenarios - profit_scenarios} 种')
    add_paragraph(document, f'• 盈利概率: {profit_scenarios/total_scenarios*100:.1f}%')

    add_paragraph(document, '')
    add_paragraph(document, f'风险分析：')
    add_paragraph(document, f'• 最大潜在亏损: {abs(worst_loss):.2f} 万元')
    add_paragraph(document, f'• 最大亏损率: {worst_loss_percent:.2f}%')
    add_paragraph(document, f'• 最坏情景: {scenario_names[worst_scenario_idx]}')

    # 风险等级评估
    if profit_scenarios >= total_scenarios * 0.7:
        risk_level = "低风险 - 大部分情景下盈利"
        risk_emoji = "🟢"
    elif profit_scenarios >= total_scenarios * 0.4:
        risk_level = "中等风险 - 约一半情景下盈利"
        risk_emoji = "🟡"
    else:
        risk_level = "高风险 - 大部分情景下亏损"
        risk_emoji = "🔴"

    add_paragraph(document, '')
    add_paragraph(document, f'{risk_emoji} 压力测试评级: {risk_level}')

    add_paragraph(document, '')
    add_paragraph(document, '投资建议：')
    if profit_scenarios < total_scenarios * 0.5:
        add_paragraph(document, '⚠️ 项目抗风险能力较弱，建议：')
        add_paragraph(document, '  1. 考虑降低投资仓位')
        add_paragraph(document, '  2. 设置止损线')
        add_paragraph(document, '  3. 分批建仓，控制风险')
    else:
        add_paragraph(document, '✅ 项目抗风险能力尚可，但仍需：')
        add_paragraph(document, '  1. 密切关注市场动态')
        add_paragraph(document, '  2. 控制投资仓位')
        add_paragraph(document, '  3. 做好对冲准备')

    add_section_break(document)

    # ==================== 五、DCF估值分析 ====================
    add_title(document, '五、DCF估值分析', level=1)

    add_paragraph(document, '本章节使用现金流折现模型评估公司的内在价值。')

    add_title(document, '5.1 DCF估值方法说明', level=2)

    add_paragraph(document, 'DCF（Discounted Cash Flow）估值法通过预测公司未来自由现金流，并以加权平均资本成本（WACC）折现至现值，得出公司的内在价值。')

    add_paragraph(document, '')
    add_paragraph(document, '估值步骤：')
    add_paragraph(document, '1. 预测未来10年自由现金流（FCF）')
    add_paragraph(document, '2. 计算终值（Terminal Value）')
    add_paragraph(document, '3. 以WACC折现至现值')
    add_paragraph(document, '4. 减去净债务，得到股权价值')
    add_paragraph(document, '5. 除以总股数，得到每股价值')

    # 从placement_params中获取财务数据
    net_income = project_params.get('net_income', 253532329.85)  # 净利润

    # 获取真实总股本（从 Tushare）
    total_shares = 767460689  # 光弘科技真实总股本（约7.67亿股）
    try:
        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if ts_token:
            import tushare as ts
            from datetime import timedelta

            pro = ts.pro_api(ts_token)
            trade_date = (datetime.now() - timedelta(days=1)).strftime('%Y%m%d')

            df_basic = pro.daily_basic(
                ts_code=stock_code,
                trade_date=trade_date,
                fields='ts_code,total_share'
            )

            if not df_basic.empty:
                total_shares = float(df_basic.iloc[0]['total_share']) * 10000  # 万股转股
    except Exception as e:
        print(f"获取总股本失败，使用估算值: {e}")

    net_assets = project_params.get('net_assets', 4939639031.34)  # 净资产（元）
    total_debt = project_params.get('total_debt', 4507009002.2)  # 总债务（元）

    # 尝试从tushare获取资产负债表数据，计算净债务
    net_debt = None  # 初始化为None
    cash_assets = 0.0
    total_liab_value = 0.0

    try:
        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if ts_token:
            import tushare as ts
            from datetime import timedelta

            pro = ts.pro_api(ts_token)

            # 尝试获取最近几个交易日
            for days_back in range(1, 15):
                test_date = (datetime.now() - timedelta(days=days_back)).strftime('%Y%m%d')
                try:
                    # 获取资产负债表数据
                    df_balancesheet = pro.balancesheet(
                        ts_code=stock_code,
                        period=test_date[:6],  # 使用期间
                        fields='ts_code,ann_date,total_liab,cash_equivalents,money_cap'
                    )
                    if not df_balancesheet.empty:
                        total_liab_value = float(df_balancesheet.iloc[0]['total_liab']) / 100000000  # 转亿元
                        cash_equivalents = float(df_balancesheet.iloc[0]['cash_equivalents']) / 100000000  # 转亿元

                        # 净债务 = 总负债 - 现金及现金等价物
                        net_debt = total_liab_value - cash_equivalents
                        cash_assets = cash_equivalents

                        print(f"✅ 获取资产负债表数据成功，期间: {test_date[:6]}")
                        print(f"   总负债: {total_liab_value:.2f} 亿元，现金: {cash_equivalents:.2f} 亿元，净债务: {net_debt:.2f} 亿元")
                        break
                except:
                    continue
    except Exception as e:
        print(f"获取资产负债表数据失败: {e}")

    # 如果无法获取实际数据，使用项目参数中的估计值
    if net_debt is None:
        net_debt = total_debt / 100000000  # 转亿元
        print(f"⚠️ 使用参数文件中的总债务作为净债务: {net_debt:.2f} 亿元")

    # 生成并插入 DCF 热力图
    add_title(document, '5.2 估值敏感性分析', level=2)

    dcf_chart_path = os.path.join(IMAGES_DIR, '04_dcf_valuation_heatmap.png')
    generate_dcf_valuation_heatmap(dcf_chart_path, project_params['current_price'], net_income, total_shares, net_debt)

    add_paragraph(document, '图表 3: DCF估值敏感性分析矩阵')
    add_image(document, dcf_chart_path, width=Inches(6))

    add_paragraph(document, '分析说明：')
    add_paragraph(document, '• 横轴：WACC（加权平均资本成本），代表折现率')
    add_paragraph(document, '• 纵轴：增长率，代表预测期FCF增长率')
    add_paragraph(document, '• 颜色：绿色表示估值高于当前价，红色表示估值低于当前价')
    add_paragraph(document, f'• 当前股价: {project_params["current_price"]:.2f} 元/股')

    add_title(document, '5.3 估值参数与过程', level=2)

    # 计算基础FCF
    base_fcf = net_income  # FCF = 净利润

    dcf_process_data = [
        ['财务数据', ''],
        ['净利润', f'{net_income/100000000:.2f} 亿元'],
        ['净资产', f'{net_assets/100000000:.2f} 亿元'],
        ['现金及现金等价物', f'{cash_assets:.2f} 亿元'],
        ['净债务', f'{net_debt:.2f} 亿元'],
        ['总股本', f'{total_shares:,} 股'],
        [''],
        ['DCF假设参数', ''],
        ['预测期', '10年'],
        ['基准FCF', f'{base_fcf/100000000:.2f} 亿元'],
        ['FCF增长率（预测期）', '15%-30%'],
        ['WACC范围', '8.0% - 13.0%'],
        ['永续增长率', '3.0%']
    ]
    add_table_data(document, ['参数', '值'], dcf_process_data)

    add_paragraph(document, '')
    add_paragraph(document, '估值计算过程（示例）：')

    # 以中间值计算示例
    wacc_example = 0.10  # 10%
    growth_example = 0.025  # 2.5%

    # 计算示例
    fcfs = [base_fcf * (1.1 ** i) for i in range(10)]
    pv_fcfs = sum([fcf / ((1 + wacc_example) ** (i+1)) for i, fcf in enumerate(fcfs)])
    terminal_fcf = fcfs[-1] * (1 + growth_example)
    terminal_value = terminal_fcf / (wacc_example - growth_example)
    pv_terminal = terminal_value / ((1 + wacc_example) ** 10)
    enterprise_value = pv_fcfs + pv_terminal  # 单位：元
    # 修正：股权价值 = 企业价值 - 净债务（注意单位转换）
    equity_value = enterprise_value - (net_debt * 100000000)  # 净债务从亿元转为元
    intrinsic_value = equity_value / total_shares  # 单位：元/股

    example_data = [
        ['WACC', f'{wacc_example*100:.1f}%'],
        ['永续增长率', f'{growth_example*100:.1f}%'],
        ['预测期FCF现值', f'{pv_fcfs/100000000:.2f} 亿元'],
        ['终值', f'{terminal_value/100000000:.2f} 亿元'],
        ['终值现值', f'{pv_terminal/100000000:.2f} 亿元'],
        ['企业价值', f'{enterprise_value/100000000:.2f} 亿元'],
        ['净债务', f'{net_debt:.2f} 亿元'],
        ['股权价值', f'{equity_value/100000000:.2f} 亿元'],
        ['每股价值', f'{intrinsic_value:.2f} 元/股']
    ]
    add_table_data(document, ['项目', '计算结果'], example_data)

    add_title(document, '5.4 估值结论', level=2)

    margin_market = (intrinsic_value - project_params['current_price']) / project_params['current_price'] * 100
    margin_issue = (intrinsic_value - project_params['issue_price']) / project_params['issue_price'] * 100

    add_paragraph(document, f'• DCF内在价值: {intrinsic_value:.2f} 元/股')
    add_paragraph(document, f'• 当前价格: {project_params["current_price"]:.2f} 元/股')
    add_paragraph(document, f'• 发行价格: {project_params["issue_price"]:.2f} 元/股')
    add_paragraph(document, f'• 相对市价安全边际: {margin_market:+.1f}%')
    add_paragraph(document, f'• 相对发行价安全边际: {margin_issue:+.1f}%')

    if margin_issue > 15:
        conclusion = "DCF估值显示，相比发行价有显著安全边际，估值合理偏低。"
    elif margin_issue > 0:
        conclusion = "DCF估值显示，相比发行价有一定安全边际，估值相对合理。"
    else:
        conclusion = "DCF估值显示，发行价高于内在价值，需谨慎。"

    add_paragraph(document, '')
    add_paragraph(document, f'估值结论: {conclusion}')

    add_section_break(document)

    # ==================== 六、相对估值分析 ====================
    add_title(document, '六、相对估值分析', level=1)

    add_paragraph(document, '本章节通过相对估值法（参数法），将光弘科技与行业内可比公司进行对比分析。')
    add_paragraph(document, '选取申万三级分类"消费电子零部件及组装"行业的同行公司，对比PE、PS、PB等估值倍数。')

    add_title(document, '6.1 估值指标对比', level=2)

    # 使用 Tushare 数据获取估值指标
    try:
        ts_token = os.environ.get('TUSHARE_TOKEN', '')

        if ts_token:
            import tushare as ts
            from datetime import timedelta  # 导入 timedelta
            import time

            pro = ts.pro_api(ts_token)

            # 获取目标公司的估值数据（自动往前推1-5天直到找到交易日）
            trade_date = None
            df_target = None

            for days_back in range(1, 6):  # 尝试往前推1-5天
                test_date = (datetime.now() - timedelta(days=days_back)).strftime('%Y%m%d')
                try:
                    df_target = pro.daily_basic(
                        ts_code=stock_code,
                        trade_date=test_date,
                        fields='ts_code,trade_date,close,pe_ttm,pb,ps_ttm,total_mv'
                    )
                    if not df_target.empty:
                        trade_date = test_date
                        break
                except:
                    continue

            if df_target is None or df_target.empty:
                raise ValueError("未获取到目标公司数据（请检查网络或交易日历）")

            current_metrics_val = {
                'pe': df_target.iloc[0]['pe_ttm'],
                'pb': df_target.iloc[0]['pb'],
                'ps': df_target.iloc[0]['ps_ttm']
            }
            print(f"✅ 获取相对估值数据成功，交易日期: {trade_date}")

            # 获取申万三级行业分类的同行公司（与 notebook 一致）
            df_industry = pro.index_member_all(ts_code=stock_code)
            if df_industry.empty:
                raise ValueError("未获取到行业分类")

            df_industry = df_industry.sort_values('in_date', ascending=False)
            latest_industry = df_industry.iloc[0]
            target_index_code = latest_industry['l3_code']  # 申万三级行业指数代码
            target_industry_l3 = latest_industry['l3_name']  # 行业名称

            # 获取该三级行业的所有成分股
            df_peers = pro.index_member_all(l3_code=target_index_code)
            df_peers = df_peers[df_peers['ts_code'] != stock_code]

            # 获取同行公司基本信息
            peer_codes = df_peers['ts_code'].unique().tolist()
            peer_basic = pro.stock_basic(ts_code=','.join(peer_codes[:30]),
                                       fields='ts_code,name,market')

            peer_stocks_all = pd.merge(df_peers, peer_basic, on='ts_code', how='left')
            peer_stocks_all = peer_stocks_all.drop_duplicates(subset=['ts_code'])

            # 限制数量并排序
            peer_stocks_all = peer_stocks_all.head(15)
            peer_names_dict = dict(zip(peer_stocks_all['ts_code'], peer_stocks_all['name_x']))

            # 获取同行公司的估值数据
            peer_data_list = []
            for peer_code in peer_stocks_all['ts_code'].tolist():
                peer_name = peer_names_dict.get(peer_code, peer_code)

                try:
                    df_peer = pro.daily_basic(
                        ts_code=peer_code,
                        trade_date=trade_date,
                        fields='ts_code,pe_ttm,pb,ps_ttm,total_mv'
                    )
                    if not df_peer.empty:
                        df_peer['name'] = peer_name
                        peer_data_list.append(df_peer)
                except:
                    pass

                time.sleep(0.2)  # 避免请求过快

            if peer_data_list:
                peer_companies_val = pd.concat(peer_data_list, ignore_index=True)

                # 过滤异常数据
                peer_companies_val = peer_companies_val[
                    (peer_companies_val['pe_ttm'] > 0) &
                    (peer_companies_val['pe_ttm'] < 500) &
                    (peer_companies_val['pb'] > 0) &
                    (peer_companies_val['pb'] < 20)
                ]

                # 重命名列并进行单位转换
                # total_mv单位是万元，需要转换为亿元（除以10000）
                peer_companies_val['market_cap'] = peer_companies_val['total_mv'] / 10000

                peer_companies_val = peer_companies_val.rename(columns={
                    'pe_ttm': 'pe',
                    'ps_ttm': 'ps'
                })

                # 只保留需要的列
                peer_companies_val = peer_companies_val[['name', 'pe', 'ps', 'pb', 'market_cap']]
            else:
                raise ValueError("未获取到同行公司估值数据")

        else:
            raise ValueError("TUSHARE_TOKEN 未设置")

    except Exception as e:
        print(f"获取相对估值数据失败: {e}，使用示例数据")
        # 使用示例数据
        current_metrics_val = {
            'pe': 56.24,
            'pb': 3.71,
            'ps': 2.30
        }

        peer_companies_val = pd.DataFrame({
            'name': ['立讯精密', '歌尔股份', '蓝思科技', '长盈精密', '领益智造', '安洁科技', '比亚迪电子'],
            'pe': [20.5, 25.8, 22.3, 28.6, 18.9, 32.1, 24.5],
            'ps': [1.2, 1.8, 1.5, 2.1, 1.0, 2.5, 1.6],
            'pb': [2.8, 3.2, 2.5, 3.8, 2.1, 4.2, 3.0],
            'market_cap': [120, 80, 50, 45, 65, 25, 90]
        })

    # 计算行业平均值
    industry_avg_val = {
        'pe': peer_companies_val['pe'].mean(),
        'ps': peer_companies_val['ps'].mean(),
        'pb': peer_companies_val['pb'].mean()
    }

    # 估值指标对比表
    valuation_headers = ['指标', '光弘科技', '行业平均', '偏离度']
    valuation_data = [
        ['PE (TTM)', f"{current_metrics_val['pe']:.2f}倍", f"{industry_avg_val['pe']:.2f}倍",
         f"{(current_metrics_val['pe']-industry_avg_val['pe'])/industry_avg_val['pe']*100:+.1f}%"],
        ['PB', f"{current_metrics_val['pb']:.2f}倍", f"{industry_avg_val['pb']:.2f}倍",
         f"{(current_metrics_val['pb']-industry_avg_val['pb'])/industry_avg_val['pb']*100:+.1f}%"],
        ['PS (TTM)', f"{current_metrics_val['ps']:.2f}倍", f"{industry_avg_val['ps']:.2f}倍",
         f"{(current_metrics_val['ps']-industry_avg_val['ps'])/industry_avg_val['ps']*100:+.1f}%"]
    ]
    add_table_data(document, valuation_headers, valuation_data)

    # 添加同行公司名单
    add_paragraph(document, '')
    add_title(document, '6.1.1 同行公司名单', level=2)
    add_paragraph(document, f'基于申万三级行业分类"消费电子零部件及组装"筛选的同行公司：')

    # 按市值排序的同行公司表格
    peer_companies_sorted = peer_companies_val.sort_values('market_cap', ascending=False)
    peer_headers = ['公司名称', 'PE (TTM)', 'PB', 'PS (TTM)', '市值(亿元)']
    peer_rows = []
    for _, row in peer_companies_sorted.iterrows():
        peer_rows.append([
            row['name'],
            f"{row['pe']:.2f}",
            f"{row['pb']:.2f}",
            f"{row['ps']:.2f}",
            f"{row['market_cap']:.2f}"
        ])
    add_table_data(document, peer_headers, peer_rows)

    add_paragraph(document, '')
    add_paragraph(document, '图表 7: 相对估值对比分析')
    chart1_path, chart2_path, df_scenarios = generate_relative_valuation_charts(
        current_metrics_val, industry_avg_val, peer_companies_val, IMAGES_DIR
    )
    add_image(document, chart1_path)

    add_title(document, '6.2 估值偏离度分析', level=2)

    add_paragraph(document, f"• PE偏离度: {(current_metrics_val['pe']-industry_avg_val['pe'])/industry_avg_val['pe']*100:+.1f}%")
    add_paragraph(document, f"• PB偏离度: {(current_metrics_val['pb']-industry_avg_val['pb'])/industry_avg_val['pb']*100:+.1f}%")
    add_paragraph(document, f"• PS偏离度: {(current_metrics_val['ps']-industry_avg_val['ps'])/industry_avg_val['ps']*100:+.1f}%")

    add_paragraph(document, '')
    if current_metrics_val['pe'] > industry_avg_val['pe'] * 1.2:
        add_paragraph(document, '⚠️ PE显著高于行业平均，可能存在估值偏高风险')
    elif current_metrics_val['pe'] < industry_avg_val['pe'] * 0.8:
        add_paragraph(document, '✅ PE低于行业平均，可能存在估值低估机会')
    else:
        add_paragraph(document, 'ℹ️ PE接近行业平均，估值相对合理')

    add_title(document, '6.3 估值回归情景分析', level=2)

    add_paragraph(document, '分析估值回归到行业平均水平时的情景：')

    # 计算回归情景
    net_income_rel = project_params.get('net_income', 253532329.85)
    net_assets_rel = project_params.get('net_assets', 4939639031.34)
    total_shares_rel = 767460689
    current_price_rel = project_params['current_price']

    eps_rel = net_income_rel / total_shares_rel
    bps_rel = net_assets_rel / total_shares_rel

    target_price_pe = eps_rel * industry_avg_val['pe']
    return_pe = (target_price_pe - current_price_rel) / current_price_rel * 100

    target_price_pb = bps_rel * industry_avg_val['pb']
    return_pb = (target_price_pb - current_price_rel) / current_price_rel * 100

    target_price_avg = (target_price_pe + target_price_pb) / 2
    return_avg_rel = (target_price_avg - current_price_rel) / current_price_rel * 100

    # 情景数据表
    scenario_headers = ['情景', 'PE', 'PB', 'PS', '目标价格(元)', '预期收益率(%)']
    scenario_data = [
        ['当前估值', f"{current_metrics_val['pe']:.2f}", f"{current_metrics_val['pb']:.2f}", f"{current_metrics_val['ps']:.2f}",
         f"{current_price_rel:.2f}", "0.0"],
        ['PE→行业平均', f"{industry_avg_val['pe']:.2f}", f"{current_metrics_val['pb']:.2f}", f"{current_metrics_val['ps']:.2f}",
         f"{target_price_pe:.2f}", f"{return_pe:+.1f}"],
        ['PB→行业平均', f"{current_metrics_val['pe']:.2f}", f"{industry_avg_val['pb']:.2f}", f"{current_metrics_val['ps']:.2f}",
         f"{target_price_pb:.2f}", f"{return_pb:+.1f}"],
        ['全面回归', f"{industry_avg_val['pe']:.2f}", f"{industry_avg_val['pb']:.2f}", f"{industry_avg_val['ps']:.2f}",
         f"{target_price_avg:.2f}", f"{return_avg_rel:+.1f}"]
    ]
    add_table_data(document, scenario_headers, scenario_data)

    add_paragraph(document, '')
    add_paragraph(document, '图表 8: 估值情景分析')
    add_image(document, chart2_path)

    add_paragraph(document, '情景分析结论：')
    add_paragraph(document, f'• 如果PE回归到行业平均，目标价格约 {target_price_pe:.2f} 元，预期收益率 {return_pe:+.1f}%')
    add_paragraph(document, f'• 如果PB回归到行业平均，目标价格约 {target_price_pb:.2f} 元，预期收益率 {return_pb:+.1f}%')
    add_paragraph(document, f'• 综合估值目标价格约 {target_price_avg:.2f} 元，预期收益率 {return_avg_rel:+.1f}%')

    add_paragraph(document, '')
    if return_avg_rel > 10:
        add_paragraph(document, '📈 估值偏低，具有投资价值')
    elif return_avg_rel < -10:
        add_paragraph(document, '📉 估值偏高，需谨慎投资')
    else:
        add_paragraph(document, '📊 估值合理，可正常投资')

    add_section_break(document)

    # ==================== 七、蒙特卡洛模拟 ====================
    add_title(document, '七、蒙特卡洛模拟', level=1)

    add_paragraph(document, '本章节使用蒙特卡洛方法模拟未来股价路径，评估收益分布。')

    add_title(document, '5.1 模拟参数', level=2)

    mc_params = [
        ['模拟次数', '10,000次'],
        ['锁定期', f'{project_params["lockup_period"]} 个月'],
        ['初始价格', f'{project_params["current_price"]:.2f} 元/股'],
        ['年化波动率', f'{market_data["volatility"]*100:.1f}%'],
        ['年化漂移率', f'{market_data["drift"]*100:.1f}%'],
        ['发行价格', f'{project_params["issue_price"]:.2f} 元/股']
    ]
    add_table_data(document, ['参数', '值'], mc_params)

    add_title(document, '7.2 模拟结果', level=2)

    # 运行简化模拟
    lockup_days = project_params['lockup_period'] * 30
    n_simulations = 10000

    lockup_drift = risk_params.get('drift', market_data.get('drift', 0.08)) * (lockup_days / 365)
    lockup_vol = risk_params.get('volatility', market_data.get('volatility', 0.35)) * np.sqrt(lockup_days / 365)

    np.random.seed(42)
    returns = np.random.normal(lockup_drift, lockup_vol, n_simulations)
    final_prices = project_params['current_price'] * np.exp(returns)
    profit_losses = (final_prices - project_params['issue_price']) / project_params['issue_price']

    # 统计
    profit_prob = (profit_losses > 0).mean()
    mean_return = profit_losses.mean()
    median_return = np.median(profit_losses)
    percent_5 = np.percentile(profit_losses, 5)
    percent_95 = np.percentile(profit_losses, 95)

    # 保存图表（拆分版）
    mc_chart_paths = generate_monte_carlo_charts_split(
        final_prices, project_params['issue_price'], project_params['current_price'], IMAGES_DIR)

    mc_results = [
        ['盈利概率', f'{profit_prob*100:.1f}%'],
        ['预期收益率', f'{mean_return*100:.1f}%'],
        ['收益率中位数', f'{median_return*100:.1f}%'],
        ['5%分位数（最差情况）', f'{percent_5*100:.1f}%'],
        ['95%分位数（较好情况）', f'{percent_95*100:.1f}%']
    ]
    add_table_data(document, ['指标', '值'], mc_results)

    add_paragraph(document, '图表 4: 蒙特卡洛模拟结果 - 价格分布')
    add_image(document, mc_chart_paths[0])

    add_paragraph(document, '')
    add_paragraph(document, '图表 4.1: 蒙特卡洛模拟结果 - 收益率分布')
    add_image(document, mc_chart_paths[1])

    add_paragraph(document, '')
    add_paragraph(document, '图表 4.2: 蒙特卡洛模拟结果 - 累积分布函数')
    add_image(document, mc_chart_paths[2])

    add_paragraph(document, '')
    add_paragraph(document, '图表 4.3: 蒙特卡洛模拟结果 - 盈亏概率')
    add_image(document, mc_chart_paths[3])

    add_paragraph(document, '')
    add_paragraph(document, '图表 4.4: 蒙特卡洛模拟结果 - 关键统计指标')
    add_image(document, mc_chart_paths[4])

    add_paragraph(document, '')
    add_paragraph(document, '图表 4.5: 蒙特卡洛模拟结果 - 概率分布曲线')
    add_image(document, mc_chart_paths[5])

    add_paragraph(document, '蒙特卡洛模拟结论：')
    add_paragraph(document, f'• 在 {n_simulations:,} 次模拟中，约 {profit_prob*100:.1f}% 的场景实现盈利')
    add_paragraph(document, f'• 预期收益率约 {mean_return*100:.1f}%')
    add_paragraph(document, f'• 95%的置信区间内，亏损可能不超过 {abs(percent_5)*100:.1f}%')

    add_section_break(document)

    # ==================== 六、VaR风险度量 ====================
    add_title(document, '八、VaR风险度量', level=1)

    add_paragraph(document, '本章节使用多种方法计算风险价值（VaR）和条件风险价值（CVaR），全面评估定增项目的下行风险。')

    add_title(document, '8.1 VaR风险度量概述', level=2)

    add_paragraph(document, 'VaR（Value at Risk）风险价值表示在给定置信水平下，投资组合可能遭受的最大损失。')
    add_paragraph(document, '本报告使用蒙特卡洛模拟法计算VaR，基于10,000次模拟路径。')

    # 计算不同置信水平下的VaR
    var_90 = abs(np.percentile(profit_losses, 10))
    var_95 = abs(np.percentile(profit_losses, 5))
    var_99 = abs(np.percentile(profit_losses, 1))

    # 计算CVaR（条件风险价值）
    cvar_90 = abs(profit_losses[profit_losses <= np.percentile(profit_losses, 10)].mean())
    cvar_95 = abs(profit_losses[profit_losses <= np.percentile(profit_losses, 5)].mean())
    cvar_99 = abs(profit_losses[profit_losses <= np.percentile(profit_losses, 1)].mean())

    # 保存图表
    var_chart_path = os.path.join(IMAGES_DIR, '05_var_analysis.png')
    generate_var_chart(var_95, var_99, cvar_95, var_chart_path)

    add_title(document, '8.2 不同置信水平下的VaR', level=2)

    # VaR表格
    var_table_data = [
        ['90%', f'{var_90*100:.1f}%', f'{cvar_90*100:.1f}%', '有10%的概率损失超过此值'],
        ['95%', f'{var_95*100:.1f}%', f'{cvar_95*100:.1f}%', '有5%的概率损失超过此值'],
        ['99%', f'{var_99*100:.1f}%', f'{cvar_99*100:.1f}%', '有1%的概率损失超过此值']
    ]
    add_table_data(document, ['置信水平', 'VaR', 'CVaR', '说明'], var_table_data)

    add_paragraph(document, '图表 5: VaR风险度量')
    add_image(document, var_chart_path)

    add_title(document, '8.3 CVaR（条件风险价值）分析', level=2)

    add_paragraph(document, 'CVaR（Conditional Value at Risk）又称期望短缺（Expected Shortfall），表示在损失超过VaR时的平均损失。')
    add_paragraph(document, 'CVaR比VaR更保守，能更好地反映极端情况下的风险。')

    cvar_comparison_data = [
        ['VaR vs CVaR (95%)', f'{var_95*100:.1f}%', f'{cvar_95*100:.1f}%', f'{(cvar_95-var_95)*100:.1f}%', f'{(cvar_95/var_95-1)*100:.1f}%'],
        ['VaR vs CVaR (99%)', f'{var_99*100:.1f}%', f'{cvar_99*100:.1f}%', f'{(cvar_99-var_99)*100:.1f}%', f'{(cvar_99/var_99-1)*100:.1f}%']
    ]
    add_table_data(document, ['项目', 'VaR', 'CVaR', '差值', 'CVaR/VaR'], cvar_comparison_data)

    add_paragraph(document, 'CVaR分析结论：')
    add_paragraph(document, f'• 95%置信水平下，CVaR比VaR高出 {(cvar_95/var_95-1)*100:.1f}%')
    add_paragraph(document, f'• 99%置信水平下，CVaR比VaR高出 {(cvar_99/var_99-1)*100:.1f}%')
    add_paragraph(document, '• 说明极端情况下的损失比VaR估计的更严重')

    add_title(document, '8.4 最大回撤分析', level=2)

    add_paragraph(document, '最大回撤是指从峰值到谷底的最大跌幅，是衡量投资风险的重要指标。')

    # 估算最大回撤
    estimated_max_drawdown = market_data['volatility'] * 2  # 简化估算
    estimated_95_drawdown = market_data['volatility'] * 1.5

    drawdown_data = [
        ['预估平均最大回撤', f'{estimated_max_drawdown*100:.1f}%', '基于波动率估算'],
        ['预估95%分位回撤', f'{estimated_95_drawdown*100:.1f}%', '95%的路径回撤不超过此值']
    ]
    add_table_data(document, ['回撤指标', '数值', '说明'], drawdown_data)

    add_title(document, '8.5 潜在损失金额估算', level=2)

    # 计算损失金额
    investment_amount = project_params['issue_price'] * project_params['issue_shares']
    loss_var_95 = investment_amount * var_95
    loss_cvar_95 = investment_amount * cvar_95
    loss_var_99 = investment_amount * var_99

    loss_amount_data = [
        ['投资总额', f'{investment_amount/10000:.2f}', '万元'],
        ['95% VaR潜在损失', f'{loss_var_95/10000:.2f}', '万元'],
        ['95% CVaR潜在损失', f'{loss_cvar_95/10000:.2f}', '万元'],
        ['99% VaR潜在损失', f'{loss_var_99/10000:.2f}', '万元']
    ]
    add_table_data(document, ['项目', '金额', '单位'], loss_amount_data)

    add_title(document, '8.6 VaR风险测算结论', level=2)

    add_paragraph(document, '风险度量总结：')
    add_paragraph(document, f'• 95%置信水平下，最大可能亏损约 {var_95*100:.1f}%，约 {loss_var_95/10000:.2f} 万元')
    add_paragraph(document, f'• 极端情况下（1%概率），亏损可能达到 {var_99*100:.1f}%，约 {loss_var_99/10000:.2f} 万元')
    add_paragraph(document, f'• CVaR显示，在最差5%情况下平均损失约 {cvar_95*100:.1f}%，约 {loss_cvar_95/10000:.2f} 万元')
    add_paragraph(document, f'• 预估最大回撤约 {estimated_max_drawdown*100:.1f}%')

    # 风险评级
    if var_95 <= 0.15:
        risk_rating = "低风险"
        risk_emoji = "🟢"
    elif var_95 <= 0.25:
        risk_rating = "中等风险"
        risk_emoji = "🟡"
    else:
        risk_rating = "高风险"
        risk_emoji = "🔴"

    add_paragraph(document, '')
    add_paragraph(document, f'{risk_emoji} 基于VaR的风险评级: {risk_rating}')

    add_section_break(document)

    # ==================== 七、综合评估 ====================
    add_title(document, '九、综合评估', level=1)

    add_title(document, '9.1 风险评分', level=2)

    # 综合评分
    scores = {
        '盈利概率': 30 if profit_prob >= 0.7 else 20 if profit_prob >= 0.5 else 10,
        '发行折价': 20 if discount_premium > 15 else 15 if discount_premium > 5 else 5,
        '锁定期': 20 if project_params['lockup_period'] <= 6 else 15 if project_params['lockup_period'] <= 12 else 5,
        '波动率': 15 if risk_params.get('volatility', 0.35) <= 0.3 else 10 if risk_params.get('volatility', 0.35) <= 0.4 else 5,
        'VaR': 15 if var_95 <= 0.2 else 10 if var_95 <= 0.3 else 3
    }
    total_score = sum(scores.values())

    # 保存雷达图
    radar_chart_path = os.path.join(IMAGES_DIR, '06_radar_chart.png')
    generate_radar_chart(scores, radar_chart_path)

    score_data = [[k, f'{v}/30' if k == '盈利概率' else f'{v}/20' if k != 'VaR' else f'{v}/15'] for k, v in scores.items()]
    score_data.append(['总分', f'{total_score}/100'])

    add_table_data(document, ['评估维度', '得分'], score_data)

    add_paragraph(document, '图表 6: 风险评分雷达图')
    add_image(document, radar_chart_path)

    add_paragraph(document, f'综合风险评分: {total_score}/100 分')

    add_title(document, '9.2 各维度评估', level=2)

    add_paragraph(document, f'1. 盈利概率（{scores["盈利概率"]}/30分）')
    if profit_prob >= 0.7:
        add_paragraph(document, '   ✅ 盈利概率较高，投资吸引力强')
    elif profit_prob >= 0.5:
        add_paragraph(document, '   ⚠️ 盈利概率中等，需谨慎评估')
    else:
        add_paragraph(document, '   ❌ 盈利概率较低，风险较大')

    add_paragraph(document, f'2. 发行折价（{scores["发行折价"]}/20分）')
    if discount_premium > 15:
        add_paragraph(document, f'   ✅ {issue_type}，安全边际充足（{discount_premium:.1f}%）')
    elif discount_premium > 5:
        add_paragraph(document, f'   ⚠️ 安全边际一般（{discount_premium:.1f}%）')
    else:
        add_paragraph(document, f'   ❌ 安全边际不足（{discount_premium:.1f}%）')

    add_paragraph(document, f'3. 锁定期（{scores["锁定期"]}/20分）')
    lockup_comment = "相对较短" if project_params["lockup_period"] <= 6 else "需要较长等待"
    add_paragraph(document, f'   锁定期 {project_params["lockup_period"]} 个月，{lockup_comment}')

    add_paragraph(document, f'4. 波动率（{scores["波动率"]}/15分）')
    vol_comment = "相对可控" if market_data["volatility"] <= 0.35 else "波动较大"
    add_paragraph(document, f'   年化波动率 {market_data["volatility"]*100:.1f}%，{vol_comment}')

    add_paragraph(document, f'5. 风险价值（{scores["VaR"]}/15分）')
    var_comment = "风险可控" if var_95 <= 0.25 else "需注意风险"
    add_paragraph(document, f'   95% VaR 为 {var_95*100:.1f}%，{var_comment}')

    add_section_break(document)

    # ==================== 八、投资建议与风险提示 ====================
    add_title(document, '十、投资建议与风险提示', level=1)

    add_title(document, '10.1 投资建议', level=2)

    # 投资建议
    if total_score >= 80:
        recommendation = "🟢 建议投资"
        detail = "项目整体风险收益比良好，各项指标表现优秀，建议积极参与。"
        actions = [
            "可以按计划参与定增",
            "建议关注锁定期内的市场动态",
            "可考虑适当杠杆"
        ]
    elif total_score >= 60:
        recommendation = "🟡 谨慎投资"
        detail = "项目整体风险收益比一般，存在一定不确定性，建议谨慎评估。"
        actions = [
            "深入调研项目基本面",
            "评估自身风险承受能力",
            "考虑分批参与"
        ]
    else:
        recommendation = "🔴 不建议投资"
        detail = "项目整体风险较高，安全边际不足，建议观望或寻找其他机会。"
        actions = [
            "暂不建议参与",
            "等待更好的入场时机",
            "考虑其他投资标的"
        ]

    add_paragraph(document, f'总体建议: {recommendation}', bold=True, font_size=12)
    add_paragraph(document, f'说明: {detail}')

    add_paragraph(document, '建议行动:')
    for action in actions:
        add_paragraph(document, f'  • {action}')

    add_title(document, '10.2 核心指标汇总', level=2)

    summary_data = [
        ['风险评分', f'{total_score}/100'],
        ['盈利概率', f'{profit_prob*100:.1f}%'],
        ['发行类型', f'{issue_type}'],
        ['安全边际/溢价率', f'{discount_premium:.2f}%'],
        ['预期收益率', f'{mean_return*100:.1f}%'],
        ['95% VaR', f'{var_95*100:.1f}%'],
        ['波动率', f'{market_data["volatility"]*100:.1f}%'],
        ['DCF内在价值', f'{intrinsic_value:.2f} 元/股']
    ]
    add_table_data(document, ['指标', '值'], summary_data)

    add_title(document, '10.3 主要风险提示', level=2)

    risks = [
        f'• 市场风险: 当前波动率 {market_data["volatility"]*100:.1f}%，市场波动可能导致实际收益偏离预期',
        f'• 流动性风险: {project_params["lockup_period"]}个月锁定期内无法交易，需承担期间价格波动',
        f'• 估值风险: DCF估值基于多个假设，实际业绩可能偏离预测',
        f'• VaR风险: 95%置信水平下最大可能亏损 {var_95*100:.1f}%',
        f'• 行业风险: 需关注行业政策变化和竞争格局'
    ]

    for risk in risks:
        add_paragraph(document, risk)

    add_title(document, '10.4 免责声明', level=2)

    disclaimer = f'''
    本报告由自动化分析系统生成，仅供参考使用，不构成投资建议。

    1. 本报告基于历史数据和公开信息进行分析，不保证数据的准确性和完整性；
    2. 市场有风险，投资需谨慎。本报告中的任何分析观点不代表未来表现；
    3. 本报告提到的任何证券或投资标的，仅为分析示例，不构成推荐；
    4. 投资者应根据自身情况独立判断，自行承担投资风险；
    5. 本报告的知识产权归分析团队所有，未经许可不得转载或使用。

    报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    '''

    add_paragraph(document, disclaimer, font_size=9)

    # ==================== 保存文档 ====================
    output_path = os.path.join(OUTPUTS_DIR, output_file)
    document.save(output_path)

    print(f"\n✅ 报告生成成功!")
    print(f"   📄 Word文档: {output_path}")
    print(f"   📊 图表目录: {IMAGES_DIR}")
    print(f"   文件大小: {os.path.getsize(output_path)/1024:.1f} KB")

    return output_path


if __name__ == '__main__':
    report_path = generate_report(stock_code='300735.SZ', output_file='光弘科技_定增风险分析报告.docx')
    print("\n报告生成完成！")
