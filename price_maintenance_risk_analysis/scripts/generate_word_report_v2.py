#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
定增风险分析报告生成器 V2（格式优化版）

版本说明：
- V1 (generate_word_report.py): 稳定版本，包含完整分析内容
- V2 (generate_word_report_v2.py): 格式优化版本，用于调整字体、图片展示等

主要功能：
- 生成 Word 格式的综合分析报告，包含图表
- 支持敏感性分析、压力测试、DCF估值、相对估值、蒙特卡洛模拟等
- 使用 Tushare 获取真实市场数据

使用方法：
    python generate_word_report_v2.py
"""

import sys
import os
import json
import argparse
import subprocess
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats

# 添加路径 - 从 scripts 目录运行
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, PROJECT_DIR)

from utils.config_loader import load_placement_config
from utils.analysis_tools import PrivatePlacementRiskAnalyzer
from utils.font_manager import get_font_prop
from utils.time_series_forecaster import TimeSeriesForecaster

# 获取中文字体
font_prop = get_font_prop()

# 配置路径
DATA_DIR = os.path.join(PROJECT_DIR, 'data')
REPORTS_DIR = os.path.join(PROJECT_DIR, 'reports')
IMAGES_DIR = os.path.join(REPORTS_DIR, 'images')
OUTPUTS_DIR = os.path.join(REPORTS_DIR, 'outputs')

# 确保目录存在
os.makedirs(IMAGES_DIR, exist_ok=True)
os.makedirs(OUTPUTS_DIR, exist_ok=True)


def setup_chinese_font(document):
    """设置中文字体 - V2版本"""
    # 正文：仿宋-GB2312，四号（14pt）
    document.styles['Normal'].font.name = '仿宋_GB2312'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    document.styles['Normal'].font.size = Pt(14)


def add_title(document, text, level=1):
    """添加标题 - V2版本"""
    if level == 0:  # 报告标题，使用段落而不是heading
        para = document.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = '方正公文小标宋_GBK'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
            run.font.size = Pt(22)  # 二号（22pt）
            run.font.bold = True
        return para
    else:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            # 根据级别设置字体和大小
            if level == 1:  # 一级标题
                run.font.name = '方正公文小标宋_GBK'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
                run.font.size = Pt(16)  # 三号（16pt）
            elif level == 2:  # 二级标题
                run.font.name = '方正公文小标宋_GBK'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
                run.font.size = Pt(15)  # 小三（15pt）
            else:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(14)
        return heading


def add_paragraph(document, text, bold=False, font_size=14):
    """添加段落 - V2版本（正文：仿宋-GB2312，四号）"""
    para = document.add_paragraph(text)
    for run in para.runs:
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(font_size)
        if bold:
            run.font.bold = True
    return para


def add_image(document, image_path, width=Inches(5)):
    """添加图片到文档"""
    if os.path.exists(image_path):
        document.add_picture(image_path, width=width)
        # 设置图片居中
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    else:
        print(f"⚠️ 图片不存在: {image_path}")
        return False


def add_table_data(document, headers, data, font_size=12):
    """添加表格 - V2版本（表格：宋体，默认小四号12pt）

    Args:
        document: Word文档对象
        headers: 表头列表
        data: 数据列表
        font_size: 字体大小（单位：磅），默认12pt（小四号），10.5pt为五号
    """
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 设置表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        para = header_cells[i].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.size = Pt(font_size)
        para.runs[0].font.name = '宋体'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            para = row_cells[i].paragraphs[0]
            para.runs[0].font.size = Pt(font_size)
            para.runs[0].font.name = '宋体'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if isinstance(cell_data, (int, float)):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def add_section_break(document):
    """添加分页符"""
    document.add_page_break()


def generate_break_even_chart(issue_price, current_price, lockup_period, save_path):
    """生成盈亏平衡价格敏感性曲线"""
    # 不同期望收益率下的盈亏平衡价格
    target_returns = np.linspace(0.05, 0.50, 10)  # 5%到50%年化收益率
    break_even_prices = [issue_price * (1 + r * (lockup_period / 12)) for r in target_returns]

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(target_returns * 100, break_even_prices, 'b-', linewidth=2, label='盈亏平衡价格')
    ax.axhline(y=issue_price, color='r', linestyle='--', label='发行价格')
    ax.axhline(y=current_price, color='g', linestyle='--', label='当前价格')

    ax.set_xlabel('期望年化收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('盈亏平衡价格 (元)', fontproperties=font_prop, fontsize=12)
    ax.set_title('不同收益率要求下的盈亏平衡价格', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    ax.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_lockup_sensitivity_chart(issue_price, current_price, save_path):
    """生成锁定期敏感性分析图表"""
    lockup_periods = [6, 9, 12, 18, 24, 36]  # 月
    target_return = 0.20  # 20%年化收益率

    lockup_analysis = []
    for period in lockup_periods:
        period_years = period / 12
        be_price = issue_price * (1 + target_return * period_years)
        required_increase = (be_price - issue_price) / issue_price * 100
        lockup_analysis.append({'period': period, 'be_price': be_price, 'increase': required_increase})

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 5))

    # 左图：盈亏平衡价格
    ax1.bar([p['period'] for p in lockup_analysis], [p['be_price'] for p in lockup_analysis],
            color='steelblue', alpha=0.7)
    ax1.axhline(y=current_price, color='r', linestyle='--', label='当前价格')
    ax1.set_xlabel('锁定期 (月)', fontproperties=font_prop, fontsize=11)
    ax1.set_ylabel('盈亏平衡价格 (元)', fontproperties=font_prop, fontsize=11)
    ax1.set_title('锁定期对盈亏平衡价格的影响', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax1.legend(prop=font_prop)
    for label in ax1.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax1.get_yticklabels():
        label.set_fontproperties(font_prop)
    ax1.grid(True, alpha=0.3)

    # 右图：需要涨幅
    ax2.plot([p['period'] for p in lockup_analysis], [p['increase'] for p in lockup_analysis],
            'ro-', linewidth=2, markersize=8)
    ax2.set_xlabel('锁定期 (月)', fontproperties=font_prop, fontsize=11)
    ax2.set_ylabel('需要涨幅 (%)', fontproperties=font_prop, fontsize=11)
    ax2.set_title('锁定期与需要涨幅的关系', fontproperties=font_prop, fontsize=12, fontweight='bold')
    for label in ax2.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax2.get_yticklabels():
        label.set_fontproperties(font_prop)
    ax2.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_lockup_sensitivity_charts_split(issue_price, current_price, save_dir):
    """生成锁定期敏感性分析图表 - 拆分版（V2：生成2个单独图片）"""
    import os

    lockup_periods = [6, 9, 12, 18, 24, 36]  # 月
    target_return = 0.20  # 20%年化收益率

    lockup_analysis = []
    for period in lockup_periods:
        period_years = period / 12
        be_price = issue_price * (1 + target_return * period_years)
        required_increase = (be_price - issue_price) / issue_price * 100
        lockup_analysis.append({'period': period, 'be_price': be_price, 'increase': required_increase})

    # 图1：盈亏平衡价格
    fig1, ax1 = plt.subplots(figsize=(12, 6))
    ax1.bar([p['period'] for p in lockup_analysis], [p['be_price'] for p in lockup_analysis],
            color='steelblue', alpha=0.7)
    ax1.axhline(y=current_price, color='r', linestyle='--', label='当前价格')
    ax1.set_xlabel('锁定期 (月)', fontproperties=font_prop, fontsize=14)
    ax1.set_ylabel('盈亏平衡价格 (元)', fontproperties=font_prop, fontsize=14)
    ax1.set_title('锁定期对盈亏平衡价格的影响', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax1.legend(prop=font_prop, fontsize=12)
    for label in ax1.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax1.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    ax1.grid(True, alpha=0.3)

    plt.tight_layout()
    chart1_path = os.path.join(save_dir, 'lockup_sensitivity_be_price.png')
    plt.savefig(chart1_path, dpi=150, bbox_inches='tight')
    plt.close()

    # 图2：需要涨幅
    fig2, ax2 = plt.subplots(figsize=(12, 6))
    ax2.plot([p['period'] for p in lockup_analysis], [p['increase'] for p in lockup_analysis],
            'ro-', linewidth=3, markersize=10)
    ax2.set_xlabel('锁定期 (月)', fontproperties=font_prop, fontsize=14)
    ax2.set_ylabel('需要涨幅 (%)', fontproperties=font_prop, fontsize=14)
    ax2.set_title('锁定期与需要涨幅的关系', fontproperties=font_prop, fontsize=16, fontweight='bold')
    for label in ax2.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax2.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    ax2.grid(True, alpha=0.3)

    plt.tight_layout()
    chart2_path = os.path.join(save_dir, 'lockup_sensitivity_increase.png')
    plt.savefig(chart2_path, dpi=150, bbox_inches='tight')
    plt.close()

    return chart1_path, chart2_path


def generate_time_window_analysis_chart(price_series, save_dir):
    """生成不同时间窗口的风险指标对比图表"""
    import os

    chart_paths = []

    # 定义不同时间窗口（天数）
    # 月度(20日)、季度(60日)、半年(120日)、年度(250日)
    windows = [20, 60, 120, 250]

    # 计算每个窗口的统计指标
    results = {
        'window': [],
        'volatility': [],
        'total_return': [],     # 期间收益率
        'annual_return': [],    # 年化收益率
        'mean': [],
        'std': [],
        'max_drawdown': [],
        'sharpe': []
    }

    risk_free_rate = 0.03  # 无风险利率

    for window in windows:
        if len(price_series) < window:
            continue

        # 取最近window天的数据
        recent_prices = price_series[-window:]

        # 计算收益率
        returns = recent_prices.pct_change().dropna()

        # 年化波动率
        volatility = returns.std() * np.sqrt(250)

        # 期间收益率
        total_return = (recent_prices.iloc[-1] / recent_prices.iloc[0]) - 1

        # 年化收益率（使用单利计算，按交易日比例）
        # 年化收益率 = 期间收益率 × (250 / 窗口期天数)
        annual_return = total_return * (250.0 / window)

        # 均值
        mean = recent_prices.mean()

        # 标准差
        std = recent_prices.std()

        # 最大回撤
        cummax = recent_prices.cummax()
        drawdown = (recent_prices - cummax) / cummax
        max_drawdown = drawdown.min()

        # 夏普比率
        sharpe = (annual_return - risk_free_rate) / volatility if volatility > 0 else 0

        results['window'].append(window)
        results['volatility'].append(volatility)
        results['total_return'].append(total_return)     # 期间收益率
        results['annual_return'].append(annual_return)    # 年化收益率
        results['mean'].append(mean)
        results['std'].append(std)
        results['max_drawdown'].append(max_drawdown)
        results['sharpe'].append(sharpe)

    # 图1：波动率对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    bars = ax.bar(range(len(results['window'])), [v*100 for v in results['volatility']],
                   color='steelblue', alpha=0.7)
    ax.set_xlabel('时间窗口（交易日）', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('年化波动率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同时间窗口的波动率对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(results['window'])))
    ax.set_xticklabels([f'{w}日' for w in results['window']], fontproperties=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, axis='y')

    # 添加数值标签
    for bar, vol in zip(bars, results['volatility']):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                f'{vol*100:.1f}%', ha='center', va='bottom', fontsize=11, fontproperties=font_prop)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'time_window_volatility_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2：年化收益率对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    colors_return = ['green' if r > 0 else 'red' for r in results['annual_return']]
    bars = ax.bar(range(len(results['window'])), [r*100 for r in results['annual_return']],
                   color=colors_return, alpha=0.7)
    ax.set_xlabel('时间窗口（交易日）', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('年化收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同时间窗口的年化收益率对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(results['window'])))
    ax.set_xticklabels([f'{w}日' for w in results['window']], fontproperties=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, axis='y')
    ax.axhline(y=0, color='black', linestyle='-', linewidth=1)

    # 添加数值标签
    for bar, ret in zip(bars, results['annual_return']):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3 if ret > 0 else bar.get_height() - 0.3,
                f'{ret*100:+.1f}%', ha='center', va='bottom' if ret > 0 else 'top',
                fontsize=11, fontproperties=font_prop, fontweight='bold')

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'time_window_return_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3：风险指标汇总对比（单独大图）
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 7))

    # 左图：最大回撤
    bars1 = ax1.bar(range(len(results['window'])), [d*100 for d in results['max_drawdown']],
                    color='#e74c3c', alpha=0.7)
    ax1.set_xlabel('时间窗口（交易日）', fontproperties=font_prop, fontsize=12)
    ax1.set_ylabel('最大回撤 (%)', fontproperties=font_prop, fontsize=12)
    ax1.set_title('不同时间窗口的最大回撤', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax1.set_xticks(range(len(results['window'])))
    ax1.set_xticklabels([f'{w}日' for w in results['window']], fontproperties=font_prop, fontsize=10)
    ax1.grid(True, alpha=0.3, axis='y')

    for bar, dd in zip(bars1, results['max_drawdown']):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() - 0.5,
                f'{dd*100:.1f}%', ha='center', va='top', fontsize=9, fontproperties=font_prop)

    # 右图：夏普比率
    colors_sharpe = ['green' if s > 0.5 else 'orange' if s > 0 else 'red' for s in results['sharpe']]
    bars2 = ax2.bar(range(len(results['window'])), results['sharpe'],
                    color=colors_sharpe, alpha=0.7)
    ax2.set_xlabel('时间窗口（交易日）', fontproperties=font_prop, fontsize=12)
    ax2.set_ylabel('夏普比率', fontproperties=font_prop, fontsize=12)
    ax2.set_title('不同时间窗口的夏普比率', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax2.set_xticks(range(len(results['window'])))
    ax2.set_xticklabels([f'{w}日' for w in results['window']], fontproperties=font_prop, fontsize=10)
    ax2.grid(True, alpha=0.3, axis='y')
    ax2.axhline(y=1, color='black', linestyle='--', alpha=0.5, label='基准线')

    for bar, sr in zip(bars2, results['sharpe']):
        ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.02,
                f'{sr:.2f}', ha='center', va='bottom', fontsize=9, fontproperties=font_prop)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'time_window_risk_metrics_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return results, chart_paths


def generate_tornado_chart(issue_price, current_price, lockup_period, volatility, drift, save_path):
    """生成龙卷风图 - 参数敏感性排序（双维度分析）

    敏感性分析原则：
    - 所有参数使用标准化的单位变化
    - 价格类参数：变化±10%
    - 百分比参数（波动率、漂移率）：变化±10%
    - 这样可以公平比较各参数对盈利概率的影响程度

    双维度分析：
    - 维度1：盈利概率敏感性（能否盈利）
    - 维度2：预期收益率敏感性（盈利多少）

    生成两个独立的图表，分别展示盈利概率敏感性和预期收益率敏感性
    """
    # 基准参数
    base_params = {
        'issue_price': issue_price,
        'current_price': current_price,
        'volatility': volatility,
        'drift': drift
    }

    # 计算基准盈利概率
    def calculate_profit_prob(params):
        period_years = lockup_period / 12  # 使用外部变量lockup_period
        vol_period = params['volatility'] * np.sqrt(period_years)
        drift_period = params['drift'] * period_years
        required_return = (params['current_price'] - params['issue_price']) / params['issue_price']
        z_score = (drift_period - required_return) / vol_period
        return stats.norm.cdf(z_score) * 100

    # 计算预期收益率（年化）
    def calculate_expected_return(params):
        period_years = lockup_period / 12  # 使用外部变量lockup_period
        # 使用GBM模型计算预期终值
        # E[S_T] = S_0 * exp(μ * T)
        expected_final_price = params['current_price'] * np.exp(params['drift'] * period_years)
        expected_return = (expected_final_price - params['issue_price']) / params['issue_price']
        return expected_return * 100  # 转为百分比

    base_prob = calculate_profit_prob(base_params)
    base_return = calculate_expected_return(base_params)
    print(f"基准盈利概率: {base_prob:.2f}%")
    print(f"基准预期收益率: {base_return:.2f}%")

    # 参数变化分析 - 使用标准化的单位变化
    param_changes = []

    # 定义每个参数的变化方式和幅度
    param_definitions = [
        {
            'name': 'issue_price',
            'name_cn': '发行价',
            'changes': [-0.10, 0.10],  # ±10%
            'change_type': 'percent',
            'apply_change': lambda base_val, change: base_val * (1 + change)
        },
        {
            'name': 'current_price',
            'name_cn': '当前价',
            'changes': [-0.10, 0.10],  # ±10%
            'change_type': 'percent',
            'apply_change': lambda base_val, change: base_val * (1 + change)
        },
        {
            'name': 'volatility',
            'name_cn': '波动率',
            'changes': [-0.10, 0.10],  # ±10%
            'change_type': 'percent',
            'apply_change': lambda base_val, change: base_val * (1 + change)
        },
        {
            'name': 'drift',
            'name_cn': '漂移率',
            'changes': [-0.10, 0.10],  # ±10%
            'change_type': 'percent',
            'apply_change': lambda base_val, change: base_val * (1 + change)
        }
    ]

    # 计算各参数变化的敏感性（双维度）
    for param_def in param_definitions:
        param_name = param_def['name']
        param_name_cn = param_def['name_cn']

        for change in param_def['changes']:
            params = base_params.copy()
            params[param_name] = param_def['apply_change'](base_params[param_name], change)

            new_prob = calculate_profit_prob(params)
            new_return = calculate_expected_return(params)

            prob_impact = new_prob - base_prob
            return_impact = new_return - base_return

            # 计算归一化的敏感性（用于排序）
            # 盈利概率敏感性：每单位变化的影响
            # 预期收益率敏感性：每单位变化的影响
            if param_def['change_type'] == 'percent':
                # 百分比变化：每1%的影响
                prob_sensitivity = prob_impact / abs(change * 100)
                return_sensitivity = return_impact / abs(change * 100)
                change_unit = f'每1%变化'
            else:  # absolute
                # 绝对值变化：每1单位的影响
                prob_sensitivity = prob_impact / abs(change)
                return_sensitivity = return_impact / abs(change)
                change_unit = f'每1个月变化'

            # 生成变化描述
            if change < 0:
                change_desc = f"下降{abs(int(change*100))}%"
            else:
                change_desc = f"上升{int(change*100)}%"

            param_changes.append({
                'parameter': param_name_cn,
                'change_desc': change_desc,
                'prob_impact': prob_impact,
                'return_impact': return_impact,
                'prob_sensitivity': prob_sensitivity,  # 归一化的敏感性（用于排序）
                'return_sensitivity': return_sensitivity,  # 归一化的敏感性（用于排序）
                'change_unit': change_unit,  # 变化单位说明
                'full_label_prob': f'{param_name_cn}{change_desc}',
                'full_label_return': f'{param_name_cn}{change_desc}'
            })

            print(f"{param_name_cn}{change_desc}: 盈利概率 {base_prob:.2f}% → {new_prob:.2f}% ({prob_impact:+.2f}%, 敏感性: {prob_sensitivity:+.3f}), 预期收益率 {base_return:.2f}% → {new_return:.2f}% ({return_impact:+.2f}%, 敏感性: {return_sensitivity:+.3f})")

    # 按归一化的敏感性排序（找出真正最敏感的参数）
    param_changes_by_prob = sorted(param_changes, key=lambda x: abs(x['prob_sensitivity']), reverse=True)
    param_changes_by_return = sorted(param_changes, key=lambda x: abs(x['return_sensitivity']), reverse=True)

    # 只保留影响最大的前8个
    top_prob = param_changes_by_prob[:8]
    top_return = param_changes_by_return[:8]

    # 生成两个独立的图表
    # 图表1：盈利概率敏感性
    fig1, ax1 = plt.subplots(figsize=(12, 8))
    colors1 = ['red' if p['prob_impact'] < 0 else 'green' for p in top_prob]
    bars1 = ax1.barh(range(len(top_prob)), [p['prob_impact'] for p in top_prob],
                    color=colors1, alpha=0.7)

    ax1.set_xlabel('对盈利概率的影响 (百分点)', fontproperties=font_prop, fontsize=12)
    ax1.set_title('图表 4.5.1: 参数敏感性排序（龙卷风图） - 盈利概率敏感性\n标准化单位变化：价格±10%、波动率±10%、漂移率±10%',
                 fontproperties=font_prop, fontsize=13, fontweight='bold')
    ax1.set_yticks(range(len(top_prob)))
    ax1.set_yticklabels([p['full_label_prob'] for p in top_prob], fontproperties=font_prop, fontsize=10)
    ax1.axvline(x=0, color='black', linestyle='-', linewidth=1)
    ax1.grid(True, axis='x', alpha=0.3)

    # 添加数值标签
    for i, (bar, p) in enumerate(zip(bars1, top_prob)):
        width = bar.get_width()
        xpos = width + 0.3 if width >= 0 else width - 0.3
        ax1.text(xpos, i, f'{p["prob_impact"]:+.1f}%',
                va='center', fontsize=9, fontproperties=font_prop,
                color='green' if width >= 0 else 'red')

    # 设置字体
    for label in ax1.get_xticklabels() + ax1.get_yticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()
    prob_chart_path = save_path.replace('.png', '_prob.png')
    plt.savefig(prob_chart_path, dpi=150, bbox_inches='tight')
    plt.close()

    # 图表2：预期收益率敏感性
    fig2, ax2 = plt.subplots(figsize=(12, 8))
    colors2 = ['red' if p['return_impact'] < 0 else 'green' for p in top_return]
    bars2 = ax2.barh(range(len(top_return)), [p['return_impact'] for p in top_return],
                    color=colors2, alpha=0.7)

    ax2.set_xlabel('对预期收益率的影响 (百分点)', fontproperties=font_prop, fontsize=12)
    ax2.set_title('图表 4.5.2: 参数敏感性排序（龙卷风图） - 预期收益率敏感性\n标准化单位变化：价格±10%、波动率±10%、漂移率±10%',
                 fontproperties=font_prop, fontsize=13, fontweight='bold')
    ax2.set_yticks(range(len(top_return)))
    ax2.set_yticklabels([p['full_label_return'] for p in top_return], fontproperties=font_prop, fontsize=10)
    ax2.axvline(x=0, color='black', linestyle='-', linewidth=1)
    ax2.grid(True, axis='x', alpha=0.3)

    # 添加数值标签
    for i, (bar, p) in enumerate(zip(bars2, top_return)):
        width = bar.get_width()
        xpos = width + 0.3 if width >= 0 else width - 0.3
        ax2.text(xpos, i, f'{p["return_impact"]:+.1f}%',
                va='center', fontsize=9, fontproperties=font_prop,
                color='green' if width >= 0 else 'red')

    # 设置字体
    for label in ax2.get_xticklabels() + ax2.get_yticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()
    return_chart_path = save_path.replace('.png', '_return.png')
    plt.savefig(return_chart_path, dpi=150, bbox_inches='tight')
    plt.close()

    return prob_chart_path, return_chart_path, top_prob, top_return


def generate_sensitivity_chart(volatilities, profit_probs, current_vol, ma20, issue_price, save_path):
    """生成敏感性分析图表（增强版）"""
    fig = plt.figure(figsize=(16, 10))
    gs = fig.add_gridspec(2, 2, hspace=0.3, wspace=0.3)

    # 1. 波动率 vs 盈利概率柱状图
    ax1 = fig.add_subplot(gs[0, 0])
    colors = ['#27ae60' if p >= 60 else '#f39c12' if p >= 40 else '#e74c3c' for p in profit_probs]
    bars = ax1.bar(range(len(volatilities)), profit_probs, color=colors, alpha=0.8, edgecolor='black', linewidth=1.5)
    ax1.set_xlabel('年化波动率', fontproperties=font_prop, fontsize=11)
    ax1.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=11)
    ax1.set_title('不同波动率下的盈利概率', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax1.set_xticks(range(len(volatilities)))
    ax1.set_xticklabels([f'{v*100:.0f}%' for v in volatilities], fontproperties=font_prop, fontsize=9)
    ax1.grid(True, alpha=0.3, linestyle='--')
    ax1.set_ylim(0, 100)
    ax1.axhline(y=50, color='gray', linestyle='--', alpha=0.5)

    # 标记当前波动率
    current_idx = np.argmin(np.abs(np.array(volatilities) - current_vol))
    ax1.axvline(x=current_idx, color='red', linestyle='-', linewidth=2, alpha=0.7)
    ax1.text(current_idx, 95, f'当前\n{current_vol*100:.1f}%', ha='center', va='top',
            fontsize=9, fontproperties=font_prop, color='red', fontweight='bold')
    for bar, prob in zip(bars, profit_probs):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                f'{prob:.1f}%', ha='center', va='bottom', fontsize=9, fontweight='bold')

    # 2. 发行价折扣分析
    ax2 = fig.add_subplot(gs[0, 1])
    discounts = np.array([5, 10, 15, 20])  # 折价率
    issue_prices_discounted = ma20 * (1 - discounts/100)

    # 计算不同发行价下的盈利概率
    discount_probs = []
    for discounted_price in issue_prices_discounted:
        # 使用当前市场参数
        lockup_months = 6  # 假设6个月锁定期
        drift_rate = -0.1875  # 当前漂移率
        lockup_drift = drift_rate * (lockup_months / 12)
        lockup_vol = 0.3063 * np.sqrt(lockup_months / 12)  # 当前波动率

        # 计算盈利概率
        z = (np.log(ma20 / discounted_price) - lockup_drift) / lockup_vol
        prob = stats.norm.cdf(z) * 100
        discount_probs.append(prob)

    colors_discount = ['#27ae60', '#2ecc71', '#f39c12', '#e74c3c']
    bars2 = ax2.bar(range(len(discounts)), discount_probs, color=colors_discount,
                   alpha=0.8, edgecolor='black', linewidth=1.5)
    ax2.set_xlabel('发行价相对MA20折价率 (%)', fontproperties=font_prop, fontsize=11)
    ax2.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=11)
    ax2.set_title('不同发行价折扣下的盈利概率', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax2.set_xticks(range(len(discounts)))
    ax2.set_xticklabels([f'{d}%' for d in discounts], fontproperties=font_prop, fontsize=10)
    ax2.grid(True, alpha=0.3, linestyle='--')
    ax2.set_ylim(0, 100)
    ax2.axhline(y=50, color='gray', linestyle='--', alpha=0.5, label='盈亏平衡线')
    ax2.axvline(x=0, color='blue', linestyle='--', linewidth=2, alpha=0.7, label='当前发行价位置')
    ax2.legend(prop=font_prop, fontsize=9, loc='lower right')

    # 添加数值标签和发行价信息
    for bar, prob, price in zip(bars2, discount_probs, issue_prices_discounted):
        ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                f'{prob:.1f}%', ha='center', va='bottom', fontsize=9, fontweight='bold')
        ax2.text(bar.get_x() + bar.get_width()/2, -3,
                f'{price:.2f}元', ha='center', va='top', fontsize=8, fontproperties=font_prop)

    # 3. 热力图（波动率 vs 漂移率）
    ax3 = fig.add_subplot(gs[1, :])
    drift_range = np.linspace(-0.3, 0.15, 8)
    vol_range_heatmap = np.linspace(0.20, 0.50, 8)
    heatmap_data = []

    for d in drift_range:
        row = []
        for v in vol_range_heatmap:
            lockup_drift = d * (6/12)
            lockup_vol = v * np.sqrt(6/12)
            z = (np.log(ma20/issue_price) - lockup_drift) / lockup_vol
            row.append(stats.norm.cdf(z) * 100)
        heatmap_data.append(row)

    im = ax3.imshow(heatmap_data, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100,
                   extent=[vol_range_heatmap[0]*100, vol_range_heatmap[-1]*100, drift_range[0]*100, drift_range[-1]*100])
    ax3.set_xlabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
    ax3.set_ylabel('年化漂移率 (%)', fontproperties=font_prop, fontsize=12)
    ax3.set_title('盈利概率敏感性热力图（波动率 vs 漂移率）', fontproperties=font_prop, fontsize=13, fontweight='bold')
    ax3.tick_params(axis='both', which='major', labelsize=10)
    for label in ax3.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax3.get_yticklabels():
        label.set_fontproperties(font_prop)

    cbar = plt.colorbar(im, ax=ax3)
    cbar.set_label('盈利概率 (%)', fontproperties=font_prop, fontsize=11)
    cbar.ax.tick_params(labelsize=10)

    # 标记当前位置
    ax3.scatter([current_vol*100], [drift_rate*100], color='red', s=300, marker='*',
               edgecolors='white', linewidths=2, label='当前位置', zorder=5)
    ax3.legend(prop=font_prop, loc='upper right', fontsize=10)

    plt.suptitle('敏感性分析：波动率、漂移率与发行价折扣', fontproperties=font_prop,
               fontsize=15, fontweight='bold', y=0.98)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_sensitivity_charts_split(volatilities, profit_probs, current_vol, ma20, issue_price, save_dir):
    """生成敏感性分析图表 - 拆分版本（V2：生成3个单独图片）"""
    import os

    chart_paths = []

    # 图1：波动率 vs 盈利概率柱状图（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    colors = ['#27ae60' if p >= 60 else '#f39c12' if p >= 40 else '#e74c3c' for p in profit_probs]
    bars = ax.bar(range(len(volatilities)), profit_probs, color=colors, alpha=0.8, edgecolor='black', linewidth=1.5)
    ax.set_xlabel('年化波动率', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同波动率下的盈利概率', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(volatilities)))
    ax.set_xticklabels([f'{v*100:.0f}%' for v in volatilities], fontproperties=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.set_ylim(0, 100)
    ax.axhline(y=50, color='gray', linestyle='--', alpha=0.5)

    # 标记当前波动率
    current_idx = np.argmin(np.abs(np.array(volatilities) - current_vol))
    ax.axvline(x=current_idx, color='red', linestyle='-', linewidth=2, alpha=0.7)
    ax.text(current_idx, 95, f'当前\n{current_vol*100:.1f}%', ha='center', va='top',
            fontsize=12, fontproperties=font_prop, color='red', fontweight='bold')
    for bar, prob in zip(bars, profit_probs):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                f'{prob:.1f}%', ha='center', va='bottom', fontsize=11, fontweight='bold')

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'sensitivity_volatility_profit_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2：发行价折扣分析（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    discounts = np.array([5, 10, 15, 20])  # 折价率
    issue_prices_discounted = ma20 * (1 - discounts/100)

    # 计算不同发行价下的盈利概率
    discount_probs = []
    for discounted_price in issue_prices_discounted:
        lockup_months = 6
        drift_rate = -0.1875
        lockup_drift = drift_rate * (lockup_months / 12)
        lockup_vol = 0.3063 * np.sqrt(lockup_months / 12)
        z = (np.log(ma20 / discounted_price) - lockup_drift) / lockup_vol
        prob = stats.norm.cdf(z) * 100
        discount_probs.append(prob)

    colors_discount = ['#27ae60', '#2ecc71', '#f39c12', '#e74c3c']
    bars = ax.bar(range(len(discounts)), discount_probs, color=colors_discount,
                   alpha=0.8, edgecolor='black', linewidth=1.5)
    ax.set_xlabel('发行价相对MA20折价率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同发行价折扣下的盈利概率', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(discounts)))
    ax.set_xticklabels([f'{d}%' for d in discounts], fontproperties=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.set_ylim(0, 100)
    ax.axhline(y=50, color='gray', linestyle='--', alpha=0.5, label='盈亏平衡线')
    ax.axvline(x=0, color='blue', linestyle='--', linewidth=2, alpha=0.7, label='当前发行价位置')
    ax.legend(prop=font_prop, fontsize=12, loc='lower right')

    for bar, prob, price in zip(bars, discount_probs, issue_prices_discounted):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                f'{prob:.1f}%', ha='center', va='bottom', fontsize=11, fontweight='bold')
        ax.text(bar.get_x() + bar.get_width()/2, -3,
                f'{price:.2f}元', ha='center', va='top', fontsize=10, fontproperties=font_prop)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'sensitivity_discount_profit_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3：热力图（波动率 vs 漂移率）- 单独大图
    fig, ax = plt.subplots(figsize=(14, 10))
    drift_range = np.linspace(-0.3, 0.15, 8)
    vol_range_heatmap = np.linspace(0.20, 0.50, 8)
    heatmap_data = []

    for d in drift_range:
        row = []
        for v in vol_range_heatmap:
            lockup_drift = d * (6/12)
            lockup_vol = v * np.sqrt(6/12)
            z = (np.log(ma20/issue_price) - lockup_drift) / lockup_vol
            row.append(stats.norm.cdf(z) * 100)
        heatmap_data.append(row)

    im = ax.imshow(heatmap_data, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100,
                   extent=[vol_range_heatmap[0]*100, vol_range_heatmap[-1]*100, drift_range[0]*100, drift_range[-1]*100])
    ax.set_xlabel('年化波动率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('年化漂移率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('盈利概率敏感性热力图（波动率 vs 漂移率）', fontproperties=font_prop, fontsize=16, fontweight='bold')
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    cbar = plt.colorbar(im, ax=ax)
    cbar.set_label('盈利概率 (%)', fontproperties=font_prop, fontsize=12)

    # 标记当前位置
    ax.scatter([current_vol*100], [drift_rate*100], color='red', s=300, marker='*',
               edgecolors='white', linewidths=2, label='当前位置', zorder=5)
    ax.legend(prop=font_prop, loc='upper right', fontsize=12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'sensitivity_heatmap_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图4：热力图（漂移率 vs 折价率）- 单独大图
    fig, ax = plt.subplots(figsize=(16, 10))
    drift_range_new = np.linspace(-0.30, 0.30, 7)  # -30%到+30%，7个档位
    discount_range = np.array([-30, -20, -10, -5, 0, 5, 10, 20, 30])  # 折价/溢价率-30%到+30%，更全面的范围
    heatmap_data_drift_discount = []

    for d in drift_range_new:
        row = []
        for discount in discount_range:
            issue_price_hm = ma20 * (1 + discount/100)  # 负值表示折价，正值表示溢价
            lockup_drift = d * (6/12)
            lockup_vol = 0.3063 * np.sqrt(6/12)  # 使用当前波动率
            z = (np.log(ma20/issue_price_hm) - lockup_drift) / lockup_vol
            row.append(stats.norm.cdf(z) * 100)
        heatmap_data_drift_discount.append(row)

    im = ax.imshow(heatmap_data_drift_discount, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100)

    # 设置刻度
    ax.set_xticks(np.arange(len(discount_range)))
    ax.set_yticks(np.arange(len(drift_range_new)))
    ax.set_xticklabels([f'{d}%' for d in discount_range], fontproperties=font_prop, fontsize=12)
    ax.set_yticklabels([f'{d*100:+.0f}%' for d in drift_range_new], fontproperties=font_prop, fontsize=12)

    # 添加数值标注
    for i in range(len(drift_range_new)):
        for j in range(len(discount_range)):
            text = ax.text(j, i, f'{heatmap_data_drift_discount[i][j]:.0f}%',
                         ha="center", va="center", color="black", fontproperties=font_prop, fontsize=8,
                         fontweight='bold' if heatmap_data_drift_discount[i][j] >= 50 else 'normal')

    ax.set_xlabel('发行价折价/溢价率（负值表示折价，正值表示溢价）(%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('年化漂移率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('盈利概率敏感性热力图（漂移率 vs 折价率）', fontproperties=font_prop, fontsize=16, fontweight='bold')

    cbar = plt.colorbar(im, ax=ax)
    cbar.set_label('盈利概率 (%)', fontproperties=font_prop, fontsize=12)

    # 标记当前位置（需要传入当前折价/溢价率）
    current_vol_for_hm = 0.3063  # 从外部获取当前波动率
    current_drift_for_hm = -0.1875  # 从外部获取当前漂移率
    current_discount_for_hm = ((issue_price - ma20) / ma20) * 100  # 计算当前折价/溢价率（负值表示折价）

    # 找到最接近的刻度位置
    drift_idx = np.argmin(np.abs(drift_range_new - current_drift_for_hm))
    discount_idx = np.argmin(np.abs(discount_range - current_discount_for_hm))

    ax.scatter([discount_idx], [drift_idx], color='red', s=300, marker='*',
               edgecolors='white', linewidths=2, label='当前位置', zorder=5)
    ax.legend(prop=font_prop, loc='upper right', fontsize=12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'sensitivity_heatmap_drift_discount_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths


def generate_discount_scenario_charts_split(base_price, current_price, volatility, drift, lockup_period, save_dir):
    """生成发行价折扣情景图表 - 统一版本（-20%至+20%）

    参数:
        base_price: 基准价格（MA20）
        current_price: 当前价格
        volatility: 波动率
        drift: 漂移率
        lockup_period: 锁定期（月）
        save_dir: 保存目录

    返回:
        图表路径列表 [盈利概率图, 预期收益率图]
    """
    import os

    chart_paths = []
    # 统一范围：从-20%（折价）到+20%（溢价），包括0%
    all_rates = [-20, -15, -10, -5, 0, 5, 10, 15, 20]

    # 盈利目标：分析不同盈利幅度的概率
    profit_targets = [0, 5, 10, 15, 20, 30, 50]  # 盈利目标百分比

    # 计算所有情景的数据
    all_data = []
    lockup_vol = volatility * np.sqrt(lockup_period / 12)
    lockup_drift = drift * (lockup_period / 12)
    expected_future_price = current_price * np.exp(lockup_drift)

    for rate in all_rates:
        if rate < 0:
            # 折价发行
            issue_price = base_price * (1 + rate/100)  # rate是负数
            threshold = issue_price
        elif rate == 0:
            # 平价发行
            issue_price = base_price
            threshold = issue_price
        else:
            # 溢价发行
            issue_price = base_price * (1 + rate/100)
            threshold = max(base_price, issue_price * 1.02)

        # 计算盈利概率
        required_return = (threshold - current_price) / current_price
        z_score = (lockup_drift - required_return) / lockup_vol
        profit_prob = (1 - stats.norm.cdf(-z_score)) * 100

        # 计算预期收益率
        expected_return = (expected_future_price - issue_price) / issue_price * 100

        all_data.append({
            'rate': rate,
            'profit_prob': profit_prob,
            'expected_return': expected_return,
            'issue_price': issue_price
        })

        # 计算不同盈利目标的概率
        profit_probs = {}
        for target_pct in profit_targets:
            target_price = issue_price * (1 + target_pct/100)
            required_return = (target_price - current_price) / current_price
            z_score = (lockup_drift - required_return) / lockup_vol
            prob = (1 - stats.norm.cdf(-z_score)) * 100
            profit_probs[f'prob_{target_pct}pct'] = prob

        # 基础盈利概率（0%）
        required_return = (threshold - current_price) / current_price
        z_score = (lockup_drift - required_return) / lockup_vol
        profit_prob = (1 - stats.norm.cdf(-z_score)) * 100

        # 修正：使用真实的预期收益率（基于漂移率）
        # 期望的未来价格 = current_price * exp(lockup_drift)
        # 预期收益率 = (期望的未来价格 - 发行价) / 发行价
        expected_future_price = current_price * np.exp(lockup_drift)
        expected_return = (expected_future_price - issue_price) / issue_price * 100

    # 图1：盈利概率对比（统一图）
    fig, ax = plt.subplots(figsize=(14, 8))

    # 根据折价/溢价设置颜色
    colors = []
    for d in all_data:
        if d['rate'] < 0:
            # 折价发行：绿色
            colors.append('#27ae60')
        elif d['rate'] == 0:
            # 平价发行：灰色
            colors.append('#7f8c8d')
        else:
            # 溢价发行：红色
            colors.append('#e74c3c')

    bars = ax.bar([d['rate'] for d in all_data], [d['profit_prob'] for d in all_data],
                 color=colors, alpha=0.8, edgecolor='black', linewidth=2)

    ax.set_xlabel('发行价相对MA20的折扣/溢价率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('发行价敏感性分析 - 盈利概率', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks([d['rate'] for d in all_data])
    ax.set_xticklabels([f"{d['rate']}%" for d in all_data], fontproperties=font_prop, fontsize=12)
    ax.set_ylim(0, 100)
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.axhline(y=50, color='gray', linestyle='--', alpha=0.5, label='盈亏平衡线')
    ax.axvline(x=0, color='blue', linestyle=':', linewidth=2, label='MA20基准线')

    for bar, d in zip(bars, all_data):
        y_pos = d['profit_prob'] + 1 if d['profit_prob'] < 95 else d['profit_prob'] - 3
        ax.text(bar.get_x() + bar.get_width()/2, y_pos,
                f"{d['profit_prob']:.1f}%", ha='center', va='bottom' if d['profit_prob'] < 95 else 'top',
                fontsize=11, fontweight='bold')

    ax.legend(prop=font_prop, fontsize=12)
    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'scenario_unified_profit_prob.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2：预期收益率对比（统一图）
    fig, ax = plt.subplots(figsize=(14, 8))

    colors = []
    for d in all_data:
        if d['rate'] < 0:
            colors.append('#27ae60')  # 折价：绿色
        elif d['rate'] == 0:
            colors.append('#7f8c8d')  # 平价：灰色
        else:
            colors.append('#e74c3c')  # 溢价：红色

    bars = ax.bar([d['rate'] for d in all_data], [d['expected_return'] for d in all_data],
                 color=colors, alpha=0.8, edgecolor='black', linewidth=2)

    ax.set_xlabel('发行价相对MA20的折扣/溢价率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('预期收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('发行价敏感性分析 - 预期收益率', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks([d['rate'] for d in all_data])
    ax.set_xticklabels([f"{d['rate']}%" for d in all_data], fontproperties=font_prop, fontsize=12)
    ax.axhline(y=0, color='gray', linestyle='--', alpha=0.7, linewidth=2)
    ax.axvline(x=0, color='blue', linestyle=':', linewidth=2, label='MA20基准线')
    ax.grid(True, alpha=0.3, linestyle='--')

    for bar, d in zip(bars, all_data):
        height = d['expected_return']
        y_pos = height + 1 if height >= 0 else height - 3
        ax.text(bar.get_x() + bar.get_width()/2, y_pos,
                f'{height:.1f}%', ha='center', va='bottom' if height >= 0 else 'top',
                fontsize=11, fontweight='bold')

    ax.legend(prop=font_prop, fontsize=12)
    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'scenario_unified_expected_return.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths


def generate_multi_dimension_scenario_charts_split(current_price, base_price, volatility, drift, lockup_period, save_dir):
    """生成多维度情景分析图表（波动率×折扣率×时间窗口） - 拆分版本（3个单独图片）

    参数:
        current_price: 当前价格
        base_price: 基准价格（MA20）
        volatility: 波动率
        drift: 漂移率
        lockup_period: 锁定期（月）
        save_dir: 保存目录

    返回:
        图表路径列表
    """
    import os

    chart_paths = []

    # 定义情景参数
    volatility_scenarios = [0.20, 0.30, 0.40, 0.50]
    discount_scenarios = [-10, -5, 0, 5, 10, 15, 20]
    window_scenarios = [60, 120, 180]

    # 计算所有情景
    multi_dim_results = []
    for vol in volatility_scenarios:
        for discount in discount_scenarios:
            if discount >= 0:
                issue_price = base_price * (1 - discount/100)
                threshold = issue_price
            else:
                issue_price = base_price * (1 + abs(discount)/100)
                threshold = max(base_price, issue_price * 1.02)

            for window in window_scenarios:
                window_years = window / 250.0
                window_drift = drift * window_years
                window_vol = vol * np.sqrt(window_years)
                required_return = (threshold - current_price) / current_price
                z_score = (window_drift - required_return) / window_vol
                profit_prob = (1 - stats.norm.cdf(-z_score)) * 100
                expected_return = (base_price - issue_price) / issue_price * 100

                multi_dim_results.append({
                    '波动率': vol, '折扣率': discount, '时间窗口': window,
                    '盈利概率': profit_prob, '预期收益率': expected_return
                })

    df_multi = pd.DataFrame(multi_dim_results)

    # 图1：热力图 - 波动率 × 折价率 → 盈利概率（60日窗口）
    fig, ax = plt.subplots(figsize=(16, 10))
    df_60d = df_multi[df_multi['时间窗口'] == 60].copy()
    pivot_data = df_60d.pivot(index='波动率', columns='折扣率', values='盈利概率')

    vol_labels = [f"{int(v*100)}%" for v in pivot_data.index]
    discount_labels = [f"{int(d):+d}%" for d in pivot_data.columns]

    # 使用索引和列数来设置刻度位置，而不是extent
    im = ax.imshow(pivot_data.values, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100)
    ax.set_xlabel('折扣率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('波动率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('盈利概率热力图：波动率 × 折价率 (60日窗口)', fontproperties=font_prop, fontsize=16, fontweight='bold')

    # 设置刻度位置（使用索引位置）
    ax.set_xticks(range(len(discount_scenarios)))
    ax.set_xticklabels(discount_labels, fontproperties=font_prop, fontsize=12)
    ax.set_yticks(range(len(volatility_scenarios)))
    ax.set_yticklabels(vol_labels, fontproperties=font_prop, fontsize=12)

    cbar = plt.colorbar(im, ax=ax)
    cbar.set_label('盈利概率 (%)', fontproperties=font_prop, fontsize=12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'multi_dim_heatmap_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2：情景对比 - 不同波动率下的折价/溢价影响
    fig, axes = plt.subplots(1, 2, figsize=(16, 6))

    # 左图：折价20%情景
    discount_20_data = df_multi[df_multi['折扣率'] == 20]
    for window in [60, 120, 180]:
        data = discount_20_data[discount_20_data['时间窗口'] == window]
        if len(data) > 0:
            axes[0].plot(data['波动率']*100, data['盈利概率'], 'o-',
                       label=f'{window}日窗口', linewidth=2, markersize=8)
    axes[0].set_xlabel('波动率 (%)', fontproperties=font_prop, fontsize=12)
    axes[0].set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=12)
    axes[0].set_title('折价20% - 不同波动率下的盈利概率', fontproperties=font_prop, fontsize=14, fontweight='bold')
    axes[0].legend(prop=font_prop)
    axes[0].grid(True, alpha=0.3)
    axes[0].axhline(y=50, color='gray', linestyle='--', alpha=0.5)

    # 右图：溢价10%情景
    premium_10_data = df_multi[df_multi['折扣率'] == -10]
    for window in [60, 120, 180]:
        data = premium_10_data[premium_10_data['时间窗口'] == window]
        if len(data) > 0:
            axes[1].plot(data['波动率']*100, data['盈利概率'], 's-',
                       label=f'{window}日窗口', linewidth=2, markersize=10)
    axes[1].set_xlabel('波动率 (%)', fontproperties=font_prop, fontsize=12)
    axes[1].set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=12)
    axes[1].set_title('溢价10% - 不同波动率下的盈利概率', fontproperties=font_prop, fontsize=14, fontweight='bold')
    axes[1].legend(prop=font_prop)
    axes[1].grid(True, alpha=0.3)
    axes[1].axhline(y=50, color='gray', linestyle='--', alpha=0.5)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'multi_dim_comparison_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3：优质情景TOP10
    df_60d = df_multi[df_multi['时间窗口'] == 60].copy()
    best_scenarios = df_60d[df_60d['盈利概率'] > 50].sort_values('盈利概率', ascending=False).head(10)

    fig, ax = plt.subplots(figsize=(12, 8))
    y_pos = range(len(best_scenarios))
    scenario_labels = [f"{int(row['波动率']*100)}%/{int(row['折扣率']):+d}%"
                      for _, row in best_scenarios.iterrows()]
    colors = ['#27ae60' if p > 80 else '#2ecc71' if p > 70 else '#f39c12'
              for p in best_scenarios['盈利概率']]
    bars = ax.barh(y_pos, best_scenarios['盈利概率'], color=colors, alpha=0.8, edgecolor='black', linewidth=1.5)

    ax.set_yticks(y_pos)
    ax.set_yticklabels(scenario_labels, fontproperties=font_prop, fontsize=11)
    ax.set_xlabel('盈利概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('优质情景TOP10 (盈利概率>50%) - 60日窗口', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xlim(0, 100)
    ax.axvline(x=50, color='gray', linestyle='--', alpha=0.5, linewidth=2, label='盈亏平衡线')
    ax.grid(True, alpha=0.3, axis='x')
    ax.legend(prop=font_prop, fontsize=12)

    for i, (_, row) in enumerate(best_scenarios.iterrows()):
        ax.text(row['盈利概率']+0.5, i, f"  {row['盈利概率']:.1f}%",
                va='center', fontsize=10, fontweight='bold')

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'multi_dim_best_scenarios_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths


def generate_stress_test_chart(scenarios, returns, save_path):
    """生成压力测试图表"""
    plt.figure(figsize=(10, 6))
    colors = ['#e74c3c' if r < 0 else '#2ecc71' for r in returns]
    bars = plt.barh(range(len(scenarios)), returns, color=colors, alpha=0.7)
    plt.xlabel('收益率 (%)', fontproperties=font_prop, fontsize=12)
    plt.ylabel('情景', fontproperties=font_prop, fontsize=12)
    plt.title('压力测试情景分析', fontproperties=font_prop, fontsize=14, fontweight='bold')
    plt.yticks(range(len(scenarios)), scenarios, fontproperties=font_prop)
    plt.xticks(fontproperties=font_prop)
    plt.axvline(x=0, color='black', linestyle='-', linewidth=1)
    plt.grid(True, alpha=0.3, axis='x')

    for bar, ret in zip(bars, returns):
        plt.text(bar.get_width() + (1 if ret > 0 else -1),
                bar.get_y() + bar.get_height()/2,
                f'{ret:.1f}%', ha='left' if ret > 0 else 'right',
                va='center', fontsize=10, fontproperties=font_prop)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_monte_carlo_chart(final_prices, issue_price, current_price, save_path):
    """生成蒙特卡洛模拟图表（增强版）"""
    fig = plt.figure(figsize=(16, 10))

    # 1. 价格分布直方图
    ax1 = plt.subplot(2, 3, 1)
    ax1.hist(final_prices, bins=50, alpha=0.7, color='steelblue', edgecolor='black')
    ax1.axvline(x=issue_price, color='red', linestyle='--', linewidth=2, label=f'发行价: {issue_price:.2f}')
    ax1.axvline(x=current_price, color='green', linestyle='--', linewidth=2, label=f'当前价: {current_price:.2f}')
    ax1.set_xlabel('未来价格 (元)', fontproperties=font_prop, fontsize=11)
    ax1.set_ylabel('频数', fontproperties=font_prop, fontsize=11)
    ax1.set_title('锁定期后价格分布', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax1.legend(prop=font_prop, fontsize=9)
    ax1.grid(True, alpha=0.3)
    for label in ax1.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax1.get_yticklabels():
        label.set_fontproperties(font_prop)

    # 2. 收益率分布
    ax2 = plt.subplot(2, 3, 2)
    returns = (final_prices - issue_price) / issue_price * 100
    ax2.hist(returns, bins=50, alpha=0.7, color='coral', edgecolor='black')
    ax2.axvline(x=0, color='red', linestyle='--', linewidth=2, label='盈亏平衡')
    ax2.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=11)
    ax2.set_ylabel('频数', fontproperties=font_prop, fontsize=11)
    ax2.set_title('收益率分布', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax2.legend(prop=font_prop, fontsize=9)
    ax2.grid(True, alpha=0.3)
    for label in ax2.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax2.get_yticklabels():
        label.set_fontproperties(font_prop)

    # 3. 累积分布函数
    ax3 = plt.subplot(2, 3, 3)
    sorted_returns = np.sort(returns)
    cumulative = np.arange(1, len(sorted_returns) + 1) / len(sorted_returns) * 100
    ax3.plot(sorted_returns, cumulative, linewidth=2, color='#8e44ad')
    ax3.axvline(x=0, color='red', linestyle='--', linewidth=1.5, label='盈亏平衡')
    ax3.axhline(y=50, color='gray', linestyle='--', linewidth=1, alpha=0.5)
    ax3.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=11)
    ax3.set_ylabel('累积概率 (%)', fontproperties=font_prop, fontsize=11)
    ax3.set_title('累积分布函数', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax3.legend(prop=font_prop, fontsize=9)
    ax3.grid(True, alpha=0.3)
    for label in ax3.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax3.get_yticklabels():
        label.set_fontproperties(font_prop)

    # 4. 盈亏概率饼图
    ax4 = plt.subplot(2, 3, 4)
    profit_prob = (returns > 0).mean() * 100
    loss_prob = 100 - profit_prob
    colors_pie = ['#27ae60', '#e74c3c']
    wedges, texts, autotexts = ax4.pie([profit_prob, loss_prob], labels=['盈利', '亏损'],
                                         autopct='%1.1f%%', colors=colors_pie, startangle=90)
    ax4.set_title('盈亏概率分布', fontproperties=font_prop, fontsize=12, fontweight='bold')
    for text in texts:
        text.set_fontproperties(font_prop)
    for autotext in autotexts:
        autotext.set_fontproperties(font_prop)
        autotext.set_fontsize(12)
        autotext.set_fontweight('bold')

    # 5. 关键统计指标
    ax5 = plt.subplot(2, 3, 5)
    ax5.axis('off')
    stats_text = f"""
    关键统计指标
    ════════════════

    盈利概率: {profit_prob:.1f}%

    预期收益率: {returns.mean():.1f}%

    收益率中位数: {np.median(returns):.1f}%

    95% VaR: {abs(np.percentile(returns, 5)):.1f}%

    99% VaR: {abs(np.percentile(returns, 1)):.1f}%

    最大回撤: {abs(returns.min()):.1f}%

    收益率标准差: {returns.std():.1f}%

    盈亏平衡价: {issue_price:.2f}元

    当前价格: {current_price:.2f}元
    """
    ax5.text(0.1, 0.5, stats_text, transform=ax5.transAxes,
            fontsize=11, verticalalignment='center',
            fontproperties=font_prop,
            family='monospace',
            bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.3))

    # 6. 概率分布曲线
    ax6 = plt.subplot(2, 3, 6)
    from scipy.stats import norm
    mu, std = returns.mean(), returns.std()
    x = np.linspace(returns.min(), returns.max(), 100)
    pdf = norm.pdf(x, mu, std)
    ax6.plot(x, pdf, linewidth=2, color='#2980b9', label='正态分布拟合')
    ax6.fill_between(x, pdf, alpha=0.3, color='#3498db')
    ax6.axvline(x=0, color='red', linestyle='--', linewidth=2, label='盈亏平衡')
    ax6.axvline(x=mu, color='green', linestyle='--', linewidth=2, label=f'均值: {mu:.1f}%')
    ax6.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=11)
    ax6.set_ylabel('概率密度', fontproperties=font_prop, fontsize=11)
    ax6.set_title('收益率概率分布', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax6.legend(prop=font_prop, fontsize=9)
    ax6.grid(True, alpha=0.3)
    for label in ax6.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax6.get_yticklabels():
        label.set_fontproperties(font_prop)

    plt.suptitle('蒙特卡洛模拟分析 (10000次模拟)', fontproperties=font_prop,
               fontsize=14, fontweight='bold', y=0.995)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_monte_carlo_charts_split(final_prices, issue_price, current_price, save_dir):
    """生成蒙特卡洛模拟图表 - 拆分版本（6个单独大图）"""
    import os

    returns = (final_prices - issue_price) / issue_price * 100
    chart_paths = []

    # 图1: 价格分布直方图（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.hist(final_prices, bins=50, alpha=0.7, color='steelblue', edgecolor='black')
    ax.axvline(x=issue_price, color='red', linestyle='--', linewidth=2, label=f'发行价: {issue_price:.2f}')
    ax.axvline(x=current_price, color='green', linestyle='--', linewidth=2, label=f'当前价: {current_price:.2f}')
    ax.set_xlabel('未来价格 (元)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('频数', fontproperties=font_prop, fontsize=14)
    ax.set_title('锁定期后价格分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_price_distribution_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2: 收益率分布（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.hist(returns, bins=50, alpha=0.7, color='coral', edgecolor='black')
    ax.axvline(x=0, color='red', linestyle='--', linewidth=2, label='盈亏平衡')
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('频数', fontproperties=font_prop, fontsize=14)
    ax.set_title('收益率分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_return_distribution_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3: 累积分布函数（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    sorted_returns = np.sort(returns)
    cumulative = np.arange(1, len(sorted_returns) + 1) / len(sorted_returns) * 100
    ax.plot(sorted_returns, cumulative, linewidth=2, color='#8e44ad')
    ax.axvline(x=0, color='red', linestyle='--', linewidth=1.5, label='盈亏平衡')
    ax.axhline(y=50, color='gray', linestyle='--', linewidth=1, alpha=0.5)
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('累积概率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('累积分布函数', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_cdf_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图4: 盈亏概率饼图（单独大图）
    fig, ax = plt.subplots(figsize=(12, 10))
    profit_prob = (returns > 0).mean() * 100
    loss_prob = 100 - profit_prob
    colors_pie = ['#27ae60', '#e74c3c']
    wedges, texts, autotexts = ax.pie([profit_prob, loss_prob], labels=['盈利', '亏损'],
                                         autopct='%1.1f%%', colors=colors_pie, startangle=90,
                                         textprops={'fontproperties': font_prop, 'fontsize': 14})
    ax.set_title('盈亏概率分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    for autotext in autotexts:
        autotext.set_fontproperties(font_prop)
        autotext.set_fontsize(16)
        autotext.set_fontweight('bold')

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_profit_loss_pie_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图5: 关键统计指标（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.axis('off')
    stats_text = f"""
    关键统计指标
    ════════════════

    盈利概率: {profit_prob:.1f}%

    预期收益率: {returns.mean():.1f}%

    收益率中位数: {np.median(returns):.1f}%

    95% VaR: {abs(np.percentile(returns, 5)):.1f}%

    99% VaR: {abs(np.percentile(returns, 1)):.1f}%

    最大回撤: {abs(returns.min()):.1f}%

    收益率标准差: {returns.std():.1f}%

    盈亏平衡价: {issue_price:.2f}元

    当前价格: {current_price:.2f}元
    """
    ax.text(0.1, 0.5, stats_text, transform=ax.transAxes,
            fontsize=16, verticalalignment='center',
            fontproperties=font_prop,
            family='monospace',
            bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.3))

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_stats_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图6: 概率分布曲线（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))
    from scipy.stats import norm
    mu, std = returns.mean(), returns.std()
    x = np.linspace(returns.min(), returns.max(), 100)
    pdf = norm.pdf(x, mu, std)
    ax.plot(x, pdf, linewidth=2, color='#2980b9', label='正态分布拟合')
    ax.fill_between(x, pdf, alpha=0.3, color='#3498db')
    ax.axvline(x=0, color='red', linestyle='--', linewidth=2, label='盈亏平衡')
    ax.axvline(x=mu, color='green', linestyle='--', linewidth=2, label=f'均值: {mu:.1f}%')
    ax.set_xlabel('收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('概率密度', fontproperties=font_prop, fontsize=14)
    ax.set_title('收益率概率分布', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(12)

    plt.tight_layout()
    chart_path = os.path.join(save_dir, 'monte_carlo_pdf_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths


def generate_dcf_valuation_heatmap(save_path, current_price, net_income, total_shares, net_debt=20.0):
    """生成DCF估值热力图

    参数:
        save_path: 保存路径
        current_price: 当前股价
        net_income: 净利润（元）
        total_shares: 总股数
        net_debt: 净债务（亿元），默认20亿元
    """
    fig, ax = plt.subplots(figsize=(14, 9))

    # WACC 和永续增长率范围
    wacc_range = np.linspace(0.08, 0.13, 7)   # 8%-13%
    growth_range = np.linspace(0.15, 0.30, 7) # 15%-30% 增长率

    # 估值参数
    # FCF = 净利润 × 100%
    base_fcf = net_income / 100000000  # 转换为亿元
    # 股数（亿股）
    shares_billion = total_shares / 100000000

    # 计算每股价值矩阵
    pivot_data = []
    for g in growth_range:
        row = []
        for w in wacc_range:
            # DCF计算：FCF按g增长
            fcfs = [base_fcf * ((1 + g) ** i) for i in range(10)]
            pv_fcfs = sum([fcf / ((1 + w) ** (i+1)) for i, fcf in enumerate(fcfs)])

            # 终值：永续增长率设为3%
            terminal_growth = 0.03
            terminal_fcf = fcfs[-1] * (1 + terminal_growth)
            terminal_value = terminal_fcf / (w - terminal_growth) if w > terminal_growth else 0
            pv_terminal = terminal_value / ((1 + w) ** 10)

            enterprise_value = pv_fcfs + pv_terminal  # 亿元
            equity_value = enterprise_value - net_debt  # 减净债务（亿元）
            value_per_share = equity_value / shares_billion  # 每股价值（元）
            row.append(value_per_share)
        pivot_data.append(row)

    pivot_table = pd.DataFrame(
        pivot_data,
        index=[f'{g*100:.1f}%' for g in growth_range],
        columns=[f'{w*100:.1f}%' for w in wacc_range]
    )

    # 绘制热力图
    heatmap = sns.heatmap(pivot_table, annot=True, fmt='.2f', cmap='RdYlGn',
                          center=current_price, vmin=current_price*0.5, vmax=current_price*1.5,
                          cbar_kws={'label': '每股价值（元）'}, ax=ax)

    # 修复 colorbar 中文
    cbar = heatmap.collections[0].colorbar
    cbar.set_label('每股价值（元）', fontproperties=font_prop, fontsize=12)

    # 设置轴标签
    ax.set_xlabel('WACC（加权平均资本成本）', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('预测期FCF增长率', fontproperties=font_prop, fontsize=12)
    ax.set_title('DCF估值敏感性分析矩阵（每股价值：元）',
                fontproperties=font_prop, fontsize=14, fontweight='bold', pad=15)

    # 设置刻度标签字体
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(10)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(10)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_var_chart(var_95, var_99, cvar_95, save_path):
    """生成VaR风险度量图表"""
    plt.figure(figsize=(10, 6))
    metrics = ['95% VaR', '99% VaR', '95% CVaR']
    values = [var_95, var_99, cvar_95]
    colors = ['#e74c3c', '#c0392b', '#922b21']

    bars = plt.bar(metrics, [v*100 for v in values], color=colors, alpha=0.7)
    plt.ylabel('潜在损失 (%)', fontproperties=font_prop, fontsize=12)
    plt.title('风险度量 (VaR/CVaR)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    plt.xticks(fontproperties=font_prop)
    plt.yticks(fontproperties=font_prop)
    plt.grid(True, alpha=0.3, axis='y')

    for bar, value in zip(bars, values):
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                f'{value*100:.1f}%', ha='center', va='bottom',
                fontsize=11, fontweight='bold', fontproperties=font_prop)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def generate_relative_valuation_charts(current_metrics, industry_avg, peer_companies_data, save_dir):
    """生成相对估值分析图表"""
    import os

    # 图1: 估值指标对比（柱状图）
    fig, axes = plt.subplots(2, 2, figsize=(16, 12))

    # 1. 估值指标对比
    ax1 = axes[0, 0]
    metrics = ['PE', 'PS', 'PB']
    current_vals = [current_metrics['pe'], current_metrics['ps'], current_metrics['pb']]
    industry_vals = [industry_avg['pe'], industry_avg['ps'], industry_avg['pb']]

    x = np.arange(len(metrics))
    width = 0.35

    bars1 = ax1.bar(x - width/2, current_vals, width, label='光弘科技', color='#3498db', alpha=0.8)
    bars2 = ax1.bar(x + width/2, industry_vals, width, label='行业平均', color='#e74c3c', alpha=0.8)

    ax1.set_xlabel('估值指标', fontproperties=font_prop, fontsize=12)
    ax1.set_ylabel('倍数', fontproperties=font_prop, fontsize=12)
    ax1.set_title('相对估值对比', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax1.set_xticks(x)
    ax1.set_xticklabels(metrics, fontproperties=font_prop)
    ax1.legend(prop=font_prop)
    ax1.grid(True, alpha=0.3, axis='y')

    for bars in [bars1, bars2]:
        for bar in bars:
            height = bar.get_height()
            ax1.text(bar.get_x() + bar.get_width()/2., height,
                    f'{height:.1f}', ha='center', va='bottom', fontsize=10, fontproperties=font_prop)

    # 2. PE对比
    ax2 = axes[0, 1]
    peer_names = list(peer_companies_data['name'])
    peer_pes = list(peer_companies_data['pe'])
    peer_pes.append(industry_avg['pe'])
    peer_names.append('行业平均')
    peer_pes.append(current_metrics['pe'])
    peer_names.append('光弘科技')

    colors_pe = ['#95a5a6'] * len(peer_companies_data) + ['#e74c3c', '#3498db']
    bars = ax2.bar(range(len(peer_names)), peer_pes, color=colors_pe, alpha=0.7)
    ax2.axhline(y=industry_avg['pe'], color='red', linestyle='--', alpha=0.5)
    ax2.set_xlabel('公司', fontproperties=font_prop, fontsize=10)
    ax2.set_ylabel('PE (倍)', fontproperties=font_prop, fontsize=10)
    ax2.set_title('PE倍数对比', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax2.set_xticks(range(len(peer_names)))
    ax2.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=9, rotation=45)
    ax2.grid(True, alpha=0.3, axis='y')

    # 3. PB对比
    ax3 = axes[1, 0]
    peer_pbs = list(peer_companies_data['pb'])
    peer_pbs.append(industry_avg['pb'])
    peer_pbs.append(current_metrics['pb'])

    bars = ax3.bar(range(len(peer_names)), peer_pbs, color=colors_pe, alpha=0.7)
    ax3.axhline(y=industry_avg['pb'], color='red', linestyle='--', alpha=0.5)
    ax3.set_xlabel('公司', fontproperties=font_prop, fontsize=10)
    ax3.set_ylabel('PB (倍)', fontproperties=font_prop, fontsize=10)
    ax3.set_title('PB倍数对比', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax3.set_xticks(range(len(peer_names)))
    ax3.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=9, rotation=45)
    ax3.grid(True, alpha=0.3, axis='y')

    # 4. PS对比
    ax4 = axes[1, 1]
    peer_pss = list(peer_companies_data['ps'])
    peer_pss.append(industry_avg['ps'])
    peer_pss.append(current_metrics['ps'])

    bars = ax4.bar(range(len(peer_names)), peer_pss, color=colors_pe, alpha=0.7)
    ax4.axhline(y=industry_avg['ps'], color='red', linestyle='--', alpha=0.5)
    ax4.set_xlabel('公司', fontproperties=font_prop, fontsize=10)
    ax4.set_ylabel('PS (倍)', fontproperties=font_prop, fontsize=10)
    ax4.set_title('PS倍数对比', fontproperties=font_prop, fontsize=12, fontweight='bold')
    ax4.set_xticks(range(len(peer_names)))
    ax4.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=9, rotation=45)
    ax4.grid(True, alpha=0.3, axis='y')

    plt.suptitle('相对估值分析：与同行对比', fontproperties=font_prop, fontsize=16, fontweight='bold', y=0.995)
    plt.tight_layout()

    chart1_path = os.path.join(save_dir, 'relative_valuation_comparison.png')
    plt.savefig(chart1_path, dpi=150, bbox_inches='tight')
    plt.close()

    # 图2: 估值情景分析
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))

    # 计算情景数据
    net_income = 253532329.85  # 从配置获取
    net_assets = 4939639031.34
    total_shares = 767460689
    current_price = 23.88
    issue_price = 20.25

    eps = net_income / total_shares
    bps = net_assets / total_shares

    scenarios_data = []
    # 情景1: 当前估值
    scenarios_data.append(['当前估值', current_metrics['pe'], current_metrics['pb'], current_metrics['ps'], current_price, 0])
    # 情景2: PE回归到行业平均
    target_price_pe = eps * industry_avg['pe']
    return_pe = (target_price_pe - current_price) / current_price * 100
    scenarios_data.append(['PE→行业平均', industry_avg['pe'], current_metrics['pb'], current_metrics['ps'], target_price_pe, return_pe])
    # 情景3: PB回归到行业平均
    target_price_pb = bps * industry_avg['pb']
    return_pb = (target_price_pb - current_price) / current_price * 100
    scenarios_data.append(['PB→行业平均', current_metrics['pe'], industry_avg['pb'], current_metrics['ps'], target_price_pb, return_pb])
    # 情景4: 全面回归行业平均
    target_price_avg = (target_price_pe + target_price_pb) / 2
    return_avg = (target_price_avg - current_price) / current_price * 100
    scenarios_data.append(['全面回归', industry_avg['pe'], industry_avg['pb'], industry_avg['ps'], target_price_avg, return_avg])
    # 情景5: 行业最低估值
    min_pe = peer_companies_data['pe'].min()
    min_pb = peer_companies_data['pb'].min()
    target_price_pessimistic = (eps * min_pe + bps * min_pb) / 2
    return_pessimistic = (target_price_pessimistic - current_price) / current_price * 100
    scenarios_data.append(['行业最低估值', min_pe, min_pb, industry_avg['ps'], target_price_pessimistic, return_pessimistic])
    # 情景6: 行业最高估值
    max_pe = peer_companies_data['pe'].max()
    max_pb = peer_companies_data['pb'].max()
    target_price_optimistic = (eps * max_pe + bps * max_pb) / 2
    return_optimistic = (target_price_optimistic - current_price) / current_price * 100
    scenarios_data.append(['行业最高估值', max_pe, max_pb, industry_avg['ps'], target_price_optimistic, return_optimistic])

    df_scenarios = pd.DataFrame(scenarios_data, columns=['情景', 'PE', 'PB', 'PS', '目标价格(元)', '预期收益率(%)'])

    # 左图：目标价格对比
    colors_scenario = ['gray', 'blue', 'green', 'orange', 'red', 'green']
    bars1 = ax1.barh(df_scenarios['情景'], df_scenarios['目标价格(元)'], color=colors_scenario, alpha=0.7)
    ax1.axvline(x=current_price, color='black', linestyle='--', label='当前价格')
    ax1.set_xlabel('目标价格 (元)', fontproperties=font_prop, fontsize=12)
    ax1.set_title('不同估值情景下的目标价格', fontproperties=font_prop, fontsize=13, fontweight='bold')
    ax1.legend(prop=font_prop)
    ax1.grid(True, alpha=0.3, axis='x')

    # 右图：预期收益率对比
    colors_return = ['green' if r > 0 else 'red' for r in df_scenarios['预期收益率(%)']]
    bars2 = ax2.barh(df_scenarios['情景'], df_scenarios['预期收益率(%)'], color=colors_return, alpha=0.7)
    ax2.axvline(x=0, color='black', linestyle='-', linewidth=1)
    ax2.set_xlabel('预期收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax2.set_title('不同估值情景下的预期收益率', fontproperties=font_prop, fontsize=13, fontweight='bold')
    ax2.grid(True, alpha=0.3, axis='x')

    plt.tight_layout()

    chart2_path = os.path.join(save_dir, 'relative_valuation_scenarios.png')
    plt.savefig(chart2_path, dpi=150, bbox_inches='tight')
    plt.close()

    return chart1_path, chart2_path, df_scenarios


def generate_relative_valuation_charts_split(current_metrics, industry_avg, peer_companies_data, save_dir):
    """生成相对估值分析图表 - 拆分版本（单独大图）"""
    import os

    chart_paths = []

    # 图1: 估值指标对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    metrics = ['PE', 'PS', 'PB']
    current_vals = [current_metrics['pe'], current_metrics['ps'], current_metrics['pb']]
    industry_vals = [industry_avg['pe'], industry_avg['ps'], industry_avg['pb']]

    x = np.arange(len(metrics))
    width = 0.35

    bars1 = ax.bar(x - width/2, current_vals, width, label='光弘科技', color='#3498db', alpha=0.8)
    bars2 = ax.bar(x + width/2, industry_vals, width, label='行业平均', color='#e74c3c', alpha=0.8)

    ax.set_xlabel('估值指标', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('倍数', fontproperties=font_prop, fontsize=14)
    ax.set_title('相对估值对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(metrics, fontproperties=font_prop, fontsize=12)
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, axis='y')

    for bars in [bars1, bars2]:
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                    f'{height:.1f}', ha='center', va='bottom', fontsize=12, fontproperties=font_prop)

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_metrics_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图2: PE对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    peer_names = list(peer_companies_data['name'])
    peer_pes = list(peer_companies_data['pe'])
    peer_pes.append(industry_avg['pe'])
    peer_names.append('行业平均')
    peer_pes.append(current_metrics['pe'])
    peer_names.append('光弘科技')

    colors_pe = ['#95a5a6'] * len(peer_companies_data) + ['#e74c3c', '#3498db']
    bars = ax.bar(range(len(peer_names)), peer_pes, color=colors_pe, alpha=0.7)
    ax.axhline(y=industry_avg['pe'], color='red', linestyle='--', alpha=0.5)
    ax.set_xlabel('公司', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('PE (倍)', fontproperties=font_prop, fontsize=14)
    ax.set_title('PE倍数对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(peer_names)))
    ax.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=11, rotation=45)
    ax.grid(True, alpha=0.3, axis='y')

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_pe_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图3: PB对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    peer_pbs = list(peer_companies_data['pb'])
    peer_pbs.append(industry_avg['pb'])
    peer_pbs.append(current_metrics['pb'])

    bars = ax.bar(range(len(peer_names)), peer_pbs, color=colors_pe, alpha=0.7)
    ax.axhline(y=industry_avg['pb'], color='red', linestyle='--', alpha=0.5)
    ax.set_xlabel('公司', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('PB (倍)', fontproperties=font_prop, fontsize=14)
    ax.set_title('PB倍数对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(peer_names)))
    ax.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=11, rotation=45)
    ax.grid(True, alpha=0.3, axis='y')

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_pb_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图4: PS对比（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    peer_pss = list(peer_companies_data['ps'])
    peer_pss.append(industry_avg['ps'])
    peer_pss.append(current_metrics['ps'])

    bars = ax.bar(range(len(peer_names)), peer_pss, color=colors_pe, alpha=0.7)
    ax.axhline(y=industry_avg['ps'], color='red', linestyle='--', alpha=0.5)
    ax.set_xlabel('公司', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('PS (倍)', fontproperties=font_prop, fontsize=14)
    ax.set_title('PS倍数对比', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.set_xticks(range(len(peer_names)))
    ax.set_xticklabels(peer_names, fontproperties=font_prop, fontsize=11, rotation=45)
    ax.grid(True, alpha=0.3, axis='y')

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_ps_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图5: 目标价格情景（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    net_income = 253532329.85
    net_assets = 4939639031.34
    total_shares = 767460689
    current_price = 23.88

    eps = net_income / total_shares
    bps = net_assets / total_shares

    scenarios_data = []
    scenarios_data.append(['当前估值', current_metrics['pe'], current_metrics['pb'], current_metrics['ps'], current_price, 0])
    target_price_pe = eps * industry_avg['pe']
    return_pe = (target_price_pe - current_price) / current_price * 100
    scenarios_data.append(['PE→行业平均', industry_avg['pe'], current_metrics['pb'], current_metrics['ps'], target_price_pe, return_pe])
    target_price_pb = bps * industry_avg['pb']
    return_pb = (target_price_pb - current_price) / current_price * 100
    scenarios_data.append(['PB→行业平均', current_metrics['pe'], industry_avg['pb'], current_metrics['ps'], target_price_pb, return_pb])
    target_price_avg = (target_price_pe + target_price_pb) / 2
    return_avg = (target_price_avg - current_price) / current_price * 100
    scenarios_data.append(['全面回归', industry_avg['pe'], industry_avg['pb'], industry_avg['ps'], target_price_avg, return_avg])
    min_pe = peer_companies_data['pe'].min()
    min_pb = peer_companies_data['pb'].min()
    target_price_pessimistic = (eps * min_pe + bps * min_pb) / 2
    return_pessimistic = (target_price_pessimistic - current_price) / current_price * 100
    scenarios_data.append(['行业最低估值', min_pe, min_pb, industry_avg['ps'], target_price_pessimistic, return_pessimistic])
    max_pe = peer_companies_data['pe'].max()
    max_pb = peer_companies_data['pb'].max()
    target_price_optimistic = (eps * max_pe + bps * max_pb) / 2
    return_optimistic = (target_price_optimistic - current_price) / current_price * 100
    scenarios_data.append(['行业最高估值', max_pe, max_pb, industry_avg['ps'], target_price_optimistic, return_optimistic])

    df_scenarios = pd.DataFrame(scenarios_data, columns=['情景', 'PE', 'PB', 'PS', '目标价格(元)', '预期收益率(%)'])

    colors_scenario = ['gray', 'blue', 'green', 'orange', 'red', 'green']
    bars = ax.barh(df_scenarios['情景'], df_scenarios['目标价格(元)'], color=colors_scenario, alpha=0.7)
    ax.axvline(x=current_price, color='black', linestyle='--', label='当前价格')
    ax.set_xlabel('目标价格 (元)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('情景', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同估值情景下的目标价格', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, axis='x')
    # 设置y轴刻度位置和标签
    ax.set_yticks(range(len(df_scenarios)))
    ax.set_yticklabels(df_scenarios['情景'], fontproperties=font_prop, fontsize=11)

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_target_price_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    # 图6: 预期收益率情景（单独大图）
    fig, ax = plt.subplots(figsize=(14, 8))

    colors_return = ['green' if r > 0 else 'red' for r in df_scenarios['预期收益率(%)']]
    bars = ax.barh(df_scenarios['情景'], df_scenarios['预期收益率(%)'], color=colors_return, alpha=0.7)
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1)
    ax.set_xlabel('预期收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_ylabel('情景', fontproperties=font_prop, fontsize=14)
    ax.set_title('不同估值情景下的预期收益率', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.grid(True, alpha=0.3, axis='x')
    # 设置y轴刻度位置和标签
    ax.set_yticks(range(len(df_scenarios)))
    ax.set_yticklabels(df_scenarios['情景'], fontproperties=font_prop, fontsize=11)

    plt.tight_layout()

    chart_path = os.path.join(save_dir, 'relative_valuation_expected_return_split.png')
    plt.savefig(chart_path, dpi=150, bbox_inches='tight')
    plt.close()
    chart_paths.append(chart_path)

    return chart_paths, df_scenarios


def generate_stock_market_data_charts_split(market_data, price_data, volatility_data, save_dir):
    """
    生成个股市场数据图表（拆分版）

    参数:
        market_data: 市场数据字典
        price_data: 价格数据（包含close序列和均线）
        volatility_data: 波动率数据
        save_dir: 保存目录

    返回:
        图片路径列表 [股价走势图, 波动率图, 收益率分布图]
    """
    paths = []

    # 1. 股价走势与均线图
    fig, ax = plt.subplots(figsize=(14, 6))

    x_range = range(len(price_data['close']))
    ax.plot(x_range, price_data['close'], label='收盘价', linewidth=1.5, alpha=0.8, color='#2c3e50')

    colors = ['#e74c3c', '#e67e22', '#3498db', '#27ae60', '#9b59b6']
    windows = [20, 60, 120, 250]  # 与调用者提供的数据一致
    for i, window in enumerate(windows):
        ma_key = f'ma_{window}'
        if ma_key in price_data:
            ax.plot(x_range, price_data[ma_key], label=f'MA{window}', color=colors[i], linewidth=1.2, alpha=0.7)

    ax.set_xlabel('交易日', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('价格 (元)', fontproperties=font_prop, fontsize=12)
    ax.set_title(f'{market_data.get("stock_name", "N/A")} - 股价走势与均线',
                 fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, framealpha=0.9)
    ax.grid(True, alpha=0.3)

    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    path1 = os.path.join(save_dir, 'market_data_price_trend.png')
    plt.tight_layout()
    plt.savefig(path1, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path1)

    # 2. 滚动波动率图
    fig, ax = plt.subplots(figsize=(14, 5))

    colors = ['#e67e22', '#3498db', '#27ae60', '#9b59b6']
    windows = [20, 60, 120, 250]  # 与调用者提供的数据一致
    for i, window in enumerate(windows):
        vol_key = f'volatility_{window}d_series'
        if vol_key in volatility_data:
            ax.plot(x_range, volatility_data[vol_key] * 100,
                   label=f'{window}日波动率', color=colors[i], linewidth=1.2, alpha=0.7)

    ax.set_xlabel('交易日', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('滚动波动率走势', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, framealpha=0.9)
    ax.grid(True, alpha=0.3)

    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    path2 = os.path.join(save_dir, 'market_data_volatility.png')
    plt.tight_layout()
    plt.savefig(path2, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path2)

    # 3. 日收益率分布图
    fig, ax = plt.subplots(figsize=(12, 5))

    returns_pct = price_data.get('pct_chg', [])
    if len(returns_pct) > 0:
        # pct_chg 已经是百分比形式（如 2.5 表示 2.5%），不需要再乘以 100
        ax.hist(returns_pct, bins=50, color='#3498db', alpha=0.7, edgecolor='black')

        mean_return = np.mean(returns_pct)
        ax.axvline(x=mean_return, color='#e74c3c', linestyle='--',
                  label=f'平均收益率 {mean_return:.3f}%', linewidth=2)
        ax.axvline(x=0, color='#27ae60', linestyle='-', label='盈亏平衡', linewidth=1.5)

    ax.set_xlabel('日收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('频数', fontproperties=font_prop, fontsize=12)
    ax.set_title('日收益率分布', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, framealpha=0.9)
    ax.grid(True, alpha=0.3, axis='y')

    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(font_prop)

    path3 = os.path.join(save_dir, 'market_data_returns_dist.png')
    plt.tight_layout()
    plt.savefig(path3, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path3)

    return paths


def generate_industry_index_charts(industry_data, save_dir):
    """
    生成行业指数图表（走势、波动率、收益率分布）

    参数:
        industry_data: 行业数据字典
        save_dir: 保存目录

    返回:
        图片路径列表 [指数走势图, 波动率图, 收益率分布图]
    """
    import tushare as ts

    paths = []

    try:
        # 重新获取行业指数历史数据用于绘图
        index_code = industry_data.get('index_code')
        if not index_code:
            print("⚠️ 缺少行业指数代码，无法生成图表")
            return paths

        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if not ts_token:
            print("⚠️ 未设置TUSHARE_TOKEN，无法获取行业指数数据")
            return paths

        pro = ts.pro_api(ts_token)

        # 获取最近500天的数据
        end_date = datetime.now().strftime('%Y%m%d')
        start_date = (datetime.now() - timedelta(days=500*2)).strftime('%Y%m%d')

        df = pro.sw_daily(ts_code=index_code, start_date=start_date, end_date=end_date)

        if df is None or df.empty:
            print(f"⚠️ 未获取到{index_code}的历史数据")
            return paths

        # 按日期排序
        df = df.sort_values('trade_date')
        df.reset_index(drop=True, inplace=True)

        # 准备价格数据
        price_data = {'close': df['close'].values}
        pct_col = 'pct_chg' if 'pct_chg' in df.columns else 'pct_change'
        price_data['pct_chg'] = df[pct_col].values

        # 计算均线
        for window in [20, 60, 120, 250]:
            ma = df['close'].rolling(window=window).mean()
            price_data[f'ma_{window}'] = ma.values

        # 准备波动率数据
        volatility_data = {}
        for window in [20, 60, 120, 250]:
            pct_decimal = df[pct_col] / 100.0
            rolling_std = pct_decimal.rolling(window=window).std()
            rolling_vol = rolling_std * np.sqrt(250)
            volatility_data[f'volatility_{window}d_series'] = rolling_vol.values

        industry_name = industry_data.get('industry_name', industry_data.get('sw_l3_name', '行业指数'))

        # 1. 指数走势与均线图
        fig, ax = plt.subplots(figsize=(14, 6))

        x_range = range(len(price_data['close']))
        ax.plot(x_range, price_data['close'], label='收盘点位', linewidth=1.5, alpha=0.8, color='#2c3e50')

        colors = ['#e74c3c', '#e67e22', '#3498db', '#27ae60']
        windows = [20, 60, 120, 250]
        for i, window in enumerate(windows):
            ma_key = f'ma_{window}'
            if ma_key in price_data:
                ax.plot(x_range, price_data[ma_key], label=f'MA{window}', color=colors[i], linewidth=1.2, alpha=0.7)

        ax.set_xlabel('交易日', fontproperties=font_prop, fontsize=12)
        ax.set_ylabel('指数点位', fontproperties=font_prop, fontsize=12)
        ax.set_title(f'{industry_name} - 指数走势与均线',
                     fontproperties=font_prop, fontsize=14, fontweight='bold')
        ax.legend(prop=font_prop, framealpha=0.9)
        ax.grid(True, alpha=0.3)

        for label in ax.get_xticklabels() + ax.get_yticklabels():
            label.set_fontproperties(font_prop)

        path1 = os.path.join(save_dir, 'industry_index_trend.png')
        plt.tight_layout()
        plt.savefig(path1, dpi=150, bbox_inches='tight')
        plt.close()
        paths.append(path1)

        # 2. 滚动波动率图
        fig, ax = plt.subplots(figsize=(14, 5))

        colors = ['#e67e22', '#3498db', '#27ae60', '#9b59b6']
        windows = [20, 60, 120, 250]
        for i, window in enumerate(windows):
            vol_key = f'volatility_{window}d_series'
            if vol_key in volatility_data:
                ax.plot(x_range, volatility_data[vol_key] * 100,
                       label=f'{window}日波动率', color=colors[i], linewidth=1.2, alpha=0.7)

        ax.set_xlabel('交易日', fontproperties=font_prop, fontsize=12)
        ax.set_ylabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
        ax.set_title('行业指数滚动波动率走势', fontproperties=font_prop, fontsize=14, fontweight='bold')
        ax.legend(prop=font_prop, framealpha=0.9)
        ax.grid(True, alpha=0.3)

        for label in ax.get_xticklabels() + ax.get_yticklabels():
            label.set_fontproperties(font_prop)

        path2 = os.path.join(save_dir, 'industry_index_volatility.png')
        plt.tight_layout()
        plt.savefig(path2, dpi=150, bbox_inches='tight')
        plt.close()
        paths.append(path2)

        # 3. 日收益率分布图
        fig, ax = plt.subplots(figsize=(12, 5))

        returns_pct = price_data.get('pct_chg', [])
        if len(returns_pct) > 0:
            ax.hist(returns_pct, bins=50, color='#3498db', alpha=0.7, edgecolor='black')

            mean_return = np.mean(returns_pct)
            ax.axvline(x=mean_return, color='#e74c3c', linestyle='--',
                      label=f'平均收益率 {mean_return:.3f}%', linewidth=2)
            ax.axvline(x=0, color='#27ae60', linestyle='-', label='盈亏平衡', linewidth=1.5)

        ax.set_xlabel('日收益率 (%)', fontproperties=font_prop, fontsize=12)
        ax.set_ylabel('频数', fontproperties=font_prop, fontsize=12)
        ax.set_title('行业指数日收益率分布', fontproperties=font_prop, fontsize=14, fontweight='bold')
        ax.legend(prop=font_prop, framealpha=0.9)
        ax.grid(True, alpha=0.3, axis='y')

        for label in ax.get_xticklabels() + ax.get_yticklabels():
            label.set_fontproperties(font_prop)

        path3 = os.path.join(save_dir, 'industry_index_returns_dist.png')
        plt.tight_layout()
        plt.savefig(path3, dpi=150, bbox_inches='tight')
        plt.close()
        paths.append(path3)

        print(f"✅ 成功生成 {len(paths)} 个行业指数图表")

    except Exception as e:
        print(f"⚠️ 生成行业指数图表失败: {e}")
        import traceback
        traceback.print_exc()

    return paths


def generate_index_data_charts_split(indices_data, save_dir):
    """
    生成市场指数数据图表（拆分版）

    参数:
        indices_data: 指数数据字典
        save_dir: 保存目录

    返回:
        图片路径列表 [60日波动率, 60日收益率, 60日胜率, 60日技术位置,
                     120日波动率, 120日收益率, 250日波动率, 250日收益率]
    """
    paths = []

    index_names = list(indices_data.keys())

    # 提取60日数据
    volatilities_60d = [indices_data[name].get('volatility_60d', 0) * 100 for name in index_names]
    returns_60d = [indices_data[name].get('return_60d', 0) * 100 for name in index_names]
    win_rates_60d = [indices_data[name].get('win_rate_60d', 0) * 100 for name in index_names]
    ma_distances_60 = [(indices_data[name]['current_level'] /
                     indices_data[name].get('ma_60', 1) - 1) * 100 for name in index_names]

    # 提取120日数据
    volatilities_120d = [indices_data[name].get('volatility_120d', 0) * 100 for name in index_names]
    returns_120d = [indices_data[name].get('return_120d', 0) * 100 for name in index_names]
    win_rates_120d = [indices_data[name].get('win_rate_120d', 0) * 100 for name in index_names]

    # 提取250日数据
    volatilities_250d = [indices_data[name].get('volatility_250d', 0) * 100 for name in index_names]
    returns_250d = [indices_data[name].get('return_250d', 0) * 100 for name in index_names]
    win_rates_250d = [indices_data[name].get('win_rate_250d', 0) * 100 for name in index_names]

    # ====== 60日窗口图表 ======
    # 1. 波动率对比图（60日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors1 = plt.cm.Reds(np.linspace(0.4, 0.9, len(index_names)))
    bars1 = ax.barh(index_names, volatilities_60d, color=colors1, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数波动率对比 (60日窗口)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, vol) in enumerate(zip(bars1, volatilities_60d)):
        ax.text(vol + 0.5, i, f'{vol:.1f}%', va='center', fontsize=10,
               fontproperties=font_prop, fontweight='bold')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path1 = os.path.join(save_dir, 'index_volatility_60d.png')
    plt.tight_layout()
    plt.savefig(path1, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path1)

    # 2. 年化收益率对比图（60日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors2 = ['#27ae60' if r > 0 else '#e74c3c' for r in returns_60d]
    bars2 = ax.barh(index_names, returns_60d, color=colors2, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数年化收益率对比 (60日窗口)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1.2)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, ret) in enumerate(zip(bars2, returns_60d)):
        pos = ret + 1.5 if ret > 0 else ret - 2
        ha = 'left' if ret > 0 else 'right'
        ax.text(pos, i, f'{ret:+.1f}%', va='center', ha=ha, fontsize=10,
               fontproperties=font_prop, fontweight='bold', color='white')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path2 = os.path.join(save_dir, 'index_returns_60d.png')
    plt.tight_layout()
    plt.savefig(path2, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path2)

    # 3. 胜率对比图（60日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors3 = plt.cm.Blues(np.linspace(0.4, 0.9, len(index_names)))
    bars3 = ax.barh(index_names, win_rates_60d, color=colors3, alpha=0.85, edgecolor='white')

    ax.set_xlabel('胜率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数胜率对比 (60日窗口)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.set_xlim(0, 100)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, wr) in enumerate(zip(bars3, win_rates_60d)):
        ax.text(wr + 1, i, f'{wr:.1f}%', va='center', fontsize=10,
               fontproperties=font_prop, fontweight='bold')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path3 = os.path.join(save_dir, 'index_winrate_60d.png')
    plt.tight_layout()
    plt.savefig(path3, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path3)

    # 4. 技术位置对比图（60日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors4 = ['#27ae60' if d > 0 else '#e74c3c' for d in ma_distances_60]
    bars4 = ax.barh(index_names, ma_distances_60, color=colors4, alpha=0.85, edgecolor='white')

    ax.set_xlabel('相对MA60的偏离 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数技术位置对比 (60日窗口)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1.2)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, dist) in enumerate(zip(bars4, ma_distances_60)):
        pos = dist + 0.8 if dist > 0 else dist - 0.8
        ha = 'left' if dist > 0 else 'right'
        ax.text(pos, i, f'{dist:+.1f}%', va='center', ha=ha, fontsize=10,
               fontproperties=font_prop, fontweight='bold', color='white')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path4 = os.path.join(save_dir, 'index_technical_60d.png')
    plt.tight_layout()
    plt.savefig(path4, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path4)

    # ====== 120日窗口图表 ======
    # 5. 波动率对比图（120日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors5 = plt.cm.Reds(np.linspace(0.4, 0.9, len(index_names)))
    bars5 = ax.barh(index_names, volatilities_120d, color=colors5, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化波动率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数波动率对比 (120日窗口/半年线)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, vol) in enumerate(zip(bars5, volatilities_120d)):
        ax.text(vol + 0.5, i, f'{vol:.1f}%', va='center', fontsize=10,
               fontproperties=font_prop, fontweight='bold')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path5 = os.path.join(save_dir, 'index_volatility_120d.png')
    plt.tight_layout()
    plt.savefig(path5, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path5)

    # 6. 年化收益率对比图（120日）
    fig, ax = plt.subplots(figsize=(12, 6))
    colors6 = ['#27ae60' if r > 0 else '#e74c3c' for r in returns_120d]
    bars6 = ax.barh(index_names, returns_120d, color=colors6, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化收益率 (%)', fontproperties=font_prop, fontsize=12)
    ax.set_title('各指数年化收益率对比 (120日窗口/半年线)', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1.2)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, ret) in enumerate(zip(bars6, returns_120d)):
        pos = ret + 1.5 if ret > 0 else ret - 2
        ha = 'left' if ret > 0 else 'right'
        ax.text(pos, i, f'{ret:+.1f}%', va='center', ha=ha, fontsize=10,
               fontproperties=font_prop, fontweight='bold', color='white')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)

    path6 = os.path.join(save_dir, 'index_returns_120d.png')
    plt.tight_layout()
    plt.savefig(path6, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path6)

    # ====== 250日窗口图表 ======
    # 7. 波动率对比图（250日）
    fig, ax = plt.subplots(figsize=(20, 10))
    colors7 = plt.cm.Reds(np.linspace(0.4, 0.9, len(index_names)))
    bars7 = ax.barh(index_names, volatilities_250d, color=colors7, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化波动率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('各指数波动率对比 (250日窗口/年线)', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, vol) in enumerate(zip(bars7, volatilities_250d)):
        ax.text(vol + 0.5, i, f'{vol:.1f}%', va='center', fontsize=12,
               fontproperties=font_prop, fontweight='bold')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(13)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(13)

    path7 = os.path.join(save_dir, 'index_volatility_250d.png')
    plt.tight_layout()
    plt.savefig(path7, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path7)

    # 8. 年化收益率对比图（250日）
    fig, ax = plt.subplots(figsize=(20, 10))
    colors8 = ['#27ae60' if r > 0 else '#e74c3c' for r in returns_250d]
    bars8 = ax.barh(index_names, returns_250d, color=colors8, alpha=0.85, edgecolor='white')

    ax.set_xlabel('年化收益率 (%)', fontproperties=font_prop, fontsize=14)
    ax.set_title('各指数年化收益率对比 (250日窗口/年线)', fontproperties=font_prop, fontsize=16, fontweight='bold')
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1.2)
    ax.grid(True, alpha=0.3, axis='x')

    for i, (bar, ret) in enumerate(zip(bars8, returns_250d)):
        pos = ret + 1.5 if ret > 0 else ret - 2
        ha = 'left' if ret > 0 else 'right'
        ax.text(pos, i, f'{ret:+.1f}%', va='center', ha=ha, fontsize=12,
               fontproperties=font_prop, fontweight='bold', color='white')

    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(13)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(13)

    path8 = os.path.join(save_dir, 'index_returns_250d.png')
    plt.tight_layout()
    plt.savefig(path8, dpi=150, bbox_inches='tight')
    plt.close()
    paths.append(path8)

    return paths


def generate_radar_chart(scores, save_path):
    """生成风险评分雷达图"""
    categories = list(scores.keys())
    values = list(scores.values())

    # 闭合雷达图
    angles_plot = np.linspace(0, 2*np.pi, len(categories), endpoint=False).tolist()
    angles_plot += angles_plot[:1]
    values_plot = values + values[:1]

    fig, ax = plt.subplots(figsize=(10, 10), subplot_kw=dict(projection='polar'))
    ax.plot(angles_plot, values_plot, 'o-', linewidth=2, color='#3498db')
    ax.fill(angles_plot, values_plot, alpha=0.25, color='#3498db')
    ax.set_xticks(angles_plot[:-1])
    ax.set_xticklabels(categories, fontproperties=font_prop, fontsize=11)
    ax.set_ylim(0, 30)
    ax.set_title('定增项目风险雷达图', fontsize=14, fontweight='bold',
                pad=20, fontproperties=font_prop)
    ax.grid(True)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()


def get_key_strengths(scores, profit_prob, discount_premium, var_95):
    """提取项目关键优势"""
    strengths = []

    if scores['盈利概率'] == 30:
        strengths.append(f"盈利概率高({profit_prob*100:.1f}%)")

    if scores['发行折价'] == 20:
        strengths.append(f"深度折价发行({discount_premium:+.1f}%)")

    if scores['锁定期'] == 20:
        strengths.append("锁定期短(≤6个月)")

    if scores['波动率'] == 15:
        strengths.append("波动率低(≤30%)")

    if scores['VaR'] == 15:
        strengths.append(f"下行风险可控(VaR={var_95*100:.1f}%)")

    if not strengths:
        strengths.append("各维度表现均衡")

    return "、".join(strengths)


def get_key_risks(scores, profit_prob, discount_premium, var_95):
    """提取主要风险点"""
    risks = []

    if scores['盈利概率'] == 10:
        risks.append(f"盈利概率低({profit_prob*100:.1f}%)")

    if scores['发行折价'] == 5:
        risks.append(f"折价不足({discount_premium:+.1f}%)")

    if scores['锁定期'] == 5:
        risks.append("锁定期长(>12个月)")

    if scores['波动率'] == 5:
        risks.append("波动率高(>40%)")

    if scores['VaR'] == 3:
        risks.append(f"下行风险大(VaR={var_95*100:.1f}%)")

    if not risks:
        risks.append("无明显风险点")

    return "、".join(risks)


def check_data_freshness(market_data_file, max_days_old=3):
    """
    检查市场数据是否过期

    参数:
        market_data_file: 市场数据文件路径
        max_days_old: 数据最大允许天数（默认3天）

    返回:
        (is_fresh, data_date, days_old)
        is_fresh: 数据是否新鲜
        data_date: 数据日期（datetime对象）
        days_old: 数据天数差
    """
    if not os.path.exists(market_data_file):
        return False, None, None

    try:
        with open(market_data_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        analysis_date_str = data.get('analysis_date', '')
        if not analysis_date_str:
            return False, None, None

        # 解析日期字符串（格式：YYYYMMDD）
        data_date = datetime.strptime(analysis_date_str, '%Y%m%d')
        today = datetime.now()
        days_old = (today - data_date).days

        is_fresh = days_old <= max_days_old

        return is_fresh, data_date, days_old

    except Exception as e:
        print(f"⚠️ 检查数据日期失败: {e}")
        return False, None, None


def update_market_data(stock_code='300735.SZ'):
    """
    更新市场数据

    参数:
        stock_code: 股票代码

    返回:
        bool: 是否更新成功
    """
    print(f"\n📡 正在更新 {stock_code} 的市场数据...")

    update_script = os.path.join(SCRIPT_DIR, 'update_market_data.py')

    if not os.path.exists(update_script):
        print(f"❌ 未找到更新脚本: {update_script}")
        return False

    try:
        # 运行更新脚本
        result = subprocess.run(
            [sys.executable, update_script, '--stock', stock_code],
            cwd=SCRIPT_DIR,
            capture_output=True,
            text=True,
            timeout=60  # 60秒超时
        )

        if result.returncode == 0:
            print("✅ 市场数据更新成功")
            if result.stdout:
                print(result.stdout)
            return True
        else:
            print(f"❌ 市场数据更新失败")
            if result.stderr:
                print(result.stderr)
            return False

    except subprocess.TimeoutExpired:
        print("❌ 数据更新超时（>60秒）")
        return False
    except Exception as e:
        print(f"❌ 数据更新异常: {e}")
        return False



def generate_report(stock_code='300735.SZ', output_file='定增风险分析报告.docx'):
    """生成完整的定增风险分析报告"""

    print("开始生成定增风险分析报告（包含图表）...")

    # 创建文档
    document = Document()
    setup_chinese_font(document)

    # 加载配置 - 使用数据目录中的文件
    print("加载配置数据...")
    project_params, risk_params, market_data = load_placement_config(stock_code, data_dir=DATA_DIR)

    # 计算发行数量（根据融资金额和发行价）
    issue_shares = int(project_params['financing_amount'] / project_params['issue_price'])
    project_params['issue_shares'] = issue_shares
    print(f"计算得出发行数量: {issue_shares:,} 股（投资金额: {project_params['financing_amount']/100000000:.2f}亿元 ÷ 发行价: {project_params['issue_price']:.2f}元/股）")
    print(f"注：使用固定投资金额1亿元进行风险评估（与实际投资规模无关）")

    # 读取原始placement_params以获取完整财务数据
    placement_file = os.path.join(DATA_DIR, f"{stock_code.replace('.', '_')}_placement_params.json")
    if os.path.exists(placement_file):
        with open(placement_file, 'r', encoding='utf-8') as f:
            placement_params_raw = json.load(f)
            # 将财务数据合并到project_params中
            project_params.update({
                'net_assets': placement_params_raw.get('net_assets', 0),
                'total_debt': placement_params_raw.get('total_debt', 0),
                'net_income': placement_params_raw.get('net_income', 0),
                'revenue_growth': placement_params_raw.get('revenue_growth', 0.25),
                'operating_margin': placement_params_raw.get('operating_margin', 0.2),
                'beta': placement_params_raw.get('beta', 1.2)
            })

    # 发行类型评估
    ma20 = market_data.get('ma_20', 0)
    ma120 = market_data.get('ma_120', 0)
    issue_price = project_params['issue_price']
    # 计算折价/溢价率（负值表示折价，正值表示溢价）
    discount_premium = (issue_price - ma20) / ma20 * 100
    if issue_price < ma20:
        issue_type = "折价发行"
    else:
        issue_type = "溢价发行"

    # 创建分析器
    analyzer = PrivatePlacementRiskAnalyzer(
        issue_price=project_params['issue_price'],
        issue_shares=project_params['issue_shares'],
        lockup_period=project_params['lockup_period'],
        current_price=project_params['current_price'],
        risk_free_rate=project_params['risk_free_rate']
    )

    current_date = datetime.now().strftime('%Y年%m月%d日')

    # ==================== 封面 ====================
    add_title(document, '定增项目风险分析报告', level=0)
    document.add_paragraph()
    document.add_paragraph()

    # 报告基本信息
    info_table = document.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'

    info_data = [
        ['公司名称', '光弘科技 (300735.SZ)'],
        ['报告日期', current_date],
        ['发行日期', current_date],
        ['发行价格', f'{project_params["issue_price"]:.2f} 元/股'],
        ['当前价格', f'{project_params["current_price"]:.2f} 元/股'],
        ['发行类型', f'{issue_type}（安全边际/溢价率: {discount_premium:.2f}%）']
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        for cell in info_table.rows[i].cells:
            cell.paragraphs[0].runs[0].font.size = Pt(12)  # 小四号（12pt）
            cell.paragraphs[0].runs[0].font.name = '宋体'
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_section_break(document)

    # ==================== 目录 ====================
    add_title(document, '目 录', level=1)
    add_paragraph(document, '一、项目概况')
    add_paragraph(document, '二、相对估值分析')
    add_paragraph(document, '三、DCF估值分析')
    add_paragraph(document, '四、敏感性分析')
    add_paragraph(document, '五、蒙特卡洛模拟')
    add_paragraph(document, '六、情景分析')
    add_paragraph(document, '七、压力测试')
    add_paragraph(document, '八、VaR风险度量')
    add_paragraph(document, '九、综合评估')
    add_paragraph(document, '十、投资建议与风险提示')

    add_section_break(document)

    # ==================== 一、项目概况 ====================
    add_title(document, '一、项目概况', level=1)

    add_title(document, '1.1 基本信息', level=2)

    project_headers = ['指标', '数值']
    project_data = [
        ['股票代码', '300735.SZ'],
        ['公司名称', '光弘科技'],
        ['发行价格', f'{project_params["issue_price"]:.2f} 元/股'],
        ['发行数量', f'{project_params["issue_shares"]:,} 股'],
        ['锁定期', f'{project_params["lockup_period"]} 个月'],
        ['投资金额（固定）', f'{project_params["financing_amount"] / 100000000:.2f} 亿元'],
        ['当前价格', f'{project_params["current_price"]:.2f} 元/股'],
        ['MA20价格', f'{ma20:.2f} 元/股'],
        ['发行类型', issue_type],
        ['安全边际/溢价率', f'{discount_premium:.2f}%']
    ]
    add_table_data(document, project_headers, project_data)

    add_title(document, '1.2 个股数据分析', level=2)

    market_headers = ['指标', '数值']
    market_table_data = [
        ['当前价格', f'{market_data["current_price"]:.2f} 元/股（截至 {market_data.get("analysis_date", "")}）'],
        ['平均价格', f'{market_data.get("avg_price_all", 0):.2f} 元/股（{market_data.get("total_days", 0)}个交易日）'],
        ['价格标准差', f'{market_data.get("price_std", 0):.2f}'],
        ['月度波动率(20日)', f'{market_data.get("volatility_20d", 0)*100:.2f}%'],
        ['季度波动率(60日)', f'{market_data.get("volatility_60d", 0)*100:.2f}%'],
        ['半年波动率(120日)', f'{market_data.get("volatility_120d", 0)*100:.2f}%'],
        ['年度波动率(250日)', f'{market_data.get("volatility_250d", 0)*100:.2f}%'],
        ['月度区间收益率(20日)', f'{market_data.get("period_return_20d", 0)*100:+.2f}%'],
        ['月度年化收益率(20日)', f'{market_data.get("annual_return_20d", 0)*100:+.2f}%'],
        ['季度区间收益率(60日)', f'{market_data.get("period_return_60d", 0)*100:+.2f}%'],
        ['季度年化收益率(60日)', f'{market_data.get("annual_return_60d", 0)*100:+.2f}%'],
        ['半年区间收益率(120日)', f'{market_data.get("period_return_120d", 0)*100:+.2f}%'],
        ['半年年化收益率(120日)', f'{market_data.get("annual_return_120d", 0)*100:+.2f}%'],
        ['年度区间收益率(250日)', f'{market_data.get("period_return_250d", 0)*100:+.2f}%'],
        ['年度年化收益率(250日)', f'{market_data.get("annual_return_250d", 0)*100:+.2f}%'],
        ['MA20', f'{market_data.get("ma_20", 0):.2f} 元/股'],
        ['MA60', f'{market_data.get("ma_60", 0):.2f} 元/股'],
        ['MA120', f'{market_data.get("ma_120", 0):.2f} 元/股'],
        ['MA250', f'{market_data.get("ma_250", 0):.2f} 元/股'],
        ['数据来源', 'Tushare API']
    ]
    add_table_data(document, market_headers, market_table_data)

    # 1.2.1 生成个股市场数据图表
    print("生成个股市场数据图表...")
    stock_market_data_file = os.path.join(DATA_DIR, f"{stock_code.replace('.', '_')}_market_data.json")

    # 尝试从Tushare获取历史数据生成图表
    stock_charts_paths = []
    try:
        import tushare as ts

        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if ts_token:
            pro = ts.pro_api(ts_token)

            # 获取历史价格数据
            end_date = datetime.now().strftime('%Y%m%d')
            start_date = (datetime.now() - timedelta(days=500)).strftime('%Y%m%d')

            df_daily = pro.daily(ts_code=stock_code, start_date=start_date, end_date=end_date)
            if not df_daily.empty:
                df_daily = df_daily.sort_values('trade_date').reset_index(drop=True)

                # 准备价格数据
                price_data = {'close': df_daily['close'].values, 'pct_chg': df_daily['pct_chg'].values}
                for window in [20, 60, 120, 250]:
                    ma = df_daily['close'].rolling(window=window).mean()
                    price_data[f'ma_{window}'] = ma.values

                # 准备波动率数据
                volatility_data = {}
                for window in [20, 60, 120, 250]:
                    pct_decimal = df_daily['pct_chg'] / 100.0
                    rolling_std = pct_decimal.rolling(window=window).std()
                    rolling_vol = rolling_std * np.sqrt(250)
                    volatility_data[f'volatility_{window}d_series'] = rolling_vol.values

                # 生成图表
                stock_charts_paths = generate_stock_market_data_charts_split(
                    market_data, price_data, volatility_data, IMAGES_DIR
                )

                # 添加图表到文档
                add_paragraph(document, '')
                add_paragraph(document, '图表 1.1: 股价走势与均线分析')
                if len(stock_charts_paths) > 0 and os.path.exists(stock_charts_paths[0]):
                    add_image(document, stock_charts_paths[0], width=Inches(6.5))
                    add_paragraph(document, '')

                add_paragraph(document, '图表 1.2: 滚动波动率分析')
                if len(stock_charts_paths) > 1 and os.path.exists(stock_charts_paths[1]):
                    add_image(document, stock_charts_paths[1], width=Inches(6.5))
                    add_paragraph(document, '')

                add_paragraph(document, '图表 1.3: 日收益率分布')
                if len(stock_charts_paths) > 2 and os.path.exists(stock_charts_paths[2]):
                    add_image(document, stock_charts_paths[2], width=Inches(6.5))
                    add_paragraph(document, '')

                add_paragraph(document, '分析结论：')
                add_paragraph(document, f'• 当前股价{market_data["current_price"]:.2f}元，{issue_type}于MA20({ma20:.2f}元)')
                add_paragraph(document, f'• 60日年化波动率为{market_data.get("volatility_60d", 0)*100:.2f}%，{"高于" if market_data.get("volatility_60d", 0) > 0.3 else "低于"}市场平均水平')
                add_paragraph(document, f'• 60日年化收益率为{market_data.get("annual_return_60d", 0)*100:.2f}%，{"表现良好" if market_data.get("annual_return_60d", 0) > 0 else "表现不佳"}')
    except Exception as e:
        print(f"⚠️ 生成个股市场数据图表失败: {e}")

    # 1.3 市场指数分析
    add_paragraph(document, '')
    add_title(document, '1.3 市场指数分析', level=2)

    add_paragraph(document, '本章节分析主要市场指数的表现，包括波动率、收益率、胜率等指标，为项目风险评估提供市场环境参考。')

    # 加载指数数据（从data目录）
    # 优先使用新方法计算的指数数据（使用修复后的年化收益率公式）
    indices_data_file_v2 = os.path.join(DATA_DIR, 'market_indices_scenario_data_v2.json')
    indices_data_file = os.path.join(DATA_DIR, 'market_indices_scenario_data.json')
    indices_charts_paths = []

    # 优先使用新数据
    if os.path.exists(indices_data_file_v2):
        print("✅ 使用新方法计算的指数数据（修复后的年化收益率公式）")
        indices_data_file = indices_data_file_v2

    if os.path.exists(indices_data_file):
        print("加载市场指数数据...")
        try:
            with open(indices_data_file, 'r', encoding='utf-8') as f:
                indices_data = json.load(f)

            # 生成指数对比图表
            indices_charts_paths = generate_index_data_charts_split(indices_data, IMAGES_DIR)

            # 添加指数对比表格 - 60日指标
            add_paragraph(document, '')
            add_paragraph(document, '主要市场指数60日指标对比：')

            index_headers_60d = ['指数', '当前点位', '波动率', '区间收益率', '年化收益率', '胜率']
            index_table_data_60d = []
            for name, data in indices_data.items():
                # 从年化收益率反推区间收益率（使用简单年化方法）
                # 年化收益率 = 区间收益率 × (250/60)
                # 反推：区间收益率 = 年化收益率 × (60/250)
                annual_ret = data.get('return_60d', 0)
                period_ret_60d = annual_ret * (60.0 / 250.0)

                index_table_data_60d.append([
                    name,
                    f"{data.get('current_level', 0):.2f}",
                    f"{data.get('volatility_60d', 0)*100:.2f}%",
                    f"{period_ret_60d*100:+.2f}%",
                    f"{data.get('return_60d', 0)*100:+.2f}%",
                    f"{data.get('win_rate_60d', 0)*100:.1f}%"
                ])
            add_table_data(document, index_headers_60d, index_table_data_60d)

            # 添加指数对比表格 - 120日指标
            add_paragraph(document, '')
            add_paragraph(document, '主要市场指数120日指标对比（半年线）：')

            index_headers_120d = ['指数', '当前点位', '波动率', '区间收益率', '年化收益率', '胜率']
            index_table_data_120d = []
            for name, data in indices_data.items():
                # 从年化收益率反推区间收益率（使用简单年化方法）
                annual_ret = data.get('return_120d', 0)
                period_ret_120d = annual_ret * (120.0 / 250.0)

                # 如果没有win_rate_120d，使用win_rate_60d作为近似值
                win_rate_120d = data.get('win_rate_120d', 0)

                index_table_data_120d.append([
                    name,
                    f"{data.get('current_level', 0):.2f}",
                    f"{data.get('volatility_120d', 0)*100:.2f}%",
                    f"{period_ret_120d*100:+.2f}%",
                    f"{data.get('return_120d', 0)*100:+.2f}%",
                    f"{win_rate_120d*100:.1f}%"
                ])
            add_table_data(document, index_headers_120d, index_table_data_120d)

            # 添加指数对比表格 - 250日指标
            add_paragraph(document, '')
            add_paragraph(document, '主要市场指数250日指标对比（年线）：')

            index_headers_250d = ['指数', '当前点位', '波动率', '区间收益率', '年化收益率', '胜率']
            index_table_data_250d = []
            for name, data in indices_data.items():
                # 对于250日窗口，年化收益率≈区间收益率（250个交易日≈1年）
                annual_ret = data.get('return_250d', 0)
                period_ret = annual_ret  # 250日≈1年，区间收益率≈年化收益率

                # 如果没有win_rate_250d，使用win_rate_60d作为近似值
                win_rate_250d = data.get('win_rate_250d', 0)

                index_table_data_250d.append([
                    name,
                    f"{data.get('current_level', 0):.2f}",
                    f"{data.get('volatility_250d', 0)*100:.2f}%",
                    f"{period_ret*100:+.2f}%",
                    f"{annual_ret*100:+.2f}%",
                    f"{win_rate_250d*100:.1f}%"
                ])
            add_table_data(document, index_headers_250d, index_table_data_250d)

            add_paragraph(document, '')
            add_paragraph(document, '💡 说明：')
            add_paragraph(document, '• 250日（年线）包含约250个交易日，接近一年的交易日数量（约252天）')
            add_paragraph(document, '• 年化收益率≈区间收益率（因250日≈1年，无需额外年化处理）')
            add_paragraph(document, '• 胜率基于实际历史数据计算，反映各时间窗口的上涨天数占比')
            add_paragraph(document, '• 建议以沪深300为基准参考，其年化收益率反映了市场整体表现')
            add_paragraph(document, '• 个股投资应结合行业特征和市场环境综合判断')

            # 添加图表 - 60日窗口
            add_paragraph(document, '')
            add_paragraph(document, '图表 1.4: 各指数波动率对比 (60日窗口)')
            if len(indices_charts_paths) > 0 and os.path.exists(indices_charts_paths[0]):
                add_image(document, indices_charts_paths[0], width=Inches(6))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.5: 各指数年化收益率对比 (60日窗口)')
            if len(indices_charts_paths) > 1 and os.path.exists(indices_charts_paths[1]):
                add_image(document, indices_charts_paths[1], width=Inches(6))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.6: 各指数胜率对比 (60日窗口)')
            if len(indices_charts_paths) > 2 and os.path.exists(indices_charts_paths[2]):
                add_image(document, indices_charts_paths[2], width=Inches(6))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.7: 各指数技术位置对比 (60日窗口)')
            if len(indices_charts_paths) > 3 and os.path.exists(indices_charts_paths[3]):
                add_image(document, indices_charts_paths[3], width=Inches(6))
                add_paragraph(document, '')

            # 添加图表 - 120日窗口（半年线）
            add_paragraph(document, '')
            add_paragraph(document, '图表 1.8: 各指数波动率对比 (120日窗口/半年线)')
            if len(indices_charts_paths) > 4 and os.path.exists(indices_charts_paths[4]):
                add_image(document, indices_charts_paths[4], width=Inches(6))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.9: 各指数年化收益率对比 (120日窗口/半年线)')
            if len(indices_charts_paths) > 5 and os.path.exists(indices_charts_paths[5]):
                add_image(document, indices_charts_paths[5], width=Inches(6))
                add_paragraph(document, '')

            # 添加图表 - 250日窗口
            add_paragraph(document, '')
            add_paragraph(document, '图表 1.10: 各指数波动率对比 (250日年线窗口)')
            if len(indices_charts_paths) > 6 and os.path.exists(indices_charts_paths[6]):
                add_image(document, indices_charts_paths[6], width=Inches(6.3))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.11: 各指数年化收益率对比 (250日年线窗口)')
            if len(indices_charts_paths) > 7 and os.path.exists(indices_charts_paths[7]):
                add_image(document, indices_charts_paths[7], width=Inches(6.3))
                add_paragraph(document, '')

            # 添加分析结论
            add_paragraph(document, '分析结论：')
            add_paragraph(document, '• 沪深300作为市场基准，波动率相对稳定，适合作为风险参照')
            add_paragraph(document, '• 创业板指和科创50波动率较高，体现成长股的高风险高收益特征')
            add_paragraph(document, '• 60日（短趋势）反映市场近期情绪，120日（半年线）反映中期趋势，250日（年线）反映长期趋势，适合作为战略参考')
            add_paragraph(document, '• 当前市场环境下，建议关注指数技术位置对个股表现的影响')

        except Exception as e:
            print(f"⚠️ 加载指数数据失败: {e}")
            add_paragraph(document, '⚠️ 指数数据暂未加载，请先运行07_market_data_analysis.ipynb生成数据')
    else:
        add_paragraph(document, '⚠️ 指数数据文件不存在，请先运行07_market_data_analysis.ipynb生成数据')

    # 1.4 行业数据分析
    add_paragraph(document, '')
    add_title(document, '1.4 行业数据分析', level=2)

    add_paragraph(document, '本章节分析所属行业的申万三级行业指数表现，为项目风险评估提供行业环境参考。')

    # 尝试加载行业数据
    try:
        industry_data_file = os.path.join(DATA_DIR, f'{stock_code.replace(".", "_")}_industry_data.json')

        if os.path.exists(industry_data_file):
            with open(industry_data_file, 'r', encoding='utf-8') as f:
                industry_data = json.load(f)

            print(f"✅ 已加载行业数据: {industry_data_file}")

            # 1.4.1 行业基本信息
            add_paragraph(document, '')
            add_title(document, '1.4.1 行业基本信息', level=3)

            industry_headers = ['指标', '数值']
            industry_table_data = [
                ['申万一级行业', f"{industry_data.get('sw_l1_name', 'N/A')}"],
                ['申万二级行业', f"{industry_data.get('sw_l2_name', 'N/A')}"],
                ['申万三级行业', f"{industry_data.get('sw_l3_name', 'N/A')}"],
                ['行业指数代码', f"{industry_data.get('index_code', 'N/A')}"],
                ['当前点位', f"{industry_data.get('current_level', 0):.2f}"],
                ['分析日期', f"{industry_data.get('analysis_date', 'N/A')}"],
            ]
            add_table_data(document, industry_headers, industry_table_data)

            # 1.4.2 行业风险指标分析
            add_paragraph(document, '')
            add_title(document, '1.4.2 行业风险指标分析', level=3)

            # 波动率
            add_paragraph(document, '')
            add_paragraph(document, '📊 行业波动率（风险水平）：')
            add_paragraph(document, f"• 月度(20日): {industry_data.get('volatility_20d', 0)*100:.2f}%")
            add_paragraph(document, f"• 季度(60日): {industry_data.get('volatility_60d', 0)*100:.2f}%")
            add_paragraph(document, f"• 半年(120日): {industry_data.get('volatility_120d', 0)*100:.2f}%")
            add_paragraph(document, f"• 年度(250日): {industry_data.get('volatility_250d', 0)*100:.2f}%")

            # 收益率
            add_paragraph(document, '')
            add_paragraph(document, '📈 行业收益率表现：')
            add_paragraph(document, f"• 月度区间收益率(20日): {industry_data.get('period_return_20d', 0)*100:+.2f}%")
            add_paragraph(document, f"• 月度年化收益率(20日): {industry_data.get('annual_return_20d', 0)*100:+.2f}%")
            add_paragraph(document, f"• 季度区间收益率(60日): {industry_data.get('period_return_60d', 0)*100:+.2f}%")
            add_paragraph(document, f"• 季度年化收益率(60日): {industry_data.get('annual_return_60d', 0)*100:+.2f}%")
            add_paragraph(document, f"• 半年区间收益率(120日): {industry_data.get('period_return_120d', 0)*100:+.2f}%")
            add_paragraph(document, f"• 半年年化收益率(120日): {industry_data.get('annual_return_120d', 0)*100:+.2f}%")
            add_paragraph(document, f"• 年度区间收益率(250日): {industry_data.get('period_return_250d', 0)*100:+.2f}%")
            add_paragraph(document, f"• 年度年化收益率(250日): {industry_data.get('annual_return_250d', 0)*100:+.2f}%")

            # 1.4.3 行业技术位置分析
            add_paragraph(document, '')
            add_title(document, '1.4.3 行业技术位置分析', level=3)

            add_paragraph(document, '')
            add_paragraph(document, '📊 行业移动平均线（技术位置）：')
            add_paragraph(document, f"• MA20: {industry_data.get('ma_20', 0):.2f} 点")
            add_paragraph(document, f"• MA60: {industry_data.get('ma_60', 0):.2f} 点")
            add_paragraph(document, f"• MA120: {industry_data.get('ma_120', 0):.2f} 点")
            add_paragraph(document, f"• MA250: {industry_data.get('ma_250', 0):.2f} 点")

            # 胜率
            add_paragraph(document, '')
            add_paragraph(document, '🎯 行业胜率（上涨天数占比）：')

            # 获取胜率数据，如果缺少120日和250日，使用60日胜率作为近似值
            win_rate_20d = industry_data.get('win_rate_20d', 0)
            win_rate_60d = industry_data.get('win_rate_60d', 0)
            win_rate_120d = industry_data.get('win_rate_120d', 0)
            win_rate_250d = industry_data.get('win_rate_250d', 0)

            add_paragraph(document, f"• 月度(20日): {win_rate_20d*100:.1f}%")
            add_paragraph(document, f"• 季度(60日): {win_rate_60d*100:.1f}%")
            add_paragraph(document, f"• 半年(120日): {win_rate_120d*100:.1f}%")
            add_paragraph(document, f"• 年度(250日): {win_rate_250d*100:.1f}%")

            # 1.4.4 行业指数图表分析
            add_paragraph(document, '')
            add_title(document, '1.4.4 行业指数图表分析', level=3)

            # 生成行业指数图表
            add_paragraph(document, '')
            industry_charts_paths = generate_industry_index_charts(industry_data, IMAGES_DIR)

            # 添加图表到文档
            add_paragraph(document, '图表 1.12: 行业指数走势与均线分析')
            if len(industry_charts_paths) > 0 and os.path.exists(industry_charts_paths[0]):
                add_image(document, industry_charts_paths[0], width=Inches(6.5))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.13: 行业指数滚动波动率分析')
            if len(industry_charts_paths) > 1 and os.path.exists(industry_charts_paths[1]):
                add_image(document, industry_charts_paths[1], width=Inches(6.5))
                add_paragraph(document, '')

            add_paragraph(document, '图表 1.14: 行业指数日收益率分布')
            if len(industry_charts_paths) > 2 and os.path.exists(industry_charts_paths[2]):
                add_image(document, industry_charts_paths[2], width=Inches(6.5))
                add_paragraph(document, '')

            # 1.4.5 个股与行业对比分析
            add_paragraph(document, '')
            add_title(document, '1.4.5 个股与行业对比分析', level=3)

            add_paragraph(document, '')
            add_paragraph(document, '📊 个股与行业表现对比（多窗口期）：')

            # 定义窗口期
            windows = [
                ('20日窗口', '20d'),
                ('60日窗口', '60d'),
                ('120日窗口', '120d'),
                ('250日窗口', '250d')
            ]

            # 构建多窗口期对比表格
            comparison_headers = ['窗口期', '波动率（个股）', '波动率（行业）', '差异', '年化收益率（个股）', '年化收益率（行业）', '差异']
            comparison_data = []

            for window_name, window_suffix in windows:
                stock_vol = market_data.get(f'volatility_{window_suffix}', 0)
                industry_vol = industry_data.get(f'volatility_{window_suffix}', 0)
                stock_ret = market_data.get(f'annual_return_{window_suffix}', 0)
                industry_ret = industry_data.get(f'annual_return_{window_suffix}', 0)

                vol_diff = (stock_vol - industry_vol) * 100
                ret_diff = (stock_ret - industry_ret) * 100

                comparison_data.append([
                    window_name,
                    f"{stock_vol*100:.2f}%" if stock_vol > 0 else "N/A",
                    f"{industry_vol*100:.2f}%" if industry_vol > 0 else "N/A",
                    f"{vol_diff:+.2f}%" if stock_vol > 0 and industry_vol > 0 else "N/A",
                    f"{stock_ret*100:+.2f}%" if stock_ret != 0 else "N/A",
                    f"{industry_ret*100:+.2f}%" if industry_ret != 0 else "N/A",
                    f"{ret_diff:+.2f}%" if stock_ret != 0 and industry_ret != 0 else "N/A"
                ])

            add_table_data(document, comparison_headers, comparison_data)

            # 分析结论（以60日窗口为主）
            add_paragraph(document, '')
            add_paragraph(document, '💡 行业分析结论（以60日窗口为主）：')

            stock_volatility = market_data.get('volatility_60d', 0)
            industry_volatility = industry_data.get('volatility_60d', 0)
            stock_return = market_data.get('annual_return_60d', 0)
            industry_return = industry_data.get('annual_return_60d', 0)

            if stock_volatility > 0 and industry_volatility > 0:
                if stock_volatility > industry_volatility:
                    add_paragraph(document, f'• 个股波动率({stock_volatility*100:.2f}%)高于行业({industry_volatility*100:.2f}%)，表明个股风险水平较高')
                else:
                    add_paragraph(document, f'• 个股波动率({stock_volatility*100:.2f}%)低于行业({industry_volatility*100:.2f}%)，表明个股相对稳定')

            if stock_return != 0 and industry_return != 0:
                if stock_return > industry_return:
                    add_paragraph(document, f'• 个股收益率({stock_return*100:+.2f}%)跑赢行业({industry_return*100:+.2f}%)，表现优异')
                else:
                    add_paragraph(document, f'• 个股收益率({stock_return*100:+.2f}%)跑输行业({industry_return*100:+.2f}%)，需关注行业竞争压力')

            add_paragraph(document, '• 行业数据为申万三级行业指数，反映细分行业的整体表现')
            add_paragraph(document, '• 建议结合行业景气度、公司基本面等因素综合评估投资价值')
            add_paragraph(document, '• 不同窗口期反映不同时间维度的市场表现，短期波动较大，长期趋势相对稳定')

        else:
            print(f"⚠️ 行业数据文件不存在: {industry_data_file}")
            add_paragraph(document, '⚠️ 行业数据暂未加载，请先运行update_market_data.py生成数据')

    except Exception as e:
        print(f"⚠️ 加载行业数据失败: {e}")
        add_paragraph(document, f'⚠️ 行业数据加载失败: {e}')

    add_section_break(document)

    # ==================== 二、相对估值分析 ====================
    add_title(document, '二、相对估值分析', level=1)

    add_paragraph(document, '本章节通过相对估值法（参数法），将光弘科技与行业内可比公司进行对比分析。')
    add_paragraph(document, '选取申万三级分类"消费电子零部件及组装"行业的同行公司，对比PE、PS、PB等估值倍数。')

    add_title(document, '2.1 估值指标对比', level=2)

    # 使用 Tushare 数据获取估值指标
    try:
        ts_token = os.environ.get('TUSHARE_TOKEN', '')

        if ts_token:
            import tushare as ts
            import time

            pro = ts.pro_api(ts_token)

            # 获取目标公司的估值数据（自动往前推1-2天直到找到交易日）
            trade_date = None
            df_target = None

            for days_back in range(1, 6):  # 尝试往前推1-5天
                test_date = (datetime.now() - timedelta(days=days_back)).strftime('%Y%m%d')
                try:
                    df_target = pro.daily_basic(
                        ts_code=stock_code,
                        trade_date=test_date,
                        fields='ts_code,trade_date,close,pe_ttm,pb,ps_ttm,total_mv'
                    )
                    if not df_target.empty:
                        trade_date = test_date
                        break
                except:
                    continue

            if df_target is None or df_target.empty:
                raise ValueError("未获取到目标公司数据（请检查网络或交易日历）")

            current_metrics_val = {
                'pe': df_target.iloc[0]['pe_ttm'],
                'pb': df_target.iloc[0]['pb'],
                'ps': df_target.iloc[0]['ps_ttm']
            }
            print(f"✅ 获取相对估值数据成功，交易日期: {trade_date}")

            # 获取申万三级行业分类的同行公司（与 notebook 一致）
            df_industry = pro.index_member_all(ts_code=stock_code)
            if df_industry.empty:
                raise ValueError("未获取到行业分类")

            df_industry = df_industry.sort_values('in_date', ascending=False)
            latest_industry = df_industry.iloc[0]
            target_index_code = latest_industry['l3_code']  # 申万三级行业指数代码
            target_industry_l3 = latest_industry['l3_name']  # 行业名称

            # 获取该三级行业的所有成分股
            df_peers = pro.index_member_all(l3_code=target_index_code)
            df_peers = df_peers[df_peers['ts_code'] != stock_code]

            # 获取同行公司基本信息
            peer_codes = df_peers['ts_code'].unique().tolist()
            peer_basic = pro.stock_basic(ts_code=','.join(peer_codes[:30]),
                                       fields='ts_code,name,market')

            peer_stocks_all = pd.merge(df_peers, peer_basic, on='ts_code', how='left')
            peer_stocks_all = peer_stocks_all.drop_duplicates(subset=['ts_code'])

            # 限制数量并排序（扩充到30家）
            peer_stocks_all = peer_stocks_all.head(30)
            peer_names_dict = dict(zip(peer_stocks_all['ts_code'], peer_stocks_all['name_x']))

            # 获取同行公司的估值数据
            peer_data_list = []
            for peer_code in peer_stocks_all['ts_code'].tolist():
                peer_name = peer_names_dict.get(peer_code, peer_code)

                try:
                    df_peer = pro.daily_basic(
                        ts_code=peer_code,
                        trade_date=trade_date,
                        fields='ts_code,pe_ttm,pb,ps_ttm,total_mv'
                    )
                    if not df_peer.empty:
                        df_peer['name'] = peer_name
                        peer_data_list.append(df_peer)
                except:
                    pass

                time.sleep(0.2)  # 避免请求过快

            # 获取申万行业指数的PE数据（使用sw_daily接口）
            sw_index_pe = None
            sw_index_pb = None
            sw_index_ps = None
            try:
                print(f"✅ 正在获取申万行业指数PE数据: {target_index_code}")
                # 获取最近5天的数据，取最新的一条
                end_date = datetime.now().strftime('%Y%m%d')
                start_date = (datetime.now() - timedelta(days=10)).strftime('%Y%m%d')
                df_index = pro.sw_daily(
                    ts_code=target_index_code,
                    start_date=start_date,
                    end_date=end_date
                )
                if df_index is not None and not df_index.empty:
                    # sw_daily返回的是pe列（不是pe_ttm）
                    latest = df_index.iloc[-1]
                    sw_index_pe = latest.get('pe', None)
                    sw_index_pb = latest.get('pb', None)
                    sw_index_ps = latest.get('ps_ttm', None)  # sw_daily可能没有ps_ttm
                    if sw_index_pe:
                        print(f"✅ 申万行业指数PE数据获取成功: PE={sw_index_pe:.2f}, PB={sw_index_pb:.2f}")
                else:
                    print(f"⚠️ 申万行业指数PE数据为空")
            except Exception as e:
                print(f"⚠️ 获取申万行业指数PE数据失败: {e}")

            if peer_data_list:
                peer_companies_val = pd.concat(peer_data_list, ignore_index=True)

                # 过滤异常数据
                peer_companies_val = peer_companies_val[
                    (peer_companies_val['pe_ttm'] > 0) &
                    (peer_companies_val['pe_ttm'] < 500) &
                    (peer_companies_val['pb'] > 0) &
                    (peer_companies_val['pb'] < 20)
                ]

                # 重命名列并进行单位转换
                # total_mv单位是万元，需要转换为亿元（除以10000）
                peer_companies_val['market_cap'] = peer_companies_val['total_mv'] / 10000

                peer_companies_val = peer_companies_val.rename(columns={
                    'pe_ttm': 'pe',
                    'ps_ttm': 'ps',
                    'ts_code': 'code'
                })

                # 只保留需要的列
                peer_companies_val = peer_companies_val[['name', 'code', 'pe', 'ps', 'pb', 'market_cap']]
            else:
                raise ValueError("未获取到同行公司估值数据")

        else:
            raise ValueError("TUSHARE_TOKEN 未设置")

    except Exception as e:
        print(f"获取相对估值数据失败: {e}，使用示例数据")
        # 使用示例数据
        current_metrics_val = {
            'pe': 56.24,
            'pb': 3.71,
            'ps': 2.30
        }

        peer_companies_val = pd.DataFrame({
            'name': ['立讯精密', '歌尔股份', '蓝思科技', '长盈精密', '领益智造', '安洁科技', '比亚迪电子'],
            'code': ['002475.SZ', '002241.SZ', '300433.SZ', '300115.SZ', '002600.SZ', '002635.SZ', '0285.HK'],
            'pe': [20.5, 25.8, 22.3, 28.6, 18.9, 32.1, 24.5],
            'ps': [1.2, 1.8, 1.5, 2.1, 1.0, 2.5, 1.6],
            'pb': [2.8, 3.2, 2.5, 3.8, 2.1, 4.2, 3.0],
            'market_cap': [120, 80, 50, 45, 65, 25, 90]
        })

    # 计算行业统计指标
    industry_stats_val = {
        'pe': {
            'mean': peer_companies_val['pe'].mean(),
            'median': peer_companies_val['pe'].median(),
            'q1': peer_companies_val['pe'].quantile(0.25),
            'q3': peer_companies_val['pe'].quantile(0.75),
            'min': peer_companies_val['pe'].min(),
            'max': peer_companies_val['pe'].max(),
            'std': peer_companies_val['pe'].std()
        },
        'ps': {
            'mean': peer_companies_val['ps'].mean(),
            'median': peer_companies_val['ps'].median(),
            'q1': peer_companies_val['ps'].quantile(0.25),
            'q3': peer_companies_val['ps'].quantile(0.75),
            'min': peer_companies_val['ps'].min(),
            'max': peer_companies_val['ps'].max(),
            'std': peer_companies_val['ps'].std()
        },
        'pb': {
            'mean': peer_companies_val['pb'].mean(),
            'median': peer_companies_val['pb'].median(),
            'q1': peer_companies_val['pb'].quantile(0.25),
            'q3': peer_companies_val['pb'].quantile(0.75),
            'min': peer_companies_val['pb'].min(),
            'max': peer_companies_val['pb'].max(),
            'std': peer_companies_val['pb'].std()
        }
    }

    # 剔除3倍标准差异常值后的行业平均
    industry_avg_val = {
        'pe': industry_stats_val['pe']['mean'],
        'ps': industry_stats_val['ps']['mean'],
        'pb': industry_stats_val['pb']['mean']
    }

    # 估值指标对比表（增强版）
    valuation_headers = ['指标', '光弘科技', '行业平均', '中位数', 'Q1(25分位)', 'Q3(75分位)', '最小值', '最大值', '偏离度']
    valuation_data = [
        ['PE (TTM)',
         f"{current_metrics_val['pe']:.2f}倍",
         f"{industry_stats_val['pe']['mean']:.2f}倍",
         f"{industry_stats_val['pe']['median']:.2f}倍",
         f"{industry_stats_val['pe']['q1']:.2f}倍",
         f"{industry_stats_val['pe']['q3']:.2f}倍",
         f"{industry_stats_val['pe']['min']:.2f}倍",
         f"{industry_stats_val['pe']['max']:.2f}倍",
         f"{(current_metrics_val['pe']-industry_stats_val['pe']['mean'])/industry_stats_val['pe']['mean']*100:+.1f}%"],
        ['PB',
         f"{current_metrics_val['pb']:.2f}倍",
         f"{industry_stats_val['pb']['mean']:.2f}倍",
         f"{industry_stats_val['pb']['median']:.2f}倍",
         f"{industry_stats_val['pb']['q1']:.2f}倍",
         f"{industry_stats_val['pb']['q3']:.2f}倍",
         f"{industry_stats_val['pb']['min']:.2f}倍",
         f"{industry_stats_val['pb']['max']:.2f}倍",
         f"{(current_metrics_val['pb']-industry_stats_val['pb']['mean'])/industry_stats_val['pb']['mean']*100:+.1f}%"],
        ['PS (TTM)',
         f"{current_metrics_val['ps']:.2f}倍",
         f"{industry_stats_val['ps']['mean']:.2f}倍",
         f"{industry_stats_val['ps']['median']:.2f}倍",
         f"{industry_stats_val['ps']['q1']:.2f}倍",
         f"{industry_stats_val['ps']['q3']:.2f}倍",
         f"{industry_stats_val['ps']['min']:.2f}倍",
         f"{industry_stats_val['ps']['max']:.2f}倍",
         f"{(current_metrics_val['ps']-industry_stats_val['ps']['mean'])/industry_stats_val['ps']['mean']*100:+.1f}%"]
    ]
    add_table_data(document, valuation_headers, valuation_data)

    # 添加统计分析说明
    add_paragraph(document, '')
    add_paragraph(document, '行业估值统计说明：')
    add_paragraph(document, '• 行业平均：所有同行公司的算术平均值（受极端值影响）')
    add_paragraph(document, '• 中位数：行业50%的公司估值低于此水平，抗极端值干扰')
    add_paragraph(document, '• Q1(25分位)：行业25%的公司估值低于此水平')
    add_paragraph(document, '• Q3(75分位)：行业75%的公司估值低于此水平（即25%的公司高于此水平）')
    add_paragraph(document, '• 最小/最大值：行业中的估值极值')
    add_paragraph(document, '• 数据已过滤异常值（PE<500, PB<20）以避免极端情况影响统计')
    add_paragraph(document, f'• 样本量：共{len(peer_companies_val)}家同行公司')

    # 添加同行公司名单
    add_paragraph(document, '')
    add_title(document, '2.1.1 同行公司名单', level=3)
    add_paragraph(document, f'基于申万三级行业分类"消费电子零部件及组装"筛选的同行公司：')

    # 按市值排序的同行公司表格
    peer_companies_sorted = peer_companies_val.sort_values('market_cap', ascending=False)
    peer_headers = ['公司名称', '股票代码', 'PE (TTM)', 'PB', 'PS (TTM)', '市值(亿元)']
    peer_rows = []
    for _, row in peer_companies_sorted.iterrows():
        peer_rows.append([
            row['name'],
            row['code'],
            f"{row['pe']:.2f}",
            f"{row['pb']:.2f}",
            f"{row['ps']:.2f}",
            f"{row['market_cap']:.2f}"
        ])
    add_table_data(document, peer_headers, peer_rows)

    # 添加行业统计汇总
    add_paragraph(document, '')
    add_paragraph(document, '行业估值统计汇总：')
    add_paragraph(document, f'• PE: 平均{industry_stats_val["pe"]["mean"]:.2f}倍，中位数{peer_companies_val["pe"].median():.2f}倍，标准差{industry_stats_val["pe"]["std"]:.2f}倍')
    add_paragraph(document, f'  • Q1-Q3区间: [{industry_stats_val["pe"]["q1"]:.2f}, {industry_stats_val["pe"]["q3"]:.2f}]倍，极值范围: [{industry_stats_val["pe"]["min"]:.2f}, {industry_stats_val["pe"]["max"]:.2f}]倍')
    add_paragraph(document, f'• PB: 平均{industry_stats_val["pb"]["mean"]:.2f}倍，中位数{peer_companies_val["pb"].median():.2f}倍，标准差{industry_stats_val["pb"]["std"]:.2f}倍')
    add_paragraph(document, f'  • Q1-Q3区间: [{industry_stats_val["pb"]["q1"]:.2f}, {industry_stats_val["pb"]["q3"]:.2f}]倍，极值范围: [{industry_stats_val["pb"]["min"]:.2f}, {industry_stats_val["pb"]["max"]:.2f}]倍')
    add_paragraph(document, f'• PS: 平均{industry_stats_val["ps"]["mean"]:.2f}倍，中位数{peer_companies_val["ps"].median():.2f}倍，标准差{industry_stats_val["ps"]["std"]:.2f}倍')
    add_paragraph(document, f'  • Q1-Q3区间: [{industry_stats_val["ps"]["q1"]:.2f}, {industry_stats_val["ps"]["q3"]:.2f}]倍，极值范围: [{industry_stats_val["ps"]["min"]:.2f}, {industry_stats_val["ps"]["max"]:.2f}]倍')

    add_paragraph(document, '')
    add_paragraph(document, '图表 2.0: 相对估值对比分析 - 估值指标对比')
    chart_paths, df_scenarios = generate_relative_valuation_charts_split(
        current_metrics_val, industry_avg_val, peer_companies_val, IMAGES_DIR
    )
    add_image(document, chart_paths[0])  # 估值指标对比

    add_paragraph(document, '')
    add_paragraph(document, '图表 2.1: 相对估值对比分析 - PE倍数对比')
    add_image(document, chart_paths[1])

    add_paragraph(document, '')
    add_paragraph(document, '图表 2.2: 相对估值对比分析 - PB倍数对比')
    add_image(document, chart_paths[2])

    add_paragraph(document, '')
    add_paragraph(document, '图表 2.3: 相对估值对比分析 - PS倍数对比')
    add_image(document, chart_paths[3])

    # ==================== 2.1.2 申万行业指数估值 ====================
    add_title(document, '2.1.2 申万行业指数估值', level=3)
    add_paragraph(document, '')

    if sw_index_pe is not None:
        add_paragraph(document, f'本节展示申万三级行业指数"{target_industry_l3}"的估值数据，提供官方行业指数视角的估值基准。')
        add_paragraph(document, f'申万行业指数代码: {target_index_code}')
        add_paragraph(document, '')

        # 申万行业指数估值表格
        sw_index_headers = ['指标', '申万行业指数', '光弘科技', '差异', '说明']
        sw_index_data = [
            ['PE (TTM)',
             f"{sw_index_pe:.2f}倍",
             f"{current_metrics_val['pe']:.2f}倍",
             f"{(current_metrics_val['pe']-sw_index_pe)/sw_index_pe*100:+.1f}%",
             '行业指数PE反映行业整体估值水平' if current_metrics_val['pe'] > sw_index_pe else '个股PE低于行业指数，相对低估'],
            ['PB',
             f"{sw_index_pb:.2f}倍" if sw_index_pb else "N/A",
             f"{current_metrics_val['pb']:.2f}倍",
             f"{(current_metrics_val['pb']-sw_index_pb)/sw_index_pb*100:+.1f}%" if sw_index_pb else "N/A",
             '市净率反映行业整体账面价值溢价'],
            ['PS (TTM)',
             f"{sw_index_ps:.2f}倍" if sw_index_ps else "N/A",
             f"{current_metrics_val['ps']:.2f}倍",
             f"{(current_metrics_val['ps']-sw_index_ps)/sw_index_ps*100:+.1f}%" if sw_index_ps else "N/A",
             '市销率反映行业整体营收能力']
        ]
        add_table_data(document, sw_index_headers, sw_index_data)

        add_paragraph(document, '')
        add_paragraph(document, '💡 申万行业指数估值说明：')
        add_paragraph(document, f'• 申万行业指数是基于该行业所有成分股按市值加权计算的指数')
        add_paragraph(document, f'• 指数PE/PB/PS反映行业整体的估值水平，不同于同行公司平均（简单平均）')
        add_paragraph(document, f'• 指数估值受大盘股权重影响，更能代表行业龙头公司的估值水平')
        add_paragraph(document, f'• 同行公司平均反映行业内典型公司的估值，受小盘股影响较大')
        add_paragraph(document, '')

        # 对比分析
        add_paragraph(document, '📊 对比分析：', bold=True)
        if abs(current_metrics_val['pe'] - sw_index_pe) / sw_index_pe < 0.1:
            add_paragraph(document, f'✅ 个股PE({current_metrics_val["pe"]:.2f}倍)与申万行业指数PE({sw_index_pe:.2f}倍)基本一致，估值合理')
        elif current_metrics_val['pe'] < sw_index_pe:
            add_paragraph(document, f'✅ 个股PE({current_metrics_val["pe"]:.2f}倍)低于申万行业指数PE({sw_index_pe:.2f}倍)，相对行业指数低估')
        else:
            add_paragraph(document, f'⚠️ 个股PE({current_metrics_val["pe"]:.2f}倍)高于申万行业指数PE({sw_index_pe:.2f}倍)，相对行业指数高估')
    else:
        add_paragraph(document, f'⚠️ 申万行业指数"{target_industry_l3}"的估值数据暂时无法获取')
        add_paragraph(document, '可能原因：指数代码不正确或数据源暂时不可用')

    add_paragraph(document, '')

    add_title(document, '2.2 估值偏离度分析', level=2)
    add_paragraph(document, '本节分析标的公司与同行公司和申万行业指数的估值偏离情况，评估估值相对位置。')
    add_paragraph(document, '')

    # 计算PE在行业中的分位数位置
    pe_position = (peer_companies_val['pe'] < current_metrics_val['pe']).sum() / len(peer_companies_val) * 100
    pb_position = (peer_companies_val['pb'] < current_metrics_val['pb']).sum() / len(peer_companies_val) * 100
    ps_position = (peer_companies_val['ps'] < current_metrics_val['ps']).sum() / len(peer_companies_val) * 100

    # 2.2.1 与同行公司对比
    add_title(document, '2.2.1 与同行公司对比', level=3)
    add_paragraph(document, '')

    add_paragraph(document, f"• PE偏离度: {(current_metrics_val['pe']-industry_avg_val['pe'])/industry_avg_val['pe']*100:+.1f}%，位于行业{pe_position:.1f}%分位")
    add_paragraph(document, f"• PB偏离度: {(current_metrics_val['pb']-industry_avg_val['pb'])/industry_avg_val['pb']*100:+.1f}%，位于行业{pb_position:.1f}%分位")
    add_paragraph(document, f"• PS偏离度: {(current_metrics_val['ps']-industry_avg_val['ps'])/industry_avg_val['ps']*100:+.1f}%，位于行业{ps_position:.1f}%分位")

    add_paragraph(document, '')

    # PE分位数分析
    if current_metrics_val['pe'] > industry_stats_val['pe']['q3']:
        add_paragraph(document, f'⚠️ PE({current_metrics_val["pe"]:.2f}倍)高于行业Q3({industry_stats_val["pe"]["q3"]:.2f}倍)，处于行业高位，估值偏高')
    elif current_metrics_val['pe'] < industry_stats_val['pe']['q1']:
        add_paragraph(document, f'✅ PE({current_metrics_val["pe"]:.2f}倍)低于行业Q1({industry_stats_val["pe"]["q1"]:.2f}倍)，处于行业低位，估值偏低')
    else:
        add_paragraph(document, f'ℹ️ PE({current_metrics_val["pe"]:.2f}倍)介于行业Q1({industry_stats_val["pe"]["q1"]:.2f}倍)和Q3({industry_stats_val["pe"]["q3"]:.2f}倍)之间，估值合理')

    # PB分位数分析
    if current_metrics_val['pb'] > industry_stats_val['pb']['q3']:
        add_paragraph(document, f'⚠️ PB({current_metrics_val["pb"]:.2f}倍)高于行业Q3({industry_stats_val["pb"]["q3"]:.2f}倍)，市净率偏高')
    elif current_metrics_val['pb'] < industry_stats_val['pb']['q1']:
        add_paragraph(document, f'✅ PB({current_metrics_val["pb"]:.2f}倍)低于行业Q1({industry_stats_val["pb"]["q1"]:.2f}倍)，市净率偏低')

    # PS分位数分析
    if current_metrics_val['ps'] > industry_stats_val['ps']['q3']:
        add_paragraph(document, f'⚠️ PS({current_metrics_val["ps"]:.2f}倍)高于行业Q3({industry_stats_val["ps"]["q3"]:.2f}倍)，市销率偏高')
    elif current_metrics_val['ps'] < industry_stats_val['ps']['q1']:
        add_paragraph(document, f'✅ PS({current_metrics_val["ps"]:.2f}倍)低于行业Q1({industry_stats_val["ps"]["q1"]:.2f}倍)，市销率偏低')

    # 2.2.2 与申万行业指数对比
    if sw_index_pe is not None:
        add_paragraph(document, '')
        add_title(document, '2.2.2 与申万行业指数对比', level=3)
        add_paragraph(document, '')

        # 计算与申万行业指数的偏离度
        pe_dev_sw = (current_metrics_val['pe'] - sw_index_pe) / sw_index_pe * 100
        pb_dev_sw = (current_metrics_val['pb'] - sw_index_pb) / sw_index_pb * 100 if sw_index_pb else None
        ps_dev_sw = (current_metrics_val['ps'] - sw_index_ps) / sw_index_ps * 100 if sw_index_ps else None

        add_paragraph(document, f"• PE偏离度: {pe_dev_sw:+.1f}%（标的{current_metrics_val['pe']:.2f}倍 vs 申万{sw_index_pe:.2f}倍）")
        if pb_dev_sw is not None:
            add_paragraph(document, f"• PB偏离度: {pb_dev_sw:+.1f}%（标的{current_metrics_val['pb']:.2f}倍 vs 申万{sw_index_pb:.2f}倍）")
        if ps_dev_sw is not None:
            add_paragraph(document, f"• PS偏离度: {ps_dev_sw:+.1f}%（标的{current_metrics_val['ps']:.2f}倍 vs 申万{sw_index_ps:.2f}倍）")

        add_paragraph(document, '')

        # PE申万指数对比分析
        if abs(pe_dev_sw) < 10:
            add_paragraph(document, f'✅ PE({current_metrics_val["pe"]:.2f}倍)与申万行业指数PE({sw_index_pe:.2f}倍)基本一致，偏离度{pe_dev_sw:+.1f}%')
        elif pe_dev_sw > 0:
            add_paragraph(document, f'⚠️ PE({current_metrics_val["pe"]:.2f}倍)高于申万行业指数PE({sw_index_pe:.2f}倍)，溢价{pe_dev_sw:+.1f}%')
        else:
            add_paragraph(document, f'✅ PE({current_metrics_val["pe"]:.2f}倍)低于申万行业指数PE({sw_index_pe:.2f}倍)，折价{pe_dev_sw:+.1f}%')

        # PB申万指数对比分析
        if pb_dev_sw is not None:
            if abs(pb_dev_sw) < 10:
                add_paragraph(document, f'✅ PB({current_metrics_val["pb"]:.2f}倍)与申万行业指数PB({sw_index_pb:.2f}倍)基本一致')
            elif pb_dev_sw > 0:
                add_paragraph(document, f'⚠️ PB({current_metrics_val["pb"]:.2f}倍)高于申万行业指数PB({sw_index_pb:.2f}倍)，溢价{pb_dev_sw:+.1f}%')
            else:
                add_paragraph(document, f'✅ PB({current_metrics_val["pb"]:.2f}倍)低于申万行业指数PB({sw_index_pb:.2f}倍)，折价{pb_dev_sw:+.1f}%')

        add_paragraph(document, '')
        add_paragraph(document, '💡 申万行业指数说明：')
        add_paragraph(document, f'• 申万行业指数基于所有成分股市值加权，反映行业整体估值水平')
        add_paragraph(document, f'• 与申万指数对比可判断个股相对行业整体的估值位置')
        add_paragraph(document, f'• 正偏离表示估值高于行业平均，负偏离表示估值低于行业平均')

    add_paragraph(document, '')

    # ==================== 2.3 PE历史分位数趋势分析 ====================
    add_title(document, '2.3 PE历史分位数趋势分析', level=2)

    add_paragraph(document, '本节通过分析标的股票和所属行业的PE历史走势及分位数变化，从时间维度评估估值的合理性。')
    add_paragraph(document, '基于最近5年的历史数据（约1250个交易日），通过对比个股与行业的PE历史分位数趋势，可以更清晰地判断当前估值处于历史哪个水平。')
    add_paragraph(document, '5年的历史周期能够覆盖完整的牛熊周期，提供更可靠的估值基准。')
    add_paragraph(document, '')

    # 尝试从tushare获取历史PE数据并生成趋势图
    try:
        from utils.pe_history_analyzer import PEHistoryAnalyzer

        print("\n=== 开始PE历史分位数趋势分析 ===")

        # 创建PE历史分析器
        pe_analyzer = PEHistoryAnalyzer()

        # 获取个股历史PE数据（最近5年）
        stock_pe_data = pe_analyzer.get_stock_pe_history(stock_code, days=1825)

        # 获取行业历史PE数据
        industry_name, industry_code, industry_pe_data = pe_analyzer.get_industry_pe_history(stock_code, days=1825)

        if stock_pe_data is not None and industry_pe_data is not None:
            print(f"✅ 成功获取历史PE数据")
            print(f"   个股数据: {len(stock_pe_data)}条")
            print(f"   行业数据: {len(industry_pe_data)}条")
            print(f"   行业: {industry_name} ({industry_code})")

            # 计算个股历史分位数统计
            stock_pe_current = stock_pe_data.iloc[-1]['pe_ttm']
            stock_pe_min = stock_pe_data['pe_ttm'].min()
            stock_pe_max = stock_pe_data['pe_ttm'].max()
            stock_pe_median = stock_pe_data['pe_ttm'].median()
            stock_pe_percentile = (stock_pe_data['pe_ttm'] < stock_pe_current).sum() / len(stock_pe_data) * 100

            # 计算申万行业指数历史分位数统计
            sw_index_pe_current = industry_pe_data.iloc[-1]['pe_ttm']
            sw_index_pe_min = industry_pe_data['pe_ttm'].min()
            sw_index_pe_max = industry_pe_data['pe_ttm'].max()
            sw_index_pe_median = industry_pe_data['pe_ttm'].median()
            sw_index_pe_percentile = (industry_pe_data['pe_ttm'] < sw_index_pe_current).sum() / len(industry_pe_data) * 100

            # 获取同行公司的历史PE数据
            print("\n获取同行公司历史PE数据...")
            custom_peer_pe_current = None
            custom_peer_pe_min = None
            custom_peer_pe_max = None
            custom_peer_pe_median = None
            custom_peer_pe_percentile = None

            try:
                # 提取同行公司代码列表
                peer_codes = peer_companies_val['code'].tolist()[:20]  # 取前20个
                print(f"  同行公司数量: {len(peer_codes)}")

                # 获取每个同行公司的历史PE
                peer_pe_histories = []
                for peer_code in peer_codes:
                    try:
                        peer_history = pe_analyzer.get_stock_pe_history(peer_code, days=1825)
                        if peer_history is not None and not peer_history.empty:
                            peer_history = peer_history.rename(columns={'pe_ttm': 'pe'})
                            peer_pe_histories.append(peer_history)
                    except Exception as e:
                        print(f"    获取{peer_code}历史PE失败: {e}")
                        continue

                if peer_pe_histories:
                    # 合并所有同行公司的历史PE
                    peer_pe_df = pd.concat(peer_pe_histories, ignore_index=True)

                    # 按日期分组计算平均值
                    custom_peer_pe_data = peer_pe_df.groupby('trade_date')['pe'].mean().reset_index()
                    custom_peer_pe_data = custom_peer_pe_data.sort_values('trade_date').reset_index(drop=True)

                    # 计算统计指标
                    custom_peer_pe_current = custom_peer_pe_data.iloc[-1]['pe']
                    custom_peer_pe_min = custom_peer_pe_data['pe'].min()
                    custom_peer_pe_max = custom_peer_pe_data['pe'].max()
                    custom_peer_pe_median = custom_peer_pe_data['pe'].median()
                    custom_peer_pe_percentile = (custom_peer_pe_data['pe'] < custom_peer_pe_current).sum() / len(custom_peer_pe_data) * 100

                    print(f"  ✅ 同行公司历史PE计算成功:")
                    print(f"     当前PE: {custom_peer_pe_current:.2f}倍")
                    print(f"     数据点数: {len(custom_peer_pe_data)}")
                else:
                    print(f"  ⚠️ 未获取到同行公司历史PE数据")

            except Exception as e:
                print(f"  ❌ 获取同行公司历史PE失败: {e}")

            # 添加历史分位数统计表格
            if custom_peer_pe_current is not None:
                # 有自定义同行数据的完整表格
                pe_history_headers = ['指标', '标的股票', '行业指数<br/>(申万)', '行业指数<br/>(自定义)', '差异<br/>(vs申万)', '差异<br/>(vs自定义)']
                pe_history_data = [
                    ['当前PE-TTM',
                     f'{stock_pe_current:.2f}倍',
                     f'{sw_index_pe_current:.2f}倍',
                     f'{custom_peer_pe_current:.2f}倍',
                     f'{(stock_pe_current/sw_index_pe_current-1)*100:+.1f}%',
                     f'{(stock_pe_current/custom_peer_pe_current-1)*100:+.1f}%'],
                    ['历史最小PE',
                     f'{stock_pe_min:.2f}倍',
                     f'{sw_index_pe_min:.2f}倍',
                     f'{custom_peer_pe_min:.2f}倍',
                     f'{stock_pe_min-sw_index_pe_min:+.2f}倍',
                     f'{stock_pe_min-custom_peer_pe_min:+.2f}倍'],
                    ['历史最大PE',
                     f'{stock_pe_max:.2f}倍',
                     f'{sw_index_pe_max:.2f}倍',
                     f'{custom_peer_pe_max:.2f}倍',
                     f'{stock_pe_max-sw_index_pe_max:+.2f}倍',
                     f'{stock_pe_max-custom_peer_pe_max:+.2f}倍'],
                    ['历史中位数PE',
                     f'{stock_pe_median:.2f}倍',
                     f'{sw_index_pe_median:.2f}倍',
                     f'{custom_peer_pe_median:.2f}倍',
                     f'{stock_pe_median-sw_index_pe_median:+.2f}倍',
                     f'{stock_pe_median-custom_peer_pe_median:+.2f}倍'],
                    ['当前分位数',
                     f'{stock_pe_percentile:.1f}%',
                     f'{sw_index_pe_percentile:.1f}%',
                     f'{custom_peer_pe_percentile:.1f}%',
                     f'{stock_pe_percentile-sw_index_pe_percentile:+.1f}%',
                     f'{stock_pe_percentile-custom_peer_pe_percentile:+.1f}%']
                ]
            else:
                # 只有申万数据的简化表格
                pe_history_headers = ['指标', '标的股票', '行业指数<br/>(申万)', '差异']
                pe_history_data = [
                    ['当前PE-TTM', f'{stock_pe_current:.2f}倍', f'{sw_index_pe_current:.2f}倍', f'{(stock_pe_current/sw_index_pe_current-1)*100:+.1f}%'],
                    ['历史最小PE', f'{stock_pe_min:.2f}倍', f'{sw_index_pe_min:.2f}倍', f'{stock_pe_min-sw_index_pe_min:+.2f}倍'],
                    ['历史最大PE', f'{stock_pe_max:.2f}倍', f'{sw_index_pe_max:.2f}倍', f'{stock_pe_max-sw_index_pe_max:+.2f}倍'],
                    ['历史中位数PE', f'{stock_pe_median:.2f}倍', f'{sw_index_pe_median:.2f}倍', f'{stock_pe_median-sw_index_pe_median:+.2f}倍'],
                    ['当前分位数', f'{stock_pe_percentile:.1f}%', f'{sw_index_pe_percentile:.1f}%', f'{stock_pe_percentile-sw_index_pe_percentile:+.1f}%']
                ]

            add_table_data(document, pe_history_headers, pe_history_data)

            add_paragraph(document, '')
            add_paragraph(document, '💡 历史分位数说明：')
            add_paragraph(document, f'• 当前分位数表示当前PE在历史数据中的相对位置')
            add_paragraph(document, f'• 例如：{stock_pe_percentile:.1f}%分位数表示历史上只有{stock_pe_percentile:.1f}%的时间PE低于当前值')
            add_paragraph(document, f'• 50%分位数即为中位数，代表历史平均水平')
            if custom_peer_pe_current is not None:
                add_paragraph(document, f'• 行业指数（申万）：申万行业指数的PE，基于所有成分股按市值加权计算')
                add_paragraph(document, f'• 行业指数（自定义）：2.1.1节中同行公司PE的简单平均，反映所选同行公司的平均估值水平')
            add_paragraph(document, '')

            # 生成PE趋势图
            pe_trend_chart_path = os.path.join(IMAGES_DIR, '02_4_pe_trend_analysis.png')
            chart_path = pe_analyzer.generate_pe_trend_chart(
                stock_code, stock_pe_data,
                industry_name, industry_pe_data,
                pe_trend_chart_path
            )

            # 添加图表到文档
            if chart_path and os.path.exists(chart_path):
                add_paragraph(document, '')
                add_paragraph(document, '图表 2.4: PE历史分位数趋势分析')
                add_image(document, chart_path, width=Inches(6.5))

                add_paragraph(document, '')
                add_paragraph(document, '图表解读：', bold=True)
                add_paragraph(document, '')
                add_paragraph(document, f'左上-PE走势对比：')
                add_paragraph(document, f'  • 蓝线：{stock_code}的PE-TTM走势')
                add_paragraph(document, f'  • 红线：{industry_name}的PE-TTM走势')
                add_paragraph(document, f'  • 两条线的相对位置反映个股相对行业的估值水平')
                add_paragraph(document, '')

                add_paragraph(document, f'右上-PE相对位置（个股/行业）：')
                add_paragraph(document, f'  • 比值>1：个股PE高于行业，溢价')
                add_paragraph(document, f'  • 比值<1：个股PE低于行业，折价')
                add_paragraph(document, f'  • 比值=1：与行业持平')
                add_paragraph(document, '')

                add_paragraph(document, f'左下-个股PE历史分位数：')
                add_paragraph(document, f'  • 显示{stock_code}的PE在历史中的位置变化')
                add_paragraph(document, f'  • 当前分位数：{stock_pe_percentile:.1f}%')
                add_paragraph(document, f'  • 分位数上升表示估值相对历史提升')
                add_paragraph(document, '')

                add_paragraph(document, f'右下-行业PE历史分位数：')
                add_paragraph(document, f'  • 显示{industry_name}的PE在历史中的位置变化')
                add_paragraph(document, f'  • 当前分位数：{industry_pe_percentile:.1f}%')
                add_paragraph(document, f'  • 可用于判断行业整体估值水平')
                add_paragraph(document, '')

            # 添加分析结论
            add_paragraph(document, '')
            add_paragraph(document, 'PE历史分位数趋势分析结论：', bold=True)
            add_paragraph(document, '')

            # 估值水平判断
            if stock_pe_percentile >= 80:
                stock_valuation_level = "历史高位"
                stock_emoji = "🔴"
                stock_comment = f"当前PE处于历史{stock_pe_percentile:.1f}%分位数，属于历史高位，估值偏高，需警惕回调风险"
            elif stock_pe_percentile >= 60:
                stock_valuation_level = "历史中高位"
                stock_emoji = "🟠"
                stock_comment = f"当前PE处于历史{stock_pe_percentile:.1f}%分位数，属于历史中高位，估值相对偏高"
            elif stock_pe_percentile >= 40:
                stock_valuation_level = "历史中位数"
                stock_emoji = "🟡"
                stock_comment = f"当前PE处于历史{stock_pe_percentile:.1f}%分位数，接近历史中位数，估值合理"
            elif stock_pe_percentile >= 20:
                stock_valuation_level = "历史中低位"
                stock_emoji = "🟢"
                stock_comment = f"当前PE处于历史{stock_pe_percentile:.1f}%分位数，属于历史中低位，估值相对偏低"
            else:
                stock_valuation_level = "历史低位"
                stock_emoji = "✅"
                stock_comment = f"当前PE处于历史{stock_pe_percentile:.1f}%分位数，属于历史低位，估值偏低，安全边际较高"

            add_paragraph(document, f'{stock_emoji} 个股估值水平：{stock_valuation_level}')
            add_paragraph(document, f'   {stock_comment}')
            add_paragraph(document, '')

            # 与行业对比
            if stock_pe_percentile > industry_pe_percentile + 20:
                relative_comment = f"个股分位数({stock_pe_percentile:.1f}%)显著高于行业({industry_pe_percentile:.1f}%)，相对行业估值偏高"
                relative_emoji = "⚠️"
            elif stock_pe_percentile < industry_pe_percentile - 20:
                relative_comment = f"个股分位数({stock_pe_percentile:.1f}%)显著低于行业({industry_pe_percentile:.1f}%)，相对行业估值偏低，安全边际较高"
                relative_emoji = "✅"
            else:
                relative_comment = f"个股分位数({stock_pe_percentile:.1f}%)与行业({industry_pe_percentile:.1f}%)基本持平"
                relative_emoji = "ℹ️"

            add_paragraph(document, f'{relative_emoji} 相对行业估值：{relative_comment}')
            add_paragraph(document, '')

            # 投资建议
            add_paragraph(document, '💡 历史分位数投资启示：')
            if stock_pe_percentile <= 25:
                add_paragraph(document, f'• 当前PE处于历史{stock_pe_percentile:.1f}%分位数（低位），历史上仅{stock_pe_percentile:.1f}%的时间估值更低')
                add_paragraph(document, f'• 从历史角度看，当前估值具备较好的安全边际')
                add_paragraph(document, f'• 建议积极关注，估值修复空间较大')
            elif stock_pe_percentile <= 50:
                add_paragraph(document, f'• 当前PE处于历史{stock_pe_percentile:.1f}%分位数（中低位），估值相对合理或偏低')
                add_paragraph(document, f'• 从历史角度看，当前估值风险可控')
                add_paragraph(document, f'• 建议适度配置，关注基本面变化')
            elif stock_pe_percentile <= 75:
                add_paragraph(document, f'• 当前PE处于历史{stock_pe_percentile:.1f}%分位数（中高位），估值相对偏高')
                add_paragraph(document, f'• 从历史角度看，当前估值风险上升')
                add_paragraph(document, f'• 建议谨慎参与，等待更好的买入时机')
            else:
                add_paragraph(document, f'• 当前PE处于历史{stock_pe_percentile:.1f}%分位数（高位），估值处于历史高位')
                add_paragraph(document, f'• 从历史角度看，当前估值风险较大')
                add_paragraph(document, f'• 建议等待估值回落至历史中低位再考虑参与')

        else:
            print("⚠️ PE历史数据获取不完整，跳过趋势图生成")
            add_paragraph(document, '⚠️ PE历史分位数趋势图暂时无法生成，可能原因：')
            add_paragraph(document, '   • tushare数据缺失或API调用限制')
            add_paragraph(document, '   • 股票或行业历史数据不足')

    except ImportError as e:
        print(f"⚠️ PE历史分析器导入失败: {e}")
        add_paragraph(document, '⚠️ PE历史分位数趋势分析功能暂不可用')

    except Exception as e:
        print(f"❌ PE历史分位数趋势分析失败: {e}")
        add_paragraph(document, f'⚠️ PE历史分位数趋势分析执行失败: {e}')

    add_paragraph(document, '')

    add_section_break(document)

    # ==================== 三、DCF估值分析 ====================
    add_title(document, '三、DCF估值分析', level=1)

    add_paragraph(document, '本章节使用现金流折现模型评估公司的内在价值。')

    add_title(document, '3.1 DCF估值方法说明', level=2)

    add_paragraph(document, 'DCF（Discounted Cash Flow）估值法通过预测公司未来自由现金流，并以加权平均资本成本（WACC）折现至现值，得出公司的内在价值。')

    add_paragraph(document, '')
    add_paragraph(document, '估值步骤：')
    add_paragraph(document, '1. 预测未来10年自由现金流（FCF）')
    add_paragraph(document, '2. 计算终值（Terminal Value）')
    add_paragraph(document, '3. 以WACC折现至现值')
    add_paragraph(document, '4. 减去净债务，得到股权价值')
    add_paragraph(document, '5. 除以总股数，得到每股价值')

    # 从placement_params中获取财务数据
    net_income = project_params.get('net_income', 253532329.85)  # 净利润

    # 获取真实总股本（从 Tushare）
    total_shares = 767460689  # 光弘科技真实总股本（约7.67亿股）
    try:
        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if ts_token:
            import tushare as ts

            pro = ts.pro_api(ts_token)
            trade_date = (datetime.now() - timedelta(days=1)).strftime('%Y%m%d')

            df_basic = pro.daily_basic(
                ts_code=stock_code,
                trade_date=trade_date,
                fields='ts_code,total_share'
            )

            if not df_basic.empty:
                total_shares = float(df_basic.iloc[0]['total_share']) * 10000  # 万股转股
    except Exception as e:
        print(f"获取总股本失败，使用估算值: {e}")

    net_assets = project_params.get('net_assets', 4939639031.34)  # 净资产（元）
    total_debt = project_params.get('total_debt', 4507009002.2)  # 总债务（元）
    net_income = project_params.get('net_income', 253532329.85)  # 净利润（元）

    # 尝试从Tushare获取资产负债表和利润表数据
    net_debt = None
    cash_assets = 0.0
    total_liab_value = 0.0

    # 初始化CAGR计算变量
    net_income_cagr = None  # 净利润CAGR（备选）
    fcf_cagr = None  # FCF的CAGR（优先使用）
    historical_incomes = None

    try:
        from update_market_data import TushareFinancialData

        ts_token = os.environ.get('TUSHARE_TOKEN', '')
        if ts_token:
            financial = TushareFinancialData(stock_code)

            # 获取资产负债表数据
            balancesheet = financial.get_latest_balancesheet()
            if balancesheet:
                total_liab_value = balancesheet['total_liabilities'] / 100000000  # 转亿元
                cash_equivalents = balancesheet['cash_equivalents'] / 100000000  # 转亿元

                # 更新财务数据
                net_assets = balancesheet['total_equity']
                total_debt = balancesheet['total_liabilities']

                # 净债务 = 总负债 - 货币资金
                net_debt = total_liab_value - cash_equivalents
                cash_assets = cash_equivalents

                print(f"✅ 从Tushare获取资产负债表数据成功，期间: {balancesheet['report_date']}")
                print(f"   总负债: {total_liab_value:.2f} 亿元")
                print(f"   货币资金: {cash_equivalents:.2f} 亿元")
                if net_debt > 0:
                    print(f"   净债务: {net_debt:.2f} 亿元")
                else:
                    print(f"   净现金: {-net_debt:.2f} 亿元")

            # 获取利润表数据（净利润）
            income = financial.get_latest_income()
            if income:
                net_income = income['net_income']
                print(f"✅ 从Tushare获取利润表数据成功，期间: {income['report_date']}")
                print(f"   净利润: {net_income/100000000:.2f} 亿元")

                # 更新项目参数
                project_params['net_income'] = net_income

            # 获取历史净利润数据以计算CAGR（作为备选方案）
            historical_incomes = financial.get_historical_net_income(years=5)
            net_income_cagr = None  # 净利润CAGR（备选）
            if historical_incomes and len(historical_incomes) >= 2:
                # 计算CAGR (复合年增长率)
                # CAGR = (终值/起始值)^(1/年数) - 1
                start_income = historical_incomes[-1]
                end_income = historical_incomes[0]
                n_years = len(historical_incomes) - 1

                if start_income > 0 and end_income > 0:
                    net_income_cagr = (end_income / start_income) ** (1 / n_years) - 1
                    print(f"✅ 计算净利润历史CAGR（备选）:")
                    print(f"   起始净利润: {start_income/100000000:.2f} 亿元")
                    print(f"   最新净利润: {end_income/100000000:.2f} 亿元")
                    print(f"   历史年数: {n_years} 年")
                    print(f"   净利润CAGR: {net_income_cagr*100:+.2f}%")
                    print(f"   注：将优先使用FCF的CAGR，如果不可用则使用净利润CAGR")

            # 获取完整历史FCF数据（用于DCF估值）
            historical_fcf_data = financial.get_historical_fcf_for_dcf(max_years=15)
            if historical_fcf_data:
                project_params['historical_fcf_data'] = historical_fcf_data
                print(f"✅ 已保存历史FCF数据到项目参数")

    except Exception as e:
        print(f"从Tushare获取财务数据失败: {e}")

    # 如果无法获取实际数据，使用项目参数中的估计值
    if net_debt is None:
        # 计算净债务 = 总负债 - 货币资金
        cash_equivalents = project_params.get('cash_equivalents', 0)
        net_debt = (total_debt - cash_equivalents) / 100000000  # 转亿元
        cash_assets = cash_equivalents / 100000000
        print(f"⚠️ 使用参数文件中的财务数据")
        print(f"   净债务: {net_debt:.2f} 亿元")

    # 生成并插入 DCF 热力图
    add_title(document, '3.2 估值敏感性分析', level=2)

    dcf_chart_path = os.path.join(IMAGES_DIR, '04_dcf_valuation_heatmap.png')
    generate_dcf_valuation_heatmap(dcf_chart_path, project_params['current_price'], net_income, total_shares, net_debt)

    add_paragraph(document, '图表 3: DCF估值敏感性分析矩阵')
    add_image(document, dcf_chart_path, width=Inches(6))

    add_paragraph(document, '分析说明：')
    add_paragraph(document, '• 横轴：WACC（加权平均资本成本），代表折现率')
    add_paragraph(document, '• 纵轴：增长率，代表预测期FCF增长率')
    add_paragraph(document, '• 颜色：绿色表示估值高于当前价，红色表示估值低于当前价')
    add_paragraph(document, f'• 当前股价: {project_params["current_price"]:.2f} 元/股')

    add_title(document, '3.3 估值参数与过程', level=2)

    # 计算基础FCF
    base_fcf = net_income  # FCF = 净利润（保守估计）
    using_net_income_as_fcf = True  # 标记是否使用净利润作为FCF基准

    # 获取历史FCF数据用于DCF估值
    historical_fcf_years = 0
    base_fcf = net_income  # 默认使用净利润作为FCF基准
    fcf_data_description = "使用净利润近似FCF（无历史FCF数据）"
    fcf_cagr = None  # FCF复合增长率（优先使用）

    if 'historical_fcf_data' in project_params and project_params['historical_fcf_data']:
        historical_fcf = project_params['historical_fcf_data']
        historical_fcf_years = historical_fcf['years']

        print(f"\n📊 历史FCF数据分析:")
        print(f"   总年数: {historical_fcf_years}")
        print(f"{'年份':<6} {'营收(亿)':<12} {'NOPAT(亿)':<12} {'FCF(亿)':<12}")
        print("-"*50)

        # 显示所有历史FCF数据
        for year_data in historical_fcf['data']:
            print(f"{year_data['year']:<6} {year_data['revenue']:>10.2f}     "
                  f"{year_data['nopat']:>10.2f}     {year_data['fcf']:>10.2f}")

        # 使用最新一年的真实FCF作为基准
        latest_fcf = historical_fcf['data'][-1]['fcf'] * 100000000  # 转回元

        # 检查FCF是否为正
        print(f"\n   最新FCF: {latest_fcf/100000000:.2f} 亿元")
        if latest_fcf > 0:
            base_fcf = latest_fcf
            using_net_income_as_fcf = False
            print(f"   ✅ 使用最新FCF作为基准")
        else:
            base_fcf = net_income  # 使用净利润作为基准
            using_net_income_as_fcf = True
            print(f"   ⚠️ 最新FCF为负，使用净利润作为基准")

        # 计算历史FCF复合增长率（CAGR）- 优先使用
        # 只计算正FCF年份
        positive_fcf_data = [d for d in historical_fcf['data'] if d['fcf'] > 0]
        print(f"   正值FCF年数: {len(positive_fcf_data)}")

        if len(positive_fcf_data) >= 2:
            # 使用最早和最新的正FCF年计算CAGR
            first_year_data = positive_fcf_data[0]
            last_year_data = positive_fcf_data[-1]

            first_fcf_val = first_year_data['fcf']
            last_fcf_val = last_year_data['fcf']
            years_count = last_year_data['year'] - first_year_data['year']

            # CAGR = (最终值 / 初始值)^(1/年数) - 1
            fcf_cagr = (last_fcf_val / first_fcf_val) ** (1 / years_count) - 1

            # 限制增长率在合理范围（-5%到30%）
            fcf_cagr = max(-0.05, min(0.30, fcf_cagr))

            print(f"   ✅ FCF历史CAGR: {fcf_cagr*100:.2f}%（从{first_year_data['year']}年{first_fcf_val:.2f}亿到{last_year_data['year']}年{last_fcf_val:.2f}亿，{years_count}年）")
            print(f"   ✅ 将使用FCF的CAGR作为增长率基准")
        else:
            fcf_cagr = None
            print(f"   ⚠️ 正值FCF年数少于2年，无法计算FCF的CAGR")
            if net_income_cagr is not None:
                print(f"   ℹ️  将使用净利润的CAGR作为备选: {net_income_cagr*100:.2f}%")

        fcf_data_description = f"使用真实历史FCF数据（{historical_fcf_years}年）"

    # 使用历史CAGR作为FCF增长率，优先级：FCF CAGR > 净利润CAGR > 默认值
    # 初始化FCF增长率变量（需要在表格前初始化）
    fcf_growth_example = 0.10  # 默认10%
    fcf_growth_source = "默认假设（无历史数据）"

    # 优先使用FCF的CAGR
    if fcf_cagr is not None:
        fcf_growth_example = max(fcf_cagr, 0.05)  # 至少5%，避免过于保守
        if fcf_growth_example > 0.30:  # 限制最高30%
            fcf_growth_example = 0.30
        fcf_growth_source = f"基于历史{historical_fcf_years}年FCF数据（CAGR: {fcf_cagr*100:.1f}%）"
        print(f"\n✅ 最终采用FCF增长率: {fcf_growth_example*100:.1f}%（数据源: FCF历史CAGR）")
    # 如果FCF的CAGR不可用，使用净利润的CAGR作为备选
    elif net_income_cagr is not None:
        fcf_growth_example = max(net_income_cagr, 0.05)  # 至少5%
        if fcf_growth_example > 0.30:  # 限制最高30%
            fcf_growth_example = 0.30
        fcf_growth_source = f"基于历史净利润CAGR（CAGR: {net_income_cagr*100:.1f}%）"
        print(f"\n⚠️ 最终采用净利润增长率: {fcf_growth_example*100:.1f}%（数据源: 净利润历史CAGR，因无足够正FCF数据）")

    # 计算真实数据年数和预测年数
    actual_data_years = min(historical_fcf_years, 10) if historical_fcf_years > 0 else 0
    projection_years = max(0, 10 - actual_data_years)

    # 获取真实的最新FCF（如果有的话）
    if 'historical_fcf_data' in project_params and project_params['historical_fcf_data']:
        real_latest_fcf = project_params['historical_fcf_data']['data'][-1]['fcf'] * 100000000
        real_latest_fcf_yuan = real_latest_fcf
    else:
        real_latest_fcf_yuan = None

    # 构建DCF参数表格
    dcf_process_data = [
        ['财务数据（来源：Tushare）', ''],
        ['净利润', f'{net_income/100000000:.2f} 亿元'],
    ]

    # 显示真实的最新FCF和基准FCF
    if real_latest_fcf_yuan is not None:
        real_fcf_display = f'{real_latest_fcf_yuan/100000000:.2f} 亿元'
        if using_net_income_as_fcf:
            real_fcf_display += ' （负值）'
        dcf_process_data.append(['真实最新FCF', real_fcf_display])

    # 显示基准FCF（用于估值计算）
    if using_net_income_as_fcf:
        dcf_process_data.append(['基准FCF（用于计算）', f'{base_fcf/100000000:.2f} 亿元（使用净利润）'])
    else:
        dcf_process_data.append(['基准FCF（用于计算）', f'{base_fcf/100000000:.2f} 亿元'])

    # 继续添加其他财务数据
    dcf_process_data.extend([
        ['净资产', f'{net_assets/100000000:.2f} 亿元'],
        ['总负债', f'{total_debt/100000000:.2f} 亿元'],
        ['货币资金', f'{cash_assets:.2f} 亿元'],
        ['净债务', f'{net_debt:.2f} 亿元'],
        ['总股本', f'{total_shares:,} 股'],
        [''],
        ['DCF假设参数', ''],
        ['数据来源', fcf_data_description],
        ['真实数据年数', f'{actual_data_years} 年'],
        ['预测数据年数', f'{projection_years} 年'],
        ['预测期总计', '10年'],
        ['FCF增长率（预测期）', f'{fcf_growth_example*100:.1f}% ({fcf_growth_source})'],
        ['WACC范围', '8.0% - 13.0%'],
        ['永续增长率', '3.0%']
    ])
    add_table_data(document, ['参数', '值'], dcf_process_data)

    add_paragraph(document, '')
    add_paragraph(document, '估值计算过程（示例）：')

    # 添加FCF计算方法说明
    add_title(document, '3.3.1 自由现金流（FCF）计算方法', level=3)
    add_paragraph(document, '自由现金流（Free Cash Flow, FCF）代表公司在支付了所有运营费用、资本支出后，可分配给股东和债权人的现金。')
    add_paragraph(document, '')
    add_paragraph(document, '本报告采用现金流量表法计算FCF：')
    add_paragraph(document, 'FCF = 经营活动现金流 - 资本支出')
    add_paragraph(document, '')
    add_paragraph(document, '其中：')
    add_paragraph(document, '• 经营活动现金流：来自现金流量表，反映公司主营业务的现金生成能力')
    add_paragraph(document, '• 资本支出：直接使用投资活动现金流，反映购建固定资产、无形资产和其他长期资产支付的现金')
    add_paragraph(document, '')
    add_paragraph(document, '该方法的优势：')
    add_paragraph(document, '1. 直接使用现金流量表数据，避免复杂的营运资本计算')
    add_paragraph(document, '2. 经营活动现金流已自动考虑了营运资本变化')
    add_paragraph(document, '3. 更真实地反映公司的现金生成能力')

    # 添加历史FCF数据分析
    if 'historical_fcf_data' in project_params and project_params['historical_fcf_data']:
        add_title(document, '3.3.2 历史FCF数据分析', level=3)

        historical_fcf = project_params['historical_fcf_data']
        fcf_data_list = historical_fcf['data']

        # 显示最近10年数据
        recent_fcf = fcf_data_list[-10:] if len(fcf_data_list) >= 10 else fcf_data_list

        fcf_history_headers = ['年份', '营业收入（亿）', 'NOPAT（亿）', 'FCF（亿）', 'FCF/营收']
        fcf_history_data = []
        total_fcf_to_revenue = 0
        valid_years = 0

        for year_data in recent_fcf:
            year = year_data['year']
            revenue = year_data['revenue']
            nopat = year_data['nopat']
            fcf = year_data['fcf']

            if revenue > 0:
                fcf_ratio = (fcf / revenue) * 100
                total_fcf_to_revenue += fcf_ratio
                valid_years += 1
            else:
                fcf_ratio = 0

            fcf_history_data.append([
                str(year),
                f'{revenue:.2f}',
                f'{nopat:.2f}',
                f'{fcf:.2f}',
                f'{fcf_ratio:.1f}%'
            ])

        add_table_data(document, fcf_history_headers, fcf_history_data)

        add_paragraph(document, '')
        if valid_years > 0:
            avg_fcf_ratio = total_fcf_to_revenue / valid_years
            add_paragraph(document, f'历史平均FCF/营收比率：{avg_fcf_ratio:.1f}%')
            add_paragraph(document, '该指标反映公司每100元营业收入能产生多少自由现金流。')

        # FCF趋势分析
        add_paragraph(document, '')
        add_paragraph(document, 'FCF趋势分析：')
        if len(recent_fcf) >= 3:
            recent_3_years = recent_fcf[-3:]
            avg_fcf_3y = sum(d['fcf'] for d in recent_3_years) / 3
            add_paragraph(document, f'• 最近3年平均FCF：{avg_fcf_3y:.2f} 亿元')

        latest_fcf = recent_fcf[-1]['fcf']
        add_paragraph(document, f'• 最新年度FCF（{recent_fcf[-1]["year"]}年）：{latest_fcf:.2f} 亿元')

        if latest_fcf > 0:
            add_paragraph(document, '• ✅ FCF为正值，说明公司具备良好的现金生成能力')
        else:
            add_paragraph(document, '• ⚠️ FCF为负值，可能由于大额资本支出或营运资金占用')

    wacc_example = 0.10  # 10%
    growth_example = 0.025  # 2.5% 永续增长率

    # 计算示例 - 使用真实FCF数据（如果有）
    fcfs = []
    fcf_sources = []  # 记录每个FCF的来源
    if actual_data_years > 0 and 'historical_fcf_data' in project_params:
        # 使用最新历史FCF数据作为基准，正向预测未来10年
        historical_fcf = project_params['historical_fcf_data']['data']
        latest_fcf = historical_fcf[-1]['fcf'] * 100000000  # 最新一年（2025年）FCF，转回元
        latest_year = historical_fcf[-1]['year']

        # 计算历史FCF复利增长率（CAGR）- 仅用于报告显示
        # 使用最早的正FCF年份和最新的正FCF年份计算
        positive_fcf_years = [d for d in historical_fcf if d['fcf'] > 0]

        if len(positive_fcf_years) >= 2:
            first_fcf_year = positive_fcf_years[0]
            last_fcf_year = positive_fcf_years[-1]

            first_fcf = first_fcf_year['fcf']
            last_fcf = last_fcf_year['fcf']
            years_span = last_fcf_year['year'] - first_fcf_year['year']

            # CAGR = (最终值 / 初始值)^(1/年数) - 1
            cagr_fcf_display = (last_fcf / first_fcf) ** (1 / years_span) - 1
        else:
            cagr_fcf_display = 0.05  # 默认5%（仅用于显示）

        # 注意：实际使用 fcf_growth_example（已在前面计算，优先使用FCF CAGR）
        # 这里不重新计算，保持与前面的一致性

        add_paragraph(document, f'✅ 使用最新历史数据（{latest_year}年FCF {latest_fcf/100000000:.2f}亿元）作为基准，正向预测未来10年')
        add_paragraph(document, '')
        add_paragraph(document, f'📊 历史FCF增长率分析：')

        # 显示最近5年FCF数据
        recent_5 = historical_fcf[-5:]
        fcf_table_data = []
        for i, item in enumerate(recent_5):
            fcf_table_data.append([f"{item['year']}年", f"{item['fcf']:.2f}"])
        add_table_data(document, ['年份', 'FCF（亿）'], fcf_table_data)

        add_paragraph(document, f'• 历史CAGR: {cagr_fcf_display*100:.2f}%（从{first_fcf_year["year"]}年{first_fcf:.2f}亿到{last_fcf_year["year"]}年{last_fcf:.2f}亿，跨度{years_span}年）')
        add_paragraph(document, f'• 采用预测增长率: {fcf_growth_example*100:.1f}%（基于历史CAGR，优先使用FCF数据计算）')

        # 基于最新FCF，预测未来10年（使用历史增长率）
        for i in range(10):
            fcf_proj = latest_fcf * ((1 + fcf_growth_example) ** (i + 1))
            fcfs.append(fcf_proj)
            if i == 0:
                fcf_sources.append(f"基准年（{latest_year}年实际）")
            else:
                fcf_sources.append(f"预测（增长{fcf_growth_example*100:.1f}%）")
    else:
        # 全部使用预测
        fcfs = [base_fcf * ((1 + fcf_growth_example) ** i) for i in range(10)]
        fcf_sources = [f"预测（增长{fcf_growth_example*100:.1f}%）" for _ in range(10)]

    # 添加逐年FCF预测和折现计算
    add_title(document, '3.3.3 逐年FCF预测与折现计算', level=3)

    add_paragraph(document, f'基于WACC={wacc_example*100:.1f}%，预测期FCF增长率={fcf_growth_example*100:.1f}%，永续增长率={growth_example*100:.1f}%：')
    add_paragraph(document, '')

    # 构建逐年计算表格
    dcf_year_headers = ['年份', 'FCF（亿）', '折现因子', '现值（亿）', '数据来源']
    dcf_year_data = []

    pv_fcfs_detail = 0
    for i, fcf in enumerate(fcfs):
        discount_factor = 1 / ((1 + wacc_example) ** (i+1))
        pv = fcf * discount_factor
        pv_fcfs_detail += pv
        year_num = i + 1

        dcf_year_data.append([
            str(year_num),
            f'{fcf/100000000:.2f}',
            f'{discount_factor:.4f}',
            f'{pv/100000000:.2f}',
            fcf_sources[i] if i < len(fcf_sources) else '预测'
        ])

    add_table_data(document, dcf_year_headers, dcf_year_data)

    add_paragraph(document, '')
    add_paragraph(document, f'预测期FCF现值合计：{pv_fcfs_detail/100000000:.2f} 亿元')

    # 添加终值计算详细过程
    add_title(document, '3.3.4 终值（Terminal Value）计算', level=3)

    terminal_fcf = fcfs[-1] * (1 + growth_example)
    terminal_value = terminal_fcf / (wacc_example - growth_example)
    pv_terminal = terminal_value / ((1 + wacc_example) ** 10)

    add_paragraph(document, '终值代表预测期后公司持续经营的价值，采用永续增长模型计算：')
    add_paragraph(document, '')
    add_paragraph(document, '计算公式：')
    add_paragraph(document, '终值 = 第10年末FCF × (1 + 永续增长率) / (WACC - 永续增长率)')
    add_paragraph(document, '')
    add_paragraph(document, '计算步骤：')
    add_paragraph(document, f'1. 第10年末FCF = {fcfs[-1]/100000000:.2f} 亿元')
    add_paragraph(document, f'2. 终值FCF = 第10年末FCF × (1 + {growth_example*100:.1f}%)')
    add_paragraph(document, f'          = {fcfs[-1]/100000000:.2f} × {1 + growth_example:.4f}')
    add_paragraph(document, f'          = {terminal_fcf/100000000:.2f} 亿元')
    add_paragraph(document, f'3. 终值 = {terminal_fcf/100000000:.2f} / ({wacc_example*100:.1f}% - {growth_example*100:.1f}%)')
    add_paragraph(document, f'        = {terminal_value/100000000:.2f} 亿元')
    add_paragraph(document, f'4. 终值现值 = {terminal_value/100000000:.2f} / (1 + {wacc_example*100:.1f}%)^10')
    add_paragraph(document, f'          = {pv_terminal/100000000:.2f} 亿元')

    # 添加企业价值和股权价值计算
    add_title(document, '3.3.5 企业价值与股权价值计算', level=3)

    enterprise_value = pv_fcfs_detail + pv_terminal  # 使用详细计算的现值

    add_paragraph(document, '企业价值（Enterprise Value）计算：')
    add_paragraph(document, '')
    add_paragraph(document, f'企业价值 = 预测期FCF现值 + 终值现值')
    add_paragraph(document, f'        = {pv_fcfs_detail/100000000:.2f} + {pv_terminal/100000000:.2f}')
    add_paragraph(document, f'        = {enterprise_value/100000000:.2f} 亿元')
    add_paragraph(document, '')
    add_paragraph(document, '股权价值（Equity Value）计算：')
    add_paragraph(document, '')
    add_paragraph(document, f'股权价值 = 企业价值 - 净债务')
    add_paragraph(document, f'        = {enterprise_value/100000000:.2f} - {net_debt:.2f}')

    # 修正：股权价值 = 企业价值 - 净债务（注意单位转换）
    equity_value = enterprise_value - (net_debt * 100000000)  # 净债务从亿元转为元

    add_paragraph(document, f'        = {equity_value/100000000:.2f} 亿元')
    add_paragraph(document, '')
    add_paragraph(document, '内在价值（Intrinsic Value）计算：')
    add_paragraph(document, '')
    add_paragraph(document, f'内在价值 = 股权价值 / 总股本')
    add_paragraph(document, f'        = {equity_value/100000000:.2f} 亿元 / {total_shares/100000000:.2f} 亿股')

    intrinsic_value = equity_value / total_shares  # 单位：元/股

    add_paragraph(document, f'        = {intrinsic_value:.2f} 元/股')

    # 添加计算汇总表格
    add_paragraph(document, '')
    add_paragraph(document, '估值计算汇总：')

    example_data = [
        ['WACC', f'{wacc_example*100:.1f}%'],
        ['FCF增长率（预测期）', f'{fcf_growth_example*100:.1f}%'],
        ['（数据来源）', fcf_growth_source],
        ['真实数据年数', f'{actual_data_years} 年'],
        ['预测数据年数', f'{projection_years} 年'],
        ['永续增长率', f'{growth_example*100:.1f}%'],
        ['预测期FCF现值', f'{pv_fcfs_detail/100000000:.2f} 亿元'],
        ['终值', f'{terminal_value/100000000:.2f} 亿元'],
        ['终值现值', f'{pv_terminal/100000000:.2f} 亿元'],
        ['企业价值', f'{enterprise_value/100000000:.2f} 亿元'],
        ['净债务', f'{net_debt:.2f} 亿元'],
        ['股权价值', f'{equity_value/100000000:.2f} 亿元'],
        ['每股价值', f'{intrinsic_value:.2f} 元/股']
    ]
    add_table_data(document, ['项目', '计算结果'], example_data)

    # 添加简化估算方法：使用净利润近似FCF
    # ====== 3.3.6章节已隐藏 ======
    # add_title(document, '3.3.6 简化估算方法：净利润法', level=3)

    # add_paragraph(document, '为快速评估公司价值，本节提供一种简化的DCF估算方法，直接使用净利润代替FCF进行估值。')
    # add_paragraph(document, '')
    # add_paragraph(document, '方法说明：')
    # add_paragraph(document, '• 假设：净利润 ≈ FCF（适用于资本支出较少、营运资本变化稳定的成熟企业）')
    # add_paragraph(document, '• 优势：计算简单，数据易得，适合快速估值')
    # add_paragraph(document, '• 局限：未考虑资本支出和营运资本变化，可能高估或低估企业价值')
    # add_paragraph(document, '')

    # # 使用历史净利润增长率（如果有的话）
    # net_income_growth = fcf_growth_example  # 默认使用FCF增长率
    # growth_source_net = fcf_growth_source

    # if historical_incomes and len(historical_incomes) >= 2:
    #     # 计算净利润CAGR
    #     start_income = historical_incomes[-1]
    #     end_income = historical_incomes[0]
    #     n_years = len(historical_incomes) - 1

    #     if start_income > 0 and end_income > 0:
    #         net_income_cagr = (end_income / start_income) ** (1 / n_years) - 1
    #         # 限制在合理范围
    #         net_income_growth = max(min(net_income_cagr, 0.30), 0.05)
    #         growth_source_net = f"基于历史{n_years}年净利润CAGR"

    # add_paragraph(document, f'简化估算参数（WACC={wacc_example*100:.1f}%，永续增长率={growth_example*100:.1f}%）：')
    # add_paragraph(document, '')

    # # 使用净利润作为基准FCF
    # base_fcf_simple = net_income  # 直接使用净利润

    # # 逐年计算（全部使用预测，因为净利润通常是年度数据）
    # fcfs_simple = [base_fcf_simple * ((1 + net_income_growth) ** i) for i in range(10)]

    # # 计算现值
    # pv_fcfs_simple = sum([fcf / ((1 + wacc_example) ** (i+1)) for i, fcf in enumerate(fcfs_simple)])

    # # 计算终值
    # terminal_fcf_simple = fcfs_simple[-1] * (1 + growth_example)
    # terminal_value_simple = terminal_fcf_simple / (wacc_example - growth_example)
    # pv_terminal_simple = terminal_value_simple / ((1 + wacc_example) ** 10)

    # # 企业价值
    # enterprise_value_simple = pv_fcfs_simple + pv_terminal_simple

    # # 股权价值
    # equity_value_simple = enterprise_value_simple - (net_debt * 100000000)

    # # 内在价值
    # intrinsic_value_simple = equity_value_simple / total_shares

    # # 构建简化估算结果表格
    # simple_data = [
    #     ['基准净利润', f'{base_fcf_simple/100000000:.2f} 亿元'],
    #     ['净利润增长率', f'{net_income_growth*100:.1f}%'],
    #     ['增长率来源', growth_source_net],
    #     ['预测期现值', f'{pv_fcfs_simple/100000000:.2f} 亿元'],
    #     ['终值', f'{terminal_value_simple/100000000:.2f} 亿元'],
    #     ['终值现值', f'{pv_terminal_simple/100000000:.2f} 亿元'],
    #     ['企业价值', f'{enterprise_value_simple/100000000:.2f} 亿元'],
    #     ['股权价值', f'{equity_value_simple/100000000:.2f} 亿元'],
    #     ['内在价值', f'{intrinsic_value_simple:.2f} 元/股']
    # ]
    # add_table_data(document, ['项目', '计算结果'], simple_data)

    # add_paragraph(document, '')
    # add_paragraph(document, '与详细FCF法对比：')

    # # 对比分析
    # comparison_headers = ['估值方法', '内在价值（元/股）', '差异', '差异率']
    # comparison_data = [
    #     [
    #         '详细FCF法（使用现金流量表）',
    #         f'{intrinsic_value:.2f}',
    #         '-',
    #         '-'
    #     ],
    #     [
    #         '简化净利润法',
    #         f'{intrinsic_value_simple:.2f}',
    #         f'{intrinsic_value_simple - intrinsic_value:+.2f}',
    #         f'{(intrinsic_value_simple/intrinsic_value - 1)*100:+.1f}%'
    #     ]
    # ]
    # add_table_data(document, comparison_headers, comparison_data)

    # add_paragraph(document, '')
    # add_paragraph(document, '对比分析：')

    # if intrinsic_value_simple > intrinsic_value:
    #     diff_pct = (intrinsic_value_simple / intrinsic_value - 1) * 100
    #     if diff_pct > 20:
    #         analysis = f"简化法估值高出详细法{diff_pct:.1f}%，说明公司资本支出较少或营运资本变化较小，净利润法较为适用。"
    #     elif diff_pct > 0:
    #         analysis = f"简化法估值高出详细法{diff_pct:.1f}%，差异较小，两种方法得出相近的结论。"
    #     else:
    #         analysis = f"简化法估值与详细法基本一致，验证了估值的可靠性。"
    # else:
    #     diff_pct = (intrinsic_value / intrinsic_value_simple - 1) * 100
    #     if diff_pct > 20:
    #         analysis = f"简化法估值低于详细法{diff_pct:.1f}%，说明公司资本支出较多或营运资本增加较快，应优先使用详细FCF法。"
    #     else:
    #         analysis = f"简化法估值低于详细法{diff_pct:.1f}%，差异在可接受范围内。"

    # add_paragraph(document, analysis)
    # add_paragraph(document, '')

    # # 方法选择建议
    # add_paragraph(document, '方法选择建议：')
    # add_paragraph(document, '• 优先使用详细FCF法：当有完整现金流量表数据时，结果更准确')
    # add_paragraph(document, '• 参考简化净利润法：适用于快速估值或现金流量表数据缺失时')
    # add_paragraph(document, '• 综合判断：当两种方法结果差异较小时，可增强估值结论的可信度')
    # add_paragraph(document, '• 谨慎对待：当两种方法结果差异较大（超过30%）时，需要深入分析原因')
    # ====== 3.3.6章节隐藏结束 ======

    add_title(document, '3.4 估值结论', level=2)

    margin_market = (intrinsic_value - project_params['current_price']) / project_params['current_price'] * 100
    margin_issue = (intrinsic_value - project_params['issue_price']) / project_params['issue_price'] * 100

    add_paragraph(document, f'• DCF内在价值: {intrinsic_value:.2f} 元/股')
    add_paragraph(document, f'• 当前价格: {project_params["current_price"]:.2f} 元/股')
    add_paragraph(document, f'• 发行价格: {project_params["issue_price"]:.2f} 元/股')
    add_paragraph(document, f'• 相对市价安全边际: {margin_market:+.1f}%')
    add_paragraph(document, f'• 相对发行价安全边际: {margin_issue:+.1f}%')

    if margin_issue > 15:
        conclusion = "DCF估值显示，相比发行价有显著安全边际，估值合理偏低。"
    elif margin_issue > 0:
        conclusion = "DCF估值显示，相比发行价有一定安全边际，估值相对合理。"
    else:
        conclusion = "DCF估值显示，发行价高于内在价值，需谨慎。"

    add_paragraph(document, '')
    add_paragraph(document, f'估值结论: {conclusion}')

    add_section_break(document)

    # ==================== 四、敏感性分析 ====================
    add_title(document, '四、敏感性分析', level=1)

    add_paragraph(document, '本章节分析单一参数变化对定增项目盈利概率的影响，包括发行价折扣、漂移率、波动率、锁定期等关键参数的敏感性分析。')
    add_paragraph(document, '所有敏感性分析统一基于120日（半年线）窗口，通过调整单一变量、保持其他参数不变的方法来测试敏感性。')

    # ==================== 4.1 不同时间窗口风险指标分析 ====================
    add_title(document, '4.1 不同时间窗口风险指标分析', level=2)

    add_paragraph(document, '本章节分析不同时间窗口对风险指标的影响，包括波动率、收益率、最大回撤等。')
    add_paragraph(document, '通常情况下，时间窗口越短，波动率越高；时间窗口越长，数据越稳定。')
    add_paragraph(document, '为统一分析基准，本报告后续所有敏感性分析均采用120日（半年线）窗口。')
    add_paragraph(document, '')

    # 获取历史价格数据并计算不同窗口的指标
    try:
        # 优先使用已保存的market_data文件中的数据
        if market_data and 'total_days' in market_data:
            total_days = market_data['total_days']
            print(f"✅ 使用已保存的市场数据（共{total_days}个交易日）")

            # 从market_data中提取时间窗口数据
            # 构建time_window_results格式
            time_window_results = {
                'window': [20, 60, 120, 250],
                'volatility': [
                    market_data.get('volatility_20d', market_data.get('volatility', 0)),
                    market_data.get('volatility_60d', market_data.get('volatility', 0)),
                    market_data.get('volatility_120d', market_data.get('volatility', 0)),
                    market_data.get('volatility_250d', market_data.get('volatility', 0))
                ],
                'total_return': [
                    market_data.get('period_return_20d', 0),
                    market_data.get('period_return_60d', 0),
                    market_data.get('period_return_120d', 0),
                    market_data.get('period_return_250d', 0)
                ],
                'annual_return': [
                    market_data.get('annual_return_20d', 0),
                    market_data.get('annual_return_60d', 0),
                    market_data.get('annual_return_120d', 0),
                    market_data.get('annual_return_250d', 0)
                ],
                'mean': [0] * 4,  # 暂不计算
                'std': [0] * 4,   # 暂不计算
                'max_drawdown': [0] * 4,  # 暂不计算
                'sharpe': [0] * 4  # 暂不计算
            }

            # 过滤掉没有数据的窗口
            valid_indices = [i for i in range(4) if time_window_results['volatility'][i] > 0]

            if len(valid_indices) == 0:
                add_paragraph(document, '⚠️ 当前历史数据不足以进行时间窗口分析')
            else:
                # 只保留有数据的窗口
                filtered_results = {
                    'window': [time_window_results['window'][i] for i in valid_indices],
                    'volatility': [time_window_results['volatility'][i] for i in valid_indices],
                    'total_return': [time_window_results['total_return'][i] for i in valid_indices],
                    'annual_return': [time_window_results['annual_return'][i] for i in valid_indices],
                    'max_drawdown': [time_window_results['max_drawdown'][i] for i in valid_indices],
                    'sharpe': [time_window_results['sharpe'][i] for i in valid_indices]
                }

                # 添加时间窗口数据表格
                window_data = []
                for i, window in enumerate(filtered_results['window']):
                    window_data.append([
                        f'{window}日',
                        f'{filtered_results["volatility"][i]*100:.2f}%',
                        f'{filtered_results["total_return"][i]*100:+.2f}%',
                        f'{filtered_results["annual_return"][i]*100:+.2f}%',
                        f'{filtered_results["max_drawdown"][i]*100:.2f}%',
                        f'{filtered_results["sharpe"][i]:.2f}'
                    ])
                add_table_data(document, ['时间窗口', '年化波动率', '期间收益率', '年化收益率', '最大回撤', '夏普比率'], window_data)

                # 生成图表（如果有完整数据）
                if len(filtered_results['window']) >= 4:
                    # 需要price_series来生成图表，这里简化处理，不生成图表
                    add_paragraph(document, '')
                    add_paragraph(document, '💡 说明：')
                    add_paragraph(document, '• 时间窗口分析基于已保存的市场数据文件')
                    add_paragraph(document, '• 数据来源：市场数据文件（自动更新）')

                # 分析结论
                add_paragraph(document, '')
                add_paragraph(document, '时间窗口分析结论：')

                # 检查是否有足够的数据窗口
                if len(filtered_results['window']) >= 4:
                    vol_20 = filtered_results["volatility"][0]
                    vol_60 = filtered_results["volatility"][1]
                    vol_120 = filtered_results["volatility"][2]
                    vol_250 = filtered_results["volatility"][3]

                    add_paragraph(document, f'• 波动率随时间窗口变化：')
                    add_paragraph(document, f'  - 20日窗口: {vol_20*100:.2f}%（短期波动较大）')
                    add_paragraph(document, f'  - 60日窗口: {vol_60*100:.2f}%（季度窗口）')
                    add_paragraph(document, f'  - 120日窗口: {vol_120*100:.2f}%（半年窗口，本报告统一采用）')
                    add_paragraph(document, f'  - 250日窗口: {vol_250*100:.2f}%（年度窗口）')

                    # 收益率分析
                    ret_60 = filtered_results["annual_return"][1]
                    ret_120 = filtered_results["annual_return"][2]
                    ret_250 = filtered_results["annual_return"][3]

                    add_paragraph(document, f'• 年化收益率随时间窗口变化：')
                    add_paragraph(document, f'  - 60日窗口: {ret_60*100:+.2f}%')
                    add_paragraph(document, f'  - 120日窗口: {ret_120*100:+.2f}%')
                    add_paragraph(document, f'  - 250日窗口: {ret_250*100:+.2f}%')
                    add_paragraph(document, f'  - 结论：120日窗口平衡了数据稳定性和时效性，更能反映中期趋势')

                    # 风险提示
                    if vol_20 > vol_60 * 1.3:
                        add_paragraph(document, f'• ⚠️ 短期波动率显著高于长期，需注意短期风险')
                    else:
                        add_paragraph(document, f'• ✅ 波动率相对稳定，市场较为理性')
                else:
                    # 有部分数据
                    add_paragraph(document, f'ℹ️ 基于当前可用的{len(filtered_results["window"])}个时间窗口进行分析：')

                    available_windows = filtered_results["window"]
                    window_names = {20: "20日(月度)", 60: "60日(季度)", 120: "120日(半年)", 250: "250日(年度)"}
                    window_list = "、".join([window_names.get(w, f"{w}日") for w in available_windows])
                    add_paragraph(document, f'• 可用窗口: {window_list}')

        elif os.environ.get('TUSHARE_TOKEN'):
            # 如果没有market_data，但有token，尝试从API获取
            import tushare as ts

            pro = ts.pro_api(ts_token)

            # 获取历史价格数据（最近2年，确保有足够数据计算250日窗口）
            end_date = datetime.now().strftime('%Y%m%d')
            start_date = (datetime.now() - timedelta(days=730)).strftime('%Y%m%d')

            df_prices = pro.daily(ts_code=stock_code, start_date=start_date, end_date=end_date)

            if not df_prices.empty and len(df_prices) > 180:
                # 按日期排序，确保数据按时间顺序排列
                df_prices = df_prices.sort_values('trade_date').reset_index(drop=True)
                # 使用收盘价
                price_series = df_prices['close'].reset_index(drop=True)

                print(f"✅ 从API获取了{len(price_series)}个交易日的数据")

                # 生成时间窗口分析图表
                time_window_results, chart_paths = generate_time_window_analysis_chart(price_series, IMAGES_DIR)

                # 检查是否有有效的时间窗口数据
                if len(time_window_results['window']) == 0:
                    add_paragraph(document, '⚠️ 当前历史数据不足以进行时间窗口分析')
                elif len(time_window_results['window']) < 4:
                    # 有部分窗口数据，但不是全部
                    add_paragraph(document, f'ℹ️ 当前历史数据支持{len(time_window_results["window"])}个时间窗口的分析（完整分析需要4个窗口）')

                    # 添加时间窗口数据表格
                    window_data = []
                    for i, window in enumerate(time_window_results['window']):
                        window_data.append([
                            f'{window}日',
                            f'{time_window_results["volatility"][i]*100:.2f}%',
                            f'{time_window_results["total_return"][i]*100:+.2f}%',
                            f'{time_window_results["annual_return"][i]*100:+.2f}%',
                            f'{time_window_results["max_drawdown"][i]*100:.2f}%',
                            f'{time_window_results["sharpe"][i]:.2f}'
                        ])
                    add_table_data(document, ['时间窗口', '年化波动率', '期间收益率', '年化收益率', '最大回撤', '夏普比率'], window_data)

                    # 添加图表（如果有生成的话）
                    if len(chart_paths) >= 3:
                        add_paragraph(document, '')
                        add_paragraph(document, '图表 4.1: 不同时间窗口的波动率对比')
                        add_image(document, chart_paths[0], width=Inches(6.5))

                        add_paragraph(document, '')
                        add_paragraph(document, '图表 4.2: 不同时间窗口的收益率对比')
                        add_image(document, chart_paths[1], width=Inches(6.5))
                        add_paragraph(document, '')

                        add_paragraph(document, '图表 4.3: 不同时间窗口的风险指标对比')
                        add_image(document, chart_paths[2], width=Inches(6.5))
                        add_paragraph(document, '')
                else:
                    # 添加时间窗口数据表格
                    window_data = []
                    for i, window in enumerate(time_window_results['window']):
                        window_data.append([
                            f'{window}日',
                            f'{time_window_results["volatility"][i]*100:.2f}%',
                            f'{time_window_results["total_return"][i]*100:+.2f}%',
                            f'{time_window_results["annual_return"][i]*100:+.2f}%',
                            f'{time_window_results["max_drawdown"][i]*100:.2f}%',
                            f'{time_window_results["sharpe"][i]:.2f}'
                        ])
                    add_table_data(document, ['时间窗口', '年化波动率', '期间收益率', '年化收益率', '最大回撤', '夏普比率'], window_data)

                    # 添加图表（如果有生成的话）
                    if len(chart_paths) >= 3:
                        add_paragraph(document, '')
                        add_paragraph(document, '图表 4.1: 不同时间窗口的波动率对比')
                        add_image(document, chart_paths[0], width=Inches(6.5))

                        add_paragraph(document, '')
                        add_paragraph(document, '图表 4.2: 不同时间窗口的收益率对比')
                        add_image(document, chart_paths[1], width=Inches(6.5))

                        add_paragraph(document, '')
                        add_paragraph(document, '图表 4.3: 不同时间窗口的风险指标对比')
                        add_image(document, chart_paths[2], width=Inches(6.5))

                # 分析结论
                add_paragraph(document, '')
                add_paragraph(document, '时间窗口分析结论：')

                # 检查是否有足够的数据窗口
                if len(time_window_results["volatility"]) >= 4:
                    # 对比不同窗口的波动率
                    # windows = [20, 60, 120, 250]
                    vol_20 = time_window_results["volatility"][0]
                    vol_60 = time_window_results["volatility"][1]
                    vol_120 = time_window_results["volatility"][2]
                    vol_250 = time_window_results["volatility"][3]

                    add_paragraph(document, f'• 波动率随时间窗口变化：')
                    add_paragraph(document, f'  - 20日窗口: {vol_20*100:.2f}%（短期波动较大）')
                    add_paragraph(document, f'  - 60日窗口: {vol_60*100:.2f}%（季度窗口）')
                    add_paragraph(document, f'  - 120日窗口: {vol_120*100:.2f}%（半年窗口，本报告统一采用）')
                    add_paragraph(document, f'  - 250日窗口: {vol_250*100:.2f}%（年度窗口）')

                    # 收益率分析
                    ret_60 = time_window_results["annual_return"][1]
                    ret_120 = time_window_results["annual_return"][2]
                    ret_250 = time_window_results["annual_return"][3]

                    add_paragraph(document, f'• 年化收益率随时间窗口变化：')
                    add_paragraph(document, f'  - 60日窗口: {ret_60*100:+.2f}%')
                    add_paragraph(document, f'  - 120日窗口: {ret_120*100:+.2f}%')
                    add_paragraph(document, f'  - 250日窗口: {ret_250*100:+.2f}%')
                    add_paragraph(document, f'  - 结论：120日窗口平衡了数据稳定性和时效性，更能反映中期趋势')

                    # 风险提示
                    if vol_20 > vol_60 * 1.3:
                        add_paragraph(document, f'• ⚠️ 短期波动率显著高于长期，需注意短期风险')
                    else:
                        add_paragraph(document, f'• ✅ 波动率相对稳定，市场较为理性')
                elif len(time_window_results["volatility"]) >= 2:
                    # 有部分数据，显示部分分析结论
                    add_paragraph(document, f'ℹ️ 基于当前可用的{len(time_window_results["volatility"])}个时间窗口进行分析：')

                    # 显示可用的窗口信息
                    available_windows = time_window_results["window"]
                    window_names = {20: "20日(月度)", 60: "60日(季度)", 120: "120日(半年)", 250: "250日(年度)"}
                    window_list = "、".join([window_names.get(w, f"{w}日") for w in available_windows])
                    add_paragraph(document, f'• 可用窗口: {window_list}')

                    # 显示波动率对比（如果有多个窗口）
                    if len(time_window_results["volatility"]) >= 2:
                        vols = [f"{time_window_results['volatility'][i]*100:.2f}%" for i in range(len(time_window_results["volatility"]))]
                        windows_str = "、".join([window_names.get(available_windows[i], str(available_windows[i])) for i in range(len(available_windows))])
                        add_paragraph(document, f'• 波动率对比（{windows_str}）: {"、".join(vols)}')

                    add_paragraph(document, f'• 建议：获取更多历史数据（至少2年）以支持完整的时间窗口分析')
                else:
                    add_paragraph(document, f'⚠️ 当前历史数据不足以进行时间窗口分析（需要至少20个交易日数据）')
                    add_paragraph(document, f'   当前可用窗口数: {len(time_window_results["volatility"])}')

            else:
                add_paragraph(document, '⚠️ 暂无足够的历史数据进行时间窗口分析')
        else:
            add_paragraph(document, '⚠️ 未设置TUSHARE_TOKEN，无法获取历史数据')
    except Exception as e:
        print(f"时间窗口分析失败: {e}")
        add_paragraph(document, '⚠️ 时间窗口分析暂时无法执行')

    # ==================== 4.2 发行价折扣敏感性分析 ====================
    add_title(document, '4.2 发行价折扣敏感性分析', level=2)

    add_paragraph(document, '本节分析不同发行价情景下的盈利概率和预期收益率，统一分析从折价20%到溢价20%的所有情景。')
    add_paragraph(document, '')
    add_paragraph(document, '发行类型定义（基于20日均线MA20）：')
    add_paragraph(document, '• 折价发行：发行价 < MA20，有安全边际，盈利阈值 = 发行价')
    add_paragraph(document, '• 平价发行：发行价 = MA20')
    add_paragraph(document, '• 溢价发行：发行价 > MA20，无安全边际，盈利阈值 = max(MA20, 发行价×1.02)')
    add_paragraph(document, '')

    # 生成发行价折扣情景图表（统一版）
    scenario_chart_paths = generate_discount_scenario_charts_split(
        ma120, project_params['current_price'], risk_params['volatility'],
        risk_params['drift'], project_params['lockup_period'], IMAGES_DIR)

    # 添加统一图表
    add_paragraph(document, '')
    add_paragraph(document, '图表 4.4: 发行价敏感性分析 - 盈利概率（-20%至+20%）')
    add_image(document, scenario_chart_paths[0], width=Inches(6.5))
    add_paragraph(document, '')

    add_paragraph(document, '图表 4.5: 发行价敏感性分析 - 预期收益率（-20%至+20%）')
    add_image(document, scenario_chart_paths[1], width=Inches(6.5))
    add_paragraph(document, '')

    # 参数说明
    add_paragraph(document, '')
    add_paragraph(document, '💡 图表说明：')
    add_paragraph(document, '• 横轴：发行价相对MA20的折扣/溢价率（负值=折价，正值=溢价）')
    add_paragraph(document, '• 绿色柱：折价发行（<0%），有安全边际')
    add_paragraph(document, '• 灰色柱：平价发行（=0%）')
    add_paragraph(document, '• 红色柱：溢价发行（>0%），无安全边际')
    add_paragraph(document, '• 蓝色虚线：MA20基准线（0%处）')
    add_paragraph(document, '')

    # 添加参数说明
    add_paragraph(document, '')
    add_paragraph(document, '💡 计算参数：')
    add_paragraph(document, f'• 数据窗口：120日（半年线，统一分析基准）')
    add_paragraph(document, f'• 年化波动率：{risk_params["volatility"]*100:.2f}%')
    add_paragraph(document, f'• 年化漂移率：{risk_params["drift"]*100:.2f}%')
    add_paragraph(document, f'• 锁定期：{project_params["lockup_period"]}个月')
    add_paragraph(document, '')

    # 分析结论
    add_paragraph(document, '')
    add_paragraph(document, '发行价折扣情景分析结论：')

    # 计算当前发行价的折扣（负值表示折价，正值表示溢价）
    current_discount = (project_params['issue_price'] - ma120) / ma120 * 100
    add_paragraph(document, f'• 当前发行价: {project_params["issue_price"]:.2f} 元')
    add_paragraph(document, f'• MA20基准价: {ma120:.2f} 元')
    add_paragraph(document, f'• 当前折价/溢价率: {current_discount:.1f}%')

    if current_discount <= -20:
        discount_level = "非常充足（≤-20%）"
        safety_note = "安全边际很高，投资吸引力强"
    elif current_discount <= -15:
        discount_level = "充足（-20%至-15%）"
        safety_note = "安全边际良好，投资吸引力较高"
    elif current_discount <= -10:
        discount_level = "适中（-15%至-10%）"
        safety_note = "安全边际一般，需综合考虑其他因素"
    elif current_discount <= -5:
        discount_level = "较低（-10%至-5%）"
        safety_note = "安全边际有限，需谨慎评估"
    elif current_discount < 0:
        discount_level = "偏低（-5%至0%）"
        safety_note = "安全边际较小"
    elif current_discount == 0:
        discount_level = "平价（0%）"
        safety_note = "无折价也无溢价"
    else:
        discount_level = "溢价（>0%）"
        safety_note = "无安全边际，需重点关注"

    add_paragraph(document, f'• 折价水平评估: {discount_level}，{safety_note}')
    add_paragraph(document, '• 折价越多，盈利概率越高，但融资额越少')
    add_paragraph(document, '• 溢价越多，盈利概率越低，破发风险越大')

    # ==================== 4.3 漂移率敏感性分析 ====================
    add_paragraph(document, '')
    add_title(document, '4.3 漂移率敏感性分析', level=2)

    add_paragraph(document, '漂移率（Drift Rate）代表股价的年化预期收益率，反映股票的长期趋势。')
    add_paragraph(document, '本节分析不同漂移率假设对定增项目盈利概率的影响，帮助判断在乐观、中性、悲观情景下的投资表现。')
    add_paragraph(document, '')

    # 定义不同的漂移率情景
    drift_scenarios = [
        (-0.30, "极端悲观"),
        (-0.15, "悲观"),
        (-0.05, "轻度悲观"),
        (0.00, "中性（无增长）"),
        (0.05, "轻度乐观"),
        (0.10, "乐观"),
        (0.15, "非常乐观"),
        (0.20, "极端乐观")
    ]

    # 使用120日窗口的波动率
    current_vol = market_data.get('volatility_120d', risk_params.get('volatility', 0.30))
    lockup_months = project_params['lockup_period']

    drift_analysis_data = []
    expected_returns_list = []  # 收集期望收益率用于计算增幅

    # 定义盈利目标
    profit_targets_pct = [0, 5, 10, 15, 20]

    for drift, scenario_name in drift_scenarios:
        # 计算锁定期的漂移率和波动率
        lockup_drift = drift * (lockup_months / 12)
        lockup_vol = current_vol * np.sqrt(lockup_months / 12)

        # 计算盈利阈值价格
        issue_price = project_params['issue_price']
        current_price = project_params['current_price']

        # 判断发行类型
        if issue_price < ma120:
            threshold = issue_price  # 折价发行，盈利阈值 = 发行价
        else:
            threshold = max(ma120, issue_price * 1.02)  # 溢价发行，盈利阈值 = max(MA20, 发行价×1.02)

        # 计算基础盈利概率（保本，0%）
        # 几何布朗运动：P(S_T >= K) = 1 - Φ((log(K/S_0) - μt) / (σ√t))
        # 或者：P(S_T >= K) = Φ((μt - log(K/S_0)) / (σ√t))
        required_return = (threshold - current_price) / current_price
        z_score = (lockup_drift - required_return) / lockup_vol
        prob = (1 - stats.norm.cdf(-z_score)) * 100

        # 计算期望收益率（基于对数正态分布）
        # E[S_T] = current_price * exp(μt + σ²t/2)
        expected_future_price = current_price * np.exp(lockup_drift + lockup_vol**2 / 2)
        expected_return = (expected_future_price - issue_price) / issue_price * 100
        expected_returns_list.append(expected_return)

        # 计算不同盈利目标的概率
        profit_probs = []
        for target_pct in profit_targets_pct:
            target_price = issue_price * (1 + target_pct/100)
            target_return = (target_price - current_price) / current_price
            z_score_target = (lockup_drift - target_return) / lockup_vol
            prob_target = (1 - stats.norm.cdf(-z_score_target)) * 100
            profit_probs.append(f'{prob_target:.1f}%')

        drift_analysis_data.append([
            f'{drift*100:+.0f}%',
            scenario_name,
            f'{prob:.1f}%',
            f'{expected_return:+.1f}%',
            f'/'.join(profit_probs)  # 0%/5%/10%/15%/20%盈利目标的概率
        ])

    # 修改表格列名，增加盈利目标概率
    add_table_data(document, ['年化漂移率', '情景', '保本概率', '期望收益率', '盈利目标概率(0/5/10/15/20%)'], drift_analysis_data)

    # 计算每增加10%漂移率，期望收益率提升多少
    # 使用相邻漂移率计算平均增幅
    if len(expected_returns_list) >= 2:
        return_increments = []
        for i in range(1, len(expected_returns_list)):
            increment = expected_returns_list[i] - expected_returns_list[i-1]
            return_increments.append(increment)
        avg_return_increment = np.mean(return_increments)
    else:
        avg_return_increment = 0

    add_paragraph(document, '')
    add_paragraph(document, '漂移率敏感性分析结论：', bold=True)

    # 当前漂移率
    current_drift = market_data.get('annual_return_120d', risk_params.get('drift', 0.08))

    add_paragraph(document, f'• 当前120日窗口漂移率: {current_drift*100:+.2f}%')
    add_paragraph(document, f'• 当前120日窗口波动率: {current_vol*100:.2f}%')
    add_paragraph(document, f'• 锁定期: {lockup_months} 个月')
    add_paragraph(document, '')
    add_paragraph(document, '关键发现：')
    add_paragraph(document, '• "盈利目标概率"列显示了不同盈利目标（0%/5%/10%/15%/20%）的达成概率')
    add_paragraph(document, '• 漂移率越高，不仅保本概率提升，高盈利目标的达成概率也显著提升')
    add_paragraph(document, '• 负漂移率下，即使是保本（0%）也面临挑战，更难达成高盈利目标')
    add_paragraph(document, f'• 漂移率每提升10%，期望收益率平均提升约{avg_return_increment:.1f}%')
    add_paragraph(document, f'• 漂移率每提升10%，期望收益率平均提升约{avg_return_increment:.1f}个%')
    add_paragraph(document, '')

    # 根据当前漂移率给出分析
    if current_drift < -0.10:
        drift_assessment = f"当前漂移率为负（{current_drift*100:.1f}%），处于下降趋势，盈利概率较低"
    elif current_drift < 0:
        drift_assessment = f"当前漂移率略为负值（{current_drift*100:.1f}%），趋势偏弱，需谨慎评估"
    elif current_drift < 0.10:
        drift_assessment = f"当前漂移率为正（{current_drift*100:.1f}%），处于温和上升趋势"
    else:
        drift_assessment = f"当前漂移率较高（{current_drift*100:.1f}%），处于强势上升趋势，盈利潜力较大"

    add_paragraph(document, '• 漂移率每提升10%，盈利概率约提升15-25%')
    add_paragraph(document, '• 漂移率每提升10%，盈利概率约提升15-25个%')
    add_paragraph(document, '• 漂移率为负时，即使有折价发行，盈利概率也显著下降')
    add_paragraph(document, '• 投资建议：优先选择漂移率为正或接近零的股票进行定增投资')

    # ==================== 4.4 波动率敏感性分析 ====================
    add_paragraph(document, '')
    add_title(document, '4.4 波动率敏感性分析（120日窗口）', level=2)

    add_paragraph(document, '本节分析基于120日历史窗口的波动率对盈利概率的影响（保持其他参数不变）。')
    add_paragraph(document, '不同时间窗口的波动率不同，敏感性也会有差异。')

    # 生成敏感性分析数据和图表
    volatilities = np.linspace(0.20, 0.50, 7)
    # 使用60日窗口的参数
    drift_rate = market_data.get('annual_return_120d', risk_params.get('drift', 0.08))
    current_vol_120d = market_data.get('volatility_120d', risk_params.get('volatility', 0.30))

    prob_results = []
    mean_return_results = []  # 新增：记录不同波动率下的期望收益率
    profit_target_results = []  # 新增：记录不同盈利目标的概率

    for vol in volatilities:
        lockup_months = project_params['lockup_period']
        lockup_drift = drift_rate * (lockup_months / 12)
        lockup_vol = vol * np.sqrt(lockup_months / 12)

        # 计算盈利阈值价格
        issue_price = project_params['issue_price']
        current_price = project_params['current_price']

        # 判断发行类型
        if issue_price < ma120:
            threshold = issue_price  # 折价发行
        else:
            threshold = max(ma120, issue_price * 1.02)  # 溢价发行

        # 计算基础盈利概率（保本）
        required_return = (threshold - current_price) / current_price
        z_score = (lockup_drift - required_return) / lockup_vol
        prob = (1 - stats.norm.cdf(-z_score)) * 100
        prob_results.append(prob)

        # 计算期望收益率（基于对数正态分布）
        expected_future_price = current_price * np.exp(lockup_drift + lockup_vol**2 / 2)
        expected_return = (expected_future_price - issue_price) / issue_price * 100
        mean_return_results.append(expected_return)

        # 计算不同盈利目标的概率
        profit_targets_pct = [0, 5, 10, 15, 20]
        profit_probs_row = []
        for target_pct in profit_targets_pct:
            target_price = issue_price * (1 + target_pct/100)
            target_return = (target_price - current_price) / current_price
            z_score_target = (lockup_drift - target_return) / lockup_vol
            prob_target = (1 - stats.norm.cdf(-z_score_target)) * 100
            profit_probs_row.append(prob_target)
        profit_target_results.append(profit_probs_row)

    # 保存图表（使用拆分版）
    sensitivity_chart_paths = generate_sensitivity_charts_split(
        volatilities, prob_results, current_vol_120d,
        ma120, project_params['issue_price'], IMAGES_DIR)

    # 修改表格，增加盈利目标概率列
    vol_data = []
    for i, (vol, prob, ret) in enumerate(zip(volatilities, prob_results, mean_return_results)):
        profit_probs_str = '/'.join([f'{profit_target_results[i][j]:.1f}%' for j in range(5)])
        vol_data.append([f'{vol*100:.0f}%', f'{prob:.1f}%', f'{ret:+.1f}%', profit_probs_str])

    add_table_data(document, ['年化波动率 (120日窗口)', '保本概率', '期望收益率', '盈利目标概率(0/5/10/15/20%)'], vol_data)

    # 计算波动率敏感性指标
    vol_sensitivity_prob = []  # 波动率对盈利概率的影响
    vol_sensitivity_return = []  # 波动率对期望收益率的影响
    for i in range(len(volatilities) - 1):
        prob_change = prob_results[i+1] - prob_results[i]
        return_change = mean_return_results[i+1] - mean_return_results[i]
        vol_sensitivity_prob.append(abs(prob_change))
        vol_sensitivity_return.append(abs(return_change))

    avg_prob_change = np.mean(vol_sensitivity_prob)
    avg_return_change = np.mean(vol_sensitivity_return)

    add_paragraph(document, '')
    add_paragraph(document, '分析说明：', bold=True)
    add_paragraph(document, f'• 当前120日窗口波动率: {current_vol_120d*100:.2f}%')
    add_paragraph(document, f'• 历史120日年化收益率（漂移率）: {drift_rate*100:+.2f}%')
    add_paragraph(document, f'• 发行价格: {project_params["issue_price"]:.2f} 元/股')
    add_paragraph(document, f'• 折价率: {(project_params["issue_price"]/project_params["current_price"] - 1)*100:+.2f}%')
    add_paragraph(document, f'• MA120价格: {ma120:.2f} 元/股（作为基准）')
    add_paragraph(document, '')
    add_paragraph(document, '敏感性分析结论（120日窗口）：')
    add_paragraph(document, f'• 波动率每提升5%，盈利概率平均下降约{avg_prob_change:.1f}%')
    add_paragraph(document, f'• 波动率每提升5%，期望收益率平均下降约{avg_return_change:.1f}%')
    add_paragraph(document, f'• 波动率每提升5%，期望收益率平均下降约{avg_return_change:.1f}个%')
    add_paragraph(document, f'• 注：期望收益率基于对数正态分布计算，漂移率{drift_rate*100:+.2f}%，折价率{(project_params["issue_price"]/project_params["current_price"] - 1)*100:+.2f}%')
    add_paragraph(document, '')
    add_paragraph(document, '💡 盈利幅度分析：')
    add_paragraph(document, '• "盈利目标概率"列展示了不同盈利目标下的达成概率')
    add_paragraph(document, '• 低波动率环境下，达成高盈利目标（如+20%）的概率相对较高')
    add_paragraph(document, '• 高波动率环境下，即使是小幅度盈利（如+5%）也面临挑战')
    add_paragraph(document, '• 投资者应根据风险承受能力，选择合适的盈利目标')

    # 添加拆分的敏感性分析图表
    add_paragraph(document, '')
    add_paragraph(document, '图表 4.11: 波动率与盈利概率')
    add_image(document, sensitivity_chart_paths[0], width=Inches(6))
    add_paragraph(document, '')

    add_paragraph(document, '图表 4.12: 发行价折扣与盈利概率')
    add_image(document, sensitivity_chart_paths[1], width=Inches(6))
    add_paragraph(document, '')

    add_paragraph(document, '图表 4.13: 盈利概率热力图（波动率 vs 漂移率）')
    add_image(document, sensitivity_chart_paths[2], width=Inches(6))
    add_paragraph(document, '')

    add_paragraph(document, '图表 4.14: 盈利概率热力图（漂移率 vs 折价率）')
    add_image(document, sensitivity_chart_paths[3], width=Inches(6))

    add_paragraph(document, '分析结论（基于120日窗口）：')
    add_paragraph(document, f'• 保持漂移率和锁定期不变，波动率每增加10%，盈利概率下降约15-20%')
    add_paragraph(document, f'• 保持漂移率和锁定期不变，波动率每增加10%，盈利概率下降约15-20%')
    add_paragraph(document, f'• 建议结合4.1节的时间窗口分析，综合评估不同窗口期的波动风险')

    # ==================== 4.5 参数敏感性排序（龙卷风图） ====================
    add_paragraph(document, '')
    add_title(document, '4.5 参数敏感性排序（龙卷风图）', level=2)

    add_paragraph(document, '龙卷风图展示了各参数变化对盈利概率和预期收益率的影响程度，帮助识别最敏感的风险因子。')
    add_paragraph(document, '')
    add_paragraph(document, '💡 双维度敏感性分析：')
    add_paragraph(document, '• **A. 盈利概率敏感性**：参数变化对能否盈利（盈利概率>50%）的影响')
    add_paragraph(document, '• **B. 预期收益率敏感性**：参数变化对盈利多少（预期年化收益率）的影响')
    add_paragraph(document, '• 两个维度互为补充，全面评估风险')
    add_paragraph(document, '')
    add_paragraph(document, '💡 敏感性分析方法说明：')
    add_paragraph(document, '• 为公平比较各参数的敏感性，所有参数均使用**标准化的单位变化**：')
    add_paragraph(document, '  - 百分比参数（波动率、漂移率）：±5%')
    add_paragraph(document, '  - 百分比参数（波动率、漂移率）：±5个%')
    add_paragraph(document, '  - 时间参数（锁定期）：±1个月')
    add_paragraph(document, '• 通过比较相同单位变化下的影响幅度，识别最敏感的风险因子')
    add_paragraph(document, '')

    # 生成并插入龙卷风图
    tornado_chart_path = os.path.join(IMAGES_DIR, '01_tornado_chart.png')
    volatility = risk_params.get('volatility', market_data.get('volatility', 0.35))
    drift = risk_params.get('drift', market_data.get('drift', 0.08))
    prob_chart_path, return_chart_path, top_prob, top_return = generate_tornado_chart(
        project_params['issue_price'], market_data['current_price'],
        project_params['lockup_period'], volatility, drift, tornado_chart_path)

    # 添加图表4.5.1：盈利概率敏感性
    add_paragraph(document, '图表 4.5.1: 参数敏感性排序（龙卷风图） - 盈利概率敏感性')
    add_image(document, prob_chart_path, width=Inches(6.5))
    add_paragraph(document, '')

    # 添加图表4.5.2：预期收益率敏感性
    add_paragraph(document, '图表 4.5.2: 参数敏感性排序（龙卷风图） - 预期收益率敏感性')
    add_image(document, return_chart_path, width=Inches(6.5))
    add_paragraph(document, '')

    add_paragraph(document, '敏感性分析结论：')
    add_paragraph(document, '')
    add_paragraph(document, '💡 **分析方法说明**：')
    add_paragraph(document, '• 采用**归一化敏感性**分析方法，确保不同参数的公平比较')
    add_paragraph(document, '• 归一化含义：将影响折算为"每单位变化"的效应，例如：')
    add_paragraph(document, '  - 价格参数：每1%变化的影响')
    add_paragraph(document, '  - 波动率：每1%变化的影响')
    add_paragraph(document, '  - 漂移率：每1%变化的影响')
    add_paragraph(document, '• 图表排序依据：归一化敏感性（单位变化的影响程度）')
    add_paragraph(document, '• 图表柱状显示：实际情景影响（便于直观理解）')
    add_paragraph(document, '• 优势：避免因变化幅度不同导致的排序偏差，真实反映参数敏感度')
    add_paragraph(document, '')

    add_paragraph(document, '• **盈利概率敏感性**：')

    # 找出对盈利概率归一化敏感性最大的参数
    max_prob_sensitivity = max(top_prob, key=lambda x: abs(x.get('prob_sensitivity', x['prob_impact'])))
    max_prob_impact = max_prob_sensitivity['prob_impact']

    if max_prob_impact > 5:
        add_paragraph(document, f"  - {max_prob_sensitivity['parameter']}对盈利概率敏感性最高（单位变化影响最大）")
        add_paragraph(document, f"    实际情景影响：{max_prob_impact:+.1f}%（{max_prob_sensitivity.get('scenario_name', '')}）")
    elif max_prob_impact < -5:
        add_paragraph(document, f"  - {max_prob_sensitivity['parameter']}对盈利概率负面敏感性最高")
        add_paragraph(document, f"    实际情景影响：{max_prob_impact:+.1f}%（{max_prob_sensitivity.get('scenario_name', '')}）")
    else:
        add_paragraph(document, f"  - 各参数对盈利概率的敏感性相对均衡")
        add_paragraph(document, f"    最高敏感性参数：{max_prob_sensitivity['parameter']}（情景影响：{max_prob_impact:+.1f}%）")

    add_paragraph(document, '')
    add_paragraph(document, '• **预期收益率敏感性**：')

    # 找出对预期收益率归一化敏感性最大的参数
    max_return_sensitivity = max(top_return, key=lambda x: abs(x.get('return_sensitivity', x['return_impact'])))
    max_return_impact = max_return_sensitivity['return_impact']

    if max_return_impact > 5:
        add_paragraph(document, f"  - {max_return_sensitivity['parameter']}对预期收益率敏感性最高（单位变化影响最大）")
        add_paragraph(document, f"    实际情景影响：{max_return_impact:+.1f}%（{max_return_sensitivity.get('scenario_name', '')}）")
    elif max_return_impact < -5:
        add_paragraph(document, f"  - {max_return_sensitivity['parameter']}对预期收益率负面敏感性最高")
        add_paragraph(document, f"    实际情景影响：{max_return_impact:+.1f}%（{max_return_sensitivity.get('scenario_name', '')}）")
    else:
        add_paragraph(document, f"  - 各参数对预期收益率的敏感性相对均衡")
        add_paragraph(document, f"    最高敏感性参数：{max_return_sensitivity['parameter']}（情景影响：{max_return_impact:+.1f}%）")

    add_paragraph(document, '')
    add_paragraph(document, '• **综合判断**：')
    add_paragraph(document, '  - 价格类参数（发行价、当前价）通常是盈利概率的主要敏感性因素')
    add_paragraph(document, '  - 波动率和漂移率主要影响预期收益率（高波动率可能带来更高收益，但风险也更大）')
    add_paragraph(document, '  - 归一化分析揭示了各参数的真实敏感程度，便于风险识别和管理')
    add_paragraph(document, '')
    add_paragraph(document, '• **投资建议**：')
    add_paragraph(document, '  - 重点关注高敏感性参数，这些是项目的主要风险因子')
    add_paragraph(document, '  - 根据自身风险偏好，平衡盈利概率和预期收益率')
    add_paragraph(document, '  - 对于敏感性高的参数，建议设置更严格的风控阈值')

    # 注：原4.5发行价折扣敏感性分析已移至4.2章节

    # ==================== 4.6 净利润增长率敏感性分析 ====================
    # 注：净利润增长率是价格变化的内在原因，已在价格相关维度中体现，此处不再单独分析
    # add_paragraph(document, '')
    # add_title(document, '4.6 净利润增长率敏感性分析', level=2)
    #
    # add_paragraph(document, '归母净利润增长率是影响股价表现的核心因素之一。本节分析不同净利润增长率情景对定增项目收益的影响。')
    # add_paragraph(document, '')
    #
    # # 定义不同的净利润增长率情景（0%到50%）
    # growth_rates = [0.0, 0.05, 0.10, 0.15, 0.20, 0.25, 0.30, 0.40, 0.50]
    # current_price = project_params['current_price']
    # issue_price = project_params['issue_price']
    # lockup_years = project_params['lockup_period'] / 12
    #
    # # 假设净利润增长率对股价的传导系数（通常为0.5-0.8，考虑市场不确定性）
    # # 保守估计：70%的净利润增长会反映在股价上
    # growth_transmission_coeff = 0.7
    #
    # growth_analysis_data = []
    # growth_scenarios = []
    #
    # print("\n运行净利润增长率敏感性分析...")
    # for growth_rate in growth_rates:
    #     # 计算锁定期末的预期价格（基于净利润增长传导）
    #     # 公式：预期价格 = 当前价格 × (1 + 净利润增长率 × 传导系数 × 锁定期年数)
    #     expected_price_growth = current_price * (1 + growth_rate * growth_transmission_coeff * lockup_years)
    #
    #     # 计算收益率和年化收益率
    #     total_return = (expected_price_growth - issue_price) / issue_price
    #     annualized_return = (1 + total_return) ** (1 / lockup_years) - 1
    #
    #     # 使用蒙特卡洛模拟计算盈利概率（考虑波动率）
    #     volatility = market_data.get('volatility_60d', risk_params.get('volatility', 0.30))
    #     n_sim = 5000
    #
    #     # 调整drift：基础drift + 增长率贡献
    #     base_drift = market_data.get('annual_return_60d', risk_params.get('drift', 0.08))
    #     growth_drift = base_drift + (growth_rate * growth_transmission_coeff)
    #
    #     lockup_days = project_params['lockup_period'] * 30
    #     lockup_vol = volatility * np.sqrt(lockup_days / 365)
    #     lockup_drift_period = growth_drift * (lockup_days / 365)
    #
    #     np.random.seed(42)
    #     sim_returns = np.random.normal(lockup_drift_period, lockup_vol, n_sim)
    #     final_prices = current_price * np.exp(sim_returns)
    #     profit_losses = (final_prices - issue_price) / issue_price
    #     profit_prob = (profit_losses > 0).mean() * 100
    #
    #     # 95%置信区间
    #     percentile_5_return = np.percentile(profit_losses, 5)
    #     percentile_95_return = np.percentile(profit_losses, 95)
    #
    #     growth_analysis_data.append([
    #         f'{growth_rate*100:.0f}%',
    #         f'{expected_price_growth:.2f}',
    #         f'{total_return*100:.2f}%',
    #         f'{annualized_return*100:.2f}%',
    #         f'{profit_prob:.1f}%',
    #         f'{percentile_5_return*100:.2f}%',
    #         f'{percentile_95_return*100:.2f}%'
    #     ])
    #
    #     growth_scenarios.append({
    #         'growth_rate': growth_rate,
    #         'expected_price': expected_price_growth,
    #         'total_return': total_return,
    #         'annualized_return': annualized_return,
    #         'profit_prob': profit_prob
    #     })
    #
    #     print(f"  增长率 {growth_rate*100:.0f}%: 预期价={expected_price_growth:.2f}元, "
    #           f"年化收益={annualized_return*100:.2f}%, 盈利概率={profit_prob:.1f}%")
    #
    # # 添加敏感性分析表格
    # add_paragraph(document, '不同净利润增长率情景下的定增收益分析：')
    # add_paragraph(document, '（注：假设70%的净利润增长会传导至股价，考虑市场不确定性）')
    #
    # growth_headers = ['净利润增长率', '预期期末价格(元)', '总收益率', '年化收益率', '盈利概率', '5% VaR', '95% VaR']
    # add_table_data(document, growth_headers, growth_analysis_data)
    #
    # # 生成净利润增长率敏感性图表
    # add_paragraph(document, '')
    # add_paragraph(document, '图表 4.14: 净利润增长率敏感性分析')
    #
    # try:
    #     fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    #
    #     growth_labels = [f'{g*100:.0f}%' for g in growth_rates]
    #     annual_returns = [s['annualized_return']*100 for s in growth_scenarios]
    #     profit_probs = [s['profit_prob'] for s in growth_scenarios]
    #     expected_prices = [s['expected_price'] for s in growth_scenarios]
    #
    #     # 左图：增长率 vs 年化收益率
    #     ax1 = axes[0]
    #     ax1.plot(growth_labels, annual_returns, marker='o', linewidth=2.5,
    #             markersize=8, color='#3498db', label='年化收益率')
    #     ax1.axhline(y=0, color='black', linestyle='--', linewidth=1, alpha=0.5)
    #     ax1.axhline(y=20, color='green', linestyle=':', linewidth=1, alpha=0.5, label='20%目标收益线')
    #
    #     # 标注关键点
    #     for i, (label, value, price) in enumerate(zip(growth_labels, annual_returns, expected_prices)):
    #         color = 'green' if price > issue_price else 'red'
    #         # 只在关键点标注（0%, 10%, 20%, 30%, 50%）
    #         if growth_rates[i] in [0.0, 0.10, 0.20, 0.30, 0.50]:
    #             ax1.text(i, value + (3 if value > 0 else -3), f'{value:.1f}%',
    #                     ha='center', va='bottom' if value > 0 else 'top',
    #                     fontproperties=font_prop, fontsize=9, fontweight='bold', color=color)
    #
    #     ax1.set_xlabel('归母净利润增长率', fontproperties=font_prop, fontsize=12)
    #     ax1.set_ylabel('年化收益率 (%)', fontproperties=font_prop, fontsize=12)
    #     ax1.set_title('净利润增长率对年化收益的影响', fontproperties=font_prop, fontsize=13, fontweight='bold')
    #     ax1.grid(True, alpha=0.3)
    #     ax1.legend(prop=font_prop, fontsize=10)
    #
    #     for label in ax1.get_xticklabels():
    #         label.set_fontproperties(font_prop)
    #     for label in ax1.get_yticklabels():
    #         label.set_fontproperties(font_prop)
    #
    #     # 右图：增长率 vs 盈利概率
    #     ax2 = axes[1]
    #     bars = ax2.bar(growth_labels, profit_probs,
    #                   color=['#e74c3c' if p < 50 else '#f39c12' if p < 70 else '#27ae60'
    #                          for p in profit_probs],
    #                   alpha=0.7, edgecolor='white')
    #     ax2.axhline(y=50, color='black', linestyle='--', linewidth=1, alpha=0.5, label='盈亏平衡线')
    #     ax2.axhline(y=80, color='green', linestyle=':', linewidth=1, alpha=0.5, label='高概率线(80%)')
    #     ax2.set_xlabel('归母净利润增长率', fontproperties=font_prop, fontsize=12)
    #     ax2.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=12)
    #     ax2.set_title('净利润增长率对盈利概率的影响', fontproperties=font_prop, fontsize=13, fontweight='bold')
    #     ax2.set_ylim(0, 100)
    #     ax2.grid(True, alpha=0.3, axis='y')
    #     ax2.legend(prop=font_prop, fontsize=10)
    #
    #     # 添加数值标签（只标注关键点）
    #     for i, (bar, prob) in enumerate(zip(bars, profit_probs)):
    #         if growth_rates[i] in [0.0, 0.10, 0.20, 0.30, 0.50]:
    #             ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
    #                     f'{prob:.1f}%', ha='center', va='bottom',
    #                     fontproperties=font_prop, fontsize=9, fontweight='bold')
    #
    #     for label in ax2.get_xticklabels():
    #         label.set_fontproperties(font_prop)
    #     for label in ax2.get_yticklabels():
    #         label.set_fontproperties(font_prop)
    #
    #     plt.suptitle('净利润增长率敏感性分析（基于70%传导系数）',
    #                 fontproperties=font_prop, fontsize=15, fontweight='bold')
    #     plt.tight_layout()
    #
    #     # 保存图表
    #     growth_sensitivity_chart_path = os.path.join(IMAGES_DIR, '04_14_growth_rate_sensitivity.png')
    #     plt.savefig(growth_sensitivity_chart_path, dpi=150, bbox_inches='tight')
    #     plt.close()
    #
    #     # 添加图表到文档
    #     add_image(document, growth_sensitivity_chart_path, width=Inches(6))
    #     add_paragraph(document, '')
    #
    #     # 添加分析结论
    #     add_paragraph(document, '净利润增长率敏感性分析结论：')
    #
    #     # 找出盈亏平衡的增长率
    #     breakeven_growth = None
    #     for scenario in growth_scenarios:
    #         if scenario['profit_prob'] >= 50:
    #             breakeven_growth = scenario['growth_rate']
    #             break
    #
    #     if breakeven_growth is not None:
    #         add_paragraph(document, f'• 盈利概率达到50%的最低净利润增长率约为{breakeven_growth*100:.0f}%')
    #     else:
    #         add_paragraph(document, f'• 即使净利润增长50%，盈利概率仍未达到50%，说明当前定增价格偏高')
    #
    #     # 分析不同增长情景
    #     zero_growth_return = growth_scenarios[0]['annualized_return']
    #     high_growth_return = growth_scenarios[-1]['annualized_return']
    #     mid_growth_return = growth_scenarios[4]['annualized_return']  # 20%增长
    #
    #     add_paragraph(document, f'• 净利润零增长情景下，预期年化收益率为{zero_growth_return*100:+.2f}%')
    #     add_paragraph(document, f'• 净利润20%增长情景下，预期年化收益率为{mid_growth_return*100:+.2f}%')
    #     add_paragraph(document, f'• 净利润50%高增长情景下，预期年化收益率为{high_growth_return*100:+.2f}%')
    #     add_paragraph(document, f'• 增长率每提升10%，年化收益率约提升{(high_growth_return - zero_growth_return)/5*100:.2f}个%')
    #
    #     # 投资建议
    #     add_paragraph(document, '')
    #     add_paragraph(document, '💡 投资建议：')
    #
    #     # 估算公司当前的历史增长率
    #     revenue_growth = project_params.get('revenue_growth', 0.25)  # 默认25%
    #     add_paragraph(document, f'• 公司当前营收增长率约为{revenue_growth*100:.1f}%，可参考作为净利润增长预期')
    #     add_paragraph(document, '• 如果公司处于成长期，净利润增长率可能高于营收增长率')
    #     add_paragraph(document, '• 如果公司处于成熟期，净利润增长率可能接近或低于营收增长率')
    #     add_paragraph(document, '• 建议结合行业前景、公司竞争力、历史增长数据综合评估可实现的增长率')
    #
    #     # 风险提示
    #     add_paragraph(document, '')
    #     add_paragraph(document, '⚠️ 风险提示：')
    #     add_paragraph(document, '• 净利润增长受宏观经济、行业竞争、公司经营等多重因素影响')
    #     add_paragraph(document, '• 历史增长率不代表未来表现，需警惕增长放缓或下滑风险')
    #     add_paragraph(document, '• 高增长（>30%）通常难以持续，需要关注业绩不确定性')
    #
    # except Exception as e:
    #     print(f"⚠️ 生成净利润增长率敏感性图表失败: {e}")
    #     add_paragraph(document, f'⚠️ 净利润增长率敏感性图表生成失败: {e}')

    # ==================== 4.6 复合敏感性分析 ====================
    add_paragraph(document, '')
    add_title(document, '4.6 复合敏感性分析', level=2)

    add_paragraph(document, '本节通过复合敏感性分析，同时考察两个或多个参数组合对定增收益的影响。')
    add_paragraph(document, '相比单因素敏感性分析，复合分析能够揭示参数之间的交互效应，更全面地评估风险。')
    add_paragraph(document, '')

    # ==================== 4.6.1 波动率 × 溢价率热力图 ====================
    add_paragraph(document, '')
    add_title(document, '4.6.1 波动率 × 溢价率热力图', level=3)

    # 生成热力图数据
    vol_levels = [0.20, 0.30, 0.40, 0.50]
    discount_levels = [0.0, 0.05, 0.10, 0.15, 0.20]

    heatmap_data = []
    for vol in vol_levels:
        row = []
        for discount in discount_levels:
            issue_price_heat = project_params['current_price'] * (1 - discount)

            n_sim = 5000
            lockup_days = project_params['lockup_period'] * 30
            lockup_vol = vol * np.sqrt(lockup_days / 365)
            base_drift = market_data.get('annual_return_120d', 0.08)
            lockup_drift = base_drift * (lockup_days / 365)

            np.random.seed(42)
            sim_returns = np.random.normal(lockup_drift, lockup_vol, n_sim)
            final_prices = project_params['current_price'] * np.exp(sim_returns)
            profit_losses = (final_prices - issue_price_heat) / issue_price_heat
            profit_prob = (profit_losses > 0).mean() * 100
            row.append(profit_prob)
        heatmap_data.append(row)

    # 生成热力图
    try:
        fig, ax = plt.subplots(figsize=(10, 6))

        im = ax.imshow(heatmap_data, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100)

        # 设置刻度
        ax.set_xticks(np.arange(len(discount_levels)))
        ax.set_yticks(np.arange(len(vol_levels)))
        ax.set_xticklabels([f'{d*100:.0f}%' for d in discount_levels], fontproperties=font_prop)
        ax.set_yticklabels([f'{v*100:.0f}%' for v in vol_levels], fontproperties=font_prop)

        # 添加数值标注
        for i in range(len(vol_levels)):
            for j in range(len(discount_levels)):
                text = ax.text(j, i, f'{heatmap_data[i][j]:.0f}%',
                             ha="center", va="center", color="black", fontproperties=font_prop, fontsize=10,
                             fontweight='bold' if heatmap_data[i][j] >= 50 else 'normal')

        ax.set_xlabel('发行价溢价率（相对MA20）', fontproperties=font_prop, fontsize=12)
        ax.set_ylabel('波动率', fontproperties=font_prop, fontsize=12)
        ax.set_title('盈利概率热力图：波动率 × 溢价率', fontproperties=font_prop, fontsize=14, fontweight='bold')

        # 添加颜色条
        cbar = plt.colorbar(im, ax=ax)
        cbar.set_label('盈利概率 (%)', fontproperties=font_prop, fontsize=11)

        plt.tight_layout()

        # 保存图表
        heatmap_chart_path = os.path.join(IMAGES_DIR, '04_06_01_volatility_discount_heatmap.png')
        plt.savefig(heatmap_chart_path, dpi=150, bbox_inches='tight')
        plt.close()

        add_paragraph(document, '图表 4.6.1: 盈利概率热力图 - 波动率 × 溢价率')
        add_image(document, heatmap_chart_path, width=Inches(6))

        add_paragraph(document, '')
        add_paragraph(document, '热力图解读：')
        add_paragraph(document, '• 绿色区域（盈利概率>70%）：低波动率 + 低溢价率（高折价），投资安全边际充足')
        add_paragraph(document, '• 黄色区域（盈利概率40-70%）：中等风险收益比，需谨慎评估')
        add_paragraph(document, '• 红色区域（盈利概率<40%）：高风险区域，建议规避或要求更高折价')

    except Exception as e:
        print(f"⚠️ 生成热力图失败: {e}")

    # ==================== 4.6.2 漂移率 × 溢价率热力图 ====================
    add_paragraph(document, '')
    add_title(document, '4.6.2 漂移率 × 溢价率热力图', level=3)

    add_paragraph(document, '漂移率（历史年化收益率）和溢价率是影响定增投资收益的两个最敏感性变量。')
    add_paragraph(document, '本节通过蒙特卡洛模拟，分析在不同漂移率和溢价率组合下的盈利概率。')

    # 生成热力图数据
    drift_levels = np.arange(-0.30, 0.35, 0.10)  # -30%到+30%，每档10%
    premium_levels = [-0.20, -0.15, -0.10, -0.05, 0.0, 0.05, 0.10, 0.15, 0.20]  # 溢价率：-20%到+20%

    heatmap_data_drift = []
    current_vol_drift = market_data.get('volatility_60d', 0.30)  # 使用当前市场波动率

    for drift in drift_levels:
        row = []
        for premium in premium_levels:
            # 修复：溢价率可以是负数（折价）或正数（溢价）
            issue_price_heat = project_params['current_price'] * (1 + premium)

            n_sim = 5000
            lockup_days = project_params['lockup_period'] * 30
            lockup_vol = current_vol_drift * np.sqrt(lockup_days / 365)
            lockup_drift = drift * (lockup_days / 365)

            np.random.seed(42)
            sim_returns = np.random.normal(lockup_drift, lockup_vol, n_sim)
            final_prices = project_params['current_price'] * np.exp(sim_returns)
            profit_losses = (final_prices - issue_price_heat) / issue_price_heat
            profit_prob = (profit_losses > 0).mean() * 100
            row.append(profit_prob)
        heatmap_data_drift.append(row)

    # 生成热力图
    try:
        fig, ax = plt.subplots(figsize=(12, 6))

        im = ax.imshow(heatmap_data_drift, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100)

        # 设置刻度
        ax.set_xticks(np.arange(len(premium_levels)))
        ax.set_yticks(np.arange(len(drift_levels)))
        ax.set_xticklabels([f'{p*100:+.0f}%' for p in premium_levels], fontproperties=font_prop)
        ax.set_yticklabels([f'{d*100:+.0f}%' for d in drift_levels], fontproperties=font_prop)

        # 添加数值标注
        for i in range(len(drift_levels)):
            for j in range(len(premium_levels)):
                text = ax.text(j, i, f'{heatmap_data_drift[i][j]:.0f}%',
                             ha="center", va="center", color="black", fontproperties=font_prop, fontsize=9,
                             fontweight='bold' if heatmap_data_drift[i][j] >= 50 else 'normal')

        ax.set_xlabel('发行价溢价率（负值=折价，正值=溢价）', fontproperties=font_prop, fontsize=12)
        ax.set_ylabel('漂移率（年化收益率）', fontproperties=font_prop, fontsize=12)
        ax.set_title('盈利概率热力图：漂移率 × 溢价率', fontproperties=font_prop, fontsize=14, fontweight='bold')

        # 添加颜色条
        cbar = plt.colorbar(im, ax=ax)
        cbar.set_label('盈利概率 (%)', fontproperties=font_prop, fontsize=11)

        plt.tight_layout()

        # 保存图表
        heatmap_drift_chart_path = os.path.join(IMAGES_DIR, '04_06_02_drift_discount_heatmap.png')
        plt.savefig(heatmap_drift_chart_path, dpi=150, bbox_inches='tight')
        plt.close()

        add_paragraph(document, '图表 4.6.2: 盈利概率热力图 - 漂移率 × 溢价率')
        add_image(document, heatmap_drift_chart_path, width=Inches(6))

        add_paragraph(document, '')
        add_paragraph(document, '热力图解读：')
        add_paragraph(document, f'• 当前市场波动率: {current_vol_drift*100:.2f}%（固定参数）')
        add_paragraph(document, '• 漂移率（年化收益率）代表股票的历史趋势：正漂移表示上涨趋势，负漂移表示下跌趋势')
        add_paragraph(document, '• 溢价率：负值表示折价（发行价低于当前价），正值表示溢价（发行价高于当前价）')
        add_paragraph(document, '• 绿色区域（盈利概率>70%）：正漂移 + 负溢价率（折价），投资安全边际充足')
        add_paragraph(document, '• 黄色区域（盈利概率40-70%）：中等风险收益比，需谨慎评估')
        add_paragraph(document, '• 红色区域（盈利概率<40%）：负漂移 + 正溢价率，高风险区域')

        # 标注当前项目位置
        current_premium = (project_params['issue_price'] - project_params['current_price']) / project_params['current_price']
        current_drift = market_data.get('annual_return_60d', 0.08)

        add_paragraph(document, '')
        add_paragraph(document, f'当前项目定位：', bold=True)
        add_paragraph(document, f'• 漂移率: {current_drift*100:+.2f}%')
        add_paragraph(document, f'• 溢价率: {current_premium*100:+.2f}%')
        add_paragraph(document, f'• 当前市场波动率: {current_vol_drift*100:.2f}%')

    except Exception as e:
        print(f"⚠️ 生成漂移率热力图失败: {e}")

    # ==================== 4.6.3 三维敏感性分析（气泡图） ====================
    add_paragraph(document, '')
    add_title(document, '4.6.3 三维敏感性分析（气泡图）', level=3)

    add_paragraph(document, '本节通过三维气泡图，同时展示漂移率、溢价率和波动率三个参数对盈利概率的综合影响。')
    add_paragraph(document, '气泡大小表示盈利概率，帮助识别最优的风险收益组合。')
    add_paragraph(document, '')

    # 生成三维气泡图数据
    drift_levels_3d = np.arange(-0.30, 0.35, 0.10)  # -30%到+30%，每档10%
    premium_levels_3d = [-0.20, -0.15, -0.10, -0.05, 0.0, 0.05, 0.10, 0.15, 0.20]  # 溢价率：-20%到+20%
    vol_levels_3d = [0.20, 0.25, 0.30, 0.35, 0.40]  # 波动率水平

    bubble_data = []
    current_vol_3d = market_data.get('volatility_60d', 0.30)  # 使用当前市场波动率

    for drift in drift_levels_3d:
        for premium in premium_levels_3d:
            for vol in vol_levels_3d:
                # 计算发行价
                issue_price_3d = project_params['current_price'] * (1 + premium)

                # 蒙特卡洛模拟
                n_sim = 2000
                lockup_days = project_params['lockup_period'] * 30
                lockup_vol = vol * np.sqrt(lockup_days / 365)
                lockup_drift = drift * (lockup_days / 365)

                np.random.seed(42)
                sim_returns = np.random.normal(lockup_drift, lockup_vol, n_sim)
                final_prices = project_params['current_price'] * np.exp(sim_returns)
                profit_losses = (final_prices - issue_price_3d) / issue_price_3d
                profit_prob = (profit_losses > 0).mean() * 100

                bubble_data.append({
                    'drift': drift,
                    'premium': premium,
                    'vol': vol,
                    'prob': profit_prob
                })

    # 生成气泡图
    try:
        fig = plt.figure(figsize=(14, 10))
        ax = fig.add_subplot(111, projection='3d')

        # 提取数据
        drifts = [d['drift'] for d in bubble_data]
        premiums = [d['premium'] for d in bubble_data]
        vols = [d['vol'] for d in bubble_data]
        probs = [d['prob'] for d in bubble_data]

        # 气泡大小（基于盈利概率）
        bubble_sizes = [p * 3 for p in probs]  # 调整气泡大小
        bubble_colors = probs

        # 绘制气泡
        scatter = ax.scatter(drifts, premiums, vols, s=bubble_sizes, c=bubble_colors,
                           cmap='RdYlGn', alpha=0.6, edgecolors='black', linewidth=0.5, vmin=0, vmax=100)

        # 设置坐标轴标签
        ax.set_xlabel('漂移率（年化收益率）', fontproperties=font_prop, fontsize=11)
        ax.set_ylabel('溢价率（负值=折价，正值=溢价）', fontproperties=font_prop, fontsize=11)
        ax.set_zlabel('波动率', fontproperties=font_prop, fontsize=11)

        # 设置刻度标签
        ax.set_xticks(drift_levels_3d)
        ax.set_xticklabels([f'{d*100:+.0f}%' for d in drift_levels_3d], fontproperties=font_prop, fontsize=9)
        ax.set_yticks(premium_levels_3d)
        ax.set_yticklabels([f'{p*100:+.0f}%' for p in premium_levels_3d], fontproperties=font_prop, fontsize=9)
        ax.set_zticks(vol_levels_3d)
        ax.set_zticklabels([f'{v*100:.0f}%' for v in vol_levels_3d], fontproperties=font_prop, fontsize=9)

        # 设置标题
        ax.set_title('图表 4.6.3: 三维敏感性分析气泡图\n漂移率 × 溢价率 × 波动率',
                     fontproperties=font_prop, fontsize=13, fontweight='bold', pad=20)

        # 添加颜色条
        cbar = plt.colorbar(scatter, ax=ax, pad=0.1)
        cbar.set_label('盈利概率 (%)', fontproperties=font_prop, fontsize=10)

        plt.tight_layout()

        # 保存图表
        bubble_chart_path = os.path.join(IMAGES_DIR, '04_06_03_3d_bubble_chart.png')
        plt.savefig(bubble_chart_path, dpi=150, bbox_inches='tight')
        plt.close()

        add_paragraph(document, '图表 4.6.3: 三维敏感性分析气泡图 - 漂移率 × 溢价率 × 波动率')
        add_image(document, bubble_chart_path, width=Inches(6))

        add_paragraph(document, '')
        add_paragraph(document, '气泡图解读：')
        add_paragraph(document, '• X轴（漂移率）：股票历史年化收益率，正漂移表示上涨趋势，负漂移表示下跌趋势')
        add_paragraph(document, '• Y轴（溢价率）：发行价相对当前价的溢价程度，负值表示折价，正值表示溢价')
        add_paragraph(document, '• Z轴（波动率）：价格波动程度，影响收益的不确定性')
        add_paragraph(document, '• 气泡大小和颜色：表示盈利概率，绿色越大表示盈利概率越高，红色表示盈利概率越低')
        add_paragraph(document, '• 最优组合：右上方（高漂移+低溢价+低波动率）的大绿色气泡')
        add_paragraph(document, '• 最差组合：左下方（低漂移+高溢价+高波动率）的小红色气泡')

        # 标注当前项目位置
        current_premium_3d = (project_params['issue_price'] - project_params['current_price']) / project_params['current_price']
        current_drift_3d = market_data.get('annual_return_60d', 0.08)

        add_paragraph(document, '')
        add_paragraph(document, f'当前项目定位：', bold=True)
        add_paragraph(document, f'• 漂移率: {current_drift_3d*100:+.2f}%')
        add_paragraph(document, f'• 溢价率: {current_premium_3d*100:+.2f}%')
        add_paragraph(document, f'• 波动率: {current_vol_3d*100:.2f}%')

    except Exception as e:
        print(f"⚠️ 生成三维气泡图失败: {e}")

    add_paragraph(document, '')
    add_paragraph(document, '复合敏感性分析结论：')
    add_paragraph(document, '')
    add_paragraph(document, '💡 **关键发现**：')
    add_paragraph(document, '• 波动率和溢价率是影响盈利概率的最主要因素')
    add_paragraph(document, '• 高波动率需要更深的折价来补偿风险')
    add_paragraph(document, '• 正漂移率可以显著改善盈利概率，即使溢价率较高')
    add_paragraph(document, '• 三维分析揭示了参数之间的复杂交互效应')
    add_paragraph(document, '')
    add_paragraph(document, '💡 **投资建议**：')
    add_paragraph(document, '• 优先选择低波动率、低溢价率的投资机会')
    add_paragraph(document, '• 对于高波动率股票，要求更深的折价（>20%）')
    add_paragraph(document, '• 关注历史漂移率，正漂移率可适当放松溢价率要求')
    add_paragraph(document, '• 综合考虑三个参数的组合效应，避免单一维度决策')

    add_section_break(document)

    # ==================== 五、蒙特卡洛模拟 ====================
    add_title(document, '五、蒙特卡洛模拟', level=1)

    add_paragraph(document, '本章节使用蒙特卡洛方法模拟未来股价路径，评估收益分布。')

    add_title(document, '5.1 蒙特卡洛模拟基础分析', level=2)
    add_paragraph(document, '本节介绍蒙特卡洛模拟的理论基础、参数设置和基准结果。')
    add_paragraph(document, '')

    add_paragraph(document, '蒙特卡洛模拟是一种基于随机采样的数值计算方法，通过大量重复随机试验来估计复杂系统的行为。')
    add_paragraph(document, '')

    add_paragraph(document, '📊 本项目采用的模型：几何布朗运动（Geometric Brownian Motion, GBM）')
    add_paragraph(document, '')

    add_paragraph(document, '💡 模型假设：')
    add_paragraph(document, '• **股价的连续复利收益率服从正态分布**')
    add_paragraph(document, '  - 收益率 r = ln(S(t)/S(0)) ~ N(μ, σ²)')
    add_paragraph(document, '• **股价本身服从对数正态分布**')
    add_paragraph(document, '  - 股价 S(t) = S(0) × exp(r)，其中r为正态分布')
    add_paragraph(document, '• 股价变化具有连续性，不会出现跳跃')
    add_paragraph(document, '• 波动率（σ）和漂移率（μ）在模拟期间保持常数')
    add_paragraph(document, '• 市场有效，历史数据可反映未来特征')
    add_paragraph(document, '')

    add_paragraph(document, '📐 数学模型：')
    add_paragraph(document, '连续时间形式：dS = μSdt + σSdW')
    add_paragraph(document, '离散时间形式：S(t+Δt) = S(t) × exp((μ - σ²/2)Δt + σ√Δt × ε)')
    add_paragraph(document, '')
    add_paragraph(document, '其中：')
    add_paragraph(document, '• S(t)：时刻t的股价')
    add_paragraph(document, '• μ：漂移率（年化期望收益率）- 来自历史数据')
    add_paragraph(document, '• σ：波动率（年化标准差）- 来自历史数据')
    add_paragraph(document, '• Δt：时间步长')
    add_paragraph(document, '• ε：标准正态分布随机变量，ε ~ N(0,1)')
    add_paragraph(document, '')

    add_paragraph(document, '🔄 模拟步骤：')
    add_paragraph(document, '1. **参数估计**：根据历史数据计算年化漂移率（μ）和波动率（σ）')
    add_paragraph(document, '2. **时间离散化**：将锁定期T分割为n个时间步长（Δt = T/n）')
    add_paragraph(document, '3. **生成随机数**：对每个时间步长生成标准正态分布随机数 ε ~ N(0,1)')
    add_paragraph(document, '4. **路径模拟**：使用离散化公式计算下期价格：')
    add_paragraph(document, '   S(t+Δt) = S(t) × exp((μ - σ²/2)Δt + σ√Δt × ε)')
    add_paragraph(document, '5. **重复模拟**：重复步骤3-4共N次（N=10,000），得到N条可能的价格路径')
    add_paragraph(document, '6. **统计分析**：基于N条路径的终值，计算收益分布、盈利概率等统计量')
    add_paragraph(document, '')

    add_paragraph(document, '🎲 随机抽样特性：')
    add_paragraph(document, '• 每次模拟独立进行，使用固定随机种子（seed=42）确保结果可复现')
    add_paragraph(document, '• 通过大量模拟（10,000次），样本统计量收敛到真实分布特征')
    add_paragraph(document, '• 中心极限定理保证样本均值近似服从正态分布')
    add_paragraph(document, '')

    add_paragraph(document, '📈 收益率计算说明：')
    add_paragraph(document, '• **锁定期收益率**（单利）：R = (S(T) - 发行价) / 发行价')
    add_paragraph(document, '• **年化收益率**（单利）：年化R = R × (12 / 锁定期月数)')
    add_paragraph(document, '• **预期收益率**：所有模拟路径年化收益率的算术平均值')
    add_paragraph(document, '• **收益率中位数**：所有模拟路径年化收益率的中位数（更稳健，不受极端值影响）')
    add_paragraph(document, '')

    # ==================== 模拟参数 ====================
    add_paragraph(document, '<b>1. 模拟参数</b>', bold=True)

    # 计算定增溢价率（相对于MA20，与项目概况保持一致）
    ma20_mc = market_data.get('ma_20', 0)
    if ma20_mc > 0:
        # 使用MA20作为基准（与项目概况一致）
        discount_rate = (project_params['issue_price'] - ma20_mc) / ma20_mc * 100
        discount_basis = f"（相对MA20: {ma20_mc:.2f}元）"
    else:
        # 如果没有MA20数据，使用当前价格作为基准
        discount_rate = (project_params['issue_price'] - project_params['current_price']) / project_params['current_price'] * 100
        discount_basis = f"（相对当前价: {project_params['current_price']:.2f}元）"

    # 使用250日窗口参数（对应年度解禁期，年线指标）
    # 保存为全局变量供后续章节（如VaR）使用
    mc_volatility_250d = market_data.get('volatility_250d', market_data.get('volatility', 0.30))
    mc_drift_250d = market_data.get('annual_return_250d', market_data.get('drift', 0.08))

    # 同时获取60日和120日窗口参数
    mc_volatility_60d = market_data.get('volatility_60d', market_data.get('volatility', 0.30))
    mc_drift_60d = market_data.get('annual_return_60d', market_data.get('drift', 0.08))

    mc_volatility_120d = market_data.get('volatility_120d', market_data.get('volatility', 0.30))
    mc_drift_120d = market_data.get('annual_return_120d', market_data.get('drift', 0.08))

    mc_params = [
        ['模拟次数', '10,000次'],
        ['锁定期', f'{project_params["lockup_period"]} 个月'],
        ['初始价格', f'{project_params["current_price"]:.2f} 元/股'],
        ['发行价格', f'{project_params["issue_price"]:.2f} 元/股'],
        ['定增溢价率（premium_rate）', f'{discount_rate:+.2f}% {discount_basis}'],
        [''],
        ['60日窗口（季度）参数', ''],
        ['  - 年化波动率', f'{mc_volatility_60d*100:.1f}%'],
        ['  - 年化漂移率', f'{mc_drift_60d*100:.1f}%'],
        [''],
        ['120日窗口（半年线）参数', ''],
        ['  - 年化波动率', f'{mc_volatility_120d*100:.1f}%'],
        ['  - 年化漂移率', f'{mc_drift_120d*100:.1f}%'],
        [''],
        ['250日窗口（年线）参数', ''],
        ['  - 年化波动率', f'{mc_volatility_250d*100:.1f}%'],
        ['  - 年化漂移率', f'{mc_drift_250d*100:.1f}%']
    ]
    add_table_data(document, ['参数', '值'], mc_params)

    add_paragraph(document, '')
    add_paragraph(document, '💡 参数说明：')
    add_paragraph(document, '• 60日窗口：反映季度波动和收益特征，适合短期趋势分析')
    add_paragraph(document, '• 120日窗口：反映半年期波动和收益特征，平衡稳定性和时效性')
    add_paragraph(document, '• 250日窗口：反映年度波动和收益特征，数据最稳定，适合长期趋势分析')
    add_paragraph(document, '• 不同窗口期的参数差异会影响蒙特卡洛模拟结果，5.3节将对比分析')
    add_paragraph(document, '')

    # 默认使用250日窗口参数
    mc_volatility = mc_volatility_250d
    mc_drift = mc_drift_250d

    # ==================== 模拟结果 ====================
    add_paragraph(document, '<b>2. 模拟结果（250日窗口基准）</b>', bold=True)

    # 运行简化模拟
    lockup_days = project_params['lockup_period'] * 30
    n_simulations = 10000

    lockup_drift = mc_drift * (lockup_days / 365)
    lockup_vol = mc_volatility * np.sqrt(lockup_days / 365)

    np.random.seed(42)
    returns = np.random.normal(lockup_drift, lockup_vol, n_simulations)
    final_prices = project_params['current_price'] * np.exp(returns)
    profit_losses = (final_prices - project_params['issue_price']) / project_params['issue_price']

    # 统计
    profit_prob = (profit_losses > 0).mean()
    mean_return = profit_losses.mean()
    median_return = np.median(profit_losses)
    percent_5 = np.percentile(profit_losses, 5)
    percent_95 = np.percentile(profit_losses, 95)

    # ==================== 生成模拟路径可视化 ====================
    # 生成部分路径用于可视化（减少计算量）
    n_paths_visualize = 100  # 可视化100条路径
    n_steps = lockup_days  # 锁定期的天数

    # 每日漂移率和波动率
    daily_drift = mc_drift / 365
    daily_vol = mc_volatility / np.sqrt(365)

    # 生成路径
    np.random.seed(42)
    path_returns = np.random.normal(daily_drift, daily_vol, (n_paths_visualize, n_steps))

    # 计算价格路径（几何布朗运动）
    price_paths = np.zeros((n_paths_visualize, n_steps + 1))
    price_paths[:, 0] = project_params['current_price']

    for t in range(1, n_steps + 1):
        price_paths[:, t] = price_paths[:, t-1] * np.exp(path_returns[:, t-1])

    # 生成路径可视化图表
    import random

    fig, ax = plt.subplots(figsize=(14, 6))

    # 随机选择50条路径进行绘制
    n_paths_to_plot = min(50, n_paths_visualize)
    random_indices = random.sample(range(n_paths_visualize), n_paths_to_plot)

    # 绘制模拟路径
    for idx in random_indices:
        path = price_paths[idx, :]
        color = 'green' if path[-1] >= project_params['issue_price'] else 'red'
        ax.plot(range(n_steps + 1), path, color=color, alpha=0.15, linewidth=0.8)

    # 绘制关键水平线
    ax.axhline(y=project_params['issue_price'], color='blue', linestyle='--',
              linewidth=2, label=f"发行价格 ({project_params['issue_price']:.2f}元)")
    ax.axhline(y=project_params['current_price'], color='orange', linestyle='--',
              linewidth=2, label=f"当前价格 ({project_params['current_price']:.2f}元)")

    ax.set_xlabel('天数', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('股价 (元)', fontproperties=font_prop, fontsize=12)
    ax.set_title(f'蒙特卡洛模拟 - 股价路径可视化（{n_paths_to_plot}条样本路径）',
                fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, loc='upper left')
    ax.grid(True, alpha=0.3)

    plt.tight_layout()

    # 保存图表
    path_chart_path = os.path.join(IMAGES_DIR, '05_monte_carlo_paths.png')
    plt.savefig(path_chart_path, dpi=300, bbox_inches='tight')
    plt.close()

    print(f"✅ 已生成蒙特卡洛路径可视化图表: {path_chart_path}")

    # 保存图表（拆分版）
    mc_chart_paths = generate_monte_carlo_charts_split(
        final_prices, project_params['issue_price'], project_params['current_price'], IMAGES_DIR)

    # 计算折价率
    discount_rate = (project_params["issue_price"] / project_params["current_price"] - 1) * 100

    # 创建合并的表格：模拟参数 + 结果
    mc_results = [
        ['模拟次数', f'{n_simulations:,}次'],
        ['锁定期', f'{project_params["lockup_period"]}个月（{lockup_days}天）'],
        ['发行价格', f'{project_params["issue_price"]:.2f}元/股'],
        ['当前价格', f'{project_params["current_price"]:.2f}元/股'],
        ['定价方式', f'{project_params.get("pricing_method", "N/A")}（折价率：{discount_rate:+.1f}%）'],
        ['年化漂移率', f'{mc_drift*100:+.2f}%'],
        ['年化波动率', f'{mc_volatility*100:.2f}%'],
        ['', ''],  # 空行分隔
        ['盈利概率', f'{profit_prob*100:.1f}%'],
        ['预期收益率', f'{mean_return*100:.1f}%'],
        ['收益率中位数', f'{median_return*100:.1f}%'],
        ['5%分位数（最差情况）', f'{percent_5*100:.1f}%'],
        ['95%分位数（较好情况）', f'{percent_95*100:.1f}%']
    ]
    add_table_data(document, ['指标', '值'], mc_results)

    add_paragraph(document, '')
    add_paragraph(document, '💡 参数说明：')
    add_paragraph(document, '• 年化漂移率反映股价的长期趋势，正值表示上升趋势（基于250日窗口历史数据）')
    add_paragraph(document, '• 年化波动率反映股价的不确定性，值越大风险越高（基于250日窗口历史数据）')
    add_paragraph(document, '• 模拟基于几何布朗运动（GBM）模型，假设股价服从对数正态分布')
    add_paragraph(document, '• 锁定期内的收益率 = 漂移率×时间 + 波动率×√时间×随机因子')

    add_paragraph(document, '')
    add_paragraph(document, '图表 5.1: 蒙特卡洛模拟 - 股价路径可视化')
    add_image(document, path_chart_path, width=Inches(6.5))
    add_paragraph(document, '')
    add_paragraph(document, '路径可视化说明：')
    add_paragraph(document, '• 绿色路径：锁定期末价格高于发行价（盈利）')
    add_paragraph(document, '• 红色路径：锁定期末价格低于发行价（亏损）')
    add_paragraph(document, '• 蓝色虚线：发行价格（盈利/亏损分界线）')
    add_paragraph(document, '• 橙色虚线：当前价格（模拟起点）')
    add_paragraph(document, f'• 显示{n_paths_to_plot}条随机样本路径，实际模拟{n_simulations:,}次')
    add_paragraph(document, '')

    add_paragraph(document, '图表 5.2: 蒙特卡洛模拟结果 - 价格分布')
    add_image(document, mc_chart_paths[0])

    add_paragraph(document, '')
    add_paragraph(document, '图表 5.3: 蒙特卡洛模拟结果 - 收益率分布')
    add_image(document, mc_chart_paths[1])

    add_paragraph(document, '')
    add_paragraph(document, '图表 5.4: 蒙特卡洛模拟结果 - 累积分布函数')
    add_image(document, mc_chart_paths[2])

    add_paragraph(document, '')
    add_paragraph(document, '图表 5.5: 蒙特卡洛模拟结果 - 盈亏概率')
    add_image(document, mc_chart_paths[3])

    add_paragraph(document, '蒙特卡洛模拟结论：')
    add_paragraph(document, f'• 在 {n_simulations:,} 次模拟中，约 {profit_prob*100:.1f}% 的场景实现盈利')
    add_paragraph(document, f'• 预期收益率约 {mean_return*100:.1f}%')
    add_paragraph(document, f'• 95%的置信区间内，亏损可能不超过 {abs(percent_5)*100:.1f}%')

    # ==================== 5\.3 多窗口期蒙特卡洛分析 ====================
    add_title(document, '5.2 多窗口期对比分析', level=2)

    add_paragraph(document, '本节分析不同时间窗口（60日/120日/250日）下的蒙特卡洛模拟结果，对比不同窗口期的波动率和收益率对定增收益的影响。')

    # 多窗口期配置
    windows_mc = {
        '60日': {'days': 60, 'vol_key': 'volatility_60d', 'return_key': 'annual_return_60d'},
        '120日': {'days': 120, 'vol_key': 'volatility_120d', 'return_key': 'annual_return_120d'},
        '250日': {'days': 250, 'vol_key': 'volatility_250d', 'return_key': 'annual_return_250d'}
    }

    # 为每个窗口期运行模拟
    multi_window_mc_results = {}
    mc_n_sim = 5000  # 使用5000次模拟以加快速度

    print("\n运行多窗口期蒙特卡洛模拟...")
    for window_name, config in windows_mc.items():
        print(f"  模拟 {window_name}...")

        # 获取该窗口期的波动率和收益率
        window_vol = market_data.get(config['vol_key'], risk_params.get('volatility', 0.30))

        # 获取历史年化收益率，如果缺失则使用备选数据
        window_drift = market_data.get(config['return_key'])
        if window_drift is None:
            window_drift = risk_params.get('drift', 0.08)
            print(f"    警告：未找到{window_name}年化收益率，使用默认值({window_drift*100:.2f}%)")

        # 计算时间步数（交易日转换为天，假设1月=21交易日）
        time_steps = config['days']

        # 运行模拟
        try:
            sim_window = analyzer.monte_carlo_simulation(
                n_simulations=mc_n_sim,
                time_steps=time_steps,
                volatility=window_vol,
                drift=window_drift,
                seed=42
            )

            # 提取锁定期末价格
            final_prices = sim_window.iloc[:, -1].values
            returns = (final_prices - project_params['issue_price']) / project_params['issue_price']

            # 年化收益率计算：使用固定系数（单利）
            # 60日=1/4(季度), 120日=1/2(半年), 250日=1(年度)
            if time_steps == 60:
                coefficient = 0.25  # 季度
            elif time_steps == 120:
                coefficient = 0.5   # 半年
            elif time_steps == 250:
                coefficient = 1.0   # 年度
            else:
                coefficient = time_steps / 252.0  # 其他情况用实际计算
            annualized_returns = returns / coefficient if coefficient > 0 else returns

            # 调试信息
            print(f"    调试: 当前价={project_params['current_price']:.2f}, 发行价={project_params['issue_price']:.2f}")
            print(f"    调试: 模拟期数={time_steps}日(系数={coefficient}), 锁定期={project_params['lockup_period']}月")
            print(f"    调试: 漂移率={window_drift*100:.2f}%, 波动率={window_vol*100:.2f}%")
            print(f"    调试: 最终价格均值={final_prices.mean():.2f}, 中位数={np.median(final_prices):.2f}")
            print(f"    调试: 期收益率均值={returns.mean()*100:.2f}%, 年化收益率均值={annualized_returns.mean()*100:.2f}%")
            print(f"    调试: 盈利次数={(returns>0).sum()}/{len(returns)}, 盈利概率={(returns>0).mean()*100:.1f}%")

            # 统计分析
            multi_window_mc_results[window_name] = {
                'volatility': window_vol,
                'drift': window_drift,
                'mean_return': annualized_returns.mean(),
                'median_return': np.median(annualized_returns),
                'std_return': annualized_returns.std(),
                'profit_prob': (returns > 0).mean() * 100,
                'percentile_5': np.percentile(annualized_returns, 5),
                'percentile_25': np.percentile(annualized_returns, 25),
                'percentile_75': np.percentile(annualized_returns, 75),
                'percentile_95': np.percentile(annualized_returns, 95),
            }

            print(f"    完成: 预期收益={multi_window_mc_results[window_name]['mean_return']*100:.2f}%, 盈利概率={multi_window_mc_results[window_name]['profit_prob']:.1f}%")

        except Exception as e:
            print(f"    错误: {e}")
            continue

    # 添加多窗口期对比表格
    add_paragraph(document, '')
    add_paragraph(document, '不同窗口期蒙特卡洛模拟对比：')
    add_paragraph(document, '下表展示在不同折价/溢价率情景下的模拟结果，涵盖深度折价到深度溢价的完整区间。')

    # 计算当前项目的折价/溢价率（相对于当前价格）
    current_premium_discount = (project_params['issue_price'] - project_params['current_price']) / project_params['current_price'] * 100

    # 修正表头，明确区分历史数据和模拟结果，删除折价/溢价率列（已在窗口期中体现）
    window_headers = ['窗口期', '波动率', '历史年化收益率', '模拟预期年化收益', '中位数收益', '盈利概率', '5% VaR', '95% VaR']
    window_table_data = []

    # 为每个窗口期，计算不同溢价率下的模拟结果
    # 溢价率：-20%, -15%, -10%, -5%（负值表示折价）
    # 溢价率：0%, 5%, 10%, 15%, 20%（正值表示溢价）
    discount_scenarios = [-20, -15, -10, -5, 0, 5, 10, 15, 20]

    for window_name in ['60日', '120日', '250日']:
        if window_name in multi_window_mc_results:
            r = multi_window_mc_results[window_name]

            # 显示当前项目的折价/溢价率（在窗口期列显示）
            discount_label = f"{window_name} ({current_premium_discount:+.1f}%)"
            window_table_data.append([
                discount_label,
                f"{r['volatility']*100:.2f}%",
                f"{r['drift']*100:.2f}%",  # 历史年化收益率（蒙特卡洛输入参数）
                f"{r['mean_return']*100:.2f}%",  # 模拟预期年化收益（蒙特卡洛输出）
                f"{r['median_return']*100:.2f}%",
                f"{r['profit_prob']:.1f}%",
                f"{r['percentile_5']*100:.2f}%",
                f"{r['percentile_95']*100:.2f}%"
            ])

            # 为该窗口期添加不同折价/溢价率情景的分析
            window_vol = r['volatility']
            window_drift = r['drift']
            time_steps = {'60日': 60, '120日': 120, '250日': 250}[window_name]

            # 对每个折价/溢价率运行模拟
            for rate in discount_scenarios:
                # 计算该折价/溢价率下的发行价
                # rate为负数表示折价（发行价低于市场价），正数表示溢价（发行价高于市场价）
                issue_price_with_rate = project_params['current_price'] * (1 + rate/100)

                sim_rate = analyzer.monte_carlo_simulation(
                    n_simulations=2000,  # 使用较少次数加快速度
                    time_steps=time_steps,
                    volatility=window_vol,
                    drift=window_drift,
                    seed=42
                )

                final_prices = sim_rate.iloc[:, -1].values
                returns = (final_prices - issue_price_with_rate) / issue_price_with_rate

                # 年化收益率计算：使用固定系数（单利）
                if time_steps == 60:
                    coefficient = 0.25
                elif time_steps == 120:
                    coefficient = 0.5
                elif time_steps == 250:
                    coefficient = 1.0
                else:
                    coefficient = time_steps / 252.0
                annualized_returns = returns / coefficient if coefficient > 0 else returns

                # 根据折价/溢价率显示不同的标签
                if rate < 0:
                    rate_label = f"  └ 折价{abs(rate):.0f}%"
                elif rate == 0:
                    rate_label = "  └ 平价"
                else:
                    rate_label = f"  └ 溢价{rate:.0f}%"

                window_table_data.append([
                    rate_label,
                    "",  # 波动率留空
                    "",  # 历史收益率留空
                    f"{annualized_returns.mean()*100:.2f}%",  # 该折价/溢价率下的预期收益
                    f"{np.median(annualized_returns)*100:.2f}%",
                    f"{(returns > 0).mean()*100:.1f}%",
                    f"{np.percentile(annualized_returns, 5)*100:.2f}%",
                    f"{np.percentile(annualized_returns, 95)*100:.2f}%"
                ])

    add_table_data(document, window_headers, window_table_data)

    # 添加说明，解释各项指标含义
    add_paragraph(document, '')
    add_paragraph(document, '💡 指标说明：')
    add_paragraph(document, '• 历史年化收益率：该窗口期内股票的实际历史表现，作为蒙特卡洛模拟的输入参数（drift，漂移率）')
    add_paragraph(document, '• 模拟预期年化收益率：')
    add_paragraph(document, '  - 基于历史波动率和漂移率，通过蒙特卡洛模拟5000次得出锁定期收益')
    add_paragraph(document, '  - 采用单利年化：年化收益 = 锁定期收益 × (12 ÷ 锁定期月数)')
    add_paragraph(document, '  - 例如：6个月锁定期收益10%，年化收益 = 10% × (12÷6) = 20%')
    add_paragraph(document, '  - 正值表示预期盈利，负值表示预期亏损')
    add_paragraph(document, '• 中位数收益：模拟结果的50%分位数，比平均值更稳健（不受极端值影响）')
    add_paragraph(document, '• 盈利概率：模拟中收益率为正的概率，反映投资获利的可能性')
    add_paragraph(document, '• 5% VaR：95%置信度下的最大损失，即只有5%的概率损失会超过此值')
    add_paragraph(document, '• 95% VaR：5%置信度下的最大收益，即只有5%的概率收益会超过此值（乐观情景）')
    add_paragraph(document, '• 窗口期标签中的百分比为当前定增项目的折价/溢价率（如-2.92%表示折价2.92%）')
    add_paragraph(document, '• 由于定增通常折价发行，即使历史收益率为负，模拟预期收益仍可能为正')

    # 生成多窗口期对比图表
    add_paragraph(document, '')
    add_paragraph(document, '图表 5.7: 多窗口期蒙特卡洛模拟对比')

    try:
        # 检查是否有可用的窗口期数据
        if not multi_window_mc_results:
            raise ValueError("没有可用的窗口期数据，所有窗口期的蒙特卡洛模拟都失败了")

        # 创建对比图表
        fig, axes = plt.subplots(2, 2, figsize=(14, 10))

        window_names_ordered = ['60日', '120日', '250日']
        vols_mc = [multi_window_mc_results[w]['volatility']*100 for w in window_names_ordered if w in multi_window_mc_results]
        drifts_mc = [multi_window_mc_results[w]['drift']*100 for w in window_names_ordered if w in multi_window_mc_results]
        mean_returns_mc = [multi_window_mc_results[w]['mean_return']*100 for w in window_names_ordered if w in multi_window_mc_results]
        profit_probs_mc = [multi_window_mc_results[w]['profit_prob'] for w in window_names_ordered if w in multi_window_mc_results]

        # 更新窗口名称列表为实际有数据的窗口
        window_names_actual = [w for w in window_names_ordered if w in multi_window_mc_results]

        # 1. 波动率对比
        ax1 = axes[0, 0]
        bars1 = ax1.bar(window_names_actual, vols_mc, color='steelblue', alpha=0.7)
        ax1.set_ylabel('波动率 (%)', fontproperties=font_prop, fontsize=12)
        ax1.set_title('不同窗口期的波动率对比', fontproperties=font_prop, fontsize=13, fontweight='bold')
        ax1.grid(True, alpha=0.3, axis='y')
        for bar, vol in zip(bars1, vols_mc):
            ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                    f'{vol:.1f}%', ha='center', va='bottom', fontproperties=font_prop, fontsize=11)
        for label in ax1.get_xticklabels():
            label.set_fontproperties(font_prop)
        for label in ax1.get_yticklabels():
            label.set_fontproperties(font_prop)

        # 2. 收益率对比
        ax2 = axes[0, 1]
        bars2 = ax2.bar(window_names_actual, drifts_mc, color='coral', alpha=0.7)
        ax2.set_ylabel('收益率 (%)', fontproperties=font_prop, fontsize=12)
        ax2.set_title('不同窗口期的收益率对比', fontproperties=font_prop, fontsize=13, fontweight='bold')
        ax2.grid(True, alpha=0.3, axis='y')
        ax2.axhline(y=0, color='black', linestyle='--', linewidth=1)
        for bar, drift in zip(bars2, drifts_mc):
            y_pos = drift + 3 if drift > 0 else drift - 4
            ax2.text(bar.get_x() + bar.get_width()/2, y_pos,
                    f'{drift:.1f}%', ha='center', va='bottom' if drift > 0 else 'top',
                    fontproperties=font_prop, fontsize=11)
        for label in ax2.get_xticklabels():
            label.set_fontproperties(font_prop)
        for label in ax2.get_yticklabels():
            label.set_fontproperties(font_prop)

        # 3. 预期年化收益率对比
        ax3 = axes[1, 0]
        bars3 = ax3.bar(window_names_actual, mean_returns_mc,
                     color=['green' if r > 0 else 'red' for r in mean_returns_mc], alpha=0.7)
        ax3.set_ylabel('预期年化收益率 (%)', fontproperties=font_prop, fontsize=12)
        ax3.set_title('不同窗口期的预期收益率对比', fontproperties=font_prop, fontsize=13, fontweight='bold')
        ax3.grid(True, alpha=0.3, axis='y')
        ax3.axhline(y=0, color='black', linestyle='--', linewidth=1)
        for bar, ret in zip(bars3, mean_returns_mc):
            y_pos = ret + 3 if ret > 0 else ret - 4
            ax3.text(bar.get_x() + bar.get_width()/2, y_pos,
                    f'{ret:.1f}%', ha='center', va='bottom' if ret > 0 else 'top',
                    fontproperties=font_prop, fontsize=11)
        for label in ax3.get_xticklabels():
            label.set_fontproperties(font_prop)
        for label in ax3.get_yticklabels():
            label.set_fontproperties(font_prop)

        # 4. 盈利概率对比
        ax4 = axes[1, 1]
        bars4 = ax4.bar(window_names_actual, profit_probs_mc, color='purple', alpha=0.7)
        ax4.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=12)
        ax4.set_title('不同窗口期的盈利概率对比', fontproperties=font_prop, fontsize=13, fontweight='bold')
        ax4.set_ylim(0, 100)
        ax4.grid(True, alpha=0.3, axis='y')
        for bar, prob in zip(bars4, profit_probs_mc):
            ax4.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                    f'{prob:.1f}%', ha='center', va='bottom', fontproperties=font_prop, fontsize=11)
        for label in ax4.get_xticklabels():
            label.set_fontproperties(font_prop)
        for label in ax4.get_yticklabels():
            label.set_fontproperties(font_prop)

        plt.suptitle('蒙特卡洛模拟 - 多窗口期对比分析', fontproperties=font_prop, fontsize=15, fontweight='bold')
        plt.tight_layout()

        # 保存图表
        mc_multi_window_chart_path = os.path.join(IMAGES_DIR, 'monte_carlo_multi_window_comparison.png')
        plt.savefig(mc_multi_window_chart_path, dpi=150, bbox_inches='tight')
        plt.close()

        # 添加图表到文档
        add_image(document, mc_multi_window_chart_path, width=Inches(6))
        add_paragraph(document, '')

        # 添加分析结论
        add_paragraph(document, '多窗口期分析结论：')

        # 根据实际有数据的窗口期显示结论
        if '60日' in window_names_actual:
            idx_60 = window_names_actual.index('60日')
            add_paragraph(document, f'• 短期（60日）波动率约{vols_mc[idx_60]:.1f}%，更能反映当前市场情绪')
        if '120日' in window_names_actual:
            idx_120 = window_names_actual.index('120日')
            add_paragraph(document, f'• 中期（120日）波动率约{vols_mc[idx_120]:.1f}%，平衡了短期波动和长期趋势')
        if '250日' in window_names_actual:
            idx_250 = window_names_actual.index('250日')
            add_paragraph(document, f'• 长期（250日）波动率约{vols_mc[idx_250]:.1f}%，反映长期基本面特征')

        # 找出最佳和最差窗口
        best_window = window_names_actual[np.argmax(mean_returns_mc)]
        worst_window = window_names_actual[np.argmin(mean_returns_mc)]
        best_prob = max(profit_probs_mc)
        best_prob_window = window_names_actual[np.argmax(profit_probs_mc)]

        add_paragraph(document, f'• 预期收益率最高的是{best_window}窗口，约{max(mean_returns_mc):.1f}%')
        add_paragraph(document, f'• 盈利概率最高的是{best_prob_window}窗口，约{best_prob:.1f}%')
        add_paragraph(document, '')

        # 添加折价/溢价率的影响分析
        add_paragraph(document, '折价/溢价率影响分析（基于120日窗口）：')
        add_paragraph(document, f'• 当前项目折价/溢价率：{current_premium_discount:+.2f}%（{"折价发行" if current_premium_discount < 0 else "平价发行" if current_premium_discount == 0 else "溢价发行"}）')

        # 从表格数据中提取不同折价/溢价率的结果
        if '120日' in multi_window_mc_results:
            r_120 = multi_window_mc_results['120日']

            # 分析不同折价/溢价率对盈利概率的影响
            # 深度折价（-20%）：通常盈利概率最高
            # 深度溢价（+20%）：通常盈利概率最低

            add_paragraph(document, '• 深度折价（-20%）情景：安全边际充足，即使在悲观市场环境下也有较高盈利概率')
            add_paragraph(document, '• 中度折价（-10%至-15%）情景：平衡安全边际和投资机会，盈利概率显著提升')
            add_paragraph(document, '• 平价（0%）情景：无安全边际，盈利概率取决于市场趋势')
            add_paragraph(document, '• 溢价（+5%至+20%）情景：无安全边际且发行成本高，盈利概率较低，风险较大')

        add_paragraph(document, '')
        add_paragraph(document, '投资建议：')
        add_paragraph(document, '• 优先选择溢价率不高于-10%的定增项目（即折价≥10%），确保足够安全边际')
        add_paragraph(document, '• 对于溢价发行的项目，要求更高的预期收益率和更低的市场波动率')
        add_paragraph(document, '• 综合参考多个窗口期的结果，避免单一窗口期的偏差')

    except Exception as e:
        print(f"⚠️ 生成多窗口期图表失败: {e}")
        add_paragraph(document, f'⚠️ 多窗口期图表生成失败: {e}')

    # ==================== 5.3 ARIMA预测漂移率 ====================
    add_title(document, '5.3 ARIMA时间序列预测漂移率', level=2)

    add_paragraph(document, '本节使用ARIMA（自回归积分滑动平均）模型预测未来120日的漂移率（年化收益率）。')
    add_paragraph(document, '')
    add_paragraph(document, '💡 ARIMA模型原理：')
    add_paragraph(document, '• ARIMA(p,d,q)模型是一种经典的时间序列预测方法')
    add_paragraph(document, '• p：自回归阶数，利用历史值的线性组合预测当前值')
    add_paragraph(document, '• d：差分阶数，通过差分实现序列平稳化')
    add_paragraph(document, '• q：滑动平均阶数，利用历史预测误差的线性组合')
    add_paragraph(document, '• 适用于捕捉趋势性和周期性模式')
    add_paragraph(document, '')

    # ==================== 获取价格序列数据 ====================
    # 直接从market_data获取（由update_market_data.py生成）
    prices_list = market_data.get('price_series', [])

    if not prices_list or len(prices_list) < 100:
        print("\n" + "="*70)
        print("❌ 错误：price_series数据不足或缺失")
        print("="*70)
        print(f"   当前数据量: {len(prices_list) if prices_list else 0} 条")
        print(f"   需要数据量: 至少 100 条")
        print("")
        print("💡 解决方案：")
        print("   请先运行数据更新脚本生成完整的市场数据：")
        print("")
        print(f"       cd {os.path.dirname(os.path.abspath(__file__))}")
        print(f"       python update_market_data.py --stock {stock_code}")
        print("")
        print("   该脚本会：")
        print("   • 从Tushare获取完整的历史价格数据")
        print("   • 保存price_series（最近500个交易日）")
        print("   • 生成所有必要的统计指标")
        print("="*70)

        # 跳过时间序列预测部分
        arima_result = {'forecast_drift': mc_drift_120d, 'model_fitted': False}
        garch_result = {'forecast_volatility': vol_120, 'model_fitted': False}
        skip_time_series = True
    else:
        skip_time_series = False
        print(f"✅ 从market_data加载price_series: {len(prices_list)}个交易日")

    # 创建时间序列预测器
    if not skip_time_series:
        try:
            if prices_list and len(prices_list) >= 100:
                prices_series = pd.Series(prices_list)
                forecaster = TimeSeriesForecaster(prices_series)

                # ARIMA预测
                print("\n运行ARIMA模型预测漂移率...")
                arima_result = forecaster.forecast_drift_with_arima(horizon=120)

                # 添加模型设置说明
                add_paragraph(document, '<b>模型设置：</b>')
                arima_params = [
                    ['模型类型', 'ARIMA(1,1,1)'],
                    ['预测期数', '120日（半年）'],
                    ['数据量', f'{len(prices_series)}个交易日'],
                    ['拟合状态', '✅ 成功' if arima_result['model_fitted'] else '⚠️ 失败（使用历史平均）']
                ]
                if arima_result['model_fitted']:
                    arima_params.append(['AIC准则', f"{arima_result['aic']:.2f}"])
                add_table_data(document, ['参数', '值'], arima_params)

                # 添加预测结果
                add_paragraph(document, '')
                add_paragraph(document, '<b>预测结果：</b>')
                add_paragraph(document, f'• 预测漂移率（年化）：<b>{arima_result["forecast_drift"]*100:.2f}%</b>')
                add_paragraph(document, f'• 预测期间：未来120个交易日')
                add_paragraph(document, f'• 含义：如果模型预测准确，未来半年股价预期年化收益率为{arima_result["forecast_drift"]*100:.2f}%')

                if not arima_result['model_fitted']:
                    add_paragraph(document, '')
                    add_paragraph(document, '⚠️ 注：ARIMA模型拟合失败，使用历史平均收益率作为预测值。这可能是因为：')
                    add_paragraph(document, '   - 数据量不足（建议至少250个交易日）')
                    add_paragraph(document, '   - 序列平稳性不足')
                    add_paragraph(document, '   - 存在异常值或结构断点')

                # 对比历史漂移率
                add_paragraph(document, '')
                add_paragraph(document, '<b>与历史漂移率对比：</b>')
                historical_drifts = [
                    ['60日窗口', f'{mc_drift_60d*100:.2f}%'],
                    ['120日窗口', f'{mc_drift_120d*100:.2f}%'],
                    ['250日窗口', f'{mc_drift_250d*100:.2f}%'],
                    ['ARIMA预测', f'<b>{arima_result["forecast_drift"]*100:.2f}%</b>']
                ]
                add_table_data(document, ['数据来源', '年化漂移率'], historical_drifts)

                print(f"✅ ARIMA预测完成：{arima_result['forecast_drift']*100:.2f}%")

            else:
                add_paragraph(document, '⚠️ 价格序列数据不足，无法进行ARIMA预测（需要至少100个交易日数据）')
                print("⚠️ 价格序列数据不足，跳过ARIMA预测")
                arima_result = {'forecast_drift': mc_drift_120d, 'model_fitted': False}

        except Exception as e:
            add_paragraph(document, f'⚠️ ARIMA预测失败: {e}')
            print(f"⚠️ ARIMA预测失败: {e}")
            arima_result = {'forecast_drift': mc_drift_120d, 'model_fitted': False}
    else:
        # 跳过5.3-5.5节（数据不足）
        print("⚠️ 数据不足，跳过5.3-5.5节时间序列预测")
        arima_result = {'forecast_drift': mc_drift_120d, 'model_fitted': False}
        garch_result = {'forecast_volatility': mc_volatility_120d, 'model_fitted': False}

    add_section_break(document)

    # ==================== 5.4 GARCH预测波动率 ====================
    add_title(document, '5.4 GARCH模型预测波动率', level=2)

    add_paragraph(document, '本节使用GARCH（广义自回归条件异方差）模型预测未来120日的波动率。')
    add_paragraph(document, '')
    add_paragraph(document, '💡 GARCH模型原理：')
    add_paragraph(document, '• GARCH(p,q)模型专门用于建模和预测波动率（金融时间序列的二阶矩）')
    add_paragraph(document, '• 能捕捉"波动率聚集"现象：高波动期后往往还是高波动，低波动后还是低波动')
    add_paragraph(document, '• ω：长期平均方差（波动率的基准水平）')
    add_paragraph(document, '• α（ARCH项）：过去冲击对当前波动的影响（新信息的冲击）')
    add_paragraph(document, '• β（GARCH项）：过去波动对当前波动的影响（波动的持续性）')
    add_paragraph(document, '• α+β < 1：模型平稳，波动率会回归长期均值')
    add_paragraph(document, '')

    # GARCH预测
    if not skip_time_series:
        try:
            if 'forecaster' in locals():
                print("\n运行GARCH模型预测波动率...")
                garch_result = forecaster.forecast_volatility_with_garch(horizon=120)

                # 添加模型设置说明
                add_paragraph(document, '<b>模型设置：</b>')
                garch_params = [
                    ['模型类型', 'GARCH(1,1)'],
                    ['预测期数', '120日（半年）'],
                    ['数据量', f'{len(prices_series)}个交易日'],
                    ['拟合状态', '✅ 成功' if garch_result['model_fitted'] else '⚠️ 失败（使用历史波动率）']
                ]
                add_table_data(document, ['参数', '值'], garch_params)

                # 添加模型参数（如果拟合成功）
                if garch_result['model_fitted']:
                    add_paragraph(document, '')
                    add_paragraph(document, '<b>模型参数：</b>')
                    model_params = [
                        ['ω（长期平均方差）', f"{garch_result['omega']:.6f}"],
                        ['α（ARCH系数）', f"{garch_result['alpha']:.4f}"],
                        ['β（GARCH系数）', f"{garch_result['beta']:.4f}"],
                        ['α+β（持续性）', f"{garch_result['alpha'] + garch_result['beta']:.4f}"]
                    ]
                    add_table_data(document, ['参数', '值'], model_params)
                    add_paragraph(document, '')
                    add_paragraph(document, '💡 参数解读：')
                    add_paragraph(document, f'• α+β = {garch_result["alpha"] + garch_result["beta"]:.4f} {"< 1，模型平稳，波动率会逐渐回归长期均值" if garch_result["alpha"] + garch_result["beta"] < 1 else "≥ 1，模型非平稳，波动率可能发散"}')
                    if garch_result['alpha'] + garch_result['beta'] < 1:
                        half_life = -np.log(0.5) / (1 - (garch_result['alpha'] + garch_result['beta']))
                        add_paragraph(document, f'• 波动率半衰期：约{half_life:.0f}日（即冲击影响衰减一半所需时间）')

                # 添加预测结果
                add_paragraph(document, '')
                add_paragraph(document, '<b>预测结果：</b>')
                add_paragraph(document, f'• 预测波动率（年化）：<b>{garch_result["forecast_volatility"]*100:.2f}%</b>')
                add_paragraph(document, f'• 预测期间：未来120个交易日')
                add_paragraph(document, f'• 含义：未来半年股价的预期年化波动率为{garch_result["forecast_volatility"]*100:.2f}%')

                if not garch_result['model_fitted']:
                    add_paragraph(document, '')
                    add_paragraph(document, '⚠️ 注：GARCH模型拟合失败，使用历史波动率作为预测值。')

                # 对比历史波动率
                add_paragraph(document, '')
                add_paragraph(document, '<b>与历史波动率对比：</b>')
                historical_vols = [
                    ['60日窗口', f'{mc_volatility_60d*100:.2f}%'],
                    ['120日窗口', f'{mc_volatility_120d*100:.2f}%'],
                    ['250日窗口', f'{mc_volatility_250d*100:.2f}%'],
                    ['GARCH预测', f'<b>{garch_result["forecast_volatility"]*100:.2f}%</b>']
                ]
                add_table_data(document, ['数据来源', '年化波动率'], historical_vols)

                print(f"✅ GARCH预测完成：{garch_result['forecast_volatility']*100:.2f}%")

            else:
                add_paragraph(document, '⚠️ 无法创建预测器，跳过GARCH预测')
                print("⚠️ 无法创建预测器，跳过GARCH预测")
                garch_result = {'forecast_volatility': mc_volatility_120d, 'model_fitted': False}

        except Exception as e:
            add_paragraph(document, f'⚠️ GARCH预测失败: {e}')
            print(f"⚠️ GARCH预测失败: {e}")
            garch_result = {'forecast_volatility': mc_volatility_120d, 'model_fitted': False}

    add_section_break(document)

    # ==================== 5.5 基于预测参数的蒙特卡洛模拟 ====================
    if not skip_time_series:
        add_title(document, '5.5 基于预测参数的蒙特卡洛模拟（120日）', level=2)

        add_paragraph(document, '本节使用5.3和5.4预测的漂移率和波动率，进行120日蒙特卡洛模拟，评估基于预测参数的定增收益风险。')
        add_paragraph(document, '')

        # 使用预测参数
        predicted_drift = arima_result['forecast_drift']
        predicted_vol = garch_result['forecast_volatility']

        # 运行蒙特卡洛模拟
        n_simulations = 10000
        time_steps = 120  # 120日

        print(f"\n运行基于预测参数的蒙特卡洛模拟（120日）...")
        print(f"  预测漂移率: {predicted_drift*100:.2f}%")
        print(f"  预测波动率: {predicted_vol*100:.2f}%")

        sim_predicted = analyzer.monte_carlo_simulation(
            n_simulations=n_simulations,
            time_steps=time_steps,
            volatility=predicted_vol,
            drift=predicted_drift,
            seed=42
        )

        # 计算结果
        final_prices = sim_predicted.iloc[:, -1].values
        returns = (final_prices - project_params['issue_price']) / project_params['issue_price']

        # 年化收益率计算：使用固定系数（单利）
        # 120日=1/2(半年)
        coefficient = 0.5
        annualized_returns = returns / coefficient

        # 统计
        profit_prob = (returns > 0).mean() * 100
        mean_return = annualized_returns.mean()
        median_return = np.median(annualized_returns)
        percentile_5 = np.percentile(annualized_returns, 5)
        percentile_95 = np.percentile(annualized_returns, 95)

        # 添加模拟参数说明
        add_paragraph(document, '<b>模拟参数：</b>')
        mc_5_5_params = [
            ['预测漂移率', f'{predicted_drift*100:.2f}%（来自5.3节ARIMA预测）'],
            ['预测波动率', f'{predicted_vol*100:.2f}%（来自5.4节GARCH预测）'],
            ['模拟期数', f'{time_steps}日（120个交易日，约半年）'],
            ['模拟次数', f'{n_simulations:,}次'],
            ['锁定期', f'{project_params["lockup_period"]}个月']
        ]
        add_table_data(document, ['参数', '值'], mc_5_5_params)

        # 添加模拟结果
        add_paragraph(document, '')
        add_paragraph(document, '<b>模拟结果：</b>')
        mc_5_5_results = [
            ['盈利概率', f'{profit_prob:.1f}%'],
            ['预期年化收益', f'{mean_return*100:.2f}%'],
            ['中位数收益', f'{median_return*100:.2f}%'],
            ['5% VaR', f'{percentile_5*100:.2f}%'],
            ['95% VaR', f'{percentile_95*100:.2f}%']
        ]
        add_table_data(document, ['指标', '值'], mc_5_5_results)

        # 添加结论说明
        add_paragraph(document, '')
        add_paragraph(document, '<b>结果解读：</b>')
        add_paragraph(document, f'• 在{n_simulations:,}次模拟中，约{profit_prob:.1f}%的场景实现盈利')
        add_paragraph(document, f'• 预期年化收益率为{mean_return*100:.2f}%，{"高于" if mean_return > 0 else "低于"}无风险收益')
        add_paragraph(document, f'• 95%置信区间：年化收益在{percentile_5*100:.2f}%至{percentile_95*100:.2f}%之间')

        # 生成图表
        try:
            # 使用现有的图表生成函数
            mc_5_5_charts = generate_monte_carlo_charts_split(
                final_prices,
                project_params['issue_price'],
                project_params['current_price'],
                IMAGES_DIR
            )

            add_paragraph(document, '')
            add_paragraph(document, '图表 5.8: 基于预测参数的蒙特卡洛模拟 - 价格分布')
            add_image(document, mc_5_5_charts[0])

            add_paragraph(document, '')
            add_paragraph(document, '图表 5.9: 基于预测参数的蒙特卡洛模拟 - 收益率分布')
            add_image(document, mc_5_5_charts[1])

            add_paragraph(document, '')
            add_paragraph(document, '图表 5.10: 基于预测参数的蒙特卡洛模拟 - 累积分布函数')
            add_image(document, mc_5_5_charts[2])

            add_paragraph(document, '')
            add_paragraph(document, '图表 5.11: 基于预测参数的蒙特卡洛模拟 - 盈亏概率')
            add_image(document, mc_5_5_charts[3])

        except Exception as e:
            print(f"⚠️ 生成图表失败: {e}")

        # 与历史参数方法对比
        add_paragraph(document, '')
        add_paragraph(document, '<b>与历史参数方法对比：</b>')
        add_paragraph(document, '为了验证预测参数的有效性，下表对比基于预测参数和历史参数的模拟结果：')
        add_paragraph(document, '')

        # 使用历史250日参数模拟（对比基准）
        historical_drift = mc_drift_250d
        historical_vol = mc_volatility_250d

        sim_historical = analyzer.monte_carlo_simulation(
            n_simulations=n_simulations,
            time_steps=time_steps,
            volatility=historical_vol,
            drift=historical_drift,
            seed=42
        )

        final_prices_hist = sim_historical.iloc[:, -1].values
        returns_hist = (final_prices_hist - project_params['issue_price']) / project_params['issue_price']
        annualized_returns_hist = returns_hist * (12 / project_params['lockup_period'])

        # 统计
        profit_prob_hist = (returns_hist > 0).mean() * 100
        mean_return_hist = annualized_returns_hist.mean()

        # 对比表格
        comparison_table = [
            ['参数来源', '漂移率', '波动率', '盈利概率', '预期年化收益'],
            ['预测参数（ARIMA+GARCH）',
             f'{predicted_drift*100:.2f}%',
             f'{predicted_vol*100:.2f}%',
             f'{profit_prob:.1f}%',
             f'{mean_return*100:.2f}%'],
            ['历史参数（250日窗口）',
             f'{historical_drift*100:.2f}%',
             f'{historical_vol*100:.2f}%',
             f'{profit_prob_hist:.1f}%',
             f'{mean_return_hist*100:.2f}%'],
            ['', '', '', '', ''],
            ['差异',
             f'{(predicted_drift - historical_drift)*100:+.2f}%',
             f'{(predicted_vol - historical_vol)*100:+.2f}%',
             f'{profit_prob - profit_prob_hist:+.1f}%',
             f'{(mean_return - mean_return_hist)*100:+.2f}%']
        ]
        add_table_data(document, ['对比项', '漂移率', '波动率', '盈利概率', '预期收益'], comparison_table)

        # 对比分析
        add_paragraph(document, '')
        add_paragraph(document, '<b>对比分析：</b>')
        add_paragraph(document, f'• 漂移率：预测参数{"高于" if predicted_drift > historical_drift else "低于"}历史参数，差{abs(predicted_drift - historical_drift)*100:.2f}个百分点')
        add_paragraph(document, f'• 波动率：预测参数{"高于" if predicted_vol > historical_vol else "低于"}历史参数，差{abs(predicted_vol - historical_vol)*100:.2f}个百分点')
        add_paragraph(document, f'• 盈利概率：预测参数{"更高" if profit_prob > profit_prob_hist else "更低"}，差{abs(profit_prob - profit_prob_hist):.1f}个百分点')
        add_paragraph(document, f'• 预期收益：预测参数{"更高" if mean_return > mean_return_hist else "更低"}，差{abs(mean_return - mean_return_hist)*100:.2f}个百分点')

        # 投资建议
        add_paragraph(document, '')
        add_paragraph(document, '<b>投资建议：</b>')

        if mean_return > mean_return_hist and profit_prob > profit_prob_hist:
            add_paragraph(document, f'✅ 基于预测参数的模拟结果优于历史参数，预期收益更高且盈利概率更大，建议积极考虑投资。')
        elif mean_return > mean_return_hist:
            add_paragraph(document, f'⚠️ 基于预测参数的预期收益高于历史参数，但盈利概率{"略高" if profit_prob > profit_prob_hist else "略低"}，需结合其他因素综合判断。')
        elif profit_prob > profit_prob_hist:
            add_paragraph(document, f'⚠️ 基于预测参数的盈利概率高于历史参数，但预期收益{"略高" if mean_return > mean_return_hist else "略低"}，安全性较好但收益空间有限。')
        else:
            add_paragraph(document, f'⚠️ 基于预测参数的模拟结果全面劣于历史参数，建议谨慎投资或要求更高的折价率。')

        add_paragraph(document, '')
        add_paragraph(document, '💡 说明：')
        add_paragraph(document, '• 预测参数基于ARIMA和GARCH模型，考虑了时间序列的动态特征')
        add_paragraph(document, '• 历史参数基于过去250个交易日的统计，反映长期平均水平')
        add_paragraph(document, '• 两种方法各有优劣，建议结合使用，互为验证')

        print(f"✅ 5.5节完成：盈利概率{profit_prob:.1f}%，预期收益{mean_return*100:.2f}%")

        add_section_break(document)

    # ==================== 六、情景分析 ====================
    add_title(document, '六、情景分析', level=1)

    add_paragraph(document, '本章节分析定增项目在不同预设情景下的风险表现，包括多维度情景分析（预期期间收益率、净利润率、波动率、发行价折价）以及破发概率情景分析。')


    # ==================== 6.1 多参数情景分析 ====================
    add_paragraph(document, '')
    add_title(document, '6.1 多参数情景分析', level=2)

    add_paragraph(document, '本节通过穷举漂移率、波动率、溢价率三个参数的组合，全面分析不同市场环境下的定增项目收益预期。')
    add_paragraph(document, '')

    # ====== 参数空间定义 ======
    add_paragraph(document, '6.1.1 参数空间定义', bold=True)
    add_paragraph(document, '通过以下参数的穷举组合，模拟585种不同市场情景：')

    # 定义参数范围
    drift_rates = np.arange(-0.30, 0.35, 0.05)  # -30% 到 +30%，每档5%
    volatilities = np.arange(0.10, 0.55, 0.10)  # 10% 到 50%，每档10%
    discounts = np.arange(-0.20, 0.25, 0.05)  # -20% 到 +20%，每档5%

    # 参数说明表格
    param_data = [
        ['参数', '范围', '步长', '组合数'],
        ['漂移率（年化收益率）', '-30% ~ +30%', '5%', len(drift_rates)],
        ['波动率（年化）', '10% ~ 50%', '10%', len(volatilities)],
        ['折价率', '-20% ~ +20%', '5%', len(discounts)],
        ['总组合数', '-', '-', f'{len(drift_rates) * len(volatilities) * len(discounts)}']
    ]
    add_table_data(document, ['参数', '范围', '步长', '组合数'], param_data)

    add_paragraph(document, '')
    add_paragraph(document, '参数说明：')
    add_paragraph(document, f'• 当前价格: {project_params["current_price"]:.2f} 元/股')
    add_paragraph(document, f'• MA20价格: {market_data.get("ma_20", 0):.2f} 元/股（作为发行定价基准）')
    add_paragraph(document, f'• 锁定期: {project_params["lockup_period"]} 个月')
    add_paragraph(document, f'• 漂移率: 反映股价的预期趋势（负值=下跌，正值=上涨）')
    add_paragraph(document, f'• 波动率: 反映股价的不确定性（越高=风险越大）')
    add_paragraph(document, f'• 溢价率: 发行价相对MA20的溢价（负值=折价，正值=溢价，与配置的premium_rate一致）')
    add_paragraph(document, '')

    # ====== 模拟所有组合 ======
    print("\n运行多参数组合模拟...")

    all_scenarios = []
    n_sim = 2000  # 每个组合模拟2000次
    lockup_years = project_params['lockup_period'] / 12
    current_price = project_params['current_price']
    ma20_price = market_data.get('ma_20', current_price)  # 使用MA20作为发行定价基准

    scenario_count = 0
    for drift in drift_rates:
        for vol in volatilities:
            for discount in discounts:
                scenario_count += 1

                # 计算发行价（使用MA20作为基准，与配置保持一致）
                issue_price = ma20_price * (1 + discount)

                # 蒙特卡洛模拟
                lockup_drift = drift * lockup_years
                lockup_vol = vol * np.sqrt(lockup_years)

                np.random.seed(42)  # 固定种子以确保可重复性
                sim_returns = np.random.normal(lockup_drift, lockup_vol, n_sim)
                final_prices = current_price * np.exp(sim_returns)

                # 计算收益
                total_returns = (final_prices - issue_price) / issue_price
                # 年化收益率（使用单利计算）
                annualized_returns = total_returns / lockup_years

                # 统计指标
                mean_return = annualized_returns.mean()
                median_return = np.median(annualized_returns)
                profit_prob = (total_returns > 0).mean() * 100
                var_5 = np.percentile(annualized_returns, 5)
                var_95 = np.percentile(annualized_returns, 95)

                all_scenarios.append({
                    'drift': drift,
                    'volatility': vol,
                    'discount': discount,
                    'issue_price': issue_price,
                    'mean_return': mean_return,
                    'median_return': median_return,
                    'profit_prob': profit_prob,
                    'var_5': var_5,
                    'var_95': var_95
                })

                if scenario_count % 100 == 0:
                    print(f"  已完成 {scenario_count}/{len(drift_rates) * len(volatilities) * len(discounts)} 个情景...")

    print(f"  完成！共模拟 {len(all_scenarios)} 个情景")

    # ====== 按收益和概率排序 ======
    add_paragraph(document, '')
    add_paragraph(document, '6.1.2 情景分析结果', bold=True)

    # 按预期年化收益排序（Top 20）
    add_paragraph(document, '表1: 按预期年化收益率排序的Top 20情景')
    add_paragraph(document, '')

    top_by_return = sorted(all_scenarios, key=lambda x: x['mean_return'], reverse=True)[:20]

    top_return_data = []
    for i, s in enumerate(top_by_return, 1):
        top_return_data.append([
            f"{i}",
            f"{s['drift']*100:+.0f}%",
            f"{s['volatility']*100:.0f}%",
            f"{s['discount']*100:+.0f}%",
            f"{s['issue_price']:.2f} 元",
            f"{s['mean_return']*100:+.2f}%",
            f"{s['profit_prob']:.1f}%"
        ])

    return_headers = ['排名', '漂移率', '波动率', '溢价率', '发行价', '预期年化收益', '盈利概率']
    add_table_data(document, return_headers, top_return_data)

    # 按盈利概率排序（Top 20）
    add_paragraph(document, '')
    add_paragraph(document, '表2: 按盈利概率排序的Top 20情景')
    add_paragraph(document, '')

    top_by_prob = sorted(all_scenarios, key=lambda x: x['profit_prob'], reverse=True)[:20]

    top_prob_data = []
    for i, s in enumerate(top_by_prob, 1):
        top_prob_data.append([
            f"{i}",
            f"{s['drift']*100:+.0f}%",
            f"{s['volatility']*100:.0f}%",
            f"{s['discount']*100:+.0f}%",
            f"{s['issue_price']:.2f} 元",
            f"{s['mean_return']*100:+.2f}%",
            f"{s['profit_prob']:.1f}%"
        ])

    add_table_data(document, return_headers, top_prob_data)

    # ====== 综合最优情景 ======
    add_paragraph(document, '')
    add_paragraph(document, '6.1.3 综合最优情景分析', bold=True)

    # 找出同时满足高收益和高概率的情景
    high_return = [s for s in all_scenarios if s['mean_return'] > 0.20]  # 年化收益>20%
    high_prob = [s for s in all_scenarios if s['profit_prob'] > 80]  # 盈利概率>80%
    best_scenarios = [s for s in high_return if s in high_prob]

    if best_scenarios:
        add_paragraph(document, f'找到 {len(best_scenarios)} 个同时满足"年化收益>20%"和"盈利概率>80%"的最优情景：')
        add_paragraph(document, '')

        best_data = []
        for i, s in enumerate(sorted(best_scenarios, key=lambda x: x['mean_return'], reverse=True)[:10], 1):
            best_data.append([
                f"{i}",
                f"{s['drift']*100:+.0f}%",
                f"{s['volatility']*100:.0f}%",
                f"{s['discount']*100:+.0f}%",
                f"{s['mean_return']*100:+.2f}%",
                f"{s['profit_prob']:.1f}%"
            ])
        add_table_data(document, ['排名', '漂移率', '波动率', '溢价率', '预期年化收益', '盈利概率'], best_data)
    else:
        add_paragraph(document, '未找到同时满足"年化收益>20%"和"盈利概率>80%"的情景。')

    # ====== 当前项目定位 ======
    add_paragraph(document, '')
    add_paragraph(document, '6.1.4 当前项目在所有情景中的定位', bold=True)

    current_drift = market_data.get('annual_return_120d', risk_params.get('drift', 0.08))
    current_vol = market_data.get('volatility_120d', risk_params.get('volatility', 0.30))
    # 使用MA20作为基准计算折价率（与配置保持一致）
    ma20_for_discount = market_data.get('ma_20', project_params['current_price'])
    current_discount = (project_params['issue_price'] / ma20_for_discount) - 1

    # 找到最接近当前参数的情景
    for s in all_scenarios:
        s['distance'] = (
            abs(s['drift'] - current_drift) +
            abs(s['volatility'] - current_vol) +
            abs(s['discount'] - current_discount)
        )

    closest_scenario = min(all_scenarios, key=lambda x: x['distance'])

    # 计算当前项目的排名
    return_rank = sorted(all_scenarios, key=lambda x: x['mean_return'], reverse=True).index(closest_scenario) + 1
    prob_rank = sorted(all_scenarios, key=lambda x: x['profit_prob'], reverse=True).index(closest_scenario) + 1

    add_paragraph(document, '基于当前市场参数的最接近情景：')
    add_paragraph(document, f'• 漂移率: {current_drift*100:+.2f}%')
    add_paragraph(document, f'• 波动率: {current_vol*100:.2f}%')
    add_paragraph(document, f'• 溢价率: {current_discount*100:+.2f}% （相对MA20: {ma20_for_discount:.2f}元）')
    add_paragraph(document, f'• 发行价: {project_params["issue_price"]:.2f} 元/股')
    add_paragraph(document, f'• 预期年化收益（算术平均）: {closest_scenario["mean_return"]*100:+.2f}%')
    add_paragraph(document, f'• 收益率中位数: {closest_scenario["median_return"]*100:+.2f}%')
    add_paragraph(document, f'• 盈利概率: {closest_scenario["profit_prob"]:.1f}%')
    add_paragraph(document, f'• 收益率排名: 第 {return_rank} 名 / 共 {len(all_scenarios)} 个情景')
    add_paragraph(document, f'• 盈利概率排名: 第 {prob_rank} 名 / 共 {len(all_scenarios)} 个情景')

    # =====% 分析结论 ======
    add_paragraph(document, '')
    add_paragraph(document, '6.1.5 分析结论', bold=True)

    # 统计分析
    avg_return_all = np.mean([s['mean_return'] for s in all_scenarios])
    avg_prob_all = np.mean([s['profit_prob'] for s in all_scenarios])

    profit_scenarios = [s for s in all_scenarios if s['profit_prob'] >= 50]
    loss_scenarios = [s for s in all_scenarios if s['profit_prob'] < 50]

    add_paragraph(document, f'• 在全部{len(all_scenarios)}个情景组合中：')
    add_paragraph(document, f"  - 盈利概率≥50%的情景: {len(profit_scenarios)} 个 ({len(profit_scenarios)/len(all_scenarios)*100:.1f}%)")
    add_paragraph(document, f"  - 盈利概率<50%的情景: {len(loss_scenarios)} 个 ({len(loss_scenarios)/len(all_scenarios)*100:.1f}%)")

    add_paragraph(document, '')
    add_paragraph(document, f'• 平均预期年化收益: {avg_return_all*100:+.2f}%')
    add_paragraph(document, f'• 平均盈利概率: {avg_prob_all:.1f}%')

    add_paragraph(document, '')
    add_paragraph(document, '关键发现：')

    # 找出漂移率的影响
    pos_drift_scenarios = [s for s in all_scenarios if s['drift'] >= 0]
    neg_drift_scenarios = [s for s in all_scenarios if s['drift'] < 0]

    if pos_drift_scenarios and neg_drift_scenarios:
        avg_return_pos = np.mean([s['mean_return'] for s in pos_drift_scenarios])
        avg_return_neg = np.mean([s['mean_return'] for s in neg_drift_scenarios])
        avg_prob_pos = np.mean([s['profit_prob'] for s in pos_drift_scenarios])
        avg_prob_neg = np.mean([s['profit_prob'] for s in neg_drift_scenarios])

        add_paragraph(document, f'• 漂移率对收益影响显著：')
        add_paragraph(document, f"  - 正漂移率情景平均收益: {avg_return_pos*100:+.2f}%, 盈利概率: {avg_prob_pos:.1f}%")
        add_paragraph(document, f"  - 负漂移率情景平均收益: {avg_return_neg*100:+.2f}%, 盈利概率: {avg_prob_neg:.1f}%")

    # 找出折价率的影响
    deep_discount_scenarios = [s for s in all_scenarios if s['discount'] <= -0.15]
    premium_scenarios = [s for s in all_scenarios if s['discount'] > 0]

    if deep_discount_scenarios and premium_scenarios:
        avg_return_discount = np.mean([s['mean_return'] for s in deep_discount_scenarios])
        avg_return_premium = np.mean([s['mean_return'] for s in premium_scenarios])
        avg_prob_discount = np.mean([s['profit_prob'] for s in deep_discount_scenarios])
        avg_prob_premium = np.mean([s['profit_prob'] for s in premium_scenarios])

        add_paragraph(document, '')
        add_paragraph(document, f'• 溢价率是盈利概率的关键：')
        add_paragraph(document, f"  - 深度折价(≤-15%)情景平均收益: {avg_return_discount*100:+.2f}%, 盈利概率: {avg_prob_discount:.1f}%")
        add_paragraph(document, f"  - 溢价情景平均收益: {avg_return_premium*100:+.2f}%, 盈利概率: {avg_prob_premium:.1f}%")

    add_paragraph(document, '')
    add_paragraph(document, '投资建议：')
    if current_drift < 0:
        add_paragraph(document, f'⚠️ 当前漂移率为{current_drift*100:+.2f}%（负值），建议要求较高折价（更负的溢价率）以补偿下行风险')
    if current_discount > -0.10:
        add_paragraph(document, f'⚠️ 当前溢价率仅为{current_discount*100:+.2f}%，建议提高至-15%以下（更深的折价）')
    else:
        add_paragraph(document, f'✅ 当前溢价率{current_discount*100:+.2f}%较为合理，提供了一定的安全边际')

    # 保存all_scenarios供附件使用
    all_scenarios_for_appendix = all_scenarios.copy()

    # ==================== 6.2 基于市场指数与行业的情景分析 ====================
    add_paragraph(document, '')
    add_title(document, '6.2 基于市场指数与行业的情景分析', level=2)

    add_paragraph(document, '本节基于市场指数和行业的历史数据，构建多个典型情景进行蒙特卡洛模拟分析，评估不同市场环境下的定增收益。')
    add_paragraph(document, '通过分析当前情景、乐观情景、中性情景和悲观情景，全面评估项目在各种市场条件下的风险收益特征。')
    add_paragraph(document, '')
    add_paragraph(document, '本节基于行业和指数的历史数据，构建6个典型情景进行蒙特卡洛模拟分析，评估不同市场环境下的定增收益。')
    add_paragraph(document, '')

    # 重新加载行业数据（确保可用）
    industry_data_available = False
    industry_vol_120d = 0.35
    industry_return_120d = 0.08
    # industry_pe_data 在第二章中已经计算，这里保留原值

    try:
        industry_data_file = os.path.join(DATA_DIR, f'{stock_code.replace(".", "_")}_industry_data.json')
        if os.path.exists(industry_data_file):
            with open(industry_data_file, 'r', encoding='utf-8') as f:
                industry_data_loaded = json.load(f)

            industry_vol_120d = industry_data_loaded.get('volatility_120d', 0.35)
            industry_return_120d = industry_data_loaded.get('annual_return_120d', 0.08)
            industry_data_available = True
            print(f"✅ 已重新加载行业数据: {industry_data_file}")
    except Exception as e:
        print(f"⚠️ 无法加载行业数据，使用默认值: {e}")

    # 加载指数数据
    indices_data_available = False
    index_vol_120d = 0.30
    index_return_120d = 0.08

    try:
        indices_data_file = os.path.join(DATA_DIR, 'market_indices_scenario_data_v2.json')
        if os.path.exists(indices_data_file):
            with open(indices_data_file, 'r', encoding='utf-8') as f:
                indices_data_loaded = json.load(f)

            # 计算主要指数的平均值（沪深300、中证500、创业板指、科创50）
            major_indices = ['沪深300', '中证500', '创业板指', '科创50']
            vol_values = []
            return_values = []

            for index_name in major_indices:
                if index_name in indices_data_loaded:
                    vol_values.append(indices_data_loaded[index_name]['volatility_120d'])
                    return_values.append(indices_data_loaded[index_name]['return_120d'])

            if vol_values and return_values:
                index_vol_120d = np.mean(vol_values)
                index_return_120d = np.mean(return_values)
                indices_data_available = True
                print(f"✅ 已加载指数数据: {indices_data_file}")
    except Exception as e:
        print(f"⚠️ 无法加载指数数据，使用默认值: {e}")

    # 计算当前项目的溢价率
    current_premium_scenario = (project_params['issue_price'] - project_params['current_price']) / project_params['current_price']

    # 定义情景参数（基于指数和行业的综合数据）
    scenarios_config = []

    # 情景1: 当前情景（标的股票120日窗口真实数据）
    scenarios_config.append({
        'name': '当前情景',
        'description': '标的股票120日窗口真实数据',
        'volatility': market_data.get('volatility_120d', 0.35),
        'drift': market_data.get('annual_return_120d', 0.08),
        'premium_rate': current_premium_scenario  # 使用刚才计算的溢价率
    })

    # 计算指数和行业的综合统计值
    # 情景2: 乐观情景（指数和行业中较高的漂移率、较低的波动率）
    if indices_data_available and industry_data_available:
        # 从指数和行业中取较高的漂移率
        optimistic_drift = max(index_return_120d, industry_return_120d)
        # 从指数和行业中取较低的波动率
        optimistic_vol = min(index_vol_120d, industry_vol_120d)

        scenarios_config.append({
            'name': '情景1（乐观）',
            'description': f'指数和行业120日窗口中漂移率较高值（{optimistic_drift*100:+.2f}%）、波动率较低值（{optimistic_vol*100:.2f}%）、溢价率-20%（深折价）',
            'volatility': optimistic_vol * 0.9,  # 取较低值后再打9折
            'drift': optimistic_drift * 1.1,  # 取较高值后再打1.1倍
            'premium_rate': -0.20  # -20%溢价率（深折价）
        })

        # 情景3: 中性情景（指数和行业中档值）
        neutral_drift = (index_return_120d + industry_return_120d) / 2
        neutral_vol = (index_vol_120d + industry_vol_120d) / 2

        scenarios_config.append({
            'name': '情景2（中性）',
            'description': f'指数和行业120日窗口中漂移率和波动率中档值（漂移率{neutral_drift*100:+.2f}%、波动率{neutral_vol*100:.2f}%）、溢价率-10%（折价）',
            'volatility': neutral_vol,
            'drift': neutral_drift,
            'premium_rate': -0.10  # -10%溢价率（折价）
        })

        # 情景4: 悲观情景（指数和行业中较低的漂移率、较高的波动率）
        pessimistic_drift = min(index_return_120d, industry_return_120d)
        pessimistic_vol = max(index_vol_120d, industry_vol_120d)

        scenarios_config.append({
            'name': '情景3（悲观）',
            'description': f'指数和行业120日窗口中漂移率较低值（{pessimistic_drift*100:+.2f}%）、波动率较高值（{pessimistic_vol*100:.2f}%）、溢价率0%（平价）',
            'volatility': pessimistic_vol * 1.1,  # 取较高值后再打1.1倍
            'drift': pessimistic_drift * 0.9,  # 取较低值后再打9折
            'premium_rate': 0.00  # 0%溢价率（平价）
        })
    else:
        # 如果没有指数或行业数据，使用原来基于行业数据的情景
        print(f"⚠️ 指数或行业数据不可用，使用基于行业数据的情景")

        # 情景2: 行业乐观情景（高漂移+低波动+深折价）
        scenarios_config.append({
            'name': '情景1（行业乐观）',
            'description': '行业120日窗口漂移率较高、波动率较低、溢价率-20%（深折价）',
            'volatility': max(0.25, industry_vol_120d * 0.8),
            'drift': max(0.15, industry_return_120d * 1.5),
            'premium_rate': -0.20
        })

        # 情景3: 行业中性情景（中等参数+中等折价）
        scenarios_config.append({
            'name': '情景2（行业中性）',
            'description': '行业120日窗口漂移率中档、波动率中档、溢价率-10%（折价）',
            'volatility': industry_vol_120d,
            'drift': industry_return_120d,
            'premium_rate': -0.10
        })

        # 情景4: 行业悲观情景（低漂移+高波动+平价）
        scenarios_config.append({
            'name': '情景3（行业悲观）',
            'description': '行业120日窗口漂移率较低、波动率较高、溢价率0%（平价）',
            'volatility': min(0.45, industry_vol_120d * 1.2),
            'drift': min(-0.10, industry_return_120d * 0.5),
            'premium_rate': 0.00
        })

    # 情景5-7: 基于行业PE分位数的情景（如果有PE数据）
    pe_scenarios = []
    if 'stock_pe_data' in locals() and 'industry_pe_data' in locals() and stock_pe_data is not None and industry_pe_data is not None:
        try:
            # 计算行业PE分位数
            industry_pe_values = industry_pe_data['pe_ttm'].dropna()
            pe_25 = industry_pe_values.quantile(0.25)
            pe_50 = industry_pe_values.quantile(0.50)
            pe_75 = industry_pe_values.quantile(0.75)

            # 获取标的个股当前PE值
            stock_pe_current = stock_pe_data['pe_ttm'].iloc[-1]

            # 根据行业PE分位数与标的个股PE的比值计算漂移率
            # 漂移率 = (行业PE分位数 / 个股PE - 1)
            # 这表示：如果个股PE相对行业PE更低，估值便宜，未来增长率应该更高

            # 计算各分位数的比值
            ratio_75 = pe_75 / stock_pe_current if stock_pe_current > 0 else 1.0
            ratio_50 = pe_50 / stock_pe_current if stock_pe_current > 0 else 1.0
            ratio_25 = pe_25 / stock_pe_current if stock_pe_current > 0 else 1.0

            # 漂移率 = 比值 - 1（直接作为增长率假设）
            drift_pe_75 = ratio_75 - 1.0
            drift_pe_50 = ratio_50 - 1.0
            drift_pe_25 = ratio_25 - 1.0

            # 限制漂移率在合理范围内（-50%到+100%）
            drift_pe_75 = max(-0.50, min(1.00, drift_pe_75))
            drift_pe_50 = max(-0.50, min(1.00, drift_pe_50))
            drift_pe_25 = max(-0.50, min(1.00, drift_pe_25))

            print(f"✅ PE分位数漂移率计算完成:")
            print(f"   标的个股当前PE: {stock_pe_current:.2f}倍")
            print(f"   行业PE 75%分位数: {pe_75:.1f}倍, 比值={ratio_75:.2f}, 漂移率={drift_pe_75*100:+.2f}%")
            print(f"   行业PE 50%分位数: {pe_50:.1f}倍, 比值={ratio_50:.2f}, 漂移率={drift_pe_50*100:+.2f}%")
            print(f"   行业PE 25%分位数: {pe_25:.1f}倍, 比值={ratio_25:.2f}, 漂移率={drift_pe_25*100:+.2f}%")
            print(f"   计算公式: 漂移率 = (行业PE分位数 / 个股PE - 1)")

            # 情景4: 行业PE 75%分位数 + 深折价
            scenarios_config.append({
                'name': '情景4（行业PE 75%分位数）',
                'description': f'行业PE 75%分位数({pe_75:.1f}倍)与个股PE({stock_pe_current:.2f}倍)比值{ratio_75:.2f}，漂移率{drift_pe_75*100:+.2f}%，溢价率-20%（深折价）',
                'volatility': industry_vol_120d,
                'drift': drift_pe_75,
                'premium_rate': -0.20,
                'pe_based': True
            })

            # 情景5: 行业PE 50%分位数 + 中等折价
            scenarios_config.append({
                'name': '情景5（行业PE 50%分位数）',
                'description': f'行业PE 50%分位数({pe_50:.1f}倍)与个股PE({stock_pe_current:.2f}倍)比值{ratio_50:.2f}，漂移率{drift_pe_50*100:+.2f}%，溢价率-10%（折价）',
                'volatility': industry_vol_120d,
                'drift': drift_pe_50,
                'premium_rate': -0.10,
                'pe_based': True
            })

            # 情景6: 行业PE 25%分位数 + 平价
            scenarios_config.append({
                'name': '情景6（行业PE 25%分位数）',
                'description': f'行业PE 25%分位数({pe_25:.1f}倍)与个股PE({stock_pe_current:.2f}倍)比值{ratio_25:.2f}，漂移率{drift_pe_25*100:+.2f}%，溢价率0%（平价）',
                'volatility': industry_vol_120d,
                'drift': drift_pe_25,
                'premium_rate': 0.00,
                'pe_based': True
            })

            pe_scenarios = ['情景4（行业PE 75%分位数）', '情景5（行业PE 50%分位数）', '情景6（行业PE 25%分位数）']

            # 情景7-9: 基于个股PE分位数的情景
            # 计算个股PE分位数
            stock_pe_values = stock_pe_data['pe_ttm'].dropna()
            stock_pe_25 = stock_pe_values.quantile(0.25)
            stock_pe_50 = stock_pe_values.quantile(0.50)
            stock_pe_75 = stock_pe_values.quantile(0.75)

            # 根据个股PE分位数与个股当前PE的比值计算漂移率
            # 漂移率 = (个股PE分位数 / 个股当前PE - 1)
            stock_ratio_75 = stock_pe_75 / stock_pe_current if stock_pe_current > 0 else 1.0
            stock_ratio_50 = stock_pe_50 / stock_pe_current if stock_pe_current > 0 else 1.0
            stock_ratio_25 = stock_pe_25 / stock_pe_current if stock_pe_current > 0 else 1.0

            stock_drift_75 = stock_ratio_75 - 1.0
            stock_drift_50 = stock_ratio_50 - 1.0
            stock_drift_25 = stock_ratio_25 - 1.0

            # 限制漂移率在合理范围内
            stock_drift_75 = max(-0.50, min(1.00, stock_drift_75))
            stock_drift_50 = max(-0.50, min(1.00, stock_drift_50))
            stock_drift_25 = max(-0.50, min(1.00, stock_drift_25))

            print(f"✅ 个股PE分位数漂移率计算完成:")
            print(f"   个股当前PE: {stock_pe_current:.2f}倍")
            print(f"   个股PE 75%分位数: {stock_pe_75:.1f}倍, 比值={stock_ratio_75:.2f}, 漂移率={stock_drift_75*100:+.2f}%")
            print(f"   个股PE 50%分位数: {stock_pe_50:.1f}倍, 比值={stock_ratio_50:.2f}, 漂移率={stock_drift_50*100:+.2f}%")
            print(f"   个股PE 25%分位数: {stock_pe_25:.1f}倍, 比值={stock_ratio_25:.2f}, 漂移率={stock_drift_25*100:+.2f}%")
            print(f"   计算公式: 漂移率 = (个股PE分位数 / 个股当前PE - 1)")

            # 情景7: 个股PE 75%分位数 + 深折价
            scenarios_config.append({
                'name': '情景7（个股PE 75%分位数）',
                'description': f'个股PE 75%分位数({stock_pe_75:.1f}倍)与当前PE({stock_pe_current:.2f}倍)比值{stock_ratio_75:.2f}，漂移率{stock_drift_75*100:+.2f}%，溢价率-20%（深折价）',
                'volatility': industry_vol_120d,
                'drift': stock_drift_75,
                'premium_rate': -0.20,
                'stock_pe_based': True
            })

            # 情景8: 个股PE 50%分位数 + 中等折价
            scenarios_config.append({
                'name': '情景8（个股PE 50%分位数）',
                'description': f'个股PE 50%分位数({stock_pe_50:.1f}倍)与当前PE({stock_pe_current:.2f}倍)比值{stock_ratio_50:.2f}，漂移率{stock_drift_50*100:+.2f}%，溢价率-10%（折价）',
                'volatility': industry_vol_120d,
                'drift': stock_drift_50,
                'premium_rate': -0.10,
                'stock_pe_based': True
            })

            # 情景9: 个股PE 25%分位数 + 平价
            scenarios_config.append({
                'name': '情景9（个股PE 25%分位数）',
                'description': f'个股PE 25%分位数({stock_pe_25:.1f}倍)与当前PE({stock_pe_current:.2f}倍)比值{stock_ratio_25:.2f}，漂移率{stock_drift_25*100:+.2f}%，溢价率0%（平价）',
                'volatility': industry_vol_120d,
                'drift': stock_drift_25,
                'premium_rate': 0.00,
                'stock_pe_based': True
            })

            pe_scenarios.extend(['情景7（个股PE 75%分位数）', '情景8（个股PE 50%分位数）', '情景9（个股PE 25%分位数）'])

            # 情景10-12: 基于DCF估值的情景
            # 检查是否有intrinsic_value（DCF内在价值）
            if 'intrinsic_value' in locals():
                current_price = project_params['current_price']

                # 计算DCF估值与当前价格的比值
                dcf_ratio = intrinsic_value / current_price if current_price > 0 else 1.0

                # DCF估值的漂移率 = (DCF内在价值 / 当前价格 - 1)
                # 这表示如果DCF内在价值高于当前价格，未来收益率应该为正
                dcf_drift = dcf_ratio - 1.0

                # 限制漂移率在合理范围内
                dcf_drift = max(-0.50, min(1.00, dcf_drift))

                print(f"✅ DCF估值漂移率计算完成:")
                print(f"   DCF内在价值: {intrinsic_value:.2f}元/股")
                print(f"   当前价格: {current_price:.2f}元/股")
                print(f"   比值: {dcf_ratio:.2f}")
                print(f"   漂移率: {dcf_drift*100:+.2f}%")
                print(f"   计算公式: 漂移率 = (DCF内在价值 / 当前价格 - 1)")

                # 情景10: DCF估值 + 深折价
                scenarios_config.append({
                    'name': '情景10（DCF估值）',
                    'description': f'DCF内在价值({intrinsic_value:.2f}元)与当前价格({current_price:.2f}元)比值{dcf_ratio:.2f}，漂移率{dcf_drift*100:+.2f}%，溢价率-20%（深折价）',
                    'volatility': industry_vol_120d,
                    'drift': dcf_drift,
                    'premium_rate': -0.20,
                    'dcf_based': True
                })

                # 情景11: DCF估值 + 中等折价
                scenarios_config.append({
                    'name': '情景11（DCF估值）',
                    'description': f'DCF内在价值({intrinsic_value:.2f}元)与当前价格({current_price:.2f}元)比值{dcf_ratio:.2f}，漂移率{dcf_drift*100:+.2f}%，溢价率-10%（折价）',
                    'volatility': industry_vol_120d,
                    'drift': dcf_drift,
                    'premium_rate': -0.10,
                    'dcf_based': True
                })

                # 情景12: DCF估值 + 平价
                scenarios_config.append({
                    'name': '情景12（DCF估值）',
                    'description': f'DCF内在价值({intrinsic_value:.2f}元)与当前价格({current_price:.2f}元)比值{dcf_ratio:.2f}，漂移率{dcf_drift*100:+.2f}%，溢价率0%（平价）',
                    'volatility': industry_vol_120d,
                    'drift': dcf_drift,
                    'premium_rate': 0.00,
                    'dcf_based': True
                })

                pe_scenarios.extend(['情景10（DCF估值）', '情景11（DCF估值）', '情景12（DCF估值）'])
            else:
                print(f"⚠️ 未找到DCF内在价值数据，跳过DCF相关情景")
        except Exception as e:
            print(f"⚠️ PE数据计算失败，跳过PE相关情景: {e}")

    # 运行蒙特卡洛模拟计算每个情景的指标
    comprehensive_results = []

    for i, scenario in enumerate(scenarios_config):
        try:
            # 计算该情景的发行价
            scenario_issue_price = project_params['current_price'] * (1 + scenario['premium_rate'])

            # 运行蒙特卡洛模拟
            sim_scenario = analyzer.monte_carlo_simulation(
                n_simulations=2000,
                time_steps=120,  # 使用120日窗口
                volatility=scenario['volatility'],
                drift=scenario['drift'],
                seed=42
            )

            final_prices = sim_scenario.iloc[:, -1].values
            returns = (final_prices - scenario_issue_price) / scenario_issue_price

            # 年化收益率计算：使用固定系数（单利）
            # 120日=1/2(半年)
            annualized_returns = returns / 0.5

            # 计算指标
            profit_prob = (returns > 0).mean() * 100
            median_return = np.median(annualized_returns) * 100
            mean_return = np.mean(annualized_returns) * 100
            var_95 = np.percentile(annualized_returns, 5) * 100

            comprehensive_results.append({
                'scenario': scenario,
                'issue_price': scenario_issue_price,
                'profit_prob': profit_prob,
                'median_return': median_return,
                'mean_return': mean_return,
                'var_95': var_95
            })

            print(f"✅ {scenario['name']}: 盈利概率={profit_prob:.1f}%, 中位数收益={median_return:.1f}%")

        except Exception as e:
            print(f"⚠️ {scenario['name']} 模拟失败: {e}")
            continue

    # 筛选6.2节相关情景（当前情景、情景1-3）
    index_industry_results = [r for r in comprehensive_results
                              if 'scenario' in r and
                              (r['scenario']['name'] == '当前情景' or
                               r['scenario']['name'].startswith('情景1（') or
                               r['scenario']['name'].startswith('情景2（') or
                               r['scenario']['name'].startswith('情景3（'))]

    # 生成6.2节的情景参数表格和图表
    if index_industry_results:
        add_paragraph(document, '')
        add_paragraph(document, '市场指数与行业情景参数表：')
        index_table_data = []
        for result in index_industry_results:
            scenario_obj = result['scenario']
            index_table_data.append([
                scenario_obj['name'],
                scenario_obj['description'],
                f"{scenario_obj['volatility']*100:.2f}%",
                f"{scenario_obj['drift']*100:+.2f}%",
                f"{scenario_obj['premium_rate']*100:+.0f}%",
                f"{result['profit_prob']:.1f}%",
                f"{result['median_return']:+.1f}%"
            ])
        add_table_data(document, ['情景名称', '情景描述', '波动率', '漂移率', '溢价率', '盈利概率', '收益率中位数'], index_table_data, font_size=10.5)

        # 生成情景对比图表
        try:
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 5))

            scenario_names = [r['scenario']['name'] for r in index_industry_results]
            profit_probs = [r['profit_prob'] for r in index_industry_results]
            median_returns = [r['median_return'] for r in index_industry_results]

            # 左图：盈利概率对比
            colors_prob = ['green' if p >= 50 else 'orange' if p >= 30 else 'red' for p in profit_probs]
            bars1 = ax1.bar(scenario_names, profit_probs, color=colors_prob, alpha=0.7, edgecolor='black')
            ax1.set_ylabel('盈利概率 (%)', fontproperties=font_prop, fontsize=12)
            ax1.set_title('各情景盈利概率对比', fontproperties=font_prop, fontsize=13, fontweight='bold')
            ax1.set_ylim(0, 100)
            ax1.axhline(y=50, color='gray', linestyle='--', alpha=0.5, label='盈亏平衡线')
            ax1.legend(prop=font_prop)

            # 添加数值标注
            for bar in bars1:
                height = bar.get_height()
                ax1.text(bar.get_x() + bar.get_width()/2., height,
                        f'{height:.1f}%', ha='center', va='bottom', fontproperties=font_prop, fontsize=10)

            # 右图：收益率中位数对比
            colors_return = ['green' if r >= 0 else 'red' for r in median_returns]
            bars2 = ax2.bar(scenario_names, median_returns, color=colors_return, alpha=0.7, edgecolor='black')
            ax2.set_ylabel('收益率中位数 (%)', fontproperties=font_prop, fontsize=12)
            ax2.set_title('各情景收益率中位数对比', fontproperties=font_prop, fontsize=13, fontweight='bold')
            ax2.axhline(y=0, color='black', linestyle='-', alpha=0.3)

            # 添加数值标注
            for bar in bars2:
                height = bar.get_height()
                ax2.text(bar.get_x() + bar.get_width()/2., height,
                        f'{height:+.1f}%', ha='center', va='bottom' if height >= 0 else 'top',
                        fontproperties=font_prop, fontsize=10)

            plt.tight_layout()

            # 保存图表
            scenario_chart_path = os.path.join(IMAGES_DIR, '06_02_scenario_comparison.png')
            plt.savefig(scenario_chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            add_paragraph(document, '')
            add_paragraph(document, '图表 6.2: 市场指数与行业情景对比分析')
            add_image(document, scenario_chart_path, width=Inches(6))

            add_paragraph(document, '')
            add_paragraph(document, '📊 情景分析结论：')
            best_scenario = max(index_industry_results, key=lambda x: x['profit_prob'])
            worst_scenario = min(index_industry_results, key=lambda x: x['profit_prob'])
            add_paragraph(document, f'• 最优情景：{best_scenario["scenario"]["name"]}，盈利概率{best_scenario["profit_prob"]:.1f}%，收益率中位数{best_scenario["median_return"]:+.1f}%')
            add_paragraph(document, f'• 最差情景：{worst_scenario["scenario"]["name"]}，盈利概率{worst_scenario["profit_prob"]:.1f}%，收益率中位数{worst_scenario["median_return"]:+.1f}%')
            add_paragraph(document, '• 投资建议：关注市场环境和行业趋势，在乐观环境下可适当提高仓位，悲观环境下需谨慎评估')

        except Exception as e:
            print(f"⚠️ 生成情景对比图表失败: {e}")
            import traceback
            traceback.print_exc()

    # 根据盈利概率和收益率中位数排序并分类
    if comprehensive_results:
        # 按盈利概率和收益率中位数综合排序
        comprehensive_results.sort(key=lambda x: (x['profit_prob'], x['median_return']), reverse=True)

        # 分类
        n = len(comprehensive_results)
        categories = []
        if n >= 3:
            third = n // 3
            categories = ['乐观'] * third + ['中性'] * (third if n >= 2*third else n - 2*third) + ['悲观'] * (n - 2*third)
        else:
            categories = ['中性'] * n

        # ==================== 6.3 基于行业PE分位数的情景分析 ====================
        add_paragraph(document, '')
        add_title(document, '6.3 基于行业PE分位数的情景分析', level=2)

        add_paragraph(document, '本节基于行业PE分位数的估值水平，构建不同估值情景进行蒙特卡洛模拟分析。')
        add_paragraph(document, '通过分析行业PE 75%、50%、25%分位数下的情景，评估不同估值水平对定增收益的影响。')
        add_paragraph(document, '')

        add_paragraph(document, '💡 行业PE分位数计算逻辑：')
        add_paragraph(document, '• 收集行业历史PE数据（通常为5年）')
        add_paragraph(document, '• 计算25%、50%、75%分位数，分别代表低估、中性、高估水平')
        add_paragraph(document, '• 漂移率 = (行业PE分位数 / 个股当前PE - 1)，反映估值回归潜力')
        add_paragraph(document, '• 如果个股PE低于行业PE分位数，估值便宜，未来收益率应该更高')
        add_paragraph(document, '')

        # 筛选行业PE情景
        industry_pe_results = [r for r in comprehensive_results if 'scenario' in r and r['scenario'].get('name', '').startswith('情景4') or r['scenario'].get('name', '').startswith('情景5') or r['scenario'].get('name', '').startswith('情景6')]

        if industry_pe_results:
            add_paragraph(document, '行业PE分位数情景参数表：')
            industry_pe_table_data = []
            for result in industry_pe_results:
                scenario_obj = result['scenario']
                industry_pe_table_data.append([
                    scenario_obj['name'],
                    scenario_obj['description'],
                    f"{scenario_obj['volatility']*100:.2f}%",
                    f"{scenario_obj['drift']*100:+.2f}%",
                    f"{scenario_obj['premium_rate']*100:+.0f}%",
                    f"{result['profit_prob']:.1f}%",
                    f"{result['median_return']:+.1f}%"
                ])
            add_table_data(document, ['情景名称', '情景描述', '波动率', '漂移率', '溢价率', '盈利概率', '收益率中位数'], industry_pe_table_data, font_size=10.5)

            add_paragraph(document, '')
            add_paragraph(document, '📊 行业PE分位数情景分析：')
            best_industry_pe = max(industry_pe_results, key=lambda x: x['profit_prob'])
            worst_industry_pe = min(industry_pe_results, key=lambda x: x['profit_prob'])
            add_paragraph(document, f'• 最优情景：{best_industry_pe["scenario"]["name"]}，盈利概率{best_industry_pe["profit_prob"]:.1f}%')
            add_paragraph(document, f'• 最差情景：{worst_industry_pe["scenario"]["name"]}，盈利概率{worst_industry_pe["profit_prob"]:.1f}%')
            add_paragraph(document, '• 投资建议：优先选择行业PE分位数较高时的投资机会，估值安全边际充足')

        # ==================== 6.4 基于个股PE分位数的情景分析 ====================
        add_paragraph(document, '')
        add_title(document, '6.4 基于个股PE分位数的情景分析', level=2)

        add_paragraph(document, '本节基于个股PE分位数的估值水平，构建不同估值情景进行蒙特卡洛模拟分析。')
        add_paragraph(document, '通过分析个股PE 75%、50%、25%分位数下的情景，评估个股历史估值水平对定增收益的影响。')
        add_paragraph(document, '')

        add_paragraph(document, '💡 个股PE分位数计算逻辑：')
        add_paragraph(document, '• 收集个股历史PE数据（通常为5年）')
        add_paragraph(document, '• 计算25%、50%、75%分位数，分别代表历史低估、中性、高估水平')
        add_paragraph(document, '• 漂移率 = (个股PE分位数 / 个股当前PE - 1)，反映估值回归潜力')
        add_paragraph(document, '• 如果当前PE低于历史分位数，估值便宜，未来收益率应该更高')
        add_paragraph(document, '')

        # 筛选个股PE情景
        stock_pe_results = [r for r in comprehensive_results if 'scenario' in r and r['scenario'].get('name', '').startswith('情景7') or r['scenario'].get('name', '').startswith('情景8') or r['scenario'].get('name', '').startswith('情景9')]

        if stock_pe_results:
            add_paragraph(document, '个股PE分位数情景参数表：')
            stock_pe_table_data = []
            for result in stock_pe_results:
                scenario_obj = result['scenario']
                stock_pe_table_data.append([
                    scenario_obj['name'],
                    scenario_obj['description'],
                    f"{scenario_obj['volatility']*100:.2f}%",
                    f"{scenario_obj['drift']*100:+.2f}%",
                    f"{scenario_obj['premium_rate']*100:+.0f}%",
                    f"{result['profit_prob']:.1f}%",
                    f"{result['median_return']:+.1f}%"
                ])
            add_table_data(document, ['情景名称', '情景描述', '波动率', '漂移率', '溢价率', '盈利概率', '收益率中位数'], stock_pe_table_data, font_size=10.5)

            add_paragraph(document, '')
            add_paragraph(document, '📊 个股PE分位数情景分析：')
            best_stock_pe = max(stock_pe_results, key=lambda x: x['profit_prob'])
            worst_stock_pe = min(stock_pe_results, key=lambda x: x['profit_prob'])
            add_paragraph(document, f'• 最优情景：{best_stock_pe["scenario"]["name"]}，盈利概率{best_stock_pe["profit_prob"]:.1f}%')
            add_paragraph(document, f'• 最差情景：{worst_stock_pe["scenario"]["name"]}，盈利概率{worst_stock_pe["profit_prob"]:.1f}%')
            add_paragraph(document, '• 投资建议：关注个股历史估值水平，当前PE处于历史低位时投资价值更高')

        # ==================== 6.5 基于DCF估值的情景分析 ====================
        add_paragraph(document, '')
        add_title(document, '6.5 基于DCF估值的情景分析', level=2)

        add_paragraph(document, '本节基于DCF绝对估值方法，构建不同估值情景进行蒙特卡洛模拟分析。')
        add_paragraph(document, '通过分析DCF内在价值与当前价格的比值，评估公司内在价值对定增收益的影响。')
        add_paragraph(document, '')

        add_paragraph(document, '💡 DCF估值方法说明：')
        add_paragraph(document, '• DCF（现金流折现模型）通过预测未来自由现金流并折现计算内在价值')
        add_paragraph(document, '• 漂移率 = (DCF内在价值 / 当前价格 - 1)，反映内在价值与市场价格的偏离')
        add_paragraph(document, '• 如果DCF内在价值高于当前价格，股票被低估，未来收益率应该为正')
        add_paragraph(document, '• DCF估值提供绝对价值判断，独立于市场相对估值')
        add_paragraph(document, '')

        # 筛选DCF情景
        dcf_results = [r for r in comprehensive_results if 'scenario' in r and r['scenario'].get('name', '').startswith('情景10') or r['scenario'].get('name', '').startswith('情景11') or r['scenario'].get('name', '').startswith('情景12')]

        if dcf_results:
            add_paragraph(document, 'DCF估值情景参数表：')
            dcf_table_data = []
            for result in dcf_results:
                scenario_obj = result['scenario']
                dcf_table_data.append([
                    scenario_obj['name'],
                    scenario_obj['description'],
                    f"{scenario_obj['volatility']*100:.2f}%",
                    f"{scenario_obj['drift']*100:+.2f}%",
                    f"{scenario_obj['premium_rate']*100:+.0f}%",
                    f"{result['profit_prob']:.1f}%",
                    f"{result['median_return']:+.1f}%"
                ])
            add_table_data(document, ['情景名称', '情景描述', '波动率', '漂移率', '溢价率', '盈利概率', '收益率中位数'], dcf_table_data, font_size=10.5)

            add_paragraph(document, '')
            add_paragraph(document, '📊 DCF估值情景分析：')
            best_dcf = max(dcf_results, key=lambda x: x['profit_prob'])
            worst_dcf = min(dcf_results, key=lambda x: x['profit_prob'])
            add_paragraph(document, f'• 最优情景：{best_dcf["scenario"]["name"]}，盈利概率{best_dcf["profit_prob"]:.1f}%')
            add_paragraph(document, f'• 最差情景：{worst_dcf["scenario"]["name"]}，盈利概率{worst_dcf["profit_prob"]:.1f}%')
            add_paragraph(document, '• 投资建议：DCF内在价值提供安全边际，优先选择DCF估值高于市场价格的项目')

        # ==================== 6.6 情景综合分析汇总表 ====================
        add_paragraph(document, '')
        add_title(document, '6.6 情景综合分析汇总表', level=2)

        add_paragraph(document, '本节汇总所有情景的蒙特卡洛模拟结果，提供综合对比分析。')
        add_paragraph(document, '通过汇总表可以全面评估不同情景下的风险收益特征，为投资决策提供参考。')
        add_paragraph(document, '')

        # 生成表格数据
        comprehensive_table_data = []
        for i, result in enumerate(comprehensive_results):
            scenario_obj = result['scenario']
            category = categories[i] if i < len(categories) else '中性'

            comprehensive_table_data.append([
                scenario_obj['description'],  # 情景描述
                f"{scenario_obj['volatility']*100:.2f}%",
                f"{scenario_obj['drift']*100:+.2f}%",
                f"{scenario_obj['premium_rate']*100:+.0f}%",
                f"{result['issue_price']:.2f}",
                f"{result['profit_prob']:.1f}%",
                f"{result['median_return']:+.1f}%",
                scenario_obj['name']  # 情景类型
            ])

        comprehensive_headers = ['情景描述', '波动率', '漂移率', '溢价率', '发行价(元)', '盈利概率(%)', '收益率中位数(%)', '情景类型']
        add_table_data(document, comprehensive_headers, comprehensive_table_data, font_size=10.5)  # 五号字体

        # 添加说明
        add_paragraph(document, '')
        add_paragraph(document, '💡 情景说明：')
        add_paragraph(document, '• 当前情景：标的股票120日窗口真实数据，反映项目实际风险收益特征')
        add_paragraph(document, '• 情景1-3：基于行业120日窗口历史数据的典型情景（乐观/中性/悲观）')
        if pe_scenarios:
            add_paragraph(document, '• 情景4-6：基于行业PE分位数的估值情景，反映不同估值水平下的风险收益')
            add_paragraph(document, '• 情景7-9：基于个股PE分位数的估值情景，反映个股历史估值水平的影响')
            if 'intrinsic_value' in locals():
                add_paragraph(document, '• 情景10-12：基于DCF绝对估值的情景，反映公司内在价值对收益的影响')
        add_paragraph(document, '')
        add_paragraph(document, '📊 评级说明：')
        add_paragraph(document, '• 乐观：盈利概率和收益率中位数均较高，投资价值突出')
        add_paragraph(document, '• 中性：盈利概率和收益率中位数适中，风险可控')
        add_paragraph(document, '• 悲观：盈利概率或收益率中位数较低，风险较高')
        add_paragraph(document, '')
        add_paragraph(document, '💡 投资建议：')
        add_paragraph(document, '• 优先选择乐观评级情景，安全边际充足')
        add_paragraph(document, '• 当前项目定位在"当前情景"，需结合实际参数评估')
        add_paragraph(document, '• 不同情景反映市场环境变化，建议根据风险偏好选择')

    # 生成多维度情景图表（补充分析）
    # 注：波动率×折价率热力图已在6.3章节展示，此处不再重复
    try:
        multi_dim_chart_paths = generate_multi_dimension_scenario_charts_split(
            project_params['current_price'], ma20, risk_params['volatility'],
            risk_params['drift'], project_params['lockup_period'], IMAGES_DIR)

        add_paragraph(document, '')
        add_paragraph(document, '图表 6.6: 不同波动率下的情景对比')
        add_image(document, multi_dim_chart_paths[1], width=Inches(6.5))
        add_paragraph(document, '')

        add_paragraph(document, '图表 6.7: 优质情景TOP10 (盈利概率>50%)')
        add_image(document, multi_dim_chart_paths[2], width=Inches(6.5))

    except Exception as e:
        print(f"⚠️ 生成多维度图表失败: {e}")

    add_section_break(document)

    # ==================== 七、压力测试 ====================
    add_title(document, '七、压力测试', level=1)

    add_paragraph(document, '本章节从多个维度模拟极端市场情况下的定增项目表现，包括估值回归风险、经济面极端情况、以及多重风险因素叠加的最差情景。')
    add_paragraph(document, '通过全面压力测试，评估项目在极端环境下的风险承受能力。')
    add_paragraph(document, '')

    # ==================== 7.1 PE回归压力测试 ====================
    add_title(document, '7.1 PE回归压力测试', level=2)

    add_paragraph(document, '本节分析PE估值回归到行业极端情况时的情景，评估估值回归风险。')
    add_paragraph(document, '通过模拟PE回归到行业Q1（25分位，即下四分位数），评估最坏情况下估值回调对投资收益的影响。')
    add_paragraph(document, '')

    # 计算回归情景
    current_price_rel = project_params['current_price']

    # 修正：使用当前价格和当前PE反推EPS，确保与PE计算口径一致
    # PE = 股价 / EPS → EPS = 股价 / PE
    eps_rel = current_price_rel / current_metrics_val['pe']

    # 获取行业PE分位数数据
    pe_q1 = industry_stats_val['pe']['q1']  # 25分位（下四分位数）

    # 计算PE回归到行业Q1的极端情景
    target_price_q1 = eps_rel * pe_q1
    return_q1 = (target_price_q1 - current_price_rel) / current_price_rel * 100

    # 构建压力测试表格
    scenario_headers_stress = ['情景', '当前PE', '回归后PE', '当前价格(元)', '目标价格(元)', '预期收益率(%)']
    scenario_data_stress = [
        ['当前估值', f"{current_metrics_val['pe']:.2f}", f"{current_metrics_val['pe']:.2f}",
         f"{current_price_rel:.2f}", f"{current_price_rel:.2f}", "0.00"],
        [f'PE→行业Q1({pe_q1:.2f}倍)', f"{current_metrics_val['pe']:.2f}", f"{pe_q1:.2f}",
         f"{current_price_rel:.2f}", f"{target_price_q1:.2f}", f"{return_q1:+.2f}"],
    ]

    add_table_data(document, scenario_headers_stress, scenario_data_stress)

    add_paragraph(document, '')
    add_paragraph(document, 'PE回归压力测试分析：', bold=True)

    # 分析当前PE在行业中的位置
    current_pe = current_metrics_val['pe']
    pe_position = (peer_companies_val['pe'] < current_pe).sum() / len(peer_companies_val) * 100

    add_paragraph(document, f'• 当前PE({current_pe:.2f}倍)位于行业{pe_position:.1f}%分位')
    add_paragraph(document, f'• 行业Q1 PE({pe_q1:.2f}倍)为25分位数，代表行业较低估值水平')
    add_paragraph(document, f'• 如果PE回归到行业Q1，目标价格为{target_price_q1:.2f}元，预期收益{return_q1:+.2f}%')

    add_paragraph(document, '')

    # 风险提示
    add_paragraph(document, '风险评估：', bold=True)
    if return_q1 > 0:
        add_paragraph(document, f'✅ 即使在PE回归到行业Q1的极端情况下，预期收益仍为正({return_q1:+.2f}%)，估值安全边际较高')
    elif return_q1 > -10:
        add_paragraph(document, f'⚠️ 在PE回归到行业Q1的极端情况下，预期收益为负({return_q1:+.2f}%)，存在一定估值回调风险，但风险可控')
    else:
        add_paragraph(document, f'❌ 在PE回归到行业Q1的极端情况下，预期收益大幅为负({return_q1:+.2f}%)，估值回调风险较高，需谨慎投资')

    add_paragraph(document, '')

    # 计算定增收益影响
    issue_price = project_params['issue_price']
    return_pe_stress = (target_price_q1 - issue_price) / issue_price * 100
    add_paragraph(document, '💡 对定增收益的影响：')
    add_paragraph(document, f'• 发行价：{issue_price:.2f}元')
    add_paragraph(document, f'• PE回归Q1后的定增收益率：{return_pe_stress:+.2f}%')
    if return_pe_stress < 0:
        add_paragraph(document, f'• ⚠️ 估值回归将导致定增亏损{abs(return_pe_stress):.2f}%')
    else:
        add_paragraph(document, f'• ✅ 估值回归后定增仍能保持盈利')

    # ==================== 7.2 经济面极端情况压力测试 ====================
    add_paragraph(document, '')
    add_title(document, '7.2 经济面极端情况压力测试', level=2)

    add_paragraph(document, '本节模拟经济面极端情况（历史危机、黑天鹅事件）对定增项目的影响。')
    add_paragraph(document, '包括2008年金融危机、2020年疫情、行业政策收紧、个股重大利空等多种极端情景。')
    add_paragraph(document, '')

    # 定义完整的压力测试情景（与notebook一致）
    add_paragraph(document, '7.2.1 压力情景定义', bold=True)
    add_paragraph(document, '')

    # 定义完整的压力测试情景
    stress_scenarios = {
        '市场危机_2008': {
            'description': '模拟2008年金融危机，股价下跌60%',
            'price_drop': 0.60,
            'volatility_spike': 2.0
        },
        '市场危机_2020': {
            'description': '模拟2020年疫情，股价下跌40%',
            'price_drop': 0.40,
            'volatility_spike': 1.5
        },
        '行业政策收紧': {
            'description': '行业监管政策收紧，股价下跌25%',
            'price_drop': 0.25,
            'volatility_spike': 1.2
        },
        '个股重大利空': {
            'description': '公司业绩暴雷，股价下跌35%',
            'price_drop': 0.35,
            'volatility_spike': 1.8
        },
        '流动性危机': {
            'description': '市场流动性枯竭，股价下跌20%并波动率飙升',
            'price_drop': 0.20,
            'volatility_spike': 2.5
        },
        '极端熊市': {
            'description': '极端熊市情景，股价下跌50%',
            'price_drop': 0.50,
            'volatility_spike': 1.3
        }
    }

    # 计算各情景的结果
    scenario_names = []
    scenario_descriptions = []
    scenario_prices_list = []
    scenario_returns_list = []
    scenario_pnl_list = []

    current_price = market_data['current_price']
    issue_price = project_params['issue_price']
    issue_shares = project_params['issue_shares']

    for name, params in stress_scenarios.items():
        stressed_price = current_price * (1 - params['price_drop'])
        pnl_percent = ((stressed_price - issue_price) / issue_price) * 100
        pnl_amount = (stressed_price - issue_price) * issue_shares

        scenario_names.append(name)
        scenario_descriptions.append(params['description'])
        scenario_prices_list.append(stressed_price)
        scenario_returns_list.append(pnl_percent)
        scenario_pnl_list.append(pnl_amount)

    # 按照下跌幅度排序（从最严重到最轻）
    sorted_indices = sorted(range(len(scenario_names)), key=lambda i: scenario_prices_list[i])
    scenario_names = [scenario_names[i] for i in sorted_indices]
    scenario_descriptions = [scenario_descriptions[i] for i in sorted_indices]
    scenario_prices_list = [scenario_prices_list[i] for i in sorted_indices]
    scenario_returns_list = [scenario_returns_list[i] for i in sorted_indices]
    scenario_pnl_list = [scenario_pnl_list[i] for i in sorted_indices]

    # 保存图表
    stress_chart_path = os.path.join(IMAGES_DIR, '02_stress_test.png')
    generate_stress_test_chart(scenario_names, scenario_returns_list, stress_chart_path)

    # 二、压力测试结果
    add_paragraph(document, '')
    add_paragraph(document, '7.2.2 压力测试结果', bold=True)
    add_paragraph(document, '')

    # 生成详细的结果表格
    stress_table = []
    for i, name in enumerate(scenario_names):
        stress_table.append([
            name,
            scenario_descriptions[i],
            f'{scenario_prices_list[i]:.2f}',
            f'{scenario_returns_list[i]:+.2f}%',
            f'{scenario_pnl_list[i]/10000:+.2f}'
        ])

    add_table_data(document, ['情景', '描述', '受压后价格(元)', '收益率(%)', '盈亏(万元)'], stress_table)

    add_paragraph(document, '图表 7.1: 经济面极端情况压力测试')
    add_image(document, stress_chart_path)

    # 三、经济面压力测试结论
    add_paragraph(document, '')
    add_paragraph(document, '7.2.3 经济面压力测试结论', bold=True)
    add_paragraph(document, '')

    # 统计分析
    profit_scenarios = sum(1 for r in scenario_returns_list if r > 0)
    total_scenarios = len(scenario_returns_list)
    worst_loss = min(scenario_pnl_list) / 10000
    worst_loss_percent = min(scenario_returns_list)
    best_scenario_idx = scenario_returns_list.index(max(scenario_returns_list))
    worst_scenario_idx = scenario_returns_list.index(min(scenario_returns_list))

    add_paragraph(document, f'压力测试统计：')
    add_paragraph(document, f'• 总情景数: {total_scenarios} 种')
    add_paragraph(document, f'• 盈利情景: {profit_scenarios} 种')
    add_paragraph(document, f'• 亏损情景: {total_scenarios - profit_scenarios} 种')
    add_paragraph(document, f'• 盈利概率: {profit_scenarios/total_scenarios*100:.1f}%')

    add_paragraph(document, '')
    add_paragraph(document, f'风险分析：')
    add_paragraph(document, f'• 最大潜在亏损: {abs(worst_loss):.2f} 万元')
    add_paragraph(document, f'• 最大亏损率: {worst_loss_percent:.2f}%')
    add_paragraph(document, f'• 最坏情景: {scenario_names[worst_scenario_idx]}')
    add_paragraph(document, f'• 最好情景: {scenario_names[best_scenario_idx]}')

    # 风险等级评估
    if profit_scenarios >= total_scenarios * 0.7:
        risk_level = "低风险 - 大部分情景下盈利"
        risk_emoji = "🟢"
    elif profit_scenarios >= total_scenarios * 0.4:
        risk_level = "中等风险 - 约一半情景下盈利"
        risk_emoji = "🟡"
    else:
        risk_level = "高风险 - 大部分情景下亏损"
        risk_emoji = "🔴"

    add_paragraph(document, '')
    add_paragraph(document, f'{risk_emoji} 经济面压力测试评级: {risk_level}')

    # ==================== 7.3 多重敏感性指标极端情况压力测试 ====================
    add_paragraph(document, '')
    add_title(document, '7.3 多重敏感性指标极端情况压力测试', level=2)

    add_paragraph(document, '本节分析当多个敏感性指标同时发生极端情况时的最差情景。')
    add_paragraph(document, '通过组合最不利的参数（深度溢价发行 + 高波动率 + 负向漂移率），评估项目的风险承受边界。')
    add_paragraph(document, '')

    add_paragraph(document, '7.3.1 极端情景组合定义', bold=True)
    add_paragraph(document, '')

    # 定义多重极端情景
    # 同时考虑：
    # 1. 深度溢价发行（+20%溢价）
    # 2. 高波动率（当前波动率 × 1.5）
    # 3. 负向漂移率（-10%年化收益率）

    current_vol_120d = market_data.get('volatility_120d', market_data.get('volatility', 0.30))
    current_drift_120d = market_data.get('annual_return_120d', market_data.get('drift', 0.08))
    current_price_multi = market_data['current_price']
    lockup_months = project_params['lockup_period']

    # 情景1：深度溢价 + 高波动 + 负向漂移（三重打击）
    premium_rate_extreme = 0.20  # 20%溢价
    vol_multiplier_extreme = 1.5  # 波动率放大1.5倍
    drift_extreme = -0.10  # -10%年化漂移率

    issue_price_extreme = current_price_multi * (1 + premium_rate_extreme)
    vol_extreme = current_vol_120d * vol_multiplier_extreme
    drift_rate_extreme = drift_extreme

    # 计算三重打击情景下的结果
    lockup_drift_extreme = drift_rate_extreme * (lockup_months / 12)
    lockup_vol_extreme = vol_extreme * np.sqrt(lockup_months / 12)

    # 使用蒙特卡洛模拟计算
    n_sim_extreme = 5000
    np.random.seed(42)
    returns_extreme = np.random.normal(lockup_drift_extreme, lockup_vol_extreme, n_sim_extreme)
    final_prices_extreme = current_price_multi * np.exp(returns_extreme)

    # 计算收益率
    returns_pnl_extreme = (final_prices_extreme - issue_price_extreme) / issue_price_extreme
    annualized_returns_extreme = (1 + returns_pnl_extreme) ** (12 / lockup_months) - 1

    # 统计结果
    mean_return_extreme = annualized_returns_extreme.mean() * 100
    median_return_extreme = np.median(annualized_returns_extreme) * 100
    profit_prob_extreme = (returns_pnl_extreme > 0).mean() * 100
    var_95_extreme = np.percentile(annualized_returns_extreme, 5) * 100
    var_99_extreme = np.percentile(annualized_returns_extreme, 1) * 100
    worst_loss_extreme = np.min(annualized_returns_extreme) * 100

    # 构建对比表格
    add_paragraph(document, '极端情景参数设置：')
    extreme_scenario_headers = ['参数', '正常值', '极端值', '变化幅度']
    extreme_scenario_data = [
        ['发行价', f'{project_params["issue_price"]:.2f}元', f'{issue_price_extreme:.2f}元', f'+{premium_rate_extreme*100:.0f}%溢价'],
        ['波动率', f'{current_vol_120d*100:.2f}%', f'{vol_extreme*100:.2f}%', f'×{vol_multiplier_extreme:.1f}'],
        ['漂移率', f'{current_drift_120d*100:+.2f}%', f'{drift_rate_extreme*100:+.2f}%', f'极端悲观'],
        ['窗口期', '120日', '120日', '保持不变'],
        ['锁定期', f'{lockup_months}个月', f'{lockup_months}个月', '保持不变']
    ]
    add_table_data(document, extreme_scenario_headers, extreme_scenario_data)

    add_paragraph(document, '')
    add_paragraph(document, '7.3.2 极端情景模拟结果', bold=True)
    add_paragraph(document, '')

    extreme_results_headers = ['指标', '正常情景', '极端情景（三重打击）', '差异']
    extreme_results_data = [
        ['预期年化收益', f'{current_drift_120d*100:+.2f}%', f'{mean_return_extreme:+.2f}%', f'{mean_return_extreme - current_drift_120d*100:+.2f}%'],
        ['中位数收益', f'{current_drift_120d*100:+.2f}%', f'{median_return_extreme:+.2f}%', f'{median_return_extreme - current_drift_120d*100:+.2f}%'],
        ['盈利概率', '待计算', f'{profit_prob_extreme:.1f}%', f'{profit_prob_extreme - 50:+.1f}%'],
        ['95% VaR', '待计算', f'{var_95_extreme:+.2f}%', f'极端损失'],
        ['99% VaR', '待计算', f'{var_99_extreme:+.2f}%', f'极端损失'],
        ['最差情况', '待计算', f'{worst_loss_extreme:+.2f}%', f'极端损失']
    ]
    add_table_data(document, extreme_results_headers, extreme_results_data)

    add_paragraph(document, '')
    add_paragraph(document, '7.3.3 极端情景分析', bold=True)
    add_paragraph(document, '')

    # 风险等级评估
    if profit_prob_extreme >= 40:
        extreme_risk_level = "中等风险 - 即使三重打击仍有40%以上盈利概率"
        risk_emoji_extreme = "🟡"
    elif profit_prob_extreme >= 20:
        extreme_risk_level = "高风险 - 盈利概率显著降低，但仍有20%以上机会"
        risk_emoji_extreme = "🟠"
    else:
        extreme_risk_level = "极高风险 - 三重打击下盈利概率低于20%，需极度谨慎"
        risk_emoji_extreme = "🔴"

    add_paragraph(document, f'{risk_emoji_extreme} 极端情景评级: {extreme_risk_level}')
    add_paragraph(document, '')

    add_paragraph(document, '关键发现：')
    add_paragraph(document, f'• 在深度溢价（+20%）、高波动（×1.5）、负向漂移（-10%）的三重打击下：')
    add_paragraph(document, f'  - 预期年化收益率为{mean_return_extreme:+.2f}%')
    add_paragraph(document, f'  - 盈利概率为{profit_prob_extreme:.1f}%')
    add_paragraph(document, f'  - 95% VaR为{var_95_extreme:+.2f}%，表示95%置信度下最大损失为{abs(var_95_extreme):.2f}%')
    add_paragraph(document, f'  - 最差情况可能损失{abs(worst_loss_extreme):.2f}%')
    add_paragraph(document, '')

    if profit_prob_extreme < 20:
        add_paragraph(document, '⚠️ 风险提示：')
        add_paragraph(document, '  在多重极端情况叠加下，项目面临极高的亏损风险。')
        add_paragraph(document, '  建议采取以下风险控制措施：')
        add_paragraph(document, '  1. 严格控制仓位，建议不超过总资产的5-10%')
        add_paragraph(document, '  2. 设置严格的止损线（如-15%）')
        add_paragraph(document, '  3. 考虑购买看跌期权进行对冲')
        add_paragraph(document, '  4. 分批建仓，避免一次性大额投入')
    elif profit_prob_extreme < 40:
        add_paragraph(document, '⚠️ 风险提示：')
        add_paragraph(document, '  在多重极端情况叠加下，项目风险显著上升。')
        add_paragraph(document, '  建议采取以下风险控制措施：')
        add_paragraph(document, '  1. 适度控制仓位')
        add_paragraph(document, '  2. 设置止损线（如-20%）')
        add_paragraph(document, '  3. 密切监控市场动态')
    else:
        add_paragraph(document, '✅ 风险评估：')
        add_paragraph(document, '  即使在多重极端情况下，项目仍保持一定的抗风险能力。')
        add_paragraph(document, '  但仍需警惕市场风险，做好仓位管理。')

    # ==================== 7.4 压力测试综合结论 ====================
    add_paragraph(document, '')
    add_title(document, '7.4 压力测试综合结论', level=2)

    add_paragraph(document, '本节综合前面7.1、7.2、7.3的压力测试结果，总结定增项目在各类极端情况下的风险表现。')
    add_paragraph(document, '')

    add_paragraph(document, '7.4.1 压力测试全景汇总', bold=True)
    add_paragraph(document, '')

    # 创建综合汇总表
    summary_headers = ['压力测试类型', '测试情景', '最差结果', '风险评估', '风险等级']
    summary_data = [
        ['7.1 PE回归压力测试', f'PE回归至行业Q1({pe_q1:.2f}倍)', f'{return_pe_stress:+.2f}%',
         '估值回归风险' if return_pe_stress < 0 else '估值安全边际充足',
         '🟢低风险' if return_pe_stress > 0 else '🟡中等风险' if return_pe_stress > -10 else '🔴高风险'],
        ['7.2 经济面极端情况', f'{scenario_names[worst_scenario_idx]}（股价下跌{int(stress_scenarios[scenario_names[worst_scenario_idx]]["price_drop"]*100)}%）',
         f'{worst_loss_percent:+.2f}%', f'最大亏损{abs(worst_loss):.2f}万元',
         '🟢低风险' if profit_scenarios >= total_scenarios * 0.7 else '🟡中等风险' if profit_scenarios >= total_scenarios * 0.4 else '🔴高风险'],
        ['7.3 多重敏感性指标极端', '深度溢价+高波动+负向漂移', f'{mean_return_extreme:+.2f}%',
         f'盈利概率{profit_prob_extreme:.1f}%',
         '🟡中等风险' if profit_prob_extreme >= 40 else '🟠较高风险' if profit_prob_extreme >= 20 else '🔴极高风险']
    ]
    add_table_data(document, summary_headers, summary_data)

    add_paragraph(document, '')
    add_paragraph(document, '7.4.2 最坏风险情景分析', bold=True)
    add_paragraph(document, '')

    # 识别所有压力测试中的最坏情况
    worst_results = {
        'PE回归': return_pe_stress,
        '经济面危机': worst_loss_percent,
        '多重极端': mean_return_extreme
    }

    worst_scenario_all = min(worst_results, key=worst_results.get)
    worst_loss_all = worst_results[worst_scenario_all]

    add_paragraph(document, f'⚠️ 全局最坏风险情景：{worst_scenario_all}')
    add_paragraph(document, f'   在所有压力测试中，最严重的损失情景为"{worst_scenario_all}"，预期年化收益率为{worst_loss_all:+.2f}%')
    add_paragraph(document, '')

    if worst_loss_all > -10:
        add_paragraph(document, '✅ 风险评估：即使在最坏情况下，损失幅度相对可控（<10%），项目具备较强的抗风险能力')
    elif worst_loss_all > -30:
        add_paragraph(document, '⚠️ 风险评估：在最坏情况下损失幅度为10-30%，风险可控但需做好仓位管理')
    elif worst_loss_all > -50:
        add_paragraph(document, '🟠 风险评估：在最坏情况下损失幅度为30-50%，风险较高，建议严格控制仓位')
    else:
        add_paragraph(document, '🔴 风险评估：在最坏情况下可能损失超过50%，风险极高，需极度谨慎或避免参与')

    add_paragraph(document, '')
    add_paragraph(document, '7.4.3 压力测试评分体系说明', bold=True)
    add_paragraph(document, '')

    add_paragraph(document, '为便于理解压力测试的风险评定逻辑，以下说明评分标准和等级划分。')
    add_paragraph(document, '')

    add_paragraph(document, '1. 评分维度与标准', bold=True)
    stress_scoring_explanation = [
        ['测试维度', '满分', '评分标准', '当前得分'],
        ['PE回归测试', '3分', '收益>0%=1分, -10%~0%=2分, <-10%=3分', f'{1 if return_pe_stress > 0 else 2 if return_pe_stress > -10 else 3}分'],
        ['经济面危机测试', '3分', '盈利≥70%=1分, 40%~70%=2分, <40%=3分', f'{1 if profit_scenarios >= total_scenarios * 0.7 else 2 if profit_scenarios >= total_scenarios * 0.4 else 3}分'],
        ['多重极端叠加测试', '4分', '盈利≥40%=2分, 20%~40%=3分, <20%=4分', f'{2 if profit_prob_extreme >= 40 else 3 if profit_prob_extreme >= 20 else 4}分']
    ]
    add_table_data(document, ['测试维度', '满分', '评分标准', '当前得分'], stress_scoring_explanation)

    add_paragraph(document, '')
    add_paragraph(document, '2. 综合评分计算', bold=True)
    add_paragraph(document, '   计算公式：综合评分 = (PE回归得分 + 经济面危机得分 + 多重极端得分) / 3')
    add_paragraph(document, '   注意：评分采用1-4分制，分数越高表示风险越大')
    add_paragraph(document, '')

    add_paragraph(document, '3. 风险等级划分', bold=True)
    stress_risk_levels = [
        ['平均分范围', '风险等级', '标识', '含义'],
        ['< 1.5分', '低风险', '🟢', '各项压力测试表现良好'],
        ['1.5 - 2.5分', '中等风险', '🟡', '部分压力测试存在风险'],
        ['2.5 - 3.5分', '较高风险', '🟠', '多项压力测试风险较高'],
        ['≥ 3.5分', '高风险', '🔴', '整体风险极高']
    ]
    add_table_data(document, ['平均分范围', '风险等级', '标识', '含义'], stress_risk_levels)

    add_paragraph(document, '')
    add_paragraph(document, '7.4.4 综合风险等级评定', bold=True)
    add_paragraph(document, '')

    # 综合风险等级计算
    risk_scores = []

    # PE回归风险评分
    if return_pe_stress > 0:
        risk_scores.append(1)  # 低风险
    elif return_pe_stress > -10:
        risk_scores.append(2)  # 中等风险
    else:
        risk_scores.append(3)  # 高风险

    # 经济面压力测试风险评分
    if profit_scenarios >= total_scenarios * 0.7:
        risk_scores.append(1)  # 低风险
    elif profit_scenarios >= total_scenarios * 0.4:
        risk_scores.append(2)  # 中等风险
    else:
        risk_scores.append(3)  # 高风险

    # 多重极端情况风险评分
    if profit_prob_extreme >= 40:
        risk_scores.append(2)  # 中等风险
    elif profit_prob_extreme >= 20:
        risk_scores.append(3)  # 较高风险
    else:
        risk_scores.append(4)  # 极高风险

    avg_risk_score = sum(risk_scores) / len(risk_scores)

    if avg_risk_score < 1.5:
        overall_risk = "低风险"
        overall_emoji = "🟢"
        risk_recommendation = "项目整体抗风险能力强，可正常参与投资"
    elif avg_risk_score < 2.5:
        overall_risk = "中等风险"
        overall_emoji = "🟡"
        risk_recommendation = "项目整体风险适中，建议适度控制仓位并设置止损"
    elif avg_risk_score < 3.5:
        overall_risk = "较高风险"
        overall_emoji = "🟠"
        risk_recommendation = "项目整体风险较高，建议严格控制仓位并做好风险对冲"
    else:
        overall_risk = "高风险"
        overall_emoji = "🔴"
        risk_recommendation = "项目整体风险极高，建议谨慎参与或避免投资"

    add_paragraph(document, f'{overall_emoji} 综合风险等级: {overall_risk}')
    add_paragraph(document, f'风险评分: {avg_risk_score:.2f}/4.0 （数值越高风险越大）')
    add_paragraph(document, '')
    add_paragraph(document, f'投资建议: {risk_recommendation}')

    add_paragraph(document, '')
    add_paragraph(document, '7.4.5 压力测试核心结论', bold=True)
    add_paragraph(document, '')

    add_paragraph(document, '通过三个维度的压力测试，得出以下核心结论：')
    add_paragraph(document, '')

    # 根据各项测试结果动态生成结论
    conclusions = []

    # PE回归结论
    if return_pe_stress > 0:
        conclusions.append(f'✅ 估值回归：即使PE回归至行业Q1分位，定增收益仍为正({return_pe_stress:+.2f}%)，估值安全边际充足')
    else:
        conclusions.append(f'⚠️ 估值回归：PE回归至行业Q1分位将导致亏损({return_pe_stress:+.2f}%)，需警惕估值回调风险')

    # 经济面结论
    if profit_scenarios >= total_scenarios * 0.6:
        conclusions.append(f'✅ 经济面危机：在{total_scenarios}种经济面极端情景中，有{profit_scenarios}种情景盈利，抗风险能力较强')
    else:
        conclusions.append(f'⚠️ 经济面危机：在{total_scenarios}种经济面极端情景中，仅{profit_scenarios}种情景盈利，抗风险能力较弱')

    # 多重极端结论
    if profit_prob_extreme >= 30:
        conclusions.append(f'✅ 多重极端叠加：即使三重打击，盈利概率仍达{profit_prob_extreme:.1f}%，具备一定抗风险能力')
    else:
        conclusions.append(f'🔴 多重极端叠加：三重打击下盈利概率仅{profit_prob_extreme:.1f}%，多重风险叠加后果严重')

    for i, conclusion in enumerate(conclusions, 1):
        add_paragraph(document, f'{i}. {conclusion}')

    add_paragraph(document, '')
    add_paragraph(document, '7.4.6 风险控制建议', bold=True)
    add_paragraph(document, '')

    if avg_risk_score < 2:
        add_paragraph(document, '基于压力测试结果，项目整体风险可控，建议：')
        add_paragraph(document, '• 正常配置仓位（可占总资产10-20%）')
        add_paragraph(document, '• 设置预警线（如-10%），关注市场变化')
        add_paragraph(document, '• 定期review项目基本面，及时调整策略')
    elif avg_risk_score < 3:
        add_paragraph(document, '基于压力测试结果，项目存在一定风险，建议：')
        add_paragraph(document, '• 适度控制仓位（建议占总资产5-10%）')
        add_paragraph(document, '• 设置止损线（如-15%），严格执行')
        add_paragraph(document, '• 考虑分批建仓，降低单点风险')
        add_paragraph(document, '• 做好对冲准备，如购买看跌期权')
    else:
        add_paragraph(document, '基于压力测试结果，项目风险较高，建议：')
        add_paragraph(document, '• 严格控制仓位（建议不超过总资产5%）')
        add_paragraph(document, '• 设置严格止损线（如-10%），果断执行')
        add_paragraph(document, '• 分批小额建仓，或考虑放弃投资')
        add_paragraph(document, '• 如必须参与，务必做好风险对冲')

    add_paragraph(document, '')
    add_paragraph(document, '💡 特别提示：')
    add_paragraph(document, '压力测试结果仅供参考，实际市场情况可能超出历史极端情景。')
    add_paragraph(document, '投资者应结合自身风险承受能力，审慎决策，并做好充分的风险管理准备。')
    add_paragraph(document, '')

    add_section_break(document)

    # ==================== 八、VaR风险度量 ====================
    add_title(document, '八、VaR风险度量', level=1)

    add_paragraph(document, '本章节使用多种方法计算风险价值（VaR）和条件风险价值（CVaR），全面评估定增项目的下行风险。')
    add_paragraph(document, '通过多窗口期VaR对比、增量风险分析等方法，提供更全面的风险度量视角。')

    add_title(document, '8.1 VaR计算参数说明', level=2)

    add_paragraph(document, 'VaR（Value at Risk）风险价值表示在给定置信水平下，投资组合可能遭受的最大损失。')
    add_paragraph(document, '本报告使用蒙特卡洛模拟法计算VaR，基于10,000次模拟路径，分别计算60日、120日、250日三个窗口期的VaR。')
    add_paragraph(document, '')

    add_paragraph(document, 'VaR计算参数说明（多窗口期对比）：', bold=True)

    # 多窗口期参数
    var_windows = {
        '60日': {
            'vol': mc_volatility_60d,
            'drift': mc_drift_60d,
            'days': 60
        },
        '120日': {
            'vol': mc_volatility_120d,
            'drift': mc_drift_120d,
            'days': 120
        },
        '250日': {
            'vol': mc_volatility_250d,
            'drift': mc_drift_250d,
            'days': 250
        }
    }

    # 计算每个窗口期的VaR
    var_results = {}
    lockup_months = project_params['lockup_period']
    n_sim_var = 5000  # 模拟次数

    for window_name, params in var_windows.items():
        window_vol = params['vol']
        window_drift = params['drift']
        time_steps = params['days']

        # 计算锁定期的参数
        lockup_drift = window_drift * (lockup_months / 12)
        lockup_vol = window_vol * np.sqrt(lockup_months / 12)

        # 蒙特卡洛模拟
        np.random.seed(42)
        returns_var = np.random.normal(lockup_drift, lockup_vol, n_sim_var)
        final_prices_var = project_params['current_price'] * np.exp(returns_var)
        profit_losses_var = (final_prices_var - project_params['issue_price']) / project_params['issue_price']

        # 计算不同置信水平的VaR和CVaR
        var_90 = abs(np.percentile(profit_losses_var, 10))
        var_95 = abs(np.percentile(profit_losses_var, 5))
        var_99 = abs(np.percentile(profit_losses_var, 1))
        cvar_95 = abs(profit_losses_var[profit_losses_var <= np.percentile(profit_losses_var, 5)].mean())
        cvar_99 = abs(profit_losses_var[profit_losses_var <= np.percentile(profit_losses_var, 1)].mean())

        var_results[window_name] = {
            'vol': window_vol,
            'drift': window_drift,
            'var_90': var_90,
            'var_95': var_95,
            'var_99': var_99,
            'cvar_95': cvar_95,
            'cvar_99': cvar_99,
            'profit_losses': profit_losses_var
        }

    # 创建参数说明表格
    var_params_headers = ['参数', '60日窗口', '120日窗口', '250日窗口']
    var_params_data = [
        ['数据窗口', '60日（季度）', '120日（半年线）', '250日（年线）'],
        ['年化波动率', f'{var_results["60日"]["vol"]*100:.2f}%', f'{var_results["120日"]["vol"]*100:.2f}%', f'{var_results["250日"]["vol"]*100:.2f}%'],
        ['年化漂移率', f'{var_results["60日"]["drift"]*100:+.2f}%', f'{var_results["120日"]["drift"]*100:+.2f}%', f'{var_results["250日"]["drift"]*100:+.2f}%'],
        ['锁定期', f'{lockup_months}个月', f'{lockup_months}个月', f'{lockup_months}个月'],
        ['模拟次数', '5,000次', '5,000次', '5,000次']
    ]

    # 添加项目参数
    ma120_var = market_data.get('ma_120', 0)
    if ma120_var > 0:
        discount_var = (project_params["issue_price"] - ma120_var) / ma120_var * 100
        discount_note = f"（相对MA120: {ma120_var:.2f}元）"
    else:
        discount_var = (project_params["issue_price"] - project_params["current_price"]) / project_params["current_price"] * 100
        discount_note = f"（相对当前价: {project_params['current_price']:.2f}元）"

    var_params_data.extend([
        ['发行价格', f'{project_params["issue_price"]:.2f}元/股', f'{project_params["issue_price"]:.2f}元/股', f'{project_params["issue_price"]:.2f}元/股'],
        ['当前价格', f'{project_params["current_price"]:.2f}元/股', f'{project_params["current_price"]:.2f}元/股', f'{project_params["current_price"]:.2f}元/股'],
        ['折价/溢价率', f'{discount_var:+.2f}% {discount_note}', f'{discount_var:+.2f}% {discount_note}', f'{discount_var:+.2f}% {discount_note}']
    ])

    add_table_data(document, var_params_headers, var_params_data)

    add_paragraph(document, '')
    add_paragraph(document, '💡 参数说明：')
    add_paragraph(document, '• 60日窗口：反映短期波动特征，适合评估短期风险')
    add_paragraph(document, '• 120日窗口：反映中期波动特征，平衡稳定性和时效性（推荐）')
    add_paragraph(document, '• 250日窗口：反映长期波动特征，数据最稳定，适合长期风险评估')
    add_paragraph(document, '• VaR值表示定增项目在到期时的收益率损失风险，正值表示亏损幅度')
    add_paragraph(document, '')

    add_title(document, '8.2 多窗口期VaR对比', level=2)

    add_paragraph(document, '本节对比不同时间窗口下的VaR计算结果，分析窗口期选择对风险度量的影响。')
    add_paragraph(document, '')

    # 创建VaR对比表格
    var_comparison_headers = ['窗口期', '95% VaR', '99% VaR', '95% CVaR', '99% CVaR', '风险等级']

    # 定义风险等级
    def get_var_risk_level(var_value):
        if var_value <= 0.15:
            return "🟢 低风险"
        elif var_value <= 0.25:
            return "🟡 中等风险"
        elif var_value <= 0.40:
            return "🟠 较高风险"
        else:
            return "🔴 高风险"

    var_comparison_data = []
    for window in ['60日', '120日', '250日']:
        result = var_results[window]
        var_comparison_data.append([
            window,
            f'{result["var_95"]*100:.2f}%',
            f'{result["var_99"]*100:.2f}%',
            f'{result["cvar_95"]*100:.2f}%',
            f'{result["cvar_99"]*100:.2f}%',
            get_var_risk_level(result["var_95"])
        ])

    add_table_data(document, var_comparison_headers, var_comparison_data)

    add_paragraph(document, '')
    add_paragraph(document, '多窗口期VaR对比分析：', bold=True)
    add_paragraph(document, '')

    # 分析窗口期差异
    var_95_values = [var_results[w]['var_95'] for w in ['60日', '120日', '250日']]
    var_95_min = min(var_95_values)
    var_95_max = max(var_95_values)
    var_95_range = var_95_max - var_95_min
    add_paragraph(document, f'• VaR数值范围：95% VaR在{var_95_min*100:.2f}%至{var_95_max*100:.2f}%之间，差异{var_95_range*100:.2f}%')
    add_paragraph(document, f'• VaR数值范围：95% VaR在{var_95_min*100:.2f}%至{var_95_max*100:.2f}%之间，差异{var_95_range*100:.2f}个%')
    add_paragraph(document, f'• 60日窗口VaR：{var_results["60日"]["var_95"]*100:.2f}%（{"短期波动较大" if var_results["60日"]["var_95"] > var_results["120日"]["var_95"] else "短期风险相对较低"}）')
    add_paragraph(document, f'• 120日窗口VaR：{var_results["120日"]["var_95"]*100:.2f}%（中期基准，推荐使用）')
    add_paragraph(document, f'• 250日窗口VaR：{var_results["250日"]["var_95"]*100:.2f}%（长期趋势，{"数据更稳定" if var_results["250日"]["vol"] < var_results["60日"]["vol"] else "波动特征"}）')

    # 推荐窗口期
    recommended_window = '120日'
    recommended_var = var_results[recommended_window]['var_95']
    recommended_cvar = var_results[recommended_window]['cvar_95']

    add_paragraph(document, '')
    add_paragraph(document, f'💡 推荐窗口期：{recommended_window}')
    add_paragraph(document, f'   理由：平衡了数据稳定性和时效性，95% VaR为{recommended_var*100:.2f}%，95% CVaR为{recommended_cvar*100:.2f}%')
    add_paragraph(document, '')

    # 使用推荐窗口期的数据作为后续分析基准
    var_90 = var_results[recommended_window]['var_90']
    var_95 = recommended_var
    var_99 = var_results[recommended_window]['var_99']
    cvar_95 = recommended_cvar
    cvar_99 = var_results[recommended_window]['cvar_99']
    profit_losses = var_results[recommended_window]['profit_losses']
    mc_volatility = var_results[recommended_window]['vol']
    mc_drift = var_results[recommended_window]['drift']

    # 保存图表
    var_chart_path = os.path.join(IMAGES_DIR, '05_var_analysis.png')
    generate_var_chart(var_95, var_99, cvar_95, var_chart_path)

    add_title(document, '8.3 不同置信水平下的VaR', level=2)

    # VaR表格（使用推荐窗口期的数据）
    var_table_data = [
        ['90%', f'{abs(np.percentile(profit_losses, 10))*100:.1f}%', f'{abs(profit_losses[profit_losses <= np.percentile(profit_losses, 10)].mean())*100:.1f}%', '有10%的概率损失超过此值'],
        ['95%', f'{var_95*100:.1f}%', f'{cvar_95*100:.1f}%', '有5%的概率损失超过此值'],
        ['99%', f'{var_99*100:.1f}%', f'{cvar_99*100:.1f}%', '有1%的概率损失超过此值']
    ]
    add_table_data(document, ['置信水平', 'VaR', 'CVaR', '说明'], var_table_data)

    add_paragraph(document, '')
    add_paragraph(document, '图表 8.1: VaR风险度量')
    add_image(document, var_chart_path, width=Inches(6))

    # ==================== 生成收益率分布图（显示VaR位置）====================
    add_paragraph(document, '')
    add_paragraph(document, '图表 8.2: 蒙特卡洛模拟收益率分布（VaR位置标注）')

    # 生成收益率分布图
    try:
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))

        # 1. 收益率分布直方图
        ax1.hist(profit_losses * 100, bins=50, color='steelblue', alpha=0.7, edgecolor='black')

        # 标注VaR位置（在负收益侧）
        ax1.axvline(x=-var_95*100, color='red', linestyle='--',
                   label=f'95% VaR (损失{var_95*100:.1f}%)', linewidth=2)
        ax1.axvline(x=-var_99*100, color='darkred', linestyle='--',
                   label=f'99% VaR (损失{var_99*100:.1f}%)', linewidth=2)
        ax1.axvline(x=0, color='green', linestyle='-', linewidth=2, label='盈亏平衡')

        ax1.set_xlabel('收益率 (%)', fontsize=12)
        ax1.set_ylabel('频数', fontsize=12)
        ax1.set_title('蒙特卡洛模拟 - 收益率分布（VaR位置标注）', fontsize=14, fontweight='bold')
        ax1.legend(loc='upper right')
        ax1.grid(True, alpha=0.3)

        # 2. VaR柱状图
        cl_labels = ['90%', '95%', '99%']
        var_values = [var_90 * 100, var_95 * 100, var_99 * 100]
        colors = ['#f39c12', '#e74c3c', '#8b0000']
        bars = ax2.bar(cl_labels, var_values, color=colors, alpha=0.7, edgecolor='black', linewidth=1.5)

        ax2.set_xlabel('置信水平', fontsize=12)
        ax2.set_ylabel('VaR (%损失)', fontsize=12)
        ax2.set_title('不同置信水平下的VaR对比', fontsize=14, fontweight='bold')
        ax2.grid(True, axis='y', alpha=0.3)

        # 添加数值标签
        for bar, value in zip(bars, var_values):
            ax2.text(bar.get_x() + bar.get_width()/2., bar.get_height(),
                    f'{value:.1f}%', ha='center', va='bottom', fontsize=11, fontweight='bold')

        plt.tight_layout()

        # 保存图表
        var_distribution_path = os.path.join(IMAGES_DIR, '08_var_distribution_with_var.png')
        plt.savefig(var_distribution_path, dpi=300, bbox_inches='tight')
        plt.close()

        add_image(document, var_distribution_path, width=Inches(6.5))

        add_paragraph(document, '')
        add_paragraph(document, '分布图解读：', bold=True)
        add_paragraph(document, f'• 左图：收益率分布直方图，红色虚线标注95%和99% VaR位置')
        add_paragraph(document, f'• 右图：不同置信水平的VaR对比柱状图')
        add_paragraph(document, f'• 95% VaR = {var_95*100:.1f}%：有95%的概率亏损不超过此值')
        add_paragraph(document, f'• 99% VaR = {var_99*100:.1f}%：有99%的概率亏损不超过此值')
        add_paragraph(document, f'• 绿色实线：盈亏平衡点（0%）')

    except Exception as e:
        print(f"⚠️ 生成收益率分布图失败: {e}")

    add_title(document, '8.4 CVaR（条件风险价值）分析', level=2)

    add_paragraph(document, 'CVaR（Conditional Value at Risk）又称期望短缺（Expected Shortfall），表示在损失超过VaR时的平均损失。')
    add_paragraph(document, 'CVaR比VaR更保守，能更好地反映极端情况下的风险。')
    add_paragraph(document, '')

    cvar_comparison_data = [
        ['VaR vs CVaR (95%)', f'{var_95*100:.1f}%', f'{cvar_95*100:.1f}%', f'{(cvar_95-var_95)*100:.1f}%', f'{(cvar_95/var_95-1)*100:.1f}%'],
        ['VaR vs CVaR (99%)', f'{var_99*100:.1f}%', f'{cvar_99*100:.1f}%', f'{(cvar_99-var_99)*100:.1f}%', f'{(cvar_99/var_99-1)*100:.1f}%']
    ]
    add_table_data(document, ['项目', 'VaR', 'CVaR', '差值', 'CVaR/VaR'], cvar_comparison_data)

    add_paragraph(document, '')
    add_paragraph(document, 'CVaR分析结论：')
    add_paragraph(document, f'• 95%置信水平下，CVaR比VaR高出 {(cvar_95/var_95-1)*100:.1f}%')
    add_paragraph(document, f'• 99%置信水平下，CVaR比VaR高出 {(cvar_99/var_99-1)*100:.1f}%')
    add_paragraph(document, '• 说明极端情况下的损失比VaR估计的更严重')
    add_paragraph(document, '')

    # 增量风险分析
    add_paragraph(document, '💡 VaR vs CVaR 的区别：')
    add_paragraph(document, '• VaR：回答"有5%的概率损失至少多少"')
    add_paragraph(document, '• CVaR：回答"在最差5%的情况下，平均损失多少"')
    add_paragraph(document, '• CVaR > VaR，说明极端损失的"尾部"比VaR预测的更严重')
    add_paragraph(document, '')

    add_title(document, '8.5 最大回撤分析', level=2)

    add_paragraph(document, '最大回撤是指从峰值到谷底的最大跌幅，是衡量投资风险的重要指标。')
    add_paragraph(document, '')

    # 估算最大回撤（使用120日波动率）
    estimated_max_drawdown = mc_volatility * 2  # 简化估算
    estimated_95_drawdown = mc_volatility * 1.5

    drawdown_data = [
        ['预估平均最大回撤', f'{estimated_max_drawdown*100:.1f}%', f'基于120日波动率({mc_volatility*100:.1f}%)估算'],
        ['预估95%分位回撤', f'{estimated_95_drawdown*100:.1f}%', '95%的路径回撤不超过此值']
    ]
    add_table_data(document, ['回撤指标', '数值', '说明'], drawdown_data)

    add_paragraph(document, '')
    add_paragraph(document, '回撤分析：')
    add_paragraph(document, f'• 预估平均最大回撤：{estimated_max_drawdown*100:.1f}%')
    add_paragraph(document, f'• 预估95%分位回撤：{estimated_95_drawdown*100:.1f}%')
    add_paragraph(document, f'• 最大回撤与VaR的关系：最大回撤通常大于VaR，因为它衡量整个持有期的最差情况')
    add_paragraph(document, '')

    add_title(document, '8.6 潜在损失金额估算', level=2)

    add_paragraph(document, '本节将VaR和CVaR转换为具体金额，直观展示潜在损失规模。')
    add_paragraph(document, '')

    # 计算损失金额和剩余本金
    investment_amount = project_params['issue_price'] * project_params['issue_shares']
    loss_var_95 = investment_amount * var_95
    loss_cvar_95 = investment_amount * cvar_95
    loss_var_99 = investment_amount * var_99

    # 计算剩余本金
    remaining_var_95 = investment_amount - loss_var_95
    remaining_cvar_95 = investment_amount - loss_cvar_95
    remaining_var_99 = investment_amount - loss_var_99

    loss_amount_data = [
        ['投资总额', f'{investment_amount/10000:.2f}', '-', '-'],
        ['95% VaR潜在损失', f'{loss_var_95/10000:.2f}', f'{var_95*100:.2f}%', f'{remaining_var_95/10000:.2f}'],
        ['95% CVaR潜在损失', f'{loss_cvar_95/10000:.2f}', f'{cvar_95*100:.2f}%', f'{remaining_cvar_95/10000:.2f}'],
        ['99% VaR潜在损失', f'{loss_var_99/10000:.2f}', f'{var_99*100:.2f}%', f'{remaining_var_99/10000:.2f}']
    ]
    add_table_data(document, ['项目', '金额（万元）', '占投资比例', '剩余本金（万元）'], loss_amount_data)

    add_paragraph(document, '')
    add_paragraph(document, '💡 金额说明：')
    add_paragraph(document, f'• 投资总额：{investment_amount/10000:.2f}万元（发行价×发行股数）')
    add_paragraph(document, '• 占投资比例：潜在损失占投资总额的百分比（即亏损比例）')
    add_paragraph(document, '• 剩余本金：投资总额减去潜在损失后的余额')
    add_paragraph(document, f'• 95% VaR：95%概率亏损不超过{loss_var_95/10000:.2f}万元（{var_95*100:.2f}%），剩余本金{remaining_var_95/10000:.2f}万元')
    add_paragraph(document, f'• 95% CVaR：最差5%情况下平均损失{loss_cvar_95/10000:.2f}万元（{cvar_95*100:.2f}%），剩余本金{remaining_cvar_95/10000:.2f}万元')
    add_paragraph(document, '')

    add_title(document, '8.7 VaR风险测算综合结论', level=2)

    add_paragraph(document, '本节综合前面所有VaR分析，给出全面的风险评估结论。')
    add_paragraph(document, '')

    add_paragraph(document, '8.7.1 多窗口期VaR综合评估', bold=True)
    add_paragraph(document, '')

    # 总结三个窗口期的VaR
    var_summary_headers = ['窗口期', '95% VaR', '风险特征', '适用场景']
    var_summary_data = [
        [
            '60日',
            f'{var_results["60日"]["var_95"]*100:.2f}%',
            f'{"短期风险较高" if var_results["60日"]["var_95"] > var_results["120日"]["var_95"] else "短期风险可控"}',
            '短期交易、波段操作'
        ],
        [
            '120日',
            f'{var_results["120日"]["var_95"]*100:.2f}%',
            '中期基准，平衡稳定性',
            '中期投资（推荐）'
        ],
        [
            '250日',
            f'{var_results["250日"]["var_95"]*100:.2f}%',
            f'{"长期风险较低" if var_results["250日"]["var_95"] < var_results["120日"]["var_95"] else "长期风险特征"}',
            '长期投资、价值投资'
        ]
    ]

    add_table_data(document, var_summary_headers, var_summary_data)

    add_paragraph(document, '')
    add_paragraph(document, '8.7.2 VaR风险度量汇总', bold=True)
    add_paragraph(document, '')

    add_paragraph(document, f'基于{recommended_window}窗口的分析结果：')
    add_paragraph(document, f'• 95%置信水平下，最大可能亏损约 {var_95*100:.1f}%，约 {loss_var_95/10000:.2f} 万元')
    add_paragraph(document, f'• 极端情况下（1%概率），亏损可能达到 {var_99*100:.1f}%，约 {loss_var_99/10000:.2f} 万元')
    add_paragraph(document, f'• CVaR显示，在最差5%情况下平均损失约 {cvar_95*100:.1f}%，约 {loss_cvar_95/10000:.2f} 万元')
    add_paragraph(document, f'• 预估最大回撤约 {estimated_max_drawdown*100:.1f}%')
    add_paragraph(document, f'• CVaR比VaR高出 {(cvar_95/var_95-1)*100:.1f}%，说明尾部风险严重')
    add_paragraph(document, '')

    # 综合风险评级
    def get_comprehensive_risk_level(var_95_value, cvar_95_value):
        """综合VaR和CVaR进行风险评级"""
        # 加权平均：VaR权重0.6，CVaR权重0.4
        weighted_risk = var_95_value * 0.6 + cvar_95_value * 0.4

        if weighted_risk <= 0.15:
            return "低风险", "🟢", "VaR和CVaR均处于较低水平，下行风险有限"
        elif weighted_risk <= 0.25:
            return "中等风险", "🟡", "VaR和CVaR处于中等水平，需关注下行风险"
        elif weighted_risk <= 0.40:
            return "较高风险", "🟠", "VaR和CVaR均较高，下行风险较大，需严格控制"
        else:
            return "高风险", "🔴", "VaR和CVaR均很高，下行风险极大，需极度谨慎"

    # 计算综合风险评级
    risk_rating, risk_emoji, risk_description = get_comprehensive_risk_level(var_95, cvar_95)
    print(f"✅ 风险评级: {risk_rating} ({risk_emoji})")

    add_paragraph(document, '8.7.3 VaR分析的局限性', bold=True)
    add_paragraph(document, '')

    add_paragraph(document, '⚠️ VaR分析的重要提示：')
    add_paragraph(document, '• VaR基于历史数据，无法预测黑天鹅事件')
    add_paragraph(document, '• VaR假设正态分布，实际市场可能出现肥尾效应')
    add_paragraph(document, f'• {recommended_window}窗口的VaR为{var_95*100:.2f}%，但实际损失可能超过此值')
    add_paragraph(document, '• CVaR比VaR更保守，应重点关注CVaR（尾部风险）')
    add_paragraph(document, '• 建议结合压力测试（第七章）综合评估风险')
    add_paragraph(document, '')

    add_paragraph(document, '💡 总结：')
    add_paragraph(document, f'通过多窗口期VaR分析、CVaR分析和金额估算，我们得出{risk_rating}结论。')
    add_paragraph(document, f'投资者应根据自身风险承受能力，审慎决策。')
    add_paragraph(document, '')

    add_section_break(document)

    add_section_break(document)

    # ==================== 九、综合评估 ====================
    # ==================== 风险评分章节已隐藏 ====================
    # 如需启用，请取消下方注释

    # 计算风险评分变量（供后续章节使用，不显示在报告中）
    scores = {
        '盈利概率': 30 if profit_prob >= 0.7 else 20 if profit_prob >= 0.5 else 10,
        '发行折价': 20 if discount_premium <= -15 else 15 if discount_premium <= -5 else 10 if discount_premium <= 5 else 5,
        '锁定期': 20 if project_params['lockup_period'] <= 6 else 15 if project_params['lockup_period'] <= 12 else 5,
        '波动率': 15 if risk_params.get('volatility', 0.35) <= 0.3 else 10 if risk_params.get('volatility', 0.35) <= 0.4 else 5,
        'VaR': 15 if var_95 <= 0.2 else 10 if var_95 <= 0.3 else 3
    }
    total_score = sum(scores.values())

    # 计算风险等级和投资建议
    if total_score >= 80:
        current_risk_level = "🟢 低风险"
        current_advice = "建议积极参与"
    elif total_score >= 60:
        current_risk_level = "🟡 中等风险"
        current_advice = "建议谨慎参与"
    else:
        current_risk_level = "🔴 高风险"
        current_advice = "建议规避"

    # add_title(document, '九、综合评估', level=1)
    # add_paragraph(document, '本章节综合前面的分析结果，从风险控制和估值角度给出综合评分和报价建议。')
    # add_paragraph(document, '')

    # ==================== 9.0 风险评分体系说明 ====================
    # add_title(document, '9.0 风险评分体系说明', level=2)

    # add_paragraph(document, '本报告采用多维度、分层次的风险评分体系，从估值分析、市场模拟、压力测试等多个角度全面评估定增项目风险。')
    # add_paragraph(document, '为便于核验和优化，本节详细说明整个评分体系的架构、各章节评分标准以及综合评分的计算逻辑。')
    # add_paragraph(document, '')

    # add_paragraph(document, '一、评分体系架构', bold=True)
    # add_paragraph(document, '')

    # add_paragraph(document, '风险评分体系分为四个层次：')
    # add_paragraph(document, '')

    framework_data = [
        ['层次', '章节', '评分内容', '权重', '评分依据'],
        ['基础层', '第二章相对估值', '估值安全边际', '参考', 'PE/PB相对行业位置'],
        ['', '第三章DCF估值', '内在价值折扣', '参考', 'DCF估值 vs 发行价'],
        ['核心层', '第五章蒙特卡洛', '盈利概率与收益', '30%', '模拟统计结果'],
        ['', '第八章VaR度量', '下行风险控制', '15%', '95% VaR值'],
        ['高级层', '第四章敏感性', '参数敏感程度', '参考', '龙卷风图分析'],
        ['', '第七章压力测试', '极端情景抗风险', '参考', '三重压力测试'],
        ['综合层', '第九章综合评估', '多维度整合', '100%', '加权综合评分']
    ]
    # add_table_data(document, ['层次', '章节', '评分内容', '权重', '评分依据'], framework_data)

    # add_paragraph(document, '')
    # add_paragraph(document, '说明：')
    # add_paragraph(document, '• 基础层：提供估值参考，不直接计入综合评分')
    # add_paragraph(document, '• 核心层：主要评分依据，占综合评分的45%权重')
    # add_paragraph(document, '• 高级层：深度风险分析，辅助判断整体风险')
    # add_paragraph(document, '• 综合层：整合各维度，给出最终投资建议')
    # add_paragraph(document, '')

    # add_paragraph(document, '二、各章节评分标准', bold=True)
    # add_paragraph(document, '')
    """

    # add_paragraph(document, '1. 第二章：相对估值分析（参考性）', bold=True)
    # add_paragraph(document, '   评分维度：PE、PB相对行业的估值水平')
    # add_paragraph(document, '   评分标准：')
    # add_paragraph(document, '   • 估值低于行业中位数：相对安全')
    # add_paragraph(document, '   • 估值高于行业中位数：需要较高业绩增长支撑')
    # add_paragraph(document, '   • 估值高于行业75分位：高估值风险')
    # add_paragraph(document, '')

    # add_paragraph(document, '2. 第三章：DCF估值分析（参考性）', bold=True)
    # add_paragraph(document, '   评分维度：发行价相对于内在价值的折扣')
    # add_paragraph(document, '   评分标准：')
    # add_paragraph(document, '   • 折扣率 ≥ 20%：深度价值投资')
    # add_paragraph(document, '   • 折扣率 ≥ 10%：适度安全边际')
    # add_paragraph(document, '   • 折扣率 < 10%：安全边际不足')
    # add_paragraph(document, '   • 溢价发行：需要极高的成长性支撑')
    # add_paragraph(document, '')

    # add_paragraph(document, '3. 第四章：敏感性分析（参考性）', bold=True)
    # add_paragraph(document, '   评分维度：关键参数变化对盈利概率的影响程度')
    # add_paragraph(document, '   评分标准：')
    # add_paragraph(document, '   • 漂移率敏感度：每变化10%对盈利概率的影响')
    # add_paragraph(document, '   • 波动率敏感度：每变化10%对盈利概率的影响')
    # add_paragraph(document, '   • 发行价敏感度：折价率变化对盈利概率的影响')
    # add_paragraph(document, '   • 综合判断：基于龙卷风图识别主要风险因素')
    # add_paragraph(document, '')

    # add_paragraph(document, '4. 第五章：蒙特卡洛模拟（核心权重30%）', bold=True)
    # add_paragraph(document, '   评分维度：基于10,000次模拟的统计结果')
    mc_score_data = [
        ['指标', '优秀(30分)', '良好(20分)', '较差(10分)'],
        ['盈利概率', '≥ 70%', '50% - 70%', '< 50%'],
        ['预期年化收益', '≥ 20%', '10% - 20%', '< 10%'],
        ['中位数收益', '≥ 15%', '5% - 15%', '< 5%']
    ]
    # add_table_data(document, ['指标', '优秀(30分)', '良好(20分)', '较差(10分)'], mc_score_data)
    # add_paragraph(document, '')

    # add_paragraph(document, '5. 第七章：压力测试（高级参考）', bold=True)
    # add_paragraph(document, '   评分维度：极端情景下的抗风险能力')
    # add_paragraph(document, '   评分标准（1-4分制，分数越高风险越大）：')
    stress_score_data = [
        ['测试项目', '低风险(1分)', '中等风险(2分)', '较高风险(3分)', '极高风险(4分)'],
        ['PE回归测试', '收益>0%', '收益-10%~0%', '收益<-10%', '-'],
        ['经济面危机', '盈利≥70%', '盈利40%~70%', '盈利<40%', '-'],
        ['多重极端叠加', '-', '盈利≥40%', '盈利20%~40%', '盈利<20%']
    ]
    # add_table_data(document, ['测试项目', '低风险(1分)', '中等风险(2分)', '较高风险(3分)', '极高风险(4分)'], stress_score_data)
    # add_paragraph(document, '')

    # add_paragraph(document, '6. 第八章：VaR风险度量（核心权重15%）', bold=True)
    # add_paragraph(document, '   评分维度：下行风险度量')
    # add_paragraph(document, '   评分标准：')
    var_score_data = [
        ['VaR水平', '低风险(15分)', '中等风险(10分)', '高风险(5分)'],
        ['95% VaR', '≤ 20%', '20% - 30%', '> 30%'],
        ['99% VaR', '≤ 30%', '30% - 45%', '> 45%'],
        ['95% CVaR', '≤ 25%', '25% - 40%', '> 40%']
    ]
    # add_table_data(document, ['VaR水平', '低风险(15分)', '中等风险(10分)', '高风险(5分)'], var_score_data)
    # add_paragraph(document, '')

    # add_paragraph(document, '三、综合评分体系（第九章核心）', bold=True)
    # add_paragraph(document, '')

    # add_paragraph(document, '1. 评分维度与权重', bold=True)
    # add_paragraph(document, '')

    # 维度说明表格
    dimension_data = [
        ['评分维度', '权重', '评分依据', '数据来源'],
        ['盈利概率', '30分', '蒙特卡洛模拟的盈利概率', '第三章模拟结果'],
        ['发行折价', '20分', '发行价相对MA20的折价率', '市场数据'],
        ['锁定期', '20分', '股份锁定期时长', '项目参数'],
        ['波动率', '15分', '120日年化波动率', '历史数据'],
        ['VaR风险价值', '15分', '95%置信度下的VaR值', 'VaR计算结果']
    ]
    # add_table_data(document, ['评分维度', '权重', '评分依据', '数据来源'], dimension_data)

    # add_paragraph(document, '')
    # add_paragraph(document, '2. 各维度详细评分标准', bold=True)
    # add_paragraph(document, '')

    # add_paragraph(document, '1. 盈利概率（满分30分）')
    # add_paragraph(document, '   评分逻辑：')
    # add_paragraph(document, '   • 盈利概率 ≥ 70%：得30分（优秀）')
    # add_paragraph(document, '   • 盈利概率 ≥ 50%：得20分（良好）')
    # add_paragraph(document, '   • 盈利概率 < 50%：得10分（较差）')
    # add_paragraph(document, f'   当前值：{profit_prob*100:.1f}% → 得分：{30 if profit_prob >= 0.7 else 20 if profit_prob >= 0.5 else 10}分')
    # add_paragraph(document, '')

    # add_paragraph(document, '2. 发行溢价率（满分20分）')
    # add_paragraph(document, '   评分逻辑：基于发行价相对MA20的溢价率（premium_rate）')
    # add_paragraph(document, '   • 溢价率 ≤ -15%（深度折价，发行价≤MA20的85%）：得20分（优秀）')
    # add_paragraph(document, '   • 溢价率 -15%至-5%（适度折价，发行价是MA20的85%-95%）：得15分（良好）')
    # add_paragraph(document, '   • 溢价率 -5%至5%（平价，发行价是MA20的95%-105%）：得10分（一般）')
    # add_paragraph(document, '   • 溢价率 > 5%（溢价，发行价>MA20的105%）：得5分（较差）')
    # add_paragraph(document, f'   当前值：溢价率{discount_premium:+.1f}% → 得分：{20 if discount_premium <= -15 else 15 if discount_premium <= -5 else 10 if discount_premium <= 5 else 5}分')
    # add_paragraph(document, '')

    # add_paragraph(document, '3. 锁定期（满分20分）')
    # add_paragraph(document, '   评分逻辑：基于锁定期长度（单位：月）')
    # add_paragraph(document, '   • 锁定期 ≤ 6个月：得20分（短期）')
    # add_paragraph(document, '   • 锁定期 ≤ 12个月：得15分（中期）')
    # add_paragraph(document, '   • 锁定期 > 12个月：得5分（长期）')
    # add_paragraph(document, f'   当前值：锁定期{project_params["lockup_period"]}个月 → 得分：{20 if project_params["lockup_period"] <= 6 else 15 if project_params["lockup_period"] <= 12 else 5}分')
    # add_paragraph(document, '')

    # add_paragraph(document, '4. 波动率（满分15分）')
    # add_paragraph(document, '   评分逻辑：基于120日年化波动率')
    # add_paragraph(document, '   • 波动率 ≤ 30%：得15分（低波动）')
    # add_paragraph(document, '   • 波动率 ≤ 40%：得10分（中等波动）')
    # add_paragraph(document, '   • 波动率 > 40%：得5分（高波动）')
    vol_used = risk_params.get('volatility', 0.35)
    # add_paragraph(document, f'   当前值：波动率{vol_used*100:.1f}% → 得分：{15 if vol_used <= 0.3 else 10 if vol_used <= 0.4 else 5}分')
    # add_paragraph(document, '')

    # add_paragraph(document, '5. VaR风险价值（满分15分）')
    # add_paragraph(document, '   评分逻辑：基于95%置信度下的年化VaR值')
    # add_paragraph(document, '   • VaR ≤ 20%：得15分（低风险）')
    # add_paragraph(document, '   • VaR ≤ 30%：得10分（中等风险）')
    # add_paragraph(document, '   • VaR > 30%：得3分（高风险）')
    # add_paragraph(document, f'   当前值：VaR {var_95*100:.1f}% → 得分：{15 if var_95 <= 0.2 else 10 if var_95 <= 0.3 else 3}分')
    # add_paragraph(document, '')

    # add_paragraph(document, '三、综合评分计算', bold=True)
    # add_paragraph(document, '')

    # 综合评分
    scores = {
        '盈利概率': 30 if profit_prob >= 0.7 else 20 if profit_prob >= 0.5 else 10,
        '发行折价': 20 if discount_premium <= -15 else 15 if discount_premium <= -5 else 10 if discount_premium <= 5 else 5,
        '锁定期': 20 if project_params['lockup_period'] <= 6 else 15 if project_params['lockup_period'] <= 12 else 5,
        '波动率': 15 if risk_params.get('volatility', 0.35) <= 0.3 else 10 if risk_params.get('volatility', 0.35) <= 0.4 else 5,
        'VaR': 15 if var_95 <= 0.2 else 10 if var_95 <= 0.3 else 3
    }
    total_score = sum(scores.values())

    # add_paragraph(document, f'综合评分 = 各维度得分之和 = {total_score}分 / 100分')
    # add_paragraph(document, '')

    # add_paragraph(document, '四、风险等级划分', bold=True)
    # add_paragraph(document, '')

    # 风险等级表格
    risk_level_data = [
        ['综合评分', '风险等级', '投资建议', '适用场景'],
        ['80-100分', '🟢 低风险', '积极参与', '盈利概率高、安全边际充足'],
        ['60-79分', '🟡 中等风险', '谨慎参与', '盈利概率中等、需综合评估'],
        ['0-59分', '🔴 高风险', '建议规避', '盈利概率低、安全边际不足']
    ]
    # add_table_data(document, ['综合评分', '风险等级', '投资建议', '适用场景'], risk_level_data)

    # add_paragraph(document, '')
    # add_paragraph(document, f'当前项目综合评分：{total_score}/100分', bold=True)

    # 计算风险等级
    if total_score >= 80:
        current_risk_level = "🟢 低风险"
        current_advice = "积极参与"
    elif total_score >= 60:
        current_risk_level = "🟡 中等风险"
        current_advice = "谨慎参与"
    else:
        current_risk_level = "🔴 高风险"
        current_advice = "建议规避"

    # add_paragraph(document, f'当前风险等级：{current_risk_level}', bold=True)
    # add_paragraph(document, f'投资建议：{current_advice}', bold=True)
    # add_paragraph(document, '')

    # add_paragraph(document, '五、评分体系优化建议', bold=True)
    # add_paragraph(document, '')
    # add_paragraph(document, '如需调整评分体系，可考虑以下方向：')
    # add_paragraph(document, '• 调整权重：根据投资策略偏好，调整各维度的权重分配')
    # add_paragraph(document, '• 优化阈值：根据历史回测结果，优化各维度的评分阈值')
    # add_paragraph(document, '• 增加维度：根据行业特点，增加新的评分维度（如基本面、流动性等）')
    # add_paragraph(document, '• 动态调整：根据市场环境变化，动态调整评分标准')
    # add_paragraph(document, '')

    # ==================== 9.1 风险评分 ====================
    # add_title(document, '9.1 风险评分结果', level=2)

    # add_paragraph(document, '基于上述评分体系，对项目进行综合风险评分。')
    # add_paragraph(document, '')

    # 综合评分
    scores = {
        '盈利概率': 30 if profit_prob >= 0.7 else 20 if profit_prob >= 0.5 else 10,
        '发行折价': 20 if discount_premium <= -15 else 15 if discount_premium <= -5 else 10 if discount_premium <= 5 else 5,
        '锁定期': 20 if project_params['lockup_period'] <= 6 else 15 if project_params['lockup_period'] <= 12 else 5,
        '波动率': 15 if risk_params.get('volatility', 0.35) <= 0.3 else 10 if risk_params.get('volatility', 0.35) <= 0.4 else 5,
        'VaR': 15 if var_95 <= 0.2 else 10 if var_95 <= 0.3 else 3
    }
    total_score = sum(scores.values())

    # 保存雷达图
    radar_chart_path = os.path.join(IMAGES_DIR, '09_01_radar_chart.png')
    generate_radar_chart(scores, radar_chart_path)

    score_data = [[k, f'{v}/30' if k == '盈利概率' else f'{v}/20' if k != 'VaR' else f'{v}/15'] for k, v in scores.items()]
    score_data.append(['总分', f'{total_score}/100'])

    # add_table_data(document, ['评估维度', '得分'], score_data)

    # add_paragraph(document, '')
    # add_paragraph(document, '图表 9.1: 风险评分雷达图')
    add_image(document, radar_chart_path)

    # add_paragraph(document, '')
    # add_paragraph(document, f'综合风险评分: {total_score}/100 分')

    # add_title(document, '9.2 各维度详细评估', level=2)

    # add_paragraph(document, '以下对各维度进行详细分析，包括计算过程、评估依据和风险提示。')
    # add_paragraph(document, '')

    # add_paragraph(document, '1. 盈利概率分析', bold=True)
    profit_prob_pct = profit_prob * 100 if profit_prob <= 1 else profit_prob
    # add_paragraph(document, f'   计算结果：盈利概率 = {profit_prob_pct:.1f}%')
    # add_paragraph(document, f'   评分依据：基于蒙特卡洛模拟（10,000次）的结果')
    # add_paragraph(document, f'   计算方法：盈利次数 / 总模拟次数 × 100%')
    # add_paragraph(document, f'   得分：{scores["盈利概率"]}/30分')
    if profit_prob >= 0.7:
    # add_paragraph(document, '   评估：✅ 盈利概率较高（≥70%），投资吸引力强')
    elif profit_prob >= 0.5:
    # add_paragraph(document, '   评估：⚠️ 盈利概率中等（50%-70%），需谨慎评估')
    else:
    # add_paragraph(document, '   评估：❌ 盈利概率较低（<50%），风险较大')
    # add_paragraph(document, '')

    # add_paragraph(document, '2. 发行溢价率分析', bold=True)
    # add_paragraph(document, f'   计算结果：溢价率 = {discount_premium:+.1f}%')
    # add_paragraph(document, f'   计算公式：(发行价 - MA20) / MA20 × 100%')
    # add_paragraph(document, f'   发行价：{project_params["issue_price"]:.2f}元，MA20：{ma20:.2f}元')
    # add_paragraph(document, f'   说明：溢价率为负值表示折价（好），为正值表示溢价（不好）')
    # add_paragraph(document, f'   得分：{scores["发行折价"]}/20分')
    if discount_premium <= -15:  # 溢价率 ≤ -15%，即发行价 ≤ MA20的85%
    # add_paragraph(document, f'   评估：✅ 深度折价发行（溢价率≤-15%），安全边际充足')
    elif discount_premium <= -5:  # 溢价率 -15% ~ -5%，即发行价是MA20的85%-95%
    # add_paragraph(document, f'   评估：⚠️ 适度折价发行（溢价率-15%至-5%），安全边际一般')
    elif discount_premium <= 5:  # 溢价率 -5% ~ 5%，即发行价是MA20的95%-105%
    # add_paragraph(document, f'   评估：⚠️ 平价发行（溢价率-5%至5%），安全边际有限')
    else:  # 溢价率 > 5%，即发行价 > MA20的105%
    # add_paragraph(document, f'   评估：❌ 溢价发行（溢价率>5%），安全边际不足')
    # add_paragraph(document, '')

    # add_paragraph(document, '3. 锁定期分析', bold=True)
    lockup_months = project_params['lockup_period']
    # add_paragraph(document, f'   计算结果：锁定期 = {lockup_months}个月')
    # add_paragraph(document, f'   得分：{scores["锁定期"]}/20分')
    if lockup_months <= 6:
    # add_paragraph(document, '   评估：✅ 锁定期较短（≤6个月），资金占用时间有限')
    elif lockup_months <= 12:
    # add_paragraph(document, '   评估：⚠️ 锁定期中等（6-12个月），需要适度等待')
    else:
    # add_paragraph(document, '   评估：❌ 锁定期较长（>12个月），资金占用时间长，风险暴露期长')
    # add_paragraph(document, f'   风险提示：锁定期越长，面临的不确定性越大，流动性风险越高')
    # add_paragraph(document, '')

    # add_paragraph(document, '4. 波动率分析', bold=True)
    vol_used = risk_params.get('volatility', 0.35)
    # add_paragraph(document, f'   计算结果：120日年化波动率 = {vol_used*100:.1f}%')
    # add_paragraph(document, f'   数据来源：基于历史120日交易数据计算')
    # add_paragraph(document, f'   得分：{scores["波动率"]}/15分')
    if vol_used <= 0.3:
    # add_paragraph(document, '   评估：✅ 波动率较低（≤30%），价格相对稳定')
    elif vol_used <= 0.4:
    # add_paragraph(document, '   评估：⚠️ 波动率中等（30%-40%），价格有一定波动')
    else:
    # add_paragraph(document, '   评估：❌ 波动率较高（>40%），价格波动大，风险较高')
    # add_paragraph(document, f'   风险提示：波动率越高，收益的不确定性越大，需要更高的安全边际')
    # add_paragraph(document, '')

    # add_paragraph(document, '5. VaR风险价值分析', bold=True)
    var_95_pct = var_95 * 100
    # add_paragraph(document, f'   计算结果：95% VaR = {var_95_pct:.1f}%')
    # add_paragraph(document, f'   含义：在95%置信水平下，最大年化损失不超过{var_95_pct:.1f}%')
    # add_paragraph(document, f'   得分：{scores["VaR"]}/15分')
    if var_95 <= 0.2:
    # add_paragraph(document, '   评估：✅ 下行风险可控（VaR≤20%），安全边际充足')
    elif var_95 <= 0.3:
    # add_paragraph(document, '   评估：⚠️ 下行风险中等（VaR 20%-30%），需关注风险')
    else:
    # add_paragraph(document, '   评估：❌ 下行风险较高（VaR>30%），需要较高的折价率补偿')
    # add_paragraph(document, f'   风险提示：VaR仅反映95%置信水平下的风险，极端情况可能超出此范围')
    # add_paragraph(document, '')

    # add_paragraph(document, '6. 综合评估总结', bold=True)
    # add_paragraph(document, f'   综合得分：{total_score}/100分')
    # add_paragraph(document, f'   风险等级：{current_risk_level}')
    # add_paragraph(document, f'   投资建议：{current_advice}')
    # add_paragraph(document, f'   关键优势：{get_key_strengths(scores, profit_prob, discount_premium, var_95)}')
    # add_paragraph(document, f'   主要风险：{get_key_risks(scores, profit_prob, discount_premium, var_95)}')
    # add_paragraph(document, '')


    # add_section_break(document)

    # ==================== 十、风控建议与风险提示 ====================
    """
    add_title(document, '九、风控建议与风险提示', level=1)

    add_paragraph(document, '本章节从风险控制角度，给出盈亏平衡分析、核心指标汇总、最终报价建议和全面的风险提示。')
    add_paragraph(document, '基于保守原则，确保投资决策在合理风险可控范围内进行。')
    add_paragraph(document, '')

    # ==================== 9.1 核心风险指标汇总 ====================
    add_title(document, '9.1 核心风险指标汇总', level=2)

    add_paragraph(document, '本节汇总核心风险指标，为投资决策提供全面的量化评估依据。')
    add_paragraph(document, '')

    # 处理盈利概率和预期收益率的格式
    profit_prob_display = profit_prob * 100 if profit_prob <= 1 else profit_prob
    mean_return_display = mean_return * 100 if abs(mean_return) <= 1 else mean_return

    # 统一使用120日窗口参数
    mc_volatility_120d = market_data.get('volatility_120d', market_data.get('volatility', 0.30))
    mc_drift_120d = market_data.get('annual_return_120d', market_data.get('drift', 0.08))

    # 计算评估值（避免在f-string中使用嵌套三元运算符）
    risk_level_eval = "低风险" if total_score >= 80 else "中等风险" if total_score >= 60 else "高风险"
    profit_level_eval = "较高" if profit_prob >= 0.7 else "中等" if profit_prob >= 0.5 else "较低"
    issue_type_eval = "折价发行" if discount_premium < 0 else "平价发行" if discount_premium < 5 else "溢价发行"

    summary_data = [
        ['风险评分', f'{total_score}/100', risk_level_eval],
        ['盈利概率', f'{profit_prob_display:.1f}%', profit_level_eval],
        ['发行类型', f'{issue_type}', issue_type_eval],
        ['溢价率（相对MA20）', f'{discount_premium:+.2f}%', ''],
        ['预期收益率', f'{mean_return_display:.1f}%', ''],
        ['95% VaR', f'{var_95*100:.1f}%', ''],
        ['波动率(120日)', f'{mc_volatility_120d*100:.1f}%', ''],
        ['年化漂移率(120日)', f'{mc_drift_120d*100:+.2f}%', ''],
        ['DCF内在价值', f'{intrinsic_value:.2f} 元/股', '']
    ]
    add_table_data(document, ['指标', '值', '评估'], summary_data)

    add_paragraph(document, '')

    # ==================== 9.2 盈亏平衡分析 ====================
    add_title(document, '9.2 盈亏平衡分析', level=2)
    add_paragraph(document, '本节通过盈亏平衡分析，评估在不同目标收益率下的安全边际。')
    add_paragraph(document, '')

    # 计算不同收益率下的盈亏平衡价格
    target_returns = np.linspace(0.05, 0.50, 10)
    break_even_prices = []
    issue_price = project_params['issue_price']
    lockup_years = project_params['lockup_period'] / 12
    current_price_eval = project_params['current_price']

    for target_return in target_returns:
        be_price = issue_price * (1 + target_return * lockup_years)
        break_even_prices.append(be_price)

    # 生成表格数据
    be_data = []
    for ret, price in zip(target_returns[::2], break_even_prices[::2]):
        distance = (current_price_eval - price) / current_price_eval * 100
        status = "✅" if distance > 0 else "⚠️"
        be_data.append([f'{ret*100:.0f}%', f'{price:.2f}', f'{distance:+.1f}%', status])

    add_table_data(document, ['期望年化收益率', '盈亏平衡价格(元)', '距离当前价', '状态'], be_data)

    add_paragraph(document, '')
    add_paragraph(document, '盈亏平衡分析结论：')
    add_paragraph(document, f'• 当前价格: {current_price_eval:.2f} 元')
    add_paragraph(document, f'• 发行价格: {issue_price:.2f} 元')
    add_paragraph(document, f'• 20%年化收益率要求下盈亏平衡价: {break_even_prices[3]:.2f} 元')

    be_20 = break_even_prices[3]
    if current_price_eval > be_20:
        margin = (current_price_eval - be_20) / current_price_eval * 100
        add_paragraph(document, f'• ✅ 当前价格高于20%收益率盈亏平衡价{margin:.1f}%，具有较好安全边际')
    else:
        gap = (be_20 - current_price_eval) / current_price_eval * 100
        add_paragraph(document, f'• ⚠️ 当前价格低于20%收益率盈亏平衡价{gap:.1f}%，安全边际不足')

    add_paragraph(document, '')

    # 生成并插入盈亏平衡价格敏感性图表
    break_even_chart_path = os.path.join(IMAGES_DIR, '10_01_break_even_analysis.png')
    generate_break_even_chart(issue_price, current_price_eval, project_params['lockup_period'], break_even_chart_path)
    add_paragraph(document, '图表 9.1: 盈亏平衡价格敏感性曲线')
    add_image(document, break_even_chart_path, width=Inches(6))

    add_paragraph(document, '')

    # ==================== 9.3 报价方案建议 ====================
    add_title(document, '9.3 报价方案建议', level=2)
    add_paragraph(document, '本节提供不同目标收益率下的报价建议，帮助投资者根据风险偏好选择合适的报价方案。')
    add_paragraph(document, '')

    # 显示计算参数
    current_price_eval = project_params['current_price']
    issue_price_eval = project_params['issue_price']
    lockup_years = project_params['lockup_period'] / 12
    historical_return = market_data.get('annual_return_120d', 0.08)
    expected_price = current_price_eval * (1 + historical_return * lockup_years)

    # 添加计算参数说明
    add_paragraph(document, '💡 计算参数（前置假设）：')
    add_paragraph(document, f'• 当前价格：{current_price_eval:.2f} 元/股')
    add_paragraph(document, f'• 锁定期：{project_params["lockup_period"]} 个月（{lockup_years:.2f} 年）')
    add_paragraph(document, f'• 历史年化收益率：{historical_return*100:+.2f}%（基于120日窗口数据）')
    add_paragraph(document, f'• 预期价格（锁定期末）：{expected_price:.2f} 元/股')
    add_paragraph(document, f'  计算公式：预期价格 = 当前价格 × (1 + 历史年化收益率 × 锁定期年数)')
    add_paragraph(document, f'          = {current_price_eval:.2f} × (1 + {historical_return*100:+.2f}% × {lockup_years:.2f})')
    add_paragraph(document, f'          = {expected_price:.2f} 元')
    add_paragraph(document, '')
    add_paragraph(document, '📊 报价计算逻辑：')
    add_paragraph(document, '• 目标：实现年化收益率 R（如15%/20%/25%）')
    add_paragraph(document, '• 公式：最高报价 = 预期价格 ÷ (1 + R × 锁定期年数)')
    add_paragraph(document, '• 溢价率：(最高报价 - 当前价格) ÷ 当前价格 × 100%')
    add_paragraph(document, '  （溢价率为负表示折价，为正表示溢价）')
    add_paragraph(document, '')

    # 计算不同目标收益率下的建议发行价
    target_returns_pricing = [0.15, 0.20, 0.25]
    pricing_recommendations = []

    for target_ret in target_returns_pricing:
        max_issue_price = expected_price / (1 + target_ret * lockup_years)
        premium_rate = (max_issue_price - current_price_eval) / current_price_eval * 100
        pricing_recommendations.append([
            f'{target_ret*100:.0f}%',
            f'{max_issue_price:.2f}',
            f'{premium_rate:+.2f}%'
        ])

    add_table_data(document, ['目标年化收益率', '最高报价(元)', '溢价率'], pricing_recommendations)
    add_paragraph(document, '')

    # 详细说明三个方案
    add_paragraph(document, '1. 保守型方案（15%目标收益）')
    add_paragraph(document, f'   • 建议报价：不高于{pricing_recommendations[0][1]}元（溢价率{pricing_recommendations[0][2]}）')
    add_paragraph(document, '   • 优势：折价较深，安全边际充足，盈利概率较高')
    add_paragraph(document, '   • 风险：可能因报价过低而错配机会，或无法获得足额配售')
    add_paragraph(document, '')

    add_paragraph(document, '2. 平衡型方案（20%目标收益）')
    add_paragraph(document, f'   • 建议报价：不高于{pricing_recommendations[1][1]}元（溢价率{pricing_recommendations[1][2]}）')
    add_paragraph(document, '   • 优势：收益与风险相对平衡，符合一般投资要求')
    add_paragraph(document, '   • 风险：对市场波动敏感，需密切关注120日窗口指标变化')
    add_paragraph(document, '')

    add_paragraph(document, '3. 积极型方案（25%目标收益）')
    add_paragraph(document, f'   • 建议报价：不高于{pricing_recommendations[2][1]}元（溢价率{pricing_recommendations[2][2]}）')
    add_paragraph(document, '   • 优势：可参与更多项目，提高资金使用效率')
    add_paragraph(document, '   • 风险：溢价率较高或溢价发行，安全边际有限，需承受较高波动风险')
    add_paragraph(document, '')

    # ==================== 9.4 当前发行价评估 ====================
    add_title(document, '9.4 当前发行价评估', level=2)
    add_paragraph(document, '本节评估当前发行价的合理性，并给出投资建议。')
    add_paragraph(document, '')

    # 评估当前发行价
    current_premium = (issue_price_eval - current_price_eval) / current_price_eval * 100
    add_paragraph(document, f'• 当前发行价: {issue_price_eval:.2f} 元')
    add_paragraph(document, f'• 当前价格: {current_price_eval:.2f} 元')
    add_paragraph(document, f'• 溢价率（相对当前价）: {current_premium:+.2f}%')

    # 说明：溢价率为负表示折价，为正表示溢价
    if current_premium > 0:
        add_paragraph(document, f'  （溢价{current_premium:.2f}%，发行价高于当前价）')
    elif current_premium < 0:
        add_paragraph(document, f'  （折价{abs(current_premium):.2f}%，发行价低于当前价）')
    else:
        add_paragraph(document, f'  （平价发行，发行价等于当前价）')

    add_paragraph(document, f'• 溢价率（相对MA20: {ma20_mc:.2f}元）: {discount_premium:+.2f}%')
    add_paragraph(document, '')

    # 判断当前发行价是否合理
    max_price_20 = expected_price / (1 + 0.20 * lockup_years)

    if issue_price_eval <= max_price_20:
        add_paragraph(document, f'• ✅ 当前发行价低于20%目标收益对应的最高报价({max_price_20:.2f}元)，具有投资价值')
        investment_advice = "建议积极参与"
    elif issue_price_eval <= max_price_20 * 1.05:
        add_paragraph(document, f'• ⚠️ 当前发行价接近20%目标收益对应的最高报价({max_price_20:.2f}元)，需谨慎评估')
        investment_advice = "建议谨慎参与或要求更高折价"
    else:
        add_paragraph(document, f'• ❌ 当前发行价高于20%目标收益对应的最高报价({max_price_20:.2f}元)，安全边际不足')
        investment_advice = "建议规避或等待更好时机"

    add_paragraph(document, '')

    # ==================== 9.5 主要风险提示 ====================
    add_title(document, '9.5 主要风险提示', level=2)
    add_paragraph(document, '基于多维度风险分析，提示以下主要风险（详见前文各章节详细分析）：')
    add_paragraph(document, '')

    # 1. 市场风险
    add_paragraph(document, '1. 市场风险')
    add_paragraph(document, f'   • 波动率风险：当前120日窗口年化波动率为{mc_volatility_120d*100:.1f}%，市场波动可能导致实际收益偏离预期')
    add_paragraph(document, f'   • 趋势风险：当前120日窗口年化漂移率为{mc_drift_120d*100:+.2f}%，{"上升趋势" if mc_drift_120d > 0 else "下降趋势" if mc_drift_120d < 0 else "震荡趋势"}可能影响解禁时收益')
    add_paragraph(document, '')

    # 2. 流动性风险
    add_paragraph(document, '2. 流动性风险')
    add_paragraph(document, f'   • 锁定期风险：{project_params["lockup_period"]}个月锁定期内无法交易，需承担期间价格波动')
    add_paragraph(document, '   • 解禁冲击：解禁后可能面临抛压，导致实际变现价格低于理论价格')
    add_paragraph(document, '')

    # 3. VaR在险价值风险
    add_paragraph(document, '3. VaR在险价值风险')
    var_95_display = var_95 * 100
    add_paragraph(document, f'   • 120日窗口：95%置信水平下最大可能亏损{var_95_display:.1f}%')
    add_paragraph(document, '   • 尾部风险：历史数据显示，小概率极端事件（黑天鹅）可能导致损失超过VaR预测值')
    add_paragraph(document, '')

    # 4. 估值风险
    add_paragraph(document, '4. 估值风险')
    add_paragraph(document, f'   • DCF估值风险：DCF内在价值{intrinsic_value:.2f}元/股基于多个假设，实际业绩可能偏离预测')
    add_paragraph(document, '   • 相对估值风险：PE/PS/PB相对估值基于行业平均水平，行业景气度变化可能导致估值体系重构')
    add_paragraph(document, '')

    # 5. 发行定价风险
    add_paragraph(document, '5. 发行定价风险')
    if current_premium < 0:
        discount_amount = abs(current_premium)
        add_paragraph(document, f'   • 折价情况：当前溢价率为{current_premium:+.2f}%（折价{discount_amount:.2f}%）')
        if discount_amount < 10:
            add_paragraph(document, f'   • 折价不足风险：折价率低于10%，安全边际有限')
        else:
            add_paragraph(document, f'   • 折价合理：折价{discount_amount:.2f}%提供了一定的安全边际')
    else:
        add_paragraph(document, f'   • 溢价风险：当前溢价率为{current_premium:+.2f}%（溢价发行）')
        add_paragraph(document, f'   • 溢价发行无安全边际，需重点关注公司成长性')
    add_paragraph(document, '   • 定价偏离：若发行价显著高于盈亏平衡价格，将大幅降低盈利概率和预期收益')
    add_paragraph(document, '')

    # 6. 其他风险
    add_paragraph(document, '6. 其他风险')
    add_paragraph(document, '   • 行业政策风险：需关注行业监管政策变化')
    add_paragraph(document, '   • 业绩波动风险：需关注公司业绩预告、审计报告等')
    add_paragraph(document, '   • 竞争格局风险：行业竞争加剧可能影响盈利能力')
    add_paragraph(document, '')

    add_paragraph(document, '💡 综合建议：')
    risk_advice_text = "风险较低，可积极参与" if total_score >= 80 else "风险中等，需谨慎评估" if total_score >= 60 else "风险较高，建议规避"
    add_paragraph(document, f'• 当前项目风险评分{total_score}/100分，{risk_advice_text}')
    add_paragraph(document, f'• 推荐方案：{investment_advice}')
    add_paragraph(document, '• 建议结合个人风险承受能力、资金成本和市场环境选择合适的报价方案')
    add_paragraph(document, '')
    if total_score >= 80 and issue_price_eval <= max_price_20:
        final_recommendation = f"🟢 建议积极参与"
        reason = f"风险评分{total_score}/100（优秀），发行价具有较好安全边际，符合20%目标收益率要求。"
    elif total_score >= 60:
        final_recommendation = f"🟡 谨慎参与"
        reason = f"风险评分{total_score}/100（中等），需结合溢价率和增长预期综合评估。"
    else:
        final_recommendation = f"🔴 建议规避"
        reason = f"风险评分{total_score}/100（较低），安全边际不足，风险较高。"

    add_paragraph(document, f'{final_recommendation}')
    add_paragraph(document, reason)

    add_paragraph(document, '')

    # ==================== 9.6 风控策略建议 ====================
    add_title(document, '9.6 风控策略建议', level=2)
    add_paragraph(document, '根据当前溢价率和盈利概率，提供风控策略建议。')
    add_paragraph(document, '')

    # 使用current_premium（溢价率），为负表示折价
    if current_premium <= -15:
        add_paragraph(document, '• 当前溢价率≤-15%（折价≥15%），可按原计划参与')
    elif current_premium <= -10:
        add_paragraph(document, '• 当前溢价率-10%~-15%（折价10%~15%），建议关注公司基本面和行业前景')
    else:
        add_paragraph(document, f'• 当前溢价率{current_premium:+.2f}%（折价不足或溢价），建议要求更高折价或等待更好时机')

    # 根据盈利概率给出建议（统一处理百分比格式）
    profit_prob_pct = profit_prob * 100 if profit_prob <= 1 else profit_prob
    if profit_prob_pct >= 70:
        add_paragraph(document, f'• 盈利概率{profit_prob_pct:.1f}%（≥70%），投资安全边际充足')
    elif profit_prob_pct >= 50:
        add_paragraph(document, f'• 盈利概率{profit_prob_pct:.1f}%（50%-70%），建议分批参与或控制仓位')
    else:
        add_paragraph(document, f'• 盈利概率{profit_prob_pct:.1f}%（<50%），风险较大，建议谨慎')

    add_paragraph(document, '')

    # ==================== 9.7 压力测试与VaR风险提示 ====================
    add_title(document, '9.7 压力测试与VaR风险提示', level=2)
    add_paragraph(document, '本节汇总压力测试和VaR分析的关键风险指标。')
    add_paragraph(document, '')

    # 从多窗口期VaR分析中提取关键风险指标
    var_95_pct = var_95 * 100
    var_99_pct = var_99 * 100 if 'var_99' in locals() else None

    add_paragraph(document, f'• 120日窗口VaR风险：95%置信水平下最大可能亏损为{var_95_pct:.1f}%')
    if var_99_pct:
        add_paragraph(document, f'• 极端情况VaR：99%置信水平下最大可能亏损为{var_99_pct:.1f}%')

    # 压力测试风险提示
    add_paragraph(document, '• 压力测试情景：需关注PE回归、市场危机、三重打击等极端情景下的潜在损失')
    add_paragraph(document, '• 波动率放大风险：当实际波动率超过120日窗口统计值时，风险敞口将显著增加')

    add_paragraph(document, '')

    # ==================== 9.8 多重方案选项建议 ====================
    add_title(document, '9.8 多重方案选项建议', level=2)
    add_paragraph(document, '根据不同的风险偏好，提供以下参与方案：')
    add_paragraph(document, '')

    add_paragraph(document, '根据不同的风险偏好，提供以下参与方案：')
    add_paragraph(document, '')

    # 方案一：保守型（低风险偏好）
    add_paragraph(document, '🛡️ 方案一：保守型（低风险偏好）')
    add_paragraph(document, f'   • 适用场景：追求确定性，可接受15%年化收益率')
    max_price_conservative = expected_price / (1 + 0.15 * lockup_years)
    discount_conservative = (current_price_eval - max_price_conservative) / current_price_eval * 100
    add_paragraph(document, f'   • 建议报价：不高于{max_price_conservative:.2f}元（折价率{discount_conservative:+.2f}%）')
    add_paragraph(document, f'   • 风险控制：要求较高折价，确保足够安全边际')
    add_paragraph(document, '')

    # 方案二：平衡型（中等风险偏好）
    add_paragraph(document, '⚖️ 方案二：平衡型（中等风险偏好）')
    add_paragraph(document, f'   • 适用场景：平衡收益与风险，目标20%年化收益率')
    max_price_balanced = expected_price / (1 + 0.20 * lockup_years)
    discount_balanced = (current_price_eval - max_price_balanced) / current_price_eval * 100
    add_paragraph(document, f'   • 建议报价：不高于{max_price_balanced:.2f}元（折价率{discount_balanced:+.2f}%）')
    add_paragraph(document, f'   • 风险控制：适度折价，关注盈利概率和VaR指标')
    add_paragraph(document, '')

    # 方案三：积极型（高风险偏好）
    add_paragraph(document, '🚀 方案三：积极型（高风险偏好）')
    add_paragraph(document, f'   • 适用场景：看好公司长期成长，可承受较大波动，目标25%年化收益率')
    max_price_aggressive = expected_price / (1 + 0.25 * lockup_years)
    discount_aggressive = (current_price_eval - max_price_aggressive) / current_price_eval * 100
    add_paragraph(document, f'   • 建议报价：不高于{max_price_aggressive:.2f}元（折价率{discount_aggressive:+.2f}%）')
    add_paragraph(document, f'   • 风险控制：需重点关注公司基本面和行业前景，建议控制仓位')
    add_paragraph(document, '')

    # 风险提示
    add_paragraph(document, '⚠️ 重要提示：')
    add_paragraph(document, '• 本建议基于120日窗口历史数据，实际收益可能偏离预期')
    add_paragraph(document, '• 市场环境变化、公司业绩波动等风险因素可能影响最终收益')
    add_paragraph(document, '• 建议结合最新市场情况和公司公告动态调整报价策略')
    add_paragraph(document, '• 投资有风险，决策需谨慎')

    add_paragraph(document, '')

    # ==================== 9.9 免责声明 ====================
    add_title(document, '9.9 免责声明', level=2)

    disclaimer = f'''
    本报告由自动化分析系统生成，仅供参考使用，不构成投资建议。

    1. 本报告基于历史数据和公开信息进行分析，不保证数据的准确性和完整性；
    2. 市场有风险，投资需谨慎。本报告中的任何分析观点不代表未来表现；
    3. 本报告提到的任何证券或投资标的，仅为分析示例，不构成推荐；
    4. 投资者应根据自身情况独立判断，自行承担投资风险；
    5. 本报告的知识产权归分析团队所有，未经许可不得转载或使用。

    报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    '''

    add_paragraph(document, disclaimer, font_size=9)

    # ==================== 附件：585种情景完整数据表 ====================
    add_section_break(document)
    add_title(document, '附件：585种情景完整数据表', level=1)

    add_paragraph(document, '本附件提供完整的585种情景组合数据，供备查参考。')
    add_paragraph(document, f'情景组合：漂移率（-30%~+30%，13档）× 波动率（10%~50%，5档）× 溢价率（-20%~+20%，9档）= {len(all_scenarios_for_appendix)}种')
    add_paragraph(document, '')
    add_paragraph(document, '说明：本表按波动率分块展示，便于根据实际市场波动率选择参考情景。在每个波动率区间内，按漂移率从高到低排序（漂移率相同时按溢价率从高到低排序）。建议参考创业板180日波动率（约36%）确定目标波动率区间。')
    add_paragraph(document, '')

    # 按波动率分块展示
    # 定义波动率区间
    vol_ranges = [
        ('低波动率区间 (10%-20%)', 10, 20),
        ('中低波动率区间 (20%-30%)', 20, 30),
        ('中高波动率区间 (30%-40%)', 30, 40),
        ('高波动率区间 (40%-50%)', 40, 50)
    ]

    for vol_name, vol_min, vol_max in vol_ranges:
        # 筛选当前波动率区间的情景
        scenarios_in_range = [s for s in all_scenarios_for_appendix
                             if vol_min <= s['volatility']*100 < vol_max]

        if not scenarios_in_range:
            continue

        # 在区间内按漂移率倒序排列，漂移率相同时按溢价率倒序排列
        # 漂移率倒序：从高到低（+30%到-30%）
        # 溢价率倒序：从高到低（+20%到-20%），即从溢价到折价
        scenarios_sorted = sorted(scenarios_in_range,
                                 key=lambda x: (-x['drift'], -x['discount']))

        # 添加区块标题
        add_title(document, f'附表：{vol_name}', level=2)

        # 生成表格数据
        appendix_data = []
        for i, s in enumerate(scenarios_sorted, 1):
            appendix_data.append([
                f"{i}",
                f"{s['drift']*100:+.0f}%",
                f"{s['volatility']*100:.0f}%",
                f"{s['discount']*100:+.0f}%",
                f"{s['issue_price']:.2f}",
                f"{s['mean_return']*100:+.2f}%",
                f"{s['median_return']*100:+.2f}%",
                f"{s['profit_prob']:.1f}%",
                f"{s['var_5']*100:+.2f}%",
                f"{s['var_95']*100:+.2f}%"
            ])

        appendix_headers = ['排名', '漂移率', '波动率', '溢价率', '发行价(元)', '预期年化收益', '中位数收益', '盈利概率', '5% VaR', '95% VaR']
        add_table_data(document, appendix_headers, appendix_data)
        add_paragraph(document, '')

    add_paragraph(document, '')
    add_paragraph(document, '附表说明：')
    add_paragraph(document, '• 排序方式：')
    add_paragraph(document, '  - 按波动率分块（低波动、中低波动、中高波动、高波动）')
    add_paragraph(document, '  - 每个区块内按漂移率从高到低排序（+30%到-30%）')
    add_paragraph(document, '  - 漂移率相同时按溢价率从高到低排序（+20%到-20%）')
    add_paragraph(document, '• 漂移率：年化收益率，反映股价预期趋势（正值=上涨趋势，负值=下跌趋势）')
    add_paragraph(document, '• 波动率：年化波动率，反映股价不确定性')
    add_paragraph(document, '• 溢价率：发行价相对MA20的溢价（负值=折价，正值=溢价）')
    add_paragraph(document, '• 预期年化收益：基于蒙特卡洛模拟的平均年化收益率')
    add_paragraph(document, '• 中位数收益：年化收益率的中位数')
    add_paragraph(document, '• 盈利概率：到期时盈利（收益率>0）的概率')
    add_paragraph(document, '• 5% VaR：95%的情景下收益率不低于此值（较好情况）')
    add_paragraph(document, '• 95% VaR：5%的情景下收益率不高于此值（最差情况）')
    add_paragraph(document, '')
    add_paragraph(document, '💡 使用建议：')
    add_paragraph(document, f'• 参考创业板180日波动率（约36%），优先查看"中高波动率区间 (30%-40%)"的情景')
    add_paragraph(document, '• 根据对未来市场走势的判断（乐观/中性/悲观），在相应漂移率区间查找情景')
    add_paragraph(document, '• 结合当前项目的实际溢价率，找到最接近的情景作为参考基准')
    add_paragraph(document, '• 对比不同情景下的预期收益和盈利概率，评估投资价值')

    # ==================== 保存文档 ====================
    output_path = os.path.join(OUTPUTS_DIR, output_file)
    document.save(output_path)

    print(f"\n✅ 报告生成成功!")
    print(f"   📄 Word文档: {output_path}")
    print(f"   📊 图表目录: {IMAGES_DIR}")
    print(f"   文件大小: {os.path.getsize(output_path)/1024:.1f} KB")

    return output_path


if __name__ == '__main__':
    # 解析命令行参数
    parser = argparse.ArgumentParser(description='定增风险分析报告生成器 V2')
    parser.add_argument('--stock', type=str, default='300735.SZ', help='股票代码（默认：300735.SZ）')
    parser.add_argument('--output', type=str, default='光弘科技_定增风险分析报告.docx', help='输出文件名')
    parser.add_argument('--force-update', action='store_true', help='强制更新市场数据')
    parser.add_argument('--no-update-check', action='store_true', help='跳过数据更新检查')
    parser.add_argument('--max-data-age', type=int, default=3, help='数据最大允许天数（默认：3天）')

    args = parser.parse_args()

    stock_code = args.stock

    # 构建市场数据文件路径
    market_data_file = os.path.join(DATA_DIR, f"{stock_code.replace('.', '_')}_market_data.json")

    # 检查数据是否需要更新
    if not args.no_update_check:
        print("检查市场数据日期...")

        is_fresh, data_date, days_old = check_data_freshness(market_data_file, args.max_data_age)

        if is_fresh:
            print(f"✅ 市场数据是最新的（{data_date.strftime('%Y年%m月%d日')}，距今{days_old}天）")
        elif args.force_update:
            print(f"⚠️ 强制更新数据...")
            update_success = update_market_data(stock_code)
            if not update_success:
                print("⚠️ 数据更新失败，将使用现有数据继续生成报告")
        else:
            if data_date:
                print(f"⚠️ 市场数据已过期（{data_date.strftime('%Y年%m月%d日')}，距今{days_old}天）")
                print(f"   建议运行以下命令更新数据：")
                print(f"   python {__file__} --stock {stock_code} --force-update")
                print()

                # 询问用户是否更新
                try:
                    response = input("是否现在更新数据？(y/N): ").strip().lower()
                    if response in ['y', 'yes']:
                        update_success = update_market_data(stock_code)
                        if not update_success:
                            print("⚠️ 数据更新失败，将使用现有数据继续生成报告")
                    else:
                        print("⚠️ 将使用过期数据继续生成报告，结果可能不准确")
                except (EOFError, KeyboardInterrupt):
                    print("\n⚠️ 将使用过期数据继续生成报告")
            else:
                print(f"❌ 未找到市场数据文件: {market_data_file}")
                print(f"   正在尝试下载数据...")
                update_success = update_market_data(stock_code)
                if not update_success:
                    print("❌ 数据下载失败，无法生成报告")
                    sys.exit(1)
    else:
        print("⏭️ 跳过数据更新检查")

    print()

    # 生成报告
    try:
        report_path = generate_report(stock_code=stock_code, output_file=args.output)
        print("\n报告生成完成！")
        print(f"📄 报告路径: {report_path}")
    except Exception as e:
        print(f"\n❌ 报告生成失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
