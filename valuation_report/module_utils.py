#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
上市公司估值报告 - Word格式化 + matplotlib图表生成

精简自 price_maintenance_risk_analysis 的 module_utils.py
"""

import os
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 加载中文字体
try:
    from utils.font_manager import get_font_prop
    font_prop = get_font_prop()
except Exception:
    font_prop = None


# ==================== Word文档格式化 ====================

def setup_chinese_font(document):
    """设置中文正文字体：仿宋_GB2312 四号(14pt)"""
    document.styles['Normal'].font.name = '仿宋_GB2312'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    document.styles['Normal'].font.size = Pt(14)


def add_title(document, text, level=1):
    """添加标题"""
    if level == 0:
        para = document.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = '方正公文小标宋_GBK'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
            run.font.size = Pt(22)
            run.font.bold = True
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        return para
    else:
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            if level == 1:
                run.font.name = '方正公文小标宋_GBK'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
                run.font.size = Pt(16)
            elif level == 2:
                run.font.name = '方正公文小标宋_GBK'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正公文小标宋_GBK')
                run.font.size = Pt(15)
            else:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(14)
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.line_spacing = 1.5
        return heading


def add_paragraph(document, text, bold=False, font_size=14):
    """添加正文段落"""
    if not text or text.strip() == '':
        return None
    para = document.add_paragraph(text)
    for run in para.runs:
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(font_size)
        if bold:
            run.font.bold = True
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.5
    return para


def add_image(document, image_path, width=Inches(5)):
    """添加图片"""
    if os.path.exists(image_path):
        document.add_picture(image_path, width=width)
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


def add_table_data(document, headers, data, font_size=12):
    """添加表格（宋体，默认小四号12pt）"""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        para = header_cells[i].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.size = Pt(font_size)
        para.runs[0].font.name = '宋体'
        para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data) if cell_data is not None else ''
            para = row_cells[i].paragraphs[0]
            para.runs[0].font.size = Pt(font_size)
            para.runs[0].font.name = '宋体'
            para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if isinstance(cell_data, (int, float)):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def add_section_break(document):
    """添加分页符"""
    document.add_page_break()


def setup_document_header(document, stock_name, stock_code=None):
    """设置统一页眉"""
    from datetime import datetime
    target_section = document.sections[0] if len(document.sections) <= 1 else document.sections[1]
    current_date = datetime.now().strftime('%Y年%m月%d日')
    header_title = f"{stock_name}({stock_code})估值分析报告 | {current_date}" if stock_code else f"{stock_name}估值分析报告 | {current_date}"
    header = target_section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.text = header_title
    for run in header_para.runs:
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(10.5)


def add_page_numbers(document):
    """添加页码"""
    from docx.oxml import parse_xml
    for section in document.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = parse_xml(
            r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:rsidR="00EA4B40" w:rsidRDefault="00EA4B40">'
            r'<w:pPr><w:jc w:val="center"/></w:pPr>'
            r'<w:r><w:rPr><w:rFonts w:ascii="仿宋_GB2312" w:eastAsia="仿宋_GB2312"/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t>- </w:t></w:r>'
            r'<w:r><w:rPr><w:rFonts w:ascii="仿宋_GB2312" w:eastAsia="仿宋_GB2312"/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:instrText xml:space="preserve">PAGE</w:instrText></w:r>'
            r'<w:r><w:rPr><w:rFonts w:ascii="仿宋_GB2312" w:eastAsia="仿宋_GB2312"/><w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t> -</w:t></w:r>'
            r'</w:p>'
        )
        footer_para.clear()
        footer_para._element.append(footer_text)


# ==================== 图表生成函数 ====================

def generate_price_trend_chart(price_df, save_path, stock_name='', ma_windows=(20, 60, 120)):
    """生成股价走势图+均线"""
    if price_df is None or price_df.empty:
        return None
    fig, ax = plt.subplots(figsize=(12, 6))

    dates = pd.to_datetime(price_df['trade_date'])
    ax.plot(dates, price_df['close'], 'b-', linewidth=1.5, label='收盘价')

    colors = ['orange', 'green', 'red']
    for i, w in enumerate(ma_windows):
        ma = price_df['close'].rolling(w).mean()
        ax.plot(dates, ma, color=colors[i % len(colors)], linewidth=1, label=f'MA{w}')

    ax.set_xlabel('日期', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('价格(元)', fontproperties=font_prop, fontsize=12)
    ax.set_title(f'{stock_name}股价走势与均线分析', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=10)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()
    return save_path


def generate_radar_chart(scores, labels, save_path, title='综合评分'):
    """生成雷达图"""
    categories = list(scores.keys())
    values = list(scores.values())
    n = len(categories)

    angles = np.linspace(0, 2 * np.pi, n, endpoint=False).tolist()
    values_plot = values + values[:1]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
    ax.fill(angles, values_plot, alpha=0.25, color='steelblue')
    ax.plot(angles, values_plot, 'o-', linewidth=2, color='steelblue')

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels([labels.get(c, c) for c in categories], fontproperties=font_prop, fontsize=11)
    ax.set_ylim(0, 100)
    ax.set_title(title, fontproperties=font_prop, fontsize=14, fontweight='bold', pad=20)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()
    return save_path


def generate_heatmap(data_matrix, row_labels, col_labels, save_path, title='', center=None, fmt='.2f'):
    """生成热力图"""
    fig, ax = plt.subplots(figsize=(12, 8))
    df = pd.DataFrame(data_matrix, index=row_labels, columns=col_labels)
    sns.heatmap(df, annot=True, fmt=fmt, cmap='RdYlGn', center=center, ax=ax,
                cbar_kws={'label': '每股价值(元)'})
    ax.set_title(title, fontproperties=font_prop, fontsize=14, fontweight='bold', pad=15)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(10)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
        label.set_fontsize(10)
    cbar = ax.collections[0].colorbar
    cbar.ax.yaxis.set_tick_params(right=False)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()
    return save_path


def generate_valuation_comparison_chart(current_metrics, peer_df, save_dir):
    """生成相对估值对比图（PE/PB/PS柱状图）"""
    chart_paths = []
    metrics = ['pe', 'pb', 'ps']
    metric_names = {'pe': 'PE(TTM)', 'pb': 'PB', 'ps': 'PS(TTM)'}

    for metric in metrics:
        fig, ax = plt.subplots(figsize=(14, 6))
        if peer_df is not None and not peer_df.empty and metric in peer_df.columns:
            sorted_df = peer_df.sort_values(metric, ascending=True)
            names = sorted_df['name'].values if 'name' in sorted_df.columns else sorted_df.index.astype(str)
            values = sorted_df[metric].values
            colors = ['red' if sorted_df.iloc[i].get('code', '') == '' else 'steelblue'
                      for i in range(len(sorted_df))]
            ax.barh(range(len(names)), values, color=colors, alpha=0.7)
            ax.set_yticks(range(len(names)))
            ax.set_yticklabels(names, fontproperties=font_prop, fontsize=9)
            if metric in current_metrics and current_metrics[metric]:
                ax.axvline(x=current_metrics[metric], color='red', linestyle='--', linewidth=2,
                           label=f"标的: {current_metrics[metric]:.2f}")
                ax.legend(prop=font_prop)
        ax.set_xlabel(metric_names[metric], fontproperties=font_prop, fontsize=12)
        ax.set_title(f'同行{metric_names[metric]}对比', fontproperties=font_prop, fontsize=14, fontweight='bold')
        ax.grid(True, alpha=0.3, axis='x')
        path = os.path.join(save_dir, f'peer_{metric}_comparison.png')
        plt.tight_layout()
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        chart_paths.append(path)

    return chart_paths


def generate_pe_percentile_chart(pe_series, save_path, stock_name='', window=250):
    """生成PE历史百分位图"""
    if pe_series is None or pe_series.empty:
        return None
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 10), gridspec_kw={'height_ratios': [3, 1]})

    dates = pd.to_datetime(pe_series.index) if pe_series.index.dtype == 'object' else pe_series.index
    ax1.plot(dates, pe_series, 'b-', linewidth=1, label='PE(TTM)')
    for q, color, label in [(0.25, 'green', '25%'), (0.5, 'orange', '50%'), (0.75, 'red', '75%')]:
        ax1.axhline(y=pe_series.quantile(q), color=color, linestyle='--', alpha=0.7, label=label)
    ax1.set_title(f'{stock_name} PE(TTM)历史走势', fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax1.legend(prop=font_prop, fontsize=10)
    ax1.grid(True, alpha=0.3)
    for label in ax1.get_xticklabels():
        label.set_fontproperties(font_prop)

    rolling_pct = pe_series.rolling(window).apply(lambda x: (x.iloc[-1] > x).mean() * 100)
    ax2.fill_between(dates, rolling_pct, alpha=0.3, color='steelblue')
    ax2.plot(dates, rolling_pct, 'b-', linewidth=1)
    ax2.set_ylabel('百分位(%)', fontproperties=font_prop, fontsize=12)
    ax2.set_title(f'PE滚动{window}日百分位', fontproperties=font_prop, fontsize=12)
    ax2.grid(True, alpha=0.3)
    for label in ax2.get_xticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()
    return save_path


def generate_mc_simulation_chart(paths, save_path, title='蒙特卡洛模拟'):
    """生成蒙特卡洛模拟路径图"""
    fig, ax = plt.subplots(figsize=(12, 6))
    for i in range(min(200, paths.shape[1])):
        ax.plot(paths[:, i], linewidth=0.3, alpha=0.3, color='steelblue')
    mean_path = np.mean(paths, axis=1)
    p5 = np.percentile(paths, 5, axis=1)
    p95 = np.percentile(paths, 95, axis=1)
    ax.plot(mean_path, 'r-', linewidth=2, label='均值')
    ax.plot(p5, 'g--', linewidth=1.5, label='5%分位')
    ax.plot(p95, 'g--', linewidth=1.5, label='95%分位')
    ax.set_title(title, fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.set_xlabel('交易日', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('价格(元)', fontproperties=font_prop, fontsize=12)
    ax.legend(prop=font_prop, fontsize=10)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()
    return save_path


def generate_tornado_chart(factors, base_value, save_path, title='龙卷风图'):
    """生成龙卷风图（敏感性分析）"""
    fig, ax = plt.subplots(figsize=(10, 6))
    factors_sorted = sorted(factors, key=lambda x: x['range'], reverse=True)
    names = [f['name'] for f in factors_sorted]
    lows = [f['low'] for f in factors_sorted]
    highs = [f['high'] for f in factors_sorted]

    y_pos = range(len(names))
    ax.barh(y_pos, [base_value - l for l in lows], left=lows, color='steelblue', alpha=0.7, height=0.5)
    ax.barh(y_pos, [h - base_value for h in highs], left=base_value, color='coral', alpha=0.7, height=0.5)
    ax.axvline(x=base_value, color='black', linewidth=1.5, linestyle='-')
    ax.set_yticks(y_pos)
    ax.set_yticklabels(names, fontproperties=font_prop, fontsize=10)
    ax.set_title(title, fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.set_xlabel('每股价值(元)', fontproperties=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, axis='x')
    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()
    return save_path


def generate_var_distribution_chart(returns, var_dict, save_path, title='VaR分析'):
    """生成VaR分布图"""
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.hist(returns, bins=50, density=True, alpha=0.7, color='steelblue', edgecolor='white')
    colors = {'var_90': 'green', 'var_95': 'orange', 'var_99': 'red',
              'cvar_95': 'purple', 'cvar_99': 'darkred'}
    for key, value in var_dict.items():
        if value is not None and key in colors:
            ax.axvline(x=value, color=colors[key], linestyle='--', linewidth=1.5,
                       label=f'{key.upper()}={value:.2%}')
    ax.set_title(title, fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.set_xlabel('日收益率', fontproperties=font_prop, fontsize=12)
    ax.set_ylabel('概率密度', fontproperties=font_prop, fontsize=12)
    ax.legend(prop=font_prop, fontsize=10)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()
    return save_path


def generate_drawdown_chart(prices, save_path, title='最大回撤分析'):
    """生成回撤分析图"""
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(12, 8), gridspec_kw={'height_ratios': [2, 1]})
    ax1.plot(prices, 'b-', linewidth=1)
    peak = prices.cummax()
    ax1.plot(peak, 'g--', linewidth=1, alpha=0.5, label='历史最高')
    ax1.set_title(title, fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax1.legend(prop=font_prop)
    ax1.grid(True, alpha=0.3)
    for label in ax1.get_xticklabels():
        label.set_fontproperties(font_prop)

    drawdown = (prices - peak) / peak
    ax2.fill_between(range(len(drawdown)), drawdown, alpha=0.3, color='red')
    ax2.plot(drawdown, 'r-', linewidth=0.5)
    ax2.set_ylabel('回撤(%)', fontproperties=font_prop, fontsize=12)
    ax2.grid(True, alpha=0.3)
    for label in ax2.get_xticklabels():
        label.set_fontproperties(font_prop)

    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()
    return save_path


def generate_scenario_bar_chart(scenarios, save_path, title='情景分析'):
    """生成情景分析柱状图"""
    fig, ax = plt.subplots(figsize=(10, 6))
    names = [s['name'] for s in scenarios]
    values = [s['value'] for s in scenarios]
    colors = ['green' if s.get('type') == 'bull' else 'steelblue' if s.get('type') == 'base' else 'red'
              for s in scenarios]
    ax.barh(range(len(names)), values, color=colors, alpha=0.7, height=0.5)
    ax.set_yticks(range(len(names)))
    ax.set_yticklabels(names, fontproperties=font_prop, fontsize=10)
    ax.set_title(title, fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.set_xlabel('每股价值(元)', fontproperties=font_prop, fontsize=12)
    ax.grid(True, alpha=0.3, axis='x')
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()
    return save_path


def generate_financial_trend_chart(data_dict, dates, save_path, title='财务趋势'):
    """生成财务趋势折线图"""
    fig, ax = plt.subplots(figsize=(12, 6))
    for name, values in data_dict.items():
        ax.plot(dates, values, 'o-', linewidth=2, markersize=6, label=name)
    ax.set_title(title, fontproperties=font_prop, fontsize=14, fontweight='bold')
    ax.legend(prop=font_prop, fontsize=10)
    ax.grid(True, alpha=0.3)
    for label in ax.get_xticklabels():
        label.set_fontproperties(font_prop)
        label.set_rotation(45)
    for label in ax.get_yticklabels():
        label.set_fontproperties(font_prop)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches='tight')
    plt.close()
    return save_path
